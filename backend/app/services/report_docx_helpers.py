"""Word 报告渲染工具函数。

提供报告渲染共享样式、段落、表格、页眉页脚等工具。
"""

from __future__ import annotations

from datetime import datetime

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml.ns import nsmap, qn
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import OxmlElement

# ============ 设计 token ============
FONT_CN = "微软雅黑"
FONT_EN = "Arial"
COLOR_BRAND = RGBColor(0x2E, 0x75, 0xB6)  # 益语蓝
COLOR_TEXT = RGBColor(0x33, 0x33, 0x33)
COLOR_MUTED = RGBColor(0x88, 0x88, 0x88)
COLOR_TABLE_HEADER_BG = "2E75B6"  # 蓝色表头
COLOR_TABLE_ALT_BG = "F5F8FB"  # 浅蓝交替
COLOR_BORDER = "CCCCCC"


# ============ 工具函数 ============

def set_run_font(run, size=11, bold=False, color=COLOR_TEXT, italic=False):
    """设置 Run 的字体属性，中英文都用微软雅黑/Arial。"""
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    run.font.name = FONT_EN
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), FONT_CN)
    rFonts.set(qn("w:ascii"), FONT_EN)
    rFonts.set(qn("w:hAnsi"), FONT_EN)


def add_paragraph(doc, text="", *, style=None, alignment=None, size=11,
                  bold=False, italic=False, color=COLOR_TEXT, line_spacing=1.5,
                  space_before=0, space_after=6):
    """添加段落并应用统一样式。"""
    p = doc.add_paragraph()
    if style:
        p.style = doc.styles[style]
    if alignment:
        p.alignment = alignment
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, italic=italic, color=color)
    return p


def add_heading(doc, text, level=1):
    """统一的标题样式（不依赖默认 heading 样式）。"""
    config = {
        1: dict(size=22, color=COLOR_BRAND, space_before=18, space_after=10),
        2: dict(size=16, color=COLOR_BRAND, space_before=14, space_after=8),
        3: dict(size=13, color=COLOR_TEXT, space_before=10, space_after=6),
    }[level]
    return add_paragraph(
        doc, text,
        size=config["size"],
        bold=True,
        color=config["color"],
        space_before=config["space_before"],
        space_after=config["space_after"],
        line_spacing=1.3,
    )


def add_data_source(doc, text):
    """添加'数据源'标注，左侧蓝色粗线 + 斜体灰色小字。"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Cm(0.5)
    # 左侧粗线（用边框 XML）
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left_border = OxmlElement("w:left")
    left_border.set(qn("w:val"), "single")
    left_border.set(qn("w:sz"), "12")
    left_border.set(qn("w:space"), "8")
    left_border.set(qn("w:color"), "2E75B6")
    pBdr.append(left_border)
    pPr.append(pBdr)
    # 文本
    label_run = p.add_run("数据源 ")
    set_run_font(label_run, size=9, bold=True, color=COLOR_BRAND)
    body_run = p.add_run(text)
    set_run_font(body_run, size=9, italic=True, color=COLOR_MUTED)
    return p


def add_bullet(doc, text, *, level=0, size=11, bold=False):
    """添加无序列表项（用缩进 + 圆点）。"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.6)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(3)
    bullet_run = p.add_run("• ")
    set_run_font(bullet_run, size=size, color=COLOR_BRAND, bold=True)
    text_run = p.add_run(text)
    set_run_font(text_run, size=size, bold=bold)
    return p


def add_callout(doc, text):
    """添加引用块/重点框（缩进 + 浅色背景）。"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_run_font(run, size=11, italic=True, color=COLOR_TEXT)
    # 添加左侧粗线（用边框 XML）
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left_border = OxmlElement("w:left")
    left_border.set(qn("w:val"), "single")
    left_border.set(qn("w:sz"), "24")
    left_border.set(qn("w:space"), "12")
    left_border.set(qn("w:color"), "2E75B6")
    pBdr.append(left_border)
    pPr.append(pBdr)
    return p


def set_cell_bg(cell, hex_color):
    """设置单元格背景色。"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, color=COLOR_BORDER, size="4"):
    """设置单元格四边边框（浅灰）。"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def add_table(doc, headers, rows, *, widths=None, header_color=COLOR_TABLE_HEADER_BG,
              header_text_color=RGBColor(0xFF, 0xFF, 0xFF), zebra=True):
    """添加统一样式的表格。
    headers: 表头文本列表
    rows: 二维数据
    widths: 列宽（Cm 单位的浮点数列表），可选
    """
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # 表头
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        cell = hdr_cells[i]
        set_cell_bg(cell, header_color)
        set_cell_borders(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # 清空默认段落
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, size=10.5, bold=True, color=header_text_color)
        if widths and i < len(widths):
            cell.width = Cm(widths[i])

    # 数据行
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if zebra and ri % 2 == 1:
                set_cell_bg(cell, COLOR_TABLE_ALT_BG)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(str(val))
            set_run_font(run, size=10)
            if widths and ci < len(widths):
                cell.width = Cm(widths[ci])

    # 表格后留空
    add_paragraph(doc, "", space_after=4)
    return table


def add_page_break(doc):
    """添加分页符。"""
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def add_horizontal_rule(doc):
    """添加水平分隔线。"""
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ============ 内容构建 ============
