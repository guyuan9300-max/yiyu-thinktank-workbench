"""Markdown → .docx 渲染器（python-docx）。

主要给录音转写的会议纪要用：LLM 摘要返回 markdown，后端在挂附件前把它转成 .docx，
这样用户在任务详情里双击附件能直接用 Word / Pages 打开和编辑，而不是看一份 .md 源码。

支持的 markdown subset:
  · 标题 (# ~ #####)
  · 无序列表 (- / * / +)
  · 有序列表 (1. / 2. ...)
  · 加粗 (**xx**)
  · 表格 (用户甲 5/26 加: | a | b |\n|---|---|\n| 1 | 2 |) → 真画 docx 表格
  · 普通段落

其他没识别的语法当作普通段落原样输出。
"""

from __future__ import annotations

import re
from io import BytesIO

from app.services.markdown_table import parse_markdown_tables, _TABLE_ROW_RE, _TABLE_SEP_RE

_BOLD_PATTERN = re.compile(r"\*\*(.+?)\*\*")
_BULLET_PATTERN = re.compile(r"^[-*+]\s+")
_NUMBERED_PATTERN = re.compile(r"^\d+\.\s+")


def _add_paragraph_with_inline_bold(paragraph, text: str) -> None:
    """把 **xx** 标记的部分作为 bold run 插入到段落里。"""
    cursor = 0
    for match in _BOLD_PATTERN.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor : match.start()])
        run = paragraph.add_run(match.group(1))
        run.bold = True
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def _add_docx_table(doc, headers: list[str], rows: list[list[str]]) -> None:
    """真把 markdown 表格画成 docx table:
      · 第 1 行 header 加粗 + 浅蓝填充
      · 边框, 真完整
      · cell 内 markdown 内联 (**加粗**) 真处理
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    n_rows = 1 + len(rows)
    n_cols = len(headers)
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Light Grid Accent 1"  # 真内置 style, 真有边框 + 浅蓝 header

    # 真写 header
    for col_idx, header in enumerate(headers):
        cell = table.rows[0].cells[col_idx]
        cell.text = ""  # 真清空, 用 paragraph 加粗
        paragraph = cell.paragraphs[0]
        _add_paragraph_with_inline_bold(paragraph, header)
        for run in paragraph.runs:
            run.bold = True
    # 真写 body
    for row_idx, row in enumerate(rows):
        for col_idx, cell_value in enumerate(row):
            if col_idx >= n_cols:
                break  # 真容错
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ""
            paragraph = cell.paragraphs[0]
            _add_paragraph_with_inline_bold(paragraph, cell_value)


def render_markdown_to_docx_bytes(
    markdown_text: str,
    *,
    document_title: str | None = None,
) -> bytes:
    """渲染 markdown → .docx 字节流。

    支持：
      - `# / ## / ### / #### / #####` 标题（对应 docx heading 1-5）
      - `- / * / +` 无序列表
      - `1. / 2. ...` 有序列表
      - `**加粗**` 内联加粗
      - 真表格 (| a | b |\\n|---|---|\\n| 1 | 2 |) → docx add_table 真画
      - 普通段落（含中文）
    """
    # 延迟 import 避免 backend 启动时间被 python-docx 拖慢
    from docx import Document  # type: ignore[import-not-found]

    doc = Document()
    if document_title:
        doc.add_heading(document_title, level=0)

    lines = markdown_text.splitlines()
    i = 0
    n = len(lines)
    while i < n:
        raw_line = lines[i]
        line = raw_line.rstrip()
        # 真表格识别 (优先): line 是 row + 下一行是 separator
        if _TABLE_ROW_RE.match(line) and i + 1 < n and _TABLE_SEP_RE.match(lines[i + 1]):
            # 真扫到表格末尾
            from app.services.markdown_table import _split_table_cells

            headers = _split_table_cells(line)
            rows: list[list[str]] = []
            j = i + 2
            while j < n and _TABLE_ROW_RE.match(lines[j]):
                cells = _split_table_cells(lines[j])
                if len(cells) < len(headers):
                    cells = cells + [""] * (len(headers) - len(cells))
                elif len(cells) > len(headers):
                    cells = cells[: len(headers)]
                rows.append(cells)
                j += 1
            if headers:
                _add_docx_table(doc, headers, rows)
            i = j
            continue
        if not line.strip():
            # 空行 → 段落分隔
            doc.add_paragraph()
            i += 1
            continue
        # 标题
        if line.startswith("#####"):
            doc.add_heading(line.lstrip("#").strip(), level=5)
            i += 1
            continue
        if line.startswith("####"):
            doc.add_heading(line.lstrip("#").strip(), level=4)
            i += 1
            continue
        if line.startswith("###"):
            doc.add_heading(line.lstrip("#").strip(), level=3)
            i += 1
            continue
        if line.startswith("##"):
            doc.add_heading(line.lstrip("#").strip(), level=2)
            i += 1
            continue
        if line.startswith("#"):
            doc.add_heading(line.lstrip("#").strip(), level=1)
            i += 1
            continue
        # 列表
        bullet_match = _BULLET_PATTERN.match(line)
        if bullet_match:
            text = line[bullet_match.end() :]
            paragraph = doc.add_paragraph(style="List Bullet")
            _add_paragraph_with_inline_bold(paragraph, text)
            i += 1
            continue
        numbered_match = _NUMBERED_PATTERN.match(line)
        if numbered_match:
            text = line[numbered_match.end() :]
            paragraph = doc.add_paragraph(style="List Number")
            _add_paragraph_with_inline_bold(paragraph, text)
            i += 1
            continue
        # 普通段落
        paragraph = doc.add_paragraph()
        _add_paragraph_with_inline_bold(paragraph, line)
        i += 1

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
