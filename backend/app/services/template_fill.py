from __future__ import annotations

import html
import re
from dataclasses import dataclass
from functools import lru_cache
from pathlib import Path
from typing import Literal
from urllib.parse import quote, urljoin, urlparse

from docx import Document as WordDocument
import httpx


PLACEHOLDER_PATTERN = re.compile(r"\{\{\s*([^{}]{1,120})\s*\}\}")
EMPTY_MARKERS = ("____", "待填写", "待补充", "待完善", "tbd", "todo")

# AI 生成的字段答案里可能夹带 emoji(✅🎯🚀 等),政府申报表正文不应出现 —— 写入 docx 前统一剔除。
# 只命中 emoji / 图形符号区段,保留中文、①②③ 圈号、【】、、。「」等正文标点与符号。
_EMOJI_RE = re.compile(
    "["
    "\U0001F000-\U0001FAFF"  # 高位 emoji 平面:表情/图形/交通/旗帜/补充/扩展-A(含肤色修饰)
    "\U00002600-\U000027BF"  # 杂项符号 + dingbats(☀-➿,含 ✅❌✨✂✈)
    "\U00002B00-\U00002BFF"  # 杂项符号与箭头(⭐⬛⬜)
    "\U00002300-\U000023FF"  # 技术符号(⌚⏰⏳⏩⏯)
    "\U0000FE00-\U0000FE0F"  # 变体选择符(VS1-VS16,emoji 变体)
    "\U0000200D"             # 零宽连接符(ZWJ,组合 emoji)
    "\U000020E3"             # 组合键帽
    "\U00002122\U00002139\U000024C2\U00003030\U0000303D\U00003297\U00003299"  # ™ ℹ Ⓜ 〰 〽 ㊗ ㊙
    "]"
)


def strip_emoji(text: str) -> str:
    """剔除文本中的 emoji / 图形符号,并折叠由此产生的连续空格;保留中文与常规标点。"""
    if not text:
        return text
    cleaned = _EMOJI_RE.sub("", text)
    cleaned = re.sub(r"[ \t ]{2,}", " ", cleaned)
    return cleaned.strip()
PROCESS_HINT_MARKERS = (
    "可从",
    "进一步梳理",
    "建议补",
    "建议内部核验",
    "可填写",
    "如需",
    "建议补充",
)

GENERIC_PUBLIC_HOST_MARKERS = (
    "dingtalk",
    "feishu",
    "lark",
    "qq.com",
    "weixin",
    "wechat",
    "zoom",
    "docs.google",
)
OFFICIAL_SITE_MILESTONE_LINK_HINTS = (
    "大事记",
    "发展历程",
    "里程碑",
    "关于我们",
    "我们是谁",
    "年度盛会",
    "年会",
)
OFFICIAL_SITE_ABOUT_LINK_HINTS = (
    "关于我们",
    "我们是谁",
    "机构简介",
    "平台介绍",
    "联系我们",
)
TemplateFieldType = Literal[
    "precise_fact",
    "structural_summary",
    "governance_mechanism",
    "quantitative_result",
    "attachment_material",
    "general",
]
TemplateFieldValueKind = Literal["fact", "summary", "inference", "missing"]
TABLE_TARGET_HEADER_KEYWORDS = ("填写内容", "主要内容", "重大事件/里程碑", "服务对象/覆盖对象")
TABLE_LABEL_HEADER_KEYWORDS = ("字段", "项目", "年份", "业务模块")


@dataclass(frozen=True)
class TemplateWebSource:
    title: str
    url: str
    snippet: str
    source: str = "public_web"


@dataclass
class TemplateFieldOccurrence:
    label: str
    kind: str
    paragraph_index: int | None = None
    table_index: int | None = None
    row_index: int | None = None
    cell_index: int | None = None
    placeholder: str | None = None
    # field_inventory 类专用：示例/说明列文本，作为 hint 给 LLM
    hint: str | None = None
    # field_inventory 类专用：表格小节名（如"机构基本信息" / "项目明细子表"）
    section: str | None = None


@dataclass(frozen=True)
class TemplateTableTarget:
    label: str
    row_index: int
    cell_index: int
    current_text: str


def normalize_template_label(text: str) -> str:
    cleaned = re.sub(r"[\s\u3000]+", " ", str(text or "")).strip()
    cleaned = cleaned.strip(":：-_[]【】")
    cleaned = re.sub(r"^(请填写|填写|问题|字段)[：:\s]+", "", cleaned)
    return cleaned[:120]


def extract_template_milestone_year(label: str) -> str | None:
    normalized = normalize_template_label(label)
    match = re.fullmatch(r"(20\d{2})年?重大事件(?:/|／)?里程碑", normalized)
    if not match:
        return None
    return match.group(1)


def build_template_fill_retrieval_query(
    *,
    client_name: str,
    template_name: str,
    field_label: str,
    field_type: TemplateFieldType,
) -> str:
    normalized_label = normalize_template_label(field_label)
    milestone_year = extract_template_milestone_year(normalized_label)
    if milestone_year:
        return (
            f"{milestone_year}年 重大事件 里程碑 大事记 发展历程 成立 重要项目 年会 "
            f"{normalized_label}"
        ).strip()
    if field_type == "attachment_material":
        return f"{normalized_label} 附件 材料 证明 文件".strip()
    if field_type == "precise_fact":
        return f"{normalized_label} 登记 官网 年报 章程".strip()
    return f"{template_name} 文档字段填写：{normalized_label}".strip()


def build_template_fill_web_queries(
    *,
    client_name: str,
    field_label: str,
    template_name: str,
    client_domain: str | None = None,
) -> list[str]:
    queries: list[str] = []
    normalized_label = normalize_template_label(field_label)
    milestone_year = extract_template_milestone_year(normalized_label)
    normalized_domain = normalize_template_public_domain(client_domain)
    if milestone_year:
        if normalized_domain:
            queries.append(f"{client_name} {milestone_year} 大事记 里程碑 site:{normalized_domain}")
        queries.append(f"{client_name} {milestone_year} 大事记 重大事件")
        queries.append(f"{client_name} {milestone_year} 年会 项目 公益")
        queries.append(f"{client_name} 发展历程 大事记")
    else:
        if normalized_domain:
            queries.append(f"{client_name} {normalized_label} site:{normalized_domain}")
        queries.append(f"{client_name} {normalized_label}")
        queries.append(f"{client_name} {normalized_label} {template_name}")
    deduped: list[str] = []
    for item in queries:
        candidate = re.sub(r"\s+", " ", item).strip()
        if candidate and candidate not in deduped:
            deduped.append(candidate)
    return deduped[:4]


def _clean_template_fill_public_name_candidate(text: str) -> str:
    cleaned = str(text or "").strip()
    cleaned = re.sub(r"^[一二三四五六七八九十〇零]{0,2}年", "", cleaned).strip()
    cleaned = re.sub(
        r"(年会评估报告|年度盛会|年会|组委会运行规则|运行规则|会议纪要|需求调研报告|项目评估结果反馈|品牌使用指南|项目方案|总结报告|评估报告|访谈核心要点|送审版|内部版|模板版|模拟版|预算调整更新)$",
        "",
        cleaned,
    ).strip()
    return cleaned


def _is_generic_template_fill_public_name_candidate(text: str) -> bool:
    normalized = normalize_template_label(text)
    if not normalized:
        return True
    generic_markers = (
        "工具包",
        "架构图",
        "运行规则",
        "会议纪要",
        "指南",
        "方案",
        "报告",
        "结果反馈",
        "调研",
        "总结",
        "手册",
        "计划书",
        "项目书",
        "需求",
        "评估",
        "模板",
        "清单",
        "表单",
        "预算",
    )
    if any(marker in normalized for marker in generic_markers):
        return True
    return False


def derive_template_fill_public_names(
    client_name: str,
    evidence_titles: list[str] | None = None,
    evidence_snippets: list[str] | None = None,
) -> list[str]:
    names: list[str] = []
    base_name = str(client_name or "").strip()
    if base_name:
        names.append(base_name)
    if re.search(r"[\u4e00-\u9fff]{3,}", base_name):
        return names
    snippet_candidates = evidence_snippets or []
    for snippet in snippet_candidates:
        text = str(snippet or "")
        for match in re.findall(r"([\u4e00-\u9fff]{4,24})(?=（英文名称|地址|邮箱|手机|电话|是由)", text):
            cleaned = _clean_template_fill_public_name_candidate(match)
            if cleaned and not _is_generic_template_fill_public_name_candidate(cleaned) and cleaned not in names:
                names.append(cleaned)
        for match in re.findall(r"[\u4e00-\u9fff]{4,24}", text):
            if not any(keyword in match for keyword in ("基金会", "论坛", "中心", "平台", "组织", "机构", "委员会", "协会")):
                continue
            cleaned = re.sub(r"^(英文名称|中文简称|简称|地址|邮箱|手机|电话)[:：]?", "", match).strip()
            cleaned = _clean_template_fill_public_name_candidate(cleaned)
            if cleaned and not _is_generic_template_fill_public_name_candidate(cleaned) and cleaned not in names:
                names.append(cleaned)
    title_candidates = evidence_titles or []
    for title in title_candidates:
        stem = Path(str(title or "")).stem
        for match in re.findall(r"[\u4e00-\u9fff]{4,24}", stem):
            if not any(keyword in match for keyword in ("基金会", "论坛", "中心", "平台", "组织", "机构", "委员会", "协会")):
                continue
            cleaned = _clean_template_fill_public_name_candidate(match)
            if cleaned and not _is_generic_template_fill_public_name_candidate(cleaned) and cleaned not in names:
                names.append(cleaned)
    return names[:4]


def infer_template_field_type(label: str) -> TemplateFieldType:
    normalized = normalize_template_label(label)
    if not normalized:
        return "general"
    if re.fullmatch(r"20\d{2}年?重大事件(?:/|／)?里程碑", normalized):
        return "structural_summary"
    if any(
        keyword in normalized
        for keyword in (
            "统一社会信用代码",
            "法定代表人",
            "成立年份",
            "成立日期",
            "成立时间",
            "登记管理机关",
            "注册地址",
            "办公地址",
            "联系电话",
            "联系邮箱",
            "官方网站",
            "组织全称",
            "英文名称",
        )
    ):
        return "precise_fact"
    if not normalized.startswith("是否") and any(
        keyword in normalized
        for keyword in (
            "附件",
            "登记证书",
            "章程",
            "批复",
            "备案材料",
            "台账",
            "介绍材料",
            "审计报告",
            "年度报告",
        )
    ):
        return "attachment_material"
    if any(
        keyword in normalized
        for keyword in (
            "党建",
            "党组织",
            "书记",
            "党员",
            "三会一课",
            "主题党日",
            "理事会",
            "主任办公会",
            "章程",
            "治理",
            "合规",
            "透明建设",
        )
    ):
        return "governance_mechanism"
    if any(
        keyword in normalized
        for keyword in (
            "数量",
            "人数",
            "覆盖",
            "总数",
            "规模",
            "荣誉",
            "评估等级",
            "报告",
            "出版物",
            "数据库覆盖",
        )
    ):
        return "quantitative_result"
    if any(
        keyword in normalized
        for keyword in (
            "机构定位",
            "机构性质",
            "业务模块",
            "平台介绍",
            "主要内容",
            "服务对象",
            "覆盖对象",
            "结合方式",
            "关系说明",
            "团队模块",
            "组织发展",
            "运营支持",
            "业务范围",
            "简称",
            "重大事件",
            "里程碑",
        )
    ):
        return "structural_summary"
    return "general"


def infer_template_value_kind(value: str, field_type: TemplateFieldType) -> TemplateFieldValueKind:
    cleaned = str(value or "").strip()
    if not cleaned or cleaned.startswith("【待确认】") or "待补充/待核验" in cleaned:
        return "missing"
    if field_type in {"precise_fact", "quantitative_result"}:
        return "fact"
    if field_type in {"structural_summary", "governance_mechanism"}:
        return "summary"
    if any(marker in cleaned for marker in PROCESS_HINT_MARKERS) or any(marker in cleaned for marker in ("可能", "或可", "推测", "建议")):
        return "inference"
    return "summary"


def build_template_follow_up_question(field_type: TemplateFieldType, label: str) -> str | None:
    normalized = normalize_template_label(label)
    if field_type == "precise_fact":
        return f"请补充或核验“{normalized}”对应的正式登记、官网或公开披露材料。"
    if field_type == "governance_mechanism":
        return f"请补充章程、制度文件、会议纪要或党组织运行记录，以核验“{normalized}”。"
    if field_type == "quantitative_result":
        return f"请补充可引用的统计口径、报表或公开数字，以核验“{normalized}”。"
    if field_type == "attachment_material":
        return f"请补齐“{normalized}”对应附件或材料。"
    if field_type == "structural_summary":
        return f"如需正式定稿，请补充机构介绍、项目方案或战略文本，复核“{normalized}”。"
    return f"请补充更直接支撑“{normalized}”的客户资料。"


def build_template_suggested_sources(field_type: TemplateFieldType, label: str) -> list[str]:
    if field_type == "precise_fact":
        return ["登记证书", "官网/机构公开页", "章程", "年报或信息公开材料"]
    if field_type == "governance_mechanism":
        return ["章程", "制度文件", "会议纪要", "党组织工作记录", "年度党建计划或总结"]
    if field_type == "quantitative_result":
        return ["年度报告", "活动总结", "数据平台或统计报表", "公开发布材料"]
    if field_type == "attachment_material":
        return ["登记证书", "章程", "批复/备案材料", "台账", "审计报告/年报"]
    if field_type == "structural_summary":
        return ["机构介绍", "项目方案", "战略文本", "业务介绍材料", "公开文章/报告"]
    normalized = normalize_template_label(label)
    return [f"与“{normalized}”直接相关的客户原始资料"]


def _is_table_target_header(text: str) -> bool:
    normalized = normalize_template_label(text)
    return any(keyword in normalized for keyword in TABLE_TARGET_HEADER_KEYWORDS)


def _is_table_label_header(text: str) -> bool:
    normalized = normalize_template_label(text)
    return any(keyword in normalized for keyword in TABLE_LABEL_HEADER_KEYWORDS)


FIELD_INVENTORY_NAME_KEYWORDS = ("字段名", "字段名称", "项名", "项目名称")
FIELD_INVENTORY_HINT_KEYWORDS = ("示例/说明", "示例", "说明", "备注", "示例 / 说明")


def _iter_field_inventory_entries(
    table,
    *,
    section_hint: str | None = None,
) -> list[TemplateFieldOccurrence]:
    """识别"规范字段清单表"——每行一个字段定义。

    典型 header 形如：编号 | 字段名 | 字段类型 | 是否必填 | 示例/说明 | 最大字符或范围。
    特征：表头里某一列叫"字段名"，列数 ≥ 3。
    每一行的"字段名"列 = 字段标签；"示例/说明"列 = 给 LLM 的 hint。
    """
    if len(table.rows) < 2:
        return []
    header_cells = table.rows[0].cells
    if len(header_cells) < 3:
        return []
    name_col: int | None = None
    hint_col: int | None = None
    for col, cell in enumerate(header_cells):
        text = normalize_template_label(cell.text)
        if name_col is None and any(kw in text for kw in FIELD_INVENTORY_NAME_KEYWORDS):
            name_col = col
        if hint_col is None and any(kw in text for kw in FIELD_INVENTORY_HINT_KEYWORDS):
            hint_col = col
    if name_col is None:
        return []
    entries: list[TemplateFieldOccurrence] = []
    for row in table.rows[1:]:
        cells = row.cells
        if name_col >= len(cells):
            continue
        label_text = normalize_template_label(cells[name_col].text)
        if not label_text or len(label_text) < 2:
            continue
        # 跳过被合并单元格重复出现的同名 label
        hint_text: str | None = None
        if hint_col is not None and hint_col < len(cells):
            raw_hint = str(cells[hint_col].text or "").strip()
            if raw_hint:
                hint_text = re.sub(r"\s+", " ", raw_hint)[:200]
        entries.append(
            TemplateFieldOccurrence(
                label=label_text,
                kind="field_inventory",
                hint=hint_text,
                section=section_hint,
            )
        )
    return entries


def _iter_header_driven_table_targets(table) -> list[TemplateTableTarget]:
    if len(table.rows) < 2:
        return []
    header_values = [normalize_template_label(cell.text) for cell in table.rows[0].cells]
    targets: list[TemplateTableTarget] = []
    for row_index, row in enumerate(table.rows[1:], start=1):
        cells = row.cells
        for cell_index, header_text in enumerate(header_values):
            if not _is_table_target_header(header_text):
                continue
            if cell_index >= len(cells) or cell_index == 0:
                continue
            label_header = header_values[cell_index - 1] if cell_index - 1 < len(header_values) else ""
            if _is_table_label_header(label_header):
                label_text = normalize_template_label(cells[cell_index - 1].text)
            elif "服务对象/覆盖对象" in header_text:
                row_anchor = normalize_template_label(cells[0].text)
                if not row_anchor:
                    continue
                label_text = f"{row_anchor}（服务对象/覆盖对象）"
            else:
                continue
            if not label_text:
                continue
            if re.fullmatch(r"20\d{2}", label_text):
                label_text = f"{label_text}年重大事件/里程碑"
            targets.append(
                TemplateTableTarget(
                    label=label_text,
                    row_index=row_index,
                    cell_index=cell_index,
                    current_text=str(cells[cell_index].text or "").strip(),
                )
            )
    return targets


def should_enable_template_fill_web_supplement(
    field_type: TemplateFieldType,
    evidence_count: int,
    *,
    field_label: str | None = None,
) -> bool:
    milestone_year = extract_template_milestone_year(field_label or "")
    if milestone_year and evidence_count >= 4:
        return False
    if milestone_year:
        return True
    if evidence_count >= 3:
        return False
    return field_type in {"precise_fact", "structural_summary", "quantitative_result", "attachment_material"}


def normalize_template_public_domain(value: str | None) -> str | None:
    raw = str(value or "").strip().lower()
    if not raw:
        return None
    raw = re.sub(r"^https?://", "", raw)
    raw = raw.split("/", 1)[0].strip()
    raw = raw.removeprefix("www.")
    if "." not in raw:
        return None
    return raw


def _is_generic_public_domain(domain: str | None) -> bool:
    normalized = normalize_template_public_domain(domain)
    if not normalized:
        return True
    return any(marker in normalized for marker in GENERIC_PUBLIC_HOST_MARKERS)


def derive_template_fill_public_domain(
    client_domain: str | None,
    evidence_snippets: list[str] | None = None,
    *,
    public_names: list[str] | None = None,
    client_name: str | None = None,
) -> str | None:
    normalized = normalize_template_public_domain(client_domain)
    if normalized:
        return normalized
    snippets = evidence_snippets or []
    domain_pattern = re.compile(r"\b(?:https?://)?(?:www\.)?([A-Za-z0-9.-]+\.(?:org|org\.cn|cn|com))\b", re.I)
    names = [str(item or "").strip() for item in (public_names or []) if str(item or "").strip()]
    if client_name:
        raw_client_name = str(client_name).strip()
        if raw_client_name and raw_client_name not in names:
            names.append(raw_client_name)
    candidates: list[tuple[str, int, int]] = []
    for snippet in snippets:
        snippet_text = str(snippet or "")
        normalized_snippet = normalize_template_label(snippet_text)
        name_match = any(name and name in normalized_snippet for name in names)
        for match in domain_pattern.finditer(str(snippet or "")):
            domain = normalize_template_public_domain(match.group(1))
            if not domain:
                continue
            if _is_generic_public_domain(domain):
                continue
            start = match.start()
            end = match.end()
            before = snippet_text[max(0, start - 24):start]
            after = snippet_text[end:end + 24]
            explicit_web_hint = any(keyword in (before + after) for keyword in ("官网", "官方网站", "官网地址", "网站"))
            email_like = ("@" in before[-2:]) or ("邮箱" in before and not explicit_web_hint)
            hint_score = 0
            if name_match:
                hint_score -= 2
            if client_name and str(client_name).strip() and str(client_name).strip().lower() in domain:
                hint_score -= 1
            if explicit_web_hint:
                hint_score -= 3
            if domain.endswith(".org.cn"):
                hint_score -= 1
            if email_like and not explicit_web_hint:
                hint_score += 1
                if not name_match:
                    hint_score += 1
            candidates.append((domain, hint_score, len(domain)))
    normalized_domains = {item[0] for item in candidates}
    adjusted_candidates: list[tuple[str, int, int]] = []
    for domain, hint_score, length in candidates:
        if domain.endswith(".org") and f"{domain}.cn" in normalized_domains:
            continue
        if domain.endswith(".org.cn") and domain.removesuffix(".cn") in normalized_domains:
            hint_score -= 1
        adjusted_candidates.append((domain, hint_score, length))
    ranked = sorted(
        adjusted_candidates,
        key=lambda item: (
            item[1],
            0 if item[0].endswith(".org.cn") else 1 if item[0].endswith(".org") else 2,
            item[2],
        ),
    )
    return ranked[0][0] if ranked else None


def _strip_web_html(value: str) -> str:
    text = re.sub(r"(?is)<(script|style).*?>.*?</\1>", " ", value)
    text = re.sub(r"(?is)<[^>]+>", " ", text)
    text = html.unescape(text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


@lru_cache(maxsize=128)
def _fetch_url_html(url: str) -> str:
    try:
        with httpx.Client(follow_redirects=True, timeout=6.0, headers={"User-Agent": "Mozilla/5.0"}) as client:
            response = client.get(url)
            response.raise_for_status()
    except Exception:
        return ""
    return response.text


@lru_cache(maxsize=128)
def _fetch_url_snippet(url: str) -> str:
    return _strip_web_html(_fetch_url_html(url))[:900]


def _official_site_link_keywords(field_label: str, field_type: TemplateFieldType) -> tuple[str, ...]:
    if extract_template_milestone_year(field_label):
        return OFFICIAL_SITE_MILESTONE_LINK_HINTS
    if field_type == "precise_fact":
        return OFFICIAL_SITE_ABOUT_LINK_HINTS
    if field_type == "structural_summary":
        return OFFICIAL_SITE_ABOUT_LINK_HINTS + ("业务", "项目", "平台")
    if field_type == "quantitative_result":
        return ("年度报告", "年报", "年度盛会", "年会", "数据")
    if field_type == "attachment_material":
        return ("附件", "章程", "登记", "年报", "审计")
    return OFFICIAL_SITE_ABOUT_LINK_HINTS


def _extract_official_site_links(
    homepage_html: str,
    *,
    homepage_url: str,
    normalized_domain: str,
    keywords: tuple[str, ...],
    max_items: int = 4,
) -> list[tuple[str, str]]:
    if not homepage_html or not keywords:
        return []
    pattern = re.compile(r'(?is)<a[^>]+href=["\'](?P<href>[^"\']+)["\'][^>]*>(?P<label>.*?)</a>')
    lowered_keywords = tuple(item.lower() for item in keywords)
    scored: list[tuple[int, str, str]] = []
    seen_urls: set[str] = set()
    for match in pattern.finditer(homepage_html):
        href = html.unescape(match.group("href")).strip()
        if not href or href.startswith(("#", "javascript:", "mailto:", "tel:")):
            continue
        absolute_url = urljoin(homepage_url, href)
        parsed = urlparse(absolute_url)
        host = (parsed.netloc or "").lower().removeprefix("www.")
        if not host or normalized_domain not in host:
            continue
        label = _strip_web_html(match.group("label"))[:80]
        basis = f"{label} {absolute_url}".lower()
        score = sum(2 for keyword in lowered_keywords if keyword in basis)
        if score <= 0 or absolute_url in seen_urls:
            continue
        seen_urls.add(absolute_url)
        display_label = label or parsed.path.strip("/") or absolute_url
        scored.append((-score, display_label, absolute_url))
    scored.sort(key=lambda item: (item[0], len(item[1]), item[2]))
    return [(label, url) for _, label, url in scored[:max_items]]


@lru_cache(maxsize=128)
def _search_duckduckgo_html(query: str) -> tuple[tuple[str, str, str], ...]:
    if not query.strip():
        return ()
    url = f"https://html.duckduckgo.com/html/?q={quote(query)}"
    try:
        with httpx.Client(follow_redirects=True, timeout=7.0, headers={"User-Agent": "Mozilla/5.0"}) as client:
            response = client.get(url)
            response.raise_for_status()
    except Exception:
        return ()
    content = response.text
    results: list[tuple[str, str, str]] = []
    pattern = re.compile(
        r'(?is)<a[^>]+class="[^"]*result__a[^"]*"[^>]+href="(?P<href>[^"]+)"[^>]*>(?P<title>.*?)</a>.*?'
        r'<a[^>]+class="[^"]*result__snippet[^"]*"[^>]*>(?P<snippet>.*?)</a>'
    )
    for match in pattern.finditer(content):
        href = html.unescape(match.group("href")).strip()
        if not href.startswith("http"):
            continue
        title = _strip_web_html(match.group("title"))
        snippet = _strip_web_html(match.group("snippet"))
        if title:
            results.append((title[:120], href, snippet[:260]))
        if len(results) >= 4:
            break
    return tuple(results)


def fetch_template_fill_web_sources(
    *,
    client_name: str,
    field_label: str,
    template_name: str,
    client_domain: str | None = None,
    evidence_titles: list[str] | None = None,
    evidence_snippets: list[str] | None = None,
    max_items: int = 2,
    field_type: TemplateFieldType = "general",
) -> list[TemplateWebSource]:
    sources: list[TemplateWebSource] = []
    public_names = derive_template_fill_public_names(client_name, evidence_titles, evidence_snippets)
    normalized_domain = derive_template_fill_public_domain(
        client_domain,
        evidence_snippets,
        public_names=public_names,
        client_name=client_name,
    )
    if normalized_domain:
        homepage_url = f"https://{normalized_domain}"
        homepage_html = _fetch_url_html(homepage_url)
        homepage_snippet = _strip_web_html(homepage_html)[:900]
        if homepage_snippet:
            sources.append(
                TemplateWebSource(
                    title=f"{client_name} 官网",
                    url=homepage_url,
                    snippet=homepage_snippet[:260],
                    source="official_site",
                )
            )
        for title, url in _extract_official_site_links(
            homepage_html,
            homepage_url=homepage_url,
            normalized_domain=normalized_domain,
            keywords=_official_site_link_keywords(field_label, field_type),
            max_items=max(0, max_items * 2),
        ):
            if len(sources) >= max_items:
                return sources[:max_items]
            snippet = _fetch_url_snippet(url)
            if not snippet:
                continue
            sources.append(
                TemplateWebSource(
                    title=f"{client_name} · {title}",
                    url=url,
                    snippet=snippet[:260],
                    source="official_site",
                )
            )
    queries: list[str] = []
    for candidate_name in public_names:
        queries.extend(
            build_template_fill_web_queries(
                client_name=candidate_name,
                field_label=field_label,
                template_name=template_name,
                client_domain=normalized_domain or client_domain,
            )
        )
    seen_urls = {item.url for item in sources}
    for query in queries:
        for title, url, snippet in _search_duckduckgo_html(query):
            if url in seen_urls:
                continue
            seen_urls.add(url)
            source_type = "official_site" if normalized_domain and normalized_domain in url else "public_web"
            sources.append(
                TemplateWebSource(
                    title=title,
                    url=url,
                    snippet=snippet,
                    source=source_type,
                )
            )
            if len(sources) >= max_items:
                return sources[:max_items]
    return sources[:max_items]


def extract_docx_attachment_checklist(path: Path) -> list[str]:
    document = WordDocument(path)
    attachments: list[str] = []
    for table in document.tables:
        if not table.rows:
            continue
        header_values = [(cell.text or "").strip() for cell in table.rows[0].cells]
        joined_header = " ".join(header_values)
        if "附件名称" not in joined_header:
            continue
        name_column = next((index for index, text in enumerate(header_values) if "附件名称" in text), 1)
        for row in table.rows[1:]:
            if len(row.cells) <= name_column:
                continue
            name = normalize_template_label(row.cells[name_column].text)
            if name:
                attachments.append(name)
    return list(dict.fromkeys(attachments))


def extract_docx_template_fields(path: Path) -> list[TemplateFieldOccurrence]:
    document = WordDocument(path)
    fields: list[TemplateFieldOccurrence] = []

    for paragraph_index, paragraph in enumerate(document.paragraphs):
        text = str(paragraph.text or "")
        for match in PLACEHOLDER_PATTERN.finditer(text):
            label = normalize_template_label(match.group(1))
            if not label:
                continue
            fields.append(
                TemplateFieldOccurrence(
                    label=label,
                    kind="placeholder",
                    paragraph_index=paragraph_index,
                    placeholder=match.group(0),
                )
            )

    # 段落里找 section 锚点供 field_inventory 表使用（取上一个标题段落作为 section）
    paragraph_section_anchor: str | None = None
    section_by_table_index: dict[int, str] = {}
    table_counter = 0
    for elem in document.element.body.iter():
        tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
        if tag == 'p':
            txt = ''.join(t.text or '' for t in elem.iter() if (t.tag.split('}')[-1] == 't' if '}' in t.tag else t.tag == 't')).strip()
            if txt and len(txt) <= 60 and not txt.startswith('|'):
                paragraph_section_anchor = txt
        elif tag == 'tbl':
            if paragraph_section_anchor:
                section_by_table_index[table_counter] = paragraph_section_anchor
            table_counter += 1

    for table_index, table in enumerate(document.tables):
        # 优先识别"规范字段清单表"（6 列：编号/字段名/字段类型/...）
        inventory_entries = _iter_field_inventory_entries(
            table,
            section_hint=section_by_table_index.get(table_index),
        )
        if inventory_entries:
            for entry in inventory_entries:
                # 标记 table_index 让后面 apply 知道这一格不该被改
                fields.append(
                    TemplateFieldOccurrence(
                        label=entry.label,
                        kind="field_inventory",
                        table_index=table_index,
                        hint=entry.hint,
                        section=entry.section,
                    )
                )
            continue
        header_targets = _iter_header_driven_table_targets(table)
        if header_targets:
            for item in header_targets:
                target_text = item.current_text
                if target_text and not any(marker in target_text.lower() for marker in EMPTY_MARKERS) and not PLACEHOLDER_PATTERN.search(target_text):
                    continue
                if PLACEHOLDER_PATTERN.search(target_text):
                    for match in PLACEHOLDER_PATTERN.finditer(target_text):
                        label = normalize_template_label(match.group(1)) or item.label
                        fields.append(
                            TemplateFieldOccurrence(
                                label=label,
                                kind="table_placeholder",
                                table_index=table_index,
                                row_index=item.row_index,
                                cell_index=item.cell_index,
                                placeholder=match.group(0),
                            )
                        )
                else:
                    fields.append(
                        TemplateFieldOccurrence(
                            label=item.label,
                            kind="table_cell",
                            table_index=table_index,
                            row_index=item.row_index,
                            cell_index=item.cell_index,
                        )
                    )
            continue
        for row_index, row in enumerate(table.rows):
            cells = row.cells
            if len(cells) < 2:
                continue
            label_text = normalize_template_label(cells[0].text)
            target_text = str(cells[1].text or "").strip()
            if not label_text:
                continue
            if re.fullmatch(r"20\d{2}", label_text):
                label_text = f"{label_text}年重大事件/里程碑"
            if target_text and not any(marker in target_text.lower() for marker in EMPTY_MARKERS) and not PLACEHOLDER_PATTERN.search(target_text):
                continue
            if PLACEHOLDER_PATTERN.search(target_text):
                for match in PLACEHOLDER_PATTERN.finditer(target_text):
                    label = normalize_template_label(match.group(1)) or label_text
                    fields.append(
                        TemplateFieldOccurrence(
                            label=label,
                            kind="table_placeholder",
                            table_index=table_index,
                            row_index=row_index,
                            cell_index=1,
                            placeholder=match.group(0),
                        )
                    )
            else:
                fields.append(
                    TemplateFieldOccurrence(
                        label=label_text,
                        kind="table_cell",
                        table_index=table_index,
                        row_index=row_index,
                        cell_index=1,
                    )
                )

    deduped: list[TemplateFieldOccurrence] = []
    seen: set[tuple[str, str, int | None, int | None, int | None, int | None]] = set()
    for item in fields:
        key = (item.label, item.kind, item.paragraph_index, item.table_index, item.row_index, item.cell_index)
        if not item.label or key in seen:
            continue
        seen.add(key)
        deduped.append(item)
    return deduped


def apply_docx_template_values(
    template_path: Path,
    target_path: Path,
    values: dict[str, str],
) -> tuple[int, int]:
    document = WordDocument(template_path)
    # 写入前剔除 AI 答案里的 emoji(只清填入值,不动模板正文符号);覆盖所有写入路径的唯一出口。
    values = {key: (strip_emoji(val) if isinstance(val, str) else val) for key, val in (values or {}).items()}
    applied = 0
    missing = 0

    # 先扫描，记下哪些表是 field_inventory 表，apply 时跳过、最后在 docx 末尾统一追加答案表
    inventory_table_indexes: set[int] = set()
    inventory_entries_in_order: list[TemplateFieldOccurrence] = []
    for ti, t in enumerate(document.tables):
        entries = _iter_field_inventory_entries(t)
        if entries:
            inventory_table_indexes.add(ti)
            for e in entries:
                e.table_index = ti
                inventory_entries_in_order.append(e)

    for paragraph in document.paragraphs:
        original = str(paragraph.text or "")
        updated = original
        for match in PLACEHOLDER_PATTERN.finditer(original):
            label = normalize_template_label(match.group(1))
            replacement = str(values.get(label) or "").strip()
            if replacement:
                updated = updated.replace(match.group(0), replacement)
            else:
                missing += 1
        if updated != original:
            paragraph.text = updated
            applied += 1

    for table_idx_apply, table in enumerate(document.tables):
        # field_inventory 表不在原位修改——末尾追加专门答案表
        if table_idx_apply in inventory_table_indexes:
            continue
        header_targets = _iter_header_driven_table_targets(table)
        if header_targets:
            for item in header_targets:
                row = table.rows[item.row_index]
                cell = row.cells[item.cell_index]
                current = str(cell.text or "")
                replacement = str(values.get(item.label) or "").strip()
                if PLACEHOLDER_PATTERN.search(current):
                    updated = current
                    row_applied = False
                    for match in PLACEHOLDER_PATTERN.finditer(current):
                        match_label = normalize_template_label(match.group(1)) or item.label
                        match_value = str(values.get(match_label) or "").strip()
                        if match_value:
                            updated = updated.replace(match.group(0), match_value)
                            row_applied = True
                        else:
                            missing += 1
                    if row_applied:
                        cell.text = updated
                        applied += 1
                    continue
                if not replacement:
                    missing += 1
                    continue
                current_lower = current.lower().strip()
                if current_lower and not any(marker in current_lower for marker in EMPTY_MARKERS):
                    continue
                cell.text = replacement
                applied += 1
            continue
        for row in table.rows:
            cells = row.cells
            if len(cells) < 2:
                continue
            label = normalize_template_label(cells[0].text)
            current = str(cells[1].text or "")
            replacement = str(values.get(label) or "").strip()
            if PLACEHOLDER_PATTERN.search(current):
                updated = current
                row_applied = False
                for match in PLACEHOLDER_PATTERN.finditer(current):
                    match_label = normalize_template_label(match.group(1)) or label
                    match_value = str(values.get(match_label) or "").strip()
                    if match_value:
                        updated = updated.replace(match.group(0), match_value)
                        row_applied = True
                    else:
                        missing += 1
                if row_applied:
                    cells[1].text = updated
                    applied += 1
                continue
            if not replacement:
                missing += 1
                continue
            current_lower = current.lower().strip()
            if current_lower and not any(marker in current_lower for marker in EMPTY_MARKERS):
                continue
            cells[1].text = replacement
            applied += 1

    # 在 docx 末尾追加"AI 自动填写答案"表
    if inventory_entries_in_order:
        document.add_paragraph()
        heading = document.add_paragraph()
        heading_run = heading.add_run("AI 自动填写答案")
        heading_run.bold = True
        # 设置黑体字号 14
        from docx.shared import Pt
        heading_run.font.size = Pt(14)
        intro = document.add_paragraph(
            "以下为 AI 基于客户数据库自动生成的字段答案。原文档"
            "中"
            "的字段清单结构保持不变；本表仅追加于末尾以提供答案参考。"
        )
        intro.runs[0].font.size = Pt(9)
        # 按 section 分组
        groups: dict[str, list[TemplateFieldOccurrence]] = {}
        order: list[str] = []
        for entry in inventory_entries_in_order:
            sec = entry.section or "（未命名分组）"
            if sec not in groups:
                groups[sec] = []
                order.append(sec)
            groups[sec].append(entry)
        for section_name in order:
            entries = groups[section_name]
            sec_heading = document.add_paragraph()
            sec_heading_run = sec_heading.add_run(f"▎{section_name}（{len(entries)} 项）")
            sec_heading_run.bold = True
            sec_heading_run.font.size = Pt(11)
            answer_table = document.add_table(rows=1 + len(entries), cols=2)
            answer_table.style = "Light Grid"
            hdr = answer_table.rows[0].cells
            hdr[0].text = "字段名"
            hdr[1].text = "AI 答案"
            for hi in (0, 1):
                for run in hdr[hi].paragraphs[0].runs:
                    run.bold = True
            for ri, entry in enumerate(entries, start=1):
                row = answer_table.rows[ri].cells
                row[0].text = entry.label
                replacement = str(values.get(entry.label) or "").strip()
                if replacement:
                    row[1].text = replacement
                    applied += 1
                else:
                    row[1].text = "【待确认】"
                    missing += 1

    target_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(target_path)
    return applied, missing
