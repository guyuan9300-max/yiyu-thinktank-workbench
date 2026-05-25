"""[A] M9 (2026-05-25) · plan_executor — approved AI 任务计划真执行.

背景:
  顾源源 5/24 真用 @庆华 给安然下了 3 件指令, plan inline approved 走通了, 但之后
  庆华 0 动作 — 因为 create_ai_task_plan 只记录 plan, 不执行. B 在 46-B 求我接.

设计原则:
  · ExecutorRegistry — tool_name → handler 映射, 加新 handler 不动主流程 (顾源源
    硬约束: "不要在 prompt 里写死流程")
  · 异步执行 — FastAPI BackgroundTasks, 不阻塞 approve 返回 (顾源源 5/24 "5-30 min
    任务必须看得到进度", M10 进度可视化挂在这上面)
  · 全 agent_run_log 留痕 — actor_type=internal_ai_agent, actor_id=bot.actor_id,
    每步 start + complete/failed 写两条 (顾源源 self_approve 硬禁止: actor 是 bot
    不是 user; user 只在 approval_queue 那边 decided_by)
  · 失败重试上限 3, 全部失败标 status='failed' + 记 errors[]
  · plan.execution_status 路径: not_started → pending_execute → running → success/failed/partial
"""
from __future__ import annotations

import json
import logging
import re
import time
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Callable, Protocol

logger = logging.getLogger(__name__)


# ── [B] 2026-05-25 真用反馈 · markdown 后处理 + 转 docx ──────────────────────

# LLM 元话语 — 自言自语/自我反思的行, 不该出现在最终文档
META_LINE_PATTERNS = [
    re.compile(r"^#{0,4}\s*[（(](?:哦不对|等等|让我|我先|让我想想|抱歉|不好意思|稍等|嗯|啊|这里我).*?[）)]\s*$"),
    re.compile(r"^[（(]\s*以下为.*[）)]\s*$"),
    re.compile(r"^\s*[（(](?:无编号|仅为区分|符合.*要求|不|按.*要求).*[）)]\s*$"),
    re.compile(r"^\s*[（(]\s*全文约\s*\d+\s*字\s*[）)]\s*$"),  # "(全文约 2100 字)" 这种自我说明
    re.compile(r"^[#\s]*\*?[（(]\s*[Nn][Oo][Tt][Ee][:：].*[）)]\s*$"),
]


def strip_llm_meta_lines(markdown: str) -> str:
    """删除 LLM 自言自语/自我反思的行, 留下纯文档内容."""
    if not markdown:
        return markdown
    out: list[str] = []
    for line in markdown.split("\n"):
        if any(p.match(line) for p in META_LINE_PATTERNS):
            continue
        out.append(line)
    return "\n".join(out).strip()


def _split_long_chinese_paragraph(text: str, max_chars: int = 220) -> list[str]:
    """长段落智能拆分: 按句号/分号/冒号分隔, 每 2-3 句一段.

    LLM 输出的中文段落经常 1000+ 字不换行, 阅读体验差.
    保留原句子, 不动语义, 只在 \"。\" \"；\" 之后插段落边界.
    """
    if not text or len(text) <= max_chars:
        return [text]
    # 按中文句末标点 (。 ！ ？ ；) 切, 保留标点
    sentences = re.split(r"(?<=[。！？；])", text)
    sentences = [s.strip() for s in sentences if s.strip()]
    if not sentences:
        return [text]
    out: list[str] = []
    buf = ""
    for s in sentences:
        # 当前 buf + 这一句 不超过 max_chars: 续在 buf
        if len(buf) + len(s) <= max_chars:
            buf += s
        else:
            if buf:
                out.append(buf)
            buf = s
    if buf:
        out.append(buf)
    return out


def _is_section_heading_bold_line(text: str) -> bool:
    """LLM 经常输出 `**XX**` 当章节标题 (不用 ## H2). 识别这种格式.

    条件: 整段就是 `**...**`, 长度 < 30 字 (短标题).
    """
    t = text.strip()
    m = re.match(r"^\*\*([^*\n]{1,30})\*\*$", t)
    return m is not None


def _extract_bold_text(text: str) -> str:
    """`**XX**` → `XX`"""
    m = re.match(r"^\*\*([^*\n]+)\*\*$", text.strip())
    return m.group(1) if m else text.strip()


def markdown_to_docx(md_text: str, out_path: Path, *, title: str = "") -> Path:
    """[B] 5/25 PM v2 改 (顾源源真用反馈 — 排版要更好):
    · 识别"单行 **XX**" 当 Heading 2 (LLM 真输出方式)
    · 长中文段落智能拆分 (默认 220 字/段)
    · 中文字体 微软雅黑, Title 24pt 居中, H1 18pt, H2 14pt, Normal 11pt, 1.5 倍行距
    · 列表/引用样式优化

    支持:
      · # H1 / ## H2 / ### H3 / #### H4
      · **粗体** (单行) → Heading 2 (LLM 真实风格)
      · **粗体** (混段内) → bold run
      · - / * / 1. 列表
      · > 引用
      · --- 水平分隔
      · 长段落自动按句拆分
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn

    CN_FONT = "微软雅黑"  # macOS/Windows 通用中文字体, 失败回退默认

    doc = Document()

    # 全局样式 (Normal): 字号 11pt, 1.5 倍行距, 中文字体
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    # 中文字体: rFonts.eastAsia
    rpr = normal.element.get_or_add_rPr()
    r_fonts = rpr.find(qn("w:rFonts"))
    if r_fonts is None:
        from docx.oxml import OxmlElement
        r_fonts = OxmlElement("w:rFonts")
        rpr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), CN_FONT)
    pfmt = normal.paragraph_format
    pfmt.line_spacing = 1.5
    pfmt.space_after = Pt(6)

    # Title 样式
    title_style = doc.styles["Title"]
    title_style.font.size = Pt(24)
    title_style.font.bold = True
    title_rpr = title_style.element.get_or_add_rPr()
    title_fonts = title_rpr.find(qn("w:rFonts"))
    if title_fonts is None:
        from docx.oxml import OxmlElement
        title_fonts = OxmlElement("w:rFonts")
        title_rpr.append(title_fonts)
    title_fonts.set(qn("w:eastAsia"), CN_FONT)

    # 文档大标题
    if title:
        h = doc.add_heading("", level=0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = h.add_run(title)
        run.font.size = Pt(24)
        run.font.bold = True
        # 标题后空一行
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(12)

    lines = md_text.split("\n")
    in_code_block = False
    code_buffer: list[str] = []

    bold_re = re.compile(r"\*\*([^*]+)\*\*")
    italic_re = re.compile(r"(?<!\*)\*([^*]+)\*(?!\*)")

    def _add_run_with_font(p, text: str, *, bold: bool = False, italic: bool = False,
                            size: int | None = None, color_rgb: tuple[int, int, int] | None = None) -> None:
        run = p.add_run(text)
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        if size:
            run.font.size = Pt(size)
        if color_rgb:
            run.font.color.rgb = RGBColor(*color_rgb)
        # 中文字体 (强制设, 否则部分 Word 用默认字体)
        rpr = run._r.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            from docx.oxml import OxmlElement
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:eastAsia"), CN_FONT)

    def _add_paragraph_with_inline(text: str, style: str | None = None) -> None:
        """处理 **bold** / *italic* 内联. 长段自动拆分."""
        if not text.strip():
            return
        # 长段落: 按句切, 每段单独一个 Paragraph (但保留 inline 格式)
        chunks = _split_long_chinese_paragraph(text) if style is None else [text]
        for chunk in chunks:
            p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
            cursor = 0
            # 找 **bold** + *italic*
            for m in bold_re.finditer(chunk):
                if m.start() > cursor:
                    _add_run_with_font(p, chunk[cursor:m.start()])
                _add_run_with_font(p, m.group(1), bold=True)
                cursor = m.end()
            if cursor < len(chunk):
                tail = chunk[cursor:]
                cursor2 = 0
                for im in italic_re.finditer(tail):
                    if im.start() > cursor2:
                        _add_run_with_font(p, tail[cursor2:im.start()])
                    _add_run_with_font(p, im.group(1), italic=True)
                    cursor2 = im.end()
                if cursor2 < len(tail):
                    _add_run_with_font(p, tail[cursor2:])

    def _add_heading(text: str, level: int = 2) -> None:
        h = doc.add_heading("", level=level)
        size_map = {0: 24, 1: 18, 2: 14, 3: 12, 4: 11}
        _add_run_with_font(h, text.strip(), bold=True, size=size_map.get(level, 11))

    for raw in lines:
        line = raw.rstrip()

        # 代码块边界
        if line.startswith("```"):
            if in_code_block:
                if code_buffer:
                    p = doc.add_paragraph()
                    run = p.add_run("\n".join(code_buffer))
                    run.font.name = "Courier New"
                    run.font.size = Pt(9)
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
            continue
        if in_code_block:
            code_buffer.append(line)
            continue

        # 水平分隔
        if re.match(r"^-{3,}\s*$", line):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_run_with_font(p, "─" * 30, color_rgb=(0xAA, 0xAA, 0xAA))
            continue

        # H1-H4 (markdown 标准 #)
        m = re.match(r"^(#{1,4})\s+(.+?)\s*$", line)
        if m:
            level = len(m.group(1))
            _add_heading(m.group(2), level=min(level, 4))
            continue

        # **粗体** 单独成行 → 当 Heading 2 (LLM 真实输出方式, 顾源源 5/25 PM 反馈)
        if _is_section_heading_bold_line(line):
            _add_heading(_extract_bold_text(line), level=2)
            continue

        # 无序列表
        m = re.match(r"^[-*]\s+(.+?)\s*$", line)
        if m:
            _add_paragraph_with_inline(m.group(1), style="List Bullet")
            continue

        # 有序列表
        m = re.match(r"^\d+[\.)]\s+(.+?)\s*$", line)
        if m:
            _add_paragraph_with_inline(m.group(1), style="List Number")
            continue

        # 引用块
        m = re.match(r"^>\s+(.+?)\s*$", line)
        if m:
            p = doc.add_paragraph()
            _add_run_with_font(p, "│ " + m.group(1), italic=True, color_rgb=(0x66, 0x66, 0x66))
            continue

        # 空行
        if not line.strip():
            continue

        # 普通段落 (含长段智能拆分)
        _add_paragraph_with_inline(line)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


def _parse_steps_prompt(text: str, today_str: str) -> tuple[str, str]:
    """构造 step 解析 prompt + system instruction. 单独抽出便于本地/云端复用."""
    system_instruction = (
        "你是任务结构化助手. 用户给一段自然语言指令 (可能是口语/微信对话/邮件), "
        "你要拆成 N 个步骤, 每步用三段式: action(做什么) / basis(基于什么) / deliverable(交付什么). "
        "严格输出 JSON 对象, 不要任何前言/解释/markdown 代码块包裹."
    )
    prompt = f"""# 当前日期
{today_str}

# 用户的原指令

```
{text}
```

# 你要做的事

把这段指令拆成 N 个步骤. N 是指令真包含的任务数, 不要硬凑也不要漏. 每步:
- index: 1, 2, 3, ...
- action: 庆华要做的具体动作 (动词短语, 20-60 字). 例: "写一份日慈基金会客户背景档案 (含组织介绍 + 项目清单)"
- basis: 这一步基于什么 (上下文/参考/前置 step 产出, 10-80 字). 例: "数据中心已有日慈合同 2 份 + 日慈官网公开资料 + 益语 2025-2026 服务记录"
- deliverable: 交付的东西 + 落点 + 篇幅 + 审批 (10-100 字). 例: "~8000 字 Word 文档 (.docx), 落点客户工作台, 完成进我审批"

# 严格规则

1. 用户口语连接词 ("首先/然后/接着/最后/第N个/N./一是/再来/还有/另外" 等) 都识别为 step 边界
2. 如果指令只有 1 个任务, steps 只 1 个; 多个任务每个单独成 step
3. basis 必须真分析这一步要什么输入 (客户官网/历史合同/对标资料/前置 step 输出/方法论...), 不要敷衍写"口语指令"
4. deliverable 真抓字数 (如"不少于 4000 字") + 落点 (如"客户工作台"/"我日程"/"进我审批"). 文档格式默认 Word (.docx), 不要写"markdown"
5. action 用动词短语开头 ("写一份..." / "拟一份..." / "建一个..." / "做一份分析..."), 不要写主语
6. 如果用户提到具体客户 (如"日慈基金会"), action 里要带客户名, 不要写成"撰写介绍"这种泛指

# 输出 JSON (只输出对象, 不要 markdown 代码块)

{{
  "steps": [
    {{"index": 1, "action": "...", "basis": "...", "deliverable": "..."}},
    ...
  ]
}}
"""
    return system_instruction, prompt


# 本地 Ollama 解析 endpoint (直接走 HTTP, 不经 AIService 路由)
_OLLAMA_BASE_URL = "http://127.0.0.1:11434/v1/chat/completions"
_LOCAL_PARSE_MODEL = "qwen2.5:7b"


def _parse_steps_via_ollama(system_instruction: str, prompt: str, timeout: float = 30.0) -> dict | None:
    """[B] 2026-05-25 PM · 直接调本地 Ollama qwen2.5:7b 解析 step.

    顾源源 5/25 真测: qwen2.5:7b 4.8s / qwen2.5:14b 18s, 7b 已经够用.
    选 7b 是为了"秒级解析" 体验.

    返回 dict (含 steps) 或 None (Ollama 不可用).
    """
    import urllib.error
    import urllib.request

    payload = {
        "model": _LOCAL_PARSE_MODEL,
        "messages": [
            {"role": "system", "content": system_instruction},
            {"role": "user", "content": prompt},
        ],
        "temperature": 0.2,
        "max_tokens": 2000,
        "stream": False,
    }
    try:
        req = urllib.request.Request(
            _OLLAMA_BASE_URL,
            data=json.dumps(payload).encode("utf-8"),
            headers={"Content-Type": "application/json"},
            method="POST",
        )
        with urllib.request.urlopen(req, timeout=timeout) as resp:
            body = json.loads(resp.read().decode("utf-8"))
    except (urllib.error.URLError, urllib.error.HTTPError, TimeoutError) as exc:
        logger.warning("Ollama parse-steps failed: %s", exc)
        return None
    except Exception as exc:  # noqa: BLE001
        logger.warning("Ollama parse-steps unexpected error: %s", exc)
        return None

    content = (body.get("choices") or [{}])[0].get("message", {}).get("content", "") or ""
    content = content.strip()
    # 去 markdown 代码块包裹
    if content.startswith("```"):
        # 找第二个 ```
        end = content.find("```", 3)
        if end > 0:
            content = content[3:end]
        # 去 `json\n` 前缀
        if content.lstrip().startswith("json"):
            content = content.lstrip()[4:]
    try:
        return json.loads(content)
    except Exception as exc:  # noqa: BLE001
        logger.warning("Ollama parse-steps JSON 解析失败: %s | raw=%s", exc, content[:300])
        return None


def parse_steps_with_llm(text: str, today_str: str = "") -> dict:
    """[B] 2026-05-25 PM · LLM 真解析自然口语指令 → step 三段式 list.

    顾源源真用反馈: regex 拆 step 跟不上自然语言. 必须 LLM 真理解.

    路由策略 (顾源源 5/25 拍板):
      1. 优先 Ollama 本地 qwen2.5:7b (4-5s, 隐私, 0 成本)
      2. Ollama 失败 (没启 / 模型缺) → fallback AIService 默认 (豆包云端)
      3. 全失败 → 返回空 steps + reason (前端 fallback 到 regex 解析)

    返回 {steps: [{index, action, basis, deliverable}], model_used, fallback_reason?}
    """
    today_str = today_str or datetime.now().strftime("%Y年%m月%d日")
    system_instruction, prompt = _parse_steps_prompt(text, today_str)

    # ── 1 · 优先本地 Ollama qwen2.5:7b ──
    raw = _parse_steps_via_ollama(system_instruction, prompt, timeout=30.0)
    model_used = _LOCAL_PARSE_MODEL

    # ── 2 · 失败 → fallback AIService (豆包) ──
    if raw is None:
        ai_service = get_ai_service()
        if not ai_service or not hasattr(ai_service, "_qwen_generate"):
            return {"steps": [], "fallback_reason": "ollama_unavailable + no_ai_service"}
        response_schema = {
            "type": "object",
            "properties": {
                "steps": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "index": {"type": "integer"},
                            "action": {"type": "string"},
                            "basis": {"type": "string"},
                            "deliverable": {"type": "string"},
                        },
                        "required": ["index", "action"],
                    },
                },
            },
            "required": ["steps"],
        }
        try:
            raw = ai_service._qwen_generate(  # noqa: SLF001
                prompt=prompt,
                system_instruction=system_instruction,
                response_schema=response_schema,
                timeout_seconds=30.0,
                max_tokens=2000,
                temperature=0.2,
                task_kind="fast_structured",
            )
            model_used = "doubao_fallback"
        except Exception as exc:  # noqa: BLE001
            logger.warning("parse_steps_with_llm doubao fallback failed: %s", exc)
            return {"steps": [], "fallback_reason": f"ollama_unavailable + doubao_error: {exc}"}

    if isinstance(raw, dict) and "steps" in raw:
        cleaned: list[dict] = []
        for i, s in enumerate(raw.get("steps") or []):
            if not isinstance(s, dict):
                continue
            cleaned.append({
                "index": int(s.get("index") or i + 1),
                "action": str(s.get("action") or "").strip(),
                "basis": str(s.get("basis") or "").strip(),
                "deliverable": str(s.get("deliverable") or "").strip(),
            })
        return {"steps": cleaned, "model_used": model_used}
    return {"steps": [], "fallback_reason": "llm_unstructured", "model_used": model_used}


# ── [B] 5/25 PM (顾源源 path D) · AI 同事跟人走一套接口 ─────────────────────
# 设计: bot 做的事自动生成 task + 自动完成 + 自动复盘
# 让顾源源在重点主线/日历/复盘直接看到, 不需要独立 AI 视图.


def _current_week_label() -> str:
    """ISO 8601 周标签: 2026-W22."""
    now = datetime.now()
    iso = now.isocalendar()
    return f"{iso.year}-W{iso.week:02d}"


def _ensure_weekly_review(db: _DbLike, *, user_id: str, week_label: str) -> str:
    """确保该 user 本周有 weekly_review row, 返回 review_id."""
    row = db.fetchone(
        "SELECT id FROM weekly_reviews WHERE user_id = ? AND week_label = ? LIMIT 1",
        (user_id, week_label),
    )
    if row:
        return dict(row)["id"]
    review_id = f"review_{uuid.uuid4().hex[:10]}"
    now = _now_iso()
    db.execute(
        """INSERT INTO weekly_reviews (
            id, organization_id, week_label, operator_id, user_id,
            summary, created_at, updated_at, sync_status
        ) VALUES (?, '', ?, ?, ?, '', ?, ?, 'local')""",
        (review_id, week_label, user_id, user_id, now, now),
    )
    return review_id


def _add_task_collaborator(
    db: _DbLike, *, task_id: str, user_id: str, full_name: str,
    is_owner: int = 0, order_index: int = 0,
) -> None:
    """给 task 加协作者 (让 user 在日历/任务列表看到)."""
    if not (task_id and user_id):
        return
    now = _now_iso()
    try:
        db.execute(
            """INSERT OR IGNORE INTO task_collaborators (
                task_id, organization_id, user_id, full_name, email,
                order_index, is_owner, inbox_status, created_at, updated_at
            ) VALUES (?, '', ?, ?, '', ?, ?, 'accepted', ?, ?)""",
            (task_id, user_id, full_name, order_index, is_owner, now, now),
        )
    except Exception as exc:  # noqa: BLE001
        logger.warning("add_task_collaborator failed: %s", exc)


def _auto_complete_task(db: _DbLike, task_id: str) -> None:
    """庆华做完自动 status='done'."""
    if not task_id:
        return
    now = _now_iso()
    try:
        db.execute(
            """UPDATE tasks SET status = 'done', progress_status = 'done',
               completed_at = ?, updated_at = ? WHERE id = ?""",
            (now, now, task_id),
        )
    except Exception as exc:  # noqa: BLE001
        logger.warning("auto_complete_task failed: %s", exc)


def _bot_write_self_review(
    bot_display_name: str,
    action: str,
    output_summary: str,
    evidence_used: dict,
    llm_used: str,
) -> str:
    """庆华自动写复盘 — 思考 + 不足 + 下一步. 不调 LLM (太慢), 模板式生成.

    P1 (顾源源拍可升级): 调本地 qwen2.5:7b 让庆华真写反思.
    """
    # 基于真数据生成反思 (不是空话)
    contracts = int(evidence_used.get("contracts") or 0)
    facts = int(evidence_used.get("facts_authoritative") or 0) + int(evidence_used.get("facts_candidate") or 0)
    commits = int(evidence_used.get("commitments") or 0)
    risks = int(evidence_used.get("risks") or 0)
    timelines = int(evidence_used.get("timeline_events") or 0)
    has_evidence = (contracts + facts + commits + risks + timelines) > 0

    lines: list[str] = []
    lines.append(f"【做了什么】{output_summary or action}")
    if has_evidence:
        used_parts: list[str] = []
        if contracts > 0: used_parts.append(f"合同 {contracts} 份")
        if commits > 0: used_parts.append(f"承诺事项 {commits} 条")
        if risks > 0: used_parts.append(f"风险点 {risks} 条")
        if timelines > 0: used_parts.append(f"时间轴 {timelines} 条")
        if facts > 0: used_parts.append(f"事实记录 {facts} 条")
        lines.append(f"【依据数据源】客户工作台已有材料 ({', '.join(used_parts)}), LLM={llm_used}")
    else:
        lines.append(f"【依据数据源】客户工作台无已有材料, 基于 plan_text 用户设定推理写, LLM={llm_used}")

    # 不足列表 (按场景动态生成)
    issues: list[str] = []
    if not has_evidence:
        issues.append("客户工作台暂无材料, 内容多为推理生成, 实际细节待客户确认")
    if llm_used in ("fallback_stub", "no_ai_service"):
        issues.append(f"LLM 调用未成功 ({llm_used}), 输出为占位草稿, 需重跑")
    if contracts == 0 and ("合同" in (action or "") or "协议" in (action or "")):
        issues.append("没接历史合同 RAG, 合同结构是通用法律模板, 需补真客户合同样本")
    if facts == 0 and any(k in (action or "") for k in ["背景", "档案", "介绍"]):
        issues.append("没拉到客户官方事实记录, 集团数据可能跟现实有出入")
    if not issues:
        issues.append("本轮按既定流程完成, 暂无明显不足")
    lines.append("【不足 / 待改进】")
    for i, iss in enumerate(issues, 1):
        lines.append(f"  {i}. {iss}")

    # 下一步建议
    next_steps: list[str] = []
    if not has_evidence:
        next_steps.append("等顾源源确认草稿后, 再补充导入客户官方资料到数据中心")
    if "合同" in (action or "") or "协议" in (action or ""):
        next_steps.append("接历史合同库 RAG 后, 重做合同条款 (P1 待 A 接)")
    if not next_steps:
        next_steps.append("等待顾源源审阅, 按反馈修改")
    lines.append("【下一步】")
    for i, s in enumerate(next_steps, 1):
        lines.append(f"  {i}. {s}")

    return "\n".join(lines)


def _write_review_entry(
    db: _DbLike, *,
    review_id: str,
    task_id: str,
    user_id: str,
    week_label: str,
    note: str,
    task_snapshot: dict | None = None,
) -> None:
    """给某 user 的本周复盘加一条 task entry (庆华自动复盘走这里)."""
    if not (review_id and task_id):
        return
    now = _now_iso()
    entry_id = f"wrte_{uuid.uuid4().hex[:12]}"
    try:
        db.execute(
            """INSERT OR IGNORE INTO weekly_review_task_entries (
                id, organization_id, review_id, task_id, user_id, week_label,
                content_domain, note, structured_note_json, reviewed_at,
                task_snapshot_json, created_at, updated_at
            ) VALUES (?, '', ?, ?, ?, ?, 'work', ?, '{}', ?, ?, ?, ?)""",
            (entry_id, review_id, task_id, user_id, week_label, note, now,
             json.dumps(task_snapshot or {}, ensure_ascii=False), now, now),
        )
    except Exception as exc:  # noqa: BLE001
        logger.warning("write_review_entry failed: %s", exc)


def _safe_filename_segment(text: str, max_len: int = 60) -> str:
    """文件名安全字符: 去 / \\ : * ? \" < > | 等, 中文/英文/数字保留."""
    # 替换 OS 不允许字符为 _
    cleaned = re.sub(r'[\\/:*?"<>|\s]+', '_', text)
    # 去首尾下划线
    cleaned = cleaned.strip('_.')
    if not cleaned:
        return "untitled"
    return cleaned[:max_len]


class _DbLike(Protocol):
    def execute(self, sql: str, params: tuple = ...) -> Any: ...
    def fetchone(self, sql: str, params: tuple = ...) -> Any: ...
    def fetchall(self, sql: str, params: tuple = ...) -> Any: ...


# ── [B] 2026-05-25 PM · AIService singleton (顾源源拍板"真接 LLM, 不模拟") ───────
# main.py 启动时调 set_ai_service(ai) 注入. handler 内部用 get_ai_service().
# 单例 holder, 不持有引用周期问题 (AIService 本身就是 module-level lazy 的).
_ai_service_holder: dict[str, Any] = {"svc": None}


def set_ai_service(svc: Any) -> None:
    """main.py 启动 (create_app 内) 调一次, 把 AiService 实例注入 plan_executor."""
    _ai_service_holder["svc"] = svc


def get_ai_service() -> Any:
    return _ai_service_holder.get("svc")


# ── 中文时间解析 (不依赖 LLM, 给 _handler_tasks_create 用) ─────────────────────
# 真用 case: "5 月 27 日下午 2 点" / "5/27 14:00" / "明天上午 10 点".
# 解析失败回退 None, handler 标"待确认".

_CN_NUM = {"零": 0, "一": 1, "二": 2, "三": 3, "四": 4, "五": 5,
           "六": 6, "七": 7, "八": 8, "九": 9, "十": 10, "两": 2}


def _parse_cn_num(s: str) -> int | None:
    if not s:
        return None
    s = s.strip()
    if s.isdigit():
        return int(s)
    # 简单中文数字: 十/十二/二十/二十三
    if s == "十":
        return 10
    if s.startswith("十") and len(s) == 2:
        n = _CN_NUM.get(s[1])
        return 10 + n if n is not None else None
    if "十" in s and len(s) <= 3:
        parts = s.split("十")
        tens = _CN_NUM.get(parts[0], 1) if parts[0] else 1
        ones = _CN_NUM.get(parts[1], 0) if len(parts) > 1 and parts[1] else 0
        return tens * 10 + ones
    if len(s) == 1:
        return _CN_NUM.get(s)
    return None


def parse_chinese_datetime(text: str, ref_year: int | None = None) -> str | None:
    """从中文文本里抓最像目标时间的片段, 返回 ISO 8601 (+08:00). 失败 None.

    支持:
      · "5 月 27 日下午 2 点" / "5月27日14:00" / "5/27 14:00"
      · "明天上午 10 点" / "今天下午 3 点" / "后天 9:30"
      · "2026-05-27 14:00" / "2026/5/27 14:00"
    """
    if not text:
        return None
    now = datetime.now()
    year = ref_year or now.year

    # 1. ISO-like: 2026-05-27 14:00 / 2026/5/27 14:00
    m = re.search(
        r"(20\d{2})[-/年](\d{1,2})[-/月](\d{1,2})[日号]?\s*(?:[T ]\s*(\d{1,2})[:点：](\d{1,2})?)?",
        text,
    )
    if m:
        y, mo, d = int(m.group(1)), int(m.group(2)), int(m.group(3))
        hh = int(m.group(4)) if m.group(4) else 9
        mm = int(m.group(5)) if m.group(5) else 0
        try:
            return datetime(y, mo, d, hh, mm).strftime("%Y-%m-%dT%H:%M:00+08:00")
        except ValueError:
            pass

    # 2. "X 月 Y 日下午 N 点" (无年, 用 ref_year)
    m = re.search(
        r"(\d{1,2}|[一二三四五六七八九十]+)\s*月\s*(\d{1,2}|[一二三四五六七八九十]+)\s*[日号]"
        r"(?:\s*(上午|下午|早上|晚上|中午|凌晨))?"
        r"\s*(\d{1,2}|[一二三四五六七八九十]+)?\s*[点时:：]?\s*(\d{1,2})?",
        text,
    )
    if m:
        mo = _parse_cn_num(m.group(1))
        d = _parse_cn_num(m.group(2))
        ampm = m.group(3) or ""
        hh = _parse_cn_num(m.group(4) or "") or (9 if ampm in ("上午", "早上") else 14)
        mm = int(m.group(5)) if m.group(5) and m.group(5).isdigit() else 0
        # 下午/晚上 hour < 12 → +12
        if ampm in ("下午", "晚上") and hh < 12:
            hh += 12
        elif ampm == "中午" and hh < 12:
            hh = 12
        try:
            return datetime(year, mo, d, hh, mm).strftime("%Y-%m-%dT%H:%M:00+08:00")
        except (ValueError, TypeError):
            pass

    # 3. "明天 / 今天 / 后天" + 时间
    rel_map = {"今天": 0, "明天": 1, "后天": 2, "大后天": 3}
    for kw, delta in rel_map.items():
        if kw in text:
            m = re.search(
                rf"{kw}\s*(上午|下午|早上|晚上|中午)?\s*(\d{{1,2}}|[一二三四五六七八九十]+)?\s*[点时:：]?\s*(\d{{1,2}})?",
                text,
            )
            if m:
                from datetime import timedelta
                target = (now + timedelta(days=delta))
                ampm = m.group(1) or ""
                hh = _parse_cn_num(m.group(2) or "") or (9 if ampm in ("上午", "早上") else 14)
                mm = int(m.group(3)) if m.group(3) and m.group(3).isdigit() else 0
                if ampm in ("下午", "晚上") and hh < 12:
                    hh += 12
                try:
                    return target.replace(hour=hh, minute=mm, second=0, microsecond=0).strftime(
                        "%Y-%m-%dT%H:%M:00+08:00",
                    )
                except ValueError:
                    pass
                break

    return None


# ── 中文 action → 文档类型 enum (顾源源虚构客户/真客户都覆盖) ─────────────────

_DOC_TYPE_KEYWORDS: list[tuple[str, str]] = [
    # (关键词正则, doc_type enum)
    (r"(客户)?背景档案|集团介绍|公司介绍|机构介绍|公司档案|集团档案", "client_background_archive"),
    (r"项目机会|机会分析|opportunity", "opportunity_analysis"),
    (r"品牌分析|brand.*analysis|品牌资产|品牌健康", "brand_analysis"),
    (r"战略陪伴提案|陪伴提案|合作提案|service.*proposal", "strategic_partnership_proposal"),
    (r"战略陪伴.*协议|陪伴.*协议|战略.*合作.*协议|service.*contract|战略合伙人协议", "strategic_partnership_contract"),
    (r"理事会|board.*brief|董事会简报", "board_brief"),
    (r"会议(纪要|议程|准备)|meeting.*(brief|prep)", "meeting_prep"),
    (r"复盘|review|总结报告", "review_doc"),
]


def infer_doc_type_from_action(action: str) -> str:
    """中文 step.action 文字 → documents.kind. 兜底 'general_draft'."""
    if not action:
        return "general_draft"
    for pat, kind in _DOC_TYPE_KEYWORDS:
        if re.search(pat, action):
            return kind
    return "general_draft"


# ── 文档大小预估 (从 action 抓字数) ──────────────────────────────────────
def extract_target_size_chars(action: str) -> int | None:
    """抓 'X 字' / 'X-Y 字' / '~X 字' → 中位字数. 没抓到回 None."""
    if not action:
        return None
    m = re.search(r"(\d{3,5})\s*[-–至到]\s*(\d{3,5})\s*字", action)
    if m:
        return (int(m.group(1)) + int(m.group(2))) // 2
    m = re.search(r"(?:大约|约|大概|~)?\s*(\d{3,5})\s*字", action)
    if m:
        return int(m.group(1))
    return None


def _now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


# ─── Subtask 解析 ─────────────────────────────────────────────────


def _parse_subtasks(plan: dict) -> list[dict]:
    """从 plan row 抽出 subtask list, 多源兜底.

    优先级 (高 → 低):
      1. parsed_subtasks_json (M9 新字段, 若 B 后续接入 — 最结构化)
      2. write_actions_json (现有 schema, 已经是 {tool, payload} 形式 — 优先 over steps)
      3. steps_json (粗粒度 module/action/expected_result, 推 tool 用)

    每个 subtask 至少含 {tool, payload, label}.
    """
    # 1. parsed_subtasks_json (优先, B M9 后续接入)
    raw = plan.get("parsed_subtasks_json") or ""
    if raw:
        try:
            subs = json.loads(raw)
            if isinstance(subs, list) and subs:
                return [_normalize_subtask(s, idx) for idx, s in enumerate(subs)]
        except Exception as exc:
            logger.warning("parse parsed_subtasks_json fail: %s", exc)

    # 2. write_actions_json (结构化, 已含具体 tool+payload, 优先 over steps)
    raw = plan.get("write_actions_json") or ""
    if raw:
        try:
            actions = json.loads(raw)
            if isinstance(actions, list) and actions:
                return [_action_to_subtask(a, idx, plan) for idx, a in enumerate(actions)]
        except Exception as exc:
            logger.warning("parse write_actions_json fail: %s", exc)

    # 3. steps_json (兜底, 粗粒度 module/action)
    raw = plan.get("steps_json") or ""
    if raw:
        try:
            steps = json.loads(raw)
            if isinstance(steps, list) and steps:
                return [_step_to_subtask(s, idx, plan) for idx, s in enumerate(steps)]
        except Exception as exc:
            logger.warning("parse steps_json fail: %s", exc)

    return []


def _normalize_subtask(s: dict, idx: int) -> dict:
    """parsed_subtasks_json 项 → 标准 subtask."""
    tool = (s.get("tool") or s.get("tool_name") or "").strip() or "noop"
    return {
        "index": idx,
        "tool": tool,
        "payload": s.get("payload") or {},
        "label": s.get("label") or f"{tool} #{idx + 1}",
    }


def _step_to_subtask(step: dict, idx: int, plan: dict) -> dict:
    """steps_json 项 → subtask. module + action 推 tool_name."""
    module = (step.get("module") or "").strip().lower()
    action = (step.get("action") or "").strip().lower()
    expected = step.get("expected_result") or ""

    # 推 tool 名: module/action 映射 (顾源源 hard rule: 不写死流程, 但路由表是合理的)
    if "document" in module or "document" in action or "draft" in action or "generate" in action:
        tool = "documents.generate"
    elif "task" in module or "task" in action:
        tool = "tasks.create"
    elif "import" in module or "smart_import" in action or "ingest" in action:
        tool = "smart_import"
    else:
        tool = "noop"

    return {
        "index": idx,
        "tool": tool,
        "payload": {
            "client_id": plan.get("client_id"),
            "module": module,
            "action": action,
            "expected_result": expected,
        },
        "label": f"{tool} ({module or action or 'noop'})",
    }


def _action_to_subtask(action: dict, idx: int, plan: dict) -> dict:
    """write_actions_json 项 → subtask."""
    tool = (action.get("tool") or action.get("type") or "noop").strip()
    payload = action.get("payload") or {k: v for k, v in action.items() if k not in ("tool", "type")}
    if "client_id" not in payload and plan.get("client_id"):
        payload["client_id"] = plan.get("client_id")
    return {
        "index": idx,
        "tool": tool,
        "payload": payload,
        "label": action.get("label") or f"{tool} #{idx + 1}",
    }


# ─── Executor Registry ──────────────────────────────────────────


class ExecutorRegistry:
    """tool_name → handler. 注册式, 加新 tool 不动 execute_plan 主流程."""

    def __init__(self) -> None:
        self._handlers: dict[str, Callable] = {}

    def register(self, tool_name: str, handler: Callable) -> None:
        self._handlers[tool_name] = handler

    def get(self, tool_name: str) -> Callable:
        return self._handlers.get(tool_name) or self._handlers["noop"]

    def known_tools(self) -> list[str]:
        return sorted(self._handlers.keys())


# ─── Handlers ───────────────────────────────────────────────────


def _handler_documents_generate(
    db: _DbLike, *, payload: dict, bot: dict, plan: dict,
) -> dict:
    """[B] 2026-05-25 PM 重写 (顾源源拍板"从底层架构修, 不要硬编码模拟"):
        1. 拉 evidence (build_company_brain_context, A 已接)
        2. 真接 LLM (AIService._qwen_generate) 生成 markdown
        3. 真 INSERT documents 表 (客户工作台真能看到这条)

    输入约定 (来自 _step_to_subtask, B 5/25 转换):
      payload = { client_id, module, action(中文), expected_result }
    """
    from app.services.company_brain_context_builder import (
        build_company_brain_context, summarize_for_api_response,
    )

    client_id = (payload.get("client_id") or plan.get("client_id") or "").strip()
    action = (payload.get("action") or "").strip()
    expected_result = (payload.get("expected_result") or "").strip()

    if not client_id:
        raise ValueError("documents.generate 需 client_id (plan 未绑定客户)")

    # 1. 推 doc_type + 目标字数
    doc_type = infer_doc_type_from_action(action)
    target_chars = extract_target_size_chars(action) or 2000  # 默认 ~2000 字

    # 2. 拉 evidence (真查数据中心已有材料)
    pack = build_company_brain_context(
        db, client_id=client_id,
        user_query=action or plan.get("plan_title") or "",
        task_type="strategy_narrative" if doc_type in (
            "strategic_partnership_contract", "strategic_partnership_proposal", "board_brief",
        ) else "workbench_qa",  # type: ignore
    )
    evidence_summary = summarize_for_api_response(pack).get("evidence_summary") or {}

    # 客户名 (写 prompt 用 + 写 documents.title 用)
    client_row = db.fetchone("SELECT name FROM clients WHERE id = ?", (client_id,))
    client_name = (dict(client_row).get("name") if client_row else client_id) or client_id

    # 2b. [B] 5/25 PM 真用反馈 — 必须真拉 documents 表内容, 不只统计.
    # 顾源源原话: "写任何资料, 一定是优先基于数据中心和工作台的内容,
    #             其次才是基于资讯情报站补充. 而且资讯情报站补充的内容在数据中心也有."
    # ↑ 真相: V2.1 documents 表已经把 raw_file (用户上传) + internet_source_doc
    # (资讯情报站抓的) + wechat_article_excerpt (公众号文章) + task_doc/meeting_doc
    # 都存进来了. 但当前 handler 只查 evidence_summary 统计, 没读真 excerpt.
    # → LLM 写报告时拿不到客户真材料 → 凭空写.
    # 修法: 拉 top N 份真客户文档的 excerpt, 按数据源优先级排序, inject prompt.
    user_docs: list[dict] = []
    intel_docs: list[dict] = []
    try:
        # 优先级 1: 用户上传文件 (raw_file) — 客户官方资料
        for row in db.fetchall(
            """SELECT title, kind, canonical_kind, excerpt, created_at FROM documents
               WHERE client_id = ?
                 AND content_domain = 'work'
                 AND canonical_kind = 'raw_file'
                 AND length(excerpt) > 80
               ORDER BY created_at DESC LIMIT 15""",
            (client_id,),
        ):
            user_docs.append(dict(row))
        # 优先级 2: 资讯情报站 + 公众号 + 外部摘录 (其次)
        for row in db.fetchall(
            """SELECT title, kind, canonical_kind, excerpt, created_at FROM documents
               WHERE client_id = ?
                 AND content_domain = 'work'
                 AND canonical_kind IN ('internet_source_doc', 'wechat_article_excerpt', 'project_enrichment_doc')
                 AND length(excerpt) > 80
               ORDER BY created_at DESC LIMIT 10""",
            (client_id,),
        ):
            intel_docs.append(dict(row))
    except Exception as exc:  # noqa: BLE001
        logger.warning("documents.generate fetch real content failed: %s", exc)

    # 3. 真调 LLM
    ai = get_ai_service()
    markdown: str = ""
    llm_used: str = "none"
    llm_error: str | None = None

    # [B] 5/25 PM P0 修复: 时间注入 / 字数达标 / 元话语清洁 / 真出 docx
    today_str = datetime.now().strftime("%Y年%m月%d日")
    target_words = min(max(target_chars, 1500), 10000)  # 目标字数 1500-10000

    if ai is not None and hasattr(ai, "_qwen_generate"):
        # 字数策略 v2: 中文 max_tokens 1 token ≈ 1.5 字, 目标 X 字需 ~X*0.7 tokens
        # P0-2 修复: 加 max_tokens 上限到 15000 (豆包 doubao-seed-2-0-pro 支持)
        # 让 LLM 真出 10000 字时不被 max_tokens 截
        max_tokens = min(int(target_words * 1.4) + 1000, 15000)

        system_instruction = (
            f"你是 {bot.get('display_name', '庆华')}, 益语智库 {bot.get('department_name', '战略发展部')} 的 AI 同事. "
            f"你的工作风格是严肃专业的咨询顾问. "
            f"严格规则:\n"
            f"1. 不要套程式化模板 (\"## 1. 引言 / ## 2. 背景\" 这种)\n"
            f"2. 不要输出任何自我反思 / 自我修正 / 元话语 (比如\"哦不对\"\"让我重新写\"\"(以下为...)\"\"(无编号)\"\"(全文约X字)\" 等都禁止)\n"
            f"3. 直接给最终版, 不要中途说\"等等我改\"或类似话\n"
            f"4. 数据中心 evidence 全 0 时, 基于 plan_text 设定写, 标注'数据源: 用户指令上下文'\n"
            f"5. 必须达到指定字数, 不允许敷衍/凑字, 但也不能空话连篇"
        )

        # 拼参考材料块 — 顾源源 5/25 PM 真用反馈: 必须真拉内容, 不只统计
        # 优先级 1: 数据中心/客户工作台真文档 (raw_file = 客户官方资料)
        # 优先级 2: 资讯情报站 / 公众号 / 外部摘录 (其次)
        # 优先级 3: 结构化 evidence (合同/承诺/风险/时间轴)
        # 优先级 4: 凭 plan_text 推理 (兜底)
        priority_sections: list[str] = []

        # P1: 用户官方文档真 excerpt
        if user_docs:
            section_lines = ["## 数据源 1 · 客户工作台已有官方文档 (优先级最高)"]
            section_lines.append(f"(共 {len(user_docs)} 份, 按时间倒序; 每份截 800 字)")
            for i, d in enumerate(user_docs, 1):
                ex = (d.get("excerpt") or "")[:800]
                section_lines.append(f"\n### [文档 {i}] {d.get('title', '(无标题)')[:60]} ({d.get('kind') or 'unknown'})")
                section_lines.append(ex)
            priority_sections.append("\n".join(section_lines))
        else:
            priority_sections.append("## 数据源 1 · 客户工作台官方文档\n\n- 数据中心暂无该客户已上传文件, 跳过\n")

        # P2: 资讯情报站抓的内容
        if intel_docs:
            section_lines = ["## 数据源 2 · 资讯情报站 + 公众号 + 外部摘录 (其次)"]
            section_lines.append(f"(共 {len(intel_docs)} 份, 这些内容也已存进数据中心, 但权威性低于官方资料)")
            for i, d in enumerate(intel_docs, 1):
                ex = (d.get("excerpt") or "")[:500]
                section_lines.append(f"\n### [情报 {i}] {d.get('title', '(无标题)')[:60]} ({d.get('canonical_kind') or 'intel'})")
                section_lines.append(ex)
            priority_sections.append("\n".join(section_lines))

        # P3: 结构化 evidence
        ref_lines: list[str] = []
        try:
            if pack.contracts:
                ref_lines.append(f"- 已识别合同 {len(pack.contracts)} 份 (来自合同结构化解析)")
            if pack.commitments:
                ref_lines.append(f"- 承诺事项 {len(pack.commitments)} 条")
            if pack.risks:
                ref_lines.append(f"- 风险点 {len(pack.risks)} 条")
            if pack.timeline:
                ref_lines.append(f"- 时间轴事件 {len(pack.timeline)} 条")
            if pack.external_evidence:
                ref_lines.append(f"- 外部证据卡 {len(pack.external_evidence)} 条")
            if pack.method_cards:
                ref_lines.append(f"- 方法论卡片 {len(pack.method_cards)} 张")
        except Exception:
            pass
        if ref_lines:
            priority_sections.append("## 数据源 3 · 结构化 evidence\n\n" + "\n".join(ref_lines))

        evidence_section = "\n\n---\n\n".join(priority_sections) if priority_sections else "- (数据中心 0 资料, 请基于下方 plan_text 设定写, 并明确标注'数据源: 用户指令上下文')"

        # P0-3 时间注入 + P0-4 元话语清洁 + P0-2 字数 + v2 排版规范 + v3 数据源优先级
        prompt = (
            f"# 当前日期 (必须用这个, 不要自己编日期)\n\n"
            f"{today_str}\n\n"
            f"# 顾源源给你的完整指令\n\n"
            f"```\n{plan.get('plan_text') or ''}\n```\n\n"
            f"# 你现在要交付的这一件\n\n"
            f"- 任务: {action}\n"
            f"- 预期交付: {expected_result or '按你专业判断'}\n"
            f"- 文档类型: {doc_type}\n"
            f"- 目标字数: **必须达到 {target_words} 字** (不允许少于 80%, 也不允许凑字)\n"
            f"- 关联客户: {client_name} (id={client_id})\n"
            f"- **最终交付格式: Word 文档 (.docx)**, 系统会从你输出的 markdown 自动转换\n\n"
            f"# 数据中心目前关于 {client_name} 的资料 (按优先级排序)\n\n"
            f"{evidence_section}\n\n"
            f"# 数据源使用硬规则 (顾源源 5/25 PM 钦定, 必须严格遵守)\n\n"
            f"1. **优先用数据源 1** (客户工作台官方文档) 当主要依据 — 这是客户授权的真实资料, 引用时直接陈述\n"
            f"2. **其次用数据源 2** (资讯情报站) 当补充 — 引用时标注 \"(据公开资料/{client_name}公众号)\"\n"
            f"3. **再次用数据源 3** (结构化 evidence) 做交叉验证\n"
            f"4. **最后才靠 plan_text 设定推理** — 如果上述 1/2/3 全空, 必须在文档开头标注 \"数据源: 用户指令上下文, 待客户确认\"\n"
            f"5. **禁止编造与数据源冲突的内容** — 数据源里写了什么就引用什么, 不要凭印象/通用知识改写\n"
            f"6. **数据源 1 有内容时, 优先级 100% 高于 LLM 自己的通用知识** (你不能引用过去训练数据里有但客户工作台里没有的所谓事实)\n\n"
            f"# Markdown 排版硬规则 (顾源源 5/25 真用反馈, 必须严格遵守)\n\n"
            f"1. **章节标题必须用 ## (二级标题)**, 不允许用 **粗体** 当标题\n"
            f"   正确: `## 集团历史沿革`\n"
            f"   错误: `**集团历史沿革**`\n"
            f"2. **每段不超过 200 字**, 长内容必须分多段 (空行隔开), 不允许 1000 字大段挤一团\n"
            f"3. **列举内容必须用 markdown 列表** (- 或 1./2./3.)\n"
            f"   正确: \n"
            f"   ```\n"
            f"   - 第一项: ...\n"
            f"   - 第二项: ...\n"
            f"   ```\n"
            f"   错误: 用文字流写\"第一是...第二是...第三是...\"\n"
            f"4. 重要数据用 **粗体** 内联强调 (但 不能 整段加粗或当标题)\n"
            f"5. 文档结构示例:\n"
            f"   ```\n"
            f"   # 安然集团客户背景档案\n\n"
            f"   生成者: 庆华 · 时间: {today_str} · 数据源: 用户指令上下文\n\n"
            f"   ## 集团历史沿革\n\n"
            f"   安然集团 1998 年成立于呼和浩特, 由 17 户养殖户联合发起.\n\n"
            f"   2010 年 A 股上市, 募资 32 亿元用于全国牧场布局.\n\n"
            f"   ## 业务结构\n\n"
            f"   - **上游**: 自有牧场 82 座, 原奶自给率 72%\n"
            f"   - **中游**: 加工基地 29 个, 年产能 1100 万吨\n"
            f"   - **下游**: 国内 120 万终端 + 海外 12 国\n"
            f"   ```\n\n"
            f"# 其他硬规则\n\n"
            f"1. 文档开头一行 metadata (无 # 号): `生成者: {bot.get('display_name', '庆华')} · 时间: {today_str} · 数据源: ...`\n"
            f"2. 必须达到 {target_words} 字, 章节数据要密集 (具体数字 / 真业务术语), 不要泛泛而谈\n"
            f"3. **禁止任何自我反思 / 自言自语 / 元话语** (如 \"(哦不对)\" \"(以下为...)\" \"(全文约X字)\" \"(请审阅)\" 等)\n"
            f"4. 如果是合同/协议类, 必须有完整条款结构 (合作期限/服务范围/付款节点/解约/IP/保密/争议解决), 每一条用 ## H2 区分\n"
            f"5. 直接给最终 markdown, 不要任何前言/总结/\"以下是...\" 这种引导句\n\n"
            f"现在开始写:"
        )

        try:
            raw = ai._qwen_generate(  # noqa: SLF001
                prompt=prompt,
                system_instruction=system_instruction,
                response_schema=None,   # 不要 JSON, 出纯 markdown
                timeout_seconds=300.0,  # 单份文档允许 5 min (长文档 LLM 慢)
                max_tokens=max_tokens,
                temperature=0.45,
                task_kind="default",
            )
            markdown = str(raw or "").strip()
            llm_used = ai.current_provider() if hasattr(ai, "current_provider") else "unknown"
        except Exception as exc:  # noqa: BLE001
            llm_error = str(exc)
            logger.warning("documents.generate LLM call failed: %s", exc)
            # 不抛, fallback 写一条 markdown stub 进 documents 表 (顾源源能看到失败原因)
            markdown = (
                f"# {expected_result or action}\n\n"
                f"> ⚠ LLM 生成失败: {llm_error}\n>\n"
                f"> 任务: {action}\n> 预期: {expected_result}\n> 客户: {client_name}\n\n"
                f"## 数据中心 evidence\n\n{evidence_section}\n\n"
                f"## plan_text 上下文\n\n```\n{plan.get('plan_text') or ''}\n```\n"
            )
            llm_used = "fallback_stub"
    else:
        # ai_service 没注入 (启动顺序问题), 写 fallback
        llm_used = "no_ai_service"
        markdown = (
            f"# {expected_result or action}\n\n"
            f"> ⚠ AIService 未注入 plan_executor (启动顺序问题, 见 set_ai_service)\n\n"
            f"## plan_text 上下文\n\n```\n{plan.get('plan_text') or ''}\n```\n"
        )

    # P0-4 元话语清洁: 删 LLM 自言自语行 (regex 抓 5 类常见模式)
    markdown = strip_llm_meta_lines(markdown)

    # 4. 推 title (顾源源 5/25 反馈: 不要用 expected_result, 它是落点描述)
    doc_id = f"doc_ai_{uuid.uuid4().hex[:20]}"
    now = _now_iso()
    title = ""
    m = re.search(r"(?:写一份|拟一份|起草一份|做一份|生成一份)\s*([^,，。.;；]{2,40})", action)
    if m:
        title = m.group(1).strip()
    if not title:
        cn_map = {
            "client_background_archive": "客户背景档案",
            "opportunity_analysis": "项目机会分析",
            "brand_analysis": "品牌分析",
            "strategic_partnership_proposal": "战略陪伴提案",
            "strategic_partnership_contract": "战略陪伴合作协议",
            "board_brief": "理事会简报",
            "meeting_prep": "会议准备",
            "review_doc": "复盘报告",
            "general_draft": "AI 草稿",
        }
        title = f"{client_name} · {cn_map.get(doc_type, doc_type)}"
    title = title[:80]

    # 5. P0-1 真出 .docx 文件 + 中文文件名 (顾源源 5/25 PM 反馈)
    # 文件名: {client_name}_{title}_{YYYYMMDD-HHMM}.docx
    # 比如 安然集团_客户背景档案_20260525-1247.docx
    # data_dir 从 db.db_path.parent 推 (Database.db_path 是真 sqlite 文件路径)
    docx_real_path: str = ""
    docx_save_error: str | None = None
    try:
        data_dir: Path | None = None
        if hasattr(db, "db_path"):
            data_dir = Path(getattr(db, "db_path")).parent  # type: ignore[arg-type]
        if data_dir is None:
            data_dir = Path.home() / "Library" / "Application Support" / "YiyuThinkTankWorkbench2_V21Lab"
        out_dir = data_dir / "ai_drafts" / client_id

        # 中文文件名: 客户名_文档标题_时间戳.docx, sanitize 不安全字符
        ts = datetime.now().strftime("%Y%m%d-%H%M%S")
        # title 可能含"安然集团"(跟 client_name 重复), 智能去重
        title_for_file = title
        if client_name and client_name in title_for_file:
            title_for_file = title_for_file.replace(client_name, "").strip(" ·-_")
        if not title_for_file:
            title_for_file = title
        filename = f"{_safe_filename_segment(client_name)}_{_safe_filename_segment(title_for_file)}_{ts}.docx"
        out_path = out_dir / filename
        # 同名冲突时加 _N 后缀
        counter = 1
        while out_path.exists():
            counter += 1
            stem = f"{_safe_filename_segment(client_name)}_{_safe_filename_segment(title_for_file)}_{ts}_v{counter}"
            out_path = out_dir / f"{stem}.docx"

        markdown_to_docx(markdown, out_path, title=title)
        docx_real_path = str(out_path)
    except Exception as exc:  # noqa: BLE001
        docx_save_error = str(exc)
        logger.warning("markdown_to_docx failed: %s", exc)

    # path 字段: 优先真 .docx 路径 (用户能下载), 失败时回退虚拟 .md 路径
    file_path = docx_real_path or f"docs/ai_drafts/{doc_id}.md"
    excerpt = markdown[:200].replace("\n", " ").strip() or "(空)"
    tags = ["ai_draft", "bot:" + (bot.get("handle") or "unknown"), f"plan:{plan.get('id', '')[:24]}"]
    if docx_real_path:
        tags.append("format:docx")
    source = f"ai_command:{bot.get('handle') or 'bot'}"

    # [B] 5/25 PM bug fix: 用 run_in_transaction 强制独立提交, 防嵌套事务 race 丢失.
    _doc_insert_sql = """INSERT INTO documents (
        id, client_id, folder_id, title, path, original_source_path,
        kind, source, excerpt, tags_json,
        created_at, document_family_id, canonical_kind, origin_type, origin_id,
        is_searchable, organization_id, owner_user_id, source_entity_type, source_entity_id,
        visibility_scope, content_domain, lifecycle_status, document_role
    ) VALUES (?, ?, NULL, ?, ?, NULL,
              ?, ?, ?, ?,
              ?, '', 'ai_draft', 'bot_member', ?,
              1, '', ?, 'ai_task_plan', ?,
              'project_public', 'work', 'active', 'draft')"""
    _doc_insert_params = (
        doc_id, client_id, title, file_path,
        doc_type, source, excerpt, json.dumps(tags, ensure_ascii=False),
        now, bot.get("id") or "",
        plan.get("human_initiator_id") or "",
        plan.get("id") or "",
    )

    def _do_insert_doc(conn) -> None:
        conn.execute(_doc_insert_sql, _doc_insert_params)
        # 同事务里把完整 markdown 写进 excerpt (前 8000 字), 真 .docx 在 path 字段
        conn.execute("UPDATE documents SET excerpt = ? WHERE id = ?",
                     (markdown[:8000], doc_id))

    if hasattr(db, "run_in_transaction"):
        db.run_in_transaction(_do_insert_doc)  # type: ignore[attr-defined]
    else:
        db.execute(_doc_insert_sql, _doc_insert_params)
        db.execute("UPDATE documents SET excerpt = ? WHERE id = ?",
                   (markdown[:8000], doc_id))
        if hasattr(db, "conn") and hasattr(getattr(db, "conn"), "commit"):
            getattr(db, "conn").commit()

    # ── [B] 5/25 PM · path D · 庆华跟人走一套接口 ──
    # 1. 给本次文档建一个 task (owner=庆华, related=顾源源 协作者)
    # 2. 立刻 status='done' + completed_at (因为文档已生成)
    # 3. 给庆华本周 review 加一条复盘 entry (它自己写反思)
    # 4. 同时给发起人 review 也加一条 (内容相同, 让顾源源进自己复盘也看到)
    bot_actor_id = bot.get("actor_id") or ""
    bot_display_name = bot.get("display_name") or "AI 同事"
    human_initiator_id = plan.get("human_initiator_id") or ""
    task_id_for_doc = f"task_{uuid.uuid4().hex[:24]}"
    now_iso = _now_iso()

    try:
        # 任务的 ddl/截止: 文档已完成, 用 now
        task_insert_sql = """INSERT INTO tasks (
            id, title, description, status, progress_status, priority, list_id,
            owner_id, owner_name, ddl, deadline_at, completed_at, due_date, duration_minutes,
            source_type, source_id,
            client_id, scope_mode, evidence_count,
            tags_json, tag_ids_json, sync_status,
            created_at, updated_at, creator_id
        ) VALUES (?, ?, ?, 'done', 'done', 'normal', 'list-1',
                  ?, ?, ?, ?, ?, ?, 0,
                  'ai_plan_executor', ?,
                  ?, 'COLLAB_SHARED', 0,
                  '[]', '[]', 'local',
                  ?, ?, ?)"""
        task_desc = f"AI 同事 {bot_display_name} 完成 · 文档 ID: {doc_id} · 字数 {len(markdown)} · LLM={llm_used}"
        task_insert_params = (
            task_id_for_doc, title[:80], task_desc,
            bot_actor_id, bot_display_name, now_iso, now_iso, now_iso, now_iso,
            plan.get("id") or "",
            client_id,
            now_iso, now_iso, bot_actor_id,
        )

        def _do_create_doc_task(conn) -> None:
            conn.execute(task_insert_sql, task_insert_params)
            # 加发起人协作者 — 顾源源能在日历/任务列表看到
            if human_initiator_id:
                # 查 user 真名 (兜底用 id)
                user_row = conn.execute(
                    "SELECT full_name FROM mirror_users WHERE id = ?", (human_initiator_id,),
                ).fetchone()
                human_name = (dict(user_row).get("full_name") if user_row else human_initiator_id) or human_initiator_id
                conn.execute(
                    """INSERT OR IGNORE INTO task_collaborators (
                        task_id, organization_id, user_id, full_name, email,
                        order_index, is_owner, inbox_status, created_at, updated_at
                    ) VALUES (?, '', ?, ?, '', 0, 0, 'accepted', ?, ?)""",
                    (task_id_for_doc, human_initiator_id, human_name, now_iso, now_iso),
                )

        if hasattr(db, "run_in_transaction"):
            db.run_in_transaction(_do_create_doc_task)  # type: ignore[attr-defined]
        else:
            _do_create_doc_task(getattr(db, "conn", db))
            if hasattr(db, "conn") and hasattr(getattr(db, "conn"), "commit"):
                getattr(db, "conn").commit()
    except Exception as exc:  # noqa: BLE001
        logger.warning("path_d_create_doc_task failed: %s", exc)

    # 庆华自动写复盘 + 同步给发起人
    try:
        review_note = _bot_write_self_review(
            bot_display_name=bot_display_name,
            action=action,
            output_summary=f"为 {client_name} 生成《{title}》({len(markdown)} 字, .docx 已存盘)" if docx_real_path else f"为 {client_name} 生成《{title}》({len(markdown)} 字)",
            evidence_used=evidence_summary,
            llm_used=llm_used,
        )
        week_label = _current_week_label()
        task_snapshot = {
            "title": title,
            "document_id": doc_id,
            "docx_path": docx_real_path,
            "markdown_chars": len(markdown),
            "client_id": client_id,
            "bot_actor_id": bot_actor_id,
            "bot_display_name": bot_display_name,
        }
        # 1. 庆华自己本周复盘
        bot_review_id = _ensure_weekly_review(db, user_id=bot_actor_id, week_label=week_label)
        _write_review_entry(
            db, review_id=bot_review_id, task_id=task_id_for_doc,
            user_id=bot_actor_id, week_label=week_label,
            note=review_note, task_snapshot=task_snapshot,
        )
        # 2. 发起人 (顾源源) 本周复盘 — 同一 task 也出现在他的复盘里
        if human_initiator_id:
            human_review_id = _ensure_weekly_review(db, user_id=human_initiator_id, week_label=week_label)
            _write_review_entry(
                db, review_id=human_review_id, task_id=task_id_for_doc,
                user_id=human_initiator_id, week_label=week_label,
                note=review_note, task_snapshot=task_snapshot,
            )
    except Exception as exc:  # noqa: BLE001
        logger.warning("path_d_auto_review failed: %s", exc)

    return {
        "client_id": client_id,
        "document_id": doc_id,
        "document_title": title,
        "document_type": doc_type,
        "task_id": task_id_for_doc,
        "markdown_chars": len(markdown),
        "docx_path": docx_real_path,
        "docx_save_error": docx_save_error,
        "llm_used": llm_used,
        "llm_error": llm_error,
        "target_words": target_words,
        "word_completion_pct": int(len(markdown) * 100 / target_words) if target_words else 0,
        "evidence_summary": evidence_summary,
        "summary_text": (
            f"已为 {client_name} 生成《{title}》({len(markdown)}/{target_words} 字, "
            f"{'.docx 已存盘' if docx_real_path else 'markdown only'}, LLM={llm_used}, "
            f"任务 {task_id_for_doc} 已自动 done + 复盘)"
        ),
    }


def _handler_tasks_create(
    db: _DbLike, *, payload: dict, bot: dict, plan: dict,
) -> dict:
    """[B] 2026-05-25 PM 重写 (顾源源拍板"真填字段, 不要 plan_text 前 30 字").

    输入约定 (来自 _step_to_subtask, B 5/25 转换):
      payload = { client_id, module, action(中文), expected_result }

    新逻辑:
      title       = payload.expected_result || payload.action (不再 fallback plan_text)
      description = payload.action (具体动作描述, 不留空)
      ddl/due_date = parse_chinese_datetime(plan_text) → ISO 8601 (不再"待确认")
      owner_id   = plan.human_initiator_id (顾源源真发起)
      owner_name = mirror_users 真查 (不再空字符串)
      creator_id = bot.id (庆华建的, 不是匿名)
    """
    # 1. title - 顾源源 5/25 真用 bug: payload.expected_result = step.deliverable
    # = "落点: 我日程 · 完成进我审批" (落点不是任务名). 改成从 plan_text 抓"建一个任务: XXX"
    # 真任务描述, fallback action.
    title = (payload.get("title") or "").strip()
    if not title:
        plan_text = plan.get("plan_text") or ""
        # 抓"给我建一个任务: 5月27日下午2点跟安然集团首次正式会议" 中冒号后的描述
        m = re.search(
            r"(?:给我|帮我)?建一?个?\s*(?:任务|会议|事项)[:：]?\s*([^,，。.;；\n]{4,80})",
            plan_text,
        )
        if m:
            title = m.group(1).strip()
    if not title:
        # plan_text 也抓不到, 用 action
        title = (payload.get("action") or "").strip()
    if not title:
        title = (plan.get("plan_title") or "").strip()
    if not title:
        raise ValueError("tasks.create 需 title (payload + plan_text + plan_title 都没)")
    title = title[:80]

    # 2. description - 用 action 的完整描述
    desc = (payload.get("desc")
            or payload.get("description")
            or payload.get("action")
            or payload.get("expected_result")
            or "").strip()

    # 3. 时间解析 — 从 plan_text 抓中文时间, 失败回退 None
    plan_text = plan.get("plan_text") or ""
    parsed_dt = parse_chinese_datetime(plan_text)
    due_date = payload.get("due_date") or payload.get("dueDate") or parsed_dt
    deadline_at = payload.get("deadline_at") or payload.get("deadlineAt") or parsed_dt
    ddl = deadline_at or due_date or "待确认"

    # [B] 5/25 PM (path D 顾源源洞察) · AI 当真同事 · 走人一样的接口:
    # · owner = 庆华 (它真做这件事)
    # · 顾源源 = collaborator (日历/任务列表可见, 但归属权属人)
    # · 会议任务保留 status='todo' (要真人参加, 不 auto-done)
    bot_actor_id = bot.get("actor_id") or ""
    bot_display_name = (bot.get("display_name") or "AI 同事").strip()
    owner_user_id = bot_actor_id
    owner_name = bot_display_name
    human_initiator_id = (plan.get("human_initiator_id") or "").strip()
    # 兜底: 取人名 (复盘归属用)
    human_name = ""
    if human_initiator_id:
        try:
            urow = db.fetchone(
                "SELECT full_name FROM org_members_v WHERE id = ?",
                (human_initiator_id,),
            )
            if urow:
                human_name = (dict(urow).get("full_name") or "").strip()
        except Exception:
            pass
        if not human_name:
            human_name = human_initiator_id

    list_id = (payload.get("listId") or payload.get("list_id") or "list-1").strip()
    client_id = payload.get("client_id") or plan.get("client_id")

    task_id = f"task_{uuid.uuid4().hex[:24]}"
    now = _now_iso()
    actor_id = bot.get("actor_id") or ""

    # [B] 5/25 PM bug fix: 5 个 ai 建的 task INSERT 成功但事后查 0 条 — 嵌套事务 race.
    # BackgroundTask plan_executor 跟其他 HTTP request 共享 _in_transaction flag,
    # 外层 rollback 时本 INSERT 也丢失. 用 run_in_transaction 强制独立提交.
    _insert_sql = """INSERT INTO tasks (
        id, title, description, status, priority, list_id,
        owner_id, owner_name, ddl, deadline_at, scheduled_start_at,
        scheduled_end_at, completed_at, start_date, due_date, duration_minutes,
        event_line_id, source_type, source_id,
        client_id, project_module_id, project_flow_id, scope_mode,
        business_category, current_blocker, next_action, recent_decision, evidence_count,
        tags_json, tag_ids_json, sync_status, created_at, updated_at,
        creator_id
    ) VALUES (?, ?, ?, 'todo', 'normal', ?, ?, ?, ?, ?, NULL, NULL, NULL, NULL, ?, 60, NULL, ?, ?, ?, NULL, NULL, 'COLLAB_SHARED', NULL, NULL, NULL, NULL, 0, '[]', '[]', 'local', ?, ?, ?)"""
    _insert_params = (
        task_id, title, desc, list_id, owner_user_id, owner_name,
        ddl, deadline_at,
        due_date,
        "ai_plan_executor", plan.get("id"),
        client_id, now, now,
        bot.get("id") or actor_id,
    )

    def _do_insert_task(conn) -> None:
        conn.execute(_insert_sql, _insert_params)
        # [B] path D: 加发起人为 collaborator (让顾源源在日历/任务列表看到这条会议)
        if human_initiator_id:
            conn.execute(
                """INSERT OR IGNORE INTO task_collaborators (
                    task_id, organization_id, user_id, full_name, email,
                    order_index, is_owner, inbox_status, created_at, updated_at
                ) VALUES (?, '', ?, ?, '', 0, 0, 'accepted', ?, ?)""",
                (task_id, human_initiator_id, human_name, now, now),
            )

    if hasattr(db, "run_in_transaction"):
        db.run_in_transaction(_do_insert_task)  # type: ignore[attr-defined]
    else:
        db.execute(_insert_sql, _insert_params)
        # 兜底: 不依赖嵌套事务的 db wrapper, 强制 commit (如 test fake db)
        if hasattr(db, "conn") and hasattr(getattr(db, "conn"), "commit"):
            getattr(db, "conn").commit()

    # [B] path D: 庆华自动写复盘 (会议类任务不 auto-done, 但要复盘"为什么建这个会议")
    try:
        review_note = _bot_write_self_review(
            bot_display_name=bot_display_name,
            action=payload.get("action") or title,
            output_summary=f"建任务: {title} (ddl={ddl})",
            evidence_used={},  # tasks.create 不依赖 evidence
            llm_used="rule_based",
        )
        week_label = _current_week_label()
        task_snapshot = {
            "title": title, "ddl": ddl, "due_date": due_date,
            "client_id": client_id, "bot_actor_id": bot_actor_id,
            "datetime_parsed_from_text": bool(parsed_dt),
        }
        # 1. 庆华本周复盘
        bot_review_id = _ensure_weekly_review(db, user_id=bot_actor_id, week_label=week_label)
        _write_review_entry(
            db, review_id=bot_review_id, task_id=task_id,
            user_id=bot_actor_id, week_label=week_label,
            note=review_note, task_snapshot=task_snapshot,
        )
        # 2. 发起人本周复盘
        if human_initiator_id:
            human_review_id = _ensure_weekly_review(db, user_id=human_initiator_id, week_label=week_label)
            _write_review_entry(
                db, review_id=human_review_id, task_id=task_id,
                user_id=human_initiator_id, week_label=week_label,
                note=review_note, task_snapshot=task_snapshot,
            )
    except Exception as exc:  # noqa: BLE001
        logger.warning("path_d tasks.create auto-review failed: %s", exc)

    return {
        "task_id": task_id,
        "title": title,
        "owner_id": owner_user_id,
        "owner_name": owner_name,
        "ddl": ddl,
        "due_date": due_date,
        "client_id": client_id,
        "created_by_bot": actor_id,
        "datetime_parsed_from_text": bool(parsed_dt),
        "summary_text": f"已建任务 {task_id}: {title} (owner={owner_name}, ddl={ddl})",
    }


def _handler_smart_import(
    db: _DbLike, *, payload: dict, bot: dict, plan: dict,
) -> dict:
    """smart_import 链路 — M9 暂占位, 后续接 data_center_ingest 真实现.

    不抛异常, 标 noop_unsupported 让进度走完 (顾源源约束: 不卡死).
    """
    logger.info("smart_import not yet wired in M9 — marking noop_unsupported (plan=%s)", plan.get("id"))
    return {
        "status": "noop_unsupported",
        "tool": "smart_import",
        "reason": "M9 未接入 smart_import handler, 留给后续 milestone",
        "payload_received": payload,
    }


def _handler_noop(
    db: _DbLike, *, payload: dict, bot: dict, plan: dict,
) -> dict:
    """fallback — 未注册 tool 走这里, 不崩溃, 标 unsupported."""
    return {
        "status": "unsupported_tool",
        "tool_requested": payload.get("__tool_requested"),
        "note": "未注册的 tool, 已 log 但未执行",
    }


def _build_default_registry() -> ExecutorRegistry:
    reg = ExecutorRegistry()
    reg.register("documents.generate", _handler_documents_generate)
    reg.register("tasks.create", _handler_tasks_create)
    reg.register("smart_import", _handler_smart_import)
    reg.register("noop", _handler_noop)
    return reg


# 模块级单例 (服务启动时建一次)
_REGISTRY: ExecutorRegistry | None = None


def get_registry() -> ExecutorRegistry:
    global _REGISTRY
    if _REGISTRY is None:
        _REGISTRY = _build_default_registry()
    return _REGISTRY


# ─── Plan progress 更新 ─────────────────────────────────────────


def _update_plan_progress(
    db: _DbLike, plan_id: str, *,
    execution_status: str | None = None,
    progress: dict | None = None,
    summary: dict | None = None,
    started: bool = False,
    completed: bool = False,
) -> None:
    """更新 ai_task_plans 的 execution 字段. 缺列时 schema migration 兜底."""
    now = _now_iso()
    sets: list[str] = []
    params: list = []
    if execution_status is not None:
        sets.append("execution_status = ?"); params.append(execution_status)
    if progress is not None:
        sets.append("progress_json = ?"); params.append(json.dumps(progress, ensure_ascii=False))
    if summary is not None:
        sets.append("execution_summary_json = ?"); params.append(json.dumps(summary, ensure_ascii=False))
    if started:
        sets.append("execution_started_at = COALESCE(execution_started_at, ?)")
        params.append(now)
    if completed:
        sets.append("execution_completed_at = ?"); params.append(now)
    if not sets:
        return
    sets.append("updated_at = ?"); params.append(now)
    try:
        db.execute(
            f"UPDATE ai_task_plans SET {', '.join(sets)} WHERE id = ?",
            tuple(params + [plan_id]),
        )
    except Exception as exc:
        logger.warning("update plan progress failed (likely schema not migrated): %s", exc)


def ensure_execution_schema(db: _DbLike) -> None:
    """M10 schema: ai_task_plans 加 5 列. ALTER + try/except (idempotent)."""
    alters = [
        "ALTER TABLE ai_task_plans ADD COLUMN execution_status TEXT NOT NULL DEFAULT 'not_started'",
        "ALTER TABLE ai_task_plans ADD COLUMN execution_started_at TEXT",
        "ALTER TABLE ai_task_plans ADD COLUMN execution_completed_at TEXT",
        "ALTER TABLE ai_task_plans ADD COLUMN progress_json TEXT NOT NULL DEFAULT '{}'",
        "ALTER TABLE ai_task_plans ADD COLUMN execution_summary_json TEXT NOT NULL DEFAULT '{}'",
    ]
    for sql in alters:
        try:
            db.execute(sql)
        except Exception as exc:
            # duplicate column → 已加过, 静默跳
            if "duplicate column" not in str(exc).lower():
                logger.warning("ensure_execution_schema stmt fail: %s", exc)


# ─── 主执行函数 ─────────────────────────────────────────────────


_MAX_RETRIES = 3


def execute_plan(plan_id: str, db: _DbLike) -> dict:
    """主入口 — 取 plan, 必须 approved, 解析 subtasks, 顺序执行, 全程留痕.

    返回执行结果摘要. 异常被吞包成 status='failed', 不让 BackgroundTasks 静默死.
    """
    from app.services.agent_governance import (
        log_agent_run_start, log_agent_run_complete, ensure_governance_schema,
    )
    ensure_governance_schema(db)
    ensure_execution_schema(db)

    # 1. 取 plan
    row = db.fetchone("SELECT * FROM ai_task_plans WHERE id = ?", (plan_id,))
    if not row:
        logger.error("execute_plan: plan not found id=%s", plan_id)
        return {"plan_id": plan_id, "execution_status": "failed", "error": "plan not found"}
    plan = dict(row)
    if plan.get("status") != "approved":
        logger.warning("execute_plan: plan status not approved (%s) — skip", plan.get("status"))
        return {"plan_id": plan_id, "execution_status": "skipped", "reason": f"status={plan.get('status')}"}

    # 2. 取 bot (actor_id 必须真存在, 不能 anonymous)
    from app.services.bot_members import get_bot_member
    bot = get_bot_member(db, plan.get("bot_member_id"))
    if not bot:
        logger.error("execute_plan: bot not found bot_member_id=%s", plan.get("bot_member_id"))
        _update_plan_progress(
            db, plan_id, execution_status="failed",
            completed=True,
            summary={"error": "bot not found"},
        )
        return {"plan_id": plan_id, "execution_status": "failed", "error": "bot not found"}
    actor_id = bot.get("actor_id") or ""

    # 3. 解析 subtasks
    subtasks = _parse_subtasks(plan)
    total = len(subtasks)
    client_id = plan.get("client_id")

    _update_plan_progress(
        db, plan_id, execution_status="running",
        started=True,
        progress={
            "total": total, "completed": 0, "current": "starting",
            "percent": 0, "errors": [],
        },
    )

    # plan 级 agent_run_log (一条总记录, 关联到客户)
    plan_run_id = log_agent_run_start(
        db,
        actor_type="internal_ai_agent",
        actor_id=actor_id,
        tool_name="plan_executor.run",
        client_id=client_id,
        input_payload={
            "plan_id": plan_id,
            "plan_title": plan.get("plan_title"),
            "subtask_count": total,
        },
        idempotency_key=f"plan_executor:{plan_id}",
    )

    # 4. 顺序执行每个 subtask
    registry = get_registry()
    summary_items: list[dict] = []
    errors: list[dict] = []
    success_count = 0

    for sub in subtasks:
        idx = sub["index"]
        tool = sub["tool"]
        payload = dict(sub.get("payload") or {})
        payload["__tool_requested"] = tool
        label = sub["label"]

        _update_plan_progress(
            db, plan_id,
            progress={
                "total": total, "completed": success_count,
                "current": label, "percent": int(success_count * 100 / total) if total else 0,
                "errors": list(errors),
            },
        )

        # subtask 级 agent_run_log
        sub_run_id = log_agent_run_start(
            db,
            actor_type="internal_ai_agent",
            actor_id=actor_id,
            tool_name=tool,
            client_id=client_id,
            input_payload={"plan_id": plan_id, "subtask_index": idx, "label": label, "payload": payload},
        )

        # 真 handler 调用, 失败重试上限 _MAX_RETRIES
        handler = registry.get(tool)
        result: dict = {}
        last_err: str | None = None
        ok = False
        start_ms = time.time()
        for attempt in range(1, _MAX_RETRIES + 1):
            try:
                result = handler(db, payload=payload, bot=bot, plan=plan)
                ok = True
                break
            except Exception as exc:
                last_err = f"attempt {attempt}: {exc}"
                logger.warning("subtask %s attempt %d failed: %s", tool, attempt, exc)
                time.sleep(min(0.5 * attempt, 2.0))  # 指数退避, 上限 2s

        duration_ms = int((time.time() - start_ms) * 1000)

        if ok:
            success_count += 1
            output_summary = _short_summary(result)
            log_agent_run_complete(
                db, sub_run_id,
                output_payload=result, status="success",
                duration_ms=duration_ms,
            )
            summary_items.append({
                "index": idx, "tool": tool, "status": "success",
                "output_summary": output_summary,
                "duration_ms": duration_ms,
            })
        else:
            errors.append({"index": idx, "tool": tool, "error": last_err or "unknown"})
            log_agent_run_complete(
                db, sub_run_id,
                output_payload={}, status="failed",
                error_message=last_err or "unknown",
                duration_ms=duration_ms,
            )
            summary_items.append({
                "index": idx, "tool": tool, "status": "failed",
                "error": last_err or "unknown",
                "duration_ms": duration_ms,
            })

    # 5. 终态计算
    if success_count == total and total > 0:
        final_status = "success"
    elif success_count == 0:
        final_status = "failed"
    elif total == 0:
        final_status = "success"  # 0 subtask = noop, 不算失败
        summary_items.append({"index": 0, "tool": "noop", "status": "success",
                              "output_summary": "plan 无 subtask, 标 noop"})
    else:
        final_status = "partial"

    final_progress = {
        "total": total,
        "completed": success_count,
        "current": "done",
        "percent": 100 if final_status == "success" else int(success_count * 100 / max(total, 1)),
        "errors": errors,
    }
    final_summary = {
        "subtasks": summary_items,
        "errors": errors,
        "success_count": success_count,
        "total_count": total,
    }
    _update_plan_progress(
        db, plan_id,
        execution_status=final_status,
        completed=True,
        progress=final_progress,
        summary=final_summary,
    )

    # plan 级 agent_run_log 收尾
    log_agent_run_complete(
        db, plan_run_id,
        output_payload={
            "plan_id": plan_id,
            "final_status": final_status,
            "success_count": success_count,
            "total_count": total,
            "errors": errors,
        },
        status="success" if final_status == "success" else (
            "failed" if final_status == "failed" else "success"
        ),
    )

    return {
        "plan_id": plan_id,
        "execution_status": final_status,
        "subtask_summary": summary_items,
        "errors": errors,
    }


def _short_summary(result: dict) -> str:
    """从 handler result 抽一段人话摘要 (供 UI 显示)."""
    if not isinstance(result, dict):
        return str(result)[:200]
    if "summary_text" in result:
        return str(result["summary_text"])[:200]
    if "task_id" in result:
        return f"已建任务 {result.get('task_id')} ({result.get('title', '无标题')})"
    if result.get("status") == "noop_unsupported":
        return f"未接入: {result.get('reason', '')}"
    if result.get("status") == "unsupported_tool":
        return f"未注册 tool: {result.get('tool_requested')}"
    # 兜底: 取前 200 字
    try:
        return json.dumps(result, ensure_ascii=False)[:200]
    except Exception:
        return str(result)[:200]


# ─── 查询: progress endpoint 用 ─────────────────────────────────


def get_plan_progress(db: _DbLike, plan_id: str) -> dict | None:
    """供 GET /api/v1/org/bots/task-plans/{id}/progress 用."""
    ensure_execution_schema(db)
    row = db.fetchone("SELECT * FROM ai_task_plans WHERE id = ?", (plan_id,))
    if not row:
        return None
    plan = dict(row)
    progress_raw = plan.get("progress_json") or "{}"
    summary_raw = plan.get("execution_summary_json") or "{}"
    try:
        progress = json.loads(progress_raw) if progress_raw else {}
    except Exception:
        progress = {}
    try:
        summary = json.loads(summary_raw) if summary_raw else {}
    except Exception:
        summary = {}

    return {
        "plan_id": plan_id,
        "plan_status": plan.get("status"),
        "execution_status": plan.get("execution_status") or "not_started",
        "started_at": plan.get("execution_started_at"),
        "completed_at": plan.get("execution_completed_at"),
        "progress": progress or {
            "total": 0, "completed": 0, "current": "",
            "percent": 0, "errors": [],
        },
        "subtasks": summary.get("subtasks") or [],
        "errors": summary.get("errors") or [],
    }
