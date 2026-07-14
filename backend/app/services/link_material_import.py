from __future__ import annotations

import importlib.util
import ipaddress
import json
import logging
import os
import re
import shlex
import shutil
import socket
import subprocess
import sys
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Callable, Literal
from urllib.parse import urljoin, urlsplit, urlunsplit

logger = logging.getLogger(__name__)


LinkMaterialPlatform = Literal["bilibili", "xiaohongshu", "wechat_article"]
CookieBrowser = Literal["firefox", "chrome", "edge", "safari"]


class LinkMaterialImportError(Exception):
    """User-facing error raised by the link material import pipeline."""

    def __init__(self, message: str, *, metadata: dict[str, Any] | None = None) -> None:
        super().__init__(message)
        self.metadata = metadata or {}


@dataclass(frozen=True)
class LinkMaterialDetection:
    platform: LinkMaterialPlatform
    normalized_url: str
    display_name: str


@dataclass
class LinkMaterialSource:
    platform: LinkMaterialPlatform
    source_url: str
    title: str
    transcript_text: str
    metadata: dict[str, Any] = field(default_factory=dict)
    downloaded_paths: list[Path] = field(default_factory=list)


@dataclass(frozen=True)
class LinkMaterialImportOptions:
    use_browser_cookies: bool = False
    cookie_browser: CookieBrowser = "firefox"
    stage_callback: Callable[[str, float, dict[str, object] | None], None] | None = None


@dataclass(frozen=True)
class _DownloadAttemptProfile:
    name: str
    headers_applied: bool = False
    impersonation_requested: bool = False
    use_browser_cookies: bool = False


@dataclass(frozen=True)
class LocalTranscriptEngine:
    # builtin_sensevoice = 应用内置 sherpa-onnx + ONNX 模型（和录音转写共享）
    # local_sensevoice / local_whisper = PATH 中的命令行可执行文件（subprocess）
    name: Literal["builtin_sensevoice", "local_sensevoice", "local_whisper"]
    command: list[str]
    command_template: str | None = None


class LinkMaterialPlatformAdapter:
    platform: LinkMaterialPlatform
    display_name: str

    def extract(self, source_url: str, temp_dir: Path, *, options: LinkMaterialImportOptions | None = None) -> LinkMaterialSource:
        raise NotImplementedError


_BILIBILI_BV_RE = re.compile(r"^BV[0-9A-Za-z]{8,}$", re.I)
_LINK_PLATFORM_HOST_SUFFIXES: dict[LinkMaterialPlatform, tuple[str, ...]] = {
    "bilibili": ("bilibili.com", "b23.tv"),
    "xiaohongshu": ("xiaohongshu.com", "xhslink.com", "xhs.cn"),
    "wechat_article": ("mp.weixin.qq.com",),
}
_SAFE_REDIRECT_STATUS_CODES = {301, 302, 303, 307, 308}
_MAX_SAFE_LINK_REDIRECTS = 5
_AUDIO_EXTENSIONS = {".aac", ".aiff", ".alac", ".flac", ".m4a", ".mp3", ".oga", ".ogg", ".opus", ".wav", ".weba", ".wma"}
# SenseVoice 用 soundfile(libsndfile)解码，它只认 wav/flac/ogg/aiff 等；m4a/aac/mp3/opus/wma
# 这些要先用 ffmpeg 转成 wav 才能转写。这个集合是"可直接喂 soundfile"的安全格式。
_SOUNDFILE_SAFE_AUDIO_EXTENSIONS = {".wav", ".flac", ".ogg", ".oga", ".aiff", ".aif"}
_VIDEO_EXTENSIONS = {".avi", ".flv", ".m4v", ".mkv", ".mov", ".mp4", ".mpeg", ".mpg", ".webm", ".wmv"}
_NON_MEDIA_EXTENSIONS = {".ass", ".description", ".info.json", ".json", ".part", ".srt", ".ssa", ".txt", ".vtt"}
_SUPPORTED_COOKIE_BROWSERS = ["firefox", "chrome", "edge", "safari"]
_BILIBILI_USER_AGENT = (
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
    "AppleWebKit/537.36 (KHTML, like Gecko) "
    "Chrome/124.0.0.0 Safari/537.36"
)
_BILIBILI_HEADER_ARGS = [
    "--add-header",
    "Origin:https://www.bilibili.com",
    "--add-header",
    "Referer:https://www.bilibili.com/",
    "--add-header",
    f"User-Agent:{_BILIBILI_USER_AGENT}",
]
# 小红书反爬会校验 Referer/Origin/UA。匿名场景补上这些能提高公开视频笔记的下载成功率；
# 真正登录墙/私密内容仍需浏览器登录态(cookie)。
_XHS_HEADER_ARGS = [
    "--add-header",
    "Origin:https://www.xiaohongshu.com",
    "--add-header",
    "Referer:https://www.xiaohongshu.com/",
    "--add-header",
    f"User-Agent:{_BILIBILI_USER_AGENT}",
]


# 从一段文本里抓出第一个 http(s) URL。小红书/B站 App"复制链接"复制的是整段分享文案
# (标题 + 短链 + 推广语)，必须把真正的 URL 抠出来，否则会把中文文案当 URL 丢给 yt-dlp。
# URL 在中文里常被空白或中文标点收尾，这里以空白和常见中文标点作为终止符。
_URL_IN_TEXT_RE = re.compile(r"https?://[^\s，。、；！？（）()【】「」『』""''《》：]+", re.I)


def _extract_url_from_text(text: str) -> str | None:
    match = _URL_IN_TEXT_RE.search(text)
    return match.group(0) if match else None


def _hostname_matches_suffix(hostname: str, suffix: str) -> bool:
    return hostname == suffix or hostname.endswith(f".{suffix}")


def _validate_platform_url(
    value: str,
    *,
    expected_platform: LinkMaterialPlatform | None = None,
) -> LinkMaterialDetection:
    """Validate a user-controlled source URL without making a network request."""
    if any(ord(character) < 32 or ord(character) == 127 for character in value) or "\\" in value:
        raise LinkMaterialImportError("链接格式无效。")
    try:
        parsed = urlsplit(value)
        hostname = (parsed.hostname or "").lower()
        port = parsed.port
    except ValueError as exc:
        raise LinkMaterialImportError("链接格式无效。") from exc
    if parsed.scheme.lower() != "https":
        raise LinkMaterialImportError("为保护本机与组织数据，链接导入仅支持 HTTPS 链接。")
    if not hostname or parsed.username is not None or parsed.password is not None:
        raise LinkMaterialImportError("链接格式无效。")
    if hostname.endswith(".") or port not in (None, 443):
        raise LinkMaterialImportError("链接格式无效。")

    platform: LinkMaterialPlatform | None = None
    for candidate, suffixes in _LINK_PLATFORM_HOST_SUFFIXES.items():
        hostname_allowed = (
            hostname in suffixes
            if candidate == "wechat_article"
            else any(_hostname_matches_suffix(hostname, suffix) for suffix in suffixes)
        )
        if hostname_allowed:
            platform = candidate
            break
    if platform is None:
        raise LinkMaterialImportError("暂不支持这个链接。当前仅支持 B 站链接、BV 号、小红书链接、微信公众号文章。")
    if platform == "wechat_article" and not (parsed.path == "/s" or parsed.path.startswith("/s/")):
        raise LinkMaterialImportError("暂不支持这个公众号链接，请使用公众号文章链接。")
    if expected_platform is not None and platform != expected_platform:
        raise LinkMaterialImportError("链接跳转到了不受支持的平台。")

    # Fragments never reach the remote server. Dropping them also makes queueing
    # and audit records deterministic while preserving signed query parameters.
    normalized_url = urlunsplit(("https", hostname, parsed.path, parsed.query, ""))
    display_names = {"bilibili": "B站", "xiaohongshu": "小红书", "wechat_article": "公众号"}
    return LinkMaterialDetection(
        platform=platform,
        normalized_url=normalized_url,
        display_name=display_names[platform],
    )


def _resolve_hostname_addresses(hostname: str) -> set[str]:
    try:
        results = socket.getaddrinfo(
            hostname,
            443,
            family=socket.AF_UNSPEC,
            type=socket.SOCK_STREAM,
            proto=socket.IPPROTO_TCP,
        )
    except OSError as exc:
        raise LinkMaterialImportError("链接域名暂时无法安全解析，请稍后重试。") from exc
    return {str(result[4][0]).split("%", 1)[0] for result in results if result[4]}


def _assert_public_link_target(
    value: str,
    *,
    expected_platform: LinkMaterialPlatform,
) -> str:
    """Fail closed if an allowlisted hostname resolves to any non-public IP."""
    detection = _validate_platform_url(value, expected_platform=expected_platform)
    hostname = str(urlsplit(detection.normalized_url).hostname or "")
    addresses = _resolve_hostname_addresses(hostname)
    if not addresses:
        raise LinkMaterialImportError("链接域名暂时无法安全解析，请稍后重试。")
    try:
        unsafe = any(not ipaddress.ip_address(address).is_global for address in addresses)
    except ValueError as exc:
        raise LinkMaterialImportError("链接域名返回了无效地址，已停止导入。") from exc
    if unsafe:
        raise LinkMaterialImportError("该链接指向本机或私有网络，已停止导入。")
    return detection.normalized_url


def _safe_get_with_redirects(
    client: Any,
    source_url: str,
    *,
    platform: LinkMaterialPlatform,
) -> Any:
    """GET a platform page while validating HTTPS/host/DNS before every hop."""
    current_url = _validate_platform_url(source_url, expected_platform=platform).normalized_url
    for redirect_count in range(_MAX_SAFE_LINK_REDIRECTS + 1):
        current_url = _assert_public_link_target(current_url, expected_platform=platform)
        response = client.get(current_url, follow_redirects=False)
        if int(response.status_code) not in _SAFE_REDIRECT_STATUS_CODES:
            return response
        location = str(response.headers.get("location") or "").strip()
        close = getattr(response, "close", None)
        if callable(close):
            close()
        if not location:
            raise LinkMaterialImportError("链接跳转响应缺少目标地址，已停止导入。")
        if redirect_count >= _MAX_SAFE_LINK_REDIRECTS:
            raise LinkMaterialImportError("链接跳转次数过多，已停止导入。")
        next_url = urljoin(current_url, location)
        # Validate before the next loop so a malicious Location is never requested.
        current_url = _validate_platform_url(next_url, expected_platform=platform).normalized_url
    raise LinkMaterialImportError("链接跳转次数过多，已停止导入。")


def detect_link_material(value: str) -> LinkMaterialDetection:
    raw = str(value or "").strip()
    if not raw:
        raise LinkMaterialImportError("请先粘贴 B 站 / 小红书 / 公众号链接。")
    # 纯 BV 号(没有 URL)单独处理。
    if _BILIBILI_BV_RE.match(raw):
        return LinkMaterialDetection(
            platform="bilibili",
            normalized_url=f"https://www.bilibili.com/video/{raw}",
            display_name="B站",
        )
    # 把真正的 URL 从分享文案里抠出来；抠不到再退回原文(兼容用户直接粘干净链接)。
    url = _extract_url_from_text(raw) or raw
    return _validate_platform_url(url)


def adapter_for_platform(platform: LinkMaterialPlatform) -> LinkMaterialPlatformAdapter:
    if platform == "bilibili":
        return BilibiliLinkAdapter()
    if platform == "xiaohongshu":
        return XiaohongshuLinkAdapter()
    if platform == "wechat_article":
        return WechatArticleAdapter()
    raise LinkMaterialImportError("暂不支持这个链接平台。")


def extract_link_material_source(
    source_url: str,
    temp_dir: Path,
    *,
    options: LinkMaterialImportOptions | None = None,
) -> LinkMaterialSource:
    detection = detect_link_material(source_url)
    temp_dir.mkdir(parents=True, exist_ok=True)
    return adapter_for_platform(detection.platform).extract(detection.normalized_url, temp_dir, options=options)


class _YtDlpAdapter(LinkMaterialPlatformAdapter):
    def extract(self, source_url: str, temp_dir: Path, *, options: LinkMaterialImportOptions | None = None) -> LinkMaterialSource:
        options = options or LinkMaterialImportOptions()
        source_url = _assert_public_link_target(source_url, expected_platform=self.platform)
        executable = _find_yt_dlp()
        if not executable:
            raise LinkMaterialImportError(
                f"{self.display_name}链接已识别，但当前未安装可用的链接提取器 yt-dlp，无法下载本地媒体。"
            )
        output_template = str(temp_dir / "%(id)s.%(ext)s")
        base_metadata = {
            "accessMode": "browser_cookie" if options.use_browser_cookies else "anonymous",
            "cookieBrowser": options.cookie_browser if options.use_browser_cookies else None,
            "pipelineMode": "media_first",
            "mediaDownloadMode": None,
            "mediaDownloaded": False,
            "downloadedMediaKind": "unknown",
            "tempDirStatus": "created",
            "tempMediaKind": "none",
            "subtitleAvailable": False,
            "transcriptSource": "none",
            "audioExtracted": False,
            "ffmpegAvailable": bool(find_ffmpeg()),
            "keepMedia": False,
            "curlCffiAvailable": _has_curl_cffi(),
            "ytDlpVersion": _get_yt_dlp_version(executable),
            "impersonationTarget": _get_yt_dlp_impersonate_target(executable),
            "impersonationAvailable": bool(_get_yt_dlp_impersonate_target(executable)),
            "supportedCookieBrowsers": _SUPPORTED_COOKIE_BROWSERS,
        }
        _notify_stage(options, "下载临时媒体中", 18.0, base_metadata)
        media_path, media_metadata = download_temp_media(
            executable=executable,
            source_url=source_url,
            temp_dir=temp_dir,
            output_template=output_template,
            options=options,
            base_metadata=base_metadata,
            display_name=self.display_name,
            platform=self.platform,
        )
        metadata = _read_downloaded_info_json(temp_dir)
        title = _clean_title(str(metadata.get("title") or "")) or f"{self.display_name}视频资料"
        source_metadata = {
            **base_metadata,
            **media_metadata,
            "extractor": metadata.get("extractor_key") or metadata.get("extractor"),
            "uploader": metadata.get("uploader") or metadata.get("channel"),
            "duration": metadata.get("duration"),
            "webpageUrl": metadata.get("webpage_url") or source_url,
            "keepMedia": False,
        }
        _notify_stage(options, "检查字幕中", 38.0, source_metadata)
        transcript_text = read_valid_downloaded_subtitles(temp_dir)
        downloaded_paths = [item for item in temp_dir.iterdir() if item.is_file()]
        if transcript_text:
            source_metadata.update({"subtitleAvailable": True, "transcriptSource": "downloaded_subtitle"})
        else:
            source_metadata.update({"subtitleAvailable": False, "transcriptSource": "none"})
            media_kind = str(media_metadata.get("downloadedMediaKind") or _guess_media_kind(media_path))
            if media_kind == "audio" and media_path.suffix.lower() in _SOUNDFILE_SAFE_AUDIO_EXTENSIONS:
                # 已是 soundfile 能直接读的音频(wav/flac/ogg…)，无需转码。
                audio_path = media_path
                source_metadata.update({"tempMediaKind": "audio", "audioExtracted": False})
            elif media_kind == "audio":
                # 下载到的是 m4a/aac/mp3/opus 等 soundfile 读不了的音频 → 用 ffmpeg 转成 16k 单声道 wav。
                ffmpeg = find_ffmpeg()
                source_metadata["ffmpegAvailable"] = bool(ffmpeg)
                if not ffmpeg:
                    raise LinkMaterialImportError(
                        f"已下载本地音频，但格式（{media_path.suffix or '未知'}）需要 ffmpeg 转码后才能转写，当前未检测到 ffmpeg。",
                        metadata=source_metadata,
                    )
                _notify_stage(options, "转码音频中", 44.0, source_metadata)
                audio_path = extract_audio_from_media(media_path, temp_dir, ffmpeg=ffmpeg)
                source_metadata.update({"tempMediaKind": "audio", "audioExtracted": True})
            else:
                ffmpeg = find_ffmpeg()
                source_metadata["ffmpegAvailable"] = bool(ffmpeg)
                if not ffmpeg:
                    raise LinkMaterialImportError(
                        "已下载本地视频，但当前未检测到 ffmpeg，无法抽取音频用于转写。",
                        metadata=source_metadata,
                    )
                _notify_stage(options, "抽取音频中", 44.0, source_metadata)
                audio_path = extract_audio_from_media(media_path, temp_dir, ffmpeg=ffmpeg)
                source_metadata.update({"tempMediaKind": "audio", "audioExtracted": True})
            engine = find_local_transcript_engine()
            if engine is None:
                raise LinkMaterialImportError(
                    "已下载本地媒体，但当前未检测到本地 SenseVoice / Whisper 转写引擎。",
                    metadata=source_metadata,
                )
            _notify_stage(options, "本地转写中", 48.0, {**source_metadata, "transcriptSource": engine.name, "transcriptEngine": engine.name})
            transcript_text = _transcribe_temp_audio(engine, audio_path, temp_dir)
            source_metadata.update(
                {
                    "transcriptSource": engine.name,
                    "transcriptEngine": engine.name,
                    "tempMediaKind": "audio",
                }
            )
            downloaded_paths = [item for item in temp_dir.iterdir() if item.is_file()]
        return LinkMaterialSource(
            platform=self.platform,
            source_url=source_url,
            title=title,
            transcript_text=transcript_text,
            metadata=source_metadata,
            downloaded_paths=downloaded_paths,
        )


class BilibiliLinkAdapter(_YtDlpAdapter):
    platform: LinkMaterialPlatform = "bilibili"
    display_name = "B站"


class XiaohongshuLinkAdapter(_YtDlpAdapter):
    platform: LinkMaterialPlatform = "xiaohongshu"
    display_name = "小红书"


_WECHAT_ARTICLE_USER_AGENT = (
    "Mozilla/5.0 (iPhone; CPU iPhone OS 16_6 like Mac OS X) "
    "AppleWebKit/605.1.15 (KHTML, like Gecko) Mobile/15E148 MicroMessenger/8.0.40"
)


class WechatArticleAdapter(LinkMaterialPlatformAdapter):
    """微信公众号文章：HTTP 抓 HTML → BeautifulSoup 提取正文 → markdownify。

    路径完全不走 yt-dlp / SenseVoice，因为公众号文章是文字而非视频/音频。
    """

    platform: LinkMaterialPlatform = "wechat_article"
    display_name = "公众号"

    def extract(self, source_url: str, temp_dir: Path, *, options: LinkMaterialImportOptions | None = None) -> LinkMaterialSource:
        try:
            import httpx
            from bs4 import BeautifulSoup
            import markdownify
        except ImportError as exc:
            raise LinkMaterialImportError(
                f"微信公众号链接需要 beautifulsoup4 / markdownify 依赖：{exc}"
            ) from exc

        _notify_stage(options or LinkMaterialImportOptions(), "抓取公众号文章中", 30.0, None)
        try:
            with httpx.Client(
                headers={
                    "User-Agent": _WECHAT_ARTICLE_USER_AGENT,
                    "Referer": "https://mp.weixin.qq.com/",
                    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
                    "Accept-Language": "zh-CN,zh;q=0.9",
                },
                follow_redirects=False,
                trust_env=False,
                timeout=30.0,
            ) as client:
                resp = _safe_get_with_redirects(client, source_url, platform=self.platform)
                resp.raise_for_status()
        except LinkMaterialImportError:
            raise
        except Exception as exc:
            logger.warning("Wechat article fetch failed for a validated target", exc_info=True)
            raise LinkMaterialImportError(
                "无法获取公众号文章。可能是链接已失效、文章被删除或需要登录。"
            ) from exc

        html = resp.text
        soup = BeautifulSoup(html, "html.parser")

        title_tag = soup.find("h1", class_="rich_media_title") or soup.find("h2", class_="rich_media_title")
        title = title_tag.get_text(strip=True) if title_tag else ""
        if not title:
            og_title = soup.find("meta", property="og:title")
            if og_title:
                title = (og_title.get("content") or "").strip()
        if not title:
            title = "公众号文章"

        # 作者 / 公众号名
        author = ""
        author_tag = soup.find("a", id="js_name") or soup.find("strong", class_="profile_nickname")
        if author_tag:
            author = author_tag.get_text(strip=True)

        body_div = soup.find("div", id="js_content") or soup.find("div", class_="rich_media_content")
        if body_div is None:
            raise LinkMaterialImportError(
                "解析公众号文章失败：找不到正文节点。可能是模板变更或链接指向非文章页。"
            )

        # 懒加载图片：data-src → src
        for img in body_div.find_all("img"):
            data_src = img.get("data-src")
            if data_src:
                img["src"] = data_src

        # 删掉不需要的元素：广告 / 互动 / 二维码 / 脚注按钮
        for selector in [
            'mp-common-mpaudio',
            'mpvoice',
            'iframe',
            'script',
            'style',
            'mp-style-type',
            '.qr_code_pc_outer',
            '.reward_area',
            '.weui-poi-msg',
            'mp-common-cps-card',
        ]:
            for node in body_div.select(selector):
                node.decompose()

        markdown_body = markdownify.markdownify(
            str(body_div),
            heading_style="ATX",
            bullets="-",
            strip=['span', 'section'],
        ).strip()

        # markdownify 容易产生过多连续空行，挤压一下
        markdown_body = re.sub(r"\n{3,}", "\n\n", markdown_body)
        if not markdown_body or len(markdown_body) < 20:
            raise LinkMaterialImportError("公众号文章正文为空或过短，无法生成资料。")

        source_metadata = {
            "sourcePlatform": "wechat_article",
            "sourceKind": "article",  # 给下游 ingest 流水线信号：跳过 polish，文章本身已经排版好
            "subtitleAvailable": True,
            "transcriptSource": "wechat_article_html",
            "transcriptEngine": "wechat_article_parser",
            "videoTitle": title,
            "wechatAuthor": author or None,
            "keepMedia": False,
            "tempMediaKind": "article",
        }

        return LinkMaterialSource(
            platform=self.platform,
            source_url=source_url,
            title=title,
            transcript_text=markdown_body,
            metadata={k: v for k, v in source_metadata.items() if v is not None},
            downloaded_paths=[],
        )


def polish_transcript_for_reading(
    *,
    ai_service: Any | None,
    title: str,
    source_url: str,
    cleaned_text: str,
) -> str:
    """二次润色：分段 + 关键句加粗。

    输入 cleaned_text（已去掉口头禅）→ 输出 markdown（带段落分隔 + **加粗** 标记）。
    LLM 失败时 fallback：原文按句号/感叹号粗分段，无加粗。
    """
    text = (cleaned_text or "").strip()
    if not text:
        return ""

    fallback = _simple_paragraph_split(text)

    if ai_service is None or not hasattr(ai_service, "generate_raw_evidence_response"):
        return fallback

    system_instruction = (
        "你是中文语音转写文本的排版助手。我会给你一段连续的转写文字（缺少段落分隔，可能有同音字错别字）。\n"
        "\n"
        "你只能做两件事：\n"
        "  (1) 按语义把它分成段落（每 3-6 句话一段，段落之间用空行），让人能读；\n"
        "  (2) 修正明显的同音字/听写错别字（如「形业」→「行业」、「沉淀」→「沉甸」反向），\n"
        "      但只能在错别字上动手——其余字、词、句子顺序必须原样保留。\n"
        "\n"
        "禁止：\n"
        "  - 禁止总结、概括、压缩内容；\n"
        "  - 禁止改写语序、换说法、合并句子；\n"
        "  - 禁止删除任何观点、例子、数据、口语连接词、过渡语；\n"
        "  - 禁止新增任何标题、章节号、列表符号、加粗标记、引用；\n"
        "  - 输出字数必须与输入字数大致相等（允许 ±8% 的微小波动，只够你修错别字 / 调标点 / 增加段落空行用）。\n"
        "\n"
        "如果你觉得 \"这里应该归纳一下\" \"这段可以删掉\"，请忍住——用户要的是完整原文，只是排版好看一点而已。\n"
        "只输出整理后的正文，不要任何解释、前缀、后缀。"
    )
    prompt = (
        f"视频题目：{title}\n\n"
        "请把下面的转写文字按段落排版（不要总结、不要压缩，逐字保留）："
    )
    try:
        structured = ai_service.generate_raw_evidence_response(
            prompt,
            system_instruction,
            text[:60000],
            timeout_seconds=180.0,
            max_tokens=8000,
            enable_thinking=False,
        )
        polished = str(getattr(structured, "content", "") or "").strip()
        # 严格字数守门：LLM 输出字数必须 ≥ 原文 92%（最多删 8% 的同音字/标点波动）。
        # 一旦低于此阈值，就当 LLM 偷偷做了总结 → 直接 fallback 到机械分段（零损失）。
        plain_len = len(re.sub(r"[*_`#>\s]", "", polished))
        origin_len = len(re.sub(r"\s", "", text))
        if origin_len > 0 and plain_len >= origin_len * 0.92:
            return polished
        logger.warning(
            "[link-material] LLM polished output rejected: %d chars vs origin %d (%.1f%%)",
            plain_len, origin_len, (plain_len / origin_len * 100) if origin_len else 0.0,
        )
    except Exception:
        pass
    return fallback


# 无标点长文本的兜底分段长度：每段约这么多字（优先在逗号/顿号等软停顿处断，否则硬断）。
_FALLBACK_PARAGRAPH_CHARS = 140


def _simple_paragraph_split(text: str) -> str:
    """fallback 分段（不依赖 LLM，云端断网时也能用）：

    1. 先尊重已有的空行分段（如分窗转写已给出的段落结构），不压平；
    2. 每个块内若有句末标点（。！？!?）→ 每 4 句一段；
    3. 块内无句末标点（ASR 转写常见无标点）→ 按长度粗分，优先在逗号/顿号处断，
       避免整段糊成"一整片"。
    """
    text = (text or "").strip()
    if not text:
        return text
    blocks = [b.strip() for b in re.split(r"\n\s*\n", text) if b.strip()]
    out: list[str] = []
    for block in blocks:
        out.extend(_split_one_block(block))
    return "\n\n".join(out) if out else text


def _split_one_block(block: str) -> list[str]:
    """单个文本块分段：有句末标点按 4 句一段，无标点按长度粗分。"""
    sentences = [s.strip() for s in re.split(r"(?<=[。！？!?])", block) if s.strip()]
    if len(sentences) > 1:
        paragraphs: list[str] = []
        buffer: list[str] = []
        for sentence in sentences:
            buffer.append(sentence)
            if len(buffer) >= 4:
                paragraphs.append("".join(buffer))
                buffer = []
        if buffer:
            paragraphs.append("".join(buffer))
        return paragraphs
    single = sentences[0] if sentences else block
    if len(single) <= _FALLBACK_PARAGRAPH_CHARS:
        return [single]
    return _split_block_by_length(single)


def _split_block_by_length(text: str) -> list[str]:
    """无句末标点的长文本按长度切段：优先在逗号/顿号/分号等软停顿处断，否则硬切。"""
    parts = [p for p in re.split(r"(?<=[，、；,;])", text) if p]
    paragraphs: list[str] = []
    buffer = ""
    for part in parts:
        if buffer and len(buffer) + len(part) > _FALLBACK_PARAGRAPH_CHARS:
            paragraphs.append(buffer)
            buffer = part
        else:
            buffer += part
    if buffer:
        paragraphs.append(buffer)
    # 仍有超长段（连软停顿都没有）→ 按字数硬切，绝不留"一整片"。
    final: list[str] = []
    hard_limit = int(_FALLBACK_PARAGRAPH_CHARS * 1.5)
    for paragraph in paragraphs:
        if len(paragraph) <= hard_limit:
            final.append(paragraph)
        else:
            for index in range(0, len(paragraph), _FALLBACK_PARAGRAPH_CHARS):
                final.append(paragraph[index:index + _FALLBACK_PARAGRAPH_CHARS])
    return final


_DOCX_BODY_FONT = "黑体"
_DOCX_BODY_FONT_SIZE_PT = 11
_DOCX_HEADING_FONT_SIZE_PT = {0: 22, 1: 18, 2: 14, 3: 12, 4: 11}


def _apply_unified_docx_styles(doc: Any) -> None:
    """统一锁定中文字体到 Normal + Heading 1-4 + Title style。

    python-docx 给中文字段设字体必须同时设 rFonts 的 eastAsia 属性（OOXML 怪招），
    否则 Word 仍然回退到默认。
    """
    from docx.oxml.ns import qn
    from docx.shared import Pt

    def _set_style_font(style_name: str, size_pt: int) -> None:
        if style_name not in doc.styles:
            return
        style = doc.styles[style_name]
        style.font.name = _DOCX_BODY_FONT
        style.font.size = Pt(size_pt)
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.find(qn('w:rFonts'))
        if rfonts is None:
            from docx.oxml import OxmlElement
            rfonts = OxmlElement('w:rFonts')
            rpr.append(rfonts)
        rfonts.set(qn('w:eastAsia'), _DOCX_BODY_FONT)
        rfonts.set(qn('w:ascii'), _DOCX_BODY_FONT)
        rfonts.set(qn('w:hAnsi'), _DOCX_BODY_FONT)
        rfonts.set(qn('w:cs'), _DOCX_BODY_FONT)

    _set_style_font('Normal', _DOCX_BODY_FONT_SIZE_PT)
    _set_style_font('Title', _DOCX_HEADING_FONT_SIZE_PT[0])
    for level in range(1, 5):
        _set_style_font(f'Heading {level}', _DOCX_HEADING_FONT_SIZE_PT.get(level, 12))


def _force_run_font(run: Any) -> None:
    """单 run 上再保险一次设字体，避免某些 style 继承没生效的边缘情况。"""
    from docx.oxml.ns import qn
    from docx.shared import Pt

    run.font.name = _DOCX_BODY_FONT
    run.font.size = Pt(_DOCX_BODY_FONT_SIZE_PT)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        from docx.oxml import OxmlElement
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), _DOCX_BODY_FONT)
    rfonts.set(qn('w:ascii'), _DOCX_BODY_FONT)
    rfonts.set(qn('w:hAnsi'), _DOCX_BODY_FONT)


_MD_IMAGE_PATTERN = re.compile(r"!\[[^\]]*\]\([^)]*\)")
_MD_LINK_PATTERN = re.compile(r"\[([^\]]+?)\]\(([^)]+?)\)")
_MD_BOLD_PATTERN = re.compile(r"\*\*(.+?)\*\*")
_MD_INLINE_CODE_PATTERN = re.compile(r"`([^`]+?)`")
_HTML_IMG_PATTERN = re.compile(r"<img[^>]*>", re.IGNORECASE)
_HTML_TAG_PATTERN = re.compile(r"<[^>]+>")
_BLANK_LINE_RE = re.compile(r"\n{3,}")
_LIST_ITEM_RE = re.compile(r"^(?:[-*+•]|\d+[.、])\s+(.+)$")


def _clean_markdown_artifacts_for_docx(text: str) -> str:
    """Strip markdown artifacts that don't render to clean docx text.

    - 图片占位符 `![alt](url)` / HTML `<img>`：去掉，避免在 docx 里看到一串 url
    - 多余连续空行压成一个段落分隔
    """
    cleaned = _MD_IMAGE_PATTERN.sub("", text)
    cleaned = _HTML_IMG_PATTERN.sub("", cleaned)
    # 内嵌的 #imgIndex= / wxfrom= 等参数残留（图片被去掉后图片标记里的尾部杂字）
    cleaned = re.sub(r"#imgIndex=\d+", "", cleaned)
    cleaned = _BLANK_LINE_RE.sub("\n\n", cleaned)
    return cleaned


def _add_docx_hyperlink(paragraph, url: str, text: str) -> None:
    """给段落加一个真正的可点击 hyperlink（蓝色下划线），不在文本里露 url。"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    if not text:
        text = url
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    # 颜色蓝色 + 下划线 + 黑体
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2563EB")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    r_fonts = OxmlElement("w:rFonts")
    for attr in ("w:ascii", "w:eastAsia", "w:hAnsi", "w:cs"):
        r_fonts.set(qn(attr), "黑体")
    r_pr.append(r_fonts)
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _render_inline_markdown_to_runs(paragraph, line: str) -> None:
    """渲染一行的 inline markdown：**bold**、`code`、[text](url) hyperlink。

    顺序：先扫所有 [text](url) 链接位置 → 中间按 bold/inline code 处理 → 链接段插 hyperlink。
    """
    cursor = 0
    for link_match in _MD_LINK_PATTERN.finditer(line):
        # 处理 link 之前的普通文本
        if link_match.start() > cursor:
            chunk = line[cursor:link_match.start()]
            _render_chunk_with_bold_code(paragraph, chunk)
        link_text = link_match.group(1).strip()
        link_url = link_match.group(2).strip()
        # 链接 url 必须看起来像 url，否则当普通文本处理
        if not link_url or not re.match(r"^(https?://|mailto:|/)", link_url):
            run = paragraph.add_run(link_text)
            _force_run_font(run)
        else:
            _add_docx_hyperlink(paragraph, link_url, link_text)
        cursor = link_match.end()
    if cursor < len(line):
        _render_chunk_with_bold_code(paragraph, line[cursor:])


def _render_chunk_with_bold_code(paragraph, chunk: str) -> None:
    """渲染一段文本里的 **bold** / `code`（不含链接，链接已在上层切片）。"""
    if not chunk:
        return
    # 先处理 bold；inline code 极简：当成 monospace 加底色（这里简化为普通字体）
    cursor = 0
    for match in _MD_BOLD_PATTERN.finditer(chunk):
        if match.start() > cursor:
            run = paragraph.add_run(chunk[cursor:match.start()])
            _force_run_font(run)
        bold_run = paragraph.add_run(match.group(1))
        bold_run.bold = True
        _force_run_font(bold_run)
        cursor = match.end()
    if cursor < len(chunk):
        run = paragraph.add_run(chunk[cursor:])
        _force_run_font(run)


def render_polished_markdown_to_docx(
    *,
    title: str,
    source_url: str,
    markdown_body: str,
    output_path: Path,
) -> Path:
    """把 polished markdown 渲染成 docx：
    - 标题 / 段落 / 列表 / 加粗 / 真 hyperlink
    - 图片占位符（![](url) 和 <img>）被剥掉，不在 docx 里露 url 乱码
    - 字体统一锁定黑体
    """
    from docx import Document

    doc = Document()
    _apply_unified_docx_styles(doc)

    if title:
        heading = doc.add_heading(title, level=0)
        for run in heading.runs:
            _force_run_font(run)
    if source_url:
        p = doc.add_paragraph()
        run = p.add_run("原链接：")
        run.italic = True
        _force_run_font(run)
        _add_docx_hyperlink(p, source_url, source_url)
    doc.add_paragraph("")

    sanitized = _clean_markdown_artifacts_for_docx(markdown_body.replace("\r\n", "\n"))
    # 按段落分割：一个或多个空行 = 段落分界
    blocks = re.split(r"\n\s*\n", sanitized)
    # 顾源源 5/26: 真 markdown 表格识别 (block 全是 |...| 行 + 含 separator → 画 docx 表格)
    from app.services.markdown_table import (
        _TABLE_ROW_RE as _MD_TABLE_ROW_RE,
        _TABLE_SEP_RE as _MD_TABLE_SEP_RE,
        _split_table_cells,
    )

    for raw_block in blocks:
        block = raw_block.strip()
        if not block:
            continue
        # 一个 block 内的多行可能是：连续段落文本 / 列表 / 一个标题 / 真表格
        block_lines = [line for line in (ln.rstrip() for ln in block.split("\n")) if line.strip()]
        if not block_lines:
            continue
        # 真表格识别 (第 1 行 row + 第 2 行 separator + 后续 row)
        is_table_block = (
            len(block_lines) >= 2
            and _MD_TABLE_ROW_RE.match(block_lines[0])
            and _MD_TABLE_SEP_RE.match(block_lines[1])
        )
        if is_table_block:
            headers = _split_table_cells(block_lines[0])
            rows: list[list[str]] = []
            for ln in block_lines[2:]:
                if not _MD_TABLE_ROW_RE.match(ln):
                    break
                cells = _split_table_cells(ln)
                if len(cells) < len(headers):
                    cells = cells + [""] * (len(headers) - len(cells))
                elif len(cells) > len(headers):
                    cells = cells[: len(headers)]
                rows.append(cells)
            if headers:
                n_cols = len(headers)
                table = doc.add_table(rows=1 + len(rows), cols=n_cols)
                try:
                    table.style = "Light Grid Accent 1"
                except Exception:
                    pass
                # 真写 header (加粗)
                for col_idx, header in enumerate(headers):
                    cell = table.rows[0].cells[col_idx]
                    cell.text = ""
                    para = cell.paragraphs[0]
                    _render_inline_markdown_to_runs(para, header)
                    for run in para.runs:
                        run.bold = True
                        _force_run_font(run)
                # 真写 body
                for r_idx, row in enumerate(rows):
                    for c_idx, cell_value in enumerate(row):
                        if c_idx >= n_cols:
                            break
                        cell = table.rows[r_idx + 1].cells[c_idx]
                        cell.text = ""
                        para = cell.paragraphs[0]
                        _render_inline_markdown_to_runs(para, cell_value)
                        for run in para.runs:
                            _force_run_font(run)
                continue
        # 判断整个 block 是不是列表（每行都以 - / * / 数字. 开头）
        is_list_block = all(_LIST_ITEM_RE.match(ln.strip()) for ln in block_lines)
        if is_list_block:
            for ln in block_lines:
                item_match = _LIST_ITEM_RE.match(ln.strip())
                if not item_match:
                    continue
                item_text = item_match.group(1).strip()
                # 用 List Bullet 样式（python-docx 内置）
                try:
                    paragraph = doc.add_paragraph(style="List Bullet")
                except Exception:
                    paragraph = doc.add_paragraph()
                    paragraph.add_run("• ")
                _render_inline_markdown_to_runs(paragraph, item_text)
                for run in paragraph.runs:
                    _force_run_font(run)
            continue
        # 否则按行处理：标题 / 普通段落
        for ln in block_lines:
            line = ln.strip()
            if not line:
                continue
            heading_match = re.match(r"^(#{1,3})\s+(.+)$", line)
            if heading_match:
                level = len(heading_match.group(1))
                heading = doc.add_heading(heading_match.group(2).strip(), level=min(level, 4))
                for run in heading.runs:
                    _force_run_font(run)
                continue
            paragraph = doc.add_paragraph()
            _render_inline_markdown_to_runs(paragraph, line)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def cleanup_transcript_text(*, ai_service: Any | None, title: str, source_url: str, transcript_text: str) -> str:
    cleaned = _strip_low_value_transcript_text(transcript_text)
    if len(cleaned) < 20:
        raise LinkMaterialImportError("转写文本过短，无法生成有效资料。")
    # Long transcripts must stay complete. LLM cleanup is useful for short,
    # messy snippets, but for full videos it can silently summarize or omit
    # sections. Keep long video transcripts deterministic and lossless.
    if len(cleaned) >= 3000:
        return cleaned
    if ai_service is not None and hasattr(ai_service, "generate_raw_evidence_response"):
        system_instruction = (
            "你是资料整理助手。请把视频字幕或转写稿整理成干净、可入库的中文资料正文。"
            "只输出正文，不要写标题、原链接、处理过程、时间戳、下载日志或自我介绍。"
            "必须保留原始转写中的全部有效信息、观点、结构和事实；只去掉口头禅、重复语气词和无意义寒暄。"
            "不要摘要，不要压缩，不要只保留重点，不要改写成短文。"
            "不要编造字幕里没有的内容。"
        )
        prompt = (
            f"视频题目：{title}\n"
            f"原链接：{source_url}\n\n"
            "请整理为一份可被知识库读取的完整正文资料。"
        )
        try:
            structured = ai_service.generate_raw_evidence_response(
                prompt,
                system_instruction,
                cleaned[:60000],
                timeout_seconds=180.0,
                max_tokens=8000,
                enable_thinking=False,
            )
            ai_body = _strip_low_value_transcript_text(str(getattr(structured, "content", "") or ""))
            min_preserved_chars = max(40, int(len(cleaned) * 0.7))
            if len(ai_body) >= min_preserved_chars:
                return ai_body
        except Exception:
            # Fallback to deterministic cleaning. Import should not fail solely
            # because the cleanup LLM is temporarily unavailable.
            pass
    return cleaned


def build_clean_video_markdown(*, title: str, source_url: str, body: str) -> str:
    clean_title = _clean_title(title) or "视频转文字资料"
    clean_body = _strip_low_value_transcript_text(body)
    return f"# {clean_title}\n\n原链接：{source_url}\n\n{clean_body}\n"


def cleanup_temp_dir(temp_dir: Path) -> Literal["cleaned", "failed"]:
    try:
        if temp_dir.exists():
            shutil.rmtree(temp_dir)
        return "cleaned"
    except Exception:
        return "failed"


def _notify_stage(options: LinkMaterialImportOptions, stage: str, progress: float, metadata: dict[str, object] | None = None) -> None:
    if options.stage_callback is None:
        return
    try:
        options.stage_callback(stage, progress, metadata)
    except Exception:
        # Progress reporting must never make media extraction fail.
        return


def find_local_transcript_engine() -> LocalTranscriptEngine | None:
    sensevoice_template = os.environ.get("YIYU_SENSEVOICE_CMD", "").strip()
    if sensevoice_template:
        return LocalTranscriptEngine(
            name="local_sensevoice",
            command=shlex.split(sensevoice_template),
            command_template=sensevoice_template,
        )
    # 优先用应用内置的 sherpa-onnx + ONNX 模型（用户在系统设置 → 语音识别模型 里下载的那个）
    # 这是和录音转写共用的引擎，避免架构断层导致用户"装了模型但链接转写说没装"。
    try:
        from app.services.local_asr.model_paths import is_model_ready as _is_local_asr_model_ready
        if _is_local_asr_model_ready():
            return LocalTranscriptEngine(name="builtin_sensevoice", command=[])
    except Exception:
        # 本地 ASR 模块加载失败不应阻塞 fallback 到命令行引擎
        pass
    sensevoice = shutil.which("sensevoice")
    if sensevoice and _command_available([sensevoice, "--help"]):
        return LocalTranscriptEngine(name="local_sensevoice", command=[sensevoice])
    for whisper in _find_all_executables("whisper"):
        if not _command_available([whisper, "--help"]):
            continue
        return LocalTranscriptEngine(name="local_whisper", command=[whisper])
    python = shutil.which("python3") or shutil.which("python")
    if python:
        if _command_available([python, "-m", "whisper", "--help"]):
            return LocalTranscriptEngine(name="local_whisper", command=[python, "-m", "whisper"])
    return None


def _find_all_executables(name: str) -> list[str]:
    seen: set[str] = set()
    matches: list[str] = []
    for raw_dir in os.get_exec_path():
        candidate = Path(raw_dir) / name
        if not candidate.exists() or not os.access(candidate, os.X_OK):
            continue
        resolved = str(candidate)
        if resolved in seen:
            continue
        seen.add(resolved)
        matches.append(resolved)
    return matches


def _command_available(command: list[str]) -> bool:
    try:
        completed = subprocess.run(
            command,
            capture_output=True,
            text=True,
            timeout=10,
            check=False,
        )
    except Exception:
        return False
    return completed.returncode == 0


def download_temp_media(
    *,
    executable: list[str],
    source_url: str,
    temp_dir: Path,
    output_template: str,
    options: LinkMaterialImportOptions,
    base_metadata: dict[str, Any],
    display_name: str,
    platform: LinkMaterialPlatform,
) -> tuple[Path, dict[str, Any]]:
    attempts: list[dict[str, Any]] = []
    profiles = _build_download_attempt_profiles(platform=platform, options=options)
    last_detail = ""
    last_failure_kind = "unknown"
    last_exit_code: int | None = None
    last_profile = profiles[-1].name if profiles else "default"
    for profile in profiles:
        _remove_media_candidates(temp_dir)
        video_result = _attempt_yt_dlp_media_download(
            executable=executable,
            source_url=source_url,
            temp_dir=temp_dir,
            output_template=output_template,
            options=options,
            profile=profile,
            mode="video_first",
            platform=platform,
        )
        attempts.append(video_result["attempt"])
        if video_result["media_path"] is not None:
            return video_result["media_path"], _successful_download_metadata(
                base_metadata=base_metadata,
                profile=profile,
                mode="video_first",
                media_path=video_result["media_path"],
                attempts=attempts,
            )
        last_detail = str(video_result["detail"] or "")
        last_failure_kind = str(video_result["failure_kind"] or "unknown")
        last_exit_code = int(video_result["returncode"])
        last_profile = profile.name
        if last_failure_kind in {"http_412", "login_required", "cookie_required"}:
            # Access failures happen before media format negotiation, so retry
            # with stronger request profiles instead of wasting an audio pass.
            continue

        _remove_media_candidates(temp_dir)
        audio_result = _attempt_yt_dlp_media_download(
            executable=executable,
            source_url=source_url,
            temp_dir=temp_dir,
            output_template=output_template,
            options=options,
            profile=profile,
            mode="audio_fallback",
            platform=platform,
        )
        attempts.append(audio_result["attempt"])
        if audio_result["media_path"] is not None:
            return audio_result["media_path"], _successful_download_metadata(
                base_metadata=base_metadata,
                profile=profile,
                mode="audio_fallback",
                media_path=audio_result["media_path"],
                attempts=attempts,
                downloaded_media_kind="audio",
            )
        last_detail = str(audio_result["detail"] or last_detail or "")
        last_failure_kind = str(audio_result["failure_kind"] or last_failure_kind or "unknown")
        last_exit_code = int(audio_result["returncode"])
        last_profile = profile.name

    if platform == "bilibili" and last_failure_kind in {"http_412", "login_required", "cookie_required"}:
        bbdown_result = _try_bbdown_download(
            source_url=source_url,
            temp_dir=temp_dir,
            base_metadata=base_metadata,
            attempts=attempts,
        )
        attempts.append(bbdown_result["attempt"])
        if bbdown_result["media_path"] is not None:
            media_path = bbdown_result["media_path"]
            media_kind = _guess_media_kind(media_path)
            return media_path, {
                **base_metadata,
                "mediaDownloadMode": "video_first",
                "mediaDownloaded": True,
                "downloadedMediaKind": media_kind,
                "tempMediaKind": media_kind,
                "accessFailureKind": None,
                "externalDownloader": "BBDown",
                "downloadAttemptProfile": "bbdown",
                "downloadAttemptCount": len(attempts),
                "downloadAttempts": attempts[-6:],
                "headersApplied": any(bool(item.get("headersApplied")) for item in attempts),
                "impersonationApplied": any(bool(item.get("impersonationApplied")) for item in attempts),
            }
        last_detail = str(bbdown_result["detail"] or last_detail or "")
        last_exit_code = int(bbdown_result["returncode"])
        last_profile = "bbdown"

    metadata = {
        **base_metadata,
        "mediaDownloadMode": "video_first",
        "mediaDownloaded": False,
        "accessFailureKind": last_failure_kind or "unknown",
        "downloadAttemptProfile": last_profile,
        "downloadAttemptCount": len(attempts),
        "downloadAttempts": attempts[-8:],
        "headersApplied": any(bool(item.get("headersApplied")) for item in attempts),
        "impersonationApplied": any(bool(item.get("impersonationApplied")) for item in attempts),
        "cookieMode": "browser_cookie" if any(item.get("cookieMode") == "browser_cookie" for item in attempts) else "none",
        "ytDlpExitCode": last_exit_code,
        "ytDlpErrorTail": _error_tail(last_detail),
        "externalDownloader": _summarize_external_downloader(attempts),
    }
    if platform == "bilibili" and last_failure_kind in {"http_412", "login_required", "cookie_required"}:
        if not bool(base_metadata.get("impersonationAvailable")):
            reason = "B站返回 HTTP 412。已尝试 Origin/Referer/User-Agent，但当前 yt-dlp 浏览器模拟不可用；可安装/修复 curl-cffi、启用浏览器登录态或安装 BBDown 后重试。"
        elif options.use_browser_cookies:
            reason = "B站仍拒绝访问。已尝试请求头、浏览器模拟和浏览器登录态；可检查浏览器是否已登录，或安装 BBDown 后重试。"
        else:
            reason = "B站返回 HTTP 412。已尝试请求头与浏览器模拟；可启用浏览器登录态，或安装 BBDown 后重试。"
        raise LinkMaterialImportError(reason, metadata=metadata)
    if platform == "xiaohongshu":
        if last_failure_kind == "no_video":
            reason = "这条小红书是图文笔记（没有视频），没有可转写的内容。请改用视频笔记的分享链接。"
        elif last_failure_kind in {"login_required", "cookie_required"} and not options.use_browser_cookies:
            reason = "小红书需要登录态才能访问该内容。请在弹窗里勾选「使用浏览器登录态读取链接」后重试。"
        else:
            reason = (
                "小红书视频下载失败。常见原因：分享链接已过期——请从小红书 App 用「复制链接」重新获取最新链接"
                "（旧链接里的访问令牌会失效）；若仍不行，可勾选浏览器登录态再试。"
            )
        raise LinkMaterialImportError(reason, metadata=metadata)
    raise LinkMaterialImportError(
        f"临时媒体下载失败：{_error_tail(last_detail, limit=240) or '未知错误'}",
        metadata=metadata,
    )


def _build_download_attempt_profiles(
    *,
    platform: LinkMaterialPlatform,
    options: LinkMaterialImportOptions,
) -> list[_DownloadAttemptProfile]:
    if platform == "xiaohongshu":
        # 小红书：base → 补 headers → impersonate(浏览器指纹，缓解反爬) → cookie(登录墙)。
        # 与 B 站同构的升级重试；加了只提高成功率，失败仍给清晰提示。
        xhs_profiles = [
            _DownloadAttemptProfile(name="base"),
            _DownloadAttemptProfile(name="xhs_headers", headers_applied=True),
            _DownloadAttemptProfile(name="xhs_impersonate", headers_applied=True, impersonation_requested=True),
        ]
        if options.use_browser_cookies:
            xhs_profiles.append(
                _DownloadAttemptProfile(
                    name="xhs_cookie",
                    headers_applied=True,
                    impersonation_requested=True,
                    use_browser_cookies=True,
                )
            )
        return xhs_profiles
    if platform != "bilibili":
        return [
            _DownloadAttemptProfile(
                name="default_cookie" if options.use_browser_cookies else "default",
                use_browser_cookies=options.use_browser_cookies,
            )
        ]
    profiles = [
        _DownloadAttemptProfile(name="base"),
        _DownloadAttemptProfile(name="bili_headers", headers_applied=True),
        _DownloadAttemptProfile(name="bili_impersonate", headers_applied=True, impersonation_requested=True),
    ]
    if options.use_browser_cookies:
        profiles.append(
            _DownloadAttemptProfile(
                name="bili_cookie",
                headers_applied=True,
                impersonation_requested=True,
                use_browser_cookies=True,
            )
        )
    return profiles


def _attempt_yt_dlp_media_download(
    *,
    executable: list[str],
    source_url: str,
    temp_dir: Path,
    output_template: str,
    options: LinkMaterialImportOptions,
    profile: _DownloadAttemptProfile,
    mode: Literal["video_first", "audio_fallback"],
    platform: LinkMaterialPlatform = "bilibili",
) -> dict[str, Any]:
    header_args = _XHS_HEADER_ARGS if platform == "xiaohongshu" else _BILIBILI_HEADER_ARGS
    impersonation_target = _get_yt_dlp_impersonate_target(executable) if profile.impersonation_requested else None
    format_args = (
        [
            "-f",
            "bv*+ba/best",
            "--merge-output-format",
            "mp4",
        ]
        if mode == "video_first"
        else ["-f", "ba/bestaudio/best"]
    )
    command = _build_yt_dlp_command(
        executable,
        options,
        *format_args,
        "--write-subs",
        "--write-auto-subs",
        "--sub-langs",
        "zh.*,en",
        "--sub-format",
        "vtt",
        "--write-info-json",
        "--no-playlist",
        "-o",
        output_template,
        source_url,
        headers_applied=profile.headers_applied,
        header_args=header_args,
        impersonation_requested=profile.impersonation_requested,
        impersonation_target=impersonation_target,
        use_browser_cookies=profile.use_browser_cookies,
    )
    completed = _run_yt_dlp_download(command, temp_dir=temp_dir)
    detail = (completed.stderr or completed.stdout or "").strip()
    failure_kind = _classify_yt_dlp_access_failure(detail) if completed.returncode != 0 else None
    media_path = find_downloaded_media_file(temp_dir) if completed.returncode == 0 else None
    if completed.returncode == 0 and media_path is None:
        detail = "yt-dlp returned success but no media file was found"
        failure_kind = "unknown"
    return {
        "media_path": media_path,
        "returncode": completed.returncode,
        "detail": detail,
        "failure_kind": failure_kind,
        "attempt": {
            "profile": profile.name,
            "mode": mode,
            "returncode": completed.returncode,
            "failureKind": failure_kind,
            "errorTail": _error_tail(detail),
            "headersApplied": profile.headers_applied,
            "impersonationApplied": bool(impersonation_target),
            "impersonationTarget": impersonation_target,
            "impersonationRequested": profile.impersonation_requested,
            "cookieMode": "browser_cookie" if profile.use_browser_cookies else "none",
            "cookieBrowser": options.cookie_browser if profile.use_browser_cookies else None,
        },
    }


def _successful_download_metadata(
    *,
    base_metadata: dict[str, Any],
    profile: _DownloadAttemptProfile,
    mode: Literal["video_first", "audio_fallback"],
    media_path: Path,
    attempts: list[dict[str, Any]],
    downloaded_media_kind: str | None = None,
) -> dict[str, Any]:
    media_kind = downloaded_media_kind or _guess_media_kind(media_path)
    return {
        **base_metadata,
        "mediaDownloadMode": mode,
        "mediaDownloaded": True,
        "downloadedMediaKind": media_kind,
        "tempMediaKind": media_kind,
        "accessFailureKind": None,
        "downloadAttemptProfile": profile.name,
        "downloadAttemptCount": len(attempts),
        "downloadAttempts": attempts[-6:],
        "headersApplied": any(bool(item.get("headersApplied")) for item in attempts),
        "impersonationApplied": any(bool(item.get("impersonationApplied")) for item in attempts),
        "cookieMode": "browser_cookie" if any(item.get("cookieMode") == "browser_cookie" for item in attempts) else "none",
        "ytDlpExitCode": 0,
        "ytDlpErrorTail": "",
        "externalDownloader": None,
    }


def _try_bbdown_download(
    *,
    source_url: str,
    temp_dir: Path,
    base_metadata: dict[str, Any],
    attempts: list[dict[str, Any]],
) -> dict[str, Any]:
    bbdown = _find_bbdown()
    if not bbdown:
        return {
            "media_path": None,
            "returncode": 127,
            "detail": "BBDown is not installed",
            "attempt": {
                "profile": "bbdown",
                "mode": "external_downloader",
                "returncode": 127,
                "failureKind": "not_installed",
                "errorTail": "BBDown is not installed",
                "headersApplied": any(bool(item.get("headersApplied")) for item in attempts),
                "impersonationApplied": any(bool(item.get("impersonationApplied")) for item in attempts),
                "cookieMode": "none",
                "externalDownloader": "BBDown:not_installed",
            },
        }
    _remove_media_candidates(temp_dir)
    command = [bbdown, source_url, "--work-dir", str(temp_dir)]
    completed = _run_external_downloader(command, temp_dir=temp_dir)
    detail = (completed.stderr or completed.stdout or "").strip()
    media_path = find_downloaded_media_file(temp_dir) if completed.returncode == 0 else None
    return {
        "media_path": media_path,
        "returncode": completed.returncode,
        "detail": detail,
        "attempt": {
            "profile": "bbdown",
            "mode": "external_downloader",
            "returncode": completed.returncode,
            "failureKind": None if media_path else "external_downloader_failed",
            "errorTail": _error_tail(detail),
            "headersApplied": any(bool(item.get("headersApplied")) for item in attempts),
            "impersonationApplied": any(bool(item.get("impersonationApplied")) for item in attempts),
            "cookieMode": "none",
            "externalDownloader": "BBDown",
        },
    }


def _run_external_downloader(command: list[str], *, temp_dir: Path) -> subprocess.CompletedProcess[str]:
    try:
        return subprocess.run(
            command,
            cwd=str(temp_dir),
            capture_output=True,
            text=True,
            timeout=1800,
            check=False,
        )
    except subprocess.TimeoutExpired as exc:
        return subprocess.CompletedProcess(command, 1, stdout="", stderr=f"external downloader timeout: {exc}")


def _summarize_external_downloader(attempts: list[dict[str, Any]]) -> str | None:
    for attempt in reversed(attempts):
        value = attempt.get("externalDownloader")
        if value:
            return str(value)
    return None


def build_link_material_runtime_diagnostics() -> dict[str, Any]:
    executable = _find_yt_dlp()
    impersonation_target = _get_yt_dlp_impersonate_target(executable) if executable else None
    return {
        "ytDlpVersion": _get_yt_dlp_version(executable) if executable else None,
        "curlCffiAvailable": _has_curl_cffi(),
        "impersonationAvailable": bool(impersonation_target),
        "impersonationTarget": impersonation_target,
        "ffmpegAvailable": bool(find_ffmpeg()),
        "supportedCookieBrowsers": _SUPPORTED_COOKIE_BROWSERS,
        "bbdownAvailable": bool(_find_bbdown()),
    }


def _run_yt_dlp_download(command: list[str], *, temp_dir: Path) -> subprocess.CompletedProcess[str]:
    try:
        return subprocess.run(
            command,
            cwd=str(temp_dir),
            capture_output=True,
            text=True,
            timeout=1200,
            check=False,
        )
    except subprocess.TimeoutExpired as exc:
        return subprocess.CompletedProcess(command, 1, stdout="", stderr=f"download timeout: {exc}")


def find_downloaded_media_file(temp_dir: Path) -> Path | None:
    candidates = sorted(
        (item for item in temp_dir.iterdir() if _is_media_candidate(item)),
        key=lambda item: item.stat().st_size,
        reverse=True,
    )
    return candidates[0] if candidates else None


def _remove_media_candidates(temp_dir: Path) -> None:
    for item in list(temp_dir.iterdir()):
        if not _is_media_candidate(item):
            continue
        try:
            item.unlink()
        except Exception:
            pass


def _is_media_candidate(path: Path) -> bool:
    if not path.is_file():
        return False
    if path.stat().st_size <= 0:
        return False
    name = path.name.lower()
    if name.endswith(".info.json"):
        return False
    suffix = path.suffix.lower()
    return suffix not in _NON_MEDIA_EXTENSIONS


def _guess_media_kind(path: Path) -> Literal["audio", "video", "unknown"]:
    suffix = path.suffix.lower()
    if suffix in _AUDIO_EXTENSIONS:
        return "audio"
    if suffix in _VIDEO_EXTENSIONS:
        return "video"
    return "unknown"


def find_ffmpeg() -> str | None:
    return shutil.which("ffmpeg")


def extract_audio_from_media(media_path: Path, temp_dir: Path, *, ffmpeg: str) -> Path:
    output_path = temp_dir / "audio.wav"
    command = [
        ffmpeg,
        "-y",
        "-i",
        str(media_path),
        "-vn",
        "-ac",
        "1",
        "-ar",
        "16000",
        str(output_path),
    ]
    completed = subprocess.run(
        command,
        cwd=str(temp_dir),
        capture_output=True,
        text=True,
        timeout=900,
        check=False,
    )
    if completed.returncode != 0 or not output_path.exists() or output_path.stat().st_size <= 0:
        detail = (completed.stderr or completed.stdout or "").strip()
        raise LinkMaterialImportError(
            f"本地音频抽取失败：{detail[:240] or 'ffmpeg 没有生成可用音频'}",
            metadata={"audioExtracted": False, "ffmpegAvailable": True, "tempMediaKind": "video"},
        )
    return output_path


def _transcribe_temp_audio(engine: LocalTranscriptEngine, audio_path: Path, temp_dir: Path) -> str:
    # 内置 sherpa-onnx 引擎：直接 in-process 调用，跳过 subprocess。
    if engine.name == "builtin_sensevoice":
        try:
            from app.services.local_asr.sense_voice_provider import transcribe_local_audio
            result = transcribe_local_audio(str(audio_path))
        except Exception as exc:
            raise LinkMaterialImportError(
                f"本地转写失败：{exc}",
                metadata={"transcriptSource": engine.name, "transcriptEngine": engine.name, "tempMediaKind": "audio"},
            ) from exc
        text = (result.text or "").strip()
        if not text:
            raise LinkMaterialImportError(
                "本地转写返回了空文本，请检查音频质量或重试。",
                metadata={"transcriptSource": engine.name, "transcriptEngine": engine.name, "tempMediaKind": "audio"},
            )
        return text
    output_path = temp_dir / "transcript.txt"
    if engine.command_template:
        command = [
            part.format(input=str(audio_path), output=str(output_path))
            for part in shlex.split(engine.command_template)
        ]
    elif engine.name == "local_sensevoice":
        command = [*engine.command, "--input", str(audio_path), "--output", str(output_path)]
    else:
        command = [
            *engine.command,
            str(audio_path),
            "--language",
            "Chinese",
            "--output_format",
            "txt",
            "--output_dir",
            str(temp_dir),
        ]
    try:
        completed = subprocess.run(
            command,
            cwd=str(temp_dir),
            capture_output=True,
            text=True,
            timeout=1800,
            check=False,
            env=_transcription_subprocess_env(),
        )
    except FileNotFoundError as exc:
        raise LinkMaterialImportError(
            f"本地转写引擎不可用：{exc}",
            metadata={"transcriptSource": engine.name, "transcriptEngine": engine.name, "tempMediaKind": "audio"},
        ) from exc
    if completed.returncode != 0:
        detail = (completed.stderr or completed.stdout or "").strip()
        raise LinkMaterialImportError(
            f"本地转写失败：{detail[:240] or '未知错误'}",
            metadata={"transcriptSource": engine.name, "transcriptEngine": engine.name, "tempMediaKind": "audio"},
        )
    transcript_candidates = [output_path]
    transcript_candidates.extend(sorted(temp_dir.glob("*.txt"), key=lambda item: item.stat().st_mtime, reverse=True))
    for candidate in transcript_candidates:
        if not candidate.exists():
            continue
        try:
            text = candidate.read_text(encoding="utf-8", errors="ignore")
        except Exception:
            continue
        cleaned = _strip_low_value_transcript_text(text)
        if len(cleaned) >= 20:
            return cleaned
    raise LinkMaterialImportError(
        "本地转写完成，但没有生成可用正文。",
        metadata={"transcriptSource": engine.name, "transcriptEngine": engine.name, "tempMediaKind": "audio"},
    )


def _transcription_subprocess_env() -> dict[str, str]:
    env = dict(os.environ)
    try:
        import certifi  # type: ignore

        ca_bundle = certifi.where()
    except Exception:
        ca_bundle = ""
    if ca_bundle:
        env.setdefault("SSL_CERT_FILE", ca_bundle)
        env.setdefault("REQUESTS_CA_BUNDLE", ca_bundle)
    return env


def _classify_yt_dlp_access_failure(detail: str) -> str:
    lowered = str(detail or "").lower()
    if "412" in lowered or "precondition failed" in lowered:
        return "http_412"
    if "login" in lowered or "sign in" in lowered or "登录" in lowered:
        return "login_required"
    if "cookie" in lowered:
        return "cookie_required"
    # 图文笔记/无视频流：yt-dlp 找不到可下载的视频格式(小红书图文笔记最典型)。
    if (
        "requested format is not available" in lowered
        or "no video formats" in lowered
        or "no video could be found" in lowered
        or "there is no video" in lowered
        or "no formats found" in lowered
    ):
        return "no_video"
    return "unknown"


def _error_tail(detail: str, *, limit: int = 1000) -> str:
    text = re.sub(r"\s+", " ", str(detail or "")).strip()
    if len(text) <= limit:
        return text
    return text[-limit:]


def _build_yt_dlp_command(
    executable: list[str],
    options: LinkMaterialImportOptions,
    *args: str,
    headers_applied: bool = False,
    header_args: list[str] | None = None,
    impersonation_requested: bool = False,
    impersonation_target: str | None = None,
    use_browser_cookies: bool | None = None,
) -> list[str]:
    command = [*executable]
    if headers_applied:
        command.extend(header_args if header_args is not None else _BILIBILI_HEADER_ARGS)
    if impersonation_requested and impersonation_target:
        command.extend(["--impersonate", impersonation_target])
    if bool(use_browser_cookies):
        command.extend(["--cookies-from-browser", options.cookie_browser])
    command.extend(args)
    return command


def _has_curl_cffi() -> bool:
    return importlib.util.find_spec("curl_cffi") is not None


def _get_yt_dlp_version(executable: list[str]) -> str | None:
    try:
        completed = subprocess.run(
            [*executable, "--version"],
            capture_output=True,
            text=True,
            timeout=5,
            check=False,
        )
    except Exception:
        return None
    if completed.returncode != 0:
        return None
    return (completed.stdout or completed.stderr or "").strip() or None


def _get_yt_dlp_impersonate_target(executable: list[str]) -> str | None:
    if not _has_curl_cffi():
        return None
    try:
        completed = subprocess.run(
            [*executable, "--list-impersonate-targets"],
            capture_output=True,
            text=True,
            timeout=8,
            check=False,
        )
    except Exception:
        return None
    if completed.returncode != 0:
        return None
    output = completed.stdout or completed.stderr or ""
    candidates: list[str] = []
    for raw_line in output.splitlines():
        line = raw_line.strip()
        if not line or line.startswith("[") or line.startswith("-") or line.lower().startswith("client"):
            continue
        if "unavailable" in line.lower():
            continue
        first = line.split()[0].strip().lower()
        if first in {"chrome", "edge", "firefox", "safari"} or first.startswith("chrome"):
            candidates.append(first)
    for preferred in ("chrome", "edge", "firefox", "safari"):
        if preferred in candidates:
            return preferred
    return candidates[0] if candidates else None


def _find_bbdown() -> str | None:
    return shutil.which("BBDown") or shutil.which("bbdown")


def _find_yt_dlp() -> list[str] | None:
    # 0. encodings 修复(最优先、最稳):当前【已正确初始化】的 python 若装了 yt_dlp 模块,
    #    直接用 [sys.executable, "-m", "yt_dlp"](继承本进程 PYTHONHOME/PYTHONPATH)。
    #    必须在 PATH 查找之前 —— 否则打包版里 venv/bin 在 PATH 上时会返回裸 bin/yt-dlp 脚本,
    #    按构建期 stale shebang 启错 python → 子进程找不到 encodings 标准库(init_fs_encoding 失败)。
    import importlib.util as _importlib_util
    if sys.executable and _importlib_util.find_spec("yt_dlp") is not None:
        return [sys.executable, "-m", "yt_dlp"]
    # 1. PATH 里直接找
    executable = shutil.which("yt-dlp")
    if executable:
        return [executable]
    # 2. P11 修复 + encodings 修复：打包 venv 里带了 yt-dlp(bin/yt-dlp 存在 = yt_dlp 模块已装)。
    #    但直接跑 bin/yt-dlp 这个 shebang 脚本, 会按【构建期的绝对 python 路径】启另一个解释器,
    #    打包重定位到用户机后该路径失效 → 子进程 python 找不到 encodings 标准库(报 init_fs_encoding)。
    #    改用当前【已正确初始化、且 os.environ 已带 PYTHONHOME/PYTHONPATH】的 sys.executable -m yt_dlp。
    import os as _os
    venv_bin = _os.path.dirname(sys.executable) if sys.executable else ""
    if venv_bin and sys.executable:
        candidate = _os.path.join(venv_bin, "yt-dlp")
        if _os.path.isfile(candidate):
            return [sys.executable, "-m", "yt_dlp"]
    # 3. sys.executable -m yt_dlp（当前 python 里装了 yt_dlp 包）
    if sys.executable:
        try:
            completed = subprocess.run(
                [sys.executable, "-m", "yt_dlp", "--version"],
                capture_output=True, text=True, timeout=5, check=False,
            )
            if completed.returncode == 0:
                return [sys.executable, "-m", "yt_dlp"]
        except Exception:
            pass
    # 4. 兜底：PATH 里的 python3 / python
    python = shutil.which("python3") or shutil.which("python")
    if python:
        try:
            completed = subprocess.run(
                [python, "-m", "yt_dlp", "--version"],
                capture_output=True, text=True, timeout=5, check=False,
            )
            if completed.returncode == 0:
                return [python, "-m", "yt_dlp"]
        except Exception:
            return None
    return None


def _parse_yt_dlp_json(stdout: str) -> dict[str, Any]:
    for raw_line in reversed(str(stdout or "").splitlines()):
        line = raw_line.strip()
        if not line.startswith("{"):
            continue
        try:
            parsed = json.loads(line)
            return parsed if isinstance(parsed, dict) else {}
        except json.JSONDecodeError:
            continue
    return {}


def _read_downloaded_info_json(temp_dir: Path) -> dict[str, Any]:
    info_paths = sorted(
        temp_dir.glob("*.info.json"),
        key=lambda item: item.stat().st_mtime,
        reverse=True,
    )
    for path in info_paths:
        try:
            parsed = json.loads(path.read_text(encoding="utf-8", errors="ignore"))
        except Exception:
            continue
        if isinstance(parsed, dict):
            return parsed
    return {}


def read_valid_downloaded_subtitles(temp_dir: Path) -> str:
    subtitle_paths = sorted(
        [
            item
            for item in temp_dir.glob("*")
            if item.is_file() and item.suffix.lower() in {".vtt", ".srt", ".ass", ".ssa"}
        ],
        key=lambda item: (
            0 if re.search(r"(zh|cn|hans|chinese)", item.name, re.I) else 1,
            item.name,
        ),
    )
    for path in subtitle_paths:
        try:
            text = path.read_text(encoding="utf-8", errors="ignore")
        except Exception:
            continue
        cleaned = _clean_subtitle_text(text)
        if len(cleaned) >= 20:
            return cleaned
    return ""


def _read_downloaded_subtitles(temp_dir: Path) -> str:
    return read_valid_downloaded_subtitles(temp_dir)


def _clean_subtitle_text(text: str) -> str:
    lines: list[str] = []
    previous = ""
    for raw_line in str(text or "").splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.upper().startswith("WEBVTT"):
            continue
        if re.match(r"^\d+$", line):
            continue
        if "-->" in line:
            continue
        line = re.sub(r"<[^>]+>", "", line)
        line = re.sub(r"\{\\.*?\}", "", line)
        line = re.sub(r"^\[.*?\]\s*", "", line)
        line = re.sub(r"\s+", " ", line).strip()
        if not line or line == previous:
            continue
        lines.append(line)
        previous = line
    return "\n".join(lines).strip()


def _strip_low_value_transcript_text(text: str) -> str:
    text = str(text or "").replace("\r\n", "\n")
    text = re.sub(r"(?im)^原链接：.*$", "", text)
    text = re.sub(r"(?im)^#\s+.*$", "", text)
    text = re.sub(r"(?im)^(以下是|这是|我将|作为).{0,40}$", "", text)
    text = re.sub(r"(?m)^\s*\d{1,2}:\d{2}(?::\d{2})?.*$", "", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    return text.strip()


def _clean_title(title: str) -> str:
    cleaned = re.sub(r"\s+", " ", str(title or "")).strip()
    cleaned = re.sub(r"[\\/:*?\"<>|]", " ", cleaned)
    cleaned = re.sub(r"\s{2,}", " ", cleaned).strip()
    return cleaned[:80]
