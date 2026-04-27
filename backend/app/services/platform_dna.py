from __future__ import annotations

from io import BytesIO
from pathlib import Path
import re
import zipfile
from xml.etree import ElementTree as ET

from docx import Document as WordDocument

try:
    from pypdf import PdfReader

    HAS_PYPDF = True
except Exception:  # pragma: no cover - dependency fallback
    PdfReader = None  # type: ignore[assignment]
    HAS_PYPDF = False


TEXT_EXTENSIONS = {".md", ".markdown", ".txt"}
DOCX_EXTENSIONS = {".docx"}
PDF_EXTENSIONS = {".pdf"}


def _read_plain_text(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8", errors="ignore")
    except Exception:
        return ""


def _read_docx_with_python_docx(path: Path) -> str:
    try:
        document = WordDocument(str(path))
    except Exception:
        return ""

    lines: list[str] = []
    for paragraph in document.paragraphs:
        text = re.sub(r"\s+", " ", paragraph.text or "").strip()
        if text:
            lines.append(text)

    for table in document.tables:
        for row in table.rows:
            values = [re.sub(r"\s+", " ", cell.text or "").strip() for cell in row.cells]
            values = [value for value in values if value]
            if values:
                lines.append(" | ".join(values))

    return "\n".join(lines).strip()


def _read_docx_xml_fallback(path: Path) -> str:
    texts: list[str] = []
    try:
        with zipfile.ZipFile(path) as archive:
            for name in archive.namelist():
                if not name.startswith("word/") or not name.endswith(".xml"):
                    continue
                try:
                    payload = archive.read(name)
                except KeyError:
                    continue
                try:
                    root = ET.fromstring(payload)
                except ET.ParseError:
                    continue
                texts.extend([node.strip() for node in root.itertext() if node and node.strip()])
    except Exception:
        return ""
    return "\n".join(texts).strip()


def _read_pdf(path: Path) -> str:
    if not HAS_PYPDF or PdfReader is None:
        return ""

    pages: list[str] = []
    try:
        reader = PdfReader(str(path))
    except Exception:
        return ""

    for page in getattr(reader, "pages", []):
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        normalized = re.sub(r"\s+", " ", text).strip()
        if normalized:
            pages.append(normalized)
    return "\n\n".join(pages).strip()


def extract_platform_dna_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in TEXT_EXTENSIONS:
        return _read_plain_text(path).strip()
    if suffix in DOCX_EXTENSIONS:
        return (_read_docx_with_python_docx(path) or _read_docx_xml_fallback(path)).strip()
    if suffix in PDF_EXTENSIONS:
        return _read_pdf(path).strip()
    raise ValueError(f"unsupported_extension:{suffix}")


def supported_platform_dna_extensions() -> tuple[str, ...]:
    return tuple(sorted(TEXT_EXTENSIONS | DOCX_EXTENSIONS | PDF_EXTENSIONS))
