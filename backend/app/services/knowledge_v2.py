from __future__ import annotations

import hashlib
import base64
import logging
from io import BytesIO
import re
import shutil
import zipfile
from collections import Counter, defaultdict
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any, Callable
from uuid import uuid4
from xml.etree import ElementTree as ET

from docx import Document as WordDocument

from app.db import Database, from_json, to_json
from app.services.data_center_access import (
    DataCenterAccessContext,
    build_document_access_where,
    normalize_department_ids,
)
from app.services.knowledge_base import append_file_reclass_event, get_vector_runtime_status
from app.services.workspace_relation_docs import materialize_workspace_relation_documents

logger = logging.getLogger(__name__)

try:
    from pypdf import PdfReader

    HAS_PYPDF = True
except Exception:  # pragma: no cover - runtime dependency
    PdfReader = None  # type: ignore[assignment]
    HAS_PYPDF = False

try:
    import fitz  # type: ignore[import-untyped]

    HAS_PYMUPDF = True
except Exception:  # pragma: no cover - optional runtime dependency
    fitz = None  # type: ignore[assignment]
    HAS_PYMUPDF = False

try:
    from PIL import Image, ImageOps

    HAS_PILLOW = True
except Exception:  # pragma: no cover - optional runtime dependency
    Image = None  # type: ignore[assignment]
    ImageOps = None  # type: ignore[assignment]
    HAS_PILLOW = False


V2_PIPELINE_VERSION = "v2-minimal-evidence"
MAIN_KNOWLEDGE_STATUS_JOB_TYPES = ("ingest_import", "backfill_workspace_import")
_AUXILIARY_KNOWLEDGE_JOB_TYPES = ("generate_client_dna_candidates",)
TEXT_EXTENSIONS = {".md", ".txt", ".json", ".csv"}
ARCHIVE_XML_EXTENSIONS = {".pptx", ".xlsx"}
IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".webp"}
ONLINE_TRANSCRIPT_CATEGORY = "线上转写"
UNIFIED_FOLDER_LABEL = "资料库"
SYSTEM_FOLDER_CATEGORIES = [UNIFIED_FOLDER_LABEL]
# 旧常量保留：迁移脚本需要识别这些目录作为搬运源。
LEGACY_SYSTEM_FOLDER_CATEGORIES = ["收件箱", ONLINE_TRANSCRIPT_CATEGORY, "待处理", "归档"]
LEGACY_FIXED_CATEGORIES = [
    "财务与筹款",
    "品牌与传播",
    "项目与业务",
    "组织与战略",
    "其他资料",
    ONLINE_TRANSCRIPT_CATEGORY,
    "战略陪伴",
]
HUMAN_VISIBLE_CATEGORIES = LEGACY_FIXED_CATEGORIES  # kept for backward compat; new clients use dynamic folders
EVIDENCE_CATEGORIES = [category for category in [*LEGACY_FIXED_CATEGORIES, *SYSTEM_FOLDER_CATEGORIES] if category != "战略陪伴"]
DEFAULT_INBOX_LABEL = "收件箱"
WORKSPACE_BACKFILL_EXTENSIONS = {".pdf", ".docx", ".md", ".txt", ".pptx", ".xlsx", ".json", ".csv", *IMAGE_EXTENSIONS}
QUERY_STOPWORDS = {
    "请",
    "一下",
    "一个",
    "一些",
    "结合",
    "现有",
    "资料",
    "当前",
    "这个",
    "那个",
    "我们",
    "你们",
    "他们",
    "以及",
    "需要",
    "进行",
    "问题",
    "情况",
}
INTRO_QUERY_TOKENS = ("介绍", "简介", "核心定位", "概况", "是什么", "项目特点", "怎么做", "业务", "创始人")
MEETING_QUERY_TOKENS = ("最近一次会议", "会议讲了什么", "会议纪要", "会里", "会议")
INTRO_BUCKET_KEYWORDS: dict[str, tuple[str, ...]] = {
    "institution": ("核心定位", "历史价值", "机构", "基金会", "平台", "定位", "介绍", "概况", "关于"),
    "project": ("项目", "计划", "业务", "学院", "赋能", "服务", "行动营", "平台"),
    "strategy": ("战略", "升级", "路线图", "规划", "重点事项", "第二曲线", "转型"),
    "source": ("申请书", "访谈", "调研", "对话", "转写", "word版"),
}
INTRO_OPERATIONAL_MARKERS = (
    "会议纪要",
    "会议",
    "沟通",
    "待办",
    "清单",
    "任务",
    "询价单",
    "报销",
    "联调",
    "q1",
    "q2",
    "q3",
    "q4",
    "工作坊",
    "ppt",
    "转写原文版",
)
DERIVED_TITLE_MARKERS = ("介绍", "简介", "精简版", "预览", "说明", "ppt", "工作台")
DERIVED_TEXT_MARKERS = (
    "完全基于你提供的材料",
    "为了便于你后续直接做 ppt",
    "为了便于你后续直接做ppt",
    "为了便于你后续直接做 ppt 或对外介绍",
    "为了便于你后续直接做ppt或对外介绍",
    "对外介绍",
    "下面给出",
    "可直接拆进",
    "产品策划 + 功能与价值 + 页面级设计",
    "我用“定位—交付—工作模式—业务价值—近期重点”的结构来写",
    "我用“定位-交付-工作模式-业务价值-近期重点”的结构来写",
    "一句话定义",
)
DERIVED_CITATION_PATTERN = re.compile(r"【\d{1,4}:\d{1,4}†|\[\d{1,4}:\d{1,4}†")
NOISE_FILES = {".ds_store", "thumbs.db", "~$"}
CATEGORY_KEYWORDS = {
    "财务与筹款": ("财务", "预算", "筹款", "募资", "捐赠", "赞助", "基金", "现金流", "审计", "年报"),
    "品牌与传播": ("品牌", "传播", "公关", "媒体", "内容", "公众号", "活动", "叙事", "视觉", "案例"),
    "项目与业务": ("项目", "业务", "交付", "方案", "行业", "客户", "服务", "运营", "年会", "图书馆"),
    "组织与战略": ("战略", "组织", "治理", "团队", "董事会", "理事会", "路线图", "规划", "会议纪要", "访谈"),
}
SECTION_BREAK_PATTERN = re.compile(r"\n{2,}")
CHUNK_TARGET_CHARS = 680
CHUNK_OVERLAP_CHARS = 120
INTERNAL_WORKSPACE_MARKERS = {"_v2_meta", "_imports"}
OCR_DEFAULT_MAX_PAGES = 60
OCR_DEFAULT_BATCH_SIZE = 8
SEARCHABLE_PARSE_STATUSES = {"ready", "partial_ready"}
SOURCE_AVAILABILITY_ORIGINAL = "original_available"
SOURCE_AVAILABILITY_MACHINE_ONLY = "machine_readable_only"
SOURCE_AVAILABILITY_INVALID = "invalid_source"
SOURCE_AVAILABILITY_UNKNOWN = "unknown"

PLACEHOLDER_MACHINE_TEXT_MARKERS = (
    "解析重试",
    "解析后暂无可用正文",
    "用途：兼容旧字段",
    "已进入资料缓冲池",
    "资料缓冲池说明",
)


@dataclass
class CitationMatch:
    knowledge_document_id: str
    chunk_id: str | None
    title: str
    excerpt: str
    score: float
    coverage: float
    section_label: str | None
    source_stage: str
    drillthrough_used: bool
    matched_terms: list[str]
    path: str | None
    original_path: str | None = None
    managed_path: str | None = None
    markdown_path: str | None = None
    openable_kind: str | None = None
    document_family_id: str | None = None
    canonical_kind: str | None = None
    origin_type: str | None = None
    origin_id: str | None = None
    is_searchable: bool | None = None
    source_availability: str | None = None
    original_available: bool | None = None
    machine_readable_available: bool | None = None
    open_original_disabled_reason: str | None = None
    display_path: str | None = None
    virtual_optimized_path: str | None = None
    path_optimization_status: str | None = None
    path_optimization_confidence: float | None = None
    purpose: str | None = None
    audience: str | None = None
    project_context: str | None = None
    key_topics: list[str] = field(default_factory=list)
    good_questions: list[str] = field(default_factory=list)
    risk_notes: str | None = None


@dataclass
class RetrievalBundle:
    citations: list[CitationMatch]
    coverage: float
    retrieval_summary: dict[str, Any]
    context_text: str
    matched_terms: list[str]
    failure_reason: str | None = None


@dataclass
class ExtractionMetadata:
    parse_status: str = "failed"
    parse_error: str | None = None
    failure_type: str | None = None
    total_pages: int | None = None
    attempted_pages: int = 0
    succeeded_pages: int = 0
    failed_pages: int = 0
    ocr_page_limit: int | None = None
    ocr_batch_size: int | None = None
    partial: bool = False


@dataclass
class ExtractedDocument:
    text: str
    sections: list[dict[str, str]]
    metadata: ExtractionMetadata
    # Phase 1：xlsx 走 structured parser 时，把 ParsedSheet 列表带回来，
    # 让 ingest_document_knowledge 顺手写到 structured_tables。pptx 不会用
    # 这个字段（slide 数据已经在 sections 里）。
    structured_sheets: list[Any] = field(default_factory=list)


@dataclass
class VisualOcrResult:
    text: str
    sections: list[dict[str, str]]
    attempted_pages: int = 0
    succeeded_pages: int = 0
    failed_pages: int = 0


def now_iso() -> str:
    return datetime.now().replace(microsecond=0).isoformat()


def new_id(prefix: str) -> str:
    return f"{prefix}_{uuid4().hex[:10]}"


def classify_parse_failure_type(error_text: str, *, kind: str = "") -> str:
    lowered = str(error_text or "").strip().lower()
    normalized_kind = str(kind or "").strip().lower().lstrip(".")
    if any(token in lowered for token in ("not found", "no such file", "不存在")):
        return "file_missing"
    if any(token in lowered for token in ("permission denied", "权限")):
        return "permission_denied"
    if any(token in lowered for token in ("unsupported", "不支持")):
        return "unsupported_format"
    if normalized_kind == "pdf" and any(token in lowered for token in ("empty", "为空", "无正文", "可用正文", "no text", "未能解析")):
        return "empty_pdf"
    if any(token in lowered for token in ("ocr", "scan", "扫描")):
        return "ocr_required"
    if any(token in lowered for token in ("empty", "为空", "无正文", "可用正文", "no text", "未能解析")):
        return "empty_text"
    if any(token in lowered for token in ("parser", "解析", "exception", "traceback")):
        return "parser_exception"
    return "unknown"


def parse_failure_recoverable(failure_type: str, *, kind: str = "") -> bool:
    normalized = str(failure_type or "").strip()
    normalized_kind = str(kind or "").strip().lower().lstrip(".")
    if normalized in {"empty_pdf", "ocr_required", "empty_text", "parser_exception", "unknown"}:
        return True
    if normalized == "unsupported_format" and normalized_kind in {"pdf", "jpg", "jpeg", "png", "webp", "docx", "txt", "md"}:
        return True
    return False


def _normalized_family_stem(name: str) -> str:
    stem = safe_filename(Path(str(name or "").strip()).stem or str(name or "").strip()).lower()
    stem = re.sub(r"^\d{8}[_-]?\d{0,6}[_-]*", "", stem)
    stem = re.sub(r"^\d{10,14}[_-]*", "", stem)
    stem = re.sub(r"^副本+", "", stem)
    stem = re.sub(r"^(copy|copied)_*", "", stem)
    stem = re.sub(r"^#+\s*", "", stem)
    stem = re.sub(r"^[一二三四五六七八九十0-9]+[、.．_\-\s]+", "", stem)
    stem = re.sub(r"__+[a-z0-9]{4,}$", "", stem)
    stem = re.sub(r"[\(\[]\d+[\)\]]", "", stem)
    stem = re.sub(r"[_\-\s]+", "", stem)
    return stem[:120]


def _effective_family_id_for_row(row: Any) -> str:
    stored = str(row["document_family_id"] or "").strip()
    canonical_kind = _canonical_kind_for_row(row)
    title = str(row["file_name"] or row["document_title"] or "")
    if canonical_kind == "raw_file":
        normalized = _normalized_family_stem(title)
        if normalized:
            return f"raw_file:{normalized}"
    if stored:
        return stored
    return f"doc:{row['document_id']}"


def _query_style(prompt: str) -> str:
    normalized = re.sub(r"\s+", "", str(prompt or "").lower())
    if any(token in normalized for token in (token.lower() for token in MEETING_QUERY_TOKENS)):
        return "meeting"
    if any(token in normalized for token in (token.lower() for token in INTRO_QUERY_TOKENS)):
        return "intro"
    return "general"


def _family_support_text(family: dict[str, Any]) -> str:
    lines: list[str] = []
    for item in family.get("rows", []):
        row = item.get("row") if isinstance(item, dict) else None
        if row is None:
            continue
        lines.extend(
            [
                str(item.get("title") or ""),
                str(item.get("path") or ""),
                str(row["visible_category"] or ""),
                str(row["secondary_category"] or ""),
                str(row["preview_text"] or ""),
            ]
        )
    return "\n".join(lines).lower()


def _intro_bucket_from_support_text(support_text: str) -> str | None:
    for bucket in ("source", "project", "strategy", "institution"):
        keywords = INTRO_BUCKET_KEYWORDS.get(bucket, ())
        if any(keyword.lower() in support_text for keyword in keywords):
            return bucket
    return None


def _intro_bucket_for_family(family: dict[str, Any]) -> str | None:
    if "raw_file" not in family.get("canonicalKinds", set()):
        return None
    support_text = _family_support_text(family)
    return _intro_bucket_from_support_text(support_text)


def _is_intro_primary_family(family: dict[str, Any]) -> bool:
    if "raw_file" not in family.get("canonicalKinds", set()):
        return False
    support_text = _family_support_text(family)
    if not _intro_bucket_from_support_text(support_text):
        return False
    return not any(marker.lower() in support_text for marker in INTRO_OPERATIONAL_MARKERS)


def backfill_client_document_family_metadata(db: Database, client_id: str) -> dict[str, int]:
    rows = db.fetchall(
        """
        SELECT
            d.id AS document_id,
            d.title AS document_title,
            d.document_family_id AS doc_family_id,
            d.canonical_kind AS doc_canonical_kind,
            d.origin_type AS doc_origin_type,
            d.origin_id AS doc_origin_id,
            d.is_searchable AS doc_is_searchable,
            vd.id AS v2_document_id,
            vd.file_name,
            vd.content_hash,
            vd.document_family_id AS v2_family_id,
            vd.canonical_kind AS v2_canonical_kind,
            vd.origin_type AS v2_origin_type,
            vd.origin_id AS v2_origin_id,
            vd.is_searchable AS v2_is_searchable
        FROM v2_documents vd
        JOIN documents d ON d.id = vd.document_id
        WHERE vd.client_id = ?
          AND COALESCE(vd.canonical_kind, 'raw_file') = 'raw_file'
        """,
        (client_id,),
    )
    updated_documents = 0
    updated_v2_documents = 0
    for row in rows:
        resolved_family_id = derive_document_family_id(
            title=str(row["file_name"] or row["document_title"] or ""),
            content_hash=str(row["content_hash"] or ""),
            origin_type="file_import",
            origin_id=str(row["document_id"] or ""),
        )
        if str(row["doc_family_id"] or "") != resolved_family_id:
            db.execute(
                """
                UPDATE documents
                SET document_family_id = ?, canonical_kind = 'raw_file', origin_type = 'file_import', origin_id = ?, is_searchable = 1
                WHERE id = ?
                """,
                (resolved_family_id, str(row["document_id"] or ""), str(row["document_id"] or "")),
            )
            updated_documents += 1
        if str(row["v2_family_id"] or "") != resolved_family_id:
            db.execute(
                """
                UPDATE v2_documents
                SET document_family_id = ?, canonical_kind = 'raw_file', origin_type = 'file_import', origin_id = ?, is_searchable = 1
                WHERE id = ?
                """,
                (resolved_family_id, str(row["document_id"] or ""), str(row["v2_document_id"] or "")),
            )
            updated_v2_documents += 1

    db.execute(
        """
        UPDATE document_chunks
        SET
            document_family_id = COALESCE(
                (
                    SELECT d.document_family_id
                    FROM knowledge_documents kd
                    JOIN documents d ON d.id = kd.document_id
                    WHERE kd.id = document_chunks.knowledge_document_id
                    LIMIT 1
                ),
                document_family_id
            ),
            canonical_kind = COALESCE(
                (
                    SELECT d.canonical_kind
                    FROM knowledge_documents kd
                    JOIN documents d ON d.id = kd.document_id
                    WHERE kd.id = document_chunks.knowledge_document_id
                    LIMIT 1
                ),
                canonical_kind,
                'raw_file'
            ),
            origin_type = COALESCE(
                (
                    SELECT d.origin_type
                    FROM knowledge_documents kd
                    JOIN documents d ON d.id = kd.document_id
                    WHERE kd.id = document_chunks.knowledge_document_id
                    LIMIT 1
                ),
                origin_type,
                'file_import'
            ),
            origin_id = COALESCE(
                (
                    SELECT d.origin_id
                    FROM knowledge_documents kd
                    JOIN documents d ON d.id = kd.document_id
                    WHERE kd.id = document_chunks.knowledge_document_id
                    LIMIT 1
                ),
                origin_id
            ),
            is_searchable = COALESCE(
                (
                    SELECT d.is_searchable
                    FROM knowledge_documents kd
                    JOIN documents d ON d.id = kd.document_id
                    WHERE kd.id = document_chunks.knowledge_document_id
                    LIMIT 1
                ),
                is_searchable,
                1
            )
        WHERE knowledge_document_id IN (
            SELECT id FROM knowledge_documents WHERE client_id = ?
        )
        """,
        (client_id,),
    )
    return {
        "updatedDocuments": updated_documents,
        "updatedV2Documents": updated_v2_documents,
    }


def derive_document_family_id(
    *,
    title: str,
    content_hash: str,
    origin_type: str = "file_import",
    origin_id: str = "",
) -> str:
    origin_type_clean = str(origin_type or "file_import").strip() or "file_import"
    origin_id_clean = str(origin_id or "").strip()
    if origin_type_clean != "file_import" and origin_id_clean:
        return f"{origin_type_clean}:{origin_id_clean}"
    stem = _normalized_family_stem(title)
    if stem:
        return f"{origin_type_clean}:{stem}"
    return f"{origin_type_clean}:{str(content_hash or 'unknown')[:24]}"


def _clean_ingested_text(text: str) -> str:
    normalized = normalize_text(text)
    if not normalized:
        return ""
    lines: list[str] = []
    for raw_line in normalized.splitlines():
        line = raw_line.strip()
        compact = re.sub(r"\s+", "", line).lower()
        if not compact:
            lines.append("")
            continue
        if any(marker in compact for marker in ("说明.txt", "整理说明", "重复件已移至废纸篓")):
            continue
        if "wps演示" in compact or "已用的字体" in compact or "字体列表" in compact or "模板页" in compact:
            continue
        if re.fullmatch(r"[\d\W_]+", compact):
            continue
        if len(compact) <= 2:
            continue
        lines.append(line)
    return normalize_text("\n".join(lines))


def _is_ingestion_noise_document(title: str, text: str) -> bool:
    normalized = f"{safe_filename(title).lower()}\n{str(text or '').lower()}"
    if any(marker in normalized for marker in ("说明.txt", "整理说明", "重复件已移至废纸篓")):
        return True
    if "wps 演示" in normalized and "已用的字体" in normalized:
        return True
    return False


def _system_document_path(data_dir: Path, client_id: str, canonical_kind: str, origin_id: str, title: str) -> Path:
    root = client_workspace_root(data_dir, client_id) / "_v2_meta" / "system_docs" / canonical_kind
    root.mkdir(parents=True, exist_ok=True)
    stem = safe_filename(Path(title).stem or origin_id or canonical_kind)
    return root / f"{origin_id or stem}_{stem}.md"


def serialize_retrieval_bundle(bundle: RetrievalBundle) -> dict[str, Any]:
    return {
        "citations": [
            {
                "knowledge_document_id": item.knowledge_document_id,
                "chunk_id": item.chunk_id,
                "title": item.title,
                "excerpt": item.excerpt,
                "score": item.score,
                "coverage": item.coverage,
                "section_label": item.section_label,
                "source_stage": item.source_stage,
                "drillthrough_used": item.drillthrough_used,
                "matched_terms": item.matched_terms,
                "path": item.path,
                "original_path": item.original_path,
                "managed_path": item.managed_path,
                "markdown_path": item.markdown_path,
                "openable_kind": item.openable_kind,
                "source_availability": item.source_availability,
                "original_available": item.original_available,
                "machine_readable_available": item.machine_readable_available,
                "open_original_disabled_reason": item.open_original_disabled_reason,
                "display_path": item.display_path,
                "virtual_optimized_path": item.virtual_optimized_path,
                "path_optimization_status": item.path_optimization_status,
                "path_optimization_confidence": item.path_optimization_confidence,
                "purpose": item.purpose,
                "audience": item.audience,
                "project_context": item.project_context,
                "key_topics": item.key_topics,
                "good_questions": item.good_questions,
                "risk_notes": item.risk_notes,
            }
            for item in bundle.citations
        ],
        "coverage": bundle.coverage,
        "retrieval_summary": bundle.retrieval_summary,
        "context_text": bundle.context_text,
        "matched_terms": bundle.matched_terms,
        "failure_reason": bundle.failure_reason,
    }


def deserialize_retrieval_bundle(payload: dict[str, Any]) -> RetrievalBundle:
    citations = []
    for item in payload.get("citations", []):
        if not isinstance(item, dict):
            continue
        citations.append(
            CitationMatch(
                knowledge_document_id=str(item.get("knowledge_document_id") or ""),
                chunk_id=str(item["chunk_id"]) if item.get("chunk_id") else None,
                title=str(item.get("title") or ""),
                excerpt=str(item.get("excerpt") or ""),
                score=float(item.get("score") or 0.0),
                coverage=float(item.get("coverage") or 0.0),
                section_label=str(item["section_label"]) if item.get("section_label") else None,
                source_stage=str(item.get("source_stage") or "raw_chunk"),
                drillthrough_used=bool(item.get("drillthrough_used", False)),
                matched_terms=[str(term) for term in item.get("matched_terms", []) if str(term).strip()],
                path=str(item["path"]) if item.get("path") else None,
                original_path=str(item["original_path"]) if item.get("original_path") else None,
                managed_path=str(item["managed_path"]) if item.get("managed_path") else None,
                markdown_path=str(item["markdown_path"]) if item.get("markdown_path") else None,
                openable_kind=str(item["openable_kind"]) if item.get("openable_kind") else None,
                source_availability=str(item["source_availability"]) if item.get("source_availability") else None,
                original_available=bool(item.get("original_available")) if item.get("original_available") is not None else None,
                machine_readable_available=bool(item.get("machine_readable_available")) if item.get("machine_readable_available") is not None else None,
                open_original_disabled_reason=str(item["open_original_disabled_reason"]) if item.get("open_original_disabled_reason") else None,
                display_path=str(item["display_path"]) if item.get("display_path") else None,
                virtual_optimized_path=str(item["virtual_optimized_path"]) if item.get("virtual_optimized_path") else None,
                path_optimization_status=str(item["path_optimization_status"]) if item.get("path_optimization_status") else None,
                path_optimization_confidence=float(item.get("path_optimization_confidence")) if item.get("path_optimization_confidence") is not None else None,
                purpose=str(item["purpose"]) if item.get("purpose") else None,
                audience=str(item["audience"]) if item.get("audience") else None,
                project_context=str(item["project_context"]) if item.get("project_context") else None,
                key_topics=[str(term) for term in item.get("key_topics", []) if str(term).strip()],
                good_questions=[str(term) for term in item.get("good_questions", []) if str(term).strip()],
                risk_notes=str(item["risk_notes"]) if item.get("risk_notes") else None,
            )
        )
    retrieval_summary = payload.get("retrieval_summary", {})
    return RetrievalBundle(
        citations=citations,
        coverage=float(payload.get("coverage") or 0.0),
        retrieval_summary=retrieval_summary if isinstance(retrieval_summary, dict) else {},
        context_text=str(payload.get("context_text") or ""),
        matched_terms=[str(term) for term in payload.get("matched_terms", []) if str(term).strip()],
        failure_reason=str(payload["failure_reason"]) if payload.get("failure_reason") else None,
    )


def safe_filename(name: str) -> str:
    cleaned = re.sub(r"[\\/:*?\"<>|]+", "_", name).strip()
    cleaned = re.sub(r"\s+", " ", cleaned)
    return cleaned[:180] or "untitled"


def ensure_client_folder_rows(db: Database, data_dir: Path, client_id: str) -> None:
    timestamp = now_iso()
    workspace_root = data_dir / "client_workspace" / client_id
    workspace_root.mkdir(parents=True, exist_ok=True)
    for index, label in enumerate(SYSTEM_FOLDER_CATEGORIES):
        path = workspace_root / label
        path.mkdir(parents=True, exist_ok=True)
        existing = db.fetchone(
            "SELECT id FROM client_folders WHERE client_id = ? AND label = ?",
            (client_id, label),
        )
        if existing:
            db.execute(
                """
                UPDATE client_folders
                SET path = ?, last_scanned_at = ?, folder_kind = 'system', source_type = 'system',
                    is_system = 1, is_hidden = 0, sort_order = ?
                WHERE id = ?
                """,
                (str(path), timestamp, index * 10, str(existing["id"])),
            )
        else:
            db.execute(
                """
                INSERT INTO client_folders(
                    id, client_id, label, path, file_count, last_scanned_at,
                    folder_kind, source_type, is_system, is_hidden, sort_order, created_by_rule,
                    created_at
                )
                VALUES(?, ?, ?, ?, 0, ?, 'system', 'system', 1, 0, ?, 'system_default', ?)
                """,
                (new_id("fld"), client_id, label, str(path), timestamp, index * 10, timestamp),
            )


def client_workspace_root(data_dir: Path, client_id: str) -> Path:
    root = data_dir / "client_workspace" / client_id
    root.mkdir(parents=True, exist_ok=True)
    return root


def normalize_text(text: str) -> str:
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    normalized = re.sub(r"\u00a0", " ", normalized)
    normalized = re.sub(r"[ \t]+", " ", normalized)
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)
    return normalized.strip()


def tokenize(text: str) -> list[str]:
    normalized = re.sub(r"[^\w\u4e00-\u9fff]+", " ", text.lower())
    tokens: list[str] = []
    seen: set[str] = set()
    for word in re.findall(r"[a-z0-9]{2,}", normalized):
        if word not in seen:
            seen.add(word)
            tokens.append(word)
    for block in re.findall(r"[\u4e00-\u9fff]{2,}", normalized):
        compact = "".join(char for char in block if char not in "的是了和在对与及并把将中上为等于")
        if len(compact) < 2:
            continue
        candidates = {compact}
        if len(compact) >= 2:
            candidates.update(compact[index : index + 2] for index in range(len(compact) - 1))
        if len(compact) >= 4:
            candidates.update(compact[index : index + 4] for index in range(len(compact) - 3))
        for candidate in candidates:
            if len(candidate) < 2 or candidate in QUERY_STOPWORDS or candidate in seen:
                continue
            seen.add(candidate)
            tokens.append(candidate)
    return tokens[:48]


def is_strategy_analysis_query(text: str) -> bool:
    lowered = text.lower()
    return any(
        token in lowered
        for token in ("战略", "主线", "关键取舍", "风险", "判断", "诊断", "路径", "张力", "重点", "取舍")
    )


def build_excerpt(text: str, fallback_title: str) -> str:
    cleaned = normalize_text(text)
    if not cleaned:
        return f"{fallback_title} 解析后暂无可用正文，请检查原文件是否为扫描件或图片。"
    return cleaned[:220]


def stage_import_copy(data_dir: Path, client_id: str, import_id: str, source_path: Path) -> Path:
    intake_root = client_workspace_root(data_dir, client_id) / "_imports" / import_id
    intake_root.mkdir(parents=True, exist_ok=True)
    safe_name = safe_filename(source_path.name)
    target = intake_root / safe_name
    if source_path.resolve() == target.resolve():
        return target
    if target.exists():
        stem = safe_filename(source_path.stem)
        suffix = source_path.suffix.lower()
        target = intake_root / f"{stem}__{uuid4().hex[:6]}{suffix}"
    shutil.copy2(source_path, target)
    return target


def detect_category(title: str, text: str, *, custom_labels: list[str] | None = None) -> tuple[str, str, float]:
    haystack = f"{title}\n{text[:2200]}".lower()
    scores: dict[str, int] = {}
    for category, keywords in CATEGORY_KEYWORDS.items():
        scores[category] = sum(2 if keyword.lower() in title.lower() else 1 for keyword in keywords if keyword.lower() in haystack)
    # Also score against user-created custom folder labels (simple keyword match)
    if custom_labels:
        for label in custom_labels:
            if label in CATEGORY_KEYWORDS or label == DEFAULT_INBOX_LABEL:
                continue
            label_lower = label.lower()
            # Check if the label text appears in the document
            score = 0
            for word in label_lower.replace("与", " ").replace("和", " ").replace("/", " ").split():
                if len(word) >= 2 and word in haystack:
                    score += 2 if word in title.lower() else 1
            if score > 0:
                scores[label] = score
    best_category = max(scores, key=scores.get) if scores else DEFAULT_INBOX_LABEL
    best_score = scores.get(best_category, 0)
    if best_score <= 0:
        return DEFAULT_INBOX_LABEL, "待分类", 0.35
    secondary = "核心资料" if best_score >= 4 else "一般资料"
    confidence = min(0.95, 0.5 + best_score * 0.08)
    return best_category, secondary, confidence


def detect_material_profile(title: str, text: str, primary_category: str, secondary_category: str, confidence: float) -> tuple[str, str, str, float]:
    title_text = safe_filename(Path(title).stem).lower()
    body = normalize_text(text).lower()
    head = body[:4800]
    derived_score = 0

    if any(marker in title_text for marker in DERIVED_TITLE_MARKERS):
        derived_score += 1
    if re.search(r"第\s*[0-9一二三四五六七八九十]+\s*稿", title_text):
        derived_score += 3
    if "精简版" in title_text or "预览" in title_text:
        derived_score += 4
    if "介绍" in title_text or "简介" in title_text:
        derived_score += 2
    if "说明" in title_text and primary_category != "财务与筹款":
        derived_score += 1
    if any(marker in head for marker in DERIVED_TEXT_MARKERS):
        derived_score += 3
    if "为了便于你后续直接做" in head and ("ppt" in head or "对外介绍" in head):
        derived_score += 3
    if ("下面给出" in head or "一句话定义" in head) and ("产品策划" in head or "页面级设计" in head or "功能与价值" in head):
        derived_score += 3
    if "完全基于你提供的材料" in head:
        derived_score += 4
    if DERIVED_CITATION_PATTERN.search(head):
        derived_score += 2

    if derived_score >= 4:
        return "background", "战略陪伴", "派生整理稿", max(confidence, 0.86)
    return "evidence", primary_category, secondary_category, confidence


def _archive_xml_text(path: Path) -> str:
    texts: list[str] = []
    with zipfile.ZipFile(path) as archive:
        for name in sorted(archive.namelist()):
            if not name.endswith(".xml"):
                continue
            try:
                root = ET.fromstring(archive.read(name))
            except ET.ParseError:
                continue
            for element in root.iter():
                value = (element.text or "").strip()
                if value:
                    texts.append(value)
    return normalize_text("\n".join(texts))


def _read_plain_text(path: Path) -> str:
    try:
        return normalize_text(path.read_text("utf-8"))
    except UnicodeDecodeError:
        return normalize_text(path.read_text("utf-8", errors="ignore"))


def _read_docx_text(path: Path) -> tuple[str, list[dict[str, str]]]:
    document = WordDocument(path)
    sections: list[dict[str, str]] = []
    current_title = "正文"
    buffer: list[str] = []

    def flush() -> None:
        nonlocal buffer
        content = normalize_text("\n".join(buffer))
        if content:
            sections.append({"title": current_title, "text": content})
        buffer = []

    for paragraph in document.paragraphs:
        text = normalize_text(paragraph.text)
        if not text:
            continue
        style_name = paragraph.style.name.lower() if paragraph.style and paragraph.style.name else ""
        if "heading" in style_name:
            flush()
            current_title = text
        else:
            buffer.append(text)
    for table in document.tables:
        rows: list[str] = []
        for row in table.rows:
            cells = [normalize_text(cell.text) for cell in row.cells]
            line = " | ".join(cell for cell in cells if cell)
            if line:
                rows.append(line)
        if rows:
            buffer.append("\n".join(rows))
    flush()
    if not sections:
        text = normalize_text("\n".join(paragraph.text for paragraph in document.paragraphs))
        if text:
            sections = [{"title": "正文", "text": text}]
    combined = normalize_text("\n\n".join(f"{item['title']}\n{item['text']}" for item in sections))
    return combined, sections


def _read_pdf_text(path: Path) -> tuple[str, list[dict[str, str]]]:
    if not HAS_PYPDF or PdfReader is None:
        return _read_pdf_text_with_pymupdf(path)
    try:
        reader = PdfReader(str(path))
    except Exception:
        return _read_pdf_text_with_pymupdf(path)
    sections: list[dict[str, str]] = []
    pages: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        try:
            raw = page.extract_text() or ""
        except Exception:
            raw = ""
        text = normalize_text(raw)
        if not text:
            continue
        title = f"第 {index} 页"
        sections.append({"title": title, "text": text})
        pages.append(f"{title}\n{text}")
    combined = normalize_text("\n\n".join(pages))
    if len(combined) >= 120:
        return combined, sections
    pymupdf_text, pymupdf_sections = _read_pdf_text_with_pymupdf(path)
    if len(pymupdf_text) > len(combined):
        return pymupdf_text, pymupdf_sections
    return combined, sections


def _read_pdf_text_with_pymupdf(path: Path) -> tuple[str, list[dict[str, str]]]:
    if not HAS_PYMUPDF or fitz is None:
        return "", []
    sections: list[dict[str, str]] = []
    pages: list[str] = []
    try:
        document = fitz.open(str(path))
    except Exception:
        return "", []
    for index, page in enumerate(document, start=1):
        try:
            raw = page.get_text("text") or ""
        except Exception:
            raw = ""
        text = normalize_text(raw)
        if not text:
            continue
        title = f"第 {index} 页"
        sections.append({"title": title, "text": text})
        pages.append(f"{title}\n{text}")
    return normalize_text("\n\n".join(pages)), sections


def _pdf_page_count(path: Path) -> int | None:
    if not HAS_PYMUPDF or fitz is None:
        return None
    try:
        document = fitz.open(str(path))
        return int(len(document))
    except Exception:
        return None


def inspect_pdf_page_count(path: Path) -> int | None:
    return _pdf_page_count(path)


# 视觉 API 单图大小上限（base64 编码后 ≈ 原图 * 1.37）。
# 豆包视觉 API 推荐 ≤ 4MB 原图，留 buffer 设 3.5MB
_OCR_MAX_PAGE_BYTES = 3_500_000


def _compress_page_image(raw_png: bytes) -> tuple[bytes, str]:
    """大图压缩：PNG 超 _OCR_MAX_PAGE_BYTES 时转 JPEG（q=85），仍超则按比例缩放。

    返回 (压缩后 bytes, mime_type)。如压缩失败或 PIL 不可用，原样返回。
    """
    if len(raw_png) <= _OCR_MAX_PAGE_BYTES:
        return raw_png, "image/png"
    if not HAS_PILLOW or Image is None:
        return raw_png, "image/png"
    try:
        img = Image.open(BytesIO(raw_png))
        if img.mode in ("RGBA", "LA", "P"):
            img = img.convert("RGB")
        # 第一次尝试: JPEG q=85
        buf = BytesIO()
        img.save(buf, format="JPEG", quality=85, optimize=True)
        data = buf.getvalue()
        if len(data) <= _OCR_MAX_PAGE_BYTES:
            return data, "image/jpeg"
        # 仍超：按 sqrt 比例缩放
        scale = (_OCR_MAX_PAGE_BYTES / len(data)) ** 0.5
        new_w = max(800, int(img.width * scale))
        new_h = max(800, int(img.height * scale))
        img2 = img.resize((new_w, new_h), Image.LANCZOS)
        buf2 = BytesIO()
        img2.save(buf2, format="JPEG", quality=85, optimize=True)
        data2 = buf2.getvalue()
        if len(data2) <= _OCR_MAX_PAGE_BYTES:
            return data2, "image/jpeg"
        # 再缩一次（极端大图）
        img3 = img2.resize((max(600, int(new_w * 0.75)), max(600, int(new_h * 0.75))), Image.LANCZOS)
        buf3 = BytesIO()
        img3.save(buf3, format="JPEG", quality=80, optimize=True)
        return buf3.getvalue(), "image/jpeg"
    except Exception:
        return raw_png, "image/png"


def _render_pdf_pages_for_ai_ocr(
    path: Path,
    *,
    start_page: int = 1,
    max_pages: int = OCR_DEFAULT_BATCH_SIZE,
    zoom: float = 1.6,
) -> list[dict[str, object]]:
    if not HAS_PYMUPDF or fitz is None:
        return []
    try:
        document = fitz.open(str(path))
    except Exception:
        return []
    rendered: list[dict[str, object]] = []
    matrix = fitz.Matrix(zoom, zoom)
    start_index = max(0, int(start_page or 1) - 1)
    end_index = min(len(document), start_index + max(0, int(max_pages or 0)))
    for page_index in range(start_index, end_index):
        try:
            page = document[page_index]
            pixmap = page.get_pixmap(matrix=matrix, alpha=False)
            image_bytes = pixmap.tobytes("png")
        except Exception:
            continue
        if not image_bytes:
            continue
        # 高分辨率扫描件单页可能 >4MB，超视觉 API 上限：压缩后再上传
        compressed_bytes, mime_type = _compress_page_image(image_bytes)
        rendered.append(
            {
                "pageNumber": page_index + 1,
                "mimeType": mime_type,
                "imageBase64": base64.b64encode(compressed_bytes).decode("ascii"),
            }
        )
    return rendered


def _call_visual_markdown_ocr(
    ai_service: Any,
    *,
    title: str,
    image_base64: str,
    mime_type: str,
    page_number: int | None = None,
    source_kind: str = "视觉资料",
) -> str:
    if hasattr(ai_service, "generate_visual_markdown"):
        return str(
            ai_service.generate_visual_markdown(
                title=title,
                page_number=page_number,
                image_base64=image_base64,
                mime_type=mime_type,
                source_kind=source_kind,
            )
            or ""
        )
    if hasattr(ai_service, "generate_pdf_page_markdown"):
        return str(
            ai_service.generate_pdf_page_markdown(
                title=title,
                page_number=page_number or 1,
                image_base64=image_base64,
                mime_type=mime_type,
            )
            or ""
        )
    return ""


def _read_visual_markdown_with_ai_ocr_result(
    image_items: list[dict[str, object]],
    *,
    title: str,
    ai_service: Any | None,
    section_prefix: str,
    source_kind: str,
    min_chars: int = 24,
) -> VisualOcrResult:
    if ai_service is None or not (
        hasattr(ai_service, "generate_visual_markdown") or hasattr(ai_service, "generate_pdf_page_markdown")
    ):
        return VisualOcrResult(text="", sections=[])
    if not image_items:
        return VisualOcrResult(text="", sections=[])
    sections: list[dict[str, str]] = []
    visual_texts: list[str] = []
    attempted_pages = len(image_items)
    for index, item in enumerate(image_items, start=1):
        page_number = int(item.get("pageNumber") or index)
        try:
            markdown = _call_visual_markdown_ocr(
                ai_service,
                title=title,
                page_number=page_number,
                image_base64=str(item.get("imageBase64") or ""),
                mime_type=str(item.get("mimeType") or "image/png"),
                source_kind=source_kind,
            )
        except Exception:
            markdown = ""
        cleaned = _clean_ingested_text(str(markdown or ""))
        if not cleaned:
            continue
        section_title = f"第 {page_number} 页" if section_prefix == "第" else f"{section_prefix} {page_number}".strip()
        sections.append({"title": section_title, "text": cleaned})
        visual_texts.append(f"## {section_title}\n\n{cleaned}")
    combined = normalize_text("\n\n".join(visual_texts))
    if len(combined) < min_chars:
        return VisualOcrResult(text="", sections=[], attempted_pages=attempted_pages, failed_pages=attempted_pages)
    succeeded_pages = len(sections)
    return VisualOcrResult(
        text=combined,
        sections=sections,
        attempted_pages=attempted_pages,
        succeeded_pages=succeeded_pages,
        failed_pages=max(0, attempted_pages - succeeded_pages),
    )


def _read_visual_markdown_with_ai_ocr(
    image_items: list[dict[str, object]],
    *,
    title: str,
    ai_service: Any | None,
    section_prefix: str,
    source_kind: str,
    min_chars: int = 24,
) -> tuple[str, list[dict[str, str]]]:
    result = _read_visual_markdown_with_ai_ocr_result(
        image_items,
        title=title,
        ai_service=ai_service,
        section_prefix=section_prefix,
        source_kind=source_kind,
        min_chars=min_chars,
    )
    return result.text, result.sections


def _read_pdf_markdown_with_ai_ocr(
    path: Path,
    *,
    title: str,
    ai_service: Any | None,
    start_page: int = 1,
    max_pages: int = OCR_DEFAULT_MAX_PAGES,
    batch_size: int = OCR_DEFAULT_BATCH_SIZE,
    continue_to_end: bool = False,
    ocr_progress_callback: Callable[[dict[str, object]], None] | None = None,
) -> tuple[str, list[dict[str, str]], ExtractionMetadata]:
    total_pages = _pdf_page_count(path)
    normalized_max_pages = max(1, int(max_pages or OCR_DEFAULT_MAX_PAGES))
    normalized_batch_size = max(1, int(batch_size or OCR_DEFAULT_BATCH_SIZE))
    normalized_start_page = max(1, int(start_page or 1))
    metadata = ExtractionMetadata(
        parse_status="failed",
        parse_error="未能解析出可用正文",
        failure_type="empty_pdf",
        total_pages=total_pages,
        ocr_page_limit=normalized_max_pages,
        ocr_batch_size=normalized_batch_size,
    )
    if total_pages is not None and normalized_start_page > total_pages:
        return "", [], metadata

    sections: list[dict[str, str]] = []
    visual_texts: list[str] = []
    attempted_pages = 0
    succeeded_pages = 0
    failed_pages = 0
    processed_end_page = normalized_start_page - 1
    current_window_start = normalized_start_page

    while True:
        if total_pages is not None and current_window_start > total_pages:
            break
        remaining_pages = (total_pages - current_window_start + 1) if total_pages is not None else normalized_max_pages
        current_window_limit = min(max(0, remaining_pages), normalized_max_pages)
        if current_window_limit <= 0:
            break
        current_window_end = current_window_start + current_window_limit - 1

        for batch_start in range(current_window_start, current_window_end + 1, normalized_batch_size):
            current_batch_size = min(normalized_batch_size, current_window_end - batch_start + 1)
            page_images = _render_pdf_pages_for_ai_ocr(path, start_page=batch_start, max_pages=current_batch_size)
            result = _read_visual_markdown_with_ai_ocr_result(
                page_images,
                title=title or path.name,
                ai_service=ai_service,
                section_prefix="第",
                source_kind="PDF 页面",
                min_chars=24,
            )
            attempted_pages += current_batch_size
            succeeded_pages += result.succeeded_pages
            failed_pages += max(0, current_batch_size - result.succeeded_pages)
            if result.sections:
                sections.extend(result.sections)
                visual_texts.append(result.text)
            if ocr_progress_callback is not None:
                try:
                    ocr_progress_callback(
                        {
                            "totalPages": total_pages,
                            "pageLimit": total_pages or current_window_end,
                            "attemptedPages": attempted_pages,
                            "succeededPages": succeeded_pages,
                            "failedPages": failed_pages,
                            "currentPageRange": [batch_start, batch_start + current_batch_size - 1],
                            "ocrWindowRange": [current_window_start, current_window_end],
                            "continueToEnd": bool(continue_to_end),
                        }
                    )
                except Exception:
                    pass

        processed_end_page = current_window_end
        if not continue_to_end or total_pages is None or current_window_end >= total_pages:
            break
        current_window_start = current_window_end + 1

    combined = normalize_text("\n\n".join(text for text in visual_texts if text))
    metadata.attempted_pages = attempted_pages
    metadata.succeeded_pages = succeeded_pages
    metadata.failed_pages = failed_pages
    if not combined or len(combined) < 120:
        metadata.parse_error = "未能解析出可用正文"
        metadata.failure_type = "empty_pdf"
        return "", [], metadata

    truncated_by_limit = bool(total_pages and processed_end_page < total_pages)
    has_failed_pages = failed_pages > 0
    metadata.partial = truncated_by_limit or has_failed_pages
    metadata.parse_status = "partial_ready" if metadata.partial else "ready"
    metadata.failure_type = None if metadata.parse_status == "ready" else "ocr_required"
    if metadata.partial:
        tail = f"，PDF 共 {total_pages} 页" if total_pages else ""
        metadata.parse_error = f"OCR 部分完成：成功 {succeeded_pages}/{attempted_pages} 页{tail}，当前已处理到第 {processed_end_page} 页。"
    else:
        metadata.parse_error = None
    return combined, sections, metadata


def _render_image_for_ai_ocr(
    path: Path,
    *,
    max_side: int = 2200,
) -> list[dict[str, object]]:
    if not HAS_PILLOW or Image is None or ImageOps is None:
        return []
    try:
        image = Image.open(path)
        image = ImageOps.exif_transpose(image)
        image = image.convert("RGB")
    except Exception:
        return []
    width, height = image.size
    longest = max(width, height)
    if longest > max_side:
        scale = max_side / float(longest)
        image = image.resize((max(1, int(width * scale)), max(1, int(height * scale))))
    buffer = BytesIO()
    try:
        image.save(buffer, format="PNG", optimize=True)
    except Exception:
        return []
    image_bytes = buffer.getvalue()
    if not image_bytes:
        return []
    return [
        {
            "pageNumber": 1,
            "mimeType": "image/png",
            "imageBase64": base64.b64encode(image_bytes).decode("ascii"),
        }
    ]


def _read_image_markdown_with_ai_ocr(path: Path, *, title: str, ai_service: Any | None) -> tuple[str, list[dict[str, str]]]:
    image_items = _render_image_for_ai_ocr(path)
    return _read_visual_markdown_with_ai_ocr(
        image_items,
        title=title or path.name,
        ai_service=ai_service,
        section_prefix="图片",
        source_kind="图片资料",
        min_chars=24,
    )


def _build_markdown_sections(text: str) -> list[dict[str, str]]:
    sections: list[dict[str, str]] = []
    current_title = "正文"
    buffer: list[str] = []
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            buffer.append("")
            continue
        if line.startswith("#"):
            content = normalize_text("\n".join(buffer))
            if content:
                sections.append({"title": current_title, "text": content})
            current_title = re.sub(r"^#+\s*", "", line) or "正文"
            buffer = []
            continue
        buffer.append(line)
    content = normalize_text("\n".join(buffer))
    if content:
        sections.append({"title": current_title, "text": content})
    return sections or [{"title": "正文", "text": normalize_text(text)}]


def _build_generic_sections(text: str) -> list[dict[str, str]]:
    normalized = normalize_text(text)
    if not normalized:
        return []
    parts = [segment.strip() for segment in SECTION_BREAK_PATTERN.split(normalized) if segment.strip()]
    if len(parts) <= 1:
        return [{"title": "正文", "text": normalized}]
    sections = []
    for index, part in enumerate(parts, start=1):
        lead = part.splitlines()[0][:36]
        sections.append({"title": lead or f"段落 {index}", "text": part})
    return sections


def extract_document_with_metadata(
    path: Path,
    *,
    title: str | None = None,
    ai_service: Any | None = None,
    ocr_start_page: int = 1,
    ocr_max_pages: int = OCR_DEFAULT_MAX_PAGES,
    ocr_batch_size: int = OCR_DEFAULT_BATCH_SIZE,
    ocr_continue_to_end: bool = False,
    force_ocr: bool = False,
    ocr_progress_callback: Callable[[dict[str, object]], None] | None = None,
) -> ExtractedDocument:
    suffix = path.suffix.lower()
    if suffix in {".md"}:
        text = _read_plain_text(path)
        cleaned = _clean_ingested_text(text)
        sections = _build_markdown_sections(cleaned)
        return ExtractedDocument(cleaned, sections, ExtractionMetadata(parse_status="ready" if cleaned and sections else "failed"))
    if suffix in {".txt", ".json", ".csv"}:
        text = _read_plain_text(path)
        cleaned = _clean_ingested_text(text)
        sections = _build_generic_sections(cleaned)
        return ExtractedDocument(cleaned, sections, ExtractionMetadata(parse_status="ready" if cleaned and sections else "failed"))
    if suffix == ".docx":
        try:
            text, sections = _read_docx_text(path)
            cleaned_sections = [
                {"title": section.get("title", "正文"), "text": _clean_ingested_text(str(section.get("text", "")))}
                for section in sections
            ]
            cleaned_sections = [section for section in cleaned_sections if section["text"]]
            cleaned = _clean_ingested_text(text)
            final_sections = cleaned_sections or _build_generic_sections(cleaned)
            return ExtractedDocument(cleaned, final_sections, ExtractionMetadata(parse_status="ready" if cleaned and final_sections else "failed"))
        except Exception:
            text = _archive_xml_text(path)
            cleaned = _clean_ingested_text(text)
            sections = _build_generic_sections(cleaned)
            return ExtractedDocument(cleaned, sections, ExtractionMetadata(parse_status="ready" if cleaned and sections else "failed"))
    if suffix == ".pdf":
        text, sections = ("", []) if force_ocr else _read_pdf_text(path)
        if text:
            cleaned_sections = [
                {"title": section.get("title", "正文"), "text": _clean_ingested_text(str(section.get("text", "")))}
                for section in sections
            ]
            cleaned_sections = [section for section in cleaned_sections if section["text"]]
            cleaned = _clean_ingested_text(text)
            final_sections = cleaned_sections or _build_generic_sections(cleaned)
            return ExtractedDocument(cleaned, final_sections, ExtractionMetadata(parse_status="ready" if cleaned and final_sections else "failed"))
        ai_text, ai_sections, metadata = _read_pdf_markdown_with_ai_ocr(
            path,
            title=title or path.name,
            ai_service=ai_service,
            start_page=ocr_start_page,
            max_pages=ocr_max_pages,
            batch_size=ocr_batch_size,
            continue_to_end=ocr_continue_to_end,
            ocr_progress_callback=ocr_progress_callback,
        )
        if ai_text:
            cleaned_sections = [
                {"title": section.get("title", "正文"), "text": _clean_ingested_text(str(section.get("text", "")))}
                for section in ai_sections
            ]
            cleaned_sections = [section for section in cleaned_sections if section["text"]]
            cleaned = _clean_ingested_text(ai_text)
            final_sections = cleaned_sections or _build_markdown_sections(cleaned)
            if not final_sections:
                metadata.parse_status = "failed"
                metadata.parse_error = "未能解析出可用正文"
                metadata.failure_type = "empty_pdf"
            return ExtractedDocument(cleaned, final_sections, metadata)
        return ExtractedDocument("", [], metadata)
    if suffix in IMAGE_EXTENSIONS:
        ai_text, ai_sections = _read_image_markdown_with_ai_ocr(path, title=title or path.name, ai_service=ai_service)
        if ai_text:
            cleaned_sections = [
                {"title": section.get("title", "正文"), "text": _clean_ingested_text(str(section.get("text", "")))}
                for section in ai_sections
            ]
            cleaned_sections = [section for section in cleaned_sections if section["text"]]
            cleaned = _clean_ingested_text(ai_text)
            final_sections = cleaned_sections or _build_markdown_sections(cleaned)
            return ExtractedDocument(cleaned, final_sections, ExtractionMetadata(parse_status="ready" if cleaned and final_sections else "failed"))
        return ExtractedDocument("", [], ExtractionMetadata(parse_status="failed", parse_error="未能解析出可用正文", failure_type="empty_text"))
    if suffix in ARCHIVE_XML_EXTENSIONS:
        # Phase 0：xlsx / pptx 走结构化解析，保留 sheet/slide 边界 + 表格 +
        # 演讲者备注。失败则 fallback 到旧的 _archive_xml_text 路径，零回归风险。
        try:
            from app.services.structured_table_parser import (
                StructuredParseError,
                parse_pptx_structured,
                parse_xlsx_structured,
            )

            if suffix == ".xlsx":
                parsed_sheets = parse_xlsx_structured(path)
                sections = [
                    {"title": f"Sheet · {s.sheet_name}", "text": s.markdown}
                    for s in parsed_sheets
                ]
                full_text = "\n\n".join(s["text"] for s in sections)
                cleaned = _clean_ingested_text(full_text)
                cleaned_sections = [
                    {"title": s["title"], "text": _clean_ingested_text(s["text"])}
                    for s in sections
                ]
                cleaned_sections = [s for s in cleaned_sections if s["text"]]
                if cleaned and cleaned_sections:
                    return ExtractedDocument(
                        cleaned,
                        cleaned_sections,
                        ExtractionMetadata(parse_status="ready"),
                        structured_sheets=list(parsed_sheets),
                    )
            elif suffix == ".pptx":
                parsed_slides = parse_pptx_structured(path)
                sections = [
                    {"title": f"Slide {s.slide_no} · {s.title or '无标题'}", "text": s.markdown}
                    for s in parsed_slides
                ]
                full_text = "\n\n".join(s["text"] for s in sections)
                cleaned = _clean_ingested_text(full_text)
                cleaned_sections = [
                    {"title": s["title"], "text": _clean_ingested_text(s["text"])}
                    for s in sections
                ]
                cleaned_sections = [s for s in cleaned_sections if s["text"]]
                if cleaned and cleaned_sections:
                    return ExtractedDocument(
                        cleaned,
                        cleaned_sections,
                        ExtractionMetadata(parse_status="ready"),
                    )
        except StructuredParseError as exc:
            logger.warning(
                "structured parser failed for %s, fallback to _archive_xml_text: %s",
                path.name,
                exc,
            )
        except Exception:
            logger.exception("unexpected error in structured parser for %s", path.name)
        # Fallback：旧路径（保底）。文件根本不是 zip 时这里也会抛，再加一层兜底。
        try:
            text = _archive_xml_text(path)
        except Exception as exc:
            return ExtractedDocument(
                "",
                [],
                ExtractionMetadata(
                    parse_status="failed",
                    parse_error=f"无法解析 {suffix} 文件：{exc}",
                    failure_type="unsupported_format",
                ),
            )
        cleaned = _clean_ingested_text(text)
        sections = _build_generic_sections(cleaned)
        return ExtractedDocument(cleaned, sections, ExtractionMetadata(parse_status="ready" if cleaned and sections else "failed"))
    return ExtractedDocument("", [], ExtractionMetadata(parse_status="failed", parse_error="不支持的文件格式", failure_type="unsupported_format"))


def extract_document(path: Path, *, title: str | None = None, ai_service: Any | None = None) -> tuple[str, list[dict[str, str]]]:
    extracted = extract_document_with_metadata(path, title=title, ai_service=ai_service)
    return extracted.text, extracted.sections


def build_chunks(section_text: str, section_title: str) -> list[dict[str, Any]]:
    normalized = normalize_text(section_text)
    if not normalized:
        return []
    paragraphs = [segment.strip() for segment in SECTION_BREAK_PATTERN.split(normalized) if segment.strip()]
    if not paragraphs:
        paragraphs = [normalized]
    chunks: list[dict[str, Any]] = []
    current = ""
    for paragraph in paragraphs:
        next_candidate = f"{current}\n\n{paragraph}".strip() if current else paragraph
        if len(next_candidate) <= CHUNK_TARGET_CHARS:
            current = next_candidate
            continue
        if current:
            chunks.append(
                {
                    "sectionLabel": section_title or "正文",
                    "content": current,
                    "tokenText": " ".join(tokenize(current)),
                    "charCount": len(current),
                }
            )
        if len(paragraph) <= CHUNK_TARGET_CHARS:
            current = paragraph
            continue
        start = 0
        while start < len(paragraph):
            end = min(len(paragraph), start + CHUNK_TARGET_CHARS)
            piece = paragraph[start:end].strip()
            if piece:
                chunks.append(
                    {
                        "sectionLabel": section_title or "正文",
                        "content": piece,
                        "tokenText": " ".join(tokenize(piece)),
                        "charCount": len(piece),
                    }
                )
            if end >= len(paragraph):
                break
            start = max(0, end - CHUNK_OVERLAP_CHARS)
        current = ""
    if current:
        chunks.append(
            {
                "sectionLabel": section_title or "正文",
                "content": current,
                "tokenText": " ".join(tokenize(current)),
                "charCount": len(current),
            }
        )
    return chunks


def _copy_into_workspace(data_dir: Path, client_id: str, source_path: Path, category: str, document_id: str) -> Path:
    workspace_root = client_workspace_root(data_dir, client_id) / category
    workspace_root.mkdir(parents=True, exist_ok=True)
    safe_name = safe_filename(source_path.name)
    target = workspace_root / safe_name
    if source_path.resolve() == target.resolve():
        return target
    if target.exists():
        stem = safe_filename(source_path.stem)
        suffix = source_path.suffix.lower()
        target = workspace_root / f"{stem}__{document_id[-6:]}{suffix}"
    source_parts = {part.lower() for part in source_path.parts}
    if source_path.exists() and "_imports" in source_parts:
        shutil.move(str(source_path), target)
    else:
        shutil.copy2(source_path, target)
    return target


def _compat_card_path(data_dir: Path, client_id: str, document_id: str, title: str) -> Path:
    meta_root = client_workspace_root(data_dir, client_id) / "_v2_meta" / "cards"
    meta_root.mkdir(parents=True, exist_ok=True)
    stem = safe_filename(Path(title).stem or document_id)
    return meta_root / f"{document_id}_{stem}.md"


def _markdown_derivative_path(data_dir: Path, client_id: str, document_id: str, title: str) -> Path:
    meta_root = client_workspace_root(data_dir, client_id) / "_v2_meta" / "markdown"
    meta_root.mkdir(parents=True, exist_ok=True)
    stem = safe_filename(Path(title).stem or document_id)
    return meta_root / f"{document_id}_{stem}.md"


def _write_compat_card_markdown(
    data_dir: Path,
    client_id: str,
    document_id: str,
    title: str,
    category: str,
    preview_text: str,
    sections: list[dict[str, str]],
) -> Path:
    target = _compat_card_path(data_dir, client_id, document_id, title)
    lines = [
        f"# {safe_filename(title)}",
        "",
        f"- V2 目录归类：{category}",
        "- 用途：兼容旧字段，不参与 V2 正式检索",
        "",
        "## 解析摘要",
        preview_text,
    ]
    if sections:
        lines.extend(["", "## 章节索引"])
        for index, section in enumerate(sections[:12], start=1):
            section_title = normalize_text(section.get("title", "")) or f"正文 {index}"
            lines.append(f"{index}. {section_title}")
    target.write_text("\n".join(lines).strip() + "\n", encoding="utf-8")
    return target


def _write_markdown_derivative(
    data_dir: Path,
    client_id: str,
    document_id: str,
    title: str,
    sections: list[dict[str, str]],
    fallback_text: str,
) -> tuple[Path, str]:
    target = _markdown_derivative_path(data_dir, client_id, document_id, title)
    lines = [f"# {safe_filename(title)}", ""]
    if sections:
        for index, section in enumerate(sections, start=1):
            section_title = normalize_text(section.get("title", "")) or f"正文 {index}"
            section_text = normalize_text(section.get("text", ""))
            if not section_text:
                continue
            heading = "##" if index == 1 else "##"
            lines.extend([f"{heading} {section_title}", "", section_text, ""])
    else:
        body = normalize_text(fallback_text)
        if body:
            lines.extend(["## 正文", "", body, ""])
    markdown_content = "\n".join(lines).strip() + "\n"
    target.write_text(markdown_content, encoding="utf-8")
    return target, markdown_content


def _content_hash(text: str, path: Path) -> str:
    hasher = hashlib.sha256()
    try:
        with path.open("rb") as handle:
            for chunk in iter(lambda: handle.read(1024 * 1024), b""):
                hasher.update(chunk)
    except Exception:
        hasher.update(text.encode("utf-8"))
    return hasher.hexdigest()


def hash_file_bytes(path: Path) -> str | None:
    """迭代 2 F3：在导入闸门期算文件 SHA-256，用于跨路径/同路径变更检测。

    返回 None 表示读不出来（路径不存在/权限不足），由调用方决定降级策略。
    """
    hasher = hashlib.sha256()
    try:
        with path.open("rb") as handle:
            for chunk in iter(lambda: handle.read(1024 * 1024), b""):
                hasher.update(chunk)
    except Exception:
        return None
    return hasher.hexdigest()


def _remove_existing_v2_rows(db: Database, v2_document_id: str) -> None:
    db.execute("DELETE FROM v2_chunks WHERE v2_document_id = ?", (v2_document_id,))
    db.execute("DELETE FROM v2_sections WHERE v2_document_id = ?", (v2_document_id,))


def _sync_legacy_knowledge_document(
    db: Database,
    *,
    client_id: str,
    import_id: str | None,
    document_id: str,
    source_path: Path,
    original_source_path: Path,
    managed_path: Path,
    kind: str,
    primary_category: str,
    secondary_category: str,
    human_folder_category: str | None = None,
    confidence: float,
    parse_status: str,
    content_hash: str,
    created_at: str,
) -> str:
    existing = db.fetchone(
        "SELECT id FROM knowledge_documents WHERE document_id = ? ORDER BY updated_at DESC LIMIT 1",
        (document_id,),
    )
    knowledge_document_id = str(existing["id"]) if existing and existing["id"] else f"kd_{document_id}"
    doc_uid = f"v2uid_{document_id}"
    original_path = str(original_source_path)
    import_source_path = str(original_source_path)
    current_human_path = str(managed_path)
    display_folder_category = human_folder_category or primary_category
    reclassified_at = created_at
    reclass_reason = f"{V2_PIPELINE_VERSION} 同步的兼容知识文档记录"
    reclass_confidence = confidence
    normalized_path = str(managed_path)
    vector_status = "chunk_indexed" if parse_status in SEARCHABLE_PARSE_STATUSES else "needs_review"
    needs_review = 0 if parse_status == "ready" else 1
    dedup_status = "unique"
    binary_hash = content_hash
    normalized_hash = content_hash
    updated_at = now_iso()
    if existing:
        db.execute(
            """
            UPDATE knowledge_documents
            SET client_id = ?, import_batch_id = ?, document_id = ?, doc_uid = ?, original_path = ?, import_source_path = ?,
                current_human_path = ?, human_folder_category = ?, reclassified_at = ?, reclass_reason = ?, reclass_confidence = ?,
                normalized_path = ?, kind = ?, primary_category = ?, secondary_category = ?, classification_confidence = ?,
                needs_review = ?, dedup_status = ?, vector_status = ?, binary_hash = ?, normalized_hash = ?, created_at = ?, updated_at = ?
            WHERE id = ?
            """,
            (
                client_id,
                import_id,
                document_id,
                doc_uid,
                original_path,
                import_source_path,
                current_human_path,
                display_folder_category,
                reclassified_at,
                reclass_reason,
                reclass_confidence,
                normalized_path,
                kind,
                primary_category,
                secondary_category,
                confidence,
                needs_review,
                dedup_status,
                vector_status,
                binary_hash,
                normalized_hash,
                created_at,
                updated_at,
                knowledge_document_id,
            ),
        )
    else:
        db.execute(
            """
            INSERT INTO knowledge_documents(
                id, client_id, import_batch_id, document_id, doc_uid, original_path, import_source_path, current_human_path,
                human_folder_category, reclassified_at, reclass_reason, reclass_confidence, normalized_path, kind,
                primary_category, secondary_category, classification_confidence, needs_review, deep_read,
                last_hit_question, dedup_status, vector_status, version, binary_hash, normalized_hash, created_at, updated_at
            )
            VALUES(?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 0, NULL, ?, ?, 1, ?, ?, ?, ?)
            """,
            (
                knowledge_document_id,
                client_id,
                import_id,
                document_id,
                doc_uid,
                original_path,
                import_source_path,
                current_human_path,
                display_folder_category,
                reclassified_at,
                reclass_reason,
                reclass_confidence,
                normalized_path,
                kind,
                primary_category,
                secondary_category,
                confidence,
                needs_review,
                dedup_status,
                vector_status,
                binary_hash,
                normalized_hash,
                created_at,
                updated_at,
            ),
        )
    existing_reclass = db.fetchone(
        """
        SELECT id
        FROM file_reclass_events
        WHERE knowledge_document_id = ? AND from_path = ? AND to_path = ? AND to_category = ?
        ORDER BY created_at DESC
        LIMIT 1
        """,
        (
            knowledge_document_id,
            original_path,
            current_human_path,
            display_folder_category,
        ),
    )
    if not existing_reclass:
        append_file_reclass_event(
            db,
            knowledge_document_id=knowledge_document_id,
            from_path=original_path,
            to_path=current_human_path,
            from_category=None,
            to_category=display_folder_category,
            reason=reclass_reason,
            confidence=reclass_confidence,
            created_at=created_at,
        )
    return knowledge_document_id


def _folder_counts(db: Database, client_id: str) -> dict[str, int]:
    rows = db.fetchall(
        """
        SELECT
            CASE
                WHEN material_layer = 'background' THEN '战略陪伴'
                ELSE COALESCE(NULLIF(visible_category, ''), ?)
            END AS category,
            COUNT(1) AS count
        FROM v2_documents
        WHERE client_id = ?
          AND parse_status IN ('ready', 'partial_ready')
        GROUP BY
            CASE
                WHEN material_layer = 'background' THEN '战略陪伴'
                ELSE COALESCE(NULLIF(visible_category, ''), ?)
            END
        """,
        (DEFAULT_INBOX_LABEL, client_id, DEFAULT_INBOX_LABEL),
    )
    return {str(row["category"]): int(row["count"]) for row in rows}


def refresh_client_folder_counts(db: Database, client_id: str) -> None:
    counts = _folder_counts(db, client_id)
    timestamp = now_iso()
    rows = db.fetchall("SELECT label FROM client_folders WHERE client_id = ?", (client_id,))
    for row in rows:
        category = str(row["label"])
        db.execute(
            "UPDATE client_folders SET file_count = ?, last_scanned_at = ? WHERE client_id = ? AND label = ?",
            (counts.get(category, 0), timestamp, client_id, category),
        )


def ingest_document_knowledge(
    db: Database,
    *,
    data_dir: Path,
    client_id: str,
    import_id: str | None,
    document_id: str,
    source_path: Path,
    original_source_path: Path | None,
    title: str,
    kind: str,
    source: str,
    fallback_excerpt: str,
    created_at: str,
    ai_service: Any | None = None,
    ocr_start_page: int = 1,
    ocr_max_pages: int = OCR_DEFAULT_MAX_PAGES,
    ocr_batch_size: int = OCR_DEFAULT_BATCH_SIZE,
    ocr_continue_to_end: bool = False,
    force_ocr: bool = False,
    ocr_progress_callback: Callable[[dict[str, object]], None] | None = None,
) -> dict[str, Any]:
    # 早期 family_id dedup：title 经 _normalized_family_stem 后能识别 xxx(1)/副本xxx/xxx_1 等语义重复。
    # 已存在同 family 的 raw_file → 当前文件直接进垃圾桶，跳过 ingest，调用方插入的 documents 行也删掉。
    candidate_family_id = derive_document_family_id(
        title=title, content_hash="", origin_type="file_import", origin_id=document_id,
    )
    if candidate_family_id and ":" in candidate_family_id:
        stem_part = candidate_family_id.split(":", 1)[1].strip()
        if stem_part and stem_part != "unknown" and len(stem_part) >= 3:
            existing = db.fetchone(
                """
                SELECT id, path FROM documents
                WHERE client_id = ? AND document_family_id = ?
                  AND canonical_kind = 'raw_file' AND id != ?
                ORDER BY created_at DESC LIMIT 1
                """,
                (client_id, candidate_family_id, document_id),
            )
            if existing:
                from app.services import trash_can as _trash_can
                _trash_can.trash_file(
                    db, data_dir,
                    source_path=source_path,
                    client_id=client_id,
                    original_document_id=document_id,
                    original_title=title,
                    reason="dedup_merge",
                )
                db.execute("DELETE FROM documents WHERE id = ?", (document_id,))
                return {
                    "knowledge_document_id": "",
                    "legacy_knowledge_document_id": "",
                    "title": title,
                    "current_human_path": str(existing["path"] or ""),
                    "compat_card_path": "",
                    "markdown_path": "",
                    "primary_category": UNIFIED_FOLDER_LABEL,
                    "secondary_category": "",
                    "short_summary": "",
                    "summary": "",
                    "classification_confidence": 1.0,
                    "material_layer": "",
                    "needs_review": False,
                    "dedup_skipped": True,
                    "dedup_kept_document_id": str(existing["id"]),
                }

    extracted = extract_document_with_metadata(
        source_path,
        title=title,
        ai_service=ai_service,
        ocr_start_page=ocr_start_page,
        ocr_max_pages=ocr_max_pages,
        ocr_batch_size=ocr_batch_size,
        ocr_continue_to_end=ocr_continue_to_end,
        force_ocr=force_ocr,
        ocr_progress_callback=ocr_progress_callback,
    )
    text, sections = extracted.text, extracted.sections
    extraction_metadata = extracted.metadata
    preview_text = build_excerpt(text or fallback_excerpt, title)
    primary_category, secondary_category, confidence = detect_category(title, text or preview_text)
    material_layer, primary_category, secondary_category, confidence = detect_material_profile(
        title,
        text or preview_text,
        primary_category,
        secondary_category,
        confidence,
    )
    # 统一文件夹策略：所有 raw_file 落到「资料库」。primary_category / confidence
    # 仍然计算，作为 v2_documents.visible_category 等字段的元数据保留供检索用。
    display_category = UNIFIED_FOLDER_LABEL
    managed_path = _copy_into_workspace(data_dir, client_id, source_path, display_category, document_id)
    compat_card_path = _write_compat_card_markdown(
        data_dir,
        client_id,
        document_id,
        title,
        primary_category,
        preview_text,
        sections,
    )
    markdown_path, markdown_content = _write_markdown_derivative(
        data_dir,
        client_id,
        document_id,
        title,
        sections,
        text or preview_text or fallback_excerpt,
    )
    v2_document_id = f"v2doc_{document_id}"
    content_hash = _content_hash(text or preview_text or title, source_path)
    document_family_id = derive_document_family_id(
        title=title,
        content_hash=content_hash,
        origin_type="file_import",
        origin_id=document_id,
    )
    canonical_kind = "raw_file"
    origin_type = "file_import"
    origin_id = document_id
    is_searchable = 0 if _is_ingestion_noise_document(title, text or preview_text) else 1
    parse_status = extraction_metadata.parse_status if text and sections else "failed"
    if parse_status not in SEARCHABLE_PARSE_STATUSES:
        parse_status = "ready" if text and sections else "failed"
    parse_error = None if parse_status == "ready" else (extraction_metadata.parse_error or "未能解析出可用正文")
    doc_tokens = tokenize(f"{title}\n{primary_category}\n{text[:2400] if text else preview_text}")
    heading_tokens = tokenize(" ".join(section["title"] for section in sections[:24]))
    doc_index_text = " ".join(dict.fromkeys([*doc_tokens, *heading_tokens]))
    section_count = len(sections)
    chunk_count = 0
    legacy_knowledge_document_id = _sync_legacy_knowledge_document(
        db,
        client_id=client_id,
        import_id=import_id,
        document_id=document_id,
        source_path=source_path,
        original_source_path=original_source_path or source_path,
        managed_path=managed_path,
        kind=kind,
        primary_category=primary_category,
        secondary_category=secondary_category,
        human_folder_category=display_category,
        confidence=confidence,
        parse_status=parse_status,
        content_hash=content_hash,
        created_at=created_at,
    )

    _remove_existing_v2_rows(db, v2_document_id)
    existing = db.fetchone("SELECT id FROM v2_documents WHERE id = ?", (v2_document_id,))
    payload = (
        v2_document_id,
        client_id,
        document_id,
        str(source_path),
        str(managed_path),
        str(markdown_path),
        safe_filename(title),
        kind,
        material_layer,
        display_category,
        secondary_category,
        parse_status,
        parse_error,
        preview_text,
        doc_index_text,
        content_hash,
        confidence,
        document_family_id,
        canonical_kind,
        origin_type,
        origin_id,
        is_searchable,
        created_at,
        now_iso(),
    )
    if existing:
        db.execute(
            """
            UPDATE v2_documents
            SET client_id = ?, document_id = ?, original_path = ?, managed_path = ?, markdown_path = ?, file_name = ?, kind = ?,
                material_layer = ?, visible_category = ?, secondary_category = ?, parse_status = ?, parse_error = ?,
                preview_text = ?, doc_index_text = ?, content_hash = ?, classification_confidence = ?,
                document_family_id = ?, canonical_kind = ?, origin_type = ?, origin_id = ?, is_searchable = ?,
                imported_at = ?, updated_at = ?
            WHERE id = ?
            """,
            payload[1:] + (v2_document_id,),
        )
    else:
        db.execute(
            """
            INSERT INTO v2_documents(
                id, client_id, document_id, original_path, managed_path, markdown_path, file_name, kind, material_layer, visible_category,
                secondary_category, parse_status, parse_error, preview_text, doc_index_text, content_hash,
                classification_confidence, document_family_id, canonical_kind, origin_type, origin_id, is_searchable, imported_at, updated_at
            )
            VALUES(?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            payload,
        )

    db.execute(
        "UPDATE v2_documents SET markdown_content = ? WHERE id = ?",
        (markdown_content if parse_status in SEARCHABLE_PARSE_STATUSES else "", v2_document_id),
    )

    db.execute(
        """
        UPDATE documents
        SET title = ?, path = ?, original_source_path = ?, kind = ?, source = ?, excerpt = ?, tags_json = ?,
            document_family_id = ?, canonical_kind = ?, origin_type = ?, origin_id = ?, is_searchable = ?
        WHERE id = ?
        """,
        (
            safe_filename(title),
            str(managed_path),
            str(original_source_path or source_path),
            kind,
            source,
            preview_text,
            to_json([kind, primary_category, secondary_category, V2_PIPELINE_VERSION]),
            document_family_id,
            canonical_kind,
            origin_type,
            origin_id,
            is_searchable,
            document_id,
        ),
    )

    if parse_status in SEARCHABLE_PARSE_STATUSES:
        for section_index, section in enumerate(sections):
            section_text = normalize_text(section["text"])
            if not section_text:
                continue
            section_id = f"v2sec_{document_id}_{section_index}"
            section_title = section["title"][:120] if section["title"] else "正文"
            section_tokens = " ".join(tokenize(f"{section_title}\n{section_text[:2400]}"))
            db.execute(
                """
                INSERT INTO v2_sections(
                    id, v2_document_id, section_index, title, content, searchable_text, char_count, created_at
                )
                VALUES(?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (section_id, v2_document_id, section_index, section_title, section_text, section_tokens, len(section_text), created_at),
            )
            for chunk_index, chunk in enumerate(build_chunks(section_text, section_title)):
                chunk_id = f"v2chunk_{document_id}_{section_index}_{chunk_index}"
                db.execute(
                    """
                    INSERT INTO v2_chunks(
                        id, v2_document_id, v2_section_id, chunk_index, section_label, content, searchable_text, char_count, created_at
                    )
                    VALUES(?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """,
                    (
                        chunk_id,
                        v2_document_id,
                        section_id,
                        chunk_index,
                        chunk["sectionLabel"],
                        chunk["content"],
                        chunk["tokenText"],
                        chunk["charCount"],
                        created_at,
                    ),
                )
                chunk_count += 1
                # 迭代 2：规则层实体抽取（不阻塞、不调用 LLM）。
                # 失败只记日志不打断入库——下次 backfill 会补上。
                try:
                    from app.services.entity_extractor import extract_entities_from_chunk
                    from app.services.entity_store import persist_chunk_entities

                    # 不能复用外层 `extracted` 变量名 —— 外层 line 1808 的 `extracted`
                    # 是 extract_document_with_metadata 返回的对象，line 2105 还要读它的
                    # .structured_sheets；这里覆盖会让外层挂在 AttributeError。
                    chunk_entities = extract_entities_from_chunk(chunk["content"])
                    if chunk_entities:
                        persist_chunk_entities(
                            db.conn,
                            client_id=client_id,
                            v2_document_id=v2_document_id,
                            v2_chunk_id=chunk_id,
                            extracted=chunk_entities,
                            now=created_at,
                        )
                except Exception:
                    logger.exception(
                        "entity extraction failed for chunk %s (doc=%s)",
                        chunk_id,
                        document_id,
                    )
                # 迭代 4：chunk 语义分类（规则层，不阻塞）
                try:
                    from app.services.semantic_classifier import classify_chunk_semantic

                    semantic = classify_chunk_semantic(chunk["content"])
                    db.execute(
                        "UPDATE v2_chunks SET semantic_type = ?, semantic_confidence = ? WHERE id = ?",
                        (semantic.semantic_type, semantic.confidence, chunk_id),
                    )
                except Exception:
                    logger.exception(
                        "semantic classification failed for chunk %s (doc=%s)",
                        chunk_id,
                        document_id,
                    )
                # 迭代 5：关系三元组抽取（规则层）。
                # 依赖实体已经入库（同一 chunk 内 entity 在前面已 persist）。
                try:
                    from app.services.relation_extractor import extract_relations_from_chunk
                    from app.services.relation_store import persist_chunk_relations

                    relations = extract_relations_from_chunk(chunk["content"])
                    if relations:
                        persist_chunk_relations(
                            db.conn,
                            client_id=client_id,
                            v2_document_id=v2_document_id,
                            v2_chunk_id=chunk_id,
                            extracted=relations,
                            now=created_at,
                        )
                except Exception:
                    logger.exception(
                        "relation extraction failed for chunk %s (doc=%s)",
                        chunk_id,
                        document_id,
                    )
                # 迭代 6：原子事实 + 矛盾检测（同 client 同 subject+attribute
                # 不同 value → 自动 detect 并写入 fact_contradictions）
                try:
                    from app.services.contradiction_detector import persist_chunk_facts
                    from app.services.fact_extractor import extract_facts_from_chunk

                    facts = extract_facts_from_chunk(chunk["content"])
                    if facts:
                        persist_chunk_facts(
                            db.conn,
                            client_id=client_id,
                            v2_document_id=v2_document_id,
                            v2_chunk_id=chunk_id,
                            facts=facts,
                            now=created_at,
                        )
                except Exception:
                    logger.exception(
                        "fact extraction / contradiction detection failed for chunk %s (doc=%s)",
                        chunk_id,
                        document_id,
                    )

    db.execute(
        """
        UPDATE v2_documents
        SET section_count = ?, chunk_count = ?, updated_at = ?
        WHERE id = ?
        """,
        (section_count, chunk_count, now_iso(), v2_document_id),
    )
    # Phase 1：把 xlsx 的 ParsedSheet 写到 structured_tables 表
    if extracted.structured_sheets:
        try:
            from app.services.structured_table_store import upsert_table_from_parsed_sheet

            for sheet_index, parsed_sheet in enumerate(extracted.structured_sheets):
                upsert_table_from_parsed_sheet(
                    db.conn,
                    client_id=client_id,
                    v2_document_id=v2_document_id,
                    knowledge_document_id=None,  # 这一步 knowledge_documents 行还没建
                    sheet_index=sheet_index,
                    parsed=parsed_sheet,
                    now=created_at,
                )
        except Exception:
            logger.exception(
                "structured_tables upsert failed for %s (doc=%s)",
                title,
                document_id,
            )
    refresh_client_folder_counts(db, client_id)

    # Phase 2：根据文档 kind 自动派发本地推理任务（如 pptx → per-slide visual_ocr）
    # 不阻塞导入主流程；任何异常吞掉只打日志
    try:
        from app.services.task_runners.router import route_document_for_local_inference

        route_result = route_document_for_local_inference(db, v2_document_id)
        if route_result.get("enqueued"):
            logger.info(
                "[router] doc=%s kind=%s enqueued=%d",
                v2_document_id,
                route_result.get("kind"),
                route_result.get("enqueued"),
            )
    except Exception:
        logger.exception("router 派发失败 (doc=%s)", v2_document_id)

    # --- 即时写入 master_index，让文档上传后立即可被检索 ---
    try:
        from app.services.knowledge_base import upsert_master_index_record

        surrogate_id = f"ks_instant_{document_id}"
        searchable = doc_index_text or preview_text or title
        retrieval_text = preview_text or fallback_excerpt or ""
        # 先创建轻量 surrogate 记录（满足外键约束）
        existing_surrogate = db.fetchone("SELECT id FROM knowledge_surrogates WHERE id = ?", (surrogate_id,))
        if not existing_surrogate:
            db.execute(
                """INSERT OR IGNORE INTO knowledge_surrogates(
                    id, knowledge_document_id, client_id, source_type, title, folder_category,
                    surrogate_md_path, overview_summary, retrieval_summary, document_role,
                    core_questions_json, query_hints_json, distinct_findings_json, entities_json,
                    time_markers_json, source_links_json, created_at, updated_at
                ) VALUES(?, ?, ?, 'v2_instant', ?, ?, ?, ?, ?, ?, '[]', '[]', '[]', '[]', '[]', '[]', ?, ?)""",
                (surrogate_id, legacy_knowledge_document_id or v2_document_id, client_id, safe_filename(title), primary_category,
                 str(compat_card_path), retrieval_text[:300], retrieval_text,
                 material_layer or kind, created_at, created_at),
            )
        upsert_master_index_record(
            db,
            data_dir=data_dir,
            entry_id=v2_document_id,
            client_id=client_id,
            surrogate_id=surrogate_id,
            title=safe_filename(title),
            folder_category=primary_category,
            document_role=material_layer or kind,
            retrieval_summary=retrieval_text,
            searchable_text=searchable,
            source_path=str(managed_path),
            surrogate_md_path=str(compat_card_path),
            timestamp=created_at,
            sync_after=True,
        )
    except Exception:
        pass  # master_index 写入失败不应阻塞文档入库

    # ─────────────────────────────────────────────────────────────────────
    # Stage B：知识扇出（Karpathy LLM Wiki #3「一资料更新影响 10-15 页面」落地）
    # 每个扇出都包在 try/except 里，失败只 log 不阻塞主 ingest 流程。
    # ─────────────────────────────────────────────────────────────────────
    try:
        n_event_lines = fanout_document_to_event_lines(
            db,
            client_id=client_id,
            v2_document_id=v2_document_id,
            document_id=document_id,
            file_name=safe_filename(title),
            created_at=created_at,
        )
        if n_event_lines > 0:
            logger.info(
                "[fanout] document=%s 关联到 %d 条 event_line",
                v2_document_id,
                n_event_lines,
            )
    except Exception:
        logger.exception("fanout_document_to_event_lines 失败 (doc=%s)", v2_document_id)

    # Stage B 扇出 #3：strategic_thought 刷新标记（含信号门槛检测）
    try:
        flagged = fanout_document_to_strategic_thoughts(
            db,
            client_id=client_id,
            v2_document_id=v2_document_id,
            file_name=safe_filename(title),
        )
        if flagged:
            logger.info("[fanout] document=%s flagged strategic_thoughts.refresh_pending", v2_document_id)
    except Exception:
        logger.exception("fanout_document_to_strategic_thoughts 失败 (doc=%s)", v2_document_id)

    # Stage B 扇出 #4：矛盾 → judgment 影响标记（纯 SQL）
    try:
        n_judgments = fanout_contradictions_to_judgment_impact(
            db,
            client_id=client_id,
            v2_document_id=v2_document_id,
        )
        if n_judgments > 0:
            logger.info("[fanout] document=%s 触发 %d 个 judgment 需要重新评估", v2_document_id, n_judgments)
    except Exception:
        logger.exception("fanout_contradictions_to_judgment_impact 失败 (doc=%s)", v2_document_id)

    # Stage B 扇出 #7：深度资料 → 客户画像复审标记（纯 SQL）
    try:
        flagged_profile = fanout_document_to_client_profile(
            db,
            client_id=client_id,
            v2_document_id=v2_document_id,
            file_name=safe_filename(title),
            chunk_count=chunk_count,
            visible_category=primary_category or "",
        )
        if flagged_profile:
            logger.info("[fanout] document=%s 标记 client=%s 画像需要复审", v2_document_id, client_id)
    except Exception:
        logger.exception("fanout_document_to_client_profile 失败 (doc=%s)", v2_document_id)

    return {
        "knowledge_document_id": v2_document_id,
        "legacy_knowledge_document_id": legacy_knowledge_document_id,
        "title": safe_filename(title),
        "current_human_path": str(managed_path),
        "compat_card_path": str(compat_card_path),
        "markdown_path": str(markdown_path),
        "primary_category": primary_category,
        "secondary_category": secondary_category,
        "short_summary": preview_text,
        "summary": preview_text,
        "classification_confidence": confidence,
        "material_layer": material_layer,
        "needs_review": parse_status != "ready",
        "parse_status": parse_status,
        "parse_error": parse_error,
        "ocr_metadata": {
            "totalPages": extraction_metadata.total_pages,
            "attemptedPages": extraction_metadata.attempted_pages,
            "succeededPages": extraction_metadata.succeeded_pages,
            "failedPages": extraction_metadata.failed_pages,
            "pageLimit": extraction_metadata.ocr_page_limit,
            "batchSize": extraction_metadata.ocr_batch_size,
            "partial": extraction_metadata.partial,
        },
        "chunk_count": chunk_count,
        "raw_text": text,
    }


def _stable_canonical_document_id(client_id: str, origin_type: str, origin_id: str) -> str:
    digest = hashlib.sha256(f"{client_id}:{origin_type}:{origin_id}".encode("utf-8")).hexdigest()[:20]
    return f"sysdoc_{digest}"


def fanout_contradictions_to_judgment_impact(
    db: Database,
    client_id: str,
    v2_document_id: str,
) -> int:
    """Stage B 扇出 #4：本次 ingest 新产生的事实矛盾，反向影响已有 confirmed judgment。

    Karpathy 启示 #10「矛盾是资产」的真实落地 —— 当一份新资料让事实打架时，
    已采纳的 confirmed judgment 标 needs_reevaluation=1，提示用户「知识结构变了，
    建议复查」。

    简化策略：因为 judgment_versions.evidence_ids_json 引用的是 evidence_cards.id
    （不是 atomic_facts.id），跨表精确匹配复杂；改为：只要本次 ingest 产生了任何
    与该客户相关的新矛盾，**所有** confirmed judgment 都标记复查。误报率可控，
    且产品语义清晰（任何矛盾都是"该 client 知识基础有变化"的信号）。
    """
    if not v2_document_id or not client_id:
        return 0

    # 找本次 ingest 产生的新 facts
    new_fact_rows = db.fetchall(
        "SELECT id FROM atomic_facts WHERE source_v2_document_id = ? AND status='active'",
        (v2_document_id,),
    )
    new_fact_ids = {str(r["id"]) for r in new_fact_rows}
    if not new_fact_ids:
        return 0

    # 看本次 ingest 的 facts 是否引发了矛盾
    placeholders = ",".join("?" for _ in new_fact_ids)
    contra_rows = db.fetchall(
        f"""
        SELECT COUNT(*) AS n
        FROM fact_contradictions
        WHERE client_id = ? AND review_status = 'pending'
          AND (fact_a_id IN ({placeholders}) OR fact_b_id IN ({placeholders}))
        """,
        (client_id, *new_fact_ids, *new_fact_ids),
    )
    contradiction_count = int(contra_rows[0]["n"] if contra_rows else 0)
    if contradiction_count == 0:
        return 0

    # 该客户所有 confirmed judgments 都标 needs_reevaluation
    judgment_rows = db.fetchall(
        """
        SELECT id, topic FROM judgment_versions
        WHERE client_id = ? AND status = 'confirmed'
          AND COALESCE(needs_reevaluation, 0) = 0
        """,
        (client_id,),
    )
    if not judgment_rows:
        return 0
    now_str = datetime.utcnow().isoformat()
    reason = f"新资料引入 {contradiction_count} 条事实矛盾，建议复查这条判断的证据链"
    impacted = 0
    for r in judgment_rows:
        jid = str(r["id"])
        db.execute(
            """
            UPDATE judgment_versions
            SET needs_reevaluation = 1,
                reevaluation_reason = ?,
                reevaluation_triggered_at = ?
            WHERE id = ?
            """,
            (reason, now_str, jid),
        )
        db.execute(
            """
            INSERT INTO activity_logs(id, actor_name, action, entity_type, entity_id, detail_json, created_at)
            VALUES(?, ?, ?, ?, ?, ?, ?)
            """,
            (
                new_id("evt"),
                "AI",
                "judgment.needs_reevaluation",
                "judgment_version",
                jid,
                to_json({
                    "v2DocumentId": v2_document_id,
                    "contradictionCount": contradiction_count,
                    "fanoutSource": "stage_b_4_contradiction_to_judgment",
                }),
                now_str,
            ),
        )
        impacted += 1
    return impacted


def fanout_document_to_client_profile(
    db: Database,
    client_id: str,
    v2_document_id: str,
    file_name: str,
    chunk_count: int,
    visible_category: str,
) -> bool:
    """Stage B 扇出 #7：深度资料 ingest 时标记客户画像需要复审。

    保守门槛：
      - chunk_count >= 30（深度资料，资料结构对画像有影响）
      - 或 visible_category 属于核心分类（战略陪伴 / 项目方案 / 组织 DNA 等）
    满足任一即标记 client_strategic_profiles.needs_review=1，让用户在客户档案页看到提示。
    纯 SQL，零 LLM 调用。
    """
    if not v2_document_id or not client_id:
        return False
    is_deep = (chunk_count or 0) >= 30
    is_core_category = (visible_category or "").strip() in {
        "战略陪伴", "项目方案", "项目与业务", "组织 DNA", "组织DNA", "核心资料", "战略资料",
    }
    if not (is_deep or is_core_category):
        return False
    # 确保 client_strategic_profiles 有一行可标
    existing = db.fetchone("SELECT client_id FROM client_strategic_profiles WHERE client_id = ?", (client_id,))
    now_str = datetime.utcnow().isoformat()
    trigger = "深度资料" if is_deep else f"核心分类「{visible_category}」"
    reason = f"AI 检测到「{file_name[:40]}」属于{trigger}（{chunk_count or 0} 段），建议复审客户画像是否需要更新"
    if existing:
        db.execute(
            """
            UPDATE client_strategic_profiles
            SET needs_review = 1,
                review_reason = ?,
                review_triggered_at = ?
            WHERE client_id = ?
            """,
            (reason, now_str, client_id),
        )
    else:
        db.execute(
            """
            INSERT INTO client_strategic_profiles (
                client_id, industry, scale, influence, current_needs, pain_points,
                strategic_value_to_yiyu, decision_chain, updated_at,
                needs_review, review_reason, review_triggered_at
            ) VALUES (?, '', '', '', '', '', '', '', ?, 1, ?, ?)
            """,
            (client_id, now_str, reason, now_str),
        )
    db.execute(
        """
        INSERT INTO activity_logs(id, actor_name, action, entity_type, entity_id, detail_json, created_at)
        VALUES(?, ?, ?, ?, ?, ?, ?)
        """,
        (
            new_id("evt"),
            "AI",
            "client_profile.needs_review",
            "client",
            client_id,
            to_json({
                "v2DocumentId": v2_document_id,
                "fileName": file_name,
                "chunkCount": chunk_count,
                "visibleCategory": visible_category,
                "triggerReason": "deep_document" if is_deep else "core_category",
                "fanoutSource": "stage_b_7_document_to_client_profile",
            }),
            now_str,
        ),
    )
    return True


def fanout_document_to_strategic_thoughts(
    db: Database,
    client_id: str,
    v2_document_id: str,
    file_name: str,
) -> bool:
    """Stage B 扇出 #3：满足门槛时，标记 strategic_thoughts 需要刷新。

    保守策略：要求文档已经被抽出 ≥3 atomic_facts 或 ≥3 entities 才标记，避免空文档触发。
    实际刷新由用户在战略陪伴 tab 打开时（或后续异步任务）真实调用 LLM 完成。
    这一步的价值是「让 AI 留下"这份资料应该被进一步思考"的待办」—— 对应 Karpathy 启示 #3
    的扇出追溯能力。
    """
    if not v2_document_id or not client_id:
        return False
    fact_row = db.fetchone(
        "SELECT COUNT(*) AS n FROM atomic_facts WHERE source_v2_document_id = ? AND status='active'",
        (v2_document_id,),
    )
    facts = int(fact_row["n"] if fact_row else 0)
    entity_row = db.fetchone(
        """
        SELECT COUNT(DISTINCT em.entity_id) AS n
        FROM entity_mentions em
        JOIN v2_chunks c ON c.id = em.v2_chunk_id
        WHERE c.v2_document_id = ?
        """,
        (v2_document_id,),
    )
    entities = int(entity_row["n"] if entity_row else 0)
    if facts < 3 and entities < 3:
        return False
    timestamp = datetime.utcnow().isoformat()
    detail = to_json({
        "triggerSource": "document_ingest",
        "v2DocumentId": v2_document_id,
        "fileName": file_name,
        "factsCount": facts,
        "entitiesCount": entities,
        "fanoutSource": "stage_b_3_document_to_strategic_thought",
    })
    db.execute(
        """
        INSERT INTO activity_logs(id, actor_name, action, entity_type, entity_id, detail_json, created_at)
        VALUES(?, ?, ?, ?, ?, ?, ?)
        """,
        (
            new_id("evt"),
            "AI",
            "strategic_thought.refresh_pending",
            "client",
            client_id,
            detail,
            timestamp,
        ),
    )
    return True


def fanout_document_to_event_lines(
    db: Database,
    client_id: str,
    v2_document_id: str,
    document_id: str,
    file_name: str,
    created_at: str,
) -> int:
    """Stage B 扇出 #2：新资料 ingest 后，通过 entity 关联自动建立和 event_lines 的活动联系。

    保守策略：至少 2 个 entity 命中同一 event_line 才触发，避免单点匹配的弱关联。
    产出的 event_line_activities is_key=0、source_type='document_ingest'，等用户在战略陪伴看到后确认。
    """
    if not v2_document_id or not client_id:
        return 0

    entity_rows = db.fetchall(
        """
        SELECT DISTINCT e.id, e.normalized_name, e.display_name
        FROM entity_mentions em
        JOIN entities e ON e.id = em.entity_id
        JOIN v2_chunks c ON c.id = em.v2_chunk_id
        WHERE c.v2_document_id = ? AND e.client_id = ? AND e.status = 'active'
        """,
        (v2_document_id, client_id),
    )
    if not entity_rows or len(entity_rows) < 2:
        return 0
    entity_names: list[str] = []
    entity_display: dict[str, str] = {}
    for r in entity_rows:
        name = str(r["normalized_name"] or "").strip().lower()
        if name:
            entity_names.append(name)
            entity_display[name] = str(r["display_name"] or name)
    if len(entity_names) < 2:
        return 0

    event_line_rows = db.fetchall(
        """
        SELECT id, name, summary, intent, primary_client_name
        FROM event_lines
        WHERE primary_client_id = ?
          AND COALESCE(status, 'active') NOT IN ('archived', 'completed')
        """,
        (client_id,),
    )

    file_name_lower = (file_name or "").lower()
    inserted = 0
    for el in event_line_rows:
        el_name = str(el["name"] or "")
        searchable = " ".join([
            el_name,
            str(el["summary"] or ""),
            str(el["intent"] or ""),
            str(el["primary_client_name"] or ""),
        ]).lower()
        matched_entities: list[str] = []
        for name in entity_names:
            if name and name in searchable:
                matched_entities.append(entity_display.get(name, name))

        # Fallback 触发逻辑：entity_extractor 抽取质量差时（常见痛点），
        # 至少让"文件名含 event_line 名 4+ 字符 substring"也能建立关联。
        # 这样像「日慈战略合作.docx」能匹配「日慈战略陪伴」事件线。
        filename_substring_match = False
        if el_name and len(el_name) >= 4 and file_name_lower:
            el_name_lower = el_name.lower()
            # 取 event_line.name 的 4 字符滑动 substring 检查
            for i in range(len(el_name_lower) - 3):
                sub = el_name_lower[i:i + 4]
                if sub in file_name_lower:
                    filename_substring_match = True
                    break

        # 触发条件：≥2 个 entity 匹配  OR  文件名 substring 命中
        should_trigger = len(matched_entities) >= 2 or filename_substring_match
        if not should_trigger:
            continue
        # 去重：同一份资料对同一条 event_line 只写一条 activity
        existing = db.fetchone(
            """
            SELECT id FROM event_line_activities
            WHERE event_line_id = ? AND source_type = 'document_ingest' AND source_id = ?
            LIMIT 1
            """,
            (str(el["id"]), v2_document_id),
        )
        if existing:
            continue
        activity_id = new_id("ela")
        metadata = to_json({
            "v2DocumentId": v2_document_id,
            "documentId": document_id,
            "fileName": file_name,
            "matchedEntities": matched_entities[:10],
            "matchedByFilename": filename_substring_match,
            "fanoutSource": "stage_b_2_document_to_event_line",
        })
        if len(matched_entities) >= 2:
            summary_text = (
                f"AI 在「{file_name[:40]}」里识别到 {len(matched_entities)} 个与本事件线相关的实体"
                f"（{'、'.join(matched_entities[:3])}{' 等' if len(matched_entities) > 3 else ''}），"
                "等你确认是否纳入正式活动。"
            )
        else:
            summary_text = (
                f"AI 根据文件名「{file_name[:40]}」推测这份资料与本事件线相关，"
                "等你确认是否纳入正式活动。"
            )
        db.execute(
            """
            INSERT INTO event_line_activities(
                id, event_line_id, source_type, source_id, happened_at,
                actor_id, actor_name, title, summary, metadata_json, is_key, created_at
            ) VALUES(?, ?, 'document_ingest', ?, ?, NULL, 'AI', ?, ?, ?, 0, ?)
            """,
            (
                activity_id,
                str(el["id"]),
                v2_document_id,
                created_at,
                f"新资料关联：{file_name[:60]}",
                summary_text,
                metadata,
                created_at,
            ),
        )
        inserted += 1
    return inserted


def _canonical_sections(text: str) -> list[dict[str, str]]:
    normalized = normalize_text(text)
    if not normalized:
        return []
    if "#" in normalized:
        return _build_markdown_sections(normalized)
    return _build_generic_sections(normalized)


def upsert_canonical_text_document(
    db: Database,
    *,
    data_dir: Path,
    client_id: str,
    canonical_kind: str,
    origin_type: str,
    origin_id: str,
    title: str,
    text: str,
    visible_category: str,
    secondary_category: str,
    created_at: str,
    updated_at: str,
    organization_id: str = "",
    department_id: str = "",
    department_ids: list[str] | tuple[str, ...] | None = None,
    owner_user_id: str = "",
    source_entity_type: str = "",
    source_entity_id: str = "",
    visibility_scope: str = "project_public",
    content_domain: str = "work",
    lifecycle_status: str = "active",
) -> dict[str, Any] | None:
    cleaned_text = _clean_ingested_text(text)
    if not cleaned_text:
        return None
    department_ids_json = to_json(list(normalize_department_ids(department_ids or department_id)))
    document_id = _stable_canonical_document_id(client_id, origin_type, origin_id)
    v2_document_id = f"v2doc_{document_id}"
    managed_path = _system_document_path(data_dir, client_id, canonical_kind, origin_id, title)
    managed_path.write_text(cleaned_text.strip() + "\n", encoding="utf-8")
    sections = _canonical_sections(cleaned_text)
    preview_text = build_excerpt(cleaned_text, title)
    content_hash = hashlib.sha256(cleaned_text.encode("utf-8")).hexdigest()
    document_family_id = derive_document_family_id(
        title=title,
        content_hash=content_hash,
        origin_type=origin_type,
        origin_id=origin_id,
    )
    doc_index_text = " ".join(
        dict.fromkeys(
            [
                *tokenize(f"{title}\n{visible_category}\n{secondary_category}\n{preview_text}"),
                *tokenize(" ".join(section.get("title", "") for section in sections[:24])),
            ]
        )
    )

    if db.fetchone("SELECT id FROM documents WHERE id = ?", (document_id,)) is None:
        db.execute(
            """
            INSERT INTO documents(
                id, client_id, folder_id, title, path, original_source_path, kind, source, excerpt, tags_json, created_at,
                document_family_id, canonical_kind, origin_type, origin_id, is_searchable,
                organization_id, department_id, department_ids_json, owner_user_id, source_entity_type, source_entity_id,
                visibility_scope, content_domain, lifecycle_status
            )
            VALUES(?, ?, NULL, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 1, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                document_id,
                client_id,
                safe_filename(title),
                str(managed_path),
                str(managed_path),
                canonical_kind,
                "workspace_native",
                preview_text,
                to_json([canonical_kind, visible_category, secondary_category, V2_PIPELINE_VERSION]),
                created_at,
                document_family_id,
                canonical_kind,
                origin_type,
                origin_id,
                organization_id,
                department_id,
                department_ids_json,
                owner_user_id,
                source_entity_type or origin_type,
                source_entity_id or origin_id,
                visibility_scope or "project_public",
                content_domain or "work",
                lifecycle_status or "active",
            ),
        )
    else:
        db.execute(
            """
            UPDATE documents
            SET client_id = ?, title = ?, path = ?, original_source_path = ?, kind = ?, source = ?, excerpt = ?, tags_json = ?,
                document_family_id = ?, canonical_kind = ?, origin_type = ?, origin_id = ?, is_searchable = ?,
                organization_id = ?, department_id = ?, department_ids_json = ?, owner_user_id = ?, source_entity_type = ?, source_entity_id = ?,
                visibility_scope = ?, content_domain = ?, lifecycle_status = ?
            WHERE id = ?
            """,
            (
                client_id,
                safe_filename(title),
                str(managed_path),
                str(managed_path),
                canonical_kind,
                "workspace_native",
                preview_text,
                to_json([canonical_kind, visible_category, secondary_category, V2_PIPELINE_VERSION]),
                document_family_id,
                canonical_kind,
                origin_type,
                origin_id,
                1 if (lifecycle_status or "active") == "active" and (visibility_scope or "project_public").lower() not in {"self", "private", "personal"} else 0,
                organization_id,
                department_id,
                department_ids_json,
                owner_user_id,
                source_entity_type or origin_type,
                source_entity_id or origin_id,
                visibility_scope or "project_public",
                content_domain or "work",
                lifecycle_status or "active",
                document_id,
            ),
        )

    _remove_existing_v2_rows(db, v2_document_id)
    payload = (
        v2_document_id,
        client_id,
        document_id,
        str(managed_path),
        str(managed_path),
        str(managed_path),
        safe_filename(title),
        canonical_kind,
        "evidence",
        visible_category,
        secondary_category,
        "ready",
        None,
        preview_text,
        doc_index_text,
        content_hash,
        0.95,
        document_family_id,
        canonical_kind,
        origin_type,
        origin_id,
        1 if (lifecycle_status or "active") == "active" and (visibility_scope or "project_public").lower() not in {"self", "private", "personal"} else 0,
        organization_id,
        department_id,
        department_ids_json,
        owner_user_id,
        source_entity_type or origin_type,
        source_entity_id or origin_id,
        visibility_scope or "project_public",
        content_domain or "work",
        lifecycle_status or "active",
        created_at,
        updated_at,
    )
    if db.fetchone("SELECT id FROM v2_documents WHERE id = ?", (v2_document_id,)) is None:
        db.execute(
            """
            INSERT INTO v2_documents(
                id, client_id, document_id, original_path, managed_path, markdown_path, file_name, kind, material_layer, visible_category,
                secondary_category, parse_status, parse_error, preview_text, doc_index_text, content_hash,
                classification_confidence, document_family_id, canonical_kind, origin_type, origin_id, is_searchable,
                organization_id, department_id, department_ids_json, owner_user_id, source_entity_type, source_entity_id,
                visibility_scope, content_domain, lifecycle_status, imported_at, updated_at
            )
            VALUES(?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            payload,
        )
    else:
        db.execute(
            """
            UPDATE v2_documents
            SET client_id = ?, document_id = ?, original_path = ?, managed_path = ?, markdown_path = ?, file_name = ?, kind = ?,
                material_layer = ?, visible_category = ?, secondary_category = ?, parse_status = ?, parse_error = ?,
                preview_text = ?, doc_index_text = ?, content_hash = ?, classification_confidence = ?,
                document_family_id = ?, canonical_kind = ?, origin_type = ?, origin_id = ?, is_searchable = ?,
                organization_id = ?, department_id = ?, department_ids_json = ?, owner_user_id = ?, source_entity_type = ?, source_entity_id = ?,
                visibility_scope = ?, content_domain = ?, lifecycle_status = ?, imported_at = ?, updated_at = ?
            WHERE id = ?
            """,
            payload[1:] + (v2_document_id,),
        )

    db.execute(
        "UPDATE v2_documents SET markdown_content = ? WHERE id = ?",
        (cleaned_text.strip() + "\n", v2_document_id),
    )

    chunk_count = 0
    for section_index, section in enumerate(sections):
        section_text = normalize_text(section.get("text", ""))
        if not section_text:
            continue
        section_title = normalize_text(section.get("title", "")) or "正文"
        section_id = f"v2sec_{document_id}_{section_index}"
        db.execute(
            """
            INSERT INTO v2_sections(id, v2_document_id, section_index, title, content, searchable_text, char_count, created_at)
            VALUES(?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                section_id,
                v2_document_id,
                section_index,
                section_title[:120],
                section_text,
                " ".join(tokenize(f"{section_title}\n{section_text[:2400]}")),
                len(section_text),
                updated_at,
            ),
        )
        for chunk_index, chunk in enumerate(build_chunks(section_text, section_title)):
            db.execute(
                """
                INSERT INTO v2_chunks(id, v2_document_id, v2_section_id, chunk_index, section_label, content, searchable_text, char_count, created_at)
                VALUES(?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    f"v2chunk_{document_id}_{section_index}_{chunk_index}",
                    v2_document_id,
                    section_id,
                    chunk_index,
                    chunk["sectionLabel"],
                    chunk["content"],
                    chunk["tokenText"],
                    chunk["charCount"],
                    updated_at,
                ),
            )
            chunk_count += 1
    db.execute(
        """
        UPDATE v2_documents
        SET section_count = ?, chunk_count = ?, updated_at = ?
        WHERE id = ?
        """,
        (len(sections), chunk_count, updated_at, v2_document_id),
    )
    return {
        "documentId": document_id,
        "v2DocumentId": v2_document_id,
        "documentFamilyId": document_family_id,
        "canonicalKind": canonical_kind,
        "originType": origin_type,
        "originId": origin_id,
    }


def _render_meeting_doc_text(
    db: Database,
    *,
    meeting_id: str,
    title: str,
    stage: str,
    scheduled_at: str,
    notes: str,
    transcript_text: str,
) -> str:
    agenda_rows = db.fetchall(
        "SELECT title, description FROM agenda_items WHERE meeting_id = ? ORDER BY sort_order ASC, rowid ASC",
        (meeting_id,),
    )
    decision_rows = db.fetchall("SELECT summary FROM decisions WHERE meeting_id = ? ORDER BY created_at DESC", (meeting_id,))
    action_rows = db.fetchall(
        "SELECT title, owner_name, due_date, publish_status FROM action_items WHERE meeting_id = ? ORDER BY created_at DESC",
        (meeting_id,),
    )
    source_rows = db.fetchall(
        "SELECT title, content_text FROM meeting_sources WHERE meeting_id = ? ORDER BY created_at DESC",
        (meeting_id,),
    )
    lines = [
        f"# 会议：{title}",
        "",
        "## 基本信息",
        f"- 阶段：{stage or '未标注'}",
        f"- 时间：{scheduled_at or '未记录'}",
    ]
    if notes:
        lines.extend(["", "## 会议纪要", "", notes])
    if transcript_text:
        lines.extend(["", "## 会议过程", "", transcript_text])
    if agenda_rows:
        lines.extend(["", "## 议程"])
        for row in agenda_rows:
            lines.append(f"- {normalize_text(str(row['title'] or '')) or '议程项'}：{normalize_text(str(row['description'] or '')) or '无描述'}")
    if decision_rows:
        lines.extend(["", "## 决议"])
        for row in decision_rows:
            summary = normalize_text(str(row["summary"] or ""))
            if summary:
                lines.append(f"- {summary}")
    if action_rows:
        lines.extend(["", "## 行动项"])
        for row in action_rows:
            lines.append(
                f"- {normalize_text(str(row['title'] or '')) or '行动项'}｜负责人：{normalize_text(str(row['owner_name'] or '')) or '待定'}｜截止：{normalize_text(str(row['due_date'] or '')) or '未定'}｜状态：{normalize_text(str(row['publish_status'] or '')) or 'draft'}"
            )
    if source_rows:
        lines.extend(["", "## 会议来源材料"])
        for row in source_rows[:8]:
            source_title = normalize_text(str(row["title"] or "")) or "来源材料"
            source_text = normalize_text(str(row["content_text"] or ""))
            if source_text:
                lines.extend([f"### {source_title}", "", source_text])
    return "\n".join(lines)


def _render_task_doc_text(db: Database, *, task_id: str, row: Any) -> str:
    note_row = db.fetchone("SELECT note FROM task_notes WHERE task_id = ?", (task_id,))
    attachment_rows = db.fetchall(
        "SELECT title, kind, path FROM task_attachments WHERE task_id = ? ORDER BY created_at DESC",
        (task_id,),
    )
    collaborator_rows = db.fetchall(
        "SELECT full_name, email, is_owner, inbox_status FROM task_collaborators WHERE task_id = ? ORDER BY order_index ASC, updated_at DESC",
        (task_id,),
    )
    lines = [
        f"# 任务：{str(row['title'] or '').strip() or task_id}",
        "",
        "## 基本信息",
        f"- 状态：{str(row['status'] or '').strip() or '未标注'}",
        f"- 优先级：{str(row['priority'] or '').strip() or '未标注'}",
        f"- 进度：{str(row['progress_status'] or '').strip() or '未标注'}",
        f"- 截止：{str(row['due_date'] or row['ddl'] or '').strip() or '未记录'}",
    ]
    for label, value in (
        ("任务描述", str(row["description"] or "")),
        ("当前阻塞", str(row["current_blocker"] or "")),
        ("下一步", str(row["next_action"] or "")),
        ("最近决策", str(row["recent_decision"] or "")),
    ):
        cleaned = normalize_text(value)
        if cleaned:
            lines.extend(["", f"## {label}", "", cleaned])
    note_text = normalize_text(str(note_row["note"] or "")) if note_row else ""
    if note_text:
        lines.extend(["", "## 任务备注", "", note_text])
    if collaborator_rows:
        lines.extend(["", "## 协作人"])
        for collaborator in collaborator_rows:
            lines.append(
                f"- {normalize_text(str(collaborator['full_name'] or '')) or '未命名'}｜邮箱：{normalize_text(str(collaborator['email'] or '')) or '未记录'}｜负责人：{'是' if int(collaborator['is_owner'] or 0) else '否'}｜收件状态：{normalize_text(str(collaborator['inbox_status'] or 'pending')) or 'pending'}"
            )
    if attachment_rows:
        lines.extend(["", "## 关联附件"])
        for attachment in attachment_rows[:10]:
            lines.append(
                f"- {normalize_text(str(attachment['title'] or '')) or '附件'}｜类型：{normalize_text(str(attachment['kind'] or '')) or '未知'}｜路径：{normalize_text(str(attachment['path'] or '')) or '未记录'}"
            )
    return "\n".join(lines)


def _render_review_doc_text(row: Any) -> str:
    lines = [
        f"# 周复盘：{str(row['week_label'] or '').strip() or str(row['id'])}",
        "",
        "## 基本信息",
        f"- 提交时间：{str(row['submitted_at'] or row['updated_at'] or row['created_at'] or '').strip() or '未记录'}",
    ]
    for label, value in (
        ("工作总结", str(row["summary"] or "")),
        ("工作进展", str(row["work_progress"] or "")),
        ("工作阻塞", str(row["work_blocker"] or "")),
        ("阻塞类型", str(row["blocker_type"] or "")),
        ("工作方向", str(row["work_direction"] or "")),
        ("下周重点", str(row["next_week_focus"] or "")),
        ("需要支持", str(row["support_needed"] or "")),
        ("自由备注", str(row["work_free_note"] or "")),
        ("成长备注", str(row["personal_growth_note"] or "")),
    ):
        cleaned = normalize_text(value)
        if cleaned:
            lines.extend(["", f"## {label}", "", cleaned])
    return "\n".join(lines)


def _render_judgment_doc_text(row: Any) -> str:
    lines = [
        f"# 判断：{str(row['topic'] or '').strip() or str(row['id'])}",
        "",
        "## 基本信息",
        f"- 状态：{str(row['status'] or '').strip() or '未标注'}",
        f"- 目标范围：{str(row['target_type'] or '').strip() or 'unknown'} / {str(row['target_id'] or '').strip() or 'unknown'}",
        f"- 风险等级：{str(row['risk_level'] or '').strip() or '未标注'}",
        f"- 置信度：{str(row['confidence'] or '').strip() or '未标注'}",
    ]
    summary = normalize_text(str(row["summary"] or ""))
    if summary:
        lines.extend(["", "## 判断内容", "", summary])
    evidence_ids = from_json(str(row["evidence_ids_json"] or "[]"), [])
    if isinstance(evidence_ids, list) and evidence_ids:
        lines.extend(["", "## 关联证据 ID"])
        for evidence_id in evidence_ids[:20]:
            lines.append(f"- {str(evidence_id)}")
    return "\n".join(lines)


def materialize_workspace_native_documents(db: Database, *, data_dir: Path, client_id: str) -> dict[str, int]:
    counts = {
        "meeting_doc": 0,
        "task_doc": 0,
        "review_doc": 0,
        "judgment_doc": 0,
        "event_line_doc": 0,
        "project_doc": 0,
        "calendar_doc": 0,
    }

    meeting_rows = db.fetchall(
        "SELECT id, title, stage, scheduled_at, transcript_text, notes, created_at, updated_at FROM meetings WHERE client_id = ?",
        (client_id,),
    )
    for row in meeting_rows:
        content = _render_meeting_doc_text(
            db,
            meeting_id=str(row["id"]),
            title=str(row["title"] or ""),
            stage=str(row["stage"] or ""),
            scheduled_at=str(row["scheduled_at"] or ""),
            notes=str(row["notes"] or ""),
            transcript_text=str(row["transcript_text"] or ""),
        )
        if upsert_canonical_text_document(
            db,
            data_dir=data_dir,
            client_id=client_id,
            canonical_kind="meeting_doc",
            origin_type="meeting",
            origin_id=str(row["id"]),
            title=str(row["title"] or "") or f"meeting_{row['id']}",
            text=content,
            visible_category="会议纪要",
            secondary_category="软件沉淀",
            created_at=str(row["created_at"] or row["updated_at"] or now_iso()),
            updated_at=str(row["updated_at"] or row["created_at"] or now_iso()),
        ):
            counts["meeting_doc"] += 1

    task_rows = db.fetchall(
        """
        SELECT *
        FROM tasks
        WHERE client_id = ?
          AND COALESCE(scope_mode, 'COLLAB_SHARED') != 'PERSONAL_ONLY'
        """,
        (client_id,),
    )
    for row in task_rows:
        content = _render_task_doc_text(db, task_id=str(row["id"]), row=row)
        if upsert_canonical_text_document(
            db,
            data_dir=data_dir,
            client_id=client_id,
            canonical_kind="task_doc",
            origin_type="task",
            origin_id=str(row["id"]),
            title=str(row["title"] or "") or f"task_{row['id']}",
            text=content,
            visible_category="任务资料",
            secondary_category="软件沉淀",
            created_at=str(row["created_at"] or row["updated_at"] or now_iso()),
            updated_at=str(row["updated_at"] or row["created_at"] or now_iso()),
        ):
            counts["task_doc"] += 1

    judgment_rows = db.fetchall(
        "SELECT * FROM judgment_versions WHERE client_id = ? ORDER BY updated_at DESC",
        (client_id,),
    )
    for row in judgment_rows:
        content = _render_judgment_doc_text(row)
        if upsert_canonical_text_document(
            db,
            data_dir=data_dir,
            client_id=client_id,
            canonical_kind="judgment_doc",
            origin_type="judgment_version",
            origin_id=str(row["id"]),
            title=str(row["topic"] or "") or f"judgment_{row['id']}",
            text=content,
            visible_category="正式判断",
            secondary_category="软件沉淀",
            created_at=str(row["created_at"] or row["updated_at"] or now_iso()),
            updated_at=str(row["updated_at"] or row["created_at"] or now_iso()),
        ):
            counts["judgment_doc"] += 1

    relation_counts = materialize_workspace_relation_documents(
        db,
        data_dir=data_dir,
        client_id=client_id,
        upsert_canonical_text_document=upsert_canonical_text_document,
        now_iso=now_iso,
    )
    for key, value in relation_counts.items():
        counts[key] = counts.get(key, 0) + int(value or 0)

    refresh_client_folder_counts(db, client_id)
    return counts


def backfill_knowledge_documents(
    db: Database,
    *,
    data_dir: Path,
    client_id: str,
    ai_service: Any | None = None,
    progress_callback: Callable[[int], None] | None = None,
) -> dict[str, Any]:
    del ai_service
    rows = db.fetchall("SELECT * FROM documents WHERE client_id = ? ORDER BY created_at ASC", (client_id,))
    processed = 0
    missing = 0
    for row in rows:
        existing_v2 = db.fetchone(
            "SELECT managed_path, original_path FROM v2_documents WHERE document_id = ? ORDER BY updated_at DESC LIMIT 1",
            (str(row["id"]),),
        )
        source_candidate = (
            str(existing_v2["managed_path"])
            if existing_v2 and existing_v2["managed_path"]
            else str(row["path"])
        )
        source_path = Path(source_candidate)
        if not source_path.exists():
            fallback_original = str(existing_v2["original_path"]) if existing_v2 and existing_v2["original_path"] else str(row["original_source_path"] or "")
            if fallback_original:
                original_fallback_path = Path(fallback_original)
                if original_fallback_path.exists():
                    source_path = original_fallback_path
                else:
                    missing += 1
                    continue
            else:
                missing += 1
                continue
        ingest_document_knowledge(
            db,
            data_dir=data_dir,
            client_id=client_id,
            import_id=None,
            document_id=str(row["id"]),
            source_path=source_path,
            original_source_path=Path(str(row["original_source_path"])) if row["original_source_path"] else source_path,
            title=str(row["title"]),
            kind=str(row["kind"]),
            source=str(row["source"]),
            fallback_excerpt=str(row["excerpt"] or row["title"]),
            created_at=str(row["created_at"]),
            ai_service=None,
        )
        processed += 1
        if progress_callback:
            progress_callback(processed)
    refresh_client_folder_counts(db, client_id)
    return {"processed": processed, "missing": missing}


def backfill_workspace_import(
    db: Database,
    *,
    data_dir: Path,
    client_id: str,
    source_root: Path | None = None,
    progress_callback: Callable[[int], None] | None = None,
) -> dict[str, Any]:
    ensure_client_folder_rows(db, data_dir, client_id)
    workspace_root = source_root or (data_dir / "client_workspace" / client_id)
    workspace_root.mkdir(parents=True, exist_ok=True)
    timestamp = now_iso()
    import_id = new_id("imp")
    job_id = new_id("kjob")
    db.execute(
        """
        INSERT INTO imports(id, client_id, source_path, mode, status, imported_count, skipped_count, created_at)
        VALUES(?, ?, ?, 'folder', 'processing', 0, 0, ?)
        """,
        (import_id, client_id, str(workspace_root), timestamp),
    )
    db.execute(
        """
        INSERT INTO knowledge_jobs(id, client_id, job_type, status, payload_json, total_items, processed_items, last_error, created_at, started_at, finished_at, updated_at)
        VALUES(?, ?, 'backfill_workspace_import', 'running', ?, 0, 0, NULL, ?, ?, NULL, ?)
        """,
        (job_id, client_id, to_json({"sourceRoot": str(workspace_root)}), timestamp, timestamp, timestamp),
    )

    existing_paths: set[str] = set()
    for row in db.fetchall("SELECT path FROM documents WHERE client_id = ?", (client_id,)):
        value = str(row["path"] or "").strip()
        if value:
            existing_paths.add(value)
    for row in db.fetchall(
        """
        SELECT original_path, managed_path
        FROM v2_documents
        WHERE client_id = ?
        """,
        (client_id,),
    ):
        for column in ("original_path", "managed_path"):
            value = str(row[column] or "").strip()
            if value:
                existing_paths.add(value)
    for row in db.fetchall(
        """
        SELECT original_path, import_source_path, current_human_path
        FROM knowledge_documents
        WHERE client_id = ?
        """,
        (client_id,),
    ):
        for column in ("original_path", "import_source_path", "current_human_path"):
            value = str(row[column] or "").strip()
            if value:
                existing_paths.add(value)

    candidates: list[Path] = []
    for path in workspace_root.rglob("*"):
        if not path.is_file():
            continue
        if any(part.startswith(".") for part in path.parts):
            continue
        if any(marker in path.parts for marker in INTERNAL_WORKSPACE_MARKERS):
            continue
        if path.name.lower() in NOISE_FILES:
            continue
        if path.suffix.lower() not in WORKSPACE_BACKFILL_EXTENSIONS:
            continue
        candidates.append(path)

    db.execute(
        "UPDATE knowledge_jobs SET total_items = ?, updated_at = ? WHERE id = ?",
        (len(candidates), now_iso(), job_id),
    )

    imported = 0
    skipped = 0
    for path in sorted(candidates):
        path_str = str(path)
        if path_str in existing_paths:
            skipped += 1
            continue
        folder_label = next((part for part in HUMAN_VISIBLE_CATEGORIES if part in path.parts), None)
        folder_row = (
            db.fetchone(
                "SELECT id FROM client_folders WHERE client_id = ? AND label = ?",
                (client_id, folder_label),
            )
            if folder_label
            else None
        )
        excerpt = build_excerpt(path.read_text(encoding="utf-8", errors="ignore"), path.name) if path.suffix.lower() in TEXT_EXTENSIONS else f"{path.name} 已进入资料缓冲池，可作为后续问答与证据引用来源。"
        document_id = new_id("doc")
        db.execute(
            """
            INSERT INTO documents(id, client_id, folder_id, title, path, kind, source, excerpt, tags_json, created_at)
            VALUES(?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                document_id,
                client_id,
                str(folder_row["id"]) if folder_row else None,
                path.name,
                path_str,
                path.suffix.lower().lstrip("."),
                "workspace_backfill",
                excerpt,
                to_json([path.suffix.lower().lstrip("."), "workspace_backfill"]),
                timestamp,
            ),
        )
        ingest_document_knowledge(
            db,
            data_dir=data_dir,
            client_id=client_id,
            import_id=import_id,
            document_id=document_id,
            source_path=path,
            original_source_path=path,
            title=path.name,
            kind=path.suffix.lower().lstrip("."),
            source="workspace_backfill",
            fallback_excerpt=excerpt,
            created_at=timestamp,
            ai_service=None,
        )
        existing_paths.add(path_str)
        imported += 1
        db.execute(
            "UPDATE imports SET imported_count = ?, skipped_count = ?, status = 'processing' WHERE id = ?",
            (imported, skipped, import_id),
        )
        db.execute(
            "UPDATE knowledge_jobs SET processed_items = ?, updated_at = ? WHERE id = ?",
            (imported, now_iso(), job_id),
        )
        if progress_callback:
            progress_callback(imported)

    refresh_client_folder_counts(db, client_id)
    finished_at = now_iso()
    db.execute(
        "UPDATE imports SET imported_count = ?, skipped_count = ?, status = 'completed' WHERE id = ?",
        (imported, skipped, import_id),
    )
    db.execute(
        """
        UPDATE knowledge_jobs
        SET status = 'completed', processed_items = ?, finished_at = ?, updated_at = ?
        WHERE id = ?
        """,
        (imported, finished_at, finished_at, job_id),
    )
    return {
        "importId": import_id,
        "jobId": job_id,
        "sourceRoot": str(workspace_root),
        "discovered": len(candidates),
        "imported": imported,
        "skipped": skipped,
    }


def compute_knowledge_status(db: Database, client_id: str, data_dir: Path | None = None) -> dict[str, Any]:
    runtime_status = (
        get_vector_runtime_status(db, data_dir=data_dir, client_id=client_id)
        if data_dir is not None
        else {
            "qdrantReady": True,
            "embeddingMode": "hash_fallback",
            "embeddingModel": None,
            "embeddingError": None,
            "embeddingProvider": None,
            "embeddingDimension": None,
            "embeddingSignature": None,
            "activeVectorCollection": None,
            "vectorIndexStatus": "ready",
            "routerEnabled": False,
            "routerModel": None,
            "rerankEnabled": False,
        }
    )
    main_job_placeholders = ",".join("?" for _ in MAIN_KNOWLEDGE_STATUS_JOB_TYPES)
    job_filter = f"client_id = ? AND job_type IN ({main_job_placeholders})"
    job_params = (client_id, *MAIN_KNOWLEDGE_STATUS_JOB_TYPES)
    document_count = int(db.scalar("SELECT COUNT(1) AS count FROM v2_documents WHERE client_id = ? AND material_layer IN ('evidence', 'external_media_transcript')", (client_id,)) or 0)
    v2_background_docs = int(db.scalar("SELECT COUNT(1) AS count FROM v2_documents WHERE client_id = ? AND material_layer = 'background'", (client_id,)) or 0)
    section_count = int(db.scalar("SELECT COUNT(1) AS count FROM v2_sections WHERE v2_document_id IN (SELECT id FROM v2_documents WHERE client_id = ? AND material_layer IN ('evidence', 'external_media_transcript'))", (client_id,)) or 0)
    chunk_count = int(db.scalar("SELECT COUNT(1) AS count FROM v2_chunks WHERE v2_document_id IN (SELECT id FROM v2_documents WHERE client_id = ? AND material_layer IN ('evidence', 'external_media_transcript'))", (client_id,)) or 0)
    failed_count = int(db.scalar("SELECT COUNT(1) AS count FROM v2_documents WHERE client_id = ? AND parse_status NOT IN ('ready', 'partial_ready')", (client_id,)) or 0)
    review_count = int(db.scalar("SELECT COUNT(1) AS count FROM v2_documents WHERE client_id = ? AND classification_confidence < 0.62", (client_id,)) or 0)
    dna_background_docs = int(db.scalar("SELECT COUNT(1) AS count FROM client_dna_documents WHERE client_id = ?", (client_id,)) or 0)
    memory_docs = int(db.scalar("SELECT COUNT(1) AS count FROM knowledge_surrogates WHERE client_id = ? AND source_type = 'memory_answer'", (client_id,)) or 0)
    pending_jobs = int(
        db.scalar(
            f"SELECT COUNT(1) AS count FROM knowledge_jobs WHERE {job_filter} AND status = 'queued'",
            job_params,
        )
        or 0
    )
    running_jobs = int(
        db.scalar(
            f"SELECT COUNT(1) AS count FROM knowledge_jobs WHERE {job_filter} AND status = 'running'",
            job_params,
        )
        or 0
    )
    last_job = db.fetchone(
        f"SELECT * FROM knowledge_jobs WHERE {job_filter} ORDER BY created_at DESC LIMIT 1",
        job_params,
    )
    last_status = str(last_job["status"]) if last_job else "idle"
    last_error = str(last_job["last_error"]) if last_job and last_job["last_error"] else None
    last_success_row = db.fetchone(
        f"""
        SELECT finished_at
        FROM knowledge_jobs
        WHERE {job_filter} AND status = 'completed'
        ORDER BY finished_at DESC
        LIMIT 1
        """,
        job_params,
    )
    last_updated_row = db.fetchone(
        """
        SELECT updated_at
        FROM v2_documents
        WHERE client_id = ?
        ORDER BY updated_at DESC
        LIMIT 1
        """,
        (client_id,),
    )
    last_success = str(last_success_row["finished_at"]) if last_success_row and last_success_row["finished_at"] else None
    last_updated = str(last_updated_row["updated_at"]) if last_updated_row and last_updated_row["updated_at"] else None
    return {
        "totalDocuments": document_count,
        "evidenceDocuments": document_count,
        "backgroundDocuments": dna_background_docs + v2_background_docs,
        "methodDocuments": 0,
        "totalSections": section_count,
        "totalChunks": chunk_count,
        "parseFailedDocuments": failed_count,
        "vectorizedDocuments": 0,
        "dedupedDocuments": 0,
        "reviewPendingDocuments": review_count,
        "surrogateCount": document_count,
        "memoryDocCount": memory_docs,
        "masterIndexCount": document_count,
        "reclassifiedDocumentCount": document_count,
        "qdrantReady": bool(runtime_status.get("qdrantReady", True)),
        "lastUpdatedAt": last_updated,
        "pendingJobs": pending_jobs,
        "runningJobs": running_jobs,
        "lastJobStatus": last_status,
        "lastJobError": last_error,
        "lastSuccessfulRunAt": last_success,
        "embeddingMode": str(runtime_status.get("embeddingMode") or "hash_fallback"),
        "embeddingModel": runtime_status.get("embeddingModel"),
        "embeddingError": runtime_status.get("embeddingError"),
        "embeddingProvider": runtime_status.get("embeddingProvider"),
        "embeddingDimension": runtime_status.get("embeddingDimension"),
        "embeddingSignature": runtime_status.get("embeddingSignature"),
        "activeVectorCollection": runtime_status.get("activeVectorCollection"),
        "vectorIndexStatus": runtime_status.get("vectorIndexStatus"),
        "routerEnabled": runtime_status.get("routerEnabled"),
        "routerModel": runtime_status.get("routerModel"),
        "rerankEnabled": runtime_status.get("rerankEnabled"),
        "lastPipelineVersion": V2_PIPELINE_VERSION,
    }


def fetch_document_cards(
    db: Database,
    client_id: str,
    data_dir: Path | None = None,
    limit: int | None = 120,
    access_context: DataCenterAccessContext | dict[str, Any] | None = None,
) -> list[dict[str, Any]]:
    access = build_document_access_where("vd", "d", access_context)
    rows = db.fetchall(
        f"""
        SELECT vd.*, d.created_at AS document_created_at, d.excerpt AS legacy_excerpt,
               dc.purpose AS card_purpose,
               dc.audience AS card_audience,
               dc.project_context AS card_project_context,
               dc.key_topics_json AS card_key_topics_json,
               dc.good_questions_json AS card_good_questions_json,
               dc.risk_notes AS card_risk_notes,
               dc.generated_model AS card_generated_model,
               dc.input_hash AS card_input_hash,
               dc.prompt_version AS card_prompt_version,
               dc.updated_at AS card_updated_at,
               dc.summary_200 AS card_summary_200,
               dc.one_line_summary AS card_one_line_summary,
               dc.keywords_json AS card_keywords_json,
               dc.tags_json AS card_tags_json,
               dpo.virtual_path AS path_virtual_path,
               dpo.classification_tags_json AS path_category_tags_json,
               dpo.client_id AS path_recommended_client_id,
               dpo.recommended_owner AS path_recommended_client_name,
               '' AS path_recommended_project_id,
               dpo.recommended_project AS path_recommended_project_name,
               dpo.virtual_path AS path_recommended_folder_label,
               dpo.confidence AS path_optimization_confidence,
               dpo.apply_status AS path_optimization_status,
               dpo.reason AS path_optimization_reasoning,
               dpo.updated_at AS path_optimization_updated_at
        FROM v2_documents vd
        JOIN documents d ON d.id = vd.document_id
        LEFT JOIN knowledge_documents kd ON kd.document_id = vd.document_id AND kd.client_id = vd.client_id
        LEFT JOIN document_cards dc ON dc.knowledge_document_id = kd.id
        LEFT JOIN document_path_optimizations dpo ON dpo.id = (
            SELECT dpo2.id
            FROM document_path_optimizations dpo2
            WHERE dpo2.knowledge_document_id = kd.id
            ORDER BY dpo2.updated_at DESC
            LIMIT 1
        )
        WHERE vd.client_id = ? AND vd.material_layer IN ('evidence', 'external_media_transcript')
          AND COALESCE(vd.canonical_kind, 'raw_file') = 'raw_file'
          AND {access.sql}
        ORDER BY vd.updated_at DESC
        LIMIT ?
        """,
        (client_id, *access.params, limit or 120),
    )
    cards: list[dict[str, Any]] = []
    for row in rows:
        compat_md_path = (
            str(_compat_card_path(data_dir, client_id, str(row["document_id"]), str(row["file_name"])))
            if data_dir is not None
            else None
        )
        card_meta = _card_metadata_for_row(row)
        physical_source_path = str(row["managed_path"] or "")
        display_path = str(card_meta.get("displayPath") or physical_source_path)
        summary_text = str(card_meta.get("cardSummary") or row["preview_text"] or row["legacy_excerpt"] or "")
        short_summary = str(card_meta.get("oneLineSummary") or row["preview_text"] or row["legacy_excerpt"] or "")
        visible_category = str(row["visible_category"] or "")
        logical_category = (
            str(card_meta.get("recommendedFolderLabel") or visible_category)
            if card_meta.get("virtualOptimizedPath")
            else visible_category
        )
        classification_reason = (
            f"32B 本地模型建议归入 {logical_category}，真实文件路径保持不变。"
            if card_meta.get("virtualOptimizedPath")
            else f"V2 根据文件标题与正文关键词归入 {visible_category}"
        )
        query_hints = _unique_texts(
            card_meta.get("cardKeywords"),
            card_meta.get("keyTopics"),
            tokenize(str(row["doc_index_text"] or "")),
        )[:10]
        keywords = _unique_texts(
            card_meta.get("cardKeywords"),
            card_meta.get("keyTopics"),
            tokenize(str(row["preview_text"] or "")),
        )[:10]
        tags = _unique_texts(
            card_meta.get("cardTags"),
            card_meta.get("virtualCategoryTags"),
            [logical_category, str(row["kind"]), V2_PIPELINE_VERSION],
        )[:10]
        retrieval_summary = "\n".join(
            str(part)
            for part in [
                card_meta.get("purpose"),
                card_meta.get("audience"),
                card_meta.get("projectContext"),
                summary_text,
            ]
            if str(part or "").strip()
        )
        cards.append(
            {
                "id": str(row["id"]).replace("v2doc_", "dock_"),
                "docId": str(row["id"]).replace("v2doc_", "dock_"),
                "clientId": str(row["client_id"]),
                "documentId": str(row["document_id"]),
                "title": str(row["file_name"]),
                "originalPath": str(row["original_path"]),
                "importSourcePath": str(row["original_path"]),
                "currentHumanPath": display_path,
                "humanFolderCategory": logical_category,
                "sourcePath": display_path,
                "logicalCategory": logical_category,
                "logicalSubcategory": str(row["secondary_category"]),
                "classificationReason": classification_reason,
                "normalizedPath": display_path,
                "surrogateMdPath": compat_md_path,
                "kind": str(row["kind"]),
                "primaryCategory": logical_category,
                "secondaryCategory": str(row["secondary_category"]),
                "shortSummary": short_summary,
                "summary": summary_text,
                "retrievalSummary": retrieval_summary,
                "documentRole": "原始证据",
                "queryHints": query_hints,
                "distinctFindings": [],
                "coreQuestions": [],
                "keywords": keywords,
                "tags": tags,
                "entities": [],
                "dateRange": None,
                "classificationConfidence": float(row["classification_confidence"] or 0.0),
                "needsReview": str(row["parse_status"]) != "ready",
                "deepRead": False,
                "lastHitQuestion": None,
                "dedupStatus": "unchecked",
                "vectorStatus": "disabled",
                "version": 1,
                "chunkCount": int(row["chunk_count"] or 0),
                "createdAt": str(row["imported_at"] or row["document_created_at"]),
                "updatedAt": str(row["updated_at"]),
                "purpose": card_meta.get("purpose") or None,
                "audience": card_meta.get("audience") or None,
                "projectContext": card_meta.get("projectContext") or None,
                "keyTopics": card_meta.get("keyTopics") or [],
                "goodQuestions": card_meta.get("goodQuestions") or [],
                "riskNotes": card_meta.get("riskNotes") or None,
                "displayPath": card_meta.get("displayPath"),
                "virtualOptimizedPath": card_meta.get("virtualOptimizedPath"),
                "pathOptimizationStatus": card_meta.get("pathOptimizationStatus"),
                "pathOptimizationConfidence": card_meta.get("pathOptimizationConfidence"),
                "generatedModel": card_meta.get("generatedModel") or None,
            }
        )
    return cards


def _score_by_terms(query_terms: list[str], *haystacks: str, title_weight: float = 1.0) -> tuple[float, list[str]]:
    matched: list[str] = []
    score = 0.0
    lowered_haystacks = [haystack.lower() for haystack in haystacks]
    for term in query_terms:
        if any(term in haystack for haystack in lowered_haystacks):
            matched.append(term)
            score += title_weight if term in lowered_haystacks[0] else 1.0
    return score, matched


def _normalize_path_for_match(path: str | None) -> str:
    raw = str(path or "").strip()
    if not raw:
        return ""
    return re.sub(r"[\\/]+", "/", raw).rstrip("/").lower()


def _normalize_title_for_match(title: str | None) -> str:
    raw = str(title or "").strip().lower()
    if not raw:
        return ""
    return re.sub(r"\s+", "", raw)


def _canonical_kind_for_row(row: Any) -> str:
    return str(row["canonical_kind"] or "raw_file").strip() or "raw_file"


def _is_software_kind(canonical_kind: str) -> bool:
    return canonical_kind not in {"", "raw_file"}


def _row_value(row: Any, key: str) -> Any:
    try:
        return row[key]
    except Exception:
        if isinstance(row, dict):
            return row.get(key)
        return getattr(row, key, None)


def _json_text_list(value: Any) -> list[str]:
    raw = from_json(value, []) if isinstance(value, str) else value
    if not isinstance(raw, list):
        return []
    result: list[str] = []
    for item in raw:
        text = str(item or "").strip()
        if text and text not in result:
            result.append(text)
    return result


def _unique_texts(*groups: Any) -> list[str]:
    result: list[str] = []
    seen: set[str] = set()
    for group in groups:
        values = group if isinstance(group, (list, tuple, set)) else [group]
        for item in values:
            text = str(item or "").strip()
            if text and text not in seen:
                seen.add(text)
                result.append(text)
    return result


def _text_value(value: Any) -> str:
    return str(value or "").strip()


def _float_or_none(value: Any) -> float | None:
    try:
        if value is None or str(value).strip() == "":
            return None
        return float(value)
    except (TypeError, ValueError):
        return None


def _card_metadata_for_row(row: Any) -> dict[str, Any]:
    key_topics = _json_text_list(_row_value(row, "card_key_topics_json"))
    good_questions = _json_text_list(_row_value(row, "card_good_questions_json"))
    keywords = _json_text_list(_row_value(row, "card_keywords_json"))
    tags = _json_text_list(_row_value(row, "card_tags_json"))
    category_tags = _json_text_list(_row_value(row, "path_category_tags_json"))
    status = _text_value(_row_value(row, "path_optimization_status"))
    raw_virtual_path = _text_value(_row_value(row, "path_virtual_path"))
    virtual_path = raw_virtual_path if status == "applied" else ""
    return {
        "purpose": _text_value(_row_value(row, "card_purpose")),
        "audience": _text_value(_row_value(row, "card_audience")),
        "projectContext": _text_value(_row_value(row, "card_project_context")),
        "keyTopics": key_topics,
        "goodQuestions": good_questions,
        "riskNotes": _text_value(_row_value(row, "card_risk_notes")),
        "cardSummary": _text_value(_row_value(row, "card_summary_200")),
        "oneLineSummary": _text_value(_row_value(row, "card_one_line_summary")),
        "cardKeywords": keywords,
        "cardTags": tags,
        "generatedModel": _text_value(_row_value(row, "card_generated_model")),
        "displayPath": virtual_path or None,
        "virtualOptimizedPath": virtual_path or None,
        "pathOptimizationStatus": status or None,
        "pathOptimizationConfidence": _float_or_none(_row_value(row, "path_optimization_confidence")),
        "virtualCategoryTags": category_tags,
        "recommendedClientId": _text_value(_row_value(row, "path_recommended_client_id")) or None,
        "recommendedClientName": _text_value(_row_value(row, "path_recommended_client_name")) or None,
        "recommendedProjectId": _text_value(_row_value(row, "path_recommended_project_id")) or None,
        "recommendedProjectName": _text_value(_row_value(row, "path_recommended_project_name")) or None,
        "recommendedFolderLabel": _text_value(_row_value(row, "path_recommended_folder_label")) or None,
        "pathOptimizationReasoning": _text_value(_row_value(row, "path_optimization_reasoning")) or None,
    }


def _card_support_text_for_row(row: Any) -> str:
    meta = _card_metadata_for_row(row)
    parts = [
        meta.get("purpose"),
        meta.get("audience"),
        meta.get("projectContext"),
        meta.get("cardSummary"),
        meta.get("oneLineSummary"),
        meta.get("riskNotes"),
        " ".join(meta.get("keyTopics") or []),
        " ".join(meta.get("goodQuestions") or []),
        " ".join(meta.get("cardKeywords") or []),
        " ".join(meta.get("virtualCategoryTags") or []),
        meta.get("virtualOptimizedPath"),
        meta.get("recommendedFolderLabel"),
        meta.get("pathOptimizationReasoning"),
    ]
    return "\n".join(str(part) for part in parts if str(part or "").strip())


def _clean_path_value(value: Any) -> str:
    return str(value or "").strip()


def _is_markdown_path(path: str | None) -> bool:
    return _clean_path_value(path).lower().endswith(".md")


def is_placeholder_machine_text(text: str) -> bool:
    """Return True for machine-readable text that is only a retry/placeholder stub."""

    normalized = normalize_text(text)
    if not normalized:
        return True
    compact = re.sub(r"\s+", "", normalized)
    if not compact:
        return True
    if any(marker in normalized for marker in PLACEHOLDER_MACHINE_TEXT_MARKERS):
        if len(compact) <= 220:
            return True
        # A document that is mostly headings plus a placeholder should still be invalid.
        body_without_headings = re.sub(r"(?m)^#{1,6}\s+.*$", "", normalized)
        body_compact = re.sub(r"\s+", "", body_without_headings)
        if len(body_compact) <= 80 and any(marker in body_without_headings for marker in PLACEHOLDER_MACHINE_TEXT_MARKERS):
            return True
    # Title-only markdown cards are not useful evidence.
    non_heading_body = re.sub(r"(?m)^#{1,6}\s+.*$", "", normalized)
    if len(compact) <= 160 and len(re.sub(r"\s+", "", non_heading_body)) <= 12:
        return True
    return False


def has_effective_machine_text(*values: Any) -> bool:
    for value in values:
        text = str(value or "").strip()
        if not text:
            continue
        compact = re.sub(r"\s+", "", normalize_text(text))
        if len(compact) >= 8 and not is_placeholder_machine_text(text):
            return True
    return False


def _read_text_sample(path: str, *, max_chars: int = 1200) -> str:
    try:
        return Path(path).read_text(encoding="utf-8", errors="ignore")[:max_chars]
    except (OSError, ValueError, UnicodeError):
        return ""


def _is_effective_machine_markdown_path(path: str) -> bool:
    cleaned = _clean_path_value(path)
    if not cleaned or not _is_markdown_path(cleaned):
        return False
    try:
        candidate = Path(cleaned)
        if not candidate.exists() or not candidate.is_file() or candidate.stat().st_size <= 0:
            return False
    except (OSError, ValueError):
        return False
    return has_effective_machine_text(_read_text_sample(cleaned))


def _first_non_markdown_path(*values: Any) -> str:
    for value in values:
        path = _clean_path_value(value)
        if path and not _is_markdown_path(path):
            return path
    return ""


def _existing_path(value: Any, *, allow_markdown: bool = True) -> str:
    path = _clean_path_value(value)
    if not path:
        return ""
    if not allow_markdown and _is_markdown_path(path):
        return ""
    try:
        candidate = Path(path)
        if candidate.exists() and candidate.is_file() and candidate.stat().st_size > 0:
            if _is_markdown_path(path) and not _is_effective_machine_markdown_path(path):
                return ""
            return path
    except (OSError, ValueError):
        return ""
    return ""


def _client_workspace_roots_from_paths(*values: Any) -> list[Path]:
    roots: list[Path] = []
    seen: set[str] = set()
    for value in values:
        path = _clean_path_value(value)
        if not path:
            continue
        try:
            parts = Path(path).parts
        except (OSError, ValueError):
            continue
        if "client_workspace" not in parts:
            continue
        index = parts.index("client_workspace")
        if len(parts) <= index + 1:
            continue
        root = Path(*parts[: index + 2])
        root_key = str(root)
        if root_key not in seen and root.exists():
            roots.append(root)
            seen.add(root_key)
    return roots


def _find_existing_same_named_file(*values: Any, allow_markdown: bool = True) -> str:
    """Recover stale workspace paths by searching the same client workspace.

    Some historical document rows kept paths under old folders while the actual
    imported file later moved to `_imports/...`. The search result should open
    the real file when an exact basename match exists inside the same client
    workspace, instead of emitting a dead path as an "original file".
    """

    candidates = [_clean_path_value(value) for value in values]
    roots = _client_workspace_roots_from_paths(*candidates)
    if not roots:
        return ""
    names: list[str] = []
    seen: set[str] = set()
    for path in candidates:
        if not path:
            continue
        if not allow_markdown and _is_markdown_path(path):
            continue
        name = Path(path).name
        if not name or name in seen:
            continue
        names.append(name)
        seen.add(name)
    for name in names:
        for root in roots:
            try:
                for match in root.rglob(name):
                    if (
                        match.is_file()
                        and match.stat().st_size > 0
                        and (allow_markdown or not _is_markdown_path(str(match)))
                        and (not _is_markdown_path(str(match)) or _is_effective_machine_markdown_path(str(match)))
                    ):
                        return str(match)
            except OSError:
                continue
    return ""


def _first_existing_path(*values: Any, allow_markdown: bool = True, recover_by_name: bool = True) -> str:
    for value in values:
        path = _existing_path(value, allow_markdown=allow_markdown)
        if path:
            return path
    if recover_by_name:
        return _find_existing_same_named_file(*values, allow_markdown=allow_markdown)
    return ""


def classify_source_availability_for_row(row: Any) -> dict[str, str | bool | None]:
    """Separate human-openable originals from machine-readable markdown.

    The retrieval layer reads markdown/sections, but the workspace UI needs to
    open editable Word/PDF/XLSX files when they exist. System generated records
    such as task/event-line docs only have markdown and must be labelled as
    system cards instead of "original documents".
    """

    canonical_kind = _canonical_kind_for_row(row)
    managed_path = _clean_path_value(_row_value(row, "managed_path"))
    original_path_raw = _clean_path_value(_row_value(row, "original_path"))
    document_path = _clean_path_value(_row_value(row, "document_path"))
    original_source_path = _clean_path_value(_row_value(row, "original_source_path"))
    markdown_path = _clean_path_value(_row_value(row, "markdown_path"))
    markdown_content = str(_row_value(row, "markdown_content") or "")
    preview_text = str(_row_value(row, "preview_text") or "")
    doc_index_text = str(_row_value(row, "doc_index_text") or "")
    origin_type = _clean_path_value(_row_value(row, "origin_type"))
    material_layer = _clean_path_value(_row_value(row, "material_layer"))
    is_online_transcript = origin_type == "video_transcript" or material_layer == "external_media_transcript"
    machine_text_available = has_effective_machine_text(markdown_content, preview_text, doc_index_text)

    if canonical_kind == "raw_file":
        original_candidates = (
            managed_path,
            original_path_raw,
            document_path,
            original_source_path,
        )
        markdown_candidates = (
            markdown_path,
            managed_path if _is_markdown_path(managed_path) else "",
            document_path if _is_markdown_path(document_path) else "",
            original_path_raw if _is_markdown_path(original_path_raw) else "",
            original_source_path if _is_markdown_path(original_source_path) else "",
        )
        if is_online_transcript:
            markdown_candidates = (
                managed_path if _is_markdown_path(managed_path) else "",
                document_path if _is_markdown_path(document_path) else "",
                original_path_raw if _is_markdown_path(original_path_raw) else "",
                original_source_path if _is_markdown_path(original_source_path) else "",
                markdown_path,
            )
        original_path = _first_existing_path(*original_candidates, allow_markdown=False)
        machine_markdown_path = _first_existing_path(*markdown_candidates, allow_markdown=True)
        machine_readable_available = bool(machine_markdown_path or machine_text_available)
        fallback_path = original_path or machine_markdown_path
        if original_path:
            source_availability = SOURCE_AVAILABILITY_ORIGINAL
            disabled_reason = None
        elif machine_readable_available:
            source_availability = SOURCE_AVAILABILITY_MACHINE_ONLY
            disabled_reason = "原文件已缺失，当前仅有机读稿。"
        else:
            source_availability = SOURCE_AVAILABILITY_INVALID
            disabled_reason = "原文件缺失或为空，且没有有效机读稿。"
        openable_kind = "original_file" if original_path else ("machine_markdown" if machine_markdown_path else "unknown")
        return {
            "path": fallback_path or None,
            "originalPath": original_path or None,
            "managedPath": original_path or machine_markdown_path or None,
            "markdownPath": machine_markdown_path or None,
            "openableKind": openable_kind,
            "sourceAvailability": source_availability,
            "originalAvailable": bool(original_path),
            "machineReadableAvailable": machine_readable_available,
            "openOriginalDisabledReason": disabled_reason,
        }

    system_path = _first_existing_path(
        markdown_path,
        managed_path,
        document_path,
        original_path_raw,
        recover_by_name=False,
    )
    machine_readable_available = bool(system_path or machine_text_available)
    return {
        "path": system_path or None,
        "originalPath": None,
        "managedPath": system_path or None,
        "markdownPath": markdown_path or system_path or None,
        "openableKind": "system_card" if system_path else "unknown",
        "sourceAvailability": SOURCE_AVAILABILITY_MACHINE_ONLY if machine_readable_available else SOURCE_AVAILABILITY_INVALID,
        "originalAvailable": False,
        "machineReadableAvailable": machine_readable_available,
        "openOriginalDisabledReason": "系统整理线索，不是原始上传文件。" if machine_readable_available else "没有有效机读稿。",
    }


def _openable_paths_for_row(row: Any) -> dict[str, str | bool | None]:
    return classify_source_availability_for_row(row)


LOW_INFORMATION_SYSTEM_KINDS = {"task_doc", "calendar_doc"}
LOW_INFORMATION_EXACT_SEGMENTS = {
    "补充录音",
    "补资料",
    "待补资料",
    "待补资料清单",
    "无",
    "暂无",
    "未记录",
}


def _compact_text_for_noise(value: str) -> str:
    return re.sub(r"\s+", "", normalize_text(value))


def _is_low_information_system_segment(canonical_kind: str, section_title: str | None, segment: str) -> bool:
    if canonical_kind not in LOW_INFORMATION_SYSTEM_KINDS:
        return False
    compact = _compact_text_for_noise(segment)
    if not compact:
        return True
    if compact in LOW_INFORMATION_EXACT_SEGMENTS:
        return True
    if compact.startswith("预期输出") and len(compact) <= 80:
        return True
    if compact.startswith("输出") and len(compact) <= 50 and ("清单" in compact or "提纲" in compact):
        return True
    section_compact = _compact_text_for_noise(section_title or "")
    if section_compact in {"下一步", "当前阻塞", "最近决策"} and len(compact) < 18:
        return True
    return False


def _split_reading_segments(text: str, *, min_chars: int = 1200, max_chars: int = 2200) -> list[str]:
    normalized = normalize_text(text)
    if not normalized:
        return []
    paragraphs = [segment.strip() for segment in SECTION_BREAK_PATTERN.split(normalized) if segment.strip()]
    if not paragraphs:
        paragraphs = [normalized]
    segments: list[str] = []
    current = ""
    for paragraph in paragraphs:
        candidate = f"{current}\n\n{paragraph}".strip() if current else paragraph
        if len(candidate) <= max_chars:
            current = candidate
            continue
        if current:
            segments.append(current)
        if len(paragraph) <= max_chars:
            current = paragraph
            continue
        start = 0
        while start < len(paragraph):
            end = min(len(paragraph), start + max_chars)
            piece = paragraph[start:end].strip()
            if len(piece) >= max(200, min_chars // 3):
                segments.append(piece)
            if end >= len(paragraph):
                break
            start = end
        current = ""
    if current:
        segments.append(current)
    return [segment for segment in segments if segment]


def _score_document_row(query_terms: list[str], prompt: str, row: Any) -> tuple[float, list[str]]:
    title = str(row["file_name"] or row["title"] or "")
    preview = str(row["preview_text"] or "")
    index_text = str(row["doc_index_text"] or "")
    category = str(row["visible_category"] or "其他资料")
    canonical_kind = _canonical_kind_for_row(row)
    card_support = _card_support_text_for_row(row)
    score, matched = _score_by_terms(
        query_terms,
        title.lower(),
        index_text.lower(),
        preview.lower(),
        category.lower(),
        canonical_kind.lower(),
        card_support.lower(),
        title_weight=2.6,
    )
    if category and category in prompt:
        score += 0.5
    virtual_category = str(_card_metadata_for_row(row).get("recommendedFolderLabel") or "")
    if virtual_category and virtual_category in prompt:
        score += 0.45
    if canonical_kind in prompt:
        score += 0.35
    return score, matched


def retrieve_knowledge_bundle(
    db: Database,
    data_dir: Path,
    client_id: str,
    prompt: str,
    access_context: DataCenterAccessContext | dict[str, Any] | None = None,
    priority_document_ids: list[str] | None = None,
) -> RetrievalBundle:
    del data_dir
    query_style = _query_style(prompt)
    query_terms = tokenize(prompt)
    if not query_terms and prompt.strip():
        query_terms = [prompt.strip().lower()]
    priority_ids = {str(item).strip() for item in (priority_document_ids or []) if str(item).strip()}
    access = build_document_access_where("vd", "d", access_context)
    rows = db.fetchall(
        f"""
        SELECT vd.*, d.title AS document_title, d.path AS document_path, d.original_source_path AS original_source_path,
               dc.purpose AS card_purpose,
               dc.audience AS card_audience,
               dc.project_context AS card_project_context,
               dc.key_topics_json AS card_key_topics_json,
               dc.good_questions_json AS card_good_questions_json,
               dc.risk_notes AS card_risk_notes,
               dc.generated_model AS card_generated_model,
               dc.input_hash AS card_input_hash,
               dc.prompt_version AS card_prompt_version,
               dc.updated_at AS card_updated_at,
               dc.summary_200 AS card_summary_200,
               dc.one_line_summary AS card_one_line_summary,
               dc.keywords_json AS card_keywords_json,
               dc.tags_json AS card_tags_json,
               dpo.virtual_path AS path_virtual_path,
               dpo.classification_tags_json AS path_category_tags_json,
               dpo.client_id AS path_recommended_client_id,
               dpo.recommended_owner AS path_recommended_client_name,
               '' AS path_recommended_project_id,
               dpo.recommended_project AS path_recommended_project_name,
               dpo.virtual_path AS path_recommended_folder_label,
               dpo.confidence AS path_optimization_confidence,
               dpo.apply_status AS path_optimization_status,
               dpo.reason AS path_optimization_reasoning,
               dpo.updated_at AS path_optimization_updated_at
        FROM v2_documents vd
        JOIN documents d ON d.id = vd.document_id
        LEFT JOIN knowledge_documents kd ON kd.document_id = vd.document_id AND kd.client_id = vd.client_id
        LEFT JOIN document_cards dc ON dc.knowledge_document_id = kd.id
        LEFT JOIN document_path_optimizations dpo ON dpo.id = (
            SELECT dpo2.id
            FROM document_path_optimizations dpo2
            WHERE dpo2.knowledge_document_id = kd.id
            ORDER BY dpo2.updated_at DESC
            LIMIT 1
        )
        WHERE vd.client_id = ?
          AND vd.material_layer IN ('evidence', 'external_media_transcript')
          AND vd.parse_status IN ('ready', 'partial_ready')
          AND COALESCE(vd.is_searchable, d.is_searchable, 1) = 1
          AND COALESCE(vd.lifecycle_status, d.lifecycle_status, 'active') = 'active'
          -- 迭代 2 F3：默认只检索 current 版本，被 superseded 的旧版不进检索池
          AND COALESCE(kd.lifecycle_status, 'current') = 'current'
          AND {access.sql}
          AND NOT (
            vd.canonical_kind = 'task_doc'
            AND vd.origin_type = 'task'
            AND NOT EXISTS (
              SELECT 1 FROM tasks t
              WHERE t.id = COALESCE(NULLIF(vd.source_entity_id, ''), vd.origin_id)
                AND COALESCE(t.scope_mode, 'COLLAB_SHARED') != 'PERSONAL_ONLY'
            )
          )
          AND NOT (
            vd.canonical_kind = 'task_note_doc'
            AND vd.origin_type = 'task_note'
            AND NOT EXISTS (
              SELECT 1 FROM tasks t
              WHERE t.id = COALESCE(NULLIF(vd.source_entity_id, ''), vd.origin_id)
                AND COALESCE(t.scope_mode, 'COLLAB_SHARED') != 'PERSONAL_ONLY'
            )
          )
          AND NOT (
            vd.canonical_kind = 'review_entry_doc'
            AND vd.origin_type = 'weekly_review_entry'
            AND NOT EXISTS (
              SELECT 1 FROM weekly_review_task_entries entry
              WHERE entry.id = COALESCE(NULLIF(vd.source_entity_id, ''), vd.origin_id)
                AND LOWER(COALESCE(entry.content_domain, 'work')) NOT IN ('personal', 'private')
            )
          )
          AND NOT (
            vd.canonical_kind = 'event_line_doc'
            AND vd.origin_type = 'event_line'
            AND NOT EXISTS (
              SELECT 1 FROM event_lines e
              WHERE e.id = COALESCE(NULLIF(vd.source_entity_id, ''), vd.origin_id)
                AND COALESCE(e.status, 'active') NOT IN ('deleted', 'archived')
                AND LOWER(COALESCE(e.visibility_scope, 'project_public')) NOT IN ('self', 'private', 'personal')
            )
          )
          AND NOT (
            vd.canonical_kind = 'event_line_update_doc'
            AND vd.origin_type = 'event_line_manual_update'
            AND NOT EXISTS (
              SELECT 1 FROM event_lines e
              WHERE e.id = COALESCE(NULLIF(vd.source_entity_id, ''), vd.origin_id)
                AND COALESCE(e.status, 'active') NOT IN ('deleted', 'archived')
                AND LOWER(COALESCE(e.visibility_scope, 'project_public')) NOT IN ('self', 'private', 'personal')
            )
          )
        ORDER BY vd.updated_at DESC
        """,
        (client_id, *access.params),
    )
    scored_docs: list[dict[str, Any]] = []
    family_pool: dict[str, dict[str, Any]] = {}
    for row in rows:
        if not has_effective_machine_text(row["markdown_content"], row["preview_text"], row["doc_index_text"]):
            continue
        openable_paths = _openable_paths_for_row(row)
        if str(openable_paths.get("sourceAvailability") or "") == SOURCE_AVAILABILITY_INVALID:
            continue
        card_meta = _card_metadata_for_row(row)
        display_path = str(card_meta.get("displayPath") or openable_paths.get("path") or "")
        score, matched = _score_document_row(query_terms, prompt, row)
        is_priority_document = bool(
            priority_ids
            and (
                str(row["document_id"] or "").strip() in priority_ids
                or str(row["id"] or "").strip() in priority_ids
            )
        )
        if is_priority_document:
            score += 50.0
            matched = list(dict.fromkeys([*matched, "working_document"]))
        canonical_kind = _canonical_kind_for_row(row)
        if query_style == "intro" and canonical_kind == "raw_file":
            intro_support_text = "\n".join(
                [
                    str(row["file_name"] or row["document_title"] or ""),
                    str(row["visible_category"] or ""),
                    str(row["secondary_category"] or ""),
                    _card_support_text_for_row(row),
                    str(row["preview_text"] or ""),
                    str(row["doc_index_text"] or ""),
                ]
            ).lower()
            intro_bucket = _intro_bucket_from_support_text(intro_support_text)
            if intro_bucket is not None:
                score += 0.35
                if not matched:
                    matched = [f"intro_bucket:{intro_bucket}"]
        if not matched and score <= 0:
            continue
        family_id = _effective_family_id_for_row(row)
        doc_payload = {
            "row": row,
            "score": score,
            "matchedTerms": matched,
            "title": str(row["file_name"] or row["document_title"] or "资料"),
            "path": str(openable_paths.get("path") or ""),
            "displayPath": display_path,
            "virtualOptimizedPath": card_meta.get("virtualOptimizedPath"),
            "pathOptimizationStatus": card_meta.get("pathOptimizationStatus"),
            "pathOptimizationConfidence": card_meta.get("pathOptimizationConfidence"),
            "purpose": card_meta.get("purpose"),
            "audience": card_meta.get("audience"),
            "projectContext": card_meta.get("projectContext"),
            "keyTopics": card_meta.get("keyTopics") or [],
            "goodQuestions": card_meta.get("goodQuestions") or [],
            "riskNotes": card_meta.get("riskNotes"),
            "originalPath": openable_paths.get("originalPath"),
            "managedPath": openable_paths.get("managedPath"),
            "markdownPath": openable_paths.get("markdownPath"),
            "openableKind": openable_paths.get("openableKind"),
            "sourceAvailability": openable_paths.get("sourceAvailability"),
            "originalAvailable": bool(openable_paths.get("originalAvailable")),
            "machineReadableAvailable": bool(openable_paths.get("machineReadableAvailable")),
            "openOriginalDisabledReason": openable_paths.get("openOriginalDisabledReason"),
            "familyId": family_id,
            "canonicalKind": canonical_kind,
            "isPriorityWorkingDocument": is_priority_document,
        }
        scored_docs.append(doc_payload)
        bucket = family_pool.setdefault(
            family_id,
            {
                "familyId": family_id,
                "bestScore": float("-inf"),
                "rows": [],
                "matchedTerms": set(),
                "canonicalKinds": set(),
            },
        )
        bucket["rows"].append(doc_payload)
        bucket["bestScore"] = max(float(bucket["bestScore"]), float(score))
        bucket["matchedTerms"].update(matched)
        bucket["canonicalKinds"].add(canonical_kind)
        bucket["isPriorityWorkingDocument"] = bool(bucket.get("isPriorityWorkingDocument")) or is_priority_document

    families = list(family_pool.values())
    families.sort(key=lambda item: (1 if item.get("isPriorityWorkingDocument") else 0, float(item["bestScore"])), reverse=True)
    if not families:
        return RetrievalBundle(
            citations=[],
            coverage=0.0,
            retrieval_summary={
                "docHitCount": 0,
                "sectionHitCount": 0,
                "rawChunkHitCount": 0,
                "masterHitCount": 0,
                "surrogateHitCount": 0,
                "readingPassCount": 0,
                "selectedDocumentFamilyCount": 0,
                "selectedCanonicalKinds": [],
                "softwareMaterialIncluded": False,
                "backgroundTrail": [],
                "workingDocumentIds": sorted(priority_ids),
                "workingDocumentHitCount": 0,
            },
            context_text="",
            matched_terms=[],
            failure_reason="no_candidate_documents",
        )

    selected_pass1: list[dict[str, Any]] = []
    selected_ids: set[str] = set()
    for family in [item for item in families if item.get("isPriorityWorkingDocument")]:
        selected_pass1.append(family)
        selected_ids.add(str(family["familyId"]))
        if len(selected_pass1) >= 8:
            break
    if query_style == "intro":
        raw_families = [item for item in families if "raw_file" in item["canonicalKinds"]]
        preferred_raw_families = [item for item in raw_families if _is_intro_primary_family(item)]
        for bucket_name in ("institution", "project", "strategy", "source"):
            candidates = [
                item
                for item in preferred_raw_families
                if _intro_bucket_for_family(item) == bucket_name and item["familyId"] not in selected_ids
            ]
            if not candidates:
                continue
            selected = sorted(candidates, key=lambda item: float(item["bestScore"]), reverse=True)[0]
            selected_pass1.append(selected)
            selected_ids.add(str(selected["familyId"]))
        for family in preferred_raw_families:
            if str(family["familyId"]) in selected_ids:
                continue
            selected_pass1.append(family)
            selected_ids.add(str(family["familyId"]))
            if len(selected_pass1) >= 8:
                break
        for family in raw_families:
            if str(family["familyId"]) in selected_ids:
                continue
            selected_pass1.append(family)
            selected_ids.add(str(family["familyId"]))
            if len(selected_pass1) >= 8:
                break
        if len(selected_pass1) < 8:
            for family in families:
                if str(family["familyId"]) in selected_ids:
                    continue
                selected_pass1.append(family)
                selected_ids.add(str(family["familyId"]))
                if len(selected_pass1) >= 8:
                    break
    elif query_style == "meeting":
        meeting_families = [item for item in families if "meeting_doc" in item["canonicalKinds"]]
        raw_families = [item for item in families if "raw_file" in item["canonicalKinds"]]
        for family in meeting_families[:4]:
            if str(family["familyId"]) in selected_ids:
                continue
            selected_pass1.append(family)
            selected_ids.add(str(family["familyId"]))
        for family in raw_families:
            if str(family["familyId"]) in selected_ids:
                continue
            selected_pass1.append(family)
            selected_ids.add(str(family["familyId"]))
            if len(selected_pass1) >= 8:
                break
        if len(selected_pass1) < 8:
            for family in families:
                if str(family["familyId"]) in selected_ids:
                    continue
                selected_pass1.append(family)
                selected_ids.add(str(family["familyId"]))
                if len(selected_pass1) >= 8:
                    break
    else:
        raw_candidate = next((item for item in families if "raw_file" in item["canonicalKinds"]), None)
        software_candidate = next((item for item in families if any(_is_software_kind(kind) for kind in item["canonicalKinds"])), None)
        if raw_candidate is not None and str(raw_candidate["familyId"]) not in selected_ids:
            selected_pass1.append(raw_candidate)
            selected_ids.add(str(raw_candidate["familyId"]))
        if software_candidate is not None and software_candidate is not raw_candidate and str(software_candidate["familyId"]) not in selected_ids:
            selected_pass1.append(software_candidate)
            selected_ids.add(str(software_candidate["familyId"]))
        for family in families:
            if str(family["familyId"]) in selected_ids:
                continue
            selected_pass1.append(family)
            selected_ids.add(str(family["familyId"]))
            if len(selected_pass1) >= 8:
                break
    selected_terms = {term for family in selected_pass1 for term in family["matchedTerms"]}
    selected_kinds = {kind for family in selected_pass1 for kind in family["canonicalKinds"]}
    uncovered_terms = [term for term in query_terms if term not in selected_terms]

    remaining_families = [family for family in families if family not in selected_pass1]
    remaining_families.sort(
        key=lambda item: (
            1 if query_style == "intro" and any(_is_software_kind(kind) for kind in item["canonicalKinds"]) else 0,
            1 if query_style == "meeting" and "meeting_doc" in item["canonicalKinds"] else 0,
            sum(1 for term in uncovered_terms if term in item["matchedTerms"]),
            len([kind for kind in item["canonicalKinds"] if kind not in selected_kinds]),
            float(item["bestScore"]),
        ),
        reverse=True,
    )
    selected_pass2: list[dict[str, Any]] = []
    for family in remaining_families[:4]:
        selected_pass2.append(family)

    selected_families = [*selected_pass1[:8], *selected_pass2]
    representative_docs = []
    for family in selected_families:
        rows_for_family = sorted(family["rows"], key=lambda item: float(item["score"]), reverse=True)
        if rows_for_family:
            representative_docs.append(rows_for_family[0])
    v2_doc_ids = [str(item["row"]["id"]) for item in representative_docs]
    section_rows = db.fetchall(
        f"""
        SELECT *
        FROM v2_sections
        WHERE v2_document_id IN ({','.join('?' for _ in v2_doc_ids)})
        ORDER BY section_index ASC
        """,
        tuple(v2_doc_ids),
    ) if v2_doc_ids else []
    sections_by_doc: dict[str, list[Any]] = defaultdict(list)
    for row in section_rows:
        sections_by_doc[str(row["v2_document_id"])].append(row)

    citations: list[CitationMatch] = []
    matched_terms: set[str] = set()
    background_trail: list[dict[str, Any]] = []
    for doc in representative_docs:
        row = doc["row"]
        doc_sections = sections_by_doc.get(str(row["id"]), [])
        scored_segments: list[tuple[float, str, str | None, list[str]]] = []
        for section in doc_sections:
            section_title = str(section["title"] or "正文")
            section_content = str(section["content"] or "")
            section_score, section_matched = _score_by_terms(
                query_terms,
                section_title.lower(),
                section_content[:3600].lower(),
                title_weight=2.0,
            )
            section_score += float(doc["score"]) * 0.35
            for segment in _split_reading_segments(section_content):
                canonical_kind = _canonical_kind_for_row(row)
                if _is_low_information_system_segment(canonical_kind, section_title, segment):
                    continue
                segment_score, segment_matched = _score_by_terms(query_terms, segment.lower(), title_weight=1.2)
                merged_terms = sorted(set([*doc["matchedTerms"], *section_matched, *segment_matched]))
                scored_segments.append((section_score + segment_score, segment, section_title, merged_terms))
        if not scored_segments:
            fallback_excerpt = str(row["preview_text"] or "")
            canonical_kind = _canonical_kind_for_row(row)
            if fallback_excerpt and not _is_low_information_system_segment(canonical_kind, "文档摘要", fallback_excerpt):
                scored_segments.append((float(doc["score"]), fallback_excerpt[:2200], "文档摘要", list(doc["matchedTerms"])))
        scored_segments.sort(key=lambda item: item[0], reverse=True)
        for score, segment, section_title, segment_terms in scored_segments[:2]:
            matched_terms.update(segment_terms)
            citations.append(
                CitationMatch(
                    knowledge_document_id=str(row["document_id"] or row["id"]),
                    chunk_id=None,
                    title=str(doc["title"]),
                    excerpt=str(segment)[:2200],
                    score=round(float(score), 4),
                    coverage=0.0,
                    section_label=section_title,
                    source_stage="raw_chunk",
                    drillthrough_used=True,
                    matched_terms=segment_terms,
                    path=str(doc["path"]),
                    original_path=str(doc.get("originalPath") or "") or None,
                    managed_path=str(doc.get("managedPath") or "") or None,
                    markdown_path=str(doc.get("markdownPath") or "") or None,
                    openable_kind=str(doc.get("openableKind") or "") or None,
                    source_availability=str(doc.get("sourceAvailability") or "") or None,
                    original_available=bool(doc.get("originalAvailable")),
                    machine_readable_available=bool(doc.get("machineReadableAvailable")),
                    open_original_disabled_reason=str(doc.get("openOriginalDisabledReason") or "") or None,
                    display_path=str(doc.get("displayPath") or "") or None,
                    virtual_optimized_path=str(doc.get("virtualOptimizedPath") or "") or None,
                    path_optimization_status=str(doc.get("pathOptimizationStatus") or "") or None,
                    path_optimization_confidence=doc.get("pathOptimizationConfidence"),
                    purpose=str(doc.get("purpose") or "") or None,
                    audience=str(doc.get("audience") or "") or None,
                    project_context=str(doc.get("projectContext") or "") or None,
                    key_topics=[str(item) for item in doc.get("keyTopics") or [] if str(item).strip()],
                    good_questions=[str(item) for item in doc.get("goodQuestions") or [] if str(item).strip()],
                    risk_notes=str(doc.get("riskNotes") or "") or None,
                    document_family_id=str(row["document_family_id"] or doc["familyId"]),
                    canonical_kind=_canonical_kind_for_row(row),
                    origin_type=str(row["origin_type"] or ""),
                    origin_id=str(row["origin_id"] or ""),
                    is_searchable=bool(int(row["is_searchable"] or 1)),
                )
            )
        background_trail.append(
            {
                "title": str(doc["title"]),
                "stage": "family_first",
                "sectionLabel": None,
                "path": str(doc["path"]),
                "originalPath": str(doc.get("originalPath") or ""),
                "managedPath": str(doc.get("managedPath") or ""),
                "markdownPath": str(doc.get("markdownPath") or ""),
                "openableKind": str(doc.get("openableKind") or ""),
                "sourceAvailability": str(doc.get("sourceAvailability") or ""),
                "originalAvailable": bool(doc.get("originalAvailable")),
                "machineReadableAvailable": bool(doc.get("machineReadableAvailable")),
                "openOriginalDisabledReason": str(doc.get("openOriginalDisabledReason") or ""),
                "displayPath": str(doc.get("displayPath") or ""),
                "virtualOptimizedPath": str(doc.get("virtualOptimizedPath") or ""),
                "pathOptimizationStatus": str(doc.get("pathOptimizationStatus") or ""),
                "pathOptimizationConfidence": doc.get("pathOptimizationConfidence"),
                "purpose": str(doc.get("purpose") or ""),
                "audience": str(doc.get("audience") or ""),
                "projectContext": str(doc.get("projectContext") or ""),
                "keyTopics": doc.get("keyTopics") or [],
                "goodQuestions": doc.get("goodQuestions") or [],
                "riskNotes": str(doc.get("riskNotes") or ""),
                "excerpt": str(row["preview_text"] or "")[:180],
                "documentFamilyId": str(row["document_family_id"] or doc["familyId"]),
                "canonicalKind": _canonical_kind_for_row(row),
            }
        )
        if len(citations) >= 24:
            break

    citations = citations[:24]
    selected_document_families = list(dict.fromkeys([item.document_family_id or item.knowledge_document_id for item in citations]))
    selected_canonical_kinds = list(dict.fromkeys([item.canonical_kind or "raw_file" for item in citations]))
    coverage = round(len(matched_terms) / len(query_terms), 2) if query_terms and citations else 0.0
    if citations:
        coverage = max(0.55, coverage)
    citations = [
        CitationMatch(
            knowledge_document_id=item.knowledge_document_id,
            chunk_id=item.chunk_id,
            title=item.title,
            excerpt=item.excerpt,
            score=item.score,
            coverage=coverage,
            section_label=item.section_label,
            source_stage=item.source_stage,
            drillthrough_used=item.drillthrough_used,
            matched_terms=item.matched_terms,
            path=item.path,
            original_path=item.original_path,
            managed_path=item.managed_path,
            markdown_path=item.markdown_path,
            openable_kind=item.openable_kind,
            source_availability=item.source_availability,
            original_available=item.original_available,
            machine_readable_available=item.machine_readable_available,
            open_original_disabled_reason=item.open_original_disabled_reason,
            display_path=item.display_path,
            virtual_optimized_path=item.virtual_optimized_path,
            path_optimization_status=item.path_optimization_status,
            path_optimization_confidence=item.path_optimization_confidence,
            purpose=item.purpose,
            audience=item.audience,
            project_context=item.project_context,
            key_topics=item.key_topics,
            good_questions=item.good_questions,
            risk_notes=item.risk_notes,
            document_family_id=item.document_family_id,
            canonical_kind=item.canonical_kind,
            origin_type=item.origin_type,
            origin_id=item.origin_id,
            is_searchable=item.is_searchable,
        )
        for item in citations
    ]
    context_lines = []
    for index, citation in enumerate(citations, start=1):
        label = citation.title
        if citation.section_label:
            label = f"{label} / {citation.section_label}"
        card_context_parts = []
        if citation.purpose:
            card_context_parts.append(f"用途：{citation.purpose}")
        if citation.audience:
            card_context_parts.append(f"服务对象：{citation.audience}")
        if citation.project_context:
            card_context_parts.append(f"项目语境：{citation.project_context}")
        if citation.key_topics:
            card_context_parts.append(f"主题：{'、'.join(citation.key_topics[:6])}")
        if citation.good_questions:
            card_context_parts.append(f"适合回答：{'、'.join(citation.good_questions[:3])}")
        if citation.risk_notes:
            card_context_parts.append(f"风险提示：{citation.risk_notes}")
        card_context = "；".join(card_context_parts)
        card_context_line = f"名片：{card_context}\n" if card_context else ""
        context_lines.append(
            f"[证据{index}] {label}\n命中要点：{'、'.join(citation.matched_terms) or '无'}\n"
            f"{card_context_line}原文：{citation.excerpt}"
        )
    retrieval_summary = {
        "docHitCount": len(families),
        "sectionHitCount": len(section_rows),
        "rawChunkHitCount": len(citations),
        "masterHitCount": len(families),
        "surrogateHitCount": 0,
        "readingPassCount": 2,
        "selectedDocumentFamilyCount": len(selected_document_families),
        "selectedCanonicalKinds": selected_canonical_kinds,
        "softwareMaterialIncluded": any(_is_software_kind(kind) for kind in selected_canonical_kinds),
        "backgroundTrail": background_trail[:12],
        "workingDocumentIds": sorted(priority_ids),
        "workingDocumentHitCount": sum(1 for item in citations if str(item.knowledge_document_id or "").strip() in priority_ids),
    }
    return RetrievalBundle(
        citations=citations,
        coverage=coverage,
        retrieval_summary=retrieval_summary,
        context_text="\n\n".join(context_lines),
        matched_terms=sorted(matched_terms),
        failure_reason=None if citations else "no_grounded_citations",
    )
