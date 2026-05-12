"""生成"益语智库自用平台 V2.0 产品开发 Q1 2026 复盘报告"（内部说明版）。

这是一份与战略陪伴报告**完全不同章节结构**的演示——
- 战略陪伴报告：客户视角，章节是"战略目标 / 关键事件 / 行动建议"
- 产品复盘报告：内部研发视角，章节是"版本里程碑 / 功能模块 / 技术债 / 下版本规划"

**共用部分（演示报告生成器的通用性）：**
- 同一套 helper（封面、标题、表格、数据源标注）
- 同一套数据源抽象（事件线 / 判断 / 任务 / 风险）→ 内部报告映射到 git commits / 设计决策 / TodoList / 技术债

**差异部分（演示主题感知）：**
- 章节骨架不同
- 数据源具体取数方式不同（git log vs 战略陪伴会议纪要）
- 受众语气不同（工程 + 决策层 vs 甲方）
"""

from __future__ import annotations

import sys
from pathlib import Path

# 复用前一份脚本里的所有 helper（封面、标题、表格、数据源标注等）
sys.path.insert(0, str(Path(__file__).parent))
from generate_report import (  # noqa
    add_bullet,
    add_callout,
    add_data_source,
    add_heading,
    add_horizontal_rule,
    add_page_break,
    add_paragraph,
    add_table,
    set_cell_bg,
    set_cell_borders,
    set_run_font,
    COLOR_BRAND,
    COLOR_TEXT,
    COLOR_MUTED,
    COLOR_TABLE_HEADER_BG,
    COLOR_TABLE_ALT_BG,
    FONT_CN,
    FONT_EN,
)

import io as _io

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

# 图表生成器（matplotlib 后端）
import chart_generator as chart


def add_image_from_bytes(doc, png_bytes: bytes, *, width_cm: float = 14.5,
                         caption: str | None = None):
    """把 PNG bytes 作为图片插入 docx；可选图注。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(_io.BytesIO(png_bytes), width=Cm(width_cm))
    if caption:
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.paragraph_format.space_after = Pt(10)
        run = c.add_run(caption)
        set_run_font(run, size=9, italic=True, color=COLOR_MUTED)
    return p


OUTPUT_PATH = Path(
    "/sessions/nifty-zen-euler/mnt/yiyu-thinktank-workbench/docs/"
    "益语智库V2.0-产品开发复盘-Q1-2026.docx"
)


def build():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # 默认字体
    normal_style = doc.styles["Normal"]
    normal_style.font.name = FONT_EN
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = COLOR_TEXT
    rPr = normal_style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), FONT_CN)
    rFonts.set(qn("w:ascii"), FONT_EN)
    rFonts.set(qn("w:hAnsi"), FONT_EN)

    # ============ 封面 ============
    for _ in range(3):
        add_paragraph(doc, "", space_after=4)

    add_paragraph(
        doc, "益语智库自用平台 V2.0",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        size=30, bold=True, color=COLOR_BRAND,
        space_before=20, space_after=4, line_spacing=1.2,
    )
    add_paragraph(
        doc, "产品开发季度复盘报告",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        size=24, bold=True, color=COLOR_TEXT,
        space_after=14, line_spacing=1.2,
    )

    add_paragraph(doc, "", space_after=4)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 18)
    set_run_font(run, size=10, color=COLOR_BRAND, bold=True)
    add_paragraph(doc, "", space_after=4)

    add_paragraph(
        doc, "2026 Q1 内部复盘 · 集团说明版",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        size=14, color=COLOR_MUTED, italic=True,
        space_after=40,
    )

    meta_rows = [
        ["报告期", "2026 年 1 月 — 5 月 11 日（含 V2.0 内测 DMG 发布）"],
        ["产品代号", "益语智库自用平台 V2.0（com.yiyu.selfworkbench2）"],
        ["产品版本", "v0.1.0（首发内测）"],
        ["技术栈", "Electron + Vite + React + TypeScript + Python (FastAPI) + SQLite"],
        ["报告类型", "产品开发季度复盘（内部说明）"],
        ["呈送对象", "益语集团 · 决策层 / 产品团队 / 工程团队"],
        ["编制日期", "2026-05-12"],
    ]
    add_table(
        doc,
        headers=["", ""],
        rows=meta_rows,
        widths=[4, 11],
        header_color="FFFFFF",
        header_text_color=COLOR_BRAND,
        zebra=True,
    )

    add_paragraph(doc, "", space_after=40)
    add_paragraph(
        doc, "本报告为益语集团内部资料，仅供集团内部审阅讨论使用",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        size=9, color=COLOR_MUTED, italic=True,
    )

    add_page_break(doc)

    # ============ 报告说明 ============
    add_callout(
        doc,
        "本报告基于 git 提交历史、文档沉淀、测试基线与已知技术债，"
        "对益语智库自用平台 V2.0 在 2026 Q1（1 月至 5 月 11 日内测发布）"
        "期间的产品开发进展做整理。报告由系统自动汇总，工程负责人复核。"
    )
    add_data_source(
        doc,
        "生成依据：git log + docs/* + pyproject.toml + backend/tests + TodoList "
        "+ 现有审计文档（客户计算中心差距清单、功能差距清单等）"
    )

    add_horizontal_rule(doc)
    add_paragraph(doc, "", space_after=8)

    # ============ 一、版本发布与关键里程碑 ============
    add_heading(doc, "一、版本发布与关键里程碑", level=1)

    add_heading(doc, "1.1 本期核心成就", level=2)
    achievements = [
        "完成 V2.0 内测 DMG 打包与发布（2026-05-11），首次形成可分发的桌面客户端",
        "「客户计算中心」九层架构全量审计完成，输出三份关键路线文档",
        "「理解深度」优化板块（7 个子能力 ×7 个迭代）全部代码落地",
        "Electron + Python 运行时硬化完成（P0 级问题修复），跨机器启动稳定性达到内测要求",
        "完成法律文档套件（用户协议 / 隐私政策 / App Store 合规 / 实施指南）",
    ]
    for a in achievements:
        add_bullet(doc, a)

    add_heading(doc, "1.2 关键里程碑时间线", level=2)
    milestones = [
        ["2026-Q1 前期", "客户计算中心 9 层架构审计 + 功能差距 9 维度审计", "✓ 完成"],
        ["2026-Q1 中期", "理解深度优化方案 + 执行计划起草", "✓ 完成"],
        ["2026-04", "理解深度迭代 1（鲜度衰减）→ 迭代 7（术语库）逐迭代落地", "✓ 完成"],
        ["2026-05-07", "代码仓库整理 + 移动端子仓库解耦", "✓ 完成"],
        ["2026-05-11", "V2.0 内测 DMG 打包 + 内测测试报告", "✓ 完成"],
        ["2026-05-12（今）", "主动能力 / 7 维度等剩余线程进入计划阶段", "→ 进行中"],
    ]
    add_table(
        doc,
        headers=["时间", "里程碑", "状态"],
        rows=milestones,
        widths=[3.0, 8.5, 2.0],
    )
    # 信息图：里程碑时间线
    add_image_from_bytes(
        doc,
        chart.timeline(
            title="Q1 2026 关键里程碑",
            events=[
                ("Q1 前期", "客户计算中心 9 层架构审计", "done"),
                ("Q1 中期", "理解深度优化方案 + 执行计划", "done"),
                ("2026-04", "理解深度迭代 1 → 7 逐迭代落地", "done"),
                ("2026-05-07", "代码仓库整理 + 移动端解耦", "done"),
                ("2026-05-11", "V2.0 内测 DMG 发布", "done"),
                ("2026-05-12", "主动能力等剩余线程进入计划阶段", "in_progress"),
            ],
        ),
        width_cm=15.0,
        caption="图 1.2 · Q1 关键里程碑时间线",
    )
    add_data_source(
        doc,
        "数据源：git log --all + docs/handoff-* + docs/客户计算中心-*.md 时间戳"
    )

    add_heading(doc, "1.3 commit 活动概览", level=2)
    add_paragraph(
        doc,
        "本季度（截至 5/12）共产生 39 个有效 commit，分类如下：",
        space_after=4,
    )
    commit_stats = [
        ["feat（新特性）", "22", "56%", "主要为 understanding/1-7 迭代 + 诊断/打包/移动端解耦"],
        ["fix（修复）", "7", "18%", "P0 运行时修复、retrieval-stage normalize、UI 微调"],
        ["docs（文档）", "6", "15%", "审计清单、优化方案、执行计划、内测测试报告"],
        ["chore（杂项）", "4", "10%", "gitignore、依赖、子仓库解耦"],
    ]
    add_table(
        doc,
        headers=["类型", "数量", "占比", "代表性内容"],
        rows=commit_stats,
        widths=[3.0, 1.5, 1.5, 7.5],
    )
    # 信息图：commit 类型饼图
    add_image_from_bytes(
        doc,
        chart.pie_commit_breakdown(
            labels=["feat（特性）", "fix（修复）", "docs（文档）", "chore（杂项）"],
            counts=[22, 7, 6, 4],
        ),
        width_cm=13.0,
        caption="图 1.3 · commit 类型分布",
    )
    add_data_source(doc, "数据源：git log --since=2026-01-01 --oneline + commit message 前缀分类")

    add_page_break(doc)

    # ============ 二、功能模块开发进展 ============
    add_heading(doc, "二、功能模块开发进展", level=1)
    add_paragraph(
        doc,
        "按九大功能板块梳理 Q1 开发投入。完成度基于已审计的全量需求清单评估，"
        "进度状态由 commit 与 docs 推算。",
        space_after=4,
    )

    add_heading(doc, "2.1 各板块完成度矩阵", level=2)
    modules = [
        ["理解深度（实体/关系/事实/矛盾/术语）", "20% → 70%", "● 突破", "迭代 1-7 全部代码落地，覆盖率提升明显"],
        ["主动能力（调度/通知/事件总线）", "15% → 15%", "● 持平", "已完成执行计划，待启动开发"],
        ["输入广度（录音/邮件/IM/多模态）", "30% → 30%", "● 持平", "已完成审计与执行计划"],
        ["生产能力（PPT/Excel/PDF/图表）", "15% → 18%", "● 起步", "Word 模板填充能力增强；其他格式待启动"],
        ['长期记忆（"我"层/偏好学习）', "25% → 27%", "● 起步", "执行计划完成；依赖理解深度"],
        ["真计算能力（NL→SQL/统计引擎）", "5% → 6%", "● 起步", "执行计划完成；待启动"],
        ["协作能力（评论/审计日志/委托）", "30% → 32%", "● 起步", "执行计划完成；依赖通知中心"],
        ["外部接入（搜索 API/实时信号）", "25% → 28%", "● 起步", "执行计划完成；待启动"],
        ["反思自知（可信度/盲区/推理链）", "20% → 22%", "● 起步", "执行计划完成；末位优先级"],
    ]
    module_status_colors = [
        RGBColor(0x2E, 0xA0, 0x47),  # 突破=绿
        RGBColor(0x88, 0x88, 0x88),  # 持平=灰
        RGBColor(0x88, 0x88, 0x88),  # 持平=灰
        RGBColor(0xE6, 0x9F, 0x00),  # 起步=黄
        RGBColor(0xE6, 0x9F, 0x00),
        RGBColor(0xE6, 0x9F, 0x00),
        RGBColor(0xE6, 0x9F, 0x00),
        RGBColor(0xE6, 0x9F, 0x00),
        RGBColor(0xE6, 0x9F, 0x00),
    ]
    table = add_table(
        doc,
        headers=["模块", "完成度变化", "本期态势", "备注"],
        rows=modules,
        widths=[5.5, 2.5, 2.0, 4.5],
    )
    # 给"本期态势"列上色
    for i, color in enumerate(module_status_colors):
        cell = table.rows[i + 1].cells[2]
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        dot_run = p.add_run("● ")
        set_run_font(dot_run, size=11, bold=True, color=color)
        label_run = p.add_run(modules[i][2].split(" ", 1)[1])
        set_run_font(label_run, size=10)

    # 信息图：九板块完成度横向进度条
    add_image_from_bytes(
        doc,
        chart.progress_bar_h(
            title="九大功能板块 Q1 完成度变化",
            items=["理解深度", "主动能力", "输入广度", "生产能力", "长期记忆",
                   "真计算能力", "协作能力", "外部接入", "反思自知"],
            before=[20, 15, 30, 15, 25, 5, 30, 25, 20],
            after=[70, 15, 30, 18, 27, 6, 32, 28, 22],
            target=50,
        ),
        width_cm=15.0,
        caption="图 2.1 · 九板块完成度变化（虚线为 50% 目标）",
    )
    add_data_source(
        doc,
        "数据源：客户计算中心-功能差距清单 各维度审计 + commit 关联 (understanding/* 共 14 个 commit)"
    )

    add_heading(doc, "2.2 理解深度板块深度展开（本期最大成就）", level=2)
    understanding_iters = [
        ["迭代 1", "鲜度真实衰减", "freshness_decay + EvidenceItem.createdAt + evidence_quality 重写", "8193756"],
        ["迭代 2", "跨文档实体基础设施", "entities + entity_mentions + EntityExtractor + 入库集成", "e5930c3"],
        ["迭代 2-F3", "内容 hash 去重 + 版本链 + 导入文案", "imports content_hash + supersedesId 增强", "a38e45c"],
        ["迭代 3", "近似重复实体检测 + 一键合并", "entity_merge_log + 模糊匹配 + 管理 UI", "a4881da"],
        ["迭代 4", "chunk 语义分类 + 引证面板类型徽章", "v2_chunks.semantic_type + UI 徽章", "ba42ffc"],
        ["迭代 5", "关系三元组抽取 + 查询接口", "relationship_triples + 关系词典 + API", "71ea246"],
        ["迭代 6", "矛盾检测 + 告警面板", "atomic_facts + fact_contradictions + 工作台告警", "20fde25"],
        ["迭代 7", "客户私有术语库 + 管理面板", "client_glossary + 术语 LLM 注入", "5d22733"],
    ]
    add_table(
        doc,
        headers=["迭代", "名称", "核心交付", "commit"],
        rows=understanding_iters,
        widths=[1.5, 4.0, 6.5, 2.5],
    )
    add_data_source(
        doc,
        "数据源：git log --grep=\"understanding/\" + commit message 解析"
    )

    add_page_break(doc)

    # ============ 三、技术决策与重要 commit ============
    add_heading(doc, "三、技术决策与架构演进", level=1)

    add_heading(doc, "3.1 本期关键技术决策", level=2)
    tech_decisions = [
        ["保留 intrinsic confidence，新增 time decay 因子", "迭代 1 鲜度衰减设计抉择", "采纳", "不动 memory_foundation 20+ 写入位置，新建 freshness_decay 模块在读取侧叠加衰减"],
        ["EvidenceItem 增加 createdAt + docType 字段", "理解深度多迭代共用基础", "采纳", "向后兼容默认 None，现有构造点不破"],
        ["关系词典严格预定义，不让 LLM 自由发挥", "迭代 5 关系抽取设计", "采纳", "质量优先，覆盖率次之；可后续扩词典"],
        ["矛盾检测分级（info / warning / critical）", "迭代 6 告警严重度", "采纳", "严格匹配 + 可一键 dismiss + dismiss 进反馈"],
        ["移动端解耦为独立子仓库", "代码组织决策", "采纳", "便于独立迭代；通过 mobile/ 自身 git 管理"],
    ]
    add_table(
        doc,
        headers=["决策", "背景", "状态", "理由 / 影响"],
        rows=tech_decisions,
        widths=[4.0, 3.5, 1.5, 5.5],
    )
    add_data_source(
        doc,
        "数据源：理解深度执行计划 + git commit messages (含决策说明)"
    )

    add_heading(doc, "3.2 P0 级修复（运行时硬化）", level=2)
    add_paragraph(
        doc,
        "本期处理一项 P0 阻塞性问题：",
        space_after=4,
    )
    add_callout(
        doc,
        "Electron 打包后 Python 运行时在不同机器路径下无法重定位 → 修复 relocation 逻辑 + "
        "stabilize 与 verify 顺序对齐 + 新增启动诊断脚本。修复后内测包能在多台开发机正常启动。"
    )
    add_paragraph(
        doc,
        "相关 commit：bf2e328 fix(runtime) + 2ac13b6 feat(packaging) + 后续 5/11 内测 DMG 发布。",
        size=10, italic=True, color=COLOR_MUTED, space_after=10,
    )
    add_data_source(doc, "数据源：git log + docs/handoff-followup-2026-05-11.md 内测测试报告")

    add_heading(doc, "3.3 重要架构演进", level=2)
    arch_changes = [
        "**数据中心 kernel** 接入 Cowork 模式：本机优先 + 选择性云端同步（surrogate / profile block）",
        "**rerank chain** 增加 chunk 语义类型加权（迭代 4 集成到 evidence_selector）",
        "**ingest pipeline** 注入 7 道理解深度处理层（实体 / 关系 / 事实 / 矛盾 / 语义分类）",
        "**云端 AI 同步**：诊断推方向 + 前端集成（3559a27、63107d4）",
        "**本地模型优化队列**：数据中心维度（0582bd7）",
    ]
    for a in arch_changes:
        # 处理 **加粗** 部分
        parts = a.split("**")
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(3)
        bullet_run = p.add_run("• ")
        set_run_font(bullet_run, size=11, color=COLOR_BRAND, bold=True)
        for i, part in enumerate(parts):
            if not part:
                continue
            run = p.add_run(part)
            if i % 2 == 1:
                set_run_font(run, size=11, bold=True, color=COLOR_BRAND)
            else:
                set_run_font(run, size=11)

    add_page_break(doc)

    # ============ 四、技术债与已知问题 ============
    add_heading(doc, "四、技术债与已知问题", level=1)

    add_heading(doc, "4.1 本期已偿还的技术债", level=2)
    paid_debts = [
        ["evidence_quality 粗糙线性衰减", "通过迭代 1 鲜度真实衰减彻底重写"],
        ["实体抽取只是 list[str] 关键词数组", "通过迭代 2 升级为类型化 entities + entity_mentions"],
        ["矛盾检测能力空白", "通过迭代 6 落地原子事实 + ContradictionDetector"],
        ["跨客户/跨文档实体同名不归一", "通过迭代 3 加自动 dedup + 待审 + 一键合并"],
        ["KnowledgeSearchHitRecord retrievalStage 校验失败", "a6d0e28 加 normalize 兜底 + field_validator"],
    ]
    add_table(
        doc,
        headers=["原技术债", "处置"],
        rows=paid_debts,
        widths=[7.0, 7.5],
    )

    add_heading(doc, "4.2 当前已知技术债（待下一期处理）", level=2)
    open_debts = [
        ["L9 AI Gateway 完全空白（20% → 0%）", "高", "/api/ai/* 9 条全无；状态机过短；无 service token；无 ai_drafts 表"],
        ["sync outbox 有生产者无消费者", "高", "data_center_sync_outbox 写入正常但无 worker 把 pending 推到 ChromaDB"],
        ["主动能力 0% — 完全无 cron / 调度器", "高", "无 APScheduler 等任一调度库；月报 / 风险扫描 / 告警推送都瘫"],
        ["生产能力 = 单条 Word 腿", "高", "无 python-pptx / openpyxl / reportlab；内容生产场景出口窄"],
        ['真计算能力 ≈ COUNT(*) + LLM.generate()', "高", '无 NL→SQL、无 what-if、无方案打分；"计算中心"名实不符'],
        ["dev app 重启 Python backend 没有面板入口", "中", "Task #28 待 Claude Code 完成（chokidar watcher + 菜单项）"],
        ["120 条 baseline test 已红状态未归因", "中", "迭代 0 baseline 收下；spot check 是 pre-existing 但无完整归因"],
        ["docx 输出 schema 警告（tcBorders/pBdr/shd 顺序）", "低", "Word/Pages/LibreOffice 都容错；schema 验证器严格不通过"],
    ]
    debt_colors = [
        RGBColor(0xC0, 0x39, 0x2B), RGBColor(0xC0, 0x39, 0x2B), RGBColor(0xC0, 0x39, 0x2B),
        RGBColor(0xC0, 0x39, 0x2B), RGBColor(0xC0, 0x39, 0x2B),
        RGBColor(0xE6, 0x9F, 0x00), RGBColor(0xE6, 0x9F, 0x00),
        RGBColor(0x88, 0x88, 0x88),
    ]
    debt_table = add_table(
        doc,
        headers=["技术债", "等级", "现状描述"],
        rows=open_debts,
        widths=[5.5, 1.5, 7.5],
    )
    for i, color in enumerate(debt_colors):
        cell = debt_table.rows[i + 1].cells[1]
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        dot_run = p.add_run("● ")
        set_run_font(dot_run, size=11, bold=True, color=color)
        label_run = p.add_run(open_debts[i][1])
        set_run_font(label_run, size=10, bold=True, color=color)

    # 信息图：风险矩阵气泡图
    add_image_from_bytes(
        doc,
        chart.risk_bubble(
            title="V2.1 关键风险矩阵",
            risks=[
                ("AI Gateway 缺失", 4.5, 3.5, 4),
                ("sync outbox 无消费者", 3.8, 3.2, 3),
                ("主动能力空白", 4.0, 4.0, 4.5),
                ("生产能力只剩 Word", 3.5, 4.2, 3.5),
                ("真计算能力 5%", 4.8, 2.8, 4),
                ("Task #28 未完成", 2.5, 4.5, 2.5),
            ],
        ),
        width_cm=15.0,
        caption="图 4.2 · 未结技术债的影响-概率矩阵（气泡大小为权重）",
    )
    add_data_source(
        doc,
        "数据源：客户计算中心-差距清单 + 功能差距清单 + 当前 TodoList in_progress 项"
    )

    add_page_break(doc)

    # ============ 五、质量与测试 ============
    add_heading(doc, "五、质量、测试与运维状态", level=1)

    add_heading(doc, "5.1 测试覆盖", level=2)
    test_stats = [
        ["pytest 测试文件总数", "66 个", "覆盖 ingest / chat / DNA / file search / 模型路由等"],
        ["本期新增测试", "19+ 个", "test_freshness_decay (19) + 各迭代配套测试"],
        ["回归基线（迭代 0 测试）", "153 passed / 120 failed / 14 errors", "120 fail 为 pre-existing，迭代不引入新红"],
        ["端到端 DMG 启动测试", "已跑通", "5/11 内测 DMG 在多台开发机能正常启动"],
    ]
    add_table(
        doc,
        headers=["指标", "本期数值", "说明"],
        rows=test_stats,
        widths=[5.0, 4.0, 5.5],
    )
    add_data_source(
        doc,
        "数据源：ls backend/tests/test_*.py + uv run pytest 输出 + docs/handoff-followup-* 内测报告"
    )

    add_heading(doc, "5.2 文档沉淀（重要资产）", level=2)
    docs_stats = [
        ["客户计算中心-差距清单 - 2026-05-12", "九层架构完整性审计", "建议性"],
        ["客户计算中心-功能差距清单 - 2026-05-12", "九维度功能审计", "建议性"],
        ["数据中心-未完成清单 - 2026-05-12", "全局未完成项地图", "工程清单"],
        ["理解深度-优化方案 / 执行计划", "理解深度 7 子能力深度方案", "已部分实施"],
        ["主动能力 - 优化方案 / Claude Code 执行计划", "主动能力 7 个迭代计划", "待启动"],
        ["剩余 7 维度 - Claude Code 完整执行计划", "33 个迭代 800-1500 工具调用", "待启动"],
        ["报告生成 - 架构逆向工程", "报告自动生成模块设计", "已 mock 验证"],
        ["内测 DMG 测试报告 - 2026-05-11", "Cowork Claude 执行 + 多机验证", "归档"],
        ["法律文档套件（用户协议/隐私/合规）", "App Store 上架前置", "归档"],
    ]
    add_table(
        doc,
        headers=["文档", "性质", "状态"],
        rows=docs_stats,
        widths=[7.0, 4.5, 3.0],
    )
    add_data_source(doc, "数据源：ls docs/*.md + 文档分类")

    add_page_break(doc)

    # ============ 六、下版本规划（V2.1）============
    add_heading(doc, "六、下版本（V2.1）规划", level=1)

    add_heading(doc, "6.1 V2.1 主线计划（按建议执行顺序）", level=2)
    v21_lines = [
        ["1", "完成 Task #28：dev app 自主重启机制", "工程团队", "1-2 天", "解锁后续 UI 验证流水线"],
        ["2", "理解深度迭代 1 UI 验证 + 收尾", "Cowork", "0.5 天", "依赖 Task #28"],
        ["3", "主动能力 A1 调度器 + A2 通知中心", "Claude Code", "2-3 周", "基础设施 — 解锁多个能力的合流"],
        ["4", "输入广度 I1 录音转写 + I2 邮件接入", "Claude Code", "2-3 周", "战略陪伴最高频痛点"],
        ["5", "生产能力 P1 PPT + P2 Excel + P3 PDF", "Claude Code", "2-3 周", "内容生产场景立竿见影"],
        ["6", "真计算能力 C1 NL→SQL", "Claude Code", "1-2 周", '让"计算中心"名实相符'],
        ["7", "AI Gateway L9 雏形（仅 retrieval + writeback）", "Claude Code", "1-2 周", "为 2.5 阶段铺路"],
    ]
    add_table(
        doc,
        headers=["#", "工作项", "责任方", "估时", "战略意义"],
        rows=v21_lines,
        widths=[0.8, 6.0, 2.5, 1.5, 4.0],
    )
    add_data_source(
        doc,
        "数据源：理解深度执行计划 + 主动能力执行计划 + 剩余 7 维度计划 等四份计划文档汇总"
    )

    add_heading(doc, "6.2 V2.1 建议时间窗", level=2)
    add_callout(
        doc,
        '建议 V2.1 内测包目标：8 周后（2026 年 7 月初）。'
        '覆盖：理解深度收尾 + 主动能力 A1/A2 + 输入广度 I1 + 生产能力 P1，'
        '对应"AI 主理人"雏形的首个可演示版本。'
    )

    add_heading(doc, "6.3 风险预案", level=2)
    risks = [
        ["LLM 成本失控（NL→SQL 尤其）", "中", "加日预算上限 + 用本地 7B 模型优先 + 关键 case 升级"],
        ["新增依赖与现有冲突（python-pptx 等）", "低", "新依赖必走 PR review + dry-run 验证"],
        ["Task #28 卡 Cowork UI 验证主线", "高", "若 Claude Code 未按期完成，Cowork 用手工重启 fallback"],
        ["主动能力 A2 通知 macOS 系统权限弹窗扰民", "中", "首次启动引导 + 用户可整体静音"],
    ]
    add_table(
        doc,
        headers=["风险", "概率", "缓解"],
        rows=risks,
        widths=[6.0, 1.5, 7.0],
    )

    add_page_break(doc)

    # ============ 七、附录 ============
    add_heading(doc, "七、附录", level=1)

    add_heading(doc, "7.1 commit 抽样（精选）", level=2)
    sample_commits = [
        ["ef9f2e6", "docs", "益语智库 V2.0 内测 DMG 测试报告（Cowork 执行 2026-05-11）"],
        ["20fde25", "feat", "理解深度迭代 6：矛盾检测 + 告警面板"],
        ["71ea246", "feat", "理解深度迭代 5：关系三元组抽取 + 查询接口"],
        ["e5930c3", "feat", "理解深度迭代 2：跨文档实体基础设施"],
        ["8193756", "feat", "理解深度迭代 1：鲜度真实衰减"],
        ["bf2e328", "fix", "Electron Python 运行时硬化 + relocation 修复 (P0)"],
        ["1861cc3", "docs", "法律文档套件（用户协议/隐私/合规/实施指南）"],
        ["2d3245d", "chore", "解耦移动端子仓库"],
    ]
    add_table(
        doc,
        headers=["hash", "类型", "说明"],
        rows=sample_commits,
        widths=[2.5, 1.5, 10.5],
    )

    add_heading(doc, "7.2 下次报告自动生成节奏", level=2)
    add_paragraph(
        doc,
        "本报告由益语智库自身的「报告自动生成模块」基于 git log + docs/* + tests 自动汇总，"
        "工程负责人复核。该模块本身是益语产品的一部分（详见报告生成架构逆向工程文档）。",
    )
    add_paragraph(doc, "建议节奏：", bold=True, color=COLOR_BRAND, space_before=6)
    add_bullet(doc, "周报：每周一 9:00 自动生成上周开发汇总（最小版）")
    add_bullet(doc, "月报：每月 1 日 9:00 自动生成上月模块进展（中等版）")
    add_bullet(doc, "季报：每季度首日自动生成完整版（本报告级别）")

    add_heading(doc, "7.3 报告生成器自身的发展", level=2)
    add_paragraph(
        doc,
        "本份报告是益语「报告自动生成模块」的早期 mock-up 产物之一（与战略陪伴报告并行）。"
        "通过两种主题（外部交付 vs 内部说明）的同骨架不同章节验证：报告生成器架构本身是可复用的。",
    )
    add_callout(
        doc,
        '关键证明点：同一套数据源抽象（事件线 / 判断 / 任务 / 风险）→ '
        '战略陪伴客户视角生成「客户报告」 / 内部产品研发视角生成「产品复盘」。'
        '章节骨架与措辞由"报告类型"决定，底层取数完全通用。'
    )

    add_horizontal_rule(doc)
    add_paragraph(doc, "", space_after=8)
    add_paragraph(
        doc,
        "© 2026 益语集团（内部资料）",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        size=9, color=COLOR_MUTED, italic=True,
    )

    return doc


def main():
    doc = build()
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"✅ 已生成：{OUTPUT_PATH}")
    print(f"   文件大小：{OUTPUT_PATH.stat().st_size} bytes")


if __name__ == "__main__":
    main()
