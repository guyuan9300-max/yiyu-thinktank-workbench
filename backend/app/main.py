from __future__ import annotations

import hashlib
import html
import json
import os
import re
import shutil
import time
from concurrent.futures import ThreadPoolExecutor
from dataclasses import dataclass
from datetime import datetime, timedelta
from pathlib import Path
from threading import Event, Thread
from time import perf_counter
from typing import Callable, Literal
from urllib.parse import quote, urlparse, urlunparse
from uuid import uuid4

import httpx
from fastapi import Body, FastAPI, File, Form, HTTPException, Query, Request, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import HTMLResponse, JSONResponse, Response
from docx import Document as WordDocument

from app.db import Database, from_json, to_json
from app.local_request_guard import ALLOWED_LOCAL_ORIGINS, ALLOWED_LOCAL_ORIGIN_REGEX, validate_local_browser_request
from app.models import (
    ActivityLogRecord,
    AgentWeeklyPlanPayload,
    AgentWeeklyPlanRecord,
    AgentWorklogResponse,
    AiTagSuggestionPayload,
    AiTagSuggestionResponse,
    AiStructuredResponse,
    AnalysisRunPayload,
    AnalysisRunRecord,
    AnalysisWorkbenchSettingsPayload,
    AnalysisWorkbenchSettingsRecord,
    AnalysisTemplateRecord,
    AnalysisToolsResponse,
    AuthLoginPayload,
    AuthRegisterPayload,
    AuthStateResponse,
    AmbiguityItem,
    AppSettingsPayload,
    AppSettingsResponse,
    BackupResponse,
    BadgeBoardResponse,
    ChatMessageRecord,
    ChatRequest,
    ChatStartResponse,
    ChatThreadDetailResponse,
    ChatThread,
    ClarificationAnswerPayload,
    ClarificationCreatePayload,
    ClarificationRecord,
    ConsultationKnowledgeProcessSummaryResponse,
    ConsultationKnowledgeRequestRecord,
    ClientDnaGeneratePayload,
    ClientDnaModuleRecord,
    ClientDnaModulesResponse,
    ClientAnalysisEvidenceSummaryRecord,
    ClientAnalysisRunRecord,
    ClientFolder,
    ClientMutationPayload,
    ClientSummary,
    ClientTemplateFillFieldRecord,
    ClientTemplateFillPayload,
    ClientTemplateFillRunRecord,
    ClientTemplateFillResponse,
    ClientTextDocumentPayload,
    ClientTextDocumentResponse,
    ClientNotebookResponse,
    ClientWorkspaceResponse,
    ClientWorkspaceSettingsPayload,
    ClientWorkspaceSettingsRecord,
    DecisionItem,
    DemoDataResponse,
    DnaTerm,
    DnaTermPayload,
    DocumentCardRecord,
    DocumentRecord,
    EvidenceItem,
    ExportAnswerPayload,
    FileReclassEventRecord,
    GoalPayload,
    GoalRecord,
    GrowthContextLinkRecord,
    GrowthAfterActionCaptureRecord,
    GrowthGenericLessonRecord,
    GrowthLedgerResponse,
    GrowthOverviewRecord,
    GrowthPendingCaptureActionPayload,
    GrowthPendingCaptureActionResponse,
    GrowthProjectGuidanceRecord,
    GrowthReasoningInputRecord,
    GrowthReasoningTraceRecord,
    GrowthRecommendationActionResponse,
    GrowthRecommendationDismissPayload,
    GrowthActionPlanItemRecord,
    GrowthMaterialRefRecord,
    GrowthProjectContextPackRecord,
    GrowthLearningSummaryRecord,
    GrowthRobotAssistRecord,
    GrowthTaskIntentRecord,
    GrowthUniversalSkillItemRecord,
    GrowthValidationActionResponse,
    GrowthValidationPayload,
    GrowthWorkbenchActionRecord,
    GrowthWorkbenchMaterialRecord,
    GrowthWorkbenchSnapshotRecord,
    GrowthWorkbenchStepRecord,
    GrowthWorkbenchSupportCopyRecord,
    GrowthWorkbenchTaskRecord,
    HandbookEntryRecord,
    HandbookEntryDetailRecord,
    HandbookReuseRecord,
    HandbookPayload,
    HandbookSettingsPayload,
    HandbookSettingsRecord,
    HandbookResponse,
    HealthAiState,
    HealthResponse,
    ImportPayload,
    ImportRecord,
    WorkspaceImportBackfillResponse,
    LegacyScanEntry,
    LegacyScanRequest,
    LegacyScanResponse,
    KnowledgeJobRecord,
    KnowledgeSearchHitRecord,
    KnowledgeSearchResponse,
    KnowledgeStatusRecord,
    MemoryStatus,
    DepartmentOptionRecord,
    MentionCandidateRecord,
    MeetingCreatePayload,
    MeetingDetail,
    MeetingIngestPayload,
    MeetingPipelineResponse,
    MeetingSummary,
    EmployeeDepartmentPayload,
    EmployeeRecord,
    EmployeeRejectPayload,
    EmployeeRolePayload,
    FeishuBotSettingsPayload,
    FeishuBotSettingsRecord,
    FeishuMeetingLaunchPayload,
    FeishuMeetingLaunchResponse,
    FeishuReceiveIdType,
    FeishuUserBindingRecord,
    FeishuUserBindingStartResponse,
    DnaReadinessQuestionRecord,
    EventLineActivityRecord,
    EventLineClarificationDraftPayload,
    EventLineClarificationDraftRecord,
    EventLineContextBundleRecord,
    EventLineContextFactRecord,
    EventLineCreatePayload,
    EventLineDetailRecord,
    EventLineJudgmentRecord,
    EventLineOpportunityCardRecord,
    EventLineMemoryResponse,
    EventLineRecord,
    EventLineRiskCardRecord,
    EventLineUpdatePayload,
    OperatorRecord,
    OrgDepartmentRecord,
    OrgEmployeeBindingRecord,
    OrgModelProfileRecord,
    OrgProfileRecord,
    OrgReportingLineRecord,
    OrgRoleTemplateRecord,
    OrgTaskControlRuleRecord,
    OrganizationDnaModuleRecord,
    OrganizationDnaResponse,
    OrganizationDnaUploadPayload,
    OrganizationNotebookSnapshot,
    ProjectFlowPayload,
    ProjectFlowRecord,
    ProjectFlowDetailRecord,
    ProjectModulePayload,
    ProjectModuleRecord,
    ProjectModuleDetailRecord,
    ProjectStructureResponse,
    ReviewDepartmentConfigRecord,
    ReviewDepartmentMemberRecord,
    ReviewDashboardCardTargetRecord,
    ReviewDashboardEvidenceRefRecord,
    ReviewDashboardDrillTargetResponse,
    ReviewGovernanceSettingsPayload,
    ReviewGovernanceSettingsRecord,
    ReviewHistoryEntryRecord,
    ReviewHistoryResponse,
    ReviewResponse,
    ReviewSimulationBundleRecord,
    RiskItem,
    SessionUserRecord,
    SettingsResponse,
    SystemAdminSettingsPayload,
    SystemAdminSettingsRecord,
    TaskActivityRecord,
    TaskAttachmentRecord,
    TaskBoardResponse,
    TaskCollaboratorRecord,
    TaskCompletionReviewPayload,
    TaskContextPreviewRecord,
    TaskListLibraryResponse,
    TaskListMutationPayload,
    TaskListRecord,
    MemoryBackfillResultRecord,
    NarrativeAnalysisRecord,
    TaskNotePayload,
    TaskContextRefreshResultRecord,
    TaskEventLineBootstrapResultRecord,
    TaskOrgBackfillResultRecord,
    TaskOrgContextRecord,
    TaskProjectContextRecord,
    TaskPlanLinkRecord,
    TaskPlanLinkUpsertPayload,
    TaskPayload,
    TaskRecord,
    TaskRejectPayload,
    TaskSettingsPayload,
    TaskSettingsRecord,
    TaskViewDefinitionRecord,
    TaskViewMutationPayload,
    TaskViewPresetRecord,
    TaskViewsResponse,
    SupportRequestCreatePayload,
    SupportRequestRecord,
    SupportRequestResolvePayload,
    TaskTagLibraryResponse,
    TaskTagMutationPayload,
    TaskTagRecord,
    TaskUpdatePayload,
    TitleSuggestionResponse,
    TopicRadarAssistPayload,
    TopicRadarAssistResponse,
    TopicRadarPreferredSourceRecord,
    TopicRadarSourceLabelPayload,
    TopicRadarSourceLabelResponse,
    TopicCandidatePayload,
    TopicCandidateChatPayload,
    TopicCandidateChatResponse,
    TopicCandidateChatMessageRecord,
    TopicCandidateInsightRecord,
    TopicCandidateRecord,
    TopicCaptureBatchResponse,
    TopicCaptureRunRecord,
    TopicTaskDraftPayload,
    TopicTaskPlanResponse,
    TopicTaskPromotionPayload,
    TopicTaskPromotionResponse,
    TopicTaskSuggestionRecord,
    TopicsSettingsPayload,
    TopicsSettingsRecord,
    TopicTitlePayload,
    TopicRadarPayload,
    TopicRadarRecord,
    TopicsResponse,
    StrategicAssetCandidateRecord,
    StrategicChangePointRecord,
    StrategicChecklistGroupRecord,
    StrategicChecklistItemRecord,
    StrategicCockpitConfirmPayload,
    StrategicCockpitSnapshotRecord,
    StrategicEvidenceCardRecord,
    StrategicEvidencePreviewRecord,
    StrategicHealthLineRecord,
    StrategicHeadlineRecord,
    StrategicJudgmentRecord,
    StrategicMeetingPackDraftRecord,
    StrategicPermissionRecord,
    StrategicReadinessRecord,
    StrategicLineRecord,
    StrategicLineDetailRecord,
    VectorizeAnswerPayload,
    WeeklyReviewAnalysisRecord,
    WeeklyReviewEventLineContextRecord,
    WeeklyReviewPayload,
    WeeklyReviewTaskEntryRecord,
    WeeklyReviewRecord,
    WeeklyReviewTaskStructuredNoteRecord,
    TrendSignalRecord,
    AgendaItem,
    KnowledgeMemoryRecord,
    LearningRecommendationRecord,
)
from app.services.ai import AiInvocationError, AiService, DEFAULT_MODEL, DEFAULT_PROVIDER
from app.services.agent_worklogs import (
    AGENT_AUTO_SOURCE_TYPE,
    build_agent_execution_task_activity,
    build_agent_execution_tasks,
    build_agent_weekly_digests,
    build_agent_weekly_plans,
    build_agent_weekly_review_items,
    build_agent_worklog_response,
    sync_agent_execution_tasks,
    upsert_agent_weekly_plan_override,
)
from app.services.department_catalog import get_department_entry, list_department_catalog
from app.services.knowledge_base import (
    create_memory_surrogate_from_answer,
    ensure_source_tree_snapshot,
    ensure_client_workspace,
    fetch_recent_knowledge_jobs,
    fetch_recent_reclass_events,
    is_finance_priority_text,
    is_finance_query,
    is_finance_statement_priority_text,
    is_finance_statement_query,
    sync_master_index_fts,
    sync_qdrant_for_client,
)
from app.services.knowledge_v2 import (
    HUMAN_VISIBLE_CATEGORIES,
    V2_PIPELINE_VERSION,
    backfill_knowledge_documents,
    backfill_workspace_import,
    compute_knowledge_status,
    deserialize_retrieval_bundle,
    fetch_document_cards,
    ingest_document_knowledge,
    is_strategy_analysis_query,
    refresh_client_folder_counts,
    retrieve_knowledge_bundle,
    safe_filename,
    serialize_retrieval_bundle,
    stage_import_copy,
    tokenize,
)
from app.services.memory_foundation import (
    answer_clarification_record,
    backfill_memory_foundation,
    create_clarification_record,
    get_client_memory_status,
    get_client_notebook_response,
    get_event_line_memory_response,
    get_task_memory_enrichment,
    list_linked_event_lines,
    record_imported_document_writeback,
    record_client_dna_writeback,
    record_meeting_publish_writeback,
    record_task_attachment_writeback,
    record_task_writeback,
    record_weekly_review_writeback,
    refresh_event_line_memory_snapshot,
    refresh_organization_notebook_snapshot,
    sanitize_memory_background_text,
)
from app.services.platform_dna import extract_platform_dna_text
from app.services.template_fill import (
    TemplateWebSource,
    apply_docx_template_values,
    build_template_follow_up_question,
    build_template_fill_retrieval_query,
    build_template_suggested_sources,
    extract_docx_attachment_checklist,
    extract_docx_template_fields,
    fetch_template_fill_web_sources,
    infer_template_field_type,
    infer_template_value_kind,
    should_enable_template_fill_web_supplement,
)
from app.services.topic_capture import fetch_topic_candidates_from_web, fetch_topic_source_excerpt
from app.services.review_analysis import _dedupe_texts, _story_evidence_refs, build_weekly_review_analysis
from app.services.review_narrative import build_weekly_overview_draft
from app.services.review_rollup import build_employee_review_report, build_executive_review_rollup
from app.services.review_simulation import build_review_simulation_bundle
from app.services.feishu import (
    FeishuApiError,
    build_user_authorize_url,
    exchange_authorization_code,
    fetch_app_access_token,
    fetch_tenant_access_token,
    fetch_user_info,
    send_text_message,
)
from app.services.badge_engine import build_badge_board
from app.services.growth_engine import (
    backfill_handbook_entries,
    build_growth_ledger,
    build_generic_learning_fallback,
    build_growth_overview,
    get_pending_capture,
    ingest_meeting_growth_candidate,
    ingest_handbook_codification,
    ingest_strategic_growth_candidate,
    ingest_task_growth_candidate,
    ingest_review_growth,
    list_learning_recommendations,
    mark_handbook_entry_reused,
    mark_recommendation_accepted,
    mark_recommendation_dismissed,
    update_pending_capture_state,
)
from app.services.secrets import MacOSKeychainSecretStore, MemorySecretStore


APP_NAME = "益语智库自用平台"
APP_VERSION = "0.1.0"
APP_BUILD_VERSION = "2026.03.15-v2-migration-1"
PROJECT_ROOT = Path(__file__).resolve().parents[2]
THREAD_SYNC_DOC_PATH = PROJECT_ROOT / "docs" / "thread-sync.md"
SUPPORTED_IMPORT_EXTENSIONS = {".pdf", ".docx", ".md", ".txt", ".pptx", ".xlsx"}
LEGACY_IMPORT_EXTENSIONS = {".json", ".csv"}
DEMO_CLIENT_IDS = ("client_cffc", "client_star")
DEMO_THREAD_IDS = ("thread_cffc", "thread_star")
DEMO_GOAL_IDS = ("goal_1", "goal_2", "goal_3")
DEMO_DNA_IDS = ("dna_1", "dna_2")
DEMO_DOCUMENT_IDS = ("doc_1", "doc_2", "doc_3")
DEMO_TASK_IDS = ("task_seed_1", "task_seed_2")
DEMO_RADAR_IDS = ("radar_ai", "radar_fund")
DEMO_CANDIDATE_IDS = ("cand_1", "cand_2")
DEMO_HANDBOOK_IDS = ("hb_1",)
DEMO_CHAT_MESSAGE_IDS = ("msg_seed_1",)
BACKEND_FEATURE_FLAGS = [
    "knowledge.vectorize-answer",
    "knowledge.reclass-events",
    "knowledge.surrogate-index",
    "knowledge.search",
    "knowledge.rebuild",
    "chat.general-answer",
    "chat.instant-send",
    "chat.async-status",
    "chat.analysis-runs",
]

ORGANIZATION_DNA_MODULES = [
    ("organization_intro", "组织介绍"),
    ("business_intro", "业务介绍"),
    ("team_intro", "团队介绍"),
    ("market_intro", "市场介绍"),
]

CLIENT_DNA_MODULES = [
    ("organization_intro", "组织介绍"),
    ("business_intro", "项目介绍"),
    ("team_intro", "团队介绍"),
    ("market_intro", "市场背景介绍"),
]

SELF_CLIENT_NAME_CANDIDATES = ["益语智库", "益语"]

DNA_READINESS_RULES: dict[str, list[dict[str, object]]] = {
    "organization_intro": [
        {
            "question": "组织定位是否清楚",
            "contentKeywords": ["定位", "战略陪伴者", "成长合伙人", "可落地的增长咨询"],
            "missingKeywords": [],
        },
        {
            "question": "组织为什么存在、主要解决什么问题是否清楚",
            "contentKeywords": ["使命", "市场不确定性", "组织效率", "数字化焦虑", "穿越不确定期"],
            "missingKeywords": [],
        },
        {
            "question": "当前升级方向或阶段是否清楚",
            "contentKeywords": ["内部引擎", "应用共建", "学习加速", "升级", "转型"],
            "missingKeywords": [],
        },
    ],
    "business_intro": [
        {
            "question": "主要服务或交付内容是否清楚",
            "contentKeywords": ["增长咨询", "战略设计", "流程陪伴", "应用共建", "学习加速", "工作平台"],
            "missingKeywords": [],
        },
        {
            "question": "服务对象和价值是否清楚",
            "contentKeywords": ["适应性组织", "企业", "持续增长", "客户痛点", "解决方案"],
            "missingKeywords": [],
        },
        {
            "question": "当前业务重点或推进路径是否清楚",
            "contentKeywords": ["技术规划", "产品", "0.5", "3.0", "GPT 5.4", "升级"],
            "missingKeywords": [],
        },
    ],
    "team_intro": [
        {
            "question": "关键成员与角色是否清楚",
            "contentKeywords": ["成员", "负责人", "角色", "团队"],
            "missingKeywords": ["成员名单", "创始人", "负责人", "履历"],
        },
        {
            "question": "分工与协作结构是否清楚",
            "contentKeywords": ["分工", "协作", "组织架构", "接口"],
            "missingKeywords": ["组织架构", "分工", "接口"],
        },
        {
            "question": "当前团队能力重点是否清楚",
            "contentKeywords": ["AI 技术", "工作平台", "技术布局", "能力", "升级"],
            "missingKeywords": [],
        },
    ],
}

STRATEGIC_PLACEHOLDER_CONTEXT_PATTERNS = [
    "当前重点仍待补充",
    "建议先明确这一阶段的核心事项",
    "当前没有特别突出的阻塞",
    "仍需盯住推进收束",
    "下一步动作：先补齐项目背景",
    "最近进展仍待补充",
]

STRATEGIC_RELATIONSHIP_TASK_KEYWORDS = [
    "吃饭",
    "见面",
    "会面",
    "拜访",
    "约",
    "沟通",
    "介绍",
    "对接",
    "维护关系",
    "关系推进",
]

STRATEGIC_CONTEXTUAL_DESCRIPTION_KEYWORDS = [
    "机构",
    "基金会",
    "负责人",
    "背景",
    "合作",
    "关系",
    "推进",
    "目标",
    "希望",
    "这次",
    "本次",
    "当前",
    "下一步",
    "方案",
]

STRATEGIC_WEEKLY_MEETING_KEYWORDS = ["周会", "周盘点", "周例会", "战略", "经营", "复盘"]


def now_iso() -> str:
    return datetime.now().replace(microsecond=0).isoformat()


APP_STARTED_AT = now_iso()


def new_id(prefix: str) -> str:
    return f"{prefix}_{uuid4().hex[:10]}"


def today_label() -> str:
    return datetime.now().strftime("%m-%d")


def normalize_markdown_text(markdown_content: str) -> str:
    text = markdown_content.replace("\r\n", "\n")
    text = re.sub(r"```.*?```", " ", text, flags=re.S)
    text = re.sub(r"`([^`]*)`", r"\1", text)
    text = re.sub(r"!\[[^\]]*\]\([^)]+\)", " ", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"^#{1,6}\s*", "", text, flags=re.M)
    text = re.sub(r"^\s*[-*+]\s+", "", text, flags=re.M)
    text = re.sub(r"^\s*\d+\.\s+", "", text, flags=re.M)
    text = re.sub(r"[>*_~#|-]", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    return text.strip()


def summarize_markdown_document(title: str, normalized_text: str) -> str:
    clean = normalized_text.strip()
    if not clean:
        return f"{title}当前还没有可用内容，请重新上传完整 Markdown 文档。"
    paragraphs = [segment.strip() for segment in re.split(r"\n{2,}", clean) if segment.strip()]
    lead = paragraphs[0] if paragraphs else clean
    detail = paragraphs[1] if len(paragraphs) > 1 else ""
    summary = lead[:160]
    if detail:
        summary = f"{summary} {detail[:120]}".strip()
    if len(summary) < 80 and len(clean) > len(summary):
        summary = f"{summary} {clean[len(summary):len(summary) + 120]}".strip()
    return summary[:320]


@dataclass
class AppState:
    data_dir: Path
    backup_dir: Path
    db: Database
    ai: AiService
    feishu_secret_store: MacOSKeychainSecretStore | MemorySecretStore
    cloud_api_url: str
    job_stop: Event
    job_thread: Thread | None = None
    topic_insight_executor: ThreadPoolExecutor | None = None
    chat_answer_executor: ThreadPoolExecutor | None = None
    template_fill_executor: ThreadPoolExecutor | None = None
    volatile_cloud_access_token: str = ""
    volatile_cloud_refresh_token: str = ""
    volatile_cloud_session_user_json: str = ""
    cloud_session_persistent: bool = False
    consultation_knowledge_sync_running: bool = False


_runtime_state: AppState | None = None


def _require_runtime_state() -> AppState:
    if _runtime_state is None:
        raise RuntimeError("App state is not initialized")
    return _runtime_state


def _parse_json_list(value: str | None) -> list[str]:
    data = from_json(value, [])
    return [str(item) for item in data] if isinstance(data, list) else []


def _parse_date_only(value: str | None) -> datetime.date | None:
    if not value:
        return None
    try:
        return datetime.fromisoformat(value).date()
    except ValueError:
        return None


def _week_bounds(week_label: str) -> tuple[datetime.date, datetime.date] | None:
    match = re.match(r"^(\d{4})-W(\d{2})$", week_label.strip())
    if not match:
        return None
    year = int(match.group(1))
    week = int(match.group(2))
    try:
        start = datetime.fromisocalendar(year, week, 1).date()
    except ValueError:
        return None
    return start, start + timedelta(days=6)


def _task_review_date(task: TaskRecord) -> datetime.date | None:
    due_date = _parse_date_only(task.dueDate)
    if due_date:
        return due_date
    return _parse_date_only(task.createdAt)


def _task_in_week(task: TaskRecord, week_label: str) -> bool:
    bounds = _week_bounds(week_label)
    if not bounds:
        return False
    review_date = _task_review_date(task)
    if not review_date:
        return False
    start, end = bounds
    return start <= review_date <= end


    def _local_task_tag_record(row) -> TaskTagRecord:
        return TaskTagRecord(
            id=str(row["id"]),
            name=str(row["name"]),
            color=str(row["color"] or ("#9CA3AF" if str(row["scope"]) == "self" else "#5B7BFE")),
            scope=str(row["scope"]),
            ownerUserId=str(row["owner_operator_id"]) if str(row["owner_operator_id"]) else None,
            createdBy=str(row["created_by"]) if str(row["created_by"]) else None,
            updatedAt=str(row["updated_at"] or row["created_at"] or now_iso()),
            archivedAt=str(row["archived_at"]) if row["archived_at"] else None,
        )

    def _local_task_list_record(row) -> TaskListRecord:
        return TaskListRecord(
            id=str(row["id"]),
            name=str(row["name"]),
            color=str(row["color"]),
            sortOrder=int(row["sort_order"] or 0),
            isDefault=bool(int(row["is_default"] or 0)),
            scope=str(row["scope"] or "org"),
            archivedAt=str(row["archived_at"]) if row["archived_at"] else None,
        )

    def _local_default_list_id() -> str | None:
        default_row = state.db.fetchone("SELECT id FROM task_lists WHERE is_default = 1 ORDER BY sort_order ASC LIMIT 1")
        if default_row:
            return str(default_row["id"])
        first_row = state.db.fetchone("SELECT id FROM task_lists ORDER BY sort_order ASC, name COLLATE NOCASE ASC LIMIT 1")
        return str(first_row["id"]) if first_row else None

    def _default_local_task_settings() -> TaskSettingsRecord:
        return TaskSettingsRecord(
            defaultListId=_local_default_list_id(),
            defaultPriority="normal",
            defaultDueDatePreset="today",
            defaultViewMode="calendar",
            listSortMode="manual",
            showCompletedTasks=False,
            defaultReviewScope="work",
            autoAssignSelf=True,
            updatedAt=now_iso(),
        )

    def _local_task_settings_record(row) -> TaskSettingsRecord:
        defaults = _default_local_task_settings()
        return TaskSettingsRecord(
            defaultListId=str(row["default_list_id"]) if row["default_list_id"] else defaults.defaultListId,
            defaultPriority=str(row["default_priority"] or defaults.defaultPriority),  # type: ignore[arg-type]
            defaultDueDatePreset=str(row["default_due_date_preset"] or defaults.defaultDueDatePreset),  # type: ignore[arg-type]
            defaultViewMode=str(row["default_view_mode"] or defaults.defaultViewMode),  # type: ignore[arg-type]
            listSortMode=str(row["list_sort_mode"] or defaults.listSortMode),  # type: ignore[arg-type]
            showCompletedTasks=bool(int(row["show_completed_tasks"] or 0)),
            defaultReviewScope=str(row["default_review_scope"] or defaults.defaultReviewScope),  # type: ignore[arg-type]
            autoAssignSelf=bool(int(row["auto_assign_self"] if row["auto_assign_self"] is not None else 1)),
            updatedAt=str(row["updated_at"] or defaults.updatedAt),
        )

    def _get_local_task_settings(operator_id: str | None = None) -> TaskSettingsRecord:
        resolved_operator_id = operator_id or str(current_operator_row()["id"])
        row = state.db.fetchone("SELECT * FROM task_settings WHERE operator_id = ?", (resolved_operator_id,))
        if row:
            return _local_task_settings_record(row)
        defaults = _default_local_task_settings()
        state.db.execute(
            """
            INSERT OR REPLACE INTO task_settings(
                operator_id, default_list_id, default_priority, default_due_date_preset,
                default_view_mode, list_sort_mode, show_completed_tasks, default_review_scope,
                auto_assign_self, updated_at
            ) VALUES(?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                resolved_operator_id,
                defaults.defaultListId,
                defaults.defaultPriority,
                defaults.defaultDueDatePreset,
                defaults.defaultViewMode,
                defaults.listSortMode,
                1 if defaults.showCompletedTasks else 0,
                defaults.defaultReviewScope,
                1 if defaults.autoAssignSelf else 0,
                defaults.updatedAt,
            ),
        )
        return defaults

    def _visible_local_task_tag_rows(db: Database, operator_id: str) -> list:
        return db.fetchall(
            """
            SELECT *
            FROM task_tags
            WHERE scope = 'org' OR owner_operator_id = ?
            ORDER BY CASE WHEN archived_at IS NULL OR archived_at = '' THEN 0 ELSE 1 END,
                     CASE scope WHEN 'org' THEN 0 ELSE 1 END,
                     name COLLATE NOCASE ASC
            """,
            (operator_id,),
        )

    def _visible_local_task_tags(db: Database, operator_id: str) -> list[TaskTagRecord]:
        return [_local_task_tag_record(row) for row in _visible_local_task_tag_rows(db, operator_id)]


    def _visible_local_task_tag_rows(db: Database, operator_id: str) -> list:
        return db.fetchall(
            """
            SELECT *
            FROM task_tags
            WHERE scope = 'org' OR owner_operator_id = ?
            ORDER BY CASE WHEN archived_at IS NULL OR archived_at = '' THEN 0 ELSE 1 END,
                     CASE scope WHEN 'org' THEN 0 ELSE 1 END,
                     name COLLATE NOCASE ASC
            """,
            (operator_id,),
        )


def _visible_local_task_tags(db: Database, operator_id: str) -> list[TaskTagRecord]:
    return [_global_local_task_tag_record(row) for row in _visible_local_task_tag_rows(db, operator_id)]


def _local_tag_rows_by_ids(db: Database, tag_ids: list[str]) -> list:
    if not tag_ids:
        return []
    rows = db.fetchall(
        f"SELECT * FROM task_tags WHERE id IN ({_sql_placeholders(tag_ids)})",
        tuple(tag_ids),
    )
    by_id = {str(row["id"]): row for row in rows}
    return [by_id[tag_id] for tag_id in tag_ids if tag_id in by_id]


def _global_local_task_tag_record(row) -> TaskTagRecord:
    return TaskTagRecord(
        id=str(row["id"]),
        name=str(row["name"]),
        color=str(row["color"] or ("#9CA3AF" if str(row["scope"]) == "self" else "#5B7BFE")),
        scope=str(row["scope"]),
        ownerUserId=str(row["owner_operator_id"]) if str(row["owner_operator_id"]) else None,
        createdBy=str(row["created_by"]) if str(row["created_by"]) else None,
        updatedAt=str(row["updated_at"] or row["created_at"] or now_iso()),
        archivedAt=str(row["archived_at"]) if row["archived_at"] else None,
    )

def _ensure_local_tag(
    db: Database,
    operator_id: str,
    name: str,
    scope: str = "org",
    color: str | None = None,
    created_by: str | None = None,
) -> TaskTagRecord:
    trimmed = name.strip()
    if not trimmed:
        raise HTTPException(status_code=400, detail="Tag name is required")
    owner_operator_id = operator_id if scope == "self" else ""
    resolved_color = color or ("#9CA3AF" if scope == "self" else "#5B7BFE")
    operator_row = db.fetchone("SELECT name FROM operators WHERE id = ?", (operator_id,))
    resolved_creator = created_by or (str(operator_row["name"]) if operator_row and operator_row["name"] else "系统")
    existing = db.fetchone(
        "SELECT * FROM task_tags WHERE name = ? AND scope = ? AND owner_operator_id = ?",
        (trimmed, scope, owner_operator_id),
    )
    timestamp = now_iso()
    if existing:
        if not str(existing["updated_at"]) or not str(existing["color"] or ""):
            db.execute(
                """
                UPDATE task_tags
                SET updated_at = ?, created_at = COALESCE(NULLIF(created_at, ''), ?),
                    color = COALESCE(NULLIF(color, ''), ?), created_by = COALESCE(NULLIF(created_by, ''), ?)
                WHERE id = ?
                """,
                (timestamp, timestamp, resolved_color, resolved_creator, str(existing["id"])),
            )
            existing = db.fetchone("SELECT * FROM task_tags WHERE id = ?", (str(existing["id"]),))
        assert existing is not None
        return _global_local_task_tag_record(existing)
    tag_id = new_id("tag")
    db.execute(
        """
        INSERT INTO task_tags(id, name, scope, color, owner_operator_id, created_by, created_at, updated_at)
        VALUES(?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (tag_id, trimmed, scope, resolved_color, owner_operator_id, resolved_creator, timestamp, timestamp),
    )
    row = db.fetchone("SELECT * FROM task_tags WHERE id = ?", (tag_id,))
    assert row is not None
    return _global_local_task_tag_record(row)


def _resolve_local_task_tags(db: Database, operator_id: str, tag_ids: list[str], legacy_names: list[str]) -> list[TaskTagRecord]:
    rows = _local_tag_rows_by_ids(db, tag_ids)
    if rows:
        return [_global_local_task_tag_record(row) for row in rows]
    if not legacy_names:
        return []
    return [_ensure_local_tag(db, operator_id, name, "org") for name in legacy_names if name.strip()]


_event_line_snapshot_cache: dict[str, dict[str, object] | None] = {}


def _invalidate_event_line_snapshot_cache(*event_line_ids: str | None) -> None:
    for event_line_id in event_line_ids:
        normalized_id = str(event_line_id or "").strip()
        if normalized_id:
            _event_line_snapshot_cache.pop(normalized_id, None)


def _event_line_snapshot_context(
    db: Database | None,
    event_line_id: str | None,
    fallback_name: str | None = None,
    *,
    cloud_resolver: Callable[[str, str | None], dict[str, object] | None] | None = None,
) -> dict[str, object] | None:
    normalized_id = (event_line_id or "").strip()
    if normalized_id and normalized_id in _event_line_snapshot_cache:
        return _event_line_snapshot_cache[normalized_id]
    if not normalized_id and not (fallback_name or "").strip():
        return None

    context: dict[str, object] | None = None
    if normalized_id:
        row = db.fetchone("SELECT * FROM event_lines WHERE id = ?", (normalized_id,)) if db is not None else None
        if row is not None:
            activity_count = int(db.scalar("SELECT COUNT(1) FROM event_line_activities WHERE event_line_id = ?", (normalized_id,)) or 0) if db is not None else 0
            context = {
                "id": str(row["id"]),
                "name": str(row["name"]),
                "businessCategory": str(row["business_category"]) if row["business_category"] else None,
                "stage": str(row["stage"]) if row["stage"] else None,
                "summary": str(row["summary"]) if row["summary"] else None,
                "intent": str(row["intent"]) if row["intent"] else None,
                "currentBlocker": str(row["current_blocker"]) if row["current_blocker"] else None,
                "recentDecision": str(row["recent_decision"]) if row["recent_decision"] else None,
                "nextStep": str(row["next_step"]) if row["next_step"] else None,
                "evidenceCount": max(int(row["evidence_count"] or 0), activity_count),
                "primaryClientId": str(row["primary_client_id"]) if row["primary_client_id"] else None,
                "primaryClientName": str(row["primary_client_name"]) if row["primary_client_name"] else None,
            }
        elif cloud_resolver is not None:
            context = cloud_resolver(normalized_id, fallback_name)
    if context is None and (fallback_name or "").strip():
        context = {
            "id": normalized_id or None,
            "name": fallback_name.strip(),
        }
    if normalized_id:
        _event_line_snapshot_cache[normalized_id] = context
    return context


def _client_workspace_relative_path(path_value: str | None) -> Path | None:
    normalized_path = str(path_value or "").strip()
    if not normalized_path:
        return None
    try:
        path = Path(normalized_path)
        parts = list(path.parts)
        workspace_index = parts.index("client_workspace")
        if len(parts) <= workspace_index + 2:
            return None
        return Path(*parts[workspace_index + 2 :])
    except ValueError:
        return None


def _resolve_client_folder_ref_by_path(db: Database, client_id: str, path_value: str | None) -> tuple[str | None, str]:
    relative_path = _client_workspace_relative_path(path_value)
    folder_label = relative_path.parts[0] if relative_path and relative_path.parts else "项目与业务"
    folder_row = db.fetchone(
        "SELECT id FROM client_folders WHERE client_id = ? AND label = ?",
        (client_id, folder_label),
    )
    return (str(folder_row["id"]) if folder_row and folder_row["id"] else None, folder_label)


def _rehome_client_workspace_path(
    data_dir: Path,
    target_client_id: str,
    current_path: str | None,
    *,
    moved_paths: dict[str, str],
    fallback_label: str = "项目与业务",
) -> str | None:
    normalized_path = str(current_path or "").strip()
    if not normalized_path:
        return None
    if normalized_path in moved_paths:
        return moved_paths[normalized_path]

    ensure_client_workspace(data_dir, target_client_id)
    client_root = data_dir / "client_workspace" / target_client_id
    relative_path = _client_workspace_relative_path(normalized_path)
    if relative_path and relative_path.parts:
        target_path = client_root / relative_path
    else:
        fallback_root = ensure_client_workspace(data_dir, target_client_id).get(fallback_label) or client_root / fallback_label
        target_path = fallback_root / safe_filename(Path(normalized_path).name or "task-attachment")

    source_path = Path(normalized_path)
    final_target = target_path
    final_target.parent.mkdir(parents=True, exist_ok=True)
    if source_path.exists():
        try:
            if source_path.resolve() != target_path.resolve():
                if target_path.exists():
                    stem = safe_filename(target_path.stem or "task-attachment")
                    final_target = target_path.with_name(f"{stem}__{uuid4().hex[:6]}{target_path.suffix.lower()}")
                shutil.move(str(source_path), str(final_target))
        except FileNotFoundError:
            final_target = target_path
    moved_paths[normalized_path] = str(final_target)
    return str(final_target)


def _sync_task_attachment_scope(
    db: Database,
    data_dir: Path,
    build_task_attachment_fn: Callable[[object], TaskAttachmentRecord],
    build_attachment_event_line_activity_fn: Callable[[TaskAttachmentRecord], EventLineActivityRecord],
    ensure_standard_client_folders_fn: Callable[[str], None],
    task_id: str,
    client_id: str | None,
    event_line_id: str | None,
    *,
    cloud: bool,
) -> None:
    normalized_client_id = str(client_id or "").strip()
    if not normalized_client_id:
        return
    normalized_event_line_id = str(event_line_id or "").strip() or None
    table_name = "task_attachments_cloud" if cloud else "task_attachments"
    attachment_rows = db.fetchall(
        f"""
        SELECT
            a.*,
            d.excerpt AS document_excerpt
        FROM {table_name} a
        LEFT JOIN documents d ON d.id = a.document_id
        WHERE a.task_id = ?
        ORDER BY a.created_at DESC
        """,
        (task_id,),
    )
    if not attachment_rows:
        return

    ensure_standard_client_folders_fn(normalized_client_id)
    moved_paths: dict[str, str] = {}
    affected_client_ids: set[str] = {normalized_client_id}
    affected_event_line_ids: set[str] = {normalized_event_line_id} if normalized_event_line_id else set()

    for row in attachment_rows:
        attachment = build_task_attachment_fn(row)
        affected_client_ids.add(attachment.clientId)
        if attachment.eventLineId:
            affected_event_line_ids.add(attachment.eventLineId)

        updated_attachment_path = _rehome_client_workspace_path(
            data_dir,
            normalized_client_id,
            attachment.path,
            moved_paths=moved_paths,
        ) or attachment.path
        db.execute(
            f"UPDATE {table_name} SET client_id = ?, event_line_id = ?, path = ? WHERE id = ?",
            (normalized_client_id, normalized_event_line_id, updated_attachment_path, attachment.id),
        )

        folder_id, folder_label = _resolve_client_folder_ref_by_path(db, normalized_client_id, updated_attachment_path)
        if attachment.documentId:
            document_row = db.fetchone("SELECT * FROM documents WHERE id = ?", (attachment.documentId,))
            if document_row:
                updated_document_path = _rehome_client_workspace_path(
                    data_dir,
                    normalized_client_id,
                    str(document_row["path"]),
                    moved_paths=moved_paths,
                ) or str(document_row["path"])
                updated_original_source_path = _rehome_client_workspace_path(
                    data_dir,
                    normalized_client_id,
                    str(document_row["original_source_path"]) if document_row["original_source_path"] else None,
                    moved_paths=moved_paths,
                )
                db.execute(
                    """
                    UPDATE documents
                    SET client_id = ?, folder_id = ?, path = ?, original_source_path = ?
                    WHERE id = ?
                    """,
                    (
                        normalized_client_id,
                        folder_id,
                        updated_document_path,
                        updated_original_source_path,
                        attachment.documentId,
                    ),
                )
                db.execute(
                    "UPDATE evidence_refs SET client_id = ?, path = ? WHERE document_id = ?",
                    (normalized_client_id, updated_document_path, attachment.documentId),
                )

            knowledge_row = db.fetchone(
                "SELECT * FROM knowledge_documents WHERE document_id = ?",
                (attachment.documentId,),
            )
            if knowledge_row:
                updated_original_path = _rehome_client_workspace_path(
                    data_dir,
                    normalized_client_id,
                    str(knowledge_row["original_path"]),
                    moved_paths=moved_paths,
                ) or str(knowledge_row["original_path"])
                updated_import_source_path = _rehome_client_workspace_path(
                    data_dir,
                    normalized_client_id,
                    str(knowledge_row["import_source_path"]) if knowledge_row["import_source_path"] else None,
                    moved_paths=moved_paths,
                )
                updated_current_human_path = _rehome_client_workspace_path(
                    data_dir,
                    normalized_client_id,
                    str(knowledge_row["current_human_path"]) if knowledge_row["current_human_path"] else None,
                    moved_paths=moved_paths,
                )
                _, human_folder_category = _resolve_client_folder_ref_by_path(db, normalized_client_id, updated_current_human_path)
                db.execute(
                    """
                    UPDATE knowledge_documents
                    SET client_id = ?, original_path = ?, import_source_path = ?, current_human_path = ?, human_folder_category = ?
                    WHERE document_id = ?
                    """,
                    (
                        normalized_client_id,
                        updated_original_path,
                        updated_import_source_path,
                        updated_current_human_path,
                        human_folder_category,
                        attachment.documentId,
                    ),
                )

            v2_row = db.fetchone("SELECT * FROM v2_documents WHERE document_id = ?", (attachment.documentId,))
            if v2_row:
                updated_v2_original_path = _rehome_client_workspace_path(
                    data_dir,
                    normalized_client_id,
                    str(v2_row["original_path"]),
                    moved_paths=moved_paths,
                ) or str(v2_row["original_path"])
                updated_managed_path = _rehome_client_workspace_path(
                    data_dir,
                    normalized_client_id,
                    str(v2_row["managed_path"]),
                    moved_paths=moved_paths,
                ) or str(v2_row["managed_path"])
                db.execute(
                    """
                    UPDATE v2_documents
                    SET client_id = ?, original_path = ?, managed_path = ?
                    WHERE document_id = ?
                    """,
                    (
                        normalized_client_id,
                        updated_v2_original_path,
                        updated_managed_path,
                        attachment.documentId,
                    ),
                )

        db.execute(
            "DELETE FROM event_line_activities WHERE source_type = 'attachment' AND source_id = ?",
            (attachment.id,),
        )
        updated_attachment = TaskAttachmentRecord(
            id=attachment.id,
            taskId=attachment.taskId,
            clientId=normalized_client_id,
            eventLineId=normalized_event_line_id,
            documentId=attachment.documentId,
            title=attachment.title,
            path=updated_attachment_path,
            kind=attachment.kind,
            source=attachment.source,
            sizeBytes=attachment.sizeBytes,
            createdAt=attachment.createdAt,
        )
        if normalized_event_line_id:
            activity = build_attachment_event_line_activity_fn(updated_attachment)
            db.execute(
                """
                INSERT INTO event_line_activities(
                    id, event_line_id, source_type, source_id, happened_at, actor_id, actor_name, title, summary, metadata_json, created_at
                ) VALUES(?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    new_id("ela"),
                    activity.eventLineId,
                    activity.sourceType,
                    activity.sourceId,
                    activity.happenedAt,
                    activity.actorId,
                    activity.actorName,
                    activity.title,
                    activity.summary,
                    to_json(activity.metadata),
                    activity.happenedAt,
                ),
            )

        attachment_source_id = f"{task_id}:{attachment.title}"
        db.execute(
            """
            DELETE FROM memory_facts
            WHERE source_type = 'task_attachment'
              AND source_id = ?
              AND scope_type = 'client'
              AND scope_id <> ?
            """,
            (attachment_source_id, normalized_client_id),
        )
        if normalized_event_line_id:
            db.execute(
                """
                DELETE FROM memory_facts
                WHERE source_type = 'task_attachment'
                  AND source_id = ?
                  AND scope_type = 'event_line'
                  AND scope_id <> ?
                """,
                (attachment_source_id, normalized_event_line_id),
            )
        else:
            db.execute(
                """
                DELETE FROM memory_facts
                WHERE source_type = 'task_attachment'
                  AND source_id = ?
                  AND scope_type = 'event_line'
                """,
                (attachment_source_id,),
            )
        record_task_attachment_writeback(
            db,
            task_id=task_id,
            client_id=normalized_client_id,
            event_line_id=normalized_event_line_id,
            attachment_title=attachment.title,
            attachment_path=updated_attachment_path,
        )

    for affected_client_id in sorted(client_id for client_id in affected_client_ids if client_id):
        refresh_client_folder_counts(db, affected_client_id)
        refresh_organization_notebook_snapshot(db, affected_client_id)
    for affected_event_line_id in sorted(event_line_id for event_line_id in affected_event_line_ids if event_line_id):
        refresh_event_line_memory_snapshot(db, affected_event_line_id)
    _invalidate_event_line_snapshot_cache(*affected_event_line_ids)


def _first_nonempty_text(*values: object) -> str | None:
    for value in values:
        if value is None:
            continue
        text = str(value).strip()
        if text:
            return text
    return None


def _contains_any_keyword(text: str, keywords: tuple[str, ...]) -> bool:
    return any(keyword in text for keyword in keywords)


def _infer_action_os_business_category(
    *,
    title: str = "",
    desc: str = "",
    source_type: str | None = None,
    event_line_name: str | None = None,
    project_module_name: str | None = None,
    project_flow_name: str | None = None,
) -> str:
    normalized = " ".join(
        part.strip()
        for part in [
            title,
            desc,
            event_line_name or "",
            project_module_name or "",
            project_flow_name or "",
            source_type or "",
        ]
        if part and part.strip()
    )
    if not normalized:
        return "专项推进"
    if _contains_any_keyword(normalized, ("模板", "标准", "手册", "知识库", "沉淀", "资料库", "归档", "官网设计", "系统设计")):
        return "产品化沉淀"
    if _contains_any_keyword(normalized, ("审批", "复核", "确认", "对齐", "同步", "协同", "会签", "回收")):
        return "组织协同"
    if _contains_any_keyword(normalized, ("流程", "机制", "规则", "自动汇总", "制度", "治理")):
        return "管理机制"
    if _contains_any_keyword(normalized, ("合作", "拜访", "约见", "介绍", "方案", "报价", "基金会", "客户", "赋能")):
        return "业务扩展"
    if _contains_any_keyword(normalized, ("交付", "上线", "实施", "演示", "需求", "开发", "系统", "网站", "官网")):
        return "项目推进"
    if _contains_any_keyword(normalized, ("伙伴", "联盟", "开源", "外部", "生态")):
        return "外部合作"
    return "专项推进"


def _resolve_task_action_os_fields(
    *,
    title: str,
    desc: str,
    source_type: str | None,
    business_category: str | None,
    current_blocker: str | None,
    next_action: str | None,
    recent_decision: str | None,
    evidence_count: int | None,
    project_context: TaskProjectContextRecord | None = None,
    event_line_context: dict[str, object] | WeeklyReviewEventLineContextRecord | None = None,
    attachment_count: int = 0,
) -> tuple[str | None, str | None, str | None, str | None, int]:
    event_line_name = None
    event_line_business_category = None
    event_line_current_blocker = None
    event_line_recent_decision = None
    event_line_next_step = None
    event_line_evidence_count = 0

    if isinstance(event_line_context, dict):
        event_line_name = _first_nonempty_text(event_line_context.get("name"))
        event_line_business_category = _first_nonempty_text(event_line_context.get("businessCategory"))
        event_line_current_blocker = _first_nonempty_text(event_line_context.get("currentBlocker"))
        event_line_recent_decision = _first_nonempty_text(event_line_context.get("recentDecision"))
        event_line_next_step = _first_nonempty_text(event_line_context.get("nextStep"))
        event_line_evidence_count = int(event_line_context.get("evidenceCount") or 0)
    elif event_line_context is not None:
        event_line_name = _first_nonempty_text(event_line_context.name)
        event_line_business_category = _first_nonempty_text(event_line_context.businessCategory)
        event_line_current_blocker = _first_nonempty_text(event_line_context.currentBlocker)
        event_line_recent_decision = _first_nonempty_text(event_line_context.recentDecision)
        event_line_next_step = _first_nonempty_text(event_line_context.nextStep)
        event_line_evidence_count = int(event_line_context.evidenceCount or 0)

    resolved_business_category = _first_nonempty_text(
        business_category,
        event_line_business_category,
        _infer_action_os_business_category(
            title=title,
            desc=desc,
            source_type=source_type,
            event_line_name=event_line_name,
            project_module_name=project_context.projectModuleName if project_context else None,
            project_flow_name=project_context.projectFlowName if project_context else None,
        ),
    )
    resolved_current_blocker = _first_nonempty_text(
        current_blocker,
        event_line_current_blocker,
        project_context.currentBlocker if project_context else None,
    )
    resolved_next_action = _first_nonempty_text(
        next_action,
        event_line_next_step,
        project_context.nextAction if project_context else None,
    )
    resolved_recent_decision = _first_nonempty_text(
        recent_decision,
        event_line_recent_decision,
        project_context.recentProgress if project_context else None,
    )
    resolved_evidence_count = max(
        int(evidence_count or 0),
        int(attachment_count or 0),
        event_line_evidence_count,
    )
    return (
        resolved_business_category,
        resolved_current_blocker,
        resolved_next_action,
        resolved_recent_decision,
        resolved_evidence_count,
    )


def _task_snapshot_from_task(task: TaskRecord, db: Database | None = None) -> dict[str, object]:
    return {
        "title": task.title,
        "status": task.status,
        "dueDate": task.dueDate,
        "createdAt": task.createdAt,
        "ownerId": task.ownerId,
        "ownerName": task.ownerName,
        "clientId": task.clientId,
        "clientName": task.clientName,
        "eventLineId": task.eventLineId,
        "eventLineName": task.eventLineName,
        "tags": [tag.model_dump() for tag in task.tags],
        "listName": task.listName,
        "listColor": task.listColor,
        "orgContext": task.orgContext.model_dump() if task.orgContext else None,
        "projectContext": task.projectContext.model_dump() if task.projectContext else None,
        "eventLineContext": _event_line_snapshot_context(db, task.eventLineId, task.eventLineName),
    }


def empty_review_structured_note() -> WeeklyReviewTaskStructuredNoteRecord:
    return WeeklyReviewTaskStructuredNoteRecord()


def _derive_reflection_text_from_legacy_structured(parsed: dict[str, object]) -> str:
    candidates = [
        str(parsed.get("reflection") or "").strip(),
        str(parsed.get("successExperience") or "").strip(),
        str(parsed.get("supportNeeded") or "").strip(),
        str(parsed.get("failureInsight") or "").strip(),
        str(parsed.get("blockerReason") or "").strip(),
        str(parsed.get("progress") or "").strip(),
        str(parsed.get("nextAction") or "").strip(),
        str(parsed.get("successReason") or "").strip(),
    ]
    return next((item for item in candidates if item), "")


def coerce_review_structured_note(value: object) -> WeeklyReviewTaskStructuredNoteRecord:
    parsed = value
    if isinstance(value, str):
        parsed = from_json(value, {})
    if isinstance(parsed, WeeklyReviewTaskStructuredNoteRecord):
        return parsed
    if isinstance(parsed, dict):
        return WeeklyReviewTaskStructuredNoteRecord(
            reflection=_derive_reflection_text_from_legacy_structured(parsed),
            lightweightTag=str(parsed.get("lightweightTag") or "").strip(),  # type: ignore[arg-type]
            planCommitment=str(parsed.get("planCommitment") or "").strip(),
            progress=str(parsed.get("progress") or "").strip(),
            completionStatus=str(parsed.get("completionStatus") or "in_progress").strip(),  # type: ignore[arg-type]
            departmentPlanId=str(parsed.get("departmentPlanId")).strip() if parsed.get("departmentPlanId") else None,
            departmentPlanAlignment=str(parsed.get("departmentPlanAlignment") or "unknown").strip(),  # type: ignore[arg-type]
            organizationPlanId=str(parsed.get("organizationPlanId")).strip() if parsed.get("organizationPlanId") else None,
            organizationPlanAlignment=str(parsed.get("organizationPlanAlignment") or "unknown").strip(),  # type: ignore[arg-type]
            successReason=str(parsed.get("successReason") or "").strip(),
            successExperience=str(parsed.get("successExperience") or "").strip(),
            blockerReason=str(parsed.get("blockerReason") or "").strip(),
            failureInsight=str(parsed.get("failureInsight") or "").strip(),
            supportNeeded=str(parsed.get("supportNeeded") or "").strip(),
            nextAction=str(parsed.get("nextAction") or "").strip(),
        )
    return empty_review_structured_note()


def compose_review_note(structured_note: WeeklyReviewTaskStructuredNoteRecord, fallback_note: str = "") -> str:
    reflection = structured_note.reflection.strip()
    if reflection:
        if structured_note.completionStatus in {"done_on_time", "done_late"}:
            return f"任务完成心得：{reflection}"
        if structured_note.lightweightTag:
            return f"需要支持 / 思考：{reflection}（当前卡点：{structured_note.lightweightTag}）"
        return f"需要支持 / 思考：{reflection}"
    if structured_note.lightweightTag and structured_note.completionStatus not in {"done_on_time", "done_late"}:
        return f"需要支持 / 思考：{structured_note.lightweightTag}"
    is_done = structured_note.completionStatus in {"done_on_time", "done_late"}
    simple_text_candidates = (
        [
            structured_note.successExperience.strip(),
            structured_note.successReason.strip(),
            structured_note.progress.strip(),
            structured_note.nextAction.strip(),
        ]
        if is_done
        else [
            structured_note.supportNeeded.strip(),
            structured_note.failureInsight.strip(),
            structured_note.blockerReason.strip(),
            structured_note.progress.strip(),
            structured_note.nextAction.strip(),
        ]
    )
    simple_text = next((item for item in simple_text_candidates if item), "")
    if simple_text:
        prefix = "任务完成心得：" if is_done else "需要支持 / 思考："
        return f"{prefix}{simple_text}"
    return fallback_note.strip()


def _review_entry_from_task(
    *,
    task: TaskRecord,
    week_label: str,
    content_domain: str,
    review_id: str | None = None,
    note: str = "",
    structured_note: WeeklyReviewTaskStructuredNoteRecord | None = None,
    reviewed_at: str | None = None,
    snapshot: dict[str, object] | None = None,
    db: Database | None = None,
) -> WeeklyReviewTaskEntryRecord:
    payload = snapshot or _task_snapshot_from_task(task, db)
    normalized_structured_note = structured_note or empty_review_structured_note()
    return WeeklyReviewTaskEntryRecord(
        id=f"draft_{task.id}" if not review_id and not reviewed_at else f"review_{task.id}_{week_label}",
        reviewId=review_id,
        taskId=task.id,
        weekLabel=week_label,
        contentDomain=content_domain,  # type: ignore[arg-type]
        note=note,
        structuredNote=normalized_structured_note,
        reviewedAt=reviewed_at,
        taskSnapshot=payload,  # type: ignore[arg-type]
    )


def _sql_placeholders(values: tuple[str, ...] | list[str]) -> str:
    return ",".join("?" for _ in values)


def demo_data_loaded(db: Database) -> bool:
    if db.get_setting("demo_data_loaded", "0") != "1":
        return False
    placeholders = _sql_placeholders(DEMO_CLIENT_IDS)
    return bool(db.scalar(f"SELECT COUNT(1) AS count FROM clients WHERE id IN ({placeholders})", DEMO_CLIENT_IDS))


def build_demo_data_response(db: Database) -> DemoDataResponse:
    client_placeholders = _sql_placeholders(DEMO_CLIENT_IDS)
    task_placeholders = _sql_placeholders(DEMO_TASK_IDS)
    radar_placeholders = _sql_placeholders(DEMO_RADAR_IDS)
    handbook_placeholders = _sql_placeholders(DEMO_HANDBOOK_IDS)
    return DemoDataResponse(
        loaded=demo_data_loaded(db),
        clients=int(db.scalar(f"SELECT COUNT(1) AS count FROM clients WHERE id IN ({client_placeholders})", DEMO_CLIENT_IDS)),
        documents=int(db.scalar(f"SELECT COUNT(1) AS count FROM documents WHERE client_id IN ({client_placeholders})", DEMO_CLIENT_IDS)),
        tasks=int(db.scalar(f"SELECT COUNT(1) AS count FROM tasks WHERE id IN ({task_placeholders})", DEMO_TASK_IDS)),
        topics=int(db.scalar(f"SELECT COUNT(1) AS count FROM topic_radars WHERE id IN ({radar_placeholders})", DEMO_RADAR_IDS)),
        handbookEntries=int(db.scalar(f"SELECT COUNT(1) AS count FROM handbook_entries WHERE id IN ({handbook_placeholders})", DEMO_HANDBOOK_IDS)),
    )


def clear_demo_dataset(state: AppState) -> DemoDataResponse:
    meeting_rows = state.db.fetchall(
        f"SELECT id FROM meetings WHERE client_id IN ({_sql_placeholders(DEMO_CLIENT_IDS)})",
        DEMO_CLIENT_IDS,
    )
    meeting_ids = tuple(str(row["id"]) for row in meeting_rows)
    source_ids = tuple([*DEMO_CLIENT_IDS, *DEMO_CANDIDATE_IDS, *meeting_ids])

    state.db.execute(
        f"DELETE FROM tasks WHERE id IN ({_sql_placeholders(DEMO_TASK_IDS)})",
        DEMO_TASK_IDS,
    )
    if source_ids:
        state.db.execute(
            f"DELETE FROM tasks WHERE source_id IN ({_sql_placeholders(source_ids)})",
            source_ids,
        )
    state.db.execute(
        f"DELETE FROM topic_candidates WHERE id IN ({_sql_placeholders(DEMO_CANDIDATE_IDS)})",
        DEMO_CANDIDATE_IDS,
    )
    state.db.execute(
        f"DELETE FROM topic_radars WHERE id IN ({_sql_placeholders(DEMO_RADAR_IDS)})",
        DEMO_RADAR_IDS,
    )
    state.db.execute(
        f"DELETE FROM handbook_entries WHERE id IN ({_sql_placeholders(DEMO_HANDBOOK_IDS)})",
        DEMO_HANDBOOK_IDS,
    )
    state.db.execute(
        f"DELETE FROM chat_messages WHERE id IN ({_sql_placeholders(DEMO_CHAT_MESSAGE_IDS)})",
        DEMO_CHAT_MESSAGE_IDS,
    )
    state.db.execute(
        f"DELETE FROM chat_threads WHERE id IN ({_sql_placeholders(DEMO_THREAD_IDS)})",
        DEMO_THREAD_IDS,
    )
    state.db.execute(
        f"DELETE FROM goal_records WHERE id IN ({_sql_placeholders(DEMO_GOAL_IDS)})",
        DEMO_GOAL_IDS,
    )
    state.db.execute(
        f"DELETE FROM dna_terms WHERE id IN ({_sql_placeholders(DEMO_DNA_IDS)})",
        DEMO_DNA_IDS,
    )
    state.db.execute(
        f"DELETE FROM documents WHERE id IN ({_sql_placeholders(DEMO_DOCUMENT_IDS)})",
        DEMO_DOCUMENT_IDS,
    )
    state.db.execute(
        f"DELETE FROM clients WHERE id IN ({_sql_placeholders(DEMO_CLIENT_IDS)})",
        DEMO_CLIENT_IDS,
    )
    state.db.set_setting("demo_data_loaded", "0")
    return build_demo_data_response(state.db)


def load_demo_dataset(state: AppState) -> DemoDataResponse:
    clear_demo_dataset(state)
    timestamp = now_iso()
    clients = [
        ("client_cffc", "为爱黔行", "黔行公益", "公益教育", "非营利项目", "聚焦山区儿童教育与志愿者体系建设。", "战略陪伴中"),
        ("client_star", "星辰科技", "星辰", "SaaS", "商业化 KA", "营销 SaaS 服务商，正在梳理增长打法。", "方案梳理"),
    ]
    state.db.executemany(
        "INSERT INTO clients(id, name, alias, domain, type, intro, stage, created_at, updated_at) VALUES(?, ?, ?, ?, ?, ?, ?, ?, ?)",
        [(item[0], item[1], item[2], item[3], item[4], item[5], item[6], timestamp, timestamp) for item in clients],
    )
    state.db.executemany(
        "INSERT INTO chat_threads(id, client_id, title, created_at, updated_at) VALUES(?, ?, ?, ?, ?)",
        [
            ("thread_cffc", "client_cffc", "默认研判线程", timestamp, timestamp),
            ("thread_star", "client_star", "默认研判线程", timestamp, timestamp),
        ],
    )
    state.db.executemany(
        "INSERT INTO goal_records(id, client_id, title, quarter, progress, owner_name, created_at, updated_at) VALUES(?, ?, ?, ?, ?, ?, ?, ?)",
        [
            ("goal_1", "client_cffc", "提升项目传播清晰度", "2026 Q2", 62, "庆华", timestamp, timestamp),
            ("goal_2", "client_cffc", "补齐捐赠人关系素材", "2026 Q2", 40, "嘉宁", timestamp, timestamp),
            ("goal_3", "client_star", "验证销售线索质量", "2026 Q2", 55, "一朔", timestamp, timestamp),
        ],
    )
    state.db.executemany(
        "INSERT INTO dna_terms(id, client_id, category, canonical_name, aliases_json, description, created_at, updated_at) VALUES(?, ?, ?, ?, ?, ?, ?, ?)",
        [
            ("dna_1", "client_cffc", "组织习惯", "田野优先", to_json(["先下场", "先到一线"]), "所有策略判断优先结合一线反馈。", timestamp, timestamp),
            ("dna_2", "client_star", "增长原则", "线索先验", to_json(["先验证线索"]), "任何市场动作都要先验证线索质量。", timestamp, timestamp),
        ],
    )
    state.db.executemany(
        "INSERT INTO documents(id, client_id, folder_id, title, path, kind, source, excerpt, tags_json, created_at) VALUES(?, ?, NULL, ?, ?, ?, ?, ?, ?, ?)",
        [
            ("doc_1", "client_cffc", "项目启动纪要.md", "/mock/client_cffc/项目启动纪要.md", "md", "file", "记录了项目目标、时间表与关键风险。", to_json(["会议", "启动"]), timestamp),
            ("doc_2", "client_cffc", "捐赠人访谈摘要.txt", "/mock/client_cffc/捐赠人访谈摘要.txt", "txt", "file", "汇总了主要捐赠人的关切、信任点与反馈。", to_json(["访谈"]), timestamp),
            ("doc_3", "client_star", "增长诊断报告.pdf", "/mock/client_star/增长诊断报告.pdf", "pdf", "file", "聚焦线索质量、转化链路和销售节奏。", to_json(["诊断"]), timestamp),
        ],
    )
    state.db.executemany(
        "INSERT INTO tasks(id, title, description, status, priority, list_id, owner_name, ddl, source_type, source_id, tags_json, created_at, updated_at) VALUES(?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
        [
            ("task_seed_1", "准备周五跨部门复盘会", "梳理客户推进中的异常转化问题。", "inbox", "high", "list-0", "庆华", "周五", "manual", "client_cffc", to_json(["会议"]), timestamp, timestamp),
            ("task_seed_2", "梳理客户反馈的 10 个核心痛点", "提炼成客户工作台的重点议题。", "doing", "normal", "list-1", "一朔", "今天", "manual", "client_star", to_json(["待跟进"]), timestamp, timestamp),
        ],
    )
    state.db.execute(
        "INSERT INTO task_notes(id, task_id, note, created_at, updated_at) VALUES(?, ?, ?, ?, ?)",
        (new_id("tn"), "task_seed_2", "用户反馈集中在试用转化与产品定位不清。", timestamp, timestamp),
    )
    state.db.executemany(
        "INSERT INTO topic_radars(id, title, prompt, time_range, created_at) VALUES(?, ?, ?, ?, ?)",
        [
            ("radar_ai", "大模型应用", "关注公益与咨询行业的大模型落地案例。", "3_days", timestamp),
            ("radar_fund", "筹资趋势", "跟踪基金会筹资、品牌传播与捐赠人运营的最新打法。", "7_days", timestamp),
        ],
    )
    state.db.executemany(
        "INSERT INTO topic_candidates(id, radar_id, title, summary, source, status, created_at, updated_at) VALUES(?, ?, ?, ?, ?, ?, ?, ?)",
        [
            ("cand_1", "radar_ai", "公益组织开始搭建内部 AI 助理", "多个案例开始从内容检索走向流程型工作台。", "行业观察", "candidate", timestamp, timestamp),
            ("cand_2", "radar_fund", "捐赠人分层运营成为主流", "筹资团队开始按关系深浅重新定义触达节奏。", "调研纪要", "tracking", timestamp, timestamp),
        ],
    )
    state.db.executemany(
        "INSERT INTO handbook_entries(id, title, summary, tags_json, source_type, client_id, created_at) VALUES(?, ?, ?, ?, ?, ?, ?)",
        [
            ("hb_1", "会议不要只产纪要", "所有会议结论必须落到负责人与时间点，否则无法闭环。", to_json(["协作规则", "会议"]), "meeting", "client_cffc", timestamp),
        ],
    )
    state.db.execute(
        "INSERT INTO chat_messages(id, thread_id, role, content, structured_data_json, model_route, evidence_json, status, created_at) VALUES(?, ?, 'assistant', ?, ?, ?, ?, 'success', ?)",
        (
            "msg_seed_1",
            "thread_cffc",
            "已为你载入为爱黔行的内部上下文。",
            to_json(
                {
                    "content": "已围绕当前客户整理出一版初始判断。",
                    "judgment": "现阶段关键不在新增素材，而在把现有资料写成推进动作。",
                    "analysis": "1. 项目目标已清楚。2. 捐赠人反馈有基础。3. 缺的是稳定的任务闭环。",
                    "actions": "先从会议发布和任务收件箱打通。",
                    "timeline": "建议本周内先跑通一次完整闭环。",
                }
            ),
            "AI · mock",
            to_json(
                [
                    {
                        "id": "seed_ev",
                        "title": "项目启动纪要.md",
                        "excerpt": "记录了项目目标、时间表与关键风险。",
                        "sourceType": "md",
                        "documentId": "doc_1",
                        "path": "/mock/client_cffc/项目启动纪要.md",
                    }
                ]
            ),
            timestamp,
        ),
    )
    state.db.set_setting("demo_data_loaded", "1")
    return build_demo_data_response(state.db)


def create_app(data_dir: Path | None = None) -> FastAPI:
    resolved_data_dir = data_dir or Path(
        os.getenv("YIYU_WORKBENCH_DATA_DIR") or (Path.home() / "Library" / "Application Support" / "YiyuThinkTankWorkbench")
    )
    resolved_data_dir.mkdir(parents=True, exist_ok=True)
    db = Database(resolved_data_dir / "app.db")
    def build_secret_store(service_name: str):
        try:
            store = MacOSKeychainSecretStore(service_name=service_name, account_name="default")
            store.get_api_key()
            return store
        except Exception:
            return MemorySecretStore()

    qwen_store = build_secret_store("com.yiyu.self-workbench.qwen")
    feishu_secret_store = build_secret_store("com.yiyu.self-workbench.feishu")
    ai = AiService(db, {"qwen": qwen_store})
    backup_dir = resolved_data_dir / "backups"
    state = AppState(
        data_dir=resolved_data_dir,
        backup_dir=backup_dir,
        db=db,
        ai=ai,
        feishu_secret_store=feishu_secret_store,
        cloud_api_url=os.environ.get("YIYU_CLOUD_API_URL", "http://127.0.0.1:47830").rstrip("/"),
        job_stop=Event(),
        topic_insight_executor=ThreadPoolExecutor(max_workers=3),
        chat_answer_executor=ThreadPoolExecutor(max_workers=2),
        template_fill_executor=ThreadPoolExecutor(max_workers=1),
    )
    seed_defaults(state)
    if state.db.get_setting("demo_data_loaded", "0") != "1":
        clear_demo_dataset(state)

    app = FastAPI(title=APP_NAME, version=APP_VERSION)
    app.state.app_state = state
    app.add_middleware(
        CORSMiddleware,
        allow_origins=list(ALLOWED_LOCAL_ORIGINS),
        allow_origin_regex=ALLOWED_LOCAL_ORIGIN_REGEX,
        allow_credentials=True,
        allow_methods=["*"],
        allow_headers=["*"],
    )

    @app.middleware("http")
    async def _local_browser_origin_guard(request: Request, call_next):
        rejection_reason = validate_local_browser_request(request.url.path, request.headers, request.method)
        if rejection_reason:
            return JSONResponse(status_code=403, content={"detail": rejection_reason})
        return await call_next(request)

    @app.on_event("startup")
    def _startup_worker() -> None:
        if state.job_thread and state.job_thread.is_alive():
            return
        recover_stale_knowledge_jobs()
        recover_stale_imports()
        recover_stale_template_fill_runs()
        recover_stale_loading_chat_messages()
        state.job_stop.clear()
        state.job_thread = Thread(target=knowledge_worker_loop, name="knowledge-worker", daemon=True)
        state.job_thread.start()

    @app.on_event("shutdown")
    def _shutdown_worker() -> None:
        state.job_stop.set()
        if state.job_thread and state.job_thread.is_alive():
            state.job_thread.join(timeout=1.5)
        if state.topic_insight_executor:
            state.topic_insight_executor.shutdown(wait=False, cancel_futures=False)
        if state.chat_answer_executor:
            state.chat_answer_executor.shutdown(wait=False, cancel_futures=False)
        if state.template_fill_executor:
            state.template_fill_executor.shutdown(wait=False, cancel_futures=False)

    def current_operator_row():
        operator_id = state.db.get_setting("current_operator_id", "")
        if operator_id:
            row = state.db.fetchone("SELECT * FROM operators WHERE id = ?", (operator_id,))
            if row:
                return row
        row = state.db.fetchone("SELECT * FROM operators ORDER BY created_at LIMIT 1")
        if not row:
            raise HTTPException(status_code=500, detail="No operator configured")
        state.db.set_setting("current_operator_id", str(row["id"]))
        state.db.execute("UPDATE operators SET is_current = CASE WHEN id = ? THEN 1 ELSE 0 END", (str(row["id"]),))
        return row

    def current_operator_name() -> str:
        row = current_operator_row()
        return str(row["name"]) if row and row["name"] else "系统"

    def _local_task_tag_record(row) -> TaskTagRecord:
        return TaskTagRecord(
            id=str(row["id"]),
            name=str(row["name"]),
            color=str(row["color"] or ("#9CA3AF" if str(row["scope"]) == "self" else "#5B7BFE")),
            scope=str(row["scope"]),
            ownerUserId=str(row["owner_operator_id"]) if str(row["owner_operator_id"]) else None,
            createdBy=str(row["created_by"]) if str(row["created_by"]) else None,
            updatedAt=str(row["updated_at"] or row["created_at"] or now_iso()),
            archivedAt=str(row["archived_at"]) if row["archived_at"] else None,
        )

    def _local_task_list_record(row) -> TaskListRecord:
        return TaskListRecord(
            id=str(row["id"]),
            name=str(row["name"]),
            color=str(row["color"]),
            sortOrder=int(row["sort_order"] or 0),
            isDefault=bool(int(row["is_default"] or 0)),
            scope=str(row["scope"] or "org"),
            archivedAt=str(row["archived_at"]) if row["archived_at"] else None,
        )

    def _local_default_list_id() -> str | None:
        default_row = state.db.fetchone("SELECT id FROM task_lists WHERE is_default = 1 ORDER BY sort_order ASC LIMIT 1")
        if default_row:
            return str(default_row["id"])
        first_row = state.db.fetchone("SELECT id FROM task_lists ORDER BY sort_order ASC, name COLLATE NOCASE ASC LIMIT 1")
        return str(first_row["id"]) if first_row else None

    def _default_local_task_settings() -> TaskSettingsRecord:
        return TaskSettingsRecord(
            defaultListId=_local_default_list_id(),
            defaultPriority="normal",
            defaultDueDatePreset="today",
            defaultViewMode="calendar",
            listSortMode="manual",
            showCompletedTasks=False,
            defaultReviewScope="work",
            autoAssignSelf=True,
            updatedAt=now_iso(),
        )

    def _local_task_settings_record(row) -> TaskSettingsRecord:
        defaults = _default_local_task_settings()
        return TaskSettingsRecord(
            defaultListId=str(row["default_list_id"]) if row["default_list_id"] else defaults.defaultListId,
            defaultPriority=str(row["default_priority"] or defaults.defaultPriority),  # type: ignore[arg-type]
            defaultDueDatePreset=str(row["default_due_date_preset"] or defaults.defaultDueDatePreset),  # type: ignore[arg-type]
            defaultViewMode=str(row["default_view_mode"] or defaults.defaultViewMode),  # type: ignore[arg-type]
            listSortMode=str(row["list_sort_mode"] or defaults.listSortMode),  # type: ignore[arg-type]
            showCompletedTasks=bool(int(row["show_completed_tasks"] or 0)),
            defaultReviewScope=str(row["default_review_scope"] or defaults.defaultReviewScope),  # type: ignore[arg-type]
            autoAssignSelf=bool(int(row["auto_assign_self"] if row["auto_assign_self"] is not None else 1)),
            updatedAt=str(row["updated_at"] or defaults.updatedAt),
        )

    def _get_local_task_settings(operator_id: str | None = None) -> TaskSettingsRecord:
        resolved_operator_id = operator_id or str(current_operator_row()["id"])
        row = state.db.fetchone("SELECT * FROM task_settings WHERE operator_id = ?", (resolved_operator_id,))
        if row:
            return _local_task_settings_record(row)
        defaults = _default_local_task_settings()
        state.db.execute(
            """
            INSERT OR REPLACE INTO task_settings(
                operator_id, default_list_id, default_priority, default_due_date_preset,
                default_view_mode, list_sort_mode, show_completed_tasks, default_review_scope,
                auto_assign_self, updated_at
            ) VALUES(?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                resolved_operator_id,
                defaults.defaultListId,
                defaults.defaultPriority,
                defaults.defaultDueDatePreset,
                defaults.defaultViewMode,
                defaults.listSortMode,
                1 if defaults.showCompletedTasks else 0,
                defaults.defaultReviewScope,
                1 if defaults.autoAssignSelf else 0,
                defaults.updatedAt,
            ),
        )
        return defaults

    def _visible_local_task_tag_rows(db: Database, operator_id: str) -> list:
        return db.fetchall(
            """
            SELECT *
            FROM task_tags
            WHERE scope = 'org' OR owner_operator_id = ?
            ORDER BY CASE WHEN archived_at IS NULL OR archived_at = '' THEN 0 ELSE 1 END,
                     CASE scope WHEN 'org' THEN 0 ELSE 1 END,
                     name COLLATE NOCASE ASC
            """,
            (operator_id,),
        )

    def _visible_local_task_tags(db: Database, operator_id: str) -> list[TaskTagRecord]:
        return [_local_task_tag_record(row) for row in _visible_local_task_tag_rows(db, operator_id)]

    def _local_tag_rows_by_ids(db: Database, tag_ids: list[str]) -> list:
        if not tag_ids:
            return []
        rows = db.fetchall(
            f"SELECT * FROM task_tags WHERE id IN ({_sql_placeholders(tag_ids)})",
            tuple(tag_ids),
        )
        by_id = {str(row["id"]): row for row in rows}
        return [by_id[tag_id] for tag_id in tag_ids if tag_id in by_id]

    def _ensure_local_tag(
        db: Database,
        operator_id: str,
        name: str,
        scope: str = "org",
        color: str | None = None,
        created_by: str | None = None,
    ) -> TaskTagRecord:
        trimmed = name.strip()
        if not trimmed:
            raise HTTPException(status_code=400, detail="Tag name is required")
        owner_operator_id = operator_id if scope == "self" else ""
        resolved_color = color or ("#9CA3AF" if scope == "self" else "#5B7BFE")
        resolved_creator = created_by or str(current_operator_row()["name"])
        existing = db.fetchone(
            "SELECT * FROM task_tags WHERE name = ? AND scope = ? AND owner_operator_id = ?",
            (trimmed, scope, owner_operator_id),
        )
        timestamp = now_iso()
        if existing:
            db.execute(
                """
                UPDATE task_tags
                SET color = COALESCE(NULLIF(?, ''), color),
                    archived_at = NULL,
                    updated_at = ?
                WHERE id = ?
                """,
                (resolved_color, timestamp, str(existing["id"])),
            )
            row = db.fetchone("SELECT * FROM task_tags WHERE id = ?", (str(existing["id"]),))
            if row:
                return _local_task_tag_record(row)
        tag_id = new_id("tag")
        db.execute(
            """
            INSERT INTO task_tags(
                id, name, color, scope, owner_operator_id, created_by, archived_at, created_at, updated_at
            ) VALUES(?, ?, ?, ?, ?, ?, NULL, ?, ?)
            """,
            (tag_id, trimmed, resolved_color, scope, owner_operator_id, resolved_creator, timestamp, timestamp),
        )
        row = db.fetchone("SELECT * FROM task_tags WHERE id = ?", (tag_id,))
        if not row:
            raise HTTPException(status_code=500, detail="Failed to create task tag")
        return _local_task_tag_record(row)

    def _resolve_local_task_tags(db: Database, operator_id: str, tag_ids: list[str], legacy_names: list[str]) -> list[TaskTagRecord]:
        rows = _local_tag_rows_by_ids(db, tag_ids)
        if rows:
            return [_local_task_tag_record(row) for row in rows]
        if not legacy_names:
            return []
        return [_ensure_local_tag(db, operator_id, name, "org") for name in legacy_names if name.strip()]

    def log_activity(action: str, entity_type: str, entity_id: str, detail: dict[str, object]) -> None:
        session_user = get_cached_session_user()
        actor_name = session_user.fullName if session_user else current_operator_row()["name"]
        state.db.execute(
            """
            INSERT INTO activity_logs(id, actor_name, action, entity_type, entity_id, detail_json, created_at)
            VALUES(?, ?, ?, ?, ?, ?, ?)
            """,
            (new_id("log"), actor_name, action, entity_type, entity_id, to_json(detail), now_iso()),
        )

    def append_knowledge_job_event(job_id: str, level: str, message: str, detail: dict[str, object] | None = None) -> None:
        state.db.execute(
            """
            INSERT INTO knowledge_job_events(id, job_id, level, message, detail_json, created_at)
            VALUES(?, ?, ?, ?, ?, ?)
            """,
            (new_id("kje"), job_id, level, message, to_json(detail or {}), now_iso()),
        )

    def enqueue_knowledge_job(client_id: str, job_type: str, payload: dict[str, object], total_items: int) -> KnowledgeJobRecord:
        job_id = new_id("kjob")
        timestamp = now_iso()
        state.db.execute(
            """
            INSERT INTO knowledge_jobs(
                id, client_id, job_type, status, payload_json, total_items, processed_items,
                last_error, created_at, started_at, finished_at, updated_at
            )
            VALUES(?, ?, ?, 'queued', ?, ?, 0, NULL, ?, NULL, NULL, ?)
            """,
            (job_id, client_id, job_type, to_json(payload), total_items, timestamp, timestamp),
        )
        append_knowledge_job_event(job_id, "info", "知识加工任务已入队", {"jobType": job_type, "totalItems": total_items})
        return KnowledgeJobRecord(
            id=job_id,
            clientId=client_id,
            jobType=job_type,
            status="queued",
            totalItems=total_items,
            processedItems=0,
            createdAt=timestamp,
            updatedAt=timestamp,
        )

    def resolve_import_id_for_job(job_row: dict[str, object] | None) -> str | None:
        if not job_row or str(job_row.get("job_type") or "") != "ingest_import":
            return None
        payload = from_json(str(job_row.get("payload_json") or "{}"), {})
        if not isinstance(payload, dict):
            return None
        import_id = payload.get("importId")
        return str(import_id) if import_id else None

    def update_import_status(import_id: str, *, status: str, imported_count: int | None = None, skipped_count: int | None = None) -> None:
        existing = state.db.fetchone("SELECT imported_count, skipped_count FROM imports WHERE id = ?", (import_id,))
        if not existing:
            return
        next_imported = int(imported_count) if imported_count is not None else int(existing["imported_count"] or 0)
        next_skipped = int(skipped_count) if skipped_count is not None else int(existing["skipped_count"] or 0)
        state.db.execute(
            """
            UPDATE imports
            SET status = ?, imported_count = ?, skipped_count = ?
            WHERE id = ?
            """,
            (status, next_imported, next_skipped, import_id),
        )

    def recover_stale_knowledge_jobs() -> None:
        stale_rows = state.db.fetchall(
            """
            SELECT id
            FROM knowledge_jobs
            WHERE status = 'running'
            ORDER BY created_at ASC
            """
        )
        if not stale_rows:
            return
        timestamp = now_iso()
        for row in stale_rows:
            job_id = str(row["id"])
            state.db.execute(
                """
                UPDATE knowledge_jobs
                SET status = 'queued', started_at = NULL, finished_at = NULL, last_error = ?, updated_at = ?
                WHERE id = ?
                """,
                ("worker_restart_requeued", timestamp, job_id),
            )
            append_knowledge_job_event(job_id, "warning", "检测到旧的运行中任务，已在启动时重新入队")

    def recover_stale_imports() -> None:
        stale_imports = state.db.fetchall(
            """
            SELECT id, client_id, status
            FROM imports
            WHERE status IN ('queued', 'processing')
            ORDER BY created_at ASC
            """
        )
        for row in stale_imports:
            import_id = str(row["id"])
            jobs = state.db.fetchall(
                """
                SELECT *
                FROM knowledge_jobs
                WHERE client_id = ? AND job_type = 'ingest_import'
                ORDER BY created_at DESC
                """,
                (str(row["client_id"]),),
            )
            matched_job: dict[str, object] | None = None
            for job in jobs:
                if resolve_import_id_for_job(dict(job)) == import_id:
                    matched_job = dict(job)
                    break
            if matched_job is None:
                if str(row["status"]) == "processing":
                    update_import_status(import_id, status="failed")
                continue
            job_status = str(matched_job.get("status") or "")
            processed_items = int(matched_job.get("processed_items") or 0)
            if job_status == "completed":
                update_import_status(import_id, status="completed", imported_count=processed_items)
            elif job_status == "failed":
                update_import_status(import_id, status="failed", imported_count=processed_items)
            elif job_status == "running":
                update_import_status(import_id, status="processing", imported_count=processed_items)
            elif job_status == "queued":
                update_import_status(import_id, status="queued", imported_count=processed_items)

    def recover_stale_template_fill_runs() -> None:
        stale_runs = state.db.fetchall(
            """
            SELECT id
            FROM client_template_fill_runs
            WHERE status IN ('queued', 'running')
            ORDER BY created_at ASC
            """
        )
        if not stale_runs:
            return
        timestamp = now_iso()
        for row in stale_runs:
            state.db.execute(
                """
                UPDATE client_template_fill_runs
                SET status = 'failed',
                    phase = 'failed',
                    progress = CASE
                        WHEN progress IS NULL THEN 8.0
                        WHEN progress > 96.0 THEN 96.0
                        WHEN progress < 8.0 THEN 8.0
                        ELSE progress
                    END,
                    stage_label = '模板填写已中断',
                    error_message = COALESCE(NULLIF(error_message, ''), '应用重启或后台中断，当前填写任务已停止，请重新发起。'),
                    updated_at = ?
                WHERE id = ?
                """,
                (timestamp, str(row["id"])),
            )

    def recover_stale_loading_chat_messages() -> None:
        stale_rows = state.db.fetchall(
            """
            SELECT id, retrieval_summary_json, created_at
            FROM chat_messages
            WHERE role = 'assistant' AND status = 'loading'
            ORDER BY created_at ASC
            """
        )
        if not stale_rows:
            return
        now = datetime.now()
        for row in stale_rows:
            created_at_raw = str(row["created_at"] or "")
            try:
                created_at = datetime.fromisoformat(created_at_raw)
            except ValueError:
                created_at = now - timedelta(minutes=10)
            age_seconds = max((now - created_at).total_seconds(), 0.0)
            if age_seconds < 30:
                continue
            summary = from_json(str(row["retrieval_summary_json"] or "{}"), {})
            if not isinstance(summary, dict):
                summary = {}
            summary.update(
                {
                    "phase": "failed",
                    "progress": 100.0,
                    "progressFloor": 100.0,
                    "progressCeiling": 100.0,
                    "stageLabel": "回答生成失败",
                    "failureReason": "后台回答任务中断或超时，已自动收口。",
                    "lastUpdatedAt": now_iso(),
                }
            )
            state.db.execute(
                """
                UPDATE chat_messages
                SET content = ?, structured_data_json = ?, model_route = ?, llm_invoked = 0, provider_used = COALESCE(provider_used, ?),
                    answer_mode = 'system_failure', evidence_status = 'none', failure_reason = ?, timing_json = ?,
                    retrieval_summary_json = ?, evidence_json = COALESCE(evidence_json, '[]'), status = 'success', created_at = ?
                WHERE id = ? AND status = 'loading'
                """,
                (
                    "庆华暂时没能完成这次回答。",
                    to_json(
                        AiStructuredResponse(
                            content="庆华暂时没能完成这次回答。",
                            judgment="后台回答任务中断或超时，本次回答已自动收口为失败态。",
                            analysis="该消息此前一直停留在生成中，没有成功写回最终结果。系统已自动把它标记为失败，避免前端一直卡在加载状态。",
                            actions="请直接重试这个问题；如果反复出现，请检查本地后端与千问配置。",
                            timeline="修复后可立即重新生成。",
                        ).model_dump()
                    ),
                    "AI 调用失败",
                    state.ai.current_provider(),
                    "stale_loading_message_recovered",
                    to_json({"totalMs": round(age_seconds * 1000, 2), "retrievalMs": 0.0, "llmMs": 0.0}),
                    to_json(summary),
                    now_iso(),
                    str(row["id"]),
                ),
            )
            state.db.execute(
                """
                UPDATE client_analysis_runs
                SET status = 'failed',
                    phase = 'failed',
                    progress = 100.0,
                    progress_floor = 100.0,
                    progress_ceiling = 100.0,
                    stage_label = '回答生成失败',
                    long_answer_status = 'failed',
                    summary_status = 'failed',
                    failure_reason = 'stale_loading_message_recovered',
                    elapsed_ms = ?,
                    timing_json = ?,
                    updated_at = ?
                WHERE assistant_message_id = ? AND status IN ('queued', 'running')
                """,
                (
                    round(age_seconds * 1000, 2),
                    to_json({"totalMs": round(age_seconds * 1000, 2), "retrievalMs": 0.0, "llmMs": 0.0}),
                    now_iso(),
                    str(row["id"]),
                ),
            )

    def claim_next_knowledge_job() -> dict[str, object] | None:
        row = state.db.fetchone(
            """
            SELECT *
            FROM knowledge_jobs
            WHERE status = 'queued'
            ORDER BY created_at ASC
            LIMIT 1
            """,
        )
        if not row:
            return None
        timestamp = now_iso()
        state.db.execute(
            "UPDATE knowledge_jobs SET status = 'running', started_at = ?, updated_at = ? WHERE id = ? AND status = 'queued'",
            (timestamp, timestamp, str(row["id"])),
        )
        claimed = state.db.fetchone("SELECT * FROM knowledge_jobs WHERE id = ?", (str(row["id"]),))
        if not claimed or str(claimed["status"]) != "running":
            return None
        append_knowledge_job_event(str(claimed["id"]), "info", "知识加工任务开始执行")
        return dict(claimed)

    def finish_knowledge_job(job_id: str, *, status: str, processed_items: int, last_error: str | None = None) -> None:
        timestamp = now_iso()
        existing_job = state.db.fetchone("SELECT * FROM knowledge_jobs WHERE id = ?", (job_id,))
        state.db.execute(
            """
            UPDATE knowledge_jobs
            SET status = ?, processed_items = ?, last_error = ?, finished_at = ?, updated_at = ?
            WHERE id = ?
            """,
            (status, processed_items, last_error, timestamp, timestamp, job_id),
        )
        append_knowledge_job_event(
            job_id,
            "error" if status == "failed" else "info",
            "知识加工任务执行失败" if status == "failed" else "知识加工任务执行完成",
            {"processedItems": processed_items, "lastError": last_error},
        )
        import_id = resolve_import_id_for_job(dict(existing_job)) if existing_job else None
        if import_id:
            update_import_status(import_id, status="failed" if status == "failed" else "completed", imported_count=processed_items)

    def update_knowledge_job_progress(job_id: str, processed_items: int, message: str) -> None:
        state.db.execute(
            "UPDATE knowledge_jobs SET processed_items = ?, updated_at = ? WHERE id = ?",
            (processed_items, now_iso(), job_id),
        )
        append_knowledge_job_event(job_id, "info", message, {"processedItems": processed_items})

    def process_knowledge_job(job: dict[str, object]) -> None:
        job_id = str(job["id"])
        client_id = str(job["client_id"])
        job_type = str(job["job_type"])
        payload = from_json(str(job["payload_json"]), {})
        processed_items = 0
        if job_type == "ingest_import":
            import_id = str(payload.get("importId"))
            documents = payload.get("documents", [])
            docs = documents if isinstance(documents, list) else []
            state.db.execute("UPDATE imports SET status = 'processing' WHERE id = ?", (import_id,))
            ensure_standard_client_folders(client_id)
            for item in docs:
                if not isinstance(item, dict):
                    continue
                document_id = str(item.get("documentId", ""))
                path = Path(str(item.get("sourcePath", "")))
                original_source_path = Path(str(item.get("originalSourcePath", ""))).expanduser() if item.get("originalSourcePath") else path
                if not document_id or not path.exists():
                    if document_id:
                        append_knowledge_job_event(
                            job_id,
                            "warning",
                            "跳过缺失的原始文件",
                            {"documentId": document_id, "sourcePath": str(path)},
                        )
                    continue
                excerpt = build_excerpt(path)
                prepared = ingest_document_knowledge(
                    state.db,
                    data_dir=state.data_dir,
                    client_id=client_id,
                    import_id=import_id,
                    document_id=document_id,
                    source_path=path,
                    original_source_path=original_source_path,
                    title=str(item.get("title", path.name)),
                    kind=str(item.get("kind", path.suffix.lower().lstrip("."))),
                    source=str(item.get("source", payload.get("mode", "file"))),
                    fallback_excerpt=excerpt,
                    created_at=str(item.get("createdAt", now_iso())),
                    ai_service=state.ai,
                )
                record_imported_document_writeback(
                    state.db,
                    client_id=client_id,
                    document_id=document_id,
                    title=str(item.get("title", path.name)),
                    prepared=prepared,
                )
                prepared_title = str(prepared.get("title") or path.name)
                prepared_category = str(prepared.get("primary_category") or "其他资料")
                target_folder = state.db.fetchone(
                    "SELECT id FROM client_folders WHERE client_id = ? AND label = ?",
                    (client_id, prepared_category),
                )
                state.db.execute(
                    """
                    UPDATE documents
                    SET folder_id = ?
                    WHERE id = ?
                    """,
                    (
                        str(target_folder["id"]) if target_folder else None,
                        document_id,
                    ),
                )
                processed_items += 1
                state.db.execute(
                    "UPDATE imports SET imported_count = ?, status = 'processing' WHERE id = ?",
                    (processed_items, import_id),
                )
                update_knowledge_job_progress(job_id, processed_items, f"已处理 {prepared_title}")
            append_knowledge_job_event(job_id, "info", f"已完成 {V2_PIPELINE_VERSION} 文档索引、章节定位与原文切块")
            ensure_standard_client_folders(client_id)
            total_items = int(job.get("total_items") or len(docs))
            existing_import = state.db.fetchone("SELECT skipped_count FROM imports WHERE id = ?", (import_id,))
            existing_skipped = int(existing_import["skipped_count"] or 0) if existing_import else 0
            skipped_items = existing_skipped + max(0, total_items - processed_items)
            state.db.execute(
                "UPDATE imports SET status = 'completed', imported_count = ?, skipped_count = ? WHERE id = ?",
                (processed_items, skipped_items, import_id),
            )
            finish_knowledge_job(job_id, status="completed", processed_items=processed_items)
            maybe_enqueue_client_dna_generation_job(client_id)
            return
        if job_type == "rebuild_client_knowledge":
            summary = backfill_knowledge_documents(
                state.db,
                data_dir=state.data_dir,
                client_id=client_id,
                ai_service=state.ai,
                progress_callback=lambda count: update_knowledge_job_progress(job_id, count, f"已回填 {count} 份文档"),
            )
            processed_items = int(summary.get("processed", 0))
            missing_items = int(summary.get("missing", 0))
            if missing_items > 0:
                append_knowledge_job_event(
                    job_id,
                    "warning",
                    "部分原始文件已丢失，已跳过缺失文件并继续重建",
                    {"missingItems": missing_items},
                )
            append_knowledge_job_event(job_id, "info", f"已完成 {V2_PIPELINE_VERSION} 文档索引、章节定位与原文切块")
            ensure_standard_client_folders(client_id)
            finish_knowledge_job(job_id, status="completed", processed_items=processed_items)
            maybe_enqueue_client_dna_generation_job(client_id)
            return
        if job_type == "generate_client_dna_candidates":
            module_keys = payload.get("moduleKeys", [])
            module_key_list = [str(item) for item in module_keys] if isinstance(module_keys, list) else []
            refresh_generated = bool(payload.get("refreshGenerated"))
            if not module_key_list:
                module_key_list = [module.moduleKey for module in resolve_client_dna_modules_for_generation(client_id, refresh_generated=refresh_generated)]
            if not module_key_list:
                finish_knowledge_job(job_id, status="completed", processed_items=0)
                return
            for module_key in module_key_list:
                markdown_content, missing_items = build_client_dna_candidate_markdown(client_id, module_key)
                save_client_dna_module(
                    client_id,
                    module_key,
                    markdown_content=markdown_content,
                    file_name=f"{safe_filename(build_client_summary(client_id).name)}-{module_key}-candidate.md",
                    source_kind="generated",
                    updated_by="系统候选生成",
                    missing_info=missing_items,
                )
                processed_items += 1
                update_knowledge_job_progress(job_id, processed_items, f"已生成 {dict(CLIENT_DNA_MODULES).get(module_key, module_key)} 候选文档")
            finish_knowledge_job(job_id, status="completed", processed_items=processed_items)
            return
        finish_knowledge_job(job_id, status="failed", processed_items=0, last_error=f"未知任务类型：{job_type}")

    def knowledge_worker_loop() -> None:
        while not state.job_stop.is_set():
            job = claim_next_knowledge_job()
            if not job:
                time.sleep(0.5)
                continue
            try:
                process_knowledge_job(job)
            except Exception as error:
                finish_knowledge_job(str(job["id"]), status="failed", processed_items=int(job.get("processed_items") or 0), last_error=str(error))
            time.sleep(0.05)

    def build_health() -> HealthResponse:
        ai_health = state.ai.get_health()
        return HealthResponse(
            appName=APP_NAME,
            appVersion=APP_VERSION,
            buildVersion=APP_BUILD_VERSION,
            startedAt=APP_STARTED_AT,
            featureFlags=BACKEND_FEATURE_FLAGS,
            dataDir=str(state.data_dir),
            stats={
                "clients": state.db.scalar("SELECT COUNT(1) AS count FROM clients"),
                "tasks": state.db.scalar("SELECT COUNT(1) AS count FROM tasks"),
                "topics": state.db.scalar("SELECT COUNT(1) AS count FROM topic_candidates"),
                "handbookEntries": state.db.scalar("SELECT COUNT(1) AS count FROM handbook_entries"),
                "analysisRuns": state.db.scalar("SELECT COUNT(1) AS count FROM analysis_runs"),
            },
            ai=HealthAiState(
                provider=ai_health.provider,  # type: ignore[arg-type]
                model=ai_health.model,
                ready=ai_health.ready,
                detail=ai_health.detail,
                credentialSource=ai_health.credential_source,
                fingerprint=ai_health.fingerprint,
            ),
        )

    def build_settings_response() -> SettingsResponse:
        operator = current_operator_row()
        ai_health = state.ai.get_health()
        settings = AppSettingsResponse(
            currentOperatorId=str(operator["id"]),
            aiProvider=state.ai.current_provider(),  # type: ignore[arg-type]
            aiModel=state.ai.current_model(),
            dataDir=str(state.data_dir),
            backupDir=str(state.backup_dir),
            cloudApiUrl=state.cloud_api_url,
            lastBackupAt=state.db.get_setting("last_backup_at", "") or None,
            foldersRootLabel=state.db.get_setting("folders_root_label", "桌面客户资料"),
            aiConfigured=bool(ai_health.fingerprint),
            aiCredentialSource=ai_health.credential_source,
            aiFingerprint=ai_health.fingerprint,
            demoDataLoaded=demo_data_loaded(state.db),
        )
        operators = [
            OperatorRecord(
                id=str(row["id"]),
                name=str(row["name"]),
                role=str(row["role"]),
                team=str(row["team"]),
                color=str(row["color"]),
                isCurrent=bool(row["is_current"]),
            )
            for row in state.db.fetchall("SELECT * FROM operators ORDER BY created_at")
        ]
        return SettingsResponse(settings=settings, operators=operators, health=build_health())

    def _has_persisted_cloud_session() -> bool:
        return bool(state.db.get_setting("cloud_access_token", "") or state.db.get_setting("cloud_refresh_token", ""))

    def get_cloud_token() -> str:
        token = state.db.get_setting("cloud_access_token", "")
        if token:
            state.cloud_session_persistent = True
            return token
        return state.volatile_cloud_access_token

    def get_cloud_refresh_token() -> str:
        token = state.db.get_setting("cloud_refresh_token", "")
        if token:
            state.cloud_session_persistent = True
            return token
        return state.volatile_cloud_refresh_token

    def set_cloud_session(token: str | None, user: SessionUserRecord | None, *, persist: bool = True) -> None:
        state.cloud_session_persistent = persist
        session_user_json = to_json(user.model_dump()) if user else ""
        if persist:
            state.db.set_setting("cloud_access_token", token or "")
            state.db.set_setting("cloud_session_user", session_user_json)
            state.volatile_cloud_access_token = ""
            state.volatile_cloud_session_user_json = ""
            return
        state.db.set_setting("cloud_access_token", "")
        state.db.set_setting("cloud_session_user", "")
        state.volatile_cloud_access_token = token or ""
        state.volatile_cloud_session_user_json = session_user_json

    def set_cloud_refresh_token(token: str | None, *, persist: bool = True) -> None:
        state.cloud_session_persistent = persist
        if persist:
            state.db.set_setting("cloud_refresh_token", token or "")
            state.volatile_cloud_refresh_token = ""
            return
        state.db.set_setting("cloud_refresh_token", "")
        state.volatile_cloud_refresh_token = token or ""

    def clear_cloud_session() -> None:
        set_cloud_session(None, None, persist=True)
        set_cloud_refresh_token(None, persist=True)
        state.volatile_cloud_access_token = ""
        state.volatile_cloud_refresh_token = ""
        state.volatile_cloud_session_user_json = ""
        state.cloud_session_persistent = False

    def get_cached_session_user() -> SessionUserRecord | None:
        raw = state.db.get_setting("cloud_session_user", "")
        if not raw:
            raw = state.volatile_cloud_session_user_json
        parsed = from_json(raw, {}) if raw else {}
        if not isinstance(parsed, dict) or not parsed:
            return None
        try:
            return SessionUserRecord(**parsed)
        except Exception:
            return None

    def current_session_is_admin() -> bool:
        session_user = get_cached_session_user()
        return bool(session_user and session_user.primaryRole == "admin")

    def _load_json_settings_record(key: str, default_factory, model_cls):
        raw = state.db.get_setting(key, "")
        if raw:
            parsed = from_json(raw, {})
            if isinstance(parsed, dict):
                try:
                    return model_cls(**parsed)
                except Exception:
                    pass
        record = default_factory()
        state.db.set_setting(key, to_json(record.model_dump()))
        return record

    def _save_json_settings_record(key: str, record) -> object:
        state.db.set_setting(key, to_json(record.model_dump()))
        return record

    def _default_client_workspace_settings() -> ClientWorkspaceSettingsRecord:
        return ClientWorkspaceSettingsRecord(
            meetingPublishDefaultListId=_local_default_list_id(),
            meetingPublishDefaultPriority="normal",
            updatedAt=now_iso(),
        )

    def _default_topics_settings() -> TopicsSettingsRecord:
        return TopicsSettingsRecord(updatedAt=now_iso())

    def _default_analysis_workbench_settings() -> AnalysisWorkbenchSettingsRecord:
        template_ids = [str(row["id"]) for row in state.db.fetchall("SELECT id FROM analysis_templates ORDER BY created_at ASC")]
        return AnalysisWorkbenchSettingsRecord(
            enabledTemplateIds=template_ids,
            defaultTemplateId=template_ids[0] if template_ids else None,
            updatedAt=now_iso(),
        )

    def _default_handbook_settings() -> HandbookSettingsRecord:
        return HandbookSettingsRecord(updatedAt=now_iso())

    def _default_system_admin_settings() -> SystemAdminSettingsRecord:
        return SystemAdminSettingsRecord(updatedAt=now_iso())

    def _normalize_brand_logo_data_url(value: str | None) -> str | None:
        normalized = (value or "").strip()
        if not normalized:
            return None
        if not normalized.startswith("data:image/png;base64,"):
            raise HTTPException(status_code=400, detail="品牌 Logo 目前只支持 PNG data URL")
        if len(normalized) > 1_500_000:
            raise HTTPException(status_code=400, detail="品牌 Logo 过大，请换更小的 PNG")
        return normalized

    def _normalize_feishu_user_binding_callback_url(value: str | None) -> str:
        normalized = (value or "").strip()
        if not normalized:
            return ""
        parsed = urlparse(normalized)
        if parsed.scheme not in {"http", "https"} or not parsed.netloc:
            raise HTTPException(status_code=400, detail="飞书个人绑定回调 URL 必须是完整的 http(s) 地址。")
        if not parsed.path:
            raise HTTPException(status_code=400, detail="飞书个人绑定回调 URL 需要包含完整回调路径。")
        if parsed.fragment:
            raise HTTPException(status_code=400, detail="飞书个人绑定回调 URL 不能包含 # 片段。")
        return normalized

    def _is_public_feishu_callback_url(callback_url: str) -> bool:
        parsed = urlparse(callback_url)
        host = (parsed.hostname or "").strip().lower()
        if parsed.scheme != "https":
            return False
        if host in {"127.0.0.1", "localhost", "::1"} or host.endswith(".local"):
            return False
        return True

    def _default_feishu_bot_settings() -> FeishuBotSettingsRecord:
        return FeishuBotSettingsRecord(updatedAt=now_iso())

    def _fixed_review_department_records() -> list[ReviewDepartmentConfigRecord]:
        return [
            ReviewDepartmentConfigRecord(
                id=item.id,
                name=item.name,
                color=item.color,
                monthlyDna="",
                weeklyFocus="",
                leaders=[],
                members=[],
            )
            for item in list_department_catalog()
        ]

    def _default_review_governance_settings() -> ReviewGovernanceSettingsRecord:
        return ReviewGovernanceSettingsRecord(departments=_fixed_review_department_records(), updatedAt=now_iso())

    def _sanitize_review_governance_settings(
        departments: list[ReviewDepartmentConfigRecord],
        *,
        keep_updated_at: str | None = None,
    ) -> ReviewGovernanceSettingsRecord:
        fixed_departments = _fixed_review_department_records()
        incoming_by_id: dict[str, ReviewDepartmentConfigRecord] = {}
        for department in departments:
            matched = get_department_entry(department.id, department.name)
            if matched:
                incoming_by_id[matched.id] = department

        sanitized_departments: list[ReviewDepartmentConfigRecord] = []
        for fixed_department in fixed_departments:
            department = incoming_by_id.get(fixed_department.id, fixed_department)
            seen_leader_names: set[str] = set()
            leaders: list[ReviewDepartmentMemberRecord] = []
            for leader in department.leaders:
                full_name = leader.fullName.strip()
                if not full_name:
                    continue
                key = full_name.lower()
                if key in seen_leader_names:
                    continue
                seen_leader_names.add(key)
                leaders.append(
                    ReviewDepartmentMemberRecord(
                        id=leader.id.strip(),
                        fullName=full_name,
                        email=leader.email.strip() if leader.email else None,
                    )
                )
            sanitized_departments.append(
                ReviewDepartmentConfigRecord(
                    id=fixed_department.id,
                    name=fixed_department.name,
                    color=department.color.strip() or fixed_department.color,
                    monthlyDna=department.monthlyDna.strip(),
                    weeklyFocus=department.weeklyFocus.strip(),
                    leaders=leaders,
                    members=[],
                )
            )
        return ReviewGovernanceSettingsRecord(
            departments=sanitized_departments,
            updatedAt=keep_updated_at or now_iso(),
        )

    def get_review_governance_settings() -> ReviewGovernanceSettingsRecord:
        current = _load_json_settings_record(
            "settings.review_governance",
            _default_review_governance_settings,
            ReviewGovernanceSettingsRecord,
        )
        sanitized = _sanitize_review_governance_settings(current.departments, keep_updated_at=current.updatedAt)
        if sanitized.model_dump() != current.model_dump():
            state.db.set_setting("settings.review_governance", to_json(sanitized.model_dump()))
        return sanitized

    def _sync_review_governance_members(
        governance: ReviewGovernanceSettingsRecord,
        employees: list[EmployeeRecord],
    ) -> ReviewGovernanceSettingsRecord:
        employees_by_department: dict[str, list[ReviewDepartmentMemberRecord]] = {department.id: [] for department in governance.departments}
        seen_keys: dict[str, set[str]] = {department.id: set() for department in governance.departments}
        for employee in employees:
            department = get_department_entry(employee.departmentId, employee.departmentName)
            if not department or department.id not in employees_by_department:
                continue
            key = employee.id.strip() or employee.fullName.strip().lower()
            if not key or key in seen_keys[department.id]:
                continue
            seen_keys[department.id].add(key)
            employees_by_department[department.id].append(
                ReviewDepartmentMemberRecord(
                    id=employee.id,
                    fullName=employee.fullName,
                    email=employee.email,
                )
            )
        return governance.model_copy(
            update={
                "departments": [
                    department.model_copy(update={"members": employees_by_department.get(department.id, [])})
                    for department in governance.departments
                ]
            }
        )

    def _load_employee_directory_from_cloud() -> list[EmployeeRecord]:
        payload = cloud_request("GET", "/api/v1/employees/directory")
        if not isinstance(payload, list):
            raise HTTPException(status_code=502, detail="Invalid employee directory payload")
        employees: list[EmployeeRecord] = []
        for item in payload:
            if not isinstance(item, dict):
                continue
            employees.append(EmployeeRecord(**item))
        return employees

    def _review_governance_with_members() -> ReviewGovernanceSettingsRecord:
        governance = get_review_governance_settings()
        token = get_cloud_token()
        if not token:
            return governance
        try:
            employees = _load_employee_directory_from_cloud()
        except HTTPException:
            return governance
        return _sync_review_governance_members(governance, employees)

    def _user_matches_department_member(member: ReviewDepartmentMemberRecord, *, user_id: str, full_name: str) -> bool:
        member_id = member.id.strip()
        member_name = member.fullName.strip().lower()
        return bool(
            (member_id and user_id and member_id == user_id)
            or (member_name and full_name and member_name == full_name.lower())
        )

    def _review_department_for_session_user(
        session_user: SessionUserRecord | None,
        governance: ReviewGovernanceSettingsRecord,
    ) -> ReviewDepartmentConfigRecord | None:
        if session_user is None:
            return None
        user_id = session_user.id.strip()
        full_name = session_user.fullName.strip()
        if not full_name and not user_id:
            return None
        for department in governance.departments:
            if any(_user_matches_department_member(leader, user_id=user_id, full_name=full_name) for leader in department.leaders):
                return department
        return None

    def _resolve_review_viewer_role(
        session_user: SessionUserRecord | None,
        governance: ReviewGovernanceSettingsRecord | None = None,
    ) -> Literal["employee", "department_lead", "admin"]:
        if session_user is None:
            return "employee"
        if session_user.primaryRole == "admin":
            return "admin"
        if governance is not None and _review_department_for_session_user(session_user, governance) is not None:
            return "department_lead"
        return "employee"

    def _normalize_department_name(value: str | None) -> str:
        return (value or "").strip().lower()

    def _resolve_agent_execution_department_scope(
        session_user: SessionUserRecord | None,
        governance: ReviewGovernanceSettingsRecord,
        requested_department: str | None,
    ) -> str | None:
        normalized_requested = _normalize_department_name(requested_department)
        if session_user and session_user.primaryRole == "admin":
            return requested_department.strip() if requested_department and requested_department.strip() else None

        viewer_department = _review_department_for_session_user(session_user, governance)
        if viewer_department is None:
            raise HTTPException(status_code=403, detail="Only department leaders or CEO can view agent execution")

        if normalized_requested and normalized_requested != _normalize_department_name(viewer_department.name):
            raise HTTPException(status_code=403, detail="Department leaders can only view their own department")

        return viewer_department.name

    def _hydrate_feishu_bot_settings(record: FeishuBotSettingsRecord) -> FeishuBotSettingsRecord:
        fingerprint: str | None = None
        source = "unconfigured"
        try:
            fingerprint = state.feishu_secret_store.get_api_key_fingerprint()
            source = state.feishu_secret_store.get_source_label() if fingerprint else "unconfigured"
        except Exception:
            source = "unavailable"
        has_app_secret = bool(fingerprint)
        return record.model_copy(
            update={
                "appId": record.appId.strip(),
                "receiverId": record.receiverId.strip(),
                "botName": record.botName.strip() or "罗茜茜",
                "userBindingCallbackUrl": _normalize_feishu_user_binding_callback_url(record.userBindingCallbackUrl),
                "ready": bool(record.appId.strip() and record.receiverId.strip() and has_app_secret),
                "hasAppSecret": has_app_secret,
                "secretSource": source,
                "secretFingerprint": fingerprint,
            }
        )

    def _persist_feishu_bot_settings(record: FeishuBotSettingsRecord) -> FeishuBotSettingsRecord:
        sanitized = record.model_copy(
            update={
                "appId": record.appId.strip(),
                "receiverId": record.receiverId.strip(),
                "botName": record.botName.strip() or "罗茜茜",
                "userBindingCallbackUrl": _normalize_feishu_user_binding_callback_url(record.userBindingCallbackUrl),
                "updatedAt": now_iso(),
            }
        )
        state.db.set_setting("settings.feishu_bot", to_json(sanitized.model_dump()))
        return _hydrate_feishu_bot_settings(sanitized)

    def get_client_workspace_settings() -> ClientWorkspaceSettingsRecord:
        return _load_json_settings_record("settings.client_workspace", _default_client_workspace_settings, ClientWorkspaceSettingsRecord)

    def _client_hidden_folders_key(client_id: str) -> str:
        return f"client.hidden_folders:{client_id}"

    def get_hidden_client_folders(client_id: str) -> set[str]:
        raw = state.db.get_setting(_client_hidden_folders_key(client_id), "")
        parsed = from_json(raw, []) if raw else []
        if not isinstance(parsed, list):
            return set()
        return {str(item).strip() for item in parsed if str(item).strip()}

    def save_hidden_client_folders(client_id: str, labels: set[str]) -> None:
        state.db.set_setting(_client_hidden_folders_key(client_id), to_json(sorted(labels)))

    def hide_client_folder_label(client_id: str, label: str) -> None:
        hidden = get_hidden_client_folders(client_id)
        hidden.add(label)
        save_hidden_client_folders(client_id, hidden)

    def unhide_client_folder_label(client_id: str, label: str) -> None:
        hidden = get_hidden_client_folders(client_id)
        if label in hidden:
            hidden.remove(label)
            save_hidden_client_folders(client_id, hidden)

    def get_topics_settings() -> TopicsSettingsRecord:
        return _load_json_settings_record("settings.topics", _default_topics_settings, TopicsSettingsRecord)

    def get_analysis_workbench_settings() -> AnalysisWorkbenchSettingsRecord:
        current = _load_json_settings_record("settings.analysis_workbench", _default_analysis_workbench_settings, AnalysisWorkbenchSettingsRecord)
        if not current.enabledTemplateIds:
            template_ids = [str(row["id"]) for row in state.db.fetchall("SELECT id FROM analysis_templates ORDER BY created_at ASC")]
            current = current.model_copy(update={"enabledTemplateIds": template_ids, "defaultTemplateId": current.defaultTemplateId or (template_ids[0] if template_ids else None)})
            _save_json_settings_record("settings.analysis_workbench", current)
        return current

    def get_handbook_settings() -> HandbookSettingsRecord:
        return _load_json_settings_record("settings.handbook", _default_handbook_settings, HandbookSettingsRecord)

    def get_system_admin_settings() -> SystemAdminSettingsRecord:
        return _load_json_settings_record("settings.system_admin", _default_system_admin_settings, SystemAdminSettingsRecord)

    def get_feishu_bot_settings() -> FeishuBotSettingsRecord:
        record = _load_json_settings_record("settings.feishu_bot", _default_feishu_bot_settings, FeishuBotSettingsRecord)
        return _hydrate_feishu_bot_settings(record)

    def _feishu_user_binding_key(user_id: str) -> str:
        return f"settings.feishu_user_binding:{user_id}"

    def _feishu_oauth_state_key(state_token: str) -> str:
        return f"settings.feishu_oauth_state:{state_token}"

    def _feishu_user_binding_pending_key(user_id: str) -> str:
        return f"settings.feishu_user_binding_pending:{user_id}"

    def _feishu_user_binding_authorization_ready() -> bool:
        settings = get_feishu_bot_settings()
        if not settings.appId.strip():
            return False
        try:
            return bool(state.feishu_secret_store.get_api_key().strip())
        except Exception:
            return False

    def _default_feishu_user_binding(user_id: str) -> FeishuUserBindingRecord:
        return FeishuUserBindingRecord(userId=user_id, appId=get_feishu_bot_settings().appId.strip())

    def get_feishu_user_binding(user_id: str) -> FeishuUserBindingRecord:
        raw = state.db.get_setting(_feishu_user_binding_key(user_id), "")
        if raw:
            parsed = from_json(raw, {})
            if isinstance(parsed, dict):
                try:
                    record = FeishuUserBindingRecord(**parsed)
                except Exception:
                    record = _default_feishu_user_binding(user_id)
            else:
                record = _default_feishu_user_binding(user_id)
        else:
            record = _default_feishu_user_binding(user_id)
        settings = get_feishu_bot_settings()
        return record.model_copy(
            update={
                "userId": user_id,
                "appId": settings.appId.strip(),
                "linked": bool(record.openId),
                "readyForAuthorization": _feishu_user_binding_authorization_ready(),
            }
        )

    def save_feishu_user_binding(record: FeishuUserBindingRecord) -> FeishuUserBindingRecord:
        sanitized = record.model_copy(
            update={
                "linked": bool(record.openId),
                "readyForAuthorization": _feishu_user_binding_authorization_ready(),
                "appId": get_feishu_bot_settings().appId.strip(),
                "lastError": record.lastError.strip() if record.lastError else None,
            }
        )
        state.db.set_setting(_feishu_user_binding_key(sanitized.userId), to_json(sanitized.model_dump()))
        return get_feishu_user_binding(sanitized.userId)

    def clear_feishu_user_binding(user_id: str) -> FeishuUserBindingRecord:
        state.db.set_setting(_feishu_user_binding_key(user_id), "")
        return get_feishu_user_binding(user_id)

    def save_feishu_oauth_state(state_token: str, user_id: str, expires_at: str) -> None:
        state.db.set_setting(
            _feishu_oauth_state_key(state_token),
            to_json({"userId": user_id, "expiresAt": expires_at, "createdAt": now_iso()}),
        )

    def pop_feishu_oauth_state(state_token: str) -> dict[str, str] | None:
        key = _feishu_oauth_state_key(state_token)
        raw = state.db.get_setting(key, "")
        state.db.set_setting(key, "")
        if not raw:
            return None
        parsed = from_json(raw, {})
        if not isinstance(parsed, dict):
            return None
        return {str(k): str(v) for k, v in parsed.items() if v is not None}

    def clear_feishu_oauth_state(state_token: str) -> None:
        state.db.set_setting(_feishu_oauth_state_key(state_token), "")

    def get_feishu_user_binding_pending(user_id: str) -> dict[str, str] | None:
        raw = state.db.get_setting(_feishu_user_binding_pending_key(user_id), "")
        if not raw:
            return None
        parsed = from_json(raw, {})
        if not isinstance(parsed, dict):
            return None
        return {str(key): str(value) for key, value in parsed.items() if value is not None}

    def save_feishu_user_binding_pending(user_id: str, *, state_token: str, expires_at: str, callback_url: str, mode: str) -> None:
        state.db.set_setting(
            _feishu_user_binding_pending_key(user_id),
            to_json(
                {
                    "state": state_token,
                    "expiresAt": expires_at,
                    "callbackUrl": callback_url,
                    "mode": mode,
                    "updatedAt": now_iso(),
                }
            ),
        )

    def clear_feishu_user_binding_pending(user_id: str) -> None:
        state.db.set_setting(_feishu_user_binding_pending_key(user_id), "")

    def _feishu_cloud_relay_callback_url() -> str:
        return f"{state.cloud_api_url.rstrip('/')}/api/v1/integrations/feishu/user-binding/callback"

    def _save_feishu_user_binding_error(user_id: str, message: str) -> FeishuUserBindingRecord:
        existing = get_feishu_user_binding(user_id)
        return save_feishu_user_binding(existing.model_copy(update={"lastError": message, "lastVerifiedAt": now_iso()}))

    def _finalize_feishu_user_binding(user_id: str, code: str) -> FeishuUserBindingRecord:
        settings = get_feishu_bot_settings()
        if not settings.appId.strip():
            raise HTTPException(status_code=400, detail="当前工作台还没有配置飞书 App ID。")
        try:
            app_secret = state.feishu_secret_store.get_api_key().strip()
        except Exception:
            app_secret = ""
        if not app_secret:
            raise HTTPException(status_code=400, detail="当前工作台还没有配置飞书 App Secret。")

        existing = get_feishu_user_binding(user_id)
        app_access_token, _ = fetch_app_access_token(app_id=settings.appId.strip(), app_secret=app_secret)
        token_payload = exchange_authorization_code(
            app_access_token=app_access_token,
            app_id=settings.appId.strip(),
            app_secret=app_secret,
            code=code.strip(),
        )
        user_access_token = str(token_payload.get("access_token") or "").strip()
        user_info = fetch_user_info(user_access_token=user_access_token)
        binding = save_feishu_user_binding(
            FeishuUserBindingRecord(
                linked=True,
                readyForAuthorization=True,
                appId=settings.appId.strip(),
                userId=user_id,
                openId=str(user_info.get("open_id") or token_payload.get("open_id") or "").strip() or None,
                unionId=str(user_info.get("union_id") or token_payload.get("union_id") or "").strip() or None,
                feishuUserId=str(user_info.get("user_id") or token_payload.get("user_id") or "").strip() or None,
                name=str(user_info.get("name") or "").strip() or None,
                enName=str(user_info.get("en_name") or "").strip() or None,
                avatarUrl=str(user_info.get("avatar_url") or "").strip() or None,
                email=str(user_info.get("email") or "").strip() or None,
                tenantKey=str(user_info.get("tenant_key") or token_payload.get("tenant_key") or "").strip() or None,
                boundAt=existing.boundAt or now_iso(),
                lastVerifiedAt=now_iso(),
                lastError=None,
            )
        )
        if not binding.openId:
            _save_feishu_user_binding_error(user_id, "飞书没有返回 open_id，无法完成绑定。")
            raise HTTPException(status_code=400, detail="飞书没有返回 open_id，当前无法把软件账号和飞书账号关联起来。")
        log_activity("feishu.user_binding.success", "settings", user_id, {"openId": binding.openId, "email": binding.email})
        return binding

    def _clear_feishu_cloud_relay_session(user_id: str) -> None:
        pending = get_feishu_user_binding_pending(user_id)
        if pending and pending.get("state"):
            clear_feishu_oauth_state(str(pending.get("state")))
        if not pending or pending.get("mode") != "cloud_relay" or not pending.get("state"):
            clear_feishu_user_binding_pending(user_id)
            return
        try:
            cloud_request("DELETE", f"/api/v1/integrations/feishu/user-binding/sessions/{pending['state']}")
        except HTTPException:
            pass
        clear_feishu_user_binding_pending(user_id)

    def sync_feishu_user_binding_from_cloud_relay(user_id: str) -> FeishuUserBindingRecord | None:
        pending = get_feishu_user_binding_pending(user_id)
        if not pending or pending.get("mode") != "cloud_relay":
            return None
        state_token = (pending.get("state") or "").strip()
        if not state_token:
            clear_feishu_user_binding_pending(user_id)
            return None
        expires_at = (pending.get("expiresAt") or "").strip()
        if expires_at:
            try:
                if datetime.fromisoformat(expires_at) <= datetime.now():
                    clear_feishu_oauth_state(state_token)
                    _clear_feishu_cloud_relay_session(user_id)
                    return _save_feishu_user_binding_error(user_id, "这次飞书扫码授权请求已经过期，请重新发起绑定。")
            except ValueError:
                clear_feishu_oauth_state(state_token)
                _clear_feishu_cloud_relay_session(user_id)
                return _save_feishu_user_binding_error(user_id, "飞书扫码授权状态已损坏，请重新发起绑定。")
        try:
            payload = cloud_request("GET", f"/api/v1/integrations/feishu/user-binding/sessions/{state_token}")
        except HTTPException as exc:
            if exc.status_code == 404:
                clear_feishu_oauth_state(state_token)
                _clear_feishu_cloud_relay_session(user_id)
                return None
            raise
        if not isinstance(payload, dict):
            return None
        status = str(payload.get("status") or "").strip()
        if status in {"", "pending"}:
            return None
        if status == "expired":
            clear_feishu_oauth_state(state_token)
            _clear_feishu_cloud_relay_session(user_id)
            return _save_feishu_user_binding_error(user_id, "这次飞书扫码授权请求已经过期，请重新发起绑定。")
        if status == "error":
            clear_feishu_oauth_state(state_token)
            _clear_feishu_cloud_relay_session(user_id)
            return _save_feishu_user_binding_error(user_id, str(payload.get("errorMessage") or "飞书扫码授权失败，请重新发起绑定。"))
        code = str(payload.get("code") or "").strip()
        if status != "authorized" or not code:
            return None
        try:
            binding = _finalize_feishu_user_binding(user_id, code)
        except FeishuApiError as exc:
            clear_feishu_oauth_state(state_token)
            _clear_feishu_cloud_relay_session(user_id)
            return _save_feishu_user_binding_error(user_id, str(exc))
        except HTTPException as exc:
            clear_feishu_oauth_state(state_token)
            _clear_feishu_cloud_relay_session(user_id)
            return _save_feishu_user_binding_error(user_id, str(exc.detail))
        clear_feishu_oauth_state(state_token)
        _clear_feishu_cloud_relay_session(user_id)
        return binding

    def _render_feishu_binding_callback_page(title: str, detail: str, *, success: bool) -> HTMLResponse:
        tone = "#16a34a" if success else "#dc2626"
        escaped_title = html.escape(title)
        escaped_detail = html.escape(detail)
        markup = f"""<!doctype html>
<html lang=\"zh-CN\">
  <head>
    <meta charset=\"utf-8\" />
    <title>{escaped_title}</title>
    <meta name=\"viewport\" content=\"width=device-width, initial-scale=1\" />
    <style>
      body {{ font-family: -apple-system, BlinkMacSystemFont, 'PingFang SC', 'Helvetica Neue', sans-serif; background: #f8fafc; color: #0f172a; margin: 0; padding: 32px; }}
      .card {{ max-width: 560px; margin: 8vh auto; background: #fff; border: 1px solid #e2e8f0; border-radius: 24px; padding: 28px; box-shadow: 0 16px 48px rgba(15, 23, 42, 0.08); }}
      .badge {{ display: inline-flex; align-items: center; gap: 8px; font-size: 12px; font-weight: 700; color: {tone}; background: rgba(91, 123, 254, 0.08); border-radius: 999px; padding: 6px 12px; }}
      h1 {{ font-size: 24px; line-height: 1.3; margin: 18px 0 12px; }}
      p {{ font-size: 14px; line-height: 1.75; color: #475569; margin: 0 0 12px; white-space: pre-wrap; }}
    </style>
  </head>
  <body>
    <div class=\"card\">
      <div class=\"badge\">{"绑定成功" if success else "绑定失败"}</div>
      <h1>{escaped_title}</h1>
      <p>{escaped_detail}</p>
      <p>现在可以回到益语智库自用平台继续操作；如果桌面端仍显示旧状态，点击一次“刷新绑定状态”。</p>
    </div>
  </body>
</html>"""
        return HTMLResponse(markup)

    def _default_feishu_test_message(bot_name: str) -> str:
        return f"{bot_name.strip() or '罗茜茜'} 已接通成功，现在可以给你发消息了。"

    def _default_feishu_inbound_reply(bot_name: str) -> str:
        resolved_name = bot_name.strip() or "罗茜茜"
        return f"我是{resolved_name}。飞书入站链路刚接通，现在先支持固定回复；客户上下文问答还没接上。"

    def update_feishu_bot_settings(payload: FeishuBotSettingsPayload) -> FeishuBotSettingsRecord:
        current = get_feishu_bot_settings()
        next_payload = current.model_dump()
        if payload.appId is not None:
            next_payload["appId"] = payload.appId.strip()
        if payload.receiveIdType is not None:
            next_payload["receiveIdType"] = payload.receiveIdType
        if payload.receiverId is not None:
            next_payload["receiverId"] = payload.receiverId.strip()
        if payload.botName is not None:
            next_payload["botName"] = payload.botName.strip() or "罗茜茜"

        if payload.userBindingCallbackUrl is not None:
            next_payload["userBindingCallbackUrl"] = _normalize_feishu_user_binding_callback_url(payload.userBindingCallbackUrl)

        next_record = FeishuBotSettingsRecord(**next_payload)

        try:
            if payload.clearAppSecret:
                state.feishu_secret_store.delete_api_key()
            elif payload.appSecret and payload.appSecret.strip():
                state.feishu_secret_store.set_api_key(payload.appSecret.strip())
        except Exception as exc:
            raise HTTPException(status_code=500, detail=f"飞书密钥保存失败：{exc}") from exc

        if payload.sendTestMessage:
            app_id = next_record.appId.strip()
            receiver_id = next_record.receiverId.strip()
            bot_name = next_record.botName.strip() or "罗茜茜"
            app_secret = state.feishu_secret_store.get_api_key().strip()
            attempted_at = now_iso()
            if not app_id:
                next_record = next_record.model_copy(
                    update={
                        "lastConnectionStatus": "failed",
                        "lastConnectionMessage": "请先填写飞书 App ID。",
                        "lastConnectedAt": attempted_at,
                    }
                )
            elif not receiver_id:
                next_record = next_record.model_copy(
                    update={
                        "lastConnectionStatus": "failed",
                        "lastConnectionMessage": "请先填写接收方标识。",
                        "lastConnectedAt": attempted_at,
                    }
                )
            elif not app_secret:
                next_record = next_record.model_copy(
                    update={
                        "lastConnectionStatus": "failed",
                        "lastConnectionMessage": "请先保存飞书 App Secret。",
                        "lastConnectedAt": attempted_at,
                    }
                )
            else:
                try:
                    tenant_access_token, _ = fetch_tenant_access_token(app_id=app_id, app_secret=app_secret)
                    send_text_message(
                        tenant_access_token=tenant_access_token,
                        receive_id_type=next_record.receiveIdType,
                        receive_id=receiver_id,
                        text=(payload.testMessage or _default_feishu_test_message(bot_name)).strip() or _default_feishu_test_message(bot_name),
                    )
                    next_record = next_record.model_copy(
                        update={
                            "lastConnectionStatus": "success",
                            "lastConnectionMessage": f"{bot_name} 已经发出测试消息。",
                            "lastConnectedAt": attempted_at,
                            "lastTestMessageAt": attempted_at,
                        }
                    )
                except FeishuApiError as exc:
                    next_record = next_record.model_copy(
                        update={
                            "lastConnectionStatus": "failed",
                            "lastConnectionMessage": str(exc),
                            "lastConnectedAt": attempted_at,
                        }
                    )

        saved = _persist_feishu_bot_settings(next_record)
        log_activity("settings.feishu_bot.update", "settings", "feishu_bot", payload.model_dump(exclude_none=True, exclude={"appSecret"}))
        return saved

    def _parse_feishu_text_content(message_payload: dict[str, object]) -> str:
        content = message_payload.get("content")
        if isinstance(content, dict):
            return str(content.get("text") or "").strip()
        if not isinstance(content, str):
            return ""
        try:
            parsed = json.loads(content)
        except ValueError:
            return content.strip()
        if isinstance(parsed, dict):
            return str(parsed.get("text") or "").strip()
        return content.strip()

    def _send_feishu_text_message(receive_id_type: FeishuReceiveIdType, receive_id: str, text: str) -> None:
        settings = get_feishu_bot_settings()
        if not settings.appId.strip():
            raise FeishuApiError("飞书 App ID 未配置。")
        app_secret = state.feishu_secret_store.get_api_key().strip()
        if not app_secret:
            raise FeishuApiError("飞书 App Secret 未配置。")
        tenant_access_token, _ = fetch_tenant_access_token(app_id=settings.appId, app_secret=app_secret)
        send_text_message(
            tenant_access_token=tenant_access_token,
            receive_id_type=receive_id_type,
            receive_id=receive_id,
            text=text,
        )

    def _send_feishu_chat_text(chat_id: str, text: str) -> None:
        _send_feishu_text_message("chat_id", chat_id, text)

    def _resolve_feishu_meeting_delivery() -> tuple[Literal["bound_user", "configured_receiver", "none"], FeishuReceiveIdType | None, str | None, str | None]:
        session_user = get_cached_session_user()
        if session_user:
            binding = get_feishu_user_binding(session_user.id)
            if binding.linked and binding.openId:
                target_label = binding.name or binding.email or binding.openId
                return "bound_user", "open_id", binding.openId, target_label
        settings = get_feishu_bot_settings()
        app_secret = ""
        try:
            app_secret = state.feishu_secret_store.get_api_key().strip()
        except Exception:
            app_secret = ""
        if settings.appId.strip() and settings.receiverId.strip() and app_secret:
            return "configured_receiver", settings.receiveIdType, settings.receiverId.strip(), settings.receiverId.strip()
        return "none", None, None, None

    def _populate_meeting_extraction(meeting_id: str, text: str) -> tuple[int, int, int, int]:
        agenda, decisions, actions, risks, ambiguities = extract_meeting_content(text)
        state.db.execute("DELETE FROM agenda_items WHERE meeting_id = ?", (meeting_id,))
        state.db.execute("DELETE FROM decisions WHERE meeting_id = ?", (meeting_id,))
        state.db.execute("DELETE FROM action_items WHERE meeting_id = ?", (meeting_id,))
        state.db.execute("DELETE FROM risks WHERE meeting_id = ?", (meeting_id,))
        state.db.execute("DELETE FROM ambiguities WHERE meeting_id = ?", (meeting_id,))
        for index, item in enumerate(agenda):
            state.db.execute("INSERT INTO agenda_items(id, meeting_id, title, description, sort_order) VALUES(?, ?, ?, ?, ?)", (new_id("agenda"), meeting_id, item[:28], "抽取后的议程点", index))
        for item in decisions:
            state.db.execute("INSERT INTO decisions(id, meeting_id, summary, created_at) VALUES(?, ?, ?, ?)", (new_id("dec"), meeting_id, item[:120], now_iso()))
        for item, owner, confidence in actions:
            state.db.execute(
                "INSERT INTO action_items(id, meeting_id, title, owner_name, due_date, confidence, publish_status, created_at) VALUES(?, ?, ?, ?, ?, ?, 'draft', ?)",
                (new_id("act"), meeting_id, item[:120], owner, "本周", confidence, now_iso()),
            )
        for item, severity in risks:
            state.db.execute("INSERT INTO risks(id, meeting_id, summary, severity, created_at) VALUES(?, ?, ?, ?, ?)", (new_id("risk"), meeting_id, item[:120], severity, now_iso()))
        for item, candidates in ambiguities:
            state.db.execute(
                "INSERT INTO ambiguities(id, meeting_id, raw_text, candidates_json, status, created_at) VALUES(?, ?, ?, ?, 'pending', ?)",
                (new_id("amb"), meeting_id, item[:120], to_json(candidates), now_iso()),
            )
        return len(decisions), len(actions), len(risks), len(ambiguities)

    def _ingest_feishu_minutes_writeback(meeting_id: str, notes_text: str) -> MeetingDetail:
        state.db.execute(
            "UPDATE meetings SET transcript_text = ?, notes = ?, stage = 'ingested', updated_at = ? WHERE id = ?",
            (notes_text, notes_text, now_iso(), meeting_id),
        )
        state.db.execute("DELETE FROM meeting_sources WHERE meeting_id = ?", (meeting_id,))
        state.db.execute(
            "INSERT INTO meeting_sources(id, meeting_id, title, content_text, created_at) VALUES(?, ?, ?, ?, ?)",
            (new_id("ms"), meeting_id, "飞书纪要回写", notes_text, now_iso()),
        )
        decision_count, action_count, risk_count, ambiguity_count = _populate_meeting_extraction(meeting_id, notes_text)
        state.db.execute("UPDATE meetings SET stage = 'extracted', updated_at = ? WHERE id = ?", (now_iso(), meeting_id))
        log_activity(
            "feishu.meeting.writeback",
            "meeting",
            meeting_id,
            {
                "decisions": decision_count,
                "actions": action_count,
                "risks": risk_count,
                "ambiguities": ambiguity_count,
            },
        )
        return build_meeting_detail(meeting_id)

    def handle_feishu_event(payload: dict[str, object]) -> dict[str, object]:
        challenge = payload.get("challenge")
        if isinstance(challenge, str) and challenge.strip():
            return {"challenge": challenge}

        header = payload.get("header")
        header_dict = header if isinstance(header, dict) else {}
        event_type = str(header_dict.get("event_type") or payload.get("type") or "")
        event = payload.get("event")
        event_dict = event if isinstance(event, dict) else {}
        sender = event_dict.get("sender")
        sender_dict = sender if isinstance(sender, dict) else {}
        if str(sender_dict.get("sender_type") or "").strip().lower() == "app":
            return {"ok": True, "ignored": "self_message"}

        message = event_dict.get("message")
        message_dict = message if isinstance(message, dict) else {}
        chat_id = str(message_dict.get("chat_id") or "").strip()
        message_type = str(message_dict.get("message_type") or "").strip()
        message_text = _parse_feishu_text_content(message_dict)
        if event_type != "im.message.receive_v1":
            return {"ok": True, "ignored": f"unsupported_event:{event_type or 'unknown'}"}
        if message_type != "text":
            return {"ok": True, "ignored": f"unsupported_message_type:{message_type or 'unknown'}"}
        if not chat_id:
            return {"ok": True, "ignored": "missing_chat_id"}

        stripped_text = message_text.strip()
        if stripped_text.startswith("纪要回写") or stripped_text.startswith("会议纪要回写"):
            first_line, _, remainder = stripped_text.partition("\n")
            parts = first_line.split(maxsplit=1)
            if len(parts) < 2:
                try:
                    _send_feishu_chat_text(chat_id, "请按“纪要回写 meeting_xxx\\n纪要内容”发送会议纪要。")
                except FeishuApiError:
                    pass
                return {"ok": False, "error": "missing_meeting_id"}
            meeting_id = parts[1].strip()
            notes_body = remainder.strip()
            meeting_row = state.db.fetchone("SELECT id FROM meetings WHERE id = ?", (meeting_id,))
            if not meeting_row:
                try:
                    _send_feishu_chat_text(chat_id, f"没有找到会议 {meeting_id}，请确认编号后重试。")
                except FeishuApiError:
                    pass
                return {"ok": False, "error": "meeting_not_found"}
            if not notes_body:
                try:
                    _send_feishu_chat_text(chat_id, "纪要内容为空，请在第二行开始粘贴纪要正文。")
                except FeishuApiError:
                    pass
                return {"ok": False, "error": "empty_notes"}
            meeting = _ingest_feishu_minutes_writeback(meeting_id, notes_body)
            try:
                _send_feishu_chat_text(
                    chat_id,
                    f"已回写《{meeting.title}》的会议纪要，并完成结构化抽取。当前识别到 {len(meeting.actionItems)} 条行动项、{len(meeting.decisions)} 条结论。",
                )
            except FeishuApiError:
                pass
            return {"ok": True, "mode": "meeting_writeback", "meetingId": meeting_id}

        try:
            settings = get_feishu_bot_settings()
            if not settings.appId.strip():
                return {"ok": True, "ignored": "missing_app_id"}
            app_secret = state.feishu_secret_store.get_api_key().strip()
            if not app_secret:
                return {"ok": True, "ignored": "missing_app_secret"}
            _send_feishu_chat_text(chat_id, _default_feishu_inbound_reply(settings.botName))
            state.db.set_setting("settings.feishu_last_event_at", now_iso())
            log_activity("feishu.inbound.reply", "channel", chat_id, {"eventType": event_type, "messageType": message_type})
            return {"ok": True}
        except FeishuApiError as exc:
            log_activity("feishu.inbound.reply_failed", "channel", chat_id, {"eventType": event_type, "error": str(exc)})
            return {"ok": False, "error": str(exc)}

    def ensure_business_settings_editable() -> None:
        session_user = get_cached_session_user()
        if not session_user:
            return
        if session_user.primaryRole == "admin":
            return
        admin_settings = get_system_admin_settings()
        if not admin_settings.allowBusinessSettingsForEmployees:
            raise HTTPException(status_code=403, detail="当前账号不能编辑业务设置")

    def ensure_org_dna_editable() -> None:
        session_user = get_cached_session_user()
        if not session_user:
            return
        if session_user.primaryRole == "admin":
            return
        admin_settings = get_system_admin_settings()
        if not admin_settings.allowOrgDnaForEmployees:
            raise HTTPException(status_code=403, detail="当前账号不能编辑组织 DNA")

    def ensure_admin_for_sensitive_settings() -> None:
        session_user = get_cached_session_user()
        if not session_user:
            return
        if session_user.primaryRole != "admin":
            raise HTTPException(status_code=403, detail="只有管理员可以编辑该设置")

    def _organization_dna_record(module_key: str, module_title: str, row=None) -> OrganizationDnaModuleRecord:
        return OrganizationDnaModuleRecord(
            moduleKey=module_key,  # type: ignore[arg-type]
            title=module_title,
            markdownContent=str(row["markdown_content"]) if row else "",
            normalizedText=str(row["normalized_text"]) if row else "",
            summary=str(row["summary"]) if row else "",
            fileName=str(row["file_name"]) if row else None,
            contentHash=str(row["content_hash"]) if row else None,
            updatedAt=str(row["updated_at"]) if row else None,
            updatedBy=str(row["updated_by"]) if row else None,
            hasDocument=bool(row),
        )

    def _find_self_client_row():
        for candidate in SELF_CLIENT_NAME_CANDIDATES:
            row = state.db.fetchone(
                """
                SELECT *
                FROM clients
                WHERE name = ? OR alias = ?
                ORDER BY updated_at DESC
                LIMIT 1
                """,
                (candidate, candidate),
            )
            if row:
                return row
        return None

    def _extract_readiness_evidence(text: str, keywords: list[str]) -> str | None:
        normalized_text = re.sub(r"\s+", " ", text).strip()
        if not normalized_text:
            return None
        for keyword in keywords:
            if not keyword:
                continue
            hit_index = normalized_text.find(keyword)
            if hit_index < 0:
                continue
            start = max(0, hit_index - 24)
            end = min(len(normalized_text), hit_index + 40)
            snippet = normalized_text[start:end].strip()
            if not snippet:
                continue
            if start > 0:
                snippet = f"…{snippet}"
            if end < len(normalized_text):
                snippet = f"{snippet}…"
            return snippet
        return None

    def _build_dna_readiness_questions(
        module_key: str,
        normalized_text: str,
        missing_info: list[str],
    ) -> list[DnaReadinessQuestionRecord]:
        rules = DNA_READINESS_RULES.get(module_key, [])
        if not rules:
            return []
        questions: list[DnaReadinessQuestionRecord] = []
        compact_text = normalized_text or ""
        for rule in rules:
            content_keywords = [str(item) for item in (rule.get("contentKeywords") or [])]
            missing_keywords = [str(item) for item in (rule.get("missingKeywords") or [])]
            has_content = any(keyword and keyword in compact_text for keyword in content_keywords)
            blocked = any(
                keyword and keyword in str(item)
                for item in missing_info
                for keyword in missing_keywords
            )
            answered = bool(has_content and not blocked)
            evidence = _extract_readiness_evidence(compact_text, content_keywords) if answered else None
            questions.append(
                DnaReadinessQuestionRecord(
                    question=str(rule.get("question") or ""),
                    answered=answered,
                    evidence=evidence,
                )
            )
        return questions

    def _build_organization_dna_readiness(
        base_module: OrganizationDnaModuleRecord,
        *,
        client_module=None,
        auto_enqueued: bool = False,
    ) -> OrganizationDnaModuleRecord:
        preferred_text = ""
        missing_info: list[str] = []
        readiness_source: str = "none"
        if client_module and client_module.hasDocument and client_module.normalizedText.strip():
            preferred_text = client_module.normalizedText
            missing_info = list(client_module.missingInfo or [])
            readiness_source = "client_dna"
        elif base_module.hasDocument and base_module.normalizedText.strip():
            preferred_text = base_module.normalizedText
            missing_info = extract_markdown_missing_info(base_module.markdownContent)
            readiness_source = "manual_document"
        elif auto_enqueued:
            readiness_source = "auto_enqueued"

        questions = _build_dna_readiness_questions(base_module.moduleKey, preferred_text, missing_info)
        answered_count = sum(1 for item in questions if item.answered)
        question_count = len(questions)
        readiness_status: Literal["ready", "missing"] = "missing"
        if question_count > 0 and answered_count >= (2 if question_count >= 3 else question_count):
            readiness_status = "ready"
        elif question_count == 0 and preferred_text.strip():
            readiness_status = "ready"

        if readiness_source == "client_dna":
            readiness_summary = f"优先采用客户 DNA，自动判定 {answered_count}/{question_count or 0} 项明确。"
        elif readiness_source == "manual_document":
            readiness_summary = f"当前采用手工上传文档，自动判定 {answered_count}/{question_count or 0} 项明确。"
        elif readiness_source == "auto_enqueued":
            readiness_summary = "客户 DNA 仍缺失，系统已自动发起补跑。"
        else:
            readiness_summary = "当前还没有客户 DNA，也没有补充文档。"

        return base_module.model_copy(
            update={
                "readinessStatus": readiness_status,
                "readinessAnsweredCount": answered_count,
                "readinessQuestionCount": question_count,
                "readinessSource": readiness_source,
                "readinessSummary": readiness_summary,
                "readinessQuestions": questions,
            }
        )

    def list_organization_dna_modules() -> list[OrganizationDnaModuleRecord]:
        records_by_key = {
            str(row["module_key"]): row
            for row in state.db.fetchall("SELECT * FROM organization_dna_documents")
        }
        base_modules = {
            module_key: _organization_dna_record(module_key, module_title, records_by_key.get(module_key))
            for module_key, module_title in ORGANIZATION_DNA_MODULES
        }
        self_client_row = _find_self_client_row()
        client_modules_by_key = {}
        auto_enqueued_keys: set[str] = set()
        if self_client_row:
            client_id = str(self_client_row["id"])
            client_modules = list_client_dna_modules(client_id)
            client_modules_by_key = {module.moduleKey: module for module in client_modules}
            required_keys = {"organization_intro", "business_intro", "team_intro"}
            missing_client_keys = {
                module_key
                for module_key in required_keys
                if not (
                    client_modules_by_key.get(module_key)
                    and client_modules_by_key[module_key].hasDocument
                    and client_modules_by_key[module_key].normalizedText.strip()
                )
            }
            if missing_client_keys:
                job = maybe_enqueue_client_dna_generation_job(client_id)
                if job is not None:
                    auto_enqueued_keys = set(missing_client_keys)

        return [
            _build_organization_dna_readiness(
                base_modules[module_key],
                client_module=client_modules_by_key.get(module_key),
                auto_enqueued=module_key in auto_enqueued_keys,
            )
            for module_key, _module_title in ORGANIZATION_DNA_MODULES
        ]

    def _is_supported_org_dna_file_name(file_name: str) -> bool:
        lower_name = file_name.strip().lower()
        return lower_name.endswith(".md") or lower_name.endswith(".markdown") or lower_name.endswith(".docx")

    def _sanitize_text_list(values: list[str] | None) -> list[str]:
        if not values:
            return []
        cleaned: list[str] = []
        for value in values:
            text = re.sub(r"\s+", " ", str(value or "")).strip()
            if not text or text in cleaned:
                continue
            cleaned.append(text)
        return cleaned

    def read_organization_document_payload(payload: OrganizationDnaUploadPayload) -> tuple[str, str]:
        document_content = (payload.markdownContent or "").strip()
        file_name = (payload.fileName or "").strip()
        if payload.filePath:
            source_path = Path(payload.filePath).expanduser()
            if not source_path.exists() or not source_path.is_file():
                raise HTTPException(status_code=400, detail="背景文件不存在")
            if not _is_supported_org_dna_file_name(source_path.name):
                raise HTTPException(status_code=400, detail="只允许上传 .md、.markdown 或 .docx 文件")
            if source_path.suffix.lower() == ".docx":
                document_content = extract_platform_dna_text(source_path).strip()
            else:
                document_content = source_path.read_text(encoding="utf-8")
            file_name = file_name or source_path.name
        if file_name and not _is_supported_org_dna_file_name(file_name):
            raise HTTPException(status_code=400, detail="只允许上传 .md、.markdown 或 .docx 文件")
        if not document_content.strip():
            raise HTTPException(status_code=400, detail="请提供可解析的背景内容")
        return document_content, file_name or "uploaded.md"

    def upsert_organization_dna_module(module_key: str, payload: OrganizationDnaUploadPayload) -> OrganizationDnaModuleRecord:
        module_map = dict(ORGANIZATION_DNA_MODULES)
        if module_key not in module_map:
            raise HTTPException(status_code=404, detail="未知的组织 DNA 模块")
        markdown_content, file_name = read_organization_document_payload(payload)
        normalized_text = normalize_markdown_text(markdown_content)
        summary = summarize_markdown_document(module_map[module_key], normalized_text)
        content_hash = hashlib.sha256(markdown_content.encode("utf-8")).hexdigest()
        session_user = get_cached_session_user()
        updated_by = session_user.fullName if session_user else str(current_operator_row()["name"])
        timestamp = now_iso()
        state.db.execute(
            """
            INSERT INTO organization_dna_documents(
                module_key, title, markdown_content, normalized_text, summary, file_name, content_hash, updated_at, updated_by
            ) VALUES(?, ?, ?, ?, ?, ?, ?, ?, ?)
            ON CONFLICT(module_key) DO UPDATE SET
                title = excluded.title,
                markdown_content = excluded.markdown_content,
                normalized_text = excluded.normalized_text,
                summary = excluded.summary,
                file_name = excluded.file_name,
                content_hash = excluded.content_hash,
                updated_at = excluded.updated_at,
                updated_by = excluded.updated_by
            """,
            (module_key, module_map[module_key], markdown_content, normalized_text, summary, file_name, content_hash, timestamp, updated_by),
        )
        row = state.db.fetchone("SELECT * FROM organization_dna_documents WHERE module_key = ?", (module_key,))
        if not row:
            raise HTTPException(status_code=500, detail="组织 DNA 保存失败")
        return _organization_dna_record(module_key, module_map[module_key], row)

    def build_organization_dna_context(max_chars: int = 2800) -> str:
        modules = [module for module in list_organization_dna_modules() if module.hasDocument and module.normalizedText.strip()]
        if not modules:
            return ""
        lines = ["组织 DNA：以下内容代表本组织的稳定背景、业务语境和市场定位。"]
        for module in modules:
            lines.append(f"[{module.title}]\n{module.normalizedText[:700]}")
        return "\n\n".join(lines)[:max_chars]

    def _client_dna_record(client_id: str, module_key: str, module_title: str, row=None) -> ClientDnaModuleRecord:
        return ClientDnaModuleRecord(
            clientId=client_id,
            moduleKey=module_key,  # type: ignore[arg-type]
            title=module_title,
            markdownContent=str(row["markdown_content"]) if row else "",
            normalizedText=str(row["normalized_text"]) if row else "",
            summary=str(row["summary"]) if row else "",
            fileName=str(row["file_name"]) if row else None,
            contentHash=str(row["content_hash"]) if row else None,
            sourceKind=str(row["source_kind"]) if row and row["source_kind"] else "manual",  # type: ignore[arg-type]
            missingInfo=_parse_json_list(row["missing_info_json"]) if row else [],
            updatedAt=str(row["updated_at"]) if row else None,
            updatedBy=str(row["updated_by"]) if row else None,
            hasDocument=bool(row),
        )

    def extract_markdown_missing_info(markdown_content: str) -> list[str]:
        sections = re.split(r"^#{1,6}\s+", markdown_content, flags=re.MULTILINE)
        target_section = ""
        for section in sections:
            stripped = section.strip()
            if not stripped:
                continue
            if stripped.startswith("缺失信息"):
                target_section = stripped[len("缺失信息"):].strip()
                break
        if not target_section:
            return []
        lines = []
        for raw_line in target_section.splitlines():
            line = re.sub(r"^\s*[-*•\d.]+\s*", "", raw_line).strip()
            if not line:
                continue
            if line.lower().startswith("暂无"):
                continue
            lines.append(line)
        return _sanitize_text_list(lines)

    def save_client_dna_module(
        client_id: str,
        module_key: str,
        *,
        markdown_content: str,
        file_name: str,
        source_kind: Literal["manual", "generated"],
        updated_by: str,
        missing_info: list[str] | None = None,
    ) -> ClientDnaModuleRecord:
        module_map = dict(CLIENT_DNA_MODULES)
        if module_key not in module_map:
            raise HTTPException(status_code=404, detail="未知的客户 DNA 模块")
        normalized_text = normalize_markdown_text(markdown_content)
        summary = summarize_markdown_document(module_map[module_key], normalized_text)
        content_hash = hashlib.sha256(markdown_content.encode("utf-8")).hexdigest()
        timestamp = now_iso()
        final_missing_info = _sanitize_text_list(missing_info if missing_info is not None else extract_markdown_missing_info(markdown_content))
        state.db.execute(
            """
            INSERT INTO client_dna_documents(
                client_id, module_key, title, markdown_content, normalized_text, summary, file_name, content_hash,
                source_kind, missing_info_json, updated_at, updated_by
            ) VALUES(?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ON CONFLICT(client_id, module_key) DO UPDATE SET
                title = excluded.title,
                markdown_content = excluded.markdown_content,
                normalized_text = excluded.normalized_text,
                summary = excluded.summary,
                file_name = excluded.file_name,
                content_hash = excluded.content_hash,
                source_kind = excluded.source_kind,
                missing_info_json = excluded.missing_info_json,
                updated_at = excluded.updated_at,
                updated_by = excluded.updated_by
            """,
            (
                client_id,
                module_key,
                module_map[module_key],
                markdown_content,
                normalized_text,
                summary,
                file_name,
                content_hash,
                source_kind,
                to_json(final_missing_info),
                timestamp,
                updated_by,
            ),
        )
        record_client_dna_writeback(
            state.db,
            client_id=client_id,
            module_key=module_key,
            summary=summary,
            file_name=file_name,
            source_kind=source_kind,
            missing_info=final_missing_info,
        )
        row = state.db.fetchone(
            "SELECT * FROM client_dna_documents WHERE client_id = ? AND module_key = ?",
            (client_id, module_key),
        )
        if not row:
            raise HTTPException(status_code=500, detail="客户 DNA 保存失败")
        return _client_dna_record(client_id, module_key, module_map[module_key], row)

    def list_client_dna_modules(client_id: str) -> list[ClientDnaModuleRecord]:
        records_by_key = {
            str(row["module_key"]): row
            for row in state.db.fetchall(
                "SELECT * FROM client_dna_documents WHERE client_id = ?",
                (client_id,),
            )
        }
        return [
            _client_dna_record(client_id, module_key, module_title, records_by_key.get(module_key))
            for module_key, module_title in CLIENT_DNA_MODULES
        ]

    def upsert_client_dna_module(client_id: str, module_key: str, payload: OrganizationDnaUploadPayload) -> ClientDnaModuleRecord:
        build_client_summary(client_id)
        markdown_content, file_name = read_organization_document_payload(payload)
        session_user = get_cached_session_user()
        updated_by = session_user.fullName if session_user else str(current_operator_row()["name"])
        return save_client_dna_module(
            client_id,
            module_key,
            markdown_content=markdown_content,
            file_name=file_name,
            source_kind="manual",
            updated_by=updated_by,
        )

    def select_client_dna_generation_cards(client_id: str, module_key: str, *, limit: int = 8) -> list[DocumentCardRecord]:
        cards = [
            build_document_card_record(item)
            for item in fetch_document_cards(state.db, client_id, data_dir=state.data_dir, limit=32)
        ]
        if not cards:
            return []
        keyword_map: dict[str, list[str]] = {
            "organization_intro": ["组织", "机构", "公司", "使命", "历史", "定位", "介绍", "战略"],
            "business_intro": ["项目", "业务", "服务", "合作", "方案", "交付", "执行", "陪伴"],
            "team_intro": ["团队", "成员", "负责人", "岗位", "组织架构", "分工", "协作", "接口"],
            "market_intro": ["市场", "行业", "竞品", "需求", "用户", "趋势", "研究", "环境"],
        }
        keywords = keyword_map.get(module_key, [])
        if not keywords:
            return cards[:limit]

        def score(card: DocumentCardRecord) -> int:
            haystack = " ".join(
                [
                    card.title,
                    card.summary,
                    card.shortSummary,
                    card.logicalCategory or "",
                    card.logicalSubcategory or "",
                    " ".join(card.keywords),
                    " ".join(card.tags),
                    " ".join(card.entities),
                ]
            )
            total = 0
            for keyword in keywords:
                if keyword and keyword in haystack:
                    total += 1
            if card.needsReview:
                total -= 1
            return total

        prioritized = sorted(cards, key=lambda item: (score(item), item.classificationConfidence), reverse=True)
        selected = [item for item in prioritized if score(item) > 0][:limit]
        return selected or prioritized[:limit]

    def build_client_dna_candidate_markdown(client_id: str, module_key: str) -> tuple[str, list[str]]:
        module_map = dict(CLIENT_DNA_MODULES)
        module_title = module_map[module_key]
        client = build_client_summary(client_id)
        cards = select_client_dna_generation_cards(client_id, module_key)
        if not cards:
            raise RuntimeError("还没有可用于生成候选文档的资料，请先导入原始资料。")
        source_lines = []
        for index, card in enumerate(cards[:8], start=1):
            source_lines.append(
                "\n".join(
                    [
                        f"[资料 {index}] {card.title}",
                        f"分类：{card.logicalCategory or card.primaryCategory} / {card.logicalSubcategory or card.secondaryCategory}",
                        f"摘要：{card.summary or card.shortSummary}",
                        f"关键词：{'、'.join(card.keywords[:8]) if card.keywords else '无'}",
                        f"补充发现：{'；'.join(card.distinctFindings[:3]) if card.distinctFindings else '无'}",
                    ]
                )
            )
        prompt = (
            f"请基于下面资料，为项目《{client.name}》生成《{module_title}》候选 Markdown。\n"
            "输出约束：\n"
            "1. judgment 字段写“执行摘要”，100-200 字。\n"
            "2. content 字段只写 Markdown 正文，不要再包代码块。\n"
            "3. analysis 字段只写缺失信息，每行一条，不要编号解释。\n"
            "4. actions 字段写本次候选文档主要依据的资料标题，每行一条。\n"
            "5. timeline 字段写一句简短提醒，说明当前候选文档是否还需要继续补资料。\n"
            "6. 不要编造资料里没有的事实，信息不足就写进缺失信息。\n\n"
            f"资料摘录：\n{chr(10).join(source_lines)}"
        )
        system_instruction = (
            "你是企业项目资料整理助手。你的任务不是写宣传文案，而是基于已有资料，生成可被任务、日历、学习和问答系统稳定引用的项目背景候选稿。"
        )
        structured = state.ai.generate_structured(prompt, system_instruction, "")
        summary_text = re.sub(r"\s+", " ", str(structured.judgment or "")).strip() or f"{client.name} 的{module_title}候选摘要待补。"
        body_markdown = str(structured.content or "").strip() or f"## 1. 待补内容\n\n当前资料还不足以稳定生成《{module_title}》正文，建议继续补原始资料后重扫。"
        missing_items = _sanitize_text_list([line for line in str(structured.analysis or "").splitlines()])
        if not missing_items:
            missing_items = ["如果你发现内容仍偏空，请继续补充原始资料后重新扫描。"]
        missing_section = "\n".join(f"- {item}" for item in missing_items)
        markdown_content = (
            f"# 执行摘要\n\n{summary_text}\n\n"
            f"# 正文\n\n{body_markdown}\n\n"
            f"# 缺失信息\n\n{missing_section}\n"
        )
        return markdown_content, missing_items

    def resolve_client_dna_modules_for_generation(client_id: str, *, refresh_generated: bool = False) -> list[ClientDnaModuleRecord]:
        modules = list_client_dna_modules(client_id)
        candidates: list[ClientDnaModuleRecord] = []
        for module in modules:
            if not module.hasDocument:
                candidates.append(module)
                continue
            if refresh_generated and module.sourceKind == "generated":
                candidates.append(module)
        return candidates

    def maybe_enqueue_client_dna_generation_job(client_id: str, *, refresh_generated: bool = False) -> KnowledgeJobRecord | None:
        target_modules = resolve_client_dna_modules_for_generation(client_id, refresh_generated=refresh_generated)
        if not target_modules:
            return None
        pending = state.db.fetchone(
            """
            SELECT * FROM knowledge_jobs
            WHERE client_id = ? AND job_type = 'generate_client_dna_candidates' AND status IN ('queued', 'running')
            ORDER BY created_at DESC LIMIT 1
            """,
            (client_id,),
        )
        if pending:
            return KnowledgeJobRecord(
                id=str(pending["id"]),
                clientId=str(pending["client_id"]),
                jobType=str(pending["job_type"]),
                status=str(pending["status"]),  # type: ignore[arg-type]
                totalItems=int(pending["total_items"]),
                processedItems=int(pending["processed_items"]),
                lastError=str(pending["last_error"]) if pending["last_error"] else None,
                createdAt=str(pending["created_at"]),
                startedAt=str(pending["started_at"]) if pending["started_at"] else None,
                finishedAt=str(pending["finished_at"]) if pending["finished_at"] else None,
                updatedAt=str(pending["updated_at"]),
            )
        return enqueue_knowledge_job(
            client_id,
            "generate_client_dna_candidates",
            {
                "clientId": client_id,
                "moduleKeys": [module.moduleKey for module in target_modules],
                "refreshGenerated": refresh_generated,
            },
            total_items=len(target_modules),
        )

    def _project_module_record(row) -> ProjectModuleRecord:
        return ProjectModuleRecord(
            id=str(row["id"]),
            clientId=str(row["client_id"]),
            name=str(row["name"]),
            alias=str(row["alias"]) if row["alias"] else None,
            goal=str(row["goal"] or ""),
            description=str(row["description"] or ""),
            ownerName=str(row["owner_name"]) if row["owner_name"] else None,
            deliverables=_parse_json_list(row["deliverables_json"]),
            keywords=_parse_json_list(row["keywords_json"]),
            createdAt=str(row["created_at"]),
            updatedAt=str(row["updated_at"]),
        )

    def list_project_modules(client_id: str) -> list[ProjectModuleRecord]:
        build_client_summary(client_id)
        return [
            _project_module_record(row)
            for row in state.db.fetchall(
                "SELECT * FROM project_modules WHERE client_id = ? ORDER BY updated_at DESC, created_at DESC",
                (client_id,),
            )
        ]

    def _project_flow_record(row, module_name: str | None = None) -> ProjectFlowRecord:
        resolved_module_name = module_name
        if not resolved_module_name and row["module_id"]:
            module_row = state.db.fetchone("SELECT name FROM project_modules WHERE id = ?", (str(row["module_id"]),))
            resolved_module_name = str(module_row["name"]) if module_row and module_row["name"] else None
        return ProjectFlowRecord(
            id=str(row["id"]),
            clientId=str(row["client_id"]),
            moduleId=str(row["module_id"]),
            moduleName=resolved_module_name,
            name=str(row["name"]),
            description=str(row["description"] or ""),
            scenario=str(row["scenario"] or ""),
            triggerCondition=str(row["trigger_condition"] or ""),
            steps=_parse_json_list(row["steps_json"]),
            inputs=_parse_json_list(row["inputs_json"]),
            outputs=_parse_json_list(row["outputs_json"]),
            collaborators=_parse_json_list(row["collaborators_json"]),
            riskPoints=_parse_json_list(row["risk_points_json"]),
            createdAt=str(row["created_at"]),
            updatedAt=str(row["updated_at"]),
        )

    def list_project_flows(client_id: str) -> list[ProjectFlowRecord]:
        build_client_summary(client_id)
        rows = state.db.fetchall(
            """
            SELECT f.*, m.name AS module_name
            FROM project_flows f
            LEFT JOIN project_modules m ON m.id = f.module_id
            WHERE f.client_id = ?
            ORDER BY f.updated_at DESC, f.created_at DESC
            """,
            (client_id,),
        )
        return [_project_flow_record(row, str(row["module_name"]) if row["module_name"] else None) for row in rows]

    def build_project_structure(client_id: str) -> ProjectStructureResponse:
        return ProjectStructureResponse(modules=list_project_modules(client_id), flows=list_project_flows(client_id))

    def get_project_module_detail(client_id: str, module_id: str) -> ProjectModuleDetailRecord:
        module_row = state.db.fetchone(
            "SELECT * FROM project_modules WHERE id = ? AND client_id = ?",
            (module_id, client_id),
        )
        if not module_row:
            raise HTTPException(status_code=404, detail="项目模块不存在")
        module_record = _project_module_record(module_row)
        flow_rows = state.db.fetchall(
            "SELECT id, name FROM project_flows WHERE client_id = ? AND module_id = ? ORDER BY updated_at DESC, created_at DESC",
            (client_id, module_id),
        )
        task_rows = state.db.fetchall(
            """
            SELECT id, title
            FROM tasks
            WHERE client_id = ? AND project_module_id = ?
            ORDER BY updated_at DESC, created_at DESC
            LIMIT 8
            """,
            (client_id, module_id),
        )
        return ProjectModuleDetailRecord(
            **module_record.model_dump(),
            relatedTaskIds=[str(row["id"]) for row in task_rows],
            relatedTaskTitles=[str(row["title"]) for row in task_rows],
            flowIds=[str(row["id"]) for row in flow_rows],
            flowNames=[str(row["name"]) for row in flow_rows],
            contextSummary=(
                f"当前模块「{module_record.name}」聚焦 {module_record.goal or '关键交付'}，"
                f"已挂接 {len(flow_rows)} 条流程、{len(task_rows)} 条相关任务。"
            ),
        )

    def get_project_flow_detail(client_id: str, flow_id: str) -> ProjectFlowDetailRecord:
        flow_row = state.db.fetchone(
            """
            SELECT f.*, m.name AS module_name
            FROM project_flows f
            LEFT JOIN project_modules m ON m.id = f.module_id
            WHERE f.id = ? AND f.client_id = ?
            """,
            (flow_id, client_id),
        )
        if not flow_row:
            raise HTTPException(status_code=404, detail="项目流程不存在")
        flow_record = _project_flow_record(flow_row, str(flow_row["module_name"]) if flow_row["module_name"] else None)
        task_rows = state.db.fetchall(
            """
            SELECT id, title
            FROM tasks
            WHERE client_id = ? AND project_flow_id = ?
            ORDER BY updated_at DESC, created_at DESC
            LIMIT 8
            """,
            (client_id, flow_id),
        )
        return ProjectFlowDetailRecord(
            **flow_record.model_dump(),
            relatedTaskIds=[str(row["id"]) for row in task_rows],
            relatedTaskTitles=[str(row["title"]) for row in task_rows],
            contextSummary=(
                f"当前流程「{flow_record.name}」位于模块「{flow_record.moduleName or '未命名模块'}」，"
                f"已挂接 {len(task_rows)} 条相关任务。"
            ),
        )

    def resolve_project_structure_refs(
        client_id: str | None,
        project_module_id: str | None,
        project_flow_id: str | None,
        *,
        strict: bool = True,
    ) -> tuple[ProjectModuleRecord | None, ProjectFlowRecord | None]:
        if not project_module_id and not project_flow_id:
            return None, None
        if not client_id:
            if strict:
                raise HTTPException(status_code=400, detail="选择模块或流程前请先关联项目")
            return None, None
        module_record: ProjectModuleRecord | None = None
        flow_record: ProjectFlowRecord | None = None
        if project_module_id:
            module_row = state.db.fetchone(
                "SELECT * FROM project_modules WHERE id = ? AND client_id = ?",
                (project_module_id, client_id),
            )
            if not module_row:
                if strict:
                    raise HTTPException(status_code=400, detail="所选任务模块不存在或不属于当前项目")
                return None, None
            module_record = _project_module_record(module_row)
        if project_flow_id:
            flow_row = state.db.fetchone(
                """
                SELECT f.*, m.name AS module_name
                FROM project_flows f
                LEFT JOIN project_modules m ON m.id = f.module_id
                WHERE f.id = ? AND f.client_id = ?
                """,
                (project_flow_id, client_id),
            )
            if not flow_row:
                if strict:
                    raise HTTPException(status_code=400, detail="所选流程不存在或不属于当前项目")
                return module_record, None
            flow_record = _project_flow_record(flow_row, str(flow_row["module_name"]) if flow_row["module_name"] else None)
            if module_record and flow_record.moduleId != module_record.id:
                if strict:
                    raise HTTPException(status_code=400, detail="所选流程不属于当前任务模块")
                return module_record, None
            if not module_record:
                module_row = state.db.fetchone("SELECT * FROM project_modules WHERE id = ?", (flow_record.moduleId,))
                module_record = _project_module_record(module_row) if module_row else None
        return module_record, flow_record

    def _tokenize_scope_text(value: str | None, min_length: int = 2, max_length: int = 18) -> list[str]:
        return [
            item
            for item in (
                part.strip().lower()
                for part in re.split(r"[，。；、,\n\s/·\-]+", value or "")
            )
            if min_length <= len(item) <= max_length
        ]

    def _infer_task_client(
        title: str,
        desc: str,
        clients: list[ClientSummary],
    ) -> tuple[str | None, str]:
        text = f"{title}\n{desc}".strip().lower()
        normalized_clients: list[tuple[ClientSummary, list[str], list[str]]] = []
        for client in clients:
            domain = client.domain.replace("https://", "").replace("http://", "").replace("www.", "").strip()
            domain_parts = [item for item in re.split(r"[/.]+", domain) if item]
            exact_tokens = [
                item
                for item in ((client.name or "").strip().lower(), (client.alias or "").strip().lower())
                if len(item) >= 2
            ]
            support_tokens = list(
                dict.fromkeys(
                    item
                    for item in [domain.lower(), *[part.lower() for part in domain_parts]]
                    if len(item) >= 2
                )
            )
            normalized_clients.append((client, exact_tokens, support_tokens))
        if not text:
            return None, "系统暂未识别到明确项目。"
        ranked: list[tuple[int, int, ClientSummary, list[str], list[str]]] = []
        for client, exact_tokens, support_tokens in normalized_clients:
            exact_hits = [token for token in exact_tokens if token in text]
            support_hits = [token for token in support_tokens if token in text]
            score = len(exact_hits) * 3 + len(support_hits)
            if score <= 0:
                continue
            ranked.append((score, len(client.name), client, exact_hits, support_hits))
        ranked.sort(key=lambda item: (-item[0], -item[1]))
        if not ranked:
            return None, "系统暂未识别到明确项目。"
        _, _, winner, exact_hits, support_hits = ranked[0]
        confidence = "high" if exact_hits else "medium" if len(support_hits) > 1 else "low"
        if confidence == "low":
            return None, f"命中项目弱信号“{(support_hits or [winner.name])[0]}”，暂不自动回填。"
        hits = [*exact_hits, *support_hits][:2]
        return winner.id, f"系统自动识别项目：命中“{'、'.join(hits) or winner.name}”。"

    def _infer_task_event_line(
        title: str,
        desc: str,
        event_lines: list[EventLineRecord],
        *,
        current_client_id: str | None = None,
    ) -> tuple[str | None, str]:
        text = f"{title}\n{desc}".strip().lower()
        scoped_event_lines = [
            item for item in event_lines if current_client_id and (item.primaryClientId or "").strip() == current_client_id
        ] if current_client_id else []
        candidate_lines = scoped_event_lines or event_lines
        if not candidate_lines:
            return None, "当前还没有可选事件线。"
        if not text:
            if current_client_id and len(scoped_event_lines) == 1:
                return scoped_event_lines[0].id, f"当前项目下仅有一条事件线，先预填为“{scoped_event_lines[0].name}”。"
            return None, f"当前范围内共有 {len(candidate_lines)} 条事件线，可继续手动调整。"
        ranked: list[tuple[int, str, EventLineRecord, list[str], list[str]]] = []
        for event_line in candidate_lines:
            exact_tokens = [item for item in [event_line.name.strip().lower()] if len(item) >= 2]
            flattened_support = list(
                dict.fromkeys(
                    token
                    for value in (event_line.summary, event_line.intent, event_line.nextStep, event_line.stage)
                    for token in _tokenize_scope_text(value, 3, 14)
                )
            )
            exact_hits = [token for token in exact_tokens if token in text]
            support_hits = [token for token in flattened_support if token in text]
            score = len(exact_hits) * 4 + len(support_hits)
            if score <= 0:
                continue
            ranked.append((score, event_line.updatedAt, event_line, exact_hits, support_hits))
        ranked.sort(key=lambda item: (item[0], item[1]), reverse=True)
        if ranked:
            _, _, winner, exact_hits, support_hits = ranked[0]
            hits = [*exact_hits, *support_hits][:2]
            scope_label = "当前项目" if current_client_id and scoped_event_lines else "可选范围"
            hit_suffix = f"，命中“{'、'.join(hits)}”" if hits else ""
            return winner.id, f"系统已在{scope_label}内匹配到事件线“{winner.name}”{hit_suffix}。"
        if current_client_id and len(scoped_event_lines) == 1:
            return scoped_event_lines[0].id, f"当前项目下仅有一条事件线，先预填为“{scoped_event_lines[0].name}”。"
        return None, f"当前范围内共有 {len(candidate_lines)} 条事件线，可继续手动调整。"

    def _infer_task_project_module(
        title: str,
        desc: str,
        modules: list[ProjectModuleRecord],
        *,
        event_line: EventLineRecord | None = None,
    ) -> tuple[str | None, str]:
        if not modules:
            return None, "当前项目下还没有任务模块。"
        text = "\n".join(
            item
            for item in [title, desc, event_line.name if event_line else None, event_line.summary if event_line else None, event_line.intent if event_line else None, event_line.nextStep if event_line else None]
            if item
        ).strip().lower()
        if not text:
            if len(modules) == 1:
                return modules[0].id, f"当前项目下仅有 1 个模块，先预填为“{modules[0].name}”。"
            return None, f"当前项目下已有 {len(modules)} 个模块，可继续手动调整。"
        ranked: list[tuple[int, str, ProjectModuleRecord, list[str], list[str]]] = []
        for module in modules:
            exact_tokens = [
                item
                for item in ((module.name or "").strip().lower(), (module.alias or "").strip().lower())
                if len(item) >= 2
            ]
            support_tokens = list(
                dict.fromkeys(
                    token
                    for value in [module.goal, module.description, module.ownerName, *module.deliverables, *module.keywords]
                    for token in _tokenize_scope_text(value, 2, 18)
                )
            )
            exact_hits = [token for token in exact_tokens if token in text]
            support_hits = [token for token in support_tokens if token in text]
            score = len(exact_hits) * 5 + len(support_hits) * 2
            if score <= 0:
                continue
            ranked.append((score, module.name, module, exact_hits, support_hits))
        ranked.sort(key=lambda item: (-item[0], item[1]))
        if ranked:
            _, _, winner, exact_hits, support_hits = ranked[0]
            hits = [*exact_hits, *support_hits][:3]
            hit_suffix = f"，命中“{'、'.join(hits)}”" if hits else ""
            return winner.id, f"系统建议挂到模块“{winner.name}”{hit_suffix}。"
        if len(modules) == 1:
            return modules[0].id, f"当前项目下仅有 1 个模块，先预填为“{modules[0].name}”。"
        return None, f"当前项目下共有 {len(modules)} 个模块，可继续手动调整。"

    def _infer_task_project_flow(
        title: str,
        desc: str,
        flows: list[ProjectFlowRecord],
        *,
        selected_module_id: str | None = None,
        event_line: EventLineRecord | None = None,
    ) -> tuple[str | None, str]:
        scoped_flows = [item for item in flows if selected_module_id and item.moduleId == selected_module_id] if selected_module_id else flows
        if not scoped_flows:
            return None, "当前模块下还没有标准流程。"
        text = "\n".join(
            item
            for item in [title, desc, event_line.name if event_line else None, event_line.summary if event_line else None, event_line.intent if event_line else None, event_line.nextStep if event_line else None]
            if item
        ).strip().lower()
        if not text:
            if len(scoped_flows) == 1:
                return scoped_flows[0].id, f"当前范围内仅有 1 条流程，先预填为“{scoped_flows[0].name}”。"
            return None, f"当前范围内已有 {len(scoped_flows)} 条流程，可继续手动调整。"
        ranked: list[tuple[int, str, ProjectFlowRecord, list[str], list[str]]] = []
        for flow in scoped_flows:
            exact_tokens = [
                item
                for item in ((flow.name or "").strip().lower(), (flow.moduleName or "").strip().lower())
                if len(item) >= 2
            ]
            support_tokens = list(
                dict.fromkeys(
                    token
                    for value in [flow.description, flow.scenario, flow.triggerCondition, *flow.steps, *flow.inputs, *flow.outputs, *flow.collaborators, *flow.riskPoints]
                    for token in _tokenize_scope_text(value, 2, 18)
                )
            )
            exact_hits = [token for token in exact_tokens if token in text]
            support_hits = [token for token in support_tokens if token in text]
            score = len(exact_hits) * 5 + len(support_hits) * 2
            if score <= 0:
                continue
            ranked.append((score, flow.name, flow, exact_hits, support_hits))
        ranked.sort(key=lambda item: (-item[0], item[1]))
        if ranked:
            _, _, winner, exact_hits, support_hits = ranked[0]
            hits = [*exact_hits, *support_hits][:3]
            hit_suffix = f"，命中“{'、'.join(hits)}”" if hits else ""
            return winner.id, f"系统建议挂到流程“{winner.name}”{hit_suffix}。"
        if len(scoped_flows) == 1:
            return scoped_flows[0].id, f"当前范围内仅有 1 条流程，先预填为“{scoped_flows[0].name}”。"
        return None, f"当前范围内共有 {len(scoped_flows)} 条流程，可继续手动调整。"

    def _normalize_task_client_and_event_line_refs(
        client_id: str | None,
        event_line_id: str | None,
    ) -> tuple[str | None, str | None]:
        normalized_client_id = (client_id or "").strip() or None
        normalized_event_line_id = (event_line_id or "").strip() or None
        if not normalized_event_line_id:
            return normalized_client_id, None
        event_line_row = state.db.fetchone(
            "SELECT id, primary_client_id FROM event_lines WHERE id = ?",
            (normalized_event_line_id,),
        )
        if not event_line_row:
            raise HTTPException(status_code=400, detail="任务绑定的事件线无效")
        event_line_client_id = (
            str(event_line_row["primary_client_id"]).strip()
            if event_line_row["primary_client_id"]
            else None
        )
        if event_line_client_id and normalized_client_id != event_line_client_id:
            normalized_client_id = event_line_client_id
        return normalized_client_id, normalized_event_line_id

    def _build_task_scope_refresh_payload(
        task: TaskRecord,
        clients: list[ClientSummary],
        event_lines: list[EventLineRecord],
        project_structures: dict[str, ProjectStructureResponse],
    ) -> dict[str, str]:
        payload: dict[str, str] = {}
        event_line_by_id = {item.id: item for item in event_lines}
        existing_event_line = event_line_by_id.get(task.eventLineId or "")
        resolved_client_id = (task.clientId or "").strip() or ((existing_event_line.primaryClientId or "").strip() if existing_event_line else "")
        if existing_event_line and existing_event_line.primaryClientId:
            event_line_client_id = existing_event_line.primaryClientId.strip()
            if event_line_client_id and resolved_client_id != event_line_client_id:
                resolved_client_id = event_line_client_id
                payload["clientId"] = event_line_client_id

        resolved_event_line_id = (task.eventLineId or "").strip()
        if not resolved_event_line_id and resolved_client_id:
            inferred_event_line_id, _ = _infer_task_event_line(
                task.title,
                task.desc,
                event_lines,
                current_client_id=resolved_client_id,
            )
            if inferred_event_line_id:
                resolved_event_line_id = inferred_event_line_id
                payload["eventLineId"] = inferred_event_line_id
        resolved_event_line = event_line_by_id.get(resolved_event_line_id or "")

        if not resolved_client_id:
            return payload

        structure = project_structures.get(resolved_client_id)
        if structure is None:
            structure = build_project_structure(resolved_client_id)
            project_structures[resolved_client_id] = structure

        resolved_module_id = (task.projectModuleId or "").strip()
        resolved_flow_id = (task.projectFlowId or "").strip()
        if resolved_flow_id and not resolved_module_id:
            derived_module, derived_flow = resolve_project_structure_refs(
                resolved_client_id,
                None,
                resolved_flow_id,
                strict=False,
            )
            if derived_flow and derived_module:
                resolved_module_id = derived_module.id
                payload["projectModuleId"] = derived_module.id

        if not resolved_module_id:
            inferred_module_id, _ = _infer_task_project_module(
                task.title,
                task.desc,
                structure.modules,
                event_line=resolved_event_line,
            )
            if inferred_module_id:
                resolved_module_id = inferred_module_id
                payload["projectModuleId"] = inferred_module_id

        if not resolved_flow_id:
            inferred_flow_id, _ = _infer_task_project_flow(
                task.title,
                task.desc,
                structure.flows,
                selected_module_id=resolved_module_id or None,
                event_line=resolved_event_line,
            )
            if inferred_flow_id:
                payload["projectFlowId"] = inferred_flow_id
        return payload

    _EVENT_LINE_BOOTSTRAP_SKIP_KEYWORDS = ("吃饭", "健身", "体检", "相机", "飞北京", "机票", "拍照", "采购")
    _EVENT_LINE_BOOTSTRAP_HINT_KEYWORDS = (
        "合作",
        "系统",
        "方案",
        "数字化",
        "官网",
        "工作坊",
        "讨论",
        "诊断",
        "提纲",
        "资料",
        "计划",
        "战略",
        "介绍",
        "开源",
        "纪要",
        "流程",
        "汇总",
        "推进",
    )

    def _task_eligible_for_event_line_bootstrap(task: TaskRecord) -> bool:
        if not (task.clientId or "").strip():
            return False
        if (task.eventLineId or "").strip():
            return False
        title = (task.title or "").strip()
        desc = (task.desc or "").strip()
        text = f"{title}\n{desc}"
        if any(keyword in text for keyword in _EVENT_LINE_BOOTSTRAP_SKIP_KEYWORDS):
            return False
        if task.sourceType == "topic_candidate":
            return False
        if task.sourceType == "pressure_seed_doc_v2":
            return True
        if any(keyword in text for keyword in _EVENT_LINE_BOOTSTRAP_HINT_KEYWORDS):
            return True
        return len(desc) >= 40

    def _derive_event_line_name_from_task(task: TaskRecord) -> str:
        title = re.sub(r"\s+", " ", (task.title or "").strip())
        client_name = (task.clientName or "").strip()
        if client_name and title.lower().startswith(client_name.lower()):
            title = title[len(client_name):].strip(" -_:：|")
        if " " in title:
            first, rest = title.split(" ", 1)
            if rest and len(first.strip()) <= 4 and re.fullmatch(r"[A-Za-z0-9\u4e00-\u9fff]+", first.strip()):
                if any(keyword in rest for keyword in _EVENT_LINE_BOOTSTRAP_HINT_KEYWORDS):
                    title = rest.strip()
        title = title.strip(" -_:：|")
        return title[:36] or (task.title or "").strip()[:36] or "未命名事件线"

    def _derive_event_line_kind_from_task(task: TaskRecord) -> str:
        text = f"{task.title}\n{task.desc}".strip()
        if any(keyword in text for keyword in ("约见", "介绍", "讨论", "会", "对接", "沟通")):
            return "coordination_line"
        if any(keyword in text for keyword in ("问题", "阻塞", "卡点", "待补", "待确认")):
            return "issue_line"
        return "project_line"

    def _build_bootstrap_event_line_payload(task: TaskRecord) -> EventLineCreatePayload:
        project_context = task.projectContext
        line_name = _derive_event_line_name_from_task(task)
        summary = (
            (project_context.currentFocus if project_context and project_context.currentFocus else None)
            or (project_context.recentProgress if project_context and project_context.recentProgress else None)
            or (task.desc.strip() if task.desc else None)
            or f"围绕“{line_name}”持续推进。"
        )
        intent = (
            (project_context.goalSummary if project_context and project_context.goalSummary else None)
            or (task.desc.strip() if task.desc else None)
            or f"把“{line_name}”这条线的任务、会议和资料沉淀到同一上下文里。"
        )
        current_blocker = (
            (project_context.currentBlocker if project_context and project_context.currentBlocker else None)
            or (project_context.riskSummary if project_context and project_context.riskSummary else None)
        )
        next_step = (
            (project_context.nextAction if project_context and project_context.nextAction else None)
            or ("继续推进并明确下一步关键动作。" if task.status != "done" else "在已有推进基础上明确下一阶段动作。")
        )
        recent_decision = (
            (project_context.recentProgress if project_context and project_context.recentProgress else None)
            or ("本周已完成一个关键动作，可据此继续推进。" if task.status == "done" else None)
        )
        stage = (
            (project_context.stage if project_context and project_context.stage else None)
            or ("推进中" if task.status != "done" else "已有阶段结果")
        )
        participant_ids = [task.ownerId] if task.ownerId else []
        return EventLineCreatePayload(
            name=line_name,
            kind=_derive_event_line_kind_from_task(task),  # type: ignore[arg-type]
            status="active",
            stage=stage,
            summary=summary[:160] if summary else None,
            intent=intent[:200] if intent else None,
            currentBlocker=current_blocker[:120] if current_blocker else None,
            recentDecision=recent_decision[:120] if recent_decision else None,
            nextStep=next_step[:120] if next_step else None,
            ownerId=task.ownerId,
            primaryClientId=task.clientId,
            participantIds=participant_ids,
        )

    def build_client_dna_context(client_id: str, prompt: str, max_chars: int = 2200) -> str:
        modules = [module for module in list_client_dna_modules(client_id) if module.hasDocument and module.normalizedText.strip()]
        rows = state.db.fetchall("SELECT * FROM dna_terms WHERE client_id = ? ORDER BY updated_at DESC, created_at DESC LIMIT 8", (client_id,))
        if not modules and not rows:
            return ""

        tokens = tokenize(prompt)

        def module_rank(module: ClientDnaModuleRecord) -> tuple[int, int, str]:
            text = f"{module.title}\n{module.summary}\n{module.normalizedText[:2400]}".lower()
            match_count = sum(1 for token in tokens if token and token in text)
            return (1 if match_count > 0 else 0, match_count, module.updatedAt or "")

        ordered_modules = sorted(modules, key=module_rank, reverse=True)
        lines = [
            "客户背景底稿（仅用于理解客户，不作为正式引证）：",
            "背景底稿使用规则=背景底稿只用于理解客户、修正语境和帮助组织分析，不作为正式引证或确定性事实来源。",
        ]
        for module in ordered_modules[:4]:
            preview = module.summary.strip() or re.sub(r"\s+", " ", module.normalizedText).strip()[:360]
            if preview:
                lines.append(f"[{module.title}] {preview}")
        if rows:
            lines.append("客户补充词条（仅用于补足背景语境）：")
            for row in rows[:8]:
                aliases = _parse_json_list(row["aliases_json"])
                alias_text = f"；别名：{'、'.join(aliases[:4])}" if aliases else ""
                lines.append(f"- [{row['category']}] {row['canonical_name']}：{row['description']}{alias_text}")
        return "\n\n".join(lines)[:max_chars]

    def build_client_dna_priority_note(client_id: str, prompt: str, max_items: int = 3) -> str:
        modules = [module for module in list_client_dna_modules(client_id) if module.hasDocument and module.normalizedText.strip()]
        if not modules:
            return ""
        tokens = tokenize(prompt)

        def module_rank(module: ClientDnaModuleRecord) -> tuple[int, int, str]:
            text = f"{module.title}\n{module.summary}\n{module.normalizedText[:1600]}".lower()
            match_count = sum(1 for token in tokens if token and token in text)
            return (1 if match_count > 0 else 0, match_count, module.updatedAt or "")

        ordered_modules = sorted(modules, key=module_rank, reverse=True)
        top_titles = [module.title for module in ordered_modules[:max_items]]
        return "本次已优先参考客户 DNA 背景：" + "、".join(top_titles)

    def build_client_memory_background_context(
        client_id: str,
        prompt: str,
        *,
        max_chars: int = 2400,
    ) -> tuple[str, dict[str, object]]:
        notebook_response = get_client_notebook_response(state.db, client_id)
        memory_status = get_client_memory_status(state.db, client_id)
        snapshot = notebook_response.organizationNotebookSnapshot
        key_facts = notebook_response.keyFacts
        linked_event_lines = notebook_response.linkedEventLines
        if not snapshot and not key_facts and not linked_event_lines:
            return "", {"memoryBackgroundUsed": False}

        prompt_tokens = tokenize(prompt)

        def compact(value: object, *, limit: int = 220) -> str:
            text = re.sub(r"\s+", " ", str(value or "")).strip()
            return text[:limit]

        def event_line_rank(memory_snapshot) -> tuple[int, int, str]:
            text = "\n".join(
                item
                for item in (
                    compact(memory_snapshot.lineName, limit=80),
                    compact(memory_snapshot.currentStage, limit=80),
                    compact(memory_snapshot.currentWork, limit=160),
                    compact(memory_snapshot.currentBlocker, limit=160),
                    compact(memory_snapshot.recentDecision, limit=160),
                    compact(memory_snapshot.nextStep, limit=160),
                )
                if item
            ).lower()
            match_count = sum(1 for token in prompt_tokens if token and token in text)
            return (1 if match_count > 0 else 0, match_count, memory_snapshot.updatedAt or "")

        event_line_snapshots: list[object] = []
        for line in linked_event_lines[:8]:
            memory_response = get_event_line_memory_response(state.db, line.id)
            if memory_response.eventLineMemorySnapshot:
                event_line_snapshots.append(memory_response.eventLineMemorySnapshot)
        ordered_event_lines = sorted(event_line_snapshots, key=event_line_rank, reverse=True)

        lines = [
            "统一记忆背景（仅用于帮助理解组织与推进，不作为正式引证）：",
            "统一记忆使用规则=统一记忆只作为背景上下文，帮助理解组织现状、事件线推进和待澄清问题；它不能作为 citation，也不能替代原始证据。",
        ]
        if snapshot:
            notebook_parts = [
                compact(snapshot.organizationIntro, limit=200),
                compact(snapshot.collaborationRelationship, limit=220),
            ]
            notebook_text = "；".join(item for item in notebook_parts if item)
            if notebook_text:
                lines.append(f"组织笔记：{notebook_text}")
            if snapshot.currentStage:
                lines.append(f"组织当前阶段：{compact(snapshot.currentStage, limit=120)}")
            if snapshot.businessModules:
                lines.append(f"主要业务模块：{'；'.join(compact(item, limit=80) for item in snapshot.businessModules[:4])}")
            if snapshot.collaborationGoals:
                lines.append(f"当前合作目标：{'；'.join(compact(item, limit=90) for item in snapshot.collaborationGoals[:3])}")
            if snapshot.currentChallenges:
                lines.append(f"当前组织困境：{'；'.join(compact(item, limit=90) for item in snapshot.currentChallenges[:3])}")
        if key_facts:
            lines.append("已确认背景事实（来自组织笔记/澄清结果，只作背景，不作引用）：")
            for fact in key_facts[:4]:
                lines.append(f"- {compact(fact.factValue, limit=180)}")
        if ordered_event_lines:
            lines.append("相关事件线记忆（只作背景，不作引用）：")
            for item in ordered_event_lines[:3]:
                event_line_parts = [
                    f"阶段：{compact(item.currentStage, limit=60)}" if item.currentStage else "",
                    f"当前事项：{compact(item.currentWork, limit=100)}" if item.currentWork else "",
                    f"当前阻塞：{compact(item.currentBlocker, limit=100)}" if item.currentBlocker else "",
                    f"最近决策：{compact(item.recentDecision, limit=100)}" if item.recentDecision else "",
                    f"下一步：{compact(item.nextStep, limit=100)}" if item.nextStep else "",
                ]
                event_line_text = "；".join(part for part in event_line_parts if part)
                if event_line_text:
                    lines.append(f"- [{item.lineName}] {event_line_text}")
        missing_slots = list(snapshot.informationGaps[:3] if snapshot else [])
        if missing_slots:
            lines.append(f"当前待澄清槽位：{'；'.join(compact(item, limit=80) for item in missing_slots)}")

        source_labels = [
            "organization_notebook" if snapshot else "",
            "key_facts" if key_facts else "",
            "event_line_memory" if ordered_event_lines else "",
        ]
        confidence_candidates = [
            float(snapshot.confidence) if snapshot else 0.0,
            *(float(item.confidence) for item in ordered_event_lines[:3]),
        ]
        meta = {
            "memoryBackgroundUsed": True,
            "memoryBackgroundSources": [item for item in source_labels if item],
            "memoryBackgroundConfidence": round(max(confidence_candidates) if confidence_candidates else 0.0, 2),
            "memoryMissingFacts": missing_slots,
            "memoryEventLineCount": len(ordered_event_lines),
            "memoryPendingClarifications": int(memory_status.pendingClarifications),
        }
        return "\n\n".join(lines)[:max_chars], meta

    def build_client_dna_retrieval_hint(client_id: str, prompt: str, max_terms: int = 6) -> list[str]:
        modules = [module for module in list_client_dna_modules(client_id) if module.hasDocument and (module.summary.strip() or module.normalizedText.strip())]
        rows = state.db.fetchall("SELECT * FROM dna_terms WHERE client_id = ? ORDER BY updated_at DESC, created_at DESC LIMIT 8", (client_id,))
        if not modules and not rows:
            return []

        prompt_tokens = set(tokenize(prompt))
        hint_terms: list[str] = []

        def append_tokens(source_text: str) -> None:
            for token in tokenize(source_text):
                if token in prompt_tokens or token in hint_terms:
                    continue
                hint_terms.append(token)
                if len(hint_terms) >= max_terms:
                    return

        for module in sorted(modules, key=lambda item: (item.updatedAt or "", item.title), reverse=True)[:3]:
            append_tokens(f"{module.title} {module.summary or module.normalizedText[:400]}")
            if len(hint_terms) >= max_terms:
                return hint_terms[:max_terms]
        for row in rows:
            append_tokens(f"{row['category']} {row['canonical_name']} {row['description']}")
            if len(hint_terms) >= max_terms:
                return hint_terms[:max_terms]
        return hint_terms[:max_terms]

    def build_client_dna_term_context(client_id: str, max_chars: int = 1200) -> str:
        rows = state.db.fetchall("SELECT * FROM dna_terms WHERE client_id = ? ORDER BY updated_at DESC, created_at DESC LIMIT 8", (client_id,))
        if not rows:
            return ""
        lines = ["客户补充 DNA：以下词条仅作为当前客户的补充语境。"]
        for row in rows:
            aliases = _parse_json_list(row["aliases_json"])
            alias_text = f"；别名：{'、'.join(aliases[:4])}" if aliases else ""
            lines.append(f"[{row['category']}] {row['canonical_name']}：{row['description']}{alias_text}")
        return "\n".join(lines)[:max_chars]

    def refresh_cloud_session() -> SessionUserRecord:
        refresh_token = get_cloud_refresh_token()
        if not refresh_token:
            raise HTTPException(status_code=401, detail="登录状态已过期，请重新登录")
        persist_session = state.cloud_session_persistent or _has_persisted_cloud_session()
        try:
            response = httpx.request(
                "POST",
                f"{state.cloud_api_url}/api/v1/auth/refresh",
                json={"refreshToken": refresh_token},
                timeout=20.0,
            )
        except httpx.HTTPError as exc:
            raise HTTPException(status_code=502, detail=f"Cloud backend unavailable: {exc}") from exc
        if response.status_code >= 400:
            try:
                payload = response.json()
                detail = payload.get("detail") if isinstance(payload, dict) else response.text
            except Exception:
                detail = response.text
            if response.status_code in {401, 403}:
                clear_cloud_session()
            raise HTTPException(status_code=response.status_code, detail=detail or f"HTTP {response.status_code}")
        payload = response.json() if response.content else {}
        if not isinstance(payload, dict):
            raise HTTPException(status_code=502, detail="Invalid cloud refresh payload")
        token = str(payload.get("accessToken", ""))
        next_refresh_token = str(payload.get("refreshToken", ""))
        user_payload = payload.get("user")
        if not token or not next_refresh_token or not isinstance(user_payload, dict):
            raise HTTPException(status_code=502, detail="Cloud refresh payload missing session data")
        user = SessionUserRecord(**user_payload)
        set_cloud_session(token, user, persist=persist_session)
        set_cloud_refresh_token(next_refresh_token, persist=persist_session)
        return user

    def cloud_request(method: str, path: str, *, json_body: dict | None = None, allow_unauthenticated: bool = False) -> object:
        def perform_request(token: str | None):
            headers = {}
            if token:
                headers["Authorization"] = f"Bearer {token}"
            return httpx.request(
                method,
                f"{state.cloud_api_url}{path}",
                json=json_body,
                headers=headers,
                timeout=20.0,
            )

        token = get_cloud_token()
        if not token and not allow_unauthenticated:
            if get_cloud_refresh_token():
                refresh_cloud_session()
                token = get_cloud_token()
            else:
                raise HTTPException(status_code=401, detail="Not authenticated")
        try:
            response = perform_request(token)
        except httpx.HTTPError as exc:
            raise HTTPException(status_code=502, detail=f"Cloud backend unavailable: {exc}") from exc
        if response.status_code == 401 and not allow_unauthenticated and get_cloud_refresh_token():
            refresh_cloud_session()
            token = get_cloud_token()
            try:
                response = perform_request(token)
            except httpx.HTTPError as exc:
                raise HTTPException(status_code=502, detail=f"Cloud backend unavailable: {exc}") from exc
        if response.status_code == 401 and not allow_unauthenticated:
            clear_cloud_session()
        if response.status_code == 403 and not allow_unauthenticated and path.startswith("/api/v1/auth/"):
            clear_cloud_session()
        if response.status_code >= 400:
            try:
                payload = response.json()
                detail = payload.get("detail") if isinstance(payload, dict) else response.text
            except Exception:
                detail = response.text
            raise HTTPException(status_code=response.status_code, detail=detail or f"HTTP {response.status_code}")
        if not response.content:
            return {}
        return response.json()

    def cloud_upload_file(
        path: str,
        *,
        file_name: str,
        file_content: bytes,
        content_type: str = "application/octet-stream",
        form_fields: dict[str, str] | None = None,
    ) -> object:
        token = get_cloud_token()
        if not token:
            raise HTTPException(status_code=401, detail="Not authenticated")
        headers = {"Authorization": f"Bearer {token}"}
        files = {"file": (file_name, file_content, content_type)}
        data = form_fields or {}
        try:
            response = httpx.post(
                f"{state.cloud_api_url}{path}",
                headers=headers,
                files=files,
                data=data,
                timeout=60.0,
            )
        except httpx.HTTPError as exc:
            raise HTTPException(status_code=502, detail=f"Cloud upload unavailable: {exc}") from exc
        if response.status_code >= 400:
            try:
                detail = response.json().get("detail", response.text) if response.content else response.text
            except Exception:
                detail = response.text
            raise HTTPException(status_code=response.status_code, detail=detail or f"HTTP {response.status_code}")
        return response.json() if response.content else {}

    def require_session_user() -> SessionUserRecord:
        payload = cloud_request("GET", "/api/v1/auth/me")
        if not isinstance(payload, dict):
            raise HTTPException(status_code=502, detail="Invalid auth response")
        user = SessionUserRecord(**payload)
        set_cloud_session(get_cloud_token(), user)
        return user

    def _parse_iso_moment(value: str | None) -> datetime | None:
        if not value:
            return None
        try:
            return datetime.fromisoformat(str(value).replace("Z", "+00:00"))
        except ValueError:
            return None

    def list_cloud_consultation_knowledge_requests(
        status_filter: Literal["pending", "processing", "completed", "failed"] | None = None,
    ) -> list[ConsultationKnowledgeRequestRecord]:
        suffix = f"?status={quote(status_filter)}" if status_filter else ""
        payload = cloud_request("GET", f"/api/v1/consultation/knowledge-requests{suffix}")
        if not isinstance(payload, list):
            raise HTTPException(status_code=502, detail="Invalid consultation knowledge request payload")
        return [ConsultationKnowledgeRequestRecord(**item) for item in payload if isinstance(item, dict)]

    def update_cloud_consultation_knowledge_request_status(
        request_id: str,
        *,
        status: Literal["processing", "completed", "failed"],
        error_message: str = "",
        local_document_id: str | None = None,
        local_document_path: str | None = None,
    ) -> ConsultationKnowledgeRequestRecord:
        payload = cloud_request(
            "POST",
            f"/api/v1/consultation/knowledge-requests/{request_id}/status",
            json_body={
                "status": status,
                "errorMessage": error_message,
                "localDocumentId": local_document_id,
                "localDocumentPath": local_document_path,
            },
        )
        if not isinstance(payload, dict):
            raise HTTPException(status_code=502, detail="Invalid consultation knowledge status payload")
        return ConsultationKnowledgeRequestRecord(**payload)

    def _should_retry_consultation_knowledge_request(request: ConsultationKnowledgeRequestRecord, *, now_ts: float) -> bool:
        if request.status == "pending":
            return True
        if request.status != "processing":
            return False
        updated_at = _parse_iso_moment(request.updatedAt)
        if updated_at is None:
            return True
        return now_ts - updated_at.timestamp() >= 300

    def resolve_growth_actor() -> tuple[str, str]:
        session_user = get_cached_session_user()
        if get_cloud_token() and session_user is None:
            session_user = require_session_user()
        if session_user:
            return session_user.id, session_user.fullName or session_user.email or "当前用户"
        operator = current_operator_row()
        return str(operator["id"]), str(operator["name"] or "当前用户")

    def resolve_growth_week_label(user_id: str, requested_week: str | None = None) -> str:
        if requested_week:
            return requested_week
        current_week = current_review_week_label()
        current_week_row = state.db.fetchone(
            """
            SELECT 1
            FROM xp_ledger
            WHERE user_id = ? AND reversed_at IS NULL AND week_label = ?
            LIMIT 1
            """,
            (user_id, current_week),
        )
        if current_week_row:
            return current_week
        latest_row = state.db.fetchone(
            """
            SELECT week_label
            FROM xp_ledger
            WHERE user_id = ? AND reversed_at IS NULL AND week_label <> ''
            ORDER BY created_at DESC
            LIMIT 1
            """,
            (user_id,),
        )
        if latest_row and str(latest_row["week_label"]).strip():
            return str(latest_row["week_label"])
        return current_week

    def task_due_label(due_date: str | None) -> str:
        if not due_date:
            return "今天"
        try:
            normalized = due_date.replace("Z", "+00:00")
            moment = datetime.fromisoformat(normalized)
        except ValueError:
            return due_date
        date = moment.date()
        today = datetime.now().date()
        has_time = bool(re.match(r"^\d{4}-\d{2}-\d{2}[T\s]\d{2}:\d{2}", due_date))
        time_label = moment.strftime("%H:%M") if has_time else ""
        if date == today:
            return f"今天 {time_label}".strip()
        return f"{date.strftime('%m-%d')} {time_label}".strip()

    def normalize_due_date_input(value: str | None) -> str | None:
        if not value:
            return None
        text = value.strip()
        if not text:
            return None
        match = re.match(r"^(\d{4}-\d{2}-\d{2})[T\s](\d{2}:\d{2})(?::\d{2})?$", text)
        if match:
            return f"{match.group(1)}T{match.group(2)}"
        today = datetime.now().date()
        weekdays = {"周一": 0, "周二": 1, "周三": 2, "周四": 3, "周五": 4, "周六": 5, "周日": 6}
        if text == "今天":
            return today.isoformat()
        if text == "本周":
            return today.isoformat()
        if text in weekdays:
            delta = (weekdays[text] - today.weekday() + 7) % 7
            return (today + timedelta(days=delta)).isoformat()
        match = re.match(r"^(\d{2})-(\d{2})$", text)
        if match:
            month = int(match.group(1))
            day = int(match.group(2))
            return datetime(today.year, month, day).date().isoformat()
        return text if re.match(r"^\d{4}-\d{2}-\d{2}$", text) else None

    def build_cloud_task_tag(payload: dict[str, object]) -> TaskTagRecord:
        return TaskTagRecord(
            id=str(payload.get("id", "")),
            name=str(payload.get("name", "")),
            color=str(payload.get("color") or ("#9CA3AF" if str(payload.get("scope", "org")) == "self" else "#5B7BFE")),
            scope=str(payload.get("scope", "org")),  # type: ignore[arg-type]
            ownerUserId=str(payload.get("ownerUserId")) if payload.get("ownerUserId") else None,
            createdBy=str(payload.get("createdBy")) if payload.get("createdBy") else None,
            updatedAt=str(payload.get("updatedAt", now_iso())),
            archivedAt=str(payload.get("archivedAt")) if payload.get("archivedAt") else None,
        )

    def build_task_attachment(row) -> TaskAttachmentRecord:
        summary_value = None
        if hasattr(row, "keys"):
            row_keys = set(row.keys())
            if "document_excerpt" in row_keys:
                summary_value = row["document_excerpt"]
            elif "excerpt" in row_keys:
                summary_value = row["excerpt"]
        return TaskAttachmentRecord(
            id=str(row["id"]),
            taskId=str(row["task_id"]),
            clientId=str(row["client_id"]),
            eventLineId=str(row["event_line_id"]) if row["event_line_id"] else None,
            documentId=str(row["document_id"]) if row["document_id"] else None,
            title=str(row["title"]),
            summary=str(summary_value) if summary_value else None,
            path=str(row["path"]),
            kind=str(row["kind"]),
            source=str(row["source"]),
            sizeBytes=int(row["size_bytes"] or 0),
            createdAt=str(row["created_at"]),
        )

    def fetch_task_attachments(task_id: str, *, cloud: bool) -> list[TaskAttachmentRecord]:
        table_name = "task_attachments_cloud" if cloud else "task_attachments"
        rows = state.db.fetchall(
            f"""
            SELECT
                a.*,
                d.excerpt AS document_excerpt
            FROM {table_name} a
            LEFT JOIN documents d ON d.id = a.document_id
            WHERE a.task_id = ?
            ORDER BY a.created_at DESC
            """,
            (task_id,),
        )
        return [build_task_attachment(row) for row in rows]

    def build_attachment_event_line_activity(attachment: TaskAttachmentRecord) -> EventLineActivityRecord:
        return EventLineActivityRecord(
            id=f"attachment-activity:{attachment.id}",
            eventLineId=attachment.eventLineId or "",
            sourceType="attachment",
            sourceId=attachment.id,
            happenedAt=attachment.createdAt,
            actorId=None,
            actorName=None,
            title=f"上传附件：{attachment.title}",
            summary=f"任务附件已进入项目资料库：{attachment.title}",
            metadata={
                "taskId": attachment.taskId,
                "documentId": attachment.documentId,
                "clientId": attachment.clientId,
                "path": attachment.path,
            },
        )

    def stage_task_attachment_upload(client_id: str, file_name: str, content: bytes) -> Path:
        folders = ensure_client_workspace(state.data_dir, client_id)
        base_root = folders.get("项目与业务") or next(iter(folders.values()))
        target_root = base_root / "任务附件"
        target_root.mkdir(parents=True, exist_ok=True)
        safe_name = safe_filename(file_name or "task-attachment")
        candidate = target_root / safe_name
        if candidate.exists():
            stem = safe_filename(Path(safe_name).stem or "task-attachment")
            candidate = target_root / f"{stem}__{uuid4().hex[:6]}{Path(safe_name).suffix.lower()}"
        candidate.write_bytes(content)
        return candidate

    def build_cloud_task(payload: dict[str, object], lists_by_id: dict[str, TaskListRecord]) -> TaskRecord:
        def ensure_local_cloud_event_line_shadow(
            event_line_id: str | None,
            *,
            fallback_name: str | None = None,
            client_id: str | None = None,
        ) -> None:
            normalized_id = (event_line_id or "").strip()
            if not normalized_id or not get_cloud_token():
                return
            try:
                response_payload = cloud_request("GET", f"/api/v1/event-lines/{normalized_id}")
            except HTTPException:
                return
            if not isinstance(response_payload, dict):
                return
            record = (
                response_payload.get("eventLine")
                if isinstance(response_payload.get("eventLine"), dict)
                else response_payload
            )
            if not isinstance(record, dict):
                return
            existing_row = state.db.fetchone("SELECT created_at FROM event_lines WHERE id = ?", (normalized_id,))
            timestamp = now_iso()
            resolved_client_id = str(record.get("primaryClientId") or client_id or "").strip() or None
            client_row = (
                state.db.fetchone("SELECT name FROM clients WHERE id = ?", (resolved_client_id,))
                if resolved_client_id
                else None
            )
            participant_ids = [
                str(item).strip()
                for item in (record.get("participantIds") or [])
                if str(item).strip()
            ]
            state.db.execute(
                """
                INSERT INTO event_lines(
                    id, name, kind, status, business_category, stage, summary, intent, current_blocker, recent_decision, next_step,
                    evidence_count, owner_id, owner_name, primary_client_id, primary_client_name, primary_department_id,
                    primary_department_name, participant_ids_json, created_at, updated_at
                ) VALUES(?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                ON CONFLICT(id) DO UPDATE SET
                    name = excluded.name,
                    kind = excluded.kind,
                    status = excluded.status,
                    business_category = excluded.business_category,
                    stage = excluded.stage,
                    summary = excluded.summary,
                    intent = excluded.intent,
                    current_blocker = excluded.current_blocker,
                    recent_decision = excluded.recent_decision,
                    next_step = excluded.next_step,
                    evidence_count = excluded.evidence_count,
                    owner_id = excluded.owner_id,
                    owner_name = excluded.owner_name,
                    primary_client_id = excluded.primary_client_id,
                    primary_client_name = excluded.primary_client_name,
                    primary_department_id = excluded.primary_department_id,
                    primary_department_name = excluded.primary_department_name,
                    participant_ids_json = excluded.participant_ids_json,
                    updated_at = excluded.updated_at
                """,
                (
                    normalized_id,
                    str(record.get("name") or fallback_name or "").strip() or normalized_id,
                    str(record.get("kind") or "custom"),
                    str(record.get("status") or "active"),
                    str(record.get("businessCategory")) if record.get("businessCategory") else None,
                    str(record.get("stage")) if record.get("stage") else None,
                    str(record.get("summary")) if record.get("summary") else None,
                    str(record.get("intent")) if record.get("intent") else None,
                    str(record.get("currentBlocker")) if record.get("currentBlocker") else None,
                    str(record.get("recentDecision")) if record.get("recentDecision") else None,
                    str(record.get("nextStep")) if record.get("nextStep") else None,
                    int(record.get("evidenceCount") or 0),
                    str(record.get("ownerId")) if record.get("ownerId") else None,
                    str(record.get("ownerName")) if record.get("ownerName") else None,
                    resolved_client_id,
                    str(client_row["name"]) if client_row and client_row["name"] else (str(record.get("primaryClientName")) if record.get("primaryClientName") else None),
                    str(record.get("primaryDepartmentId")) if record.get("primaryDepartmentId") else None,
                    str(record.get("primaryDepartmentName")) if record.get("primaryDepartmentName") else None,
                    to_json(participant_ids),
                    str(existing_row["created_at"]) if existing_row and existing_row["created_at"] else timestamp,
                    timestamp,
                ),
            )
            refresh_event_line_memory_snapshot(state.db, normalized_id)
            _invalidate_event_line_snapshot_cache(normalized_id)

        cloud_note = str(payload.get("note")) if payload.get("note") else None
        note_row = state.db.fetchone("SELECT note FROM task_notes_cloud WHERE task_id = ?", (str(payload.get("id")),))
        resolved_note = cloud_note or (str(note_row["note"]) if note_row and note_row["note"] else None)
        client_id = str(payload.get("clientId")) if payload.get("clientId") else None
        event_line_id = str(payload.get("eventLineId")) if payload.get("eventLineId") else None
        event_line_name = str(payload.get("eventLineName")) if payload.get("eventLineName") else None
        ensure_local_cloud_event_line_shadow(event_line_id, fallback_name=event_line_name, client_id=client_id)
        client_row = state.db.fetchone("SELECT name FROM clients WHERE id = ?", (client_id,)) if client_id else None
        project_module_id = str(payload.get("projectModuleId")) if payload.get("projectModuleId") else None
        project_flow_id = str(payload.get("projectFlowId")) if payload.get("projectFlowId") else None
        project_module, project_flow = resolve_project_structure_refs(client_id, project_module_id, project_flow_id, strict=False) if client_id else (None, None)
        collaborators = [
            TaskCollaboratorRecord(
                userId=str(item.get("userId", "")),
                fullName=str(item.get("fullName", "")),
                email=str(item.get("email", "")),
                orderIndex=int(item.get("orderIndex", 0)),
                isOwner=bool(item.get("isOwner", False)),
                inboxStatus=str(item.get("inboxStatus", "pending")),  # type: ignore[arg-type]
                returnReason=item.get("returnReason") if isinstance(item.get("returnReason"), str) else None,
                handledAt=item.get("handledAt") if isinstance(item.get("handledAt"), str) else None,
            )
            for item in payload.get("collaborators", []) if isinstance(item, dict)
        ]
        list_id = str(payload.get("listId", "list-0"))
        list_record = lists_by_id.get(
            list_id,
            TaskListRecord(
                id=list_id,
                name=str(payload.get("listName", "收集箱")),
                color=str(payload.get("listColor", "#888681")),
                sortOrder=0,
                isDefault=False,
                archivedAt=None,
            ),
        )
        progress_status = str(payload.get("progressStatus", "todo"))
        viewer_status = payload.get("viewerInboxStatus")
        task_status = "inbox" if viewer_status == "pending" else progress_status
        org_context_payload = payload.get("orgContext")
        org_context = (
            TaskOrgContextRecord(**org_context_payload)
            if isinstance(org_context_payload, dict)
            else None
        )
        project_context = build_task_project_context(
            client_id,
            str(payload.get("sourceType", "manual")),
            str(payload.get("sourceId")) if payload.get("sourceId") else None,
            task_title=str(payload.get("title", "")),
            task_desc=str(payload.get("description", "")),
            project_module_id=project_module.id if project_module else None,
            project_flow_id=project_flow.id if project_flow else None,
        )
        def resolve_cloud_event_line_context(
            normalized_id: str,
            fallback_name: str | None,
        ) -> dict[str, object] | None:
            if not get_cloud_token():
                return None
            try:
                event_line_payload = cloud_request("GET", f"/api/v1/event-lines/{normalized_id}")
            except HTTPException:
                return None
            if not isinstance(event_line_payload, dict):
                return None
            record = event_line_payload.get("eventLine") if isinstance(event_line_payload.get("eventLine"), dict) else event_line_payload
            if not isinstance(record, dict):
                return None
            return {
                "id": str(record.get("id") or normalized_id),
                "name": str(record.get("name") or fallback_name or "").strip(),
                "businessCategory": record.get("businessCategory"),
                "stage": record.get("stage"),
                "summary": record.get("summary"),
                "intent": record.get("intent"),
                "currentBlocker": record.get("currentBlocker"),
                "recentDecision": record.get("recentDecision"),
                "nextStep": record.get("nextStep"),
                "evidenceCount": int(record.get("evidenceCount") or 0),
                "primaryClientId": record.get("primaryClientId"),
                "primaryClientName": record.get("primaryClientName"),
            }

        event_line_context = _event_line_snapshot_context(
            state.db,
            event_line_id,
            event_line_name,
            cloud_resolver=resolve_cloud_event_line_context,
        )
        _sync_task_attachment_scope(
            state.db,
            state.data_dir,
            build_task_attachment,
            build_attachment_event_line_activity,
            ensure_standard_client_folders,
            str(payload.get("id")),
            client_id,
            event_line_id,
            cloud=True,
        )
        attachments = fetch_task_attachments(str(payload.get("id")), cloud=True)
        (
            business_category,
            current_blocker,
            next_action,
            recent_decision,
            evidence_count,
        ) = _resolve_task_action_os_fields(
            title=str(payload.get("title", "")),
            desc=str(payload.get("description", "")),
            source_type=str(payload.get("sourceType", "manual")),
            business_category=str(payload.get("businessCategory")) if payload.get("businessCategory") else None,
            current_blocker=str(payload.get("currentBlocker")) if payload.get("currentBlocker") else None,
            next_action=str(payload.get("nextAction")) if payload.get("nextAction") else None,
            recent_decision=str(payload.get("recentDecision")) if payload.get("recentDecision") else None,
            evidence_count=int(payload.get("evidenceCount") or 0),
            project_context=project_context,
            event_line_context=event_line_context,
            attachment_count=len(attachments),
        )
        memory_hints, background_readiness, linked_facts_preview = get_task_memory_enrichment(
            state.db,
            task_id=str(payload.get("id")),
            client_id=client_id,
            event_line_id=event_line_id,
        )
        return TaskRecord(
            id=str(payload.get("id")),
            title=str(payload.get("title", "")),
            desc=str(payload.get("description", "")),
            status=task_status,  # type: ignore[arg-type]
            creatorId=str(payload.get("creatorId")) if payload.get("creatorId") else None,
            creatorName=str(payload.get("creatorName")) if payload.get("creatorName") else None,
            priority=str(payload.get("priority", "normal")),  # type: ignore[arg-type]
            listId=list_record.id,
            listName=list_record.name,
            listColor=list_record.color,
            ddl=task_due_label(payload.get("dueDate") if isinstance(payload.get("dueDate"), str) else None),
            dueDate=payload.get("dueDate") if isinstance(payload.get("dueDate"), str) else None,
            durationMinutes=int(payload.get("durationMinutes") or 60),
            scopeMode=str(payload.get("scopeMode") or "COLLAB_SHARED"),  # type: ignore[arg-type]
            clientId=client_id,
            clientName=str(client_row["name"]) if client_row else None,
            eventLineId=event_line_id,
            eventLineName=event_line_name,
            projectModuleId=project_module.id if project_module else project_module_id,
            projectModuleName=project_module.name if project_module else None,
            projectFlowId=project_flow.id if project_flow else project_flow_id,
            projectFlowName=project_flow.name if project_flow else None,
            ownerId=str(payload.get("ownerId")) if payload.get("ownerId") else None,
            ownerName=str(payload.get("ownerName") or ""),
            sourceType=str(payload.get("sourceType", "manual")),
            sourceId=str(payload.get("sourceId")) if payload.get("sourceId") else None,
            businessCategory=business_category,
            currentBlocker=current_blocker,
            nextAction=next_action,
            recentDecision=recent_decision,
            evidenceCount=evidence_count,
            tags=[build_cloud_task_tag(item) for item in payload.get("tags", []) if isinstance(item, dict)],
            note=resolved_note,
            attachments=attachments,
            collaborators=collaborators,
            collaborationSummary=payload.get("collaborationSummary") if isinstance(payload.get("collaborationSummary"), dict) else {},
            viewerInboxStatus=viewer_status if isinstance(viewer_status, str) else None,
            orgContext=org_context,
            projectContext=project_context,
            memoryHints=memory_hints,
            backgroundReadiness=background_readiness,
            linkedFactsPreview=linked_facts_preview,
            createdAt=str(payload.get("createdAt", now_iso())),
            updatedAt=str(payload.get("updatedAt", now_iso())),
        )

    def cloud_task_board() -> TaskBoardResponse:
        payload = cloud_request("GET", "/api/v1/tasks")
        if not isinstance(payload, dict):
            raise HTTPException(status_code=502, detail="Invalid task board payload")
        lists = [
            TaskListRecord(
                id=str(item["id"]),
                name=str(item["name"]),
                color=str(item["color"]),
                sortOrder=int(item.get("sortOrder", 0)),
                isDefault=bool(item.get("isDefault", False)),
                scope=str(item.get("scope") or "org"),
                archivedAt=str(item.get("archivedAt")) if item.get("archivedAt") else None,
            )
            for item in payload.get("lists", [])
            if isinstance(item, dict)
        ]
        lists_by_id = {item.id: item for item in lists}
        tasks = [build_cloud_task(item, lists_by_id) for item in payload.get("tasks", []) if isinstance(item, dict)]
        cloud_tags = [build_cloud_task_tag(item) for item in payload.get("tags", []) if isinstance(item, dict)]
        cloud_common_tags = [str(item) for item in payload.get("commonTags", []) if isinstance(item, str)]
        return TaskBoardResponse(tasks=tasks, lists=lists, tags=cloud_tags, commonTags=cloud_common_tags)

    def fetch_cloud_task_by_id(task_id: str) -> TaskRecord:
        board = cloud_task_board()
        task = next((item for item in board.tasks if item.id == task_id), None)
        if not task:
            raise HTTPException(status_code=404, detail="Task not found")
        return task

    def task_lists() -> list[TaskListRecord]:
        return [
            _local_task_list_record(row)
            for row in state.db.fetchall(
                """
                SELECT *
                FROM task_lists
                ORDER BY CASE WHEN archived_at IS NULL OR archived_at = '' THEN 0 ELSE 1 END,
                         CASE WHEN is_default = 1 THEN 0 ELSE 1 END,
                         sort_order ASC,
                         name COLLATE NOCASE ASC
                """
            )
        ]

    def task_tags() -> list[TaskTagRecord]:
        operator_row = current_operator_row()
        return _visible_local_task_tags(state.db, str(operator_row["id"]))

    def build_task(row) -> TaskRecord:
        note_row = state.db.fetchone("SELECT note FROM task_notes WHERE task_id = ?", (str(row["id"]),))
        client_id = str(row["client_id"]) if row["client_id"] else None
        event_line_id = str(row["event_line_id"]) if row["event_line_id"] else None
        event_line_row = state.db.fetchone("SELECT name FROM event_lines WHERE id = ?", (event_line_id,)) if event_line_id else None
        project_module_id = str(row["project_module_id"]) if row["project_module_id"] else None
        project_flow_id = str(row["project_flow_id"]) if row["project_flow_id"] else None
        project_module, project_flow = resolve_project_structure_refs(client_id, project_module_id, project_flow_id, strict=False) if client_id else (None, None)
        project_context = build_task_project_context(
            client_id,
            str(row["source_type"]),
            str(row["source_id"]) if row["source_id"] else None,
            task_title=str(row["title"]),
            task_desc=str(row["description"]),
            project_module_id=project_module.id if project_module else project_module_id,
            project_flow_id=project_flow.id if project_flow else project_flow_id,
        )
        event_line_context = _event_line_snapshot_context(state.db, event_line_id, str(event_line_row["name"]) if event_line_row else None)
        attachments = fetch_task_attachments(str(row["id"]), cloud=False)
        (
            business_category,
            current_blocker,
            next_action,
            recent_decision,
            evidence_count,
        ) = _resolve_task_action_os_fields(
            title=str(row["title"]),
            desc=str(row["description"]),
            source_type=str(row["source_type"]),
            business_category=str(row["business_category"]) if row["business_category"] else None,
            current_blocker=str(row["current_blocker"]) if row["current_blocker"] else None,
            next_action=str(row["next_action"]) if row["next_action"] else None,
            recent_decision=str(row["recent_decision"]) if row["recent_decision"] else None,
            evidence_count=int(row["evidence_count"] or 0),
            project_context=project_context,
            event_line_context=event_line_context,
            attachment_count=len(attachments),
        )
        memory_hints, background_readiness, linked_facts_preview = get_task_memory_enrichment(
            state.db,
            task_id=str(row["id"]),
            client_id=client_id,
            event_line_id=event_line_id,
        )
        return TaskRecord(
            id=str(row["id"]),
            title=str(row["title"]),
            desc=str(row["description"]),
            status=str(row["status"]),  # type: ignore[arg-type]
            priority=str(row["priority"]),  # type: ignore[arg-type]
            listId=str(row["list_id"]),
            listName=str(row["list_name"]),
            listColor=str(row["list_color"]),
            ddl=str(row["ddl"]),
            dueDate=str(row["due_date"]) if row["due_date"] else None,
            durationMinutes=int(row["duration_minutes"] or 60),
            scopeMode=str(row["scope_mode"] or "COLLAB_SHARED"),  # type: ignore[arg-type]
            clientId=client_id,
            clientName=str(row["client_name"]) if row["client_name"] else None,
            eventLineId=event_line_id,
            eventLineName=str(event_line_row["name"]) if event_line_row else None,
            projectModuleId=project_module.id if project_module else project_module_id,
            projectModuleName=project_module.name if project_module else None,
            projectFlowId=project_flow.id if project_flow else project_flow_id,
            projectFlowName=project_flow.name if project_flow else None,
            ownerName=str(row["owner_name"]),
            sourceType=str(row["source_type"]),
            sourceId=str(row["source_id"]) if row["source_id"] else None,
            businessCategory=business_category,
            currentBlocker=current_blocker,
            nextAction=next_action,
            recentDecision=recent_decision,
            evidenceCount=evidence_count,
            tags=[],
            note=str(note_row["note"]) if note_row else None,
            attachments=attachments,
            projectContext=project_context,
            memoryHints=memory_hints,
            backgroundReadiness=background_readiness,
            linkedFactsPreview=linked_facts_preview,
            createdAt=str(row["created_at"]),
            updatedAt=str(row["updated_at"]),
        )

    def build_cloud_event_line(payload: dict[str, object]) -> EventLineRecord:
        client_id = str(payload.get("primaryClientId")) if payload.get("primaryClientId") else None
        client_row = state.db.fetchone("SELECT name FROM clients WHERE id = ?", (client_id,)) if client_id else None
        return EventLineRecord(
            id=str(payload.get("id", "")),
            name=str(payload.get("name", "")),
            kind=str(payload.get("kind", "custom")),  # type: ignore[arg-type]
            status=str(payload.get("status", "active")),  # type: ignore[arg-type]
            businessCategory=str(payload.get("businessCategory")) if payload.get("businessCategory") else None,
            stage=str(payload.get("stage")) if payload.get("stage") else None,
            summary=str(payload.get("summary")) if payload.get("summary") else None,
            intent=str(payload.get("intent")) if payload.get("intent") else None,
            currentBlocker=str(payload.get("currentBlocker")) if payload.get("currentBlocker") else None,
            recentDecision=str(payload.get("recentDecision")) if payload.get("recentDecision") else None,
            nextStep=str(payload.get("nextStep")) if payload.get("nextStep") else None,
            evidenceCount=int(payload.get("evidenceCount") or 0),
            ownerId=str(payload.get("ownerId")) if payload.get("ownerId") else None,
            ownerName=str(payload.get("ownerName")) if payload.get("ownerName") else None,
            primaryClientId=client_id,
            primaryClientName=str(client_row["name"]) if client_row else None,
            primaryDepartmentId=str(payload.get("primaryDepartmentId")) if payload.get("primaryDepartmentId") else None,
            primaryDepartmentName=str(payload.get("primaryDepartmentName")) if payload.get("primaryDepartmentName") else None,
            participantIds=[str(item) for item in payload.get("participantIds", [])] if isinstance(payload.get("participantIds"), list) else [],
            createdAt=str(payload.get("createdAt", now_iso())),
            updatedAt=str(payload.get("updatedAt", now_iso())),
        )

    def build_cloud_event_line_activity(payload: dict[str, object]) -> EventLineActivityRecord:
        metadata = payload.get("metadata") if isinstance(payload.get("metadata"), dict) else {}
        return EventLineActivityRecord(
            id=str(payload.get("id", "")),
            eventLineId=str(payload.get("eventLineId", "")),
            sourceType=str(payload.get("sourceType", "manual_note")),  # type: ignore[arg-type]
            sourceId=str(payload.get("sourceId", "")),
            happenedAt=str(payload.get("happenedAt", now_iso())),
            actorId=str(payload.get("actorId")) if payload.get("actorId") else None,
            actorName=str(payload.get("actorName")) if payload.get("actorName") else None,
            title=str(payload.get("title", "")),
            summary=str(payload.get("summary", "")),
            metadata=metadata,
        )

    def build_event_line(row) -> EventLineRecord:
        client_id = str(row["primary_client_id"]) if row["primary_client_id"] else None
        client_row = state.db.fetchone("SELECT name FROM clients WHERE id = ?", (client_id,)) if client_id else None
        activity_count = int(state.db.scalar("SELECT COUNT(1) FROM event_line_activities WHERE event_line_id = ?", (str(row["id"]),)) or 0)
        return EventLineRecord(
            id=str(row["id"]),
            name=str(row["name"]),
            kind=str(row["kind"]),  # type: ignore[arg-type]
            status=str(row["status"]),  # type: ignore[arg-type]
            businessCategory=str(row["business_category"]) if row["business_category"] else None,
            stage=str(row["stage"]) if row["stage"] else None,
            summary=str(row["summary"]) if row["summary"] else None,
            intent=str(row["intent"]) if row["intent"] else None,
            currentBlocker=str(row["current_blocker"]) if row["current_blocker"] else None,
            recentDecision=str(row["recent_decision"]) if row["recent_decision"] else None,
            nextStep=str(row["next_step"]) if row["next_step"] else None,
            evidenceCount=max(int(row["evidence_count"] or 0), activity_count),
            ownerId=str(row["owner_id"]) if row["owner_id"] else None,
            ownerName=str(row["owner_name"]) if row["owner_name"] else None,
            primaryClientId=client_id,
            primaryClientName=str(client_row["name"]) if client_row else (str(row["primary_client_name"]) if row["primary_client_name"] else None),
            primaryDepartmentId=str(row["primary_department_id"]) if row["primary_department_id"] else None,
            primaryDepartmentName=str(row["primary_department_name"]) if row["primary_department_name"] else None,
            participantIds=[str(item) for item in from_json(row["participant_ids_json"], []) if str(item)],
            createdAt=str(row["created_at"]),
            updatedAt=str(row["updated_at"]),
        )

    def build_event_line_activity(row) -> EventLineActivityRecord:
        metadata = from_json(row["metadata_json"], {})
        return EventLineActivityRecord(
            id=str(row["id"]),
            eventLineId=str(row["event_line_id"]),
            sourceType=str(row["source_type"]),  # type: ignore[arg-type]
            sourceId=str(row["source_id"]),
            happenedAt=str(row["happened_at"]),
            actorId=str(row["actor_id"]) if row["actor_id"] else None,
            actorName=str(row["actor_name"]) if row["actor_name"] else None,
            title=str(row["title"]),
            summary=str(row["summary"]),
            metadata=metadata if isinstance(metadata, dict) else {},
        )

    def build_cloud_event_line_detail(payload: dict[str, object]) -> EventLineDetailRecord:
        event_line_payload = payload.get("eventLine")
        tasks_payload = payload.get("tasks")
        activities_payload = payload.get("activities")
        event_line = (
            build_cloud_event_line(event_line_payload)
            if isinstance(event_line_payload, dict)
            else EventLineRecord(
                id="",
                name="",
                kind="custom",
                status="active",
                participantIds=[],
                createdAt=now_iso(),
                updatedAt=now_iso(),
            )
        )
        remote_activities = [build_cloud_event_line_activity(item) for item in activities_payload if isinstance(item, dict)] if isinstance(activities_payload, list) else []
        local_attachment_rows = state.db.fetchall(
            "SELECT * FROM task_attachments_cloud WHERE event_line_id = ? ORDER BY created_at DESC",
            (event_line.id,),
        ) if event_line.id else []
        attachment_activities = [
            build_attachment_event_line_activity(build_task_attachment(row))
            for row in local_attachment_rows
        ]
        combined_activities = remote_activities + attachment_activities
        combined_activities.sort(key=lambda item: (item.happenedAt, item.id), reverse=True)
        return EventLineDetailRecord(
            eventLine=event_line,
            tasks=[build_cloud_task(item, {}) for item in tasks_payload if isinstance(item, dict)] if isinstance(tasks_payload, list) else [],
            activities=combined_activities,
        )

    def build_event_line_detail(row) -> EventLineDetailRecord:
        tasks = fetch_tasks("t.event_line_id = ?", (str(row["id"]),))
        activity_rows = state.db.fetchall(
            """
            SELECT *
            FROM event_line_activities
            WHERE event_line_id = ?
            ORDER BY happened_at DESC, created_at DESC
            """,
            (str(row["id"]),),
        )
        memory_response = get_event_line_memory_response(state.db, str(row["id"]))
        return EventLineDetailRecord(
            eventLine=build_event_line(row),
            tasks=tasks,
            activities=[build_event_line_activity(item) for item in activity_rows],
            memorySnapshot=memory_response.eventLineMemorySnapshot,
            predictionReadiness=memory_response.eventLineMemorySnapshot.predictionReadiness if memory_response.eventLineMemorySnapshot else None,
            clarificationNeeds=memory_response.clarificationNeeds,
        )

    def fetch_tasks(where_clause: str = "", params: tuple = ()) -> list[TaskRecord]:
        query = """
            SELECT t.*, l.name AS list_name, l.color AS list_color, c.name AS client_name
            FROM tasks t
            JOIN task_lists l ON l.id = t.list_id
            LEFT JOIN clients c ON c.id = t.client_id
        """
        if where_clause:
            query += f" WHERE {where_clause}"
        query += " ORDER BY CASE t.status WHEN 'inbox' THEN 0 WHEN 'doing' THEN 1 WHEN 'todo' THEN 2 WHEN 'done' THEN 3 ELSE 4 END, t.updated_at DESC"
        return [build_task(row) for row in state.db.fetchall(query, params)]

    def build_growth_workbench_snapshot(week_label: str | None = None) -> GrowthWorkbenchSnapshotRecord:
        phase_blueprints = [
            ("p1", "需求接收", "明确需求来源、目标对象和优先级", ["需求来源模糊", "优先级未经确认"]),
            ("p2", "信息核对", "确认关键事实、材料和依赖项都已到位", ["输入材料不完整", "事实口径未统一"]),
            ("p3", "内部对齐", "明确会议目标、参会人及预期结论", ["未提前拉齐信息", "会议目标发散"]),
            ("p4", "方案产出", "形成结构清晰、可执行的初版方案", ["结构与受众不匹配", "缺少支撑数据"]),
            ("p5", "沟通推进", "把边界、责任人和时间线谈清楚", ["临场判断不足", "关键利益方未提前对齐"]),
            ("p6", "交付闭环", "形成明确交付物、待办与复核节点", ["只做了动作，没有闭环", "责任人和时间点不明确"]),
            ("p7", "复盘沉淀", "把本次有效做法转成可复用经验", ["只记录结果，没有方法", "经验无法迁移复用"]),
        ]
        task_kind_blueprints = [
            {
                "taskKind": "agreement_alignment",
                "keywords": ("协议", "合同", "条款", "说明迭代", "合作说明", "合作协议", "修订"),
                "riskTypes": ["boundary_risk", "commitment_risk", "negotiation_risk"],
                "requiredAbilities": ["collab", "risk", "write", "insight"],
                "defaultGoal": "把合作边界、关键争议点和本次要确认的结论谈清楚",
                "defaultDeliverable": "一版协议差异、待确认点和下一轮修改动作",
                "whyRelevant": "这类任务不是单纯沟通，而是边界与承诺对齐，稍早拍板就会留下后续风险。",
                "cards": [
                    {
                        "cardType": "动作卡",
                        "title": "沟通前先列本次必须确认的 3 个点",
                        "summary": "先把本次沟通一定要拿到的结论写清楚，再决定怎么开口。",
                        "checklist": ["本次必须确认的条款或边界", "哪些问题你能现场确认", "哪些问题需要带回内部"],
                        "talkTrack": ["这次我希望先把三件事对齐，避免双方理解继续漂移。"],
                        "templateHint": "协议沟通前置清单",
                        "expectedOutput": "本次沟通的核心议题与确认边界",
                    },
                    {
                        "cardType": "检查卡",
                        "title": "协议沟通前先排查承诺风险",
                        "summary": "把不能现场承诺的内容提前划出来，避免沟通时话说满。",
                        "checklist": ["哪些条款涉及资源/交付承诺", "哪些点需要负责人或法务兜底", "哪些说法只能表达方向不能表态"],
                        "talkTrack": [],
                        "templateHint": "风险排查清单",
                        "expectedOutput": "不能现场确认的条款清单",
                    },
                    {
                        "cardType": "话术卡",
                        "title": "先问真实顾虑，再谈条款表述",
                        "summary": "如果先急着改字句，容易错过对方真正卡住的顾虑。",
                        "checklist": [],
                        "talkTrack": ["为了避免我们只改表述不改问题，我想先确认您最担心的是哪一类合作风险。"],
                        "templateHint": "",
                        "expectedOutput": "对方真实顾虑与协商空间",
                    },
                    {
                        "cardType": "模板卡",
                        "title": "沟通后立刻沉淀版本差异与待确认项",
                        "summary": "会后不只记结论，要沉淀版本变化、待确认项和责任人。",
                        "checklist": ["版本差异", "待确认项", "责任人和时间点"],
                        "talkTrack": [],
                        "templateHint": "协议迭代纪要模板",
                        "expectedOutput": "带责任人的版本差异纪要",
                    },
                ],
            },
            {
                "taskKind": "external_communication",
                "keywords": ("沟通", "联系", "对接", "访谈", "拜访", "电话", "老师", "客户", "约访"),
                "riskTypes": ["boundary_risk", "fact_gap", "negotiation_risk"],
                "requiredAbilities": ["collab", "insight", "risk"],
                "defaultGoal": "确认对方真实诉求、边界和下一步推进条件",
                "defaultDeliverable": "一次带结论的沟通纪要和下一步动作",
                "whyRelevant": "外部沟通的关键不是把信息说完，而是拿到真实顾虑与下一步承诺。",
                "cards": [
                    {
                        "cardType": "动作卡",
                        "title": "沟通前先定目标、对象和预期结论",
                        "summary": "先回答这次为什么沟通、找谁沟通、沟通完要留下什么。",
                        "checklist": ["核心目标", "对方角色与立场", "预期结论"],
                        "talkTrack": [],
                        "templateHint": "外部沟通准备卡",
                        "expectedOutput": "明确的沟通目标和预期结论",
                    },
                    {
                        "cardType": "检查卡",
                        "title": "先补项目背景，再进入沟通",
                        "summary": "没有项目背景时，沟通容易停留在表面信息交换。",
                        "checklist": ["当前项目阶段", "最近一次相关沟通结论", "本次沟通与整体项目的关系"],
                        "talkTrack": [],
                        "templateHint": "",
                        "expectedOutput": "足够支撑沟通判断的背景包",
                    },
                    {
                        "cardType": "话术卡",
                        "title": "先确认对方最关注什么，再给方案",
                        "summary": "先问对方担心点，比上来先讲方案更容易收口。",
                        "checklist": [],
                        "talkTrack": ["为了确保这次沟通不跑偏，我想先确认一下您目前最关注的是什么。"],
                        "templateHint": "",
                        "expectedOutput": "对方最关注的问题清单",
                    },
                ],
            },
            {
                "taskKind": "cross_team_coordination",
                "keywords": ("跨部门", "协调", "资源", "协同", "对齐", "推动", "联动"),
                "riskTypes": ["boundary_risk", "fact_gap"],
                "requiredAbilities": ["collab", "exec", "risk"],
                "defaultGoal": "把协作边界、责任人和时间线收清楚",
                "defaultDeliverable": "一组已确认的协作动作和责任归属",
                "whyRelevant": "跨团队事项最容易卡在边界模糊和责任漂移。",
                "cards": [
                    {
                        "cardType": "动作卡",
                        "title": "先写清协作边界和第一责任人",
                        "summary": "没有边界和第一责任人，协作推进只会停在口头共识。",
                        "checklist": ["交付物是什么", "谁先动", "最晚时间点"],
                        "talkTrack": [],
                        "templateHint": "协作边界清单",
                        "expectedOutput": "带责任人和时间点的协作边界",
                    },
                    {
                        "cardType": "话术卡",
                        "title": "对齐资源时先谈约束，不要直接要结果",
                        "summary": "先把对方当前约束讲清楚，后面才知道怎么交换优先级。",
                        "checklist": [],
                        "talkTrack": ["为了让这件事有落地可能，我想先了解你们当前最大的排期约束是什么。"],
                        "templateHint": "",
                        "expectedOutput": "协作约束和可谈空间",
                    },
                ],
            },
            {
                "taskKind": "meeting_preparation",
                "keywords": ("会议", "议程", "纪要", "评审", "复盘会", "对齐会"),
                "riskTypes": ["fact_gap", "boundary_risk"],
                "requiredAbilities": ["collab", "write", "exec"],
                "defaultGoal": "让会议开始前就知道结论、边界和会后动作如何落地",
                "defaultDeliverable": "会议议程、参会人、预期结论和会后待办结构",
                "whyRelevant": "会前准备做得差，会议会退化成信息交换。",
                "cards": [
                    {
                        "cardType": "动作卡",
                        "title": "会前先锁定议题、参会人和预期结论",
                        "summary": "这三件事不清楚，会议就很难产出有效结论。",
                        "checklist": ["会议目标", "关键参会人", "预期结论"],
                        "talkTrack": [],
                        "templateHint": "会议准备模板",
                        "expectedOutput": "可执行的会议准备单",
                    },
                    {
                        "cardType": "模板卡",
                        "title": "会后直接转责任到人",
                        "summary": "纪要不只记结论，要能直接落到任务和负责人。",
                        "checklist": ["待办", "责任人", "截止时间"],
                        "talkTrack": [],
                        "templateHint": "会议纪要转任务模板",
                        "expectedOutput": "会后行动项清单",
                    },
                ],
            },
            {
                "taskKind": "proposal_output",
                "keywords": ("方案", "白皮书", "提案", "大纲", "汇报", "说明书", "材料"),
                "riskTypes": ["fact_gap", "commitment_risk"],
                "requiredAbilities": ["write", "analyze", "insight"],
                "defaultGoal": "形成结构清楚、面向对象、可被继续推进的输出物",
                "defaultDeliverable": "一个可继续编辑或评审的结构化版本",
                "whyRelevant": "方案类任务最怕只写内容，不先想受众、结论和支撑依据。",
                "cards": [
                    {
                        "cardType": "动作卡",
                        "title": "先定受众、目的和目录结构",
                        "summary": "结构先错了，后面只会越写越重。",
                        "checklist": ["面向谁", "想推进什么", "目录骨架"],
                        "talkTrack": [],
                        "templateHint": "方案大纲模板",
                        "expectedOutput": "清晰的目录和表达主线",
                    },
                    {
                        "cardType": "检查卡",
                        "title": "每一页都要有事实和判断的对应关系",
                        "summary": "没有支撑依据的判断，后续很难被采纳。",
                        "checklist": ["关键事实", "判断结论", "下一步动作"],
                        "talkTrack": [],
                        "templateHint": "",
                        "expectedOutput": "事实-判断-动作链条",
                    },
                ],
            },
            {
                "taskKind": "review_and_closure",
                "keywords": ("复盘", "验收", "闭环", "总结", "回顾", "沉淀"),
                "riskTypes": ["fact_gap"],
                "requiredAbilities": ["write", "analyze", "risk"],
                "defaultGoal": "把结果、原因、方法和下次动作讲清楚",
                "defaultDeliverable": "一条可复用经验或复盘结论",
                "whyRelevant": "复盘的价值不在记录结果，而在把方法和误区说清楚。",
                "cards": [
                    {
                        "cardType": "动作卡",
                        "title": "结果后面一定要补原因和改法",
                        "summary": "只有结果没有原因，这次经验很难迁移。",
                        "checklist": ["发生了什么", "为什么会这样", "下次如何更好"],
                        "talkTrack": [],
                        "templateHint": "复盘四段式模板",
                        "expectedOutput": "可复用的复盘记录",
                    },
                    {
                        "cardType": "模板卡",
                        "title": "把有效做法沉淀成经验卡",
                        "summary": "把一次有效动作沉淀出来，后面才能在相似项目里复用。",
                        "checklist": ["适用场景", "方法", "边界", "下一次提醒"],
                        "talkTrack": [],
                        "templateHint": "经验沉淀模板",
                        "expectedOutput": "一条结构完整的经验资产",
                    },
                ],
            },
        ]
        client_workspace_cache: dict[str, ClientWorkspaceResponse | None] = {}
        strategic_snapshot_cache: dict[str, StrategicCockpitSnapshotRecord | None] = {}

        def normalize_text(value: str | None) -> str:
            return (value or "").strip()

        def parse_task_date(value: str | None):
            if not value:
                return None
            candidate = f"{value}T00:00:00" if len(value) <= 10 else value
            try:
                return datetime.fromisoformat(candidate)
            except ValueError:
                return None

        def sort_updated_at(value: str | None):
            if not value:
                return datetime(1970, 1, 1)
            try:
                normalized = value.replace("Z", "+00:00")
                parsed = datetime.fromisoformat(normalized)
                if parsed.tzinfo is not None:
                    return parsed.replace(tzinfo=None)
                return parsed
            except ValueError:
                return datetime(1970, 1, 1)

        def format_deadline(task: TaskRecord) -> str:
            raw = task.dueDate or task.ddl
            if not raw:
                return "待补日期"
            date = parse_task_date(raw)
            if not date:
                return raw
            today = datetime.now()
            target = date.replace(hour=0, minute=0, second=0, microsecond=0)
            base = today.replace(hour=0, minute=0, second=0, microsecond=0)
            diff_days = round((target - base).total_seconds() / 86400)
            if diff_days < 0:
                return f"已超期 {abs(diff_days)} 天"
            if diff_days == 0:
                return "今天"
            if diff_days == 1:
                return "明天"
            if diff_days <= 7:
                return f"{diff_days} 天后"
            return f"{date.month}月{date.day}日"

        def infer_phase(task: TaskRecord) -> str:
            haystack = " ".join(
                part
                for part in (
                    task.title,
                    task.desc,
                    task.note or "",
                    task.orgContext.blockedAtStep if task.orgContext else "",
                    task.projectContext.projectFlowName if task.projectContext else "",
                )
                if part
            )
            if any(keyword in haystack for keyword in ("需求", "接收", "收件")) or task.status == "inbox":
                return "需求接收"
            if any(keyword in haystack for keyword in ("信息", "资料", "材料", "核对", "澄清")):
                return "信息核对"
            if any(keyword in haystack for keyword in ("对齐", "会议", "纪要", "评审")):
                return "内部对齐"
            if any(keyword in haystack for keyword in ("方案", "白皮书", "提案", "文档", "大纲", "写作", "输出")):
                return "方案产出"
            if any(keyword in haystack for keyword in ("沟通", "协调", "协作", "推进", "谈判", "资源")):
                return "沟通推进"
            if any(keyword in haystack for keyword in ("交付", "验收", "上线", "发布", "闭环")):
                return "交付闭环"
            if task.status == "done":
                return "复盘沉淀"
            if task.status == "doing":
                return "沟通推进" if (task.orgContext.isCrossDepartment if task.orgContext else False) else "交付闭环"
            return "内部对齐" if (task.orgContext.isCrossDepartment if task.orgContext else False) or task.collaborators else "信息核对"

        def urgency_meta(task: TaskRecord) -> tuple[str, str]:
            due_date = parse_task_date(task.dueDate or task.ddl)
            today = datetime.now().replace(hour=0, minute=0, second=0, microsecond=0)
            diff_days = round((due_date.replace(hour=0, minute=0, second=0, microsecond=0) - today).total_seconds() / 86400) if due_date else None
            if diff_days is not None and diff_days < 0:
                return "建议优先处理", "text-red-700 bg-red-100"
            if task.priority == "high" or (diff_days is not None and diff_days <= 2):
                return "建议优先处理", "text-red-700 bg-red-100"
            if task.viewerInboxStatus == "pending" or (task.orgContext.needsReview if task.orgContext else False) or (task.orgContext.blockedAtStep if task.orgContext else False):
                return "需先补关键动作", "text-orange-700 bg-orange-100"
            return "可直接推进", "text-green-700 bg-green-100"

        def risks_for_task(task: TaskRecord, phase: str) -> list[str]:
            risks: list[str] = []
            if not normalize_text(task.desc) and not normalize_text(task.note):
                risks.append("任务背景信息偏少，开始前建议先补齐目标、上下文和预期输出。")
            if not task.dueDate and not task.ddl:
                risks.append("截止时间尚未明确，推进节奏容易在中途松掉。")
            if (task.orgContext.isCrossDepartment if task.orgContext else False) or task.collaborators:
                risks.append("涉及多人或跨部门协作，如果不先对齐边界和责任人，后续容易返工。")
            if task.viewerInboxStatus == "pending" or (task.collaborationSummary.get("pending", 0) > 0):
                risks.append("仍有协作者未完成接收确认，关键动作可能停在等待。")
            if task.orgContext.needsReview if task.orgContext else False:
                risks.append("当前任务仍需要复核或审批，建议先补齐说明与证据。")
            if risks:
                return risks[:2]
            defaults = {
                "需求接收": "需求来源和目标对象还未完全确认，过早执行容易方向跑偏。",
                "信息核对": "关键信息口径若未先统一，后续材料和决策会反复返工。",
                "内部对齐": "参会人、边界和预期结论不清楚时，会议很容易变成信息交换。",
                "方案产出": "结构与受众若不匹配，方案会花很多时间在重写上。",
                "沟通推进": "关键利益方未提前识别时，推进节点最容易卡在协作博弈上。",
                "交付闭环": "只推进动作不收责任人和时间点，容易在最后一步失去闭环。",
                "复盘沉淀": "如果只记录结果不提炼方法，这次经验很难转成下次可复用资产。",
            }
            return [defaults.get(phase, "先补齐关键动作，再继续推进。")]

        def robot_assessment(task: TaskRecord, phase: str) -> tuple[bool, list[str]]:
            context_signals = len(
                [
                    item
                    for item in (
                        normalize_text(task.desc),
                        normalize_text(task.note),
                        "tags" if task.tags else "",
                        task.dueDate or task.ddl or "",
                    )
                    if item
                ]
            )
            haystack = f"{task.title}{task.desc}{task.note or ''}"
            standardizable = any(keyword in haystack for keyword in ("会议", "纪要", "清单", "模板", "方案", "提纲", "白皮书", "复盘", "风险", "对齐", "材料", "SOP", "文档"))
            human_heavy = (task.orgContext.isCrossDepartment if task.orgContext else False) or any(keyword in haystack for keyword in ("协调", "沟通", "谈判", "客户", "资源", "博弈", "冲突"))
            ready = context_signals >= 2 and standardizable and not human_heavy and task.status != "inbox"
            if ready:
                return True, ["任务上下文已补齐到可生成首稿", f"当前处在 {phase} 阶段，标准输出较明确", "可先由机器人生成准备清单或文档草稿"]
            reasons: list[str] = []
            if context_signals < 2:
                reasons.append("任务描述、备注或截止信息仍不够完整")
            if human_heavy:
                reasons.append("当前阶段强依赖跨部门或现场判断，暂不适合全自动执行")
            if not standardizable:
                reasons.append("任务输出结构还不够标准化，机器人难以稳定接手")
            return False, reasons[:3] or ["当前任务仍需要人先定调，再适合让机器人协助执行"]

        def next_advice(task: TaskRecord, phase: str) -> str:
            task_name = f"「{task.title}」"
            mapping = {
                "需求接收": f"先为{task_name}确认目标对象、优先级和成功标准，再进入执行。",
                "信息核对": f"先补齐{task_name}所需的材料、数据和关键口径，再进入下一步。",
                "内部对齐": f"建议先把{task_name}的参会人、边界和预期结论写清楚，再开始拉会或对齐。",
                "方案产出": f"已具备开始条件，建议先为{task_name}拉出结构化大纲，再补细节。",
                "沟通推进": f"不要直接硬推，先把{task_name}的责任人、协作边界和时间线谈清楚。",
                "交付闭环": f"把{task_name}的交付物、待办和复核节点一起收拢，避免最后一步失焦。",
                "复盘沉淀": f"完成{task_name}后，尽快把有效做法沉淀成一条可复用经验。",
            }
            return task.nextAction or (task.projectContext.nextAction if task.projectContext else None) or mapping.get(phase, f"先补齐{task_name}的关键动作，再继续推进。")

        def task_contexts(task: TaskRecord) -> list[GrowthContextLinkRecord]:
            contexts = [
                GrowthContextLinkRecord(
                    objectType="task",
                    objectId=task.id,
                    label=task.title,
                    subtitle=(task.projectContext.stage if task.projectContext else None) or task.eventLineName or task.clientName or task.listName,
                    tab="tasks",
                    statusLabel=task.status,
                )
            ]
            if task.eventLineId and task.eventLineName:
                contexts.append(
                    GrowthContextLinkRecord(
                        objectType="event_line",
                        objectId=task.eventLineId,
                        label=task.eventLineName,
                        subtitle=task.businessCategory or (task.projectContext.stage if task.projectContext else "") or "事件线",
                        tab="tasks",
                        statusLabel="事件线",
                    )
                )
            if task.clientId and task.clientName:
                contexts.append(
                    GrowthContextLinkRecord(
                        objectType="client",
                        objectId=task.clientId,
                        label=task.clientName,
                        subtitle=(task.projectContext.stage if task.projectContext else "") or task.businessCategory or "项目工作台",
                        tab="client_workspace",
                        statusLabel="客户项目",
                    )
                )
            project_module_id = (task.projectContext.projectModuleId if task.projectContext else None) or task.projectModuleId
            project_module_name = (task.projectContext.projectModuleName if task.projectContext else None) or task.projectModuleName
            if project_module_id and project_module_name:
                contexts.append(
                    GrowthContextLinkRecord(
                        objectType="project_module",
                        objectId=project_module_id,
                        label=project_module_name,
                        subtitle=task.clientName or task.eventLineName or "项目模块",
                        tab="tasks",
                        statusLabel="项目模块",
                    )
                )
            project_flow_id = (task.projectContext.projectFlowId if task.projectContext else None) or task.projectFlowId
            project_flow_name = (task.projectContext.projectFlowName if task.projectContext else None) or task.projectFlowName
            if project_flow_id and project_flow_name:
                contexts.append(
                    GrowthContextLinkRecord(
                        objectType="project_flow",
                        objectId=project_flow_id,
                        label=project_flow_name,
                        subtitle=(task.projectContext.stage if task.projectContext else "") or task.businessCategory or "流程节点",
                        tab="tasks",
                        statusLabel="项目流程",
                    )
                )
            strategic_snapshot = strategic_snapshot_for_client_cached(task.clientId or (task.projectContext.clientId if task.projectContext else None))
            if strategic_snapshot and strategic_snapshot.strategicLines:
                line = strategic_snapshot.strategicLines[0]
                contexts.append(
                    GrowthContextLinkRecord(
                        objectType="strategic_focus",
                        objectId=line.id,
                        label=line.title,
                        subtitle=line.stage or "战略呼应",
                        tab="strategic_accompaniment",
                        statusLabel="战略线",
                    )
                )
            return contexts

        def workspace_for_client_cached(client_id: str | None) -> ClientWorkspaceResponse | None:
            normalized = normalize_text(client_id)
            if not normalized:
                return None
            if normalized not in client_workspace_cache:
                try:
                    client_workspace_cache[normalized] = workspace_for_client(normalized)
                except HTTPException:
                    client_workspace_cache[normalized] = None
            return client_workspace_cache[normalized]

        def strategic_snapshot_for_client_cached(client_id: str | None) -> StrategicCockpitSnapshotRecord | None:
            normalized = normalize_text(client_id)
            if not normalized:
                return None
            if normalized not in strategic_snapshot_cache:
                try:
                    strategic_snapshot_cache[normalized] = build_strategic_cockpit_snapshot(normalized)
                except HTTPException:
                    strategic_snapshot_cache[normalized] = None
            return strategic_snapshot_cache[normalized]

        def primary_task_context(task: TaskRecord, contexts: list[GrowthContextLinkRecord]) -> GrowthContextLinkRecord | None:
            return (
                next((context for context in contexts if context.objectType == "task"), None)
                or next((context for context in contexts if context.objectType == "event_line"), None)
                or next((context for context in contexts if context.objectType == "client"), None)
                or (contexts[0] if contexts else None)
            )

        def infer_task_intent(task: TaskRecord, phase: str) -> GrowthTaskIntentRecord:
            haystack = normalize_text(
                " ".join(
                    part
                    for part in (
                        task.title,
                        task.desc,
                        task.note or "",
                        task.projectContext.projectFlowName if task.projectContext else "",
                        task.projectContext.projectModuleName if task.projectContext else "",
                    )
                    if part
                )
            )
            matched_blueprint = next(
                (
                    blueprint
                    for blueprint in task_kind_blueprints
                    if any(keyword in haystack for keyword in blueprint["keywords"])
                ),
                None,
            )
            if matched_blueprint is None:
                matched_blueprint = {
                    "taskKind": "general_execution",
                    "riskTypes": ["fact_gap"],
                    "requiredAbilities": ["exec", "collab"],
                    "defaultGoal": "把当前任务推进到下一个明确节点",
                    "defaultDeliverable": "一条带责任人和时间点的下一步动作",
                    "whyRelevant": f"当前任务处在「{phase}」阶段，系统会先给出最小可执行动作，再逐步补齐背景。",
                    "cards": [],
                }
            goal = (
                normalize_text(task.projectContext.goalSummary if task.projectContext else None)
                or normalize_text(task.nextAction)
                or normalize_text(task.recentDecision)
                or str(matched_blueprint["defaultGoal"])
            )
            deliverable = (
                normalize_text(task.projectContext.projectFlowSummary if task.projectContext else None)
                or normalize_text(task.projectContext.projectModuleSummary if task.projectContext else None)
                or normalize_text(task.nextAction)
                or str(matched_blueprint["defaultDeliverable"])
            )
            risk_types = list(matched_blueprint["riskTypes"])
            if not normalize_text(task.desc) and not normalize_text(task.note):
                risk_types.append("fact_gap")
            if (task.orgContext.isCrossDepartment if task.orgContext else False) or task.collaborators:
                risk_types.append("boundary_risk")
            if task.viewerInboxStatus == "pending":
                risk_types.append("coordination_gap")
            return GrowthTaskIntentRecord(
                taskKind=str(matched_blueprint["taskKind"]),
                goal=goal,
                deliverable=deliverable,
                riskTypes=list(dict.fromkeys(str(item) for item in risk_types if str(item).strip())),
                requiredAbilities=[str(item) for item in matched_blueprint["requiredAbilities"]],  # type: ignore[list-item]
                confidence=0.84 if matched_blueprint["taskKind"] != "general_execution" else 0.56,
                whyRelevant=str(matched_blueprint["whyRelevant"]),
            )

        def build_project_context_pack(task: TaskRecord, contexts: list[GrowthContextLinkRecord]) -> GrowthProjectContextPackRecord:
            workspace = workspace_for_client_cached(task.clientId or (task.projectContext.clientId if task.projectContext else None))
            strategic_snapshot = strategic_snapshot_for_client_cached(task.clientId or (task.projectContext.clientId if task.projectContext else None))
            task_notes = [
                item
                for item in (
                    normalize_text(task.desc),
                    normalize_text(task.note),
                    normalize_text(task.projectContext.backgroundSummary if task.projectContext else None),
                    normalize_text(task.projectContext.goalSummary if task.projectContext else None),
                    normalize_text(task.recentDecision),
                    normalize_text(task.nextAction),
                )
                if item
            ]
            linked_facts = [normalize_text(item.factValue) for item in task.linkedFactsPreview if normalize_text(item.factValue)]
            recent_meetings = [
                item
                for item in (
                    [
                        " · ".join(
                            part
                            for part in (
                                normalize_text(meeting.title),
                                normalize_text(meeting.stageLabel if hasattr(meeting, "stageLabel") else None),
                                normalize_text(meeting.updatedAt[:10] if getattr(meeting, "updatedAt", None) else None),
                            )
                            if part
                        )
                        for meeting in (workspace.meetings[:3] if workspace else [])
                    ]
                )
                if item
            ]
            strategic_focus: list[str] = []
            if strategic_snapshot:
                strategic_focus.extend(str(item).strip() for item in strategic_snapshot.headline.focusItems[:3] if str(item).strip())
                strategic_focus.extend(
                    item.title
                    for item in strategic_snapshot.strategicLines[:2]
                    if normalize_text(item.title)
                )
            context_gaps = [
                item
                for item in (
                    "缺任务背景说明" if not task_notes else "",
                    "缺历史会议信息" if not recent_meetings else "",
                    "缺附件或事实依据" if not task.attachments and not linked_facts else "",
                    "缺战略焦点" if not strategic_focus and task.clientId else "",
                )
                if item
            ]
            event_line_summary = "；".join(
                item
                for item in (
                    normalize_text(task.eventLineName),
                    normalize_text(task.projectContext.currentFocus if task.projectContext else None),
                    normalize_text(task.projectContext.currentBlocker if task.projectContext else None),
                )
                if item
            )
            return GrowthProjectContextPackRecord(
                title=(task.clientName or (task.projectContext.clientName if task.projectContext else None) or task.eventLineName or task.title),
                taskNotes=task_notes[:4],
                attachments=[normalize_text(item.title) for item in task.attachments[:4] if normalize_text(item.title)],
                memoryHints=[normalize_text(item) for item in task.memoryHints[:4] if normalize_text(item)],
                linkedFacts=linked_facts[:4],
                clientSummary=normalize_text(workspace.client.intro if workspace else None) or normalize_text(task.projectContext.backgroundSummary if task.projectContext else None),
                recentMeetings=recent_meetings[:3],
                eventLineSummary=event_line_summary,
                strategicFocus=list(dict.fromkeys(str(item) for item in strategic_focus if str(item).strip()))[:3],
                keyWarnings=list(dict.fromkeys(task.projectContext.riskSummary.split("；") if task.projectContext and normalize_text(task.projectContext.riskSummary) else []))[:3],
                contextGaps=context_gaps[:4],
            )

        def build_universal_skills(
            task: TaskRecord,
            *,
            task_intent: GrowthTaskIntentRecord,
            primary_context: GrowthContextLinkRecord | None,
            project_context_pack: GrowthProjectContextPackRecord,
        ) -> list[GrowthUniversalSkillItemRecord]:
            blueprint = next((item for item in task_kind_blueprints if item["taskKind"] == task_intent.taskKind), None)
            cards = list(blueprint["cards"]) if blueprint else []
            skill_items: list[GrowthUniversalSkillItemRecord] = []
            for index, card in enumerate(cards[:4]):
                skill_items.append(
                    GrowthUniversalSkillItemRecord(
                        id=f"{task.id}-skill-{index + 1}",
                        cardType=str(card["cardType"]),  # type: ignore[arg-type]
                        title=str(card["title"]),
                        summary=str(card["summary"]),
                        whyRelevant=task_intent.whyRelevant,
                        checklist=[str(item) for item in card.get("checklist", []) if str(item).strip()],
                        talkTrack=[str(item) for item in card.get("talkTrack", []) if str(item).strip()],
                        templateHint=str(card.get("templateHint") or ""),
                        sourceKind="rule",
                        expectedOutput=str(card.get("expectedOutput") or ""),
                        linkedContext=primary_context,
                    )
                )
            if len(skill_items) < 2 and project_context_pack.contextGaps:
                skill_items.append(
                    GrowthUniversalSkillItemRecord(
                        id=f"{task.id}-skill-ai-gap",
                        cardType="检查卡",
                        title="先补齐当前任务最缺的背景再继续推进",
                        summary="系统已识别到上下文缺口，先把缺的背景或事实补齐，推荐质量才会稳定。",
                        whyRelevant="当前任务背景不足时，任何泛化建议都会变空。",
                        checklist=project_context_pack.contextGaps[:3],
                        talkTrack=[],
                        templateHint="任务背景补齐清单",
                        sourceKind="ai_supplement",
                        expectedOutput="最小可执行的项目背景包",
                        linkedContext=primary_context,
                    )
                )
            return skill_items[:4]

        def build_material_refs(
            task: TaskRecord,
            *,
            contexts: list[GrowthContextLinkRecord],
            project_context_pack: GrowthProjectContextPackRecord,
        ) -> list[GrowthMaterialRefRecord]:
            refs: list[GrowthMaterialRefRecord] = []
            task_context = next((context for context in contexts if context.objectType == "task"), None)
            client_context = next((context for context in contexts if context.objectType == "client"), None)
            event_line_context = next((context for context in contexts if context.objectType == "event_line"), None)
            strategic_context = next((context for context in contexts if context.objectType == "strategic_focus"), None)
            for attachment in task.attachments[:3]:
                refs.append(
                    GrowthMaterialRefRecord(
                        id=f"attachment-{attachment.id}",
                        title=attachment.title,
                        summary="任务附件中已存在的资料，可直接进入本次动作准备。",
                        sourceKind="task_material",
                        linkedContext=task_context,
                    )
                )
            for index, meeting in enumerate(project_context_pack.recentMeetings[:2]):
                refs.append(
                    GrowthMaterialRefRecord(
                        id=f"meeting-{task.id}-{index}",
                        title=meeting,
                        summary="最近一次相关会议，可先读结论和争议点再进入本次动作。",
                        sourceKind="client_workspace",
                        linkedContext=client_context,
                    )
                )
            if project_context_pack.eventLineSummary:
                refs.append(
                    GrowthMaterialRefRecord(
                        id=f"event-line-{task.id}",
                        title=task.eventLineName or "当前事件线",
                        summary=project_context_pack.eventLineSummary,
                        sourceKind="event_line",
                        linkedContext=event_line_context,
                    )
                )
            for index, focus in enumerate(project_context_pack.strategicFocus[:2]):
                refs.append(
                    GrowthMaterialRefRecord(
                        id=f"strategic-focus-{task.id}-{index}",
                        title=focus,
                        summary="这条任务和当前战略焦点直接相关，沟通或输出时要对齐这一层目标。",
                        sourceKind="strategic_focus",
                        linkedContext=strategic_context or client_context,
                    )
                )
            if project_context_pack.clientSummary:
                refs.append(
                    GrowthMaterialRefRecord(
                        id=f"client-summary-{task.id}",
                        title=(task.clientName or "当前客户") + "背景摘要",
                        summary=project_context_pack.clientSummary,
                        sourceKind="project_context",
                        linkedContext=client_context,
                    )
                )
            return refs[:6]

        def build_action_plan(
            task: TaskRecord,
            *,
            task_intent: GrowthTaskIntentRecord,
            phase: str,
            primary_context: GrowthContextLinkRecord | None,
            project_context_pack: GrowthProjectContextPackRecord,
        ) -> list[GrowthActionPlanItemRecord]:
            return [
                GrowthActionPlanItemRecord(
                    id=f"{task.id}-plan-before-1",
                    phaseGroup="before",
                    title="开始前先确认本次要拿到的结论",
                    purpose="避免把沟通或推进做成无产出的信息交换。",
                    expectedOutput=task_intent.goal or "本次任务的目标与结论清单",
                    ifMissing="没有目标，沟通和输出都会发散，后面很难判断这次任务是否成功。",
                    actionLabel="回到当前任务补目标",
                    sourceKind="rule",
                    linkedContext=primary_context,
                ),
                GrowthActionPlanItemRecord(
                    id=f"{task.id}-plan-before-2",
                    phaseGroup="before",
                    title="先补项目背景和最近事实",
                    purpose="让当前动作建立在真实项目材料和历史结论上。",
                    expectedOutput="一份可执行的背景包：任务说明、最近会议、附件、关键事实",
                    ifMissing="背景不足时，建议会变空，沟通也容易停留在表面。",
                    actionLabel="查看项目背景包",
                    sourceKind="project_context",
                    linkedContext=primary_context,
                ),
                GrowthActionPlanItemRecord(
                    id=f"{task.id}-plan-during-1",
                    phaseGroup="during",
                    title=f"执行中围绕「{task_intent.taskKind}」收口关键问题",
                    purpose="把对方顾虑、边界和下一步动作在现场讲清楚。",
                    expectedOutput=task_intent.deliverable or "一组可执行的结论与待办",
                    ifMissing="执行时只讲信息不收结论，后续就只能靠会后补猜测。",
                    actionLabel="打开沟通/执行清单",
                    sourceKind="rule",
                    linkedContext=primary_context,
                ),
                GrowthActionPlanItemRecord(
                    id=f"{task.id}-plan-after-1",
                    phaseGroup="after",
                    title="完成后立即沉淀版本差异、待确认项和责任人",
                    purpose="把这次动作转成后续任务、会议和成长沉淀的证据。",
                    expectedOutput="带责任人的纪要、待办或经验记录",
                    ifMissing="如果会后不沉淀，这次任务里的有效判断很难进入后续成长账本。",
                    actionLabel="沉淀为经验",
                    sourceKind="project_context" if project_context_pack.attachments or project_context_pack.recentMeetings else "rule",
                    linkedContext=primary_context,
                ),
            ]

        def workbench_task_from_task(task: TaskRecord) -> GrowthWorkbenchTaskRecord:
            phase = infer_phase(task)
            urgency, urgency_color = urgency_meta(task)
            robot_ready, robot_reasons = robot_assessment(task, phase)
            contexts = task_contexts(task)
            primary_context = primary_task_context(task, contexts)
            task_intent = infer_task_intent(task, phase)
            project_context_pack = build_project_context_pack(task, contexts)
            universal_skills = build_universal_skills(
                task,
                task_intent=task_intent,
                primary_context=primary_context,
                project_context_pack=project_context_pack,
            )
            material_refs = build_material_refs(
                task,
                contexts=contexts,
                project_context_pack=project_context_pack,
            )
            action_plan = build_action_plan(
                task,
                task_intent=task_intent,
                phase=phase,
                primary_context=primary_context,
                project_context_pack=project_context_pack,
            )
            return GrowthWorkbenchTaskRecord(
                id=task.id,
                title=task.title,
                project=(task.projectContext.projectFlowName if task.projectContext else None) or (task.projectContext.projectModuleName if task.projectContext else None) or task.eventLineName or (task.projectContext.clientName if task.projectContext else None) or task.clientName or task.listName or task.ownerName or "任务执行",
                clientName=(task.projectContext.clientName if task.projectContext else None) or task.clientName,
                eventLineName=task.eventLineName,
                deadline=format_deadline(task),
                urgency=urgency,
                urgencyColor=urgency_color,
                phase=phase,
                risks=risks_for_task(task, phase),
                nextAdvice=next_advice(task, phase),
                robotReady=robot_ready,
                robotReasons=robot_reasons,
                recommendationId=None,
                linkedTaskId=task.id,
                linkedContexts=contexts,
                xpReward=28 if task.priority == "high" else 22 if task.priority == "normal" else 16,
                contextSummary=(task.projectContext.backgroundSummary if task.projectContext else "") or task.desc or task.note or "",
                projectModuleName=(task.projectContext.projectModuleName if task.projectContext else None) or task.projectModuleName,
                projectFlowName=(task.projectContext.projectFlowName if task.projectContext else None) or task.projectFlowName,
                projectStage=(task.projectContext.stage if task.projectContext else None),
                businessCategory=task.businessCategory,
                sourceEvidence=(task.projectContext.sourceEvidence if task.projectContext else []) or [],
                currentBlocker=task.currentBlocker or (task.projectContext.currentBlocker if task.projectContext else None) or (task.orgContext.blockedAtStep if task.orgContext else None),
                missingSignals=[
                    item
                    for item in (
                        "缺任务背景说明" if not normalize_text(task.desc) and not normalize_text(task.note) else "",
                        "缺明确时间点" if not task.dueDate and not task.ddl else "",
                        "缺协作边界确认" if ((task.orgContext.isCrossDepartment if task.orgContext else False) or task.collaborators) else "",
                        "缺复核说明" if (task.orgContext.needsReview if task.orgContext else False) else "",
                    )
                    if item
                ],
                hasBackground=bool(normalize_text(task.desc) or normalize_text(task.note) or (task.projectContext.backgroundSummary if task.projectContext else "")),
                hasDeadline=bool(task.dueDate or task.ddl),
                isCrossDepartment=bool((task.orgContext.isCrossDepartment if task.orgContext else False) or task.collaborators),
                needsReview=bool(task.orgContext.needsReview if task.orgContext else False),
                evidenceCount=task.evidenceCount,
                pendingCollaborations=int(task.collaborationSummary.get("pending", 0)),
                taskIntent=task_intent,
                universalSkills=universal_skills,
                projectContextPack=project_context_pack,
                actionPlan=action_plan,
                materialRefs=material_refs,
            )

        def workbench_task_from_focus(index: int, action: GrowthFocusActionRecord) -> GrowthWorkbenchTaskRecord:
            phase = next((item[1] for item in phase_blueprints if item[1] in (action.triggerNode or "") or item[1] in (action.projectStage or "") or item[1] in action.title or item[1] in action.summary), phase_blueprints[min(index + 2, len(phase_blueprints) - 1)][1])
            linked_contexts = list(action.linkedContexts)
            if action.linkedTaskId and not any(context.objectType == "task" and context.objectId == action.linkedTaskId for context in linked_contexts):
                linked_contexts.insert(
                    0,
                    GrowthContextLinkRecord(
                        objectType="task",
                        objectId=action.linkedTaskId,
                        label=action.title,
                        subtitle=action.projectStage or action.eventLineName or action.clientName or "当前焦点",
                        tab="tasks",
                        statusLabel="成长练习",
                    ),
                )
            return GrowthWorkbenchTaskRecord(
                id=f"focus-{action.id}",
                title=action.title,
                project=action.clientName or action.eventLineName or action.triggerNode or "成长焦点",
                clientName=action.clientName,
                eventLineName=action.eventLineName,
                deadline="本周补动作",
                urgency="建议优先处理" if any(keyword in action.whyNow for keyword in ("风险", "卡住", "返工", "阻塞", "现在")) else "需先补关键动作",
                urgencyColor="text-red-700 bg-red-100" if any(keyword in action.whyNow for keyword in ("风险", "卡住", "返工", "阻塞", "现在")) else "text-orange-700 bg-orange-100",
                phase=phase,
                risks=[action.whyNow or action.summary or "当前动作还没有稳定落到真实任务中。"],
                nextAdvice=action.summary or action.whyNow or f"先围绕 {action.title} 补一条可执行动作。",
                robotReady=any(keyword in f"{action.title}{action.summary}{action.whyNow}" for keyword in ("模板", "清单", "纪要", "生成", "对齐", "跟踪", "排查", "草案")),
                robotReasons=["当前动作有清晰输出", "已匹配到可复用练习或模板", "适合先让机器人生成草案再人工判断"] if any(keyword in f"{action.title}{action.summary}{action.whyNow}" for keyword in ("模板", "清单", "纪要", "生成", "对齐", "跟踪", "排查", "草案")) else ["仍需要人工结合现场判断", "当前动作更偏策略或协作博弈，不适合直接自动执行"],
                recommendationId=None,
                linkedTaskId=action.linkedTaskId,
                linkedContexts=linked_contexts,
                xpReward=20,
                contextSummary=action.summary,
                projectFlowName=action.triggerNode,
                projectStage=action.projectStage,
                sourceEvidence=[item for item in (action.whyNow, action.summary) if item],
                currentBlocker=action.whyNow or None,
                missingSignals=[action.whyNow] if action.whyNow else [],
                hasBackground=True,
                hasDeadline=False,
                isCrossDepartment=bool(action.eventLineId or action.clientId),
                needsReview=False,
                evidenceCount=1,
                pendingCollaborations=0,
            )

        def workbench_task_from_recommendation(index: int, recommendation: LearningRecommendationRecord) -> GrowthWorkbenchTaskRecord:
            phase = next((item[1] for item in phase_blueprints if item[1] in (recommendation.projectStage or "") or item[1] in (recommendation.triggerNode or "") or item[1] in recommendation.title), phase_blueprints[min(index + 2, len(phase_blueprints) - 1)][1])
            linked_contexts = list(recommendation.linkedContexts)
            if recommendation.linkedTaskId and not any(context.objectType == "task" and context.objectId == recommendation.linkedTaskId for context in linked_contexts):
                linked_contexts.insert(
                    0,
                    GrowthContextLinkRecord(
                        objectType="task",
                        objectId=recommendation.linkedTaskId,
                        label=recommendation.title,
                        subtitle=recommendation.projectStage or recommendation.eventLineName or recommendation.clientName or "成长练习",
                        tab="tasks",
                        statusLabel="成长练习",
                    ),
                )
            return GrowthWorkbenchTaskRecord(
                id=f"recommendation-{recommendation.id}",
                title=recommendation.title,
                project=recommendation.clientName or recommendation.eventLineName or recommendation.abilityLabel,
                clientName=recommendation.clientName,
                eventLineName=recommendation.eventLineName,
                deadline="本周排期" if recommendation.priority == "high" else "可安排到下周",
                urgency="建议优先处理" if recommendation.priority == "high" else "需先补关键动作",
                urgencyColor="text-red-700 bg-red-100" if recommendation.priority == "high" else "text-orange-700 bg-orange-100",
                phase=phase,
                risks=[recommendation.reason or recommendation.summary],
                nextAdvice=recommendation.practiceTask or recommendation.summary,
                robotReady=any(keyword in f"{recommendation.title}{recommendation.summary}{recommendation.practiceTask}" for keyword in ("模板", "清单", "纪要", "生成", "对齐", "跟踪", "排查")),
                robotReasons=["任务输出格式明确", "已匹配学习资产", "当前阶段可先由机器人生成首稿"] if any(keyword in f"{recommendation.title}{recommendation.summary}{recommendation.practiceTask}" for keyword in ("模板", "清单", "纪要", "生成", "对齐", "跟踪", "排查")) else ["关键判断仍需人工定调", "上下文还需要结合现场信息", "属于高博弈或高创造性动作"],
                recommendationId=recommendation.id,
                linkedTaskId=recommendation.linkedTaskId,
                linkedContexts=linked_contexts,
                xpReward=20,
                contextSummary=recommendation.reason or recommendation.summary,
                projectFlowName=recommendation.triggerNode,
                projectStage=recommendation.projectStage,
                sourceEvidence=[recommendation.summary] if recommendation.summary else [],
                currentBlocker=recommendation.reason or None,
                missingSignals=[recommendation.reason] if recommendation.reason else [],
                hasBackground=True,
                hasDeadline=False,
                isCrossDepartment=bool(recommendation.eventLineId or recommendation.clientId),
                needsReview=False,
                evidenceCount=1,
                pendingCollaborations=0,
            )

        def workbench_task_from_capture(index: int, capture: GrowthPendingCaptureRecord) -> GrowthWorkbenchTaskRecord:
            phase = next((item[1] for item in phase_blueprints if item[1] in (capture.projectStage or "") or item[1] in capture.nextActionText or item[1] in capture.summary), "复盘沉淀" if capture.sourceType == "task_attachment_candidate" else phase_blueprints[min(index + 3, len(phase_blueprints) - 1)][1])
            linked_task = next((context.objectId for context in capture.linkedContexts if context.objectType == "task"), None)
            return GrowthWorkbenchTaskRecord(
                id=f"capture-{capture.id}",
                title=capture.title,
                project=capture.clientName or capture.eventLineName or "待放大成长",
                clientName=capture.clientName,
                eventLineName=capture.eventLineName,
                deadline="等待闭环",
                urgency="需先补关键动作" if any(any(keyword in reason for keyword in ("复盘", "沉淀", "闭环")) for reason in capture.missingReasons) else "可继续推进",
                urgencyColor="text-orange-700 bg-orange-100" if any(any(keyword in reason for keyword in ("复盘", "沉淀", "闭环")) for reason in capture.missingReasons) else "text-green-700 bg-green-100",
                phase=phase,
                risks=capture.missingReasons[:2] or [capture.summary or "系统已经识别到成长信号，但还缺最终闭环。"],
                nextAdvice=capture.nextActionText or capture.summary or "先补资料、复盘或沉淀，再把这条成长放大。",
                robotReady=False,
                robotReasons=["当前更适合先由人补资料、复盘或沉淀说明", "这类信号需要解释层，不适合只靠自动执行完成"],
                recommendationId=None,
                linkedTaskId=linked_task,
                linkedContexts=list(capture.linkedContexts),
                xpReward=16,
                contextSummary=capture.summary,
                projectFlowName=capture.projectStage,
                projectStage=capture.projectStage,
                sourceEvidence=list(capture.missingReasons),
                currentBlocker=capture.missingReasons[0] if capture.missingReasons else None,
                missingSignals=list(capture.missingReasons),
                hasBackground=True,
                hasDeadline=False,
                isCrossDepartment=bool(capture.eventLineId or capture.clientId),
                needsReview=any(any(keyword in reason for keyword in ("复盘", "解释", "说明")) for reason in capture.missingReasons),
                evidenceCount=1,
                pendingCollaborations=0,
            )

        def context_key(context: GrowthContextLinkRecord) -> str:
            return f"{context.objectType}:{context.objectId}"

        def overlaps(left: list[GrowthContextLinkRecord], right: list[GrowthContextLinkRecord]) -> bool:
            if not left or not right:
                return False
            right_keys = {context_key(context) for context in right}
            return any(context_key(context) in right_keys for context in left)

        def matches_task(
            task: GrowthWorkbenchTaskRecord,
            *,
            linked_task_id: str | None = None,
            linked_contexts: list[GrowthContextLinkRecord] | None = None,
            client_name: str | None = None,
            event_line_name: str | None = None,
            project_stage: str | None = None,
        ) -> bool:
            if linked_task_id and task.linkedTaskId and linked_task_id == task.linkedTaskId:
                return True
            if overlaps(task.linkedContexts, linked_contexts or []):
                return True
            if normalize_text(event_line_name) and normalize_text(event_line_name) in normalize_text(task.project):
                return True
            if normalize_text(client_name) and normalize_text(client_name) in normalize_text(task.project):
                return True
            if normalize_text(project_stage) and normalize_text(project_stage) in normalize_text(task.phase):
                return True
            return False

        def dedupe_strings(values: list[str], *, limit: int | None = None) -> list[str]:
            seen: set[str] = set()
            output: list[str] = []
            for value in values:
                normalized = normalize_text(value)
                if not normalized or normalized in seen:
                    continue
                seen.add(normalized)
                output.append(normalized)
                if limit is not None and len(output) >= limit:
                    break
            return output

        def preferred_task_context(task: GrowthWorkbenchTaskRecord) -> GrowthContextLinkRecord | None:
            return (
                next((context for context in task.linkedContexts if context.objectType == "task"), None)
                or next((context for context in task.linkedContexts if context.objectType == "event_line"), None)
                or next((context for context in task.linkedContexts if context.objectType == "client"), None)
                or (task.linkedContexts[0] if task.linkedContexts else None)
            )

        def infer_learning_ability_keys(
            task: GrowthWorkbenchTaskRecord,
            focus_actions: list[GrowthFocusActionRecord],
            captures: list[GrowthPendingCaptureRecord],
            recommendations: list[LearningRecommendationRecord],
        ) -> list[str]:
            keys: list[str] = [item.abilityKey for item in recommendations if item.abilityKey]
            haystack = " ".join(
                item
                for item in (
                    task.title,
                    task.project,
                    task.phase,
                    task.currentBlocker or "",
                    task.contextSummary,
                    " ".join(task.risks),
                    " ".join(item.title for item in focus_actions),
                    " ".join(item.title for item in captures),
                )
                if item
            )
            if task.isCrossDepartment or any(keyword in haystack for keyword in ("沟通", "协作", "对齐", "会议", "负责人")):
                keys.append("collab")
            if task.currentBlocker or any(keyword in haystack for keyword in ("风险", "卡点", "依赖", "阻塞", "返工")):
                keys.append("risk")
            if any(keyword in haystack for keyword in ("方案", "提案", "文档", "白皮书", "输出", "写")):
                keys.append("write")
            if any(keyword in haystack for keyword in ("分析", "判断", "研究", "说明", "原因")):
                keys.append("analyze")
            if any(keyword in haystack for keyword in ("客户", "访谈", "顾虑", "诉求", "对象")):
                keys.append("insight")
            if task.phase in {"需求接收", "信息核对", "交付闭环"}:
                keys.append("exec")
            ordered = [key for key in dict.fromkeys(keys) if key in {"exec", "collab", "analyze", "insight", "risk", "write"}]
            return ordered or ["exec", "risk"]

        def build_generic_lessons(
            task: GrowthWorkbenchTaskRecord,
            focus_actions: list[GrowthFocusActionRecord],
            captures: list[GrowthPendingCaptureRecord],
            recommendations: list[LearningRecommendationRecord],
        ) -> list[GrowthGenericLessonRecord]:
            lessons: list[GrowthGenericLessonRecord] = []
            default_context = preferred_task_context(task)
            for recommendation in recommendations[:2]:
                lessons.append(
                    GrowthGenericLessonRecord(
                        id=f"recommendation-{recommendation.id}",
                        title=recommendation.title or recommendation.summary,
                        judgment=recommendation.summary or recommendation.reason or recommendation.practiceTask,
                        applicableScene=recommendation.projectStage or recommendation.triggerNode or task.phase,
                        whyItWorks=recommendation.reason or recommendation.whyNow or recommendation.body or "这条方法来自近期真实成长推荐，可直接作为当前任务的练习模板。",
                        reuseHint=recommendation.practiceTask or "把这条方法写回到任务模板、会议纪要或复盘沉淀里。",
                        linkedContext=recommendation.linkedContexts[0] if recommendation.linkedContexts else default_context,
                    )
                )
            fallback_items = build_generic_learning_fallback(
                infer_learning_ability_keys(task, focus_actions, captures, recommendations),
                limit=3,
            )
            existing_titles = {normalize_text(item.title) for item in lessons}
            for item in fallback_items:
                if normalize_text(item.title) in existing_titles:
                    continue
                lessons.append(
                    GrowthGenericLessonRecord(
                        id=f"fallback-{item.id}",
                        title=item.title,
                        judgment=item.summary or item.body,
                        applicableScene=f"当前处在「{task.phase}」阶段，适合先把动作标准压实。",
                        whyItWorks=item.body or item.summary,
                        reuseHint=item.practiceTask or "把这条方法沉淀到成长手册，后续同类任务直接复用。",
                        linkedContext=default_context,
                    )
                )
                if len(lessons) >= 3:
                    break
            return lessons[:3]

        def build_project_guidance(
            task: GrowthWorkbenchTaskRecord,
            focus_actions: list[GrowthFocusActionRecord],
            captures: list[GrowthPendingCaptureRecord],
            recommendations: list[LearningRecommendationRecord],
            *,
            source_mode_value: Literal["task", "growth_seed", "empty"],
        ) -> list[GrowthProjectGuidanceRecord]:
            guidance: list[GrowthProjectGuidanceRecord] = []
            if source_mode_value != "task":
                guidance.append(
                    GrowthProjectGuidanceRecord(
                        id=f"{task.id}-context-mode",
                        title="当前还不是完整项目判断",
                        judgment="现在更多是成长推荐或待放大信号，不是来自一条上下文完整的真实任务。",
                        whySpecial="缺少真实任务、附件证据和连续事件线时，系统只能给规则基础版建议，不能假装已经理解了整个项目。",
                        guidanceType="context_gap",
                        linkedContexts=task.linkedContexts,
                        evidenceRefs=dedupe_strings(task.missingSignals + ["缺真实任务上下文"], limit=3),
                    )
                )
            if task.eventLineName:
                guidance.append(
                    GrowthProjectGuidanceRecord(
                        id=f"{task.id}-event-line",
                        title=f"这条动作属于事件线「{task.eventLineName}」",
                        judgment=task.currentBlocker or task.nextAdvice or "这次判断标准不是把单点动作做满，而是让整条业务线继续向前。",
                        whySpecial="一旦任务已经挂到事件线上，就不能只把它当孤立事项处理，优先级应该围绕整条线的连续推进来判断。",
                        guidanceType="project_specific",
                        linkedContexts=[context for context in task.linkedContexts if context.objectType in {"event_line", "task", "client"}],
                        evidenceRefs=dedupe_strings(task.sourceEvidence + ([task.currentBlocker] if task.currentBlocker else []), limit=3),
                    )
                )
            elif task.projectFlowName or task.projectModuleName:
                subject = task.projectFlowName or task.projectModuleName or "当前项目流程"
                guidance.append(
                    GrowthProjectGuidanceRecord(
                        id=f"{task.id}-project-flow",
                        title=f"当前动作受「{subject}」约束",
                        judgment=f"这次更像「{task.phase}」阶段的推进节点，要优先满足流程继续推进的标准，而不是一次性把内容写满。",
                        whySpecial="项目模块和流程已经给了当前任务明确的判断边界，所以真正关键的是让这个节点向前，而不是追求泛化的完整产出。",
                        guidanceType="project_specific",
                        linkedContexts=[context for context in task.linkedContexts if context.objectType in {"project_flow", "project_module", "task"}],
                        evidenceRefs=dedupe_strings(task.sourceEvidence + [subject], limit=3),
                    )
                )
            if task.currentBlocker or task.risks:
                guidance.append(
                    GrowthProjectGuidanceRecord(
                        id=f"{task.id}-stage-risk",
                        title="当前阶段最容易返工的点",
                        judgment=task.currentBlocker or task.risks[0],
                        whySpecial="这条风险不是静态模板里推导出来的泛化句，而是当前任务对象已经显式暴露出的阻塞或缺口。",
                        guidanceType="stage_risk",
                        linkedContexts=[context for context in task.linkedContexts if context.objectType in {"task", "event_line", "client"}],
                        evidenceRefs=dedupe_strings(task.risks + task.missingSignals, limit=3),
                    )
                )
            if task.isCrossDepartment or task.pendingCollaborations > 0 or task.needsReview or not task.hasBackground or task.evidenceCount <= 0:
                gap_reasons = dedupe_strings(
                    [
                        "缺任务背景说明" if not task.hasBackground else "",
                        "缺附件或明确证据" if task.evidenceCount <= 0 and not task.sourceEvidence else "",
                        "缺协作边界确认" if task.isCrossDepartment or task.pendingCollaborations > 0 else "",
                        "缺复核说明或审批依据" if task.needsReview else "",
                    ],
                    limit=4,
                )
                guidance.append(
                    GrowthProjectGuidanceRecord(
                        id=f"{task.id}-context-gap",
                        title="项目特有判断还不够稳",
                        judgment="当前项目背景、证据或协作边界仍有缺口，所以系统只能先给基础建议，不能假装已经理解了全部业务语境。",
                        whySpecial="先把缺口补齐，再去看更深的项目判断，页面才能真正像助理而不是像模板卡。",
                        guidanceType="context_gap",
                        linkedContexts=task.linkedContexts,
                        evidenceRefs=gap_reasons,
                    )
                )
            return guidance[:3]

        def build_reasoning_trace(
            task: GrowthWorkbenchTaskRecord,
            focus_actions: list[GrowthFocusActionRecord],
            captures: list[GrowthPendingCaptureRecord],
            recommendations: list[LearningRecommendationRecord],
            *,
            source_mode_value: Literal["task", "growth_seed", "empty"],
        ) -> GrowthReasoningTraceRecord:
            used_inputs: list[GrowthReasoningInputRecord] = []
            for context in task.linkedContexts[:4]:
                source_type = context.objectType if context.objectType in {"task", "event_line", "client", "project_module", "project_flow"} else "rule"
                used_inputs.append(
                    GrowthReasoningInputRecord(
                        id=f"context-{context.objectType}-{context.objectId}",
                        sourceType=source_type,
                        label=context.label,
                        detail=context.subtitle or context.statusLabel or "",
                    )
                )
            for action in focus_actions[:1]:
                used_inputs.append(
                    GrowthReasoningInputRecord(
                        id=f"focus-{action.id}",
                        sourceType="focus_action",
                        label=action.title,
                        detail=action.summary or action.whyNow,
                    )
                )
            for capture in captures[:1]:
                used_inputs.append(
                    GrowthReasoningInputRecord(
                        id=f"capture-{capture.id}",
                        sourceType="pending_capture",
                        label=capture.title,
                        detail=capture.summary or capture.nextActionText,
                    )
                )
            for recommendation in recommendations[:1]:
                used_inputs.append(
                    GrowthReasoningInputRecord(
                        id=f"recommendation-{recommendation.id}",
                        sourceType="recommendation",
                        label=recommendation.title,
                        detail=recommendation.summary or recommendation.reason,
                    )
                )
            if not used_inputs:
                used_inputs.append(
                    GrowthReasoningInputRecord(
                        id="rule-only",
                        sourceType="rule",
                        label="规则推导基线",
                        detail="当前没有足够的真实对象输入，系统只能输出基础规则判断。",
                    )
                )

            evidence_refs = dedupe_strings(
                task.sourceEvidence
                + ([task.currentBlocker] if task.currentBlocker else [])
                + task.risks
                + [capture.summary for capture in captures if capture.summary]
                + [recommendation.summary for recommendation in recommendations if recommendation.summary],
                limit=6,
            )
            missing_context = dedupe_strings(
                task.missingSignals
                + [
                    "当前没有真实任务上下文" if source_mode_value != "task" else "",
                    "缺事件线连续上下文" if not any(context.objectType == "event_line" for context in task.linkedContexts) else "",
                    "缺项目模块或流程归属" if not any(context.objectType in {"project_module", "project_flow"} for context in task.linkedContexts) else "",
                    "缺附件或明确证据" if task.evidenceCount <= 0 and not task.sourceEvidence else "",
                    "缺任务背景说明" if not task.hasBackground else "",
                ],
                limit=6,
            )
            confidence: Literal["high", "medium", "low"]
            if source_mode_value == "task" and task.hasBackground and (task.evidenceCount > 0 or task.sourceEvidence) and len(missing_context) <= 1:
                confidence = "high"
            elif source_mode_value == "empty" or len(missing_context) >= 3:
                confidence = "low"
            else:
                confidence = "medium"
            return GrowthReasoningTraceRecord(
                mode="rules_only",
                usedInputs=used_inputs[:6],
                evidenceRefs=evidence_refs,
                missingContext=missing_context,
                aiContribution=[],
                modelLabel=None,
                confidence=confidence,
            )

        def build_ai_learning_synthesis(
            task: GrowthWorkbenchTaskRecord,
            generic_lessons: list[GrowthGenericLessonRecord],
            guidance_items: list[GrowthProjectGuidanceRecord],
            reasoning: GrowthReasoningTraceRecord,
            *,
            source_mode_value: Literal["task", "growth_seed", "empty"],
        ) -> tuple[
            list[GrowthGenericLessonRecord],
            list[GrowthProjectGuidanceRecord],
            GrowthReasoningTraceRecord,
            GrowthLearningSummaryRecord,
        ] | None:
            ai_health = state.ai.get_health()
            provider_name = str(ai_health.provider or "").strip().lower()
            if source_mode_value != "task" or not ai_health.ready or provider_name in {"", "mock"}:
                return None
            if not task.hasBackground or len(reasoning.missingContext) >= 4:
                return None

            def clean_ai_text(value: str | None, *, prefixes: tuple[str, ...] = ()) -> str:
                text = re.sub(r"\s+", " ", str(value or "")).strip().strip("•-")
                for prefix in prefixes:
                    if text.startswith(prefix):
                        text = text[len(prefix):].strip("：: -")
                return text

            context_lines: list[str] = [
                f"任务标题：{task.title}",
                f"项目：{task.project}",
                f"阶段：{task.phase}",
                f"当前建议动作：{task.nextAdvice}",
                f"当前阻塞：{task.currentBlocker or '暂无显式阻塞'}",
                f"上下文摘要：{task.contextSummary or '暂无'}",
                f"当前缺口：{'；'.join(reasoning.missingContext) or '暂无'}",
            ]
            if task.linkedContexts:
                context_lines.append(
                    "关联对象：" + "；".join(
                        f"{item.label}{f'（{item.subtitle}）' if item.subtitle else ''}"
                        for item in task.linkedContexts[:6]
                    )
                )
            if reasoning.evidenceRefs:
                context_lines.append("证据引用：" + "；".join(reasoning.evidenceRefs[:6]))
            if guidance_items:
                context_lines.append(
                    "当前规则判断：" + "；".join(f"{item.title}：{item.judgment}" for item in guidance_items[:3])
                )
            if generic_lessons:
                context_lines.append(
                    "当前可迁移方法：" + "；".join(f"{item.title}：{item.judgment}" for item in generic_lessons[:3])
                )

            prompt = (
                "请基于下面的真实任务上下文，为“任务学习页”补一版 AI 综合判断。\n"
                "输出约束：\n"
                "1. judgment：只写一句“这次真正要学什么”，控制在 16-36 字。\n"
                "2. analysis：第一行写“为什么值得学：…”。后面最多三行，每行分别以“项目特有：”“阶段风险：”“上下文缺口：”开头。\n"
                "3. actions：只写一句“现在先做什么”，不要解释。\n"
                "4. timeline：每行以“AI贡献：”或“置信度：”开头，最多四行。\n"
                "5. 不要假装已经掌握缺失信息；如果证据不足，必须明确写进“上下文缺口”。\n"
                "6. 保持中文、简洁、可执行，不要写套话。\n\n"
                f"{chr(10).join(context_lines)}"
            )
            system_instruction = (
                "你是任务学习页的项目判断助手。你的职责不是替代规则，而是在真实任务上下文已经足够时，"
                "补充项目特有判断、阶段风险和可迁移方法。必须如实说明缺口，不得编造。"
            )

            try:
                structured = state.ai.generate_structured(prompt, system_instruction, task.contextSummary or "")
            except Exception:
                return None

            ai_headline = clean_ai_text(structured.judgment)
            ai_action = clean_ai_text(structured.actions) or task.nextAdvice
            analysis_lines = [
                clean_ai_text(line)
                for line in re.split(r"[\r\n]+", str(structured.analysis or ""))
                if clean_ai_text(line)
            ]
            why_line = next(
                (
                    clean_ai_text(line, prefixes=("为什么值得学", "为什么值得学：", "为什么值得学:"))
                    for line in analysis_lines
                    if line.startswith("为什么值得学")
                ),
                "",
            )
            if not why_line:
                why_line = clean_ai_text(structured.analysis) or (
                    guidance_items[0].whySpecial
                    if guidance_items
                    else "AI 已基于当前真实任务、背景和证据做了一次项目特有判断补充。"
                )

            ai_guidance: list[GrowthProjectGuidanceRecord] = []
            for index, line in enumerate(analysis_lines, start=1):
                guidance_type: Literal["project_specific", "stage_risk", "context_gap"] | None = None
                title = ""
                if line.startswith("项目特有"):
                    guidance_type = "project_specific"
                    title = "AI 识别到的项目特有提醒"
                elif line.startswith("阶段风险"):
                    guidance_type = "stage_risk"
                    title = "AI 识别到的阶段风险"
                elif line.startswith("上下文缺口"):
                    guidance_type = "context_gap"
                    title = "AI 识别到的上下文缺口"
                if guidance_type is None:
                    continue
                ai_guidance.append(
                    GrowthProjectGuidanceRecord(
                        id=f"{task.id}-ai-guidance-{index}",
                        title=title,
                        judgment=clean_ai_text(line, prefixes=("项目特有", "阶段风险", "上下文缺口")),
                        whySpecial="这条判断来自 AI 对真实任务、项目上下文和当前缺口的综合归纳。",
                        guidanceType=guidance_type,
                        linkedContexts=task.linkedContexts[:4],
                        evidenceRefs=reasoning.evidenceRefs[:4],
                    )
                )

            ai_timeline_lines = [
                clean_ai_text(line)
                for line in re.split(r"[\r\n]+", str(structured.timeline or ""))
                if clean_ai_text(line)
            ]
            ai_contribution = [
                clean_ai_text(line, prefixes=("AI贡献", "AI贡献：", "AI贡献:"))
                for line in ai_timeline_lines
                if line.startswith("AI贡献")
            ]
            confidence_line = next((line for line in ai_timeline_lines if line.startswith("置信度")), "")
            confidence_value = reasoning.confidence
            if "高" in confidence_line:
                confidence_value = "high"
            elif "中" in confidence_line:
                confidence_value = "medium"
            elif "低" in confidence_line:
                confidence_value = "low"

            merged_generic_lessons = list(generic_lessons)
            if ai_headline:
                ai_generic_lesson = GrowthGenericLessonRecord(
                    id=f"{task.id}-ai-lesson",
                    title="AI 抽象出的通用方法",
                    judgment=ai_headline,
                    applicableScene=f"当前处在「{task.phase}」阶段，且已有真实任务上下文。",
                    whyItWorks=why_line,
                    reuseHint="下次遇到同类任务时，先用这条方法判断应该先补哪一步，再决定是否扩写细节。",
                    linkedContext=preferred_task_context(task),
                )
                existing_titles = {normalize_text(item.title) for item in merged_generic_lessons}
                if normalize_text(ai_generic_lesson.title) not in existing_titles:
                    merged_generic_lessons = [ai_generic_lesson, *merged_generic_lessons]

            merged_guidance = [*ai_guidance, *guidance_items]
            deduped_guidance: list[GrowthProjectGuidanceRecord] = []
            seen_guidance_keys: set[str] = set()
            for item in merged_guidance:
                dedupe_key = normalize_text(f"{item.guidanceType}:{item.judgment}")
                if dedupe_key in seen_guidance_keys:
                    continue
                seen_guidance_keys.add(dedupe_key)
                deduped_guidance.append(item)

            updated_reasoning = GrowthReasoningTraceRecord(
                mode="ai_synthesized",
                usedInputs=reasoning.usedInputs,
                evidenceRefs=reasoning.evidenceRefs,
                missingContext=reasoning.missingContext,
                aiContribution=ai_contribution or [
                    "AI 已基于真实任务、项目上下文和证据线索做了一次项目特有判断补充。",
                    "AI 没有覆盖缺失上下文，缺口仍保留在页面里。",
                ],
                modelLabel=state.ai.current_model(),
                confidence=confidence_value,
            )
            updated_summary = GrowthLearningSummaryRecord(
                headline=ai_headline or reasoning.usedInputs[0].label,
                whyItMatters=why_line,
                immediateMove=ai_action,
                generator="ai",
                confidence=confidence_value,
            )
            return merged_generic_lessons[:3], deduped_guidance[:3], updated_reasoning, updated_summary

        def build_learning_summary(
            task: GrowthWorkbenchTaskRecord,
            guidance_items: list[GrowthProjectGuidanceRecord],
            reasoning: GrowthReasoningTraceRecord,
            *,
            source_mode_value: Literal["task", "growth_seed", "empty"],
        ) -> GrowthLearningSummaryRecord:
            if source_mode_value == "empty":
                return GrowthLearningSummaryRecord(
                    headline="当前还没有可学习的真实任务，上下文还没进入任务学习页。",
                    whyItMatters="没有真实任务、客户背景或事件线时，页面最多只能展示空壳，不能给出负责任的学习判断。",
                    immediateMove="先去任务与日历、客户工作台或战略陪伴生成一个真实对象，再回来学习。",
                    generator="rules",
                    confidence="low",
                )
            if source_mode_value == "growth_seed":
                return GrowthLearningSummaryRecord(
                    headline="先把成长信号压成真实任务，再谈更深的项目判断。",
                    whyItMatters="现在还处在成长种子模式，系统能告诉你先做什么，但不能假装已经理解了整个项目语境。",
                    immediateMove=task.nextAdvice or "先把这条信号落成真实任务，并补齐背景、附件和责任人。",
                    generator="rules",
                    confidence=reasoning.confidence,
                )
            if not task.hasBackground:
                headline = "这次最该学的不是直接推进，而是先把任务背景、目标和边界补清楚。"
            elif task.isCrossDepartment:
                headline = "这次真正要学的是：多人协作里先收边界、责任人与时间线。"
            elif task.phase in {"方案产出", "交付闭环"}:
                headline = "这次真正要学的是：先把结构和交付标准拉稳，再扩写细节。"
            elif task.phase == "复盘沉淀":
                headline = "这次真正要学的是：把有效动作拆成可复用方法，而不是只记结果。"
            else:
                headline = "这次真正要学的是：先判断当前阶段最关键的一步，再推进动作。"
            why_it_matters = (
                guidance_items[0].whySpecial
                if guidance_items
                else "任务学习页的价值不在于堆动作，而在于先说清这次任务真正值得学的判断。"
            )
            return GrowthLearningSummaryRecord(
                headline=headline,
                whyItMatters=why_it_matters,
                immediateMove=task.nextAdvice or (guidance_items[0].judgment if guidance_items else "先补当前任务的关键动作。"),
                generator="rules",
                confidence=reasoning.confidence,
            )

        def build_robot_assist(task: GrowthWorkbenchTaskRecord) -> GrowthRobotAssistRecord:
            haystack = f"{task.title}{task.project}{task.contextSummary}{task.currentBlocker or ''}"
            can_delegate: list[str] = []
            if any(keyword in haystack for keyword in ("会议", "对齐", "沟通", "纪要")):
                can_delegate.extend(["会议议程初稿", "会后纪要骨架", "行动项清单"])
            if any(keyword in haystack for keyword in ("方案", "提案", "白皮书", "文档", "大纲", "写")):
                can_delegate.extend(["结构化大纲", "首版文档骨架", "待确认问题清单"])
            if any(keyword in haystack for keyword in ("复盘", "总结", "方法", "沉淀")):
                can_delegate.extend(["复盘骨架", "方法卡初稿"])
            if task.evidenceCount > 0 or task.sourceEvidence:
                can_delegate.append("材料整理与证据摘录")
            if not can_delegate:
                can_delegate.extend(["待确认问题清单", "材料整理清单"])

            must_stay_human: list[str] = []
            if task.isCrossDepartment or task.pendingCollaborations > 0:
                must_stay_human.append("跨部门边界和责任分配")
            if any(keyword in haystack for keyword in ("客户", "沟通", "谈判", "协调")):
                must_stay_human.append("关键对象口径和现场判断")
            if task.needsReview:
                must_stay_human.append("复核 / 审批结论")
            must_stay_human.append("最终优先级和是否推进的拍板")
            return GrowthRobotAssistRecord(
                ready=task.robotReady,
                canDelegate=dedupe_strings(can_delegate, limit=3),
                mustStayHuman=dedupe_strings(must_stay_human, limit=3),
                why=dedupe_strings(task.robotReasons, limit=3),
            )

        def build_after_action_capture(
            task: GrowthWorkbenchTaskRecord,
            captures: list[GrowthPendingCaptureRecord],
        ) -> GrowthAfterActionCaptureRecord:
            if captures:
                capture = captures[0]
                return GrowthAfterActionCaptureRecord(
                    title=capture.title,
                    summary=capture.summary or capture.nextActionText,
                    experienceType="待放大成长信号",
                    recommendedWriteback=(
                        f"优先写回事件线「{capture.eventLineName}」" if capture.eventLineName else f"优先写回客户「{capture.clientName}」" if capture.clientName else "写回成长手册或项目经验库"
                    ),
                )
            return GrowthAfterActionCaptureRecord(
                title=f"{task.title}：{task.phase} 阶段复盘",
                summary=f"记录这次在「{task.phase}」阶段的关键判断、有效动作、适用边界和下次可复用的方法。",
                experienceType="方法卡" if task.isCrossDepartment or task.phase in {"沟通推进", "复盘沉淀"} else "复盘卡",
                recommendedWriteback=(
                    f"优先写回事件线「{task.eventLineName}」" if task.eventLineName else f"优先写回客户「{task.project}」背景与经验库" if task.project else "写回成长手册"
                ),
            )

        user_id, user_name = resolve_growth_actor()
        resolved_week = resolve_growth_week_label(user_id, week_label)
        overview = build_growth_overview(state.db, user_id=user_id, user_name=user_name, week_label=resolved_week)

        real_tasks = (
            fetch_tasks("t.source_type != ? AND t.status NOT IN ('done', 'rejected')", (AGENT_AUTO_SOURCE_TYPE,))
        )
        real_tasks = sorted(
            real_tasks,
            key=lambda task: (
                {"doing": 0, "todo": 1, "inbox": 2, "done": 3, "rejected": 4}.get(task.status, 5),
                {"high": 0, "normal": 1, "low": 2}.get(task.priority, 3),
                parse_task_date(task.dueDate or task.ddl) or datetime.max,
                -sort_updated_at(task.updatedAt).timestamp(),
            ),
        )
        workbench_tasks = [workbench_task_from_task(task) for task in real_tasks[:3]]
        source_mode: Literal["task", "growth_seed", "empty"] = "task"
        if not workbench_tasks:
            source_mode = "growth_seed"
            if overview.currentFocusActions:
                workbench_tasks = [workbench_task_from_focus(index, action) for index, action in enumerate(overview.currentFocusActions[:2])]
            elif overview.recommendations:
                workbench_tasks = [workbench_task_from_recommendation(index, item) for index, item in enumerate(overview.recommendations[:2])]
            elif overview.pendingCaptures:
                workbench_tasks = [workbench_task_from_capture(index, item) for index, item in enumerate(overview.pendingCaptures[:2])]
        if not workbench_tasks:
            source_mode = "empty"
            process_steps = [GrowthWorkbenchStepRecord(id=step_id, name=name, output=output, bottlenecks=bottlenecks) for step_id, name, output, bottlenecks in phase_blueprints]
            active_process_id = next((step.id for step in process_steps if step.name == "信息核对"), process_steps[1].id if len(process_steps) > 1 else None)
            return GrowthWorkbenchSnapshotRecord(
                tasks=[],
                activeTaskId=None,
                learningSummary=GrowthLearningSummaryRecord(
                    headline="当前还没有可学习的真实任务。",
                    whyItMatters="没有真实任务、项目上下文或成长信号时，系统不能负责任地给出学习判断。",
                    immediateMove="先去任务与日历、客户工作台或战略陪伴形成真实对象，再回来学习。",
                    generator="rules",
                    confidence="low",
                ),
                genericLessons=[],
                projectGuidance=[
                    GrowthProjectGuidanceRecord(
                        id="empty-context",
                        title="当前只能给空白提示",
                        judgment="系统还没有拿到真实任务、附件、事件线或项目背景，所以这里不会假装在做深度分析。",
                        whySpecial="任务学习页应该建立在真实对象上，而不是建立在想象中的项目上。",
                        guidanceType="context_gap",
                        linkedContexts=[],
                        evidenceRefs=["缺真实任务", "缺项目上下文"],
                    )
                ],
                reasoningTrace=GrowthReasoningTraceRecord(
                    mode="rules_only",
                    usedInputs=[
                        GrowthReasoningInputRecord(
                            id="rule-only",
                            sourceType="rule",
                            label="空上下文保护规则",
                            detail="当前没有真实对象输入，所以系统只返回空白保护提示。",
                        )
                    ],
                    evidenceRefs=[],
                    missingContext=["缺真实任务", "缺项目背景", "缺事件线连续上下文"],
                    aiContribution=[],
                    modelLabel=None,
                    confidence="low",
                ),
                robotAssist=GrowthRobotAssistRecord(
                    ready=False,
                    canDelegate=[],
                    mustStayHuman=["先创建真实任务或业务对象"],
                    why=["没有真实对象输入前，机器人也无法给出有意义的执行包。"],
                ),
                afterActionCapture=GrowthAfterActionCaptureRecord(
                    title="当前没有可沉淀内容",
                    summary="先让真实任务进入学习页，再决定沉淀成什么。",
                    experienceType="待创建",
                    recommendedWriteback="暂不写回",
                ),
                processSteps=process_steps,
                activeProcessId=active_process_id,
                actionsBefore=[],
                actionsDuring=[],
                actionsAfter=[],
                supportMaterials=[],
                checklistItems=[],
                supportCopy=GrowthWorkbenchSupportCopyRecord(
                    title="当前没有可执行的成长上下文",
                    intro="先在任务与日历创建一条任务，或在客户工作台发布会议 / 行动项，任务学习页就会自动补齐上下文。",
                    bullets=["当前没有真实任务、事件线或成长推荐进入任务学习页。"],
                ),
                robotPlan=[],
                sourceMode=source_mode,
                updatedAt=now_iso(),
            )

        active_task = workbench_tasks[0]
        related_focus_actions = [item for item in overview.currentFocusActions if matches_task(active_task, linked_task_id=item.linkedTaskId, linked_contexts=item.linkedContexts, client_name=item.clientName, event_line_name=item.eventLineName, project_stage=item.projectStage or item.triggerNode)][:3]
        related_captures = [item for item in overview.pendingCaptures if matches_task(active_task, linked_contexts=item.linkedContexts, client_name=item.clientName, event_line_name=item.eventLineName, project_stage=item.projectStage)][:3]
        related_recommendations = [item for item in overview.recommendations if matches_task(active_task, linked_task_id=item.linkedTaskId, linked_contexts=item.linkedContexts, client_name=item.clientName, event_line_name=item.eventLineName, project_stage=item.projectStage or item.triggerNode)][:3]
        generic_lessons = build_generic_lessons(active_task, related_focus_actions, related_captures, related_recommendations)
        project_guidance = build_project_guidance(active_task, related_focus_actions, related_captures, related_recommendations, source_mode_value=source_mode)
        reasoning_trace = build_reasoning_trace(active_task, related_focus_actions, related_captures, related_recommendations, source_mode_value=source_mode)
        learning_summary = build_learning_summary(active_task, project_guidance, reasoning_trace, source_mode_value=source_mode)
        robot_assist = build_robot_assist(active_task)
        after_action_capture = build_after_action_capture(active_task, related_captures)
        ai_learning_bundle = build_ai_learning_synthesis(
            active_task,
            generic_lessons,
            project_guidance,
            reasoning_trace,
            source_mode_value=source_mode,
        )
        if ai_learning_bundle:
            generic_lessons, project_guidance, reasoning_trace, learning_summary = ai_learning_bundle

        process_steps: list[GrowthWorkbenchStepRecord] = []
        active_process_id: str | None = None
        for step_id, name, output, bottlenecks in phase_blueprints:
            if name == active_task.phase:
                process_steps.append(
                    GrowthWorkbenchStepRecord(
                        id=step_id,
                        name=name,
                        output=active_task.nextAdvice or output,
                        bottlenecks=active_task.risks[:2] or bottlenecks,
                    )
                )
                active_process_id = step_id
            elif name == "复盘沉淀" and related_captures:
                process_steps.append(
                    GrowthWorkbenchStepRecord(
                        id=step_id,
                        name=name,
                        output=related_captures[0].nextActionText or output,
                        bottlenecks=related_captures[0].missingReasons[:2] or bottlenecks,
                    )
                )
            else:
                process_steps.append(GrowthWorkbenchStepRecord(id=step_id, name=name, output=output, bottlenecks=bottlenecks))
        active_process_id = active_process_id or process_steps[0].id

        primary_context = active_task.linkedContexts[0] if active_task.linkedContexts else None
        before_actions = [
            GrowthWorkbenchActionRecord(
                id=f"{active_task.id}-before-1",
                title=related_focus_actions[0].title if related_focus_actions else f"开始前先定：{active_task.title} 的目标与边界",
                output=related_focus_actions[0].summary if related_focus_actions else f"{active_task.nextAdvice or active_task.phase}，并明确第一责任人",
                scenario=f"{active_task.phase} 开始前",
                actionLabel="排入练习" if active_task.recommendationId else "打开当前任务",
                supportTitle="查看为什么要做这一步",
                detail=related_focus_actions[0].whyNow if related_focus_actions else active_task.contextSummary,
                kind="schedule" if active_task.recommendationId else "task",
                recommendationId=active_task.recommendationId,
                linkedContext=primary_context,
            ),
            GrowthWorkbenchActionRecord(
                id=f"{active_task.id}-before-2",
                title=f"优先处理卡点：{active_task.currentBlocker}" if active_task.currentBlocker else "识别风险：先排查最可能翻车的 2 个点",
                output=related_captures[0].nextActionText if related_captures else "关键争议点 + 一条可执行预案",
                scenario="正式拉人或开工前",
                actionLabel="回到当前任务" if active_task.currentBlocker else "先做风险排查",
                supportTitle="查看常见翻车案例",
                detail=(related_captures[0].missingReasons[0] if related_captures and related_captures[0].missingReasons else "") or (active_task.risks[0] if active_task.risks else ""),
                kind="task" if active_task.currentBlocker else "support",
                linkedContext=primary_context,
            ),
        ]
        during_actions = [
            GrowthWorkbenchActionRecord(
                id=f"{active_task.id}-during-1",
                title=f"执行中关键动作：稳住{active_task.phase}",
                output="各方认同的交付物、边界与时间线" if active_task.isCrossDepartment else active_task.nextAdvice,
                scenario="讨论开始发散或推进变慢时",
                actionLabel="生成沟通话术" if active_task.isCrossDepartment else "查看节点清单",
                supportTitle="查看沟通原理" if active_task.isCrossDepartment else "查看节点标准",
                detail=active_task.contextSummary,
                kind="support",
                linkedContext=primary_context,
            )
        ]
        after_actions = [
            GrowthWorkbenchActionRecord(
                id=f"{active_task.id}-after-1",
                title=f"完成后补强：{related_captures[0].title}" if related_captures else "完成后沉淀：把这次动作转成可复用经验",
                output=related_captures[0].nextActionText if related_captures else f"一条可复用经验 + {active_task.xpReward} XP 的练习回流",
                scenario="动作完成后 2 小时内",
                actionLabel="沉淀为经验" if related_captures else "去记录经验",
                supportTitle=(related_captures[0].missingReasons[0] if related_captures and related_captures[0].missingReasons else "") or "查看标准沉淀方式",
                kind="compose",
                linkedContext=primary_context,
                seedTitle=related_captures[0].title if related_captures else active_task.title,
                seedSummary=(related_captures[0].summary or related_captures[0].nextActionText) if related_captures else active_task.nextAdvice,
            )
        ]
        support_materials: list[GrowthWorkbenchMaterialRecord] = []
        if active_task.projectFlowName or active_task.projectModuleName:
            support_materials.append(
                GrowthWorkbenchMaterialRecord(
                    id=f"{active_task.id}-flow",
                    title=active_task.projectFlowName or active_task.projectModuleName or "当前项目流程说明",
                    type="流程说明",
                    scenario=active_task.contextSummary or f"适用于当前 {active_task.phase} 阶段",
                    summary=(active_task.sourceEvidence[0] if active_task.sourceEvidence else "") or active_task.nextAdvice,
                    linkedContext=next((context for context in active_task.linkedContexts if context.objectType in {"project_flow", "project_module", "task"}), primary_context),
                )
            )
        if related_recommendations:
            support_materials.append(
                GrowthWorkbenchMaterialRecord(
                    id=f"recommendation-{related_recommendations[0].id}",
                    title=related_recommendations[0].summary or related_recommendations[0].title,
                    type="模板工具" if related_recommendations[0].contentType == "practice_card" else "流程说明",
                    scenario=related_recommendations[0].whyNow or related_recommendations[0].reason,
                    summary=related_recommendations[0].practiceTask or related_recommendations[0].summary,
                    linkedContext=related_recommendations[0].linkedContexts[0] if related_recommendations[0].linkedContexts else None,
                )
            )
        if related_captures:
            support_materials.append(
                GrowthWorkbenchMaterialRecord(
                    id=f"capture-{related_captures[0].id}",
                    title=related_captures[0].title,
                    type="经验案例",
                    scenario=related_captures[0].summary or related_captures[0].projectStage or "系统已识别到待放大的成长信号",
                    summary="；".join(related_captures[0].missingReasons) or related_captures[0].nextActionText,
                    linkedContext=related_captures[0].linkedContexts[0] if related_captures[0].linkedContexts else None,
                )
            )
        support_materials = support_materials[:3]
        checklist_items = [
            f"明确该节点的预期产出：{next((step.output for step in process_steps if step.id == active_process_id), active_task.nextAdvice)}",
            "补齐任务背景、目标和预期输出" if not active_task.hasBackground else "",
            "补齐明确的截止时间或推进节奏" if not active_task.hasDeadline else "",
            "把协作边界、责任人和时间点讲清楚" if active_task.isCrossDepartment else "",
            f"完成 {active_task.pendingCollaborations} 个待确认协作动作" if active_task.pendingCollaborations > 0 else "",
            "补复核说明、审批依据或验证证据" if active_task.needsReview else "",
            f"把「{related_focus_actions[0].title}」压进当前任务动作清单" if related_focus_actions else "",
            f"完成后处理「{related_captures[0].title}」的复盘或经验沉淀" if related_captures else "",
        ]
        checklist_items = [item for item in checklist_items if item][:5]
        support_copy = GrowthWorkbenchSupportCopyRecord(
            title="为什么这件事要先讲清边界与责任？" if active_task.isCrossDepartment else ("为什么开始前一定要先补齐上下文？" if not active_task.hasBackground else (f"为什么在「{active_task.phase}」阶段要先补关键动作？")),
            intro="这类跨部门或多人任务最容易翻车的点，不是大家不努力，而是边界、责任人和时间点没有先被讲清楚。" if active_task.isCrossDepartment else ("系统已经识别到当前任务缺少背景、目标或预期输出。没有这些上下文，后续动作看起来很忙，但很容易做偏。" if not active_task.hasBackground else "学习导航不是给你一堆资料，而是提醒你在当前节点最应该补的那一步。先把动作做对，再去扩写内容。"),
            bullets=[
                "当前任务已经有基础背景，可以直接对齐关键动作。" if active_task.hasBackground else "先写清任务目标、对象和预期交付物。",
                "当前已经有时间点，下一步重点是把责任和边界讲清楚。" if active_task.hasDeadline else "没有截止时间时，动作很容易在中途失焦。",
                "跨部门任务要优先处理协作边界，避免会后推诿返工。" if active_task.isCrossDepartment else "单点任务更要先补事实依据和当前阶段判断。",
            ],
        )
        robot_plan = [
            f"根据 {active_task.project} 的上下文，先拟一版「{active_task.phase}」阶段动作清单",
            f"围绕当前卡点「{active_task.currentBlocker}」生成一版应对草案" if active_task.currentBlocker else "",
            f"把推荐动作「{related_focus_actions[0].title}」整理成可直接执行的脚本或清单" if related_focus_actions else "",
            f"预先生成「{related_captures[0].title}」对应的复盘或经验沉淀骨架" if related_captures else "",
        ]
        robot_plan = [item for item in robot_plan if item][:3]
        return GrowthWorkbenchSnapshotRecord(
            tasks=workbench_tasks,
            activeTaskId=active_task.id,
            learningSummary=learning_summary,
            genericLessons=generic_lessons,
            projectGuidance=project_guidance,
            reasoningTrace=reasoning_trace,
            robotAssist=robot_assist,
            afterActionCapture=after_action_capture,
            processSteps=process_steps,
            activeProcessId=active_process_id,
            actionsBefore=before_actions,
            actionsDuring=during_actions,
            actionsAfter=after_actions,
            supportMaterials=support_materials,
            checklistItems=checklist_items,
            supportCopy=support_copy,
            robotPlan=robot_plan,
            sourceMode=source_mode,
            updatedAt=now_iso(),
        )

    def _builtin_task_view_blueprints() -> list[dict[str, object]]:
        visible_fields = ["title", "status", "priority", "sourceType", "businessCategory", "eventLine", "evidenceCount"]
        return [
            {
                "id": "builtin_event_line_view",
                "name": "事件线视图",
                "kind": "event_line",
                "description": "优先看已经挂到事件线的持续推进事项。",
                "calendarScope": "event_line",
                "shareability": "org",
                "sortBy": "updatedAt",
                "sortDirection": "desc",
                "visibleFields": visible_fields,
                "filterSet": {"onlyWithEventLine": True},
                "builtIn": True,
            },
            {
                "id": "builtin_risk_view",
                "name": "风险视图",
                "kind": "risk",
                "description": "优先看有阻塞、待复核、低证据或已逾期的事项。",
                "calendarScope": "risk",
                "shareability": "org",
                "sortBy": "evidenceCount",
                "sortDirection": "asc",
                "visibleFields": visible_fields,
                "filterSet": {"onlyRisky": True},
                "builtIn": True,
            },
            {
                "id": "builtin_source_view",
                "name": "来源视图",
                "kind": "source",
                "description": "按会议、周判断动作、支持请求等来源追踪任务。",
                "calendarScope": "source",
                "shareability": "org",
                "sortBy": "updatedAt",
                "sortDirection": "desc",
                "visibleFields": visible_fields,
                "filterSet": {},
                "builtIn": True,
            },
            {
                "id": "builtin_business_category_view",
                "name": "业务分类视图",
                "kind": "business_category",
                "description": "按业务扩展、项目推进、组织协同等业务类别查看任务。",
                "calendarScope": "business_category",
                "shareability": "org",
                "sortBy": "updatedAt",
                "sortDirection": "desc",
                "visibleFields": visible_fields,
                "filterSet": {},
                "builtIn": True,
            },
        ]

    def _task_view_record_from_row(row) -> TaskViewDefinitionRecord:
        visible_fields = from_json(row["visible_fields_json"], [])
        filter_set = from_json(row["filter_set_json"], {})
        return TaskViewDefinitionRecord(
            id=str(row["id"]),
            name=str(row["name"]),
            kind=str(row["kind"]),  # type: ignore[arg-type]
            description=str(row["description"] or ""),
            calendarScope=str(row["calendar_scope"] or "all"),  # type: ignore[arg-type]
            shareability=str(row["shareability"] or "private"),  # type: ignore[arg-type]
            sortBy=str(row["sort_by"] or "updatedAt"),  # type: ignore[arg-type]
            sortDirection=str(row["sort_direction"] or "desc"),  # type: ignore[arg-type]
            visibleFields=[str(item) for item in visible_fields if str(item).strip()],
            filterSet=filter_set if isinstance(filter_set, dict) else {},
            builtIn=bool(int(row["built_in"] or 0)),
            createdAt=str(row["created_at"]),
            updatedAt=str(row["updated_at"]),
        )

    def _task_view_presets(views: list[TaskViewDefinitionRecord]) -> list[TaskViewPresetRecord]:
        presets: list[TaskViewPresetRecord] = []
        for key, label, description in [
            ("event_line", "事件线视图", "按持续推进的工作线查看任务"),
            ("risk", "风险视图", "优先看阻塞、复核与低证据事项"),
            ("source", "来源视图", "按会议、支持请求、周判断动作等来源查看"),
            ("business_category", "业务分类视图", "按业务扩展、项目推进、组织协同等查看"),
        ]:
            matched = next((item for item in views if item.kind == key and item.builtIn), None)
            if matched:
                presets.append(
                    TaskViewPresetRecord(
                        key=key,  # type: ignore[arg-type]
                        label=label,
                        description=description,
                        viewId=matched.id,
                    )
                )
        return presets

    def _ensure_builtin_task_views() -> list[TaskViewDefinitionRecord]:
        timestamp = now_iso()
        for blueprint in _builtin_task_view_blueprints():
            state.db.execute(
                """
                INSERT OR IGNORE INTO task_views(
                    id, name, kind, description, calendar_scope, shareability, sort_by, sort_direction,
                    visible_fields_json, filter_set_json, built_in, created_at, updated_at
                ) VALUES(?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    str(blueprint["id"]),
                    str(blueprint["name"]),
                    str(blueprint["kind"]),
                    str(blueprint["description"]),
                    str(blueprint["calendarScope"]),
                    str(blueprint["shareability"]),
                    str(blueprint["sortBy"]),
                    str(blueprint["sortDirection"]),
                    to_json(blueprint["visibleFields"]),
                    to_json(blueprint["filterSet"]),
                    1,
                    timestamp,
                    timestamp,
                ),
            )
            state.db.execute(
                """
                UPDATE task_views
                   SET name = ?, kind = ?, description = ?, calendar_scope = ?, shareability = ?,
                       sort_by = ?, sort_direction = ?, visible_fields_json = ?, filter_set_json = ?,
                       built_in = 1, updated_at = ?
                 WHERE id = ?
                """,
                (
                    str(blueprint["name"]),
                    str(blueprint["kind"]),
                    str(blueprint["description"]),
                    str(blueprint["calendarScope"]),
                    str(blueprint["shareability"]),
                    str(blueprint["sortBy"]),
                    str(blueprint["sortDirection"]),
                    to_json(blueprint["visibleFields"]),
                    to_json(blueprint["filterSet"]),
                    timestamp,
                    str(blueprint["id"]),
                ),
            )
        rows = state.db.fetchall("SELECT * FROM task_views ORDER BY built_in DESC, updated_at DESC")
        return [_task_view_record_from_row(row) for row in rows]

    def _task_is_risky(task: TaskRecord) -> bool:
        if bool(task.orgContext and task.orgContext.needsReview):
            return True
        if (task.currentBlocker or "").strip():
            return True
        if is_task_overdue(task):
            return True
        return int(task.evidenceCount or 0) <= 1 and task.status != "done"

    def _matches_task_view(task: TaskRecord, view: TaskViewDefinitionRecord, extra_filters: dict[str, object] | None = None) -> bool:
        filters = dict(extra_filters or {})
        filter_set = view.filterSet
        source_types = list(filters.get("sourceTypes") or filter_set.sourceTypes or [])
        business_categories = list(filters.get("businessCategories") or filter_set.businessCategories or [])
        event_line_ids = list(filters.get("eventLineIds") or filter_set.eventLineIds or [])
        client_names = [str(item).strip() for item in filters.get("clientNames", []) if str(item).strip()]
        related_task_ids = [str(item).strip() for item in filters.get("relatedTaskIds", []) if str(item).strip()]
        only_risky = bool(filters.get("onlyRisky", filter_set.onlyRisky))
        only_with_event_line = bool(filters.get("onlyWithEventLine", filter_set.onlyWithEventLine))
        needs_review = filters.get("needsReview", filter_set.needsReview)
        minimum_evidence = filters.get("minimumEvidenceCount", filter_set.minimumEvidenceCount)

        if view.kind == "event_line" and not (task.eventLineId or "").strip():
            return False
        if view.kind == "risk" and not _task_is_risky(task):
            return False
        if view.kind == "source" and (task.sourceType or "").strip() in {"", "manual"}:
            return False
        if view.kind == "business_category" and not (task.businessCategory or "").strip():
            return False
        if source_types and (task.sourceType or "") not in source_types:
            return False
        if business_categories and (task.businessCategory or "") not in business_categories:
            return False
        if event_line_ids and (task.eventLineId or "") not in event_line_ids:
            return False
        if client_names and (task.projectContext.clientName if task.projectContext else "") not in client_names:
            return False
        if related_task_ids and task.id not in related_task_ids:
            return False
        if only_risky and not _task_is_risky(task):
            return False
        if only_with_event_line and not (task.eventLineId or "").strip():
            return False
        if needs_review is not None and bool(task.orgContext and task.orgContext.needsReview) != bool(needs_review):
            return False
        if minimum_evidence is not None and int(task.evidenceCount or 0) < int(minimum_evidence):
            return False
        return True

    def _sort_tasks_for_view(tasks: list[TaskRecord], view: TaskViewDefinitionRecord) -> list[TaskRecord]:
        reverse = view.sortDirection == "desc"

        def priority_rank(task: TaskRecord) -> int:
            return {"high": 0, "medium": 1, "low": 2}.get(task.priority, 3)

        def due_timestamp(task: TaskRecord) -> float:
            parsed = parse_task_date_value(task.dueDate)
            return parsed.timestamp() if parsed else 0

        if view.sortBy == "priority":
            return sorted(tasks, key=priority_rank, reverse=reverse)
        if view.sortBy == "dueDate":
            return sorted(tasks, key=due_timestamp, reverse=reverse)
        if view.sortBy == "evidenceCount":
            return sorted(tasks, key=lambda item: int(item.evidenceCount or 0), reverse=reverse)
        return sorted(tasks, key=lambda item: item.updatedAt, reverse=reverse)

    def _task_records_for_ids(task_ids: list[str]) -> list[TaskRecord]:
        wanted = {task_id for task_id in task_ids if task_id}
        if not wanted:
            return []
        records = _task_records_for_views()
        return [task for task in records if task.id in wanted]

    def _attachments_for_tasks(tasks: list[TaskRecord], *, cloud: bool) -> list[TaskAttachmentRecord]:
        attachments: dict[str, TaskAttachmentRecord] = {}
        source_modes = [False, True] if cloud else [False]
        for task in tasks:
            for use_cloud in source_modes:
                for attachment in fetch_task_attachments(task.id, cloud=use_cloud):
                    attachments[attachment.id] = attachment
        return sorted(attachments.values(), key=lambda item: item.createdAt, reverse=True)

    def _attachments_for_ids(attachment_ids: list[str]) -> list[TaskAttachmentRecord]:
        wanted = {attachment_id for attachment_id in attachment_ids if attachment_id}
        if not wanted:
            return []
        attachments: dict[str, TaskAttachmentRecord] = {}
        for table_name in ("task_attachments", "task_attachments_cloud"):
            rows = state.db.fetchall(
                f"SELECT * FROM {table_name} WHERE id IN ({_sql_placeholders(tuple(wanted))}) ORDER BY created_at DESC",
                tuple(wanted),
            )
            for row in rows:
                attachment = build_task_attachment(row)
                attachments[attachment.id] = attachment
        return sorted(attachments.values(), key=lambda item: item.createdAt, reverse=True)

    def _attachments_for_event_line(event_line_id: str) -> list[TaskAttachmentRecord]:
        if not event_line_id:
            return []
        attachments: dict[str, TaskAttachmentRecord] = {}
        for table_name in ("task_attachments", "task_attachments_cloud"):
            rows = state.db.fetchall(
                f"SELECT * FROM {table_name} WHERE event_line_id = ? ORDER BY created_at DESC",
                (event_line_id,),
            )
            for row in rows:
                attachment = build_task_attachment(row)
                attachments[attachment.id] = attachment
        return sorted(attachments.values(), key=lambda item: item.createdAt, reverse=True)

    def _meeting_summaries_for_tasks(tasks: list[TaskRecord]) -> list[MeetingSummary]:
        meeting_ids = [task.sourceId for task in tasks if task.sourceType == "meeting" and task.sourceId]
        if not meeting_ids:
            return []
        rows = state.db.fetchall(
            f"SELECT * FROM meetings WHERE id IN ({_sql_placeholders(meeting_ids)}) ORDER BY updated_at DESC",
            tuple(meeting_ids),
        )
        return [
            MeetingSummary(
                id=str(row["id"]),
                clientId=str(row["client_id"]),
                title=str(row["title"]),
                stage=str(row["stage"]),  # type: ignore[arg-type]
                scheduledAt=str(row["scheduled_at"]) if row["scheduled_at"] else None,
                updatedAt=str(row["updated_at"]),
            )
            for row in rows
        ]

    def _meeting_summary_for_id(meeting_id: str) -> MeetingSummary:
        row = state.db.fetchone("SELECT * FROM meetings WHERE id = ?", (meeting_id,))
        if not row:
            raise HTTPException(status_code=404, detail="Meeting not found")
        return build_meeting_summary(row)

    def _support_request_by_id(request_id: str) -> SupportRequestRecord:
        if not get_cloud_token():
            raise HTTPException(status_code=404, detail="Support request drill-down requires cloud support")
        response = cloud_request("GET", "/api/v1/support-requests")
        if not isinstance(response, list):
            raise HTTPException(status_code=502, detail="Invalid support request payload")
        for item in response:
            if not isinstance(item, dict):
                continue
            if str(item.get("id")) == request_id:
                return SupportRequestRecord(**item)
        raise HTTPException(status_code=404, detail="Support request not found")

    def _task_records_for_views() -> list[TaskRecord]:
        if get_cloud_token():
            try:
                return cloud_task_board().tasks
            except HTTPException:
                return []
        return fetch_tasks("t.source_type != ?", (AGENT_AUTO_SOURCE_TYPE,))

    def _load_task_view_definition(view_id: str) -> TaskViewDefinitionRecord | None:
        rows = _ensure_builtin_task_views()
        return next((item for item in rows if item.id == view_id), None)

    def _build_ad_hoc_task_view(
        *,
        target_id: str,
        target_label: str | None,
        target_filters: dict[str, object] | None,
    ) -> TaskViewDefinitionRecord:
        filter_set = target_filters if isinstance(target_filters, dict) else {}
        return TaskViewDefinitionRecord(
            id=target_id or "adhoc_task_view",
            name=target_label or "临时任务视图",
            kind="custom",
            description="从周判断卡片下钻得到的临时任务集合。",
            calendarScope="all",
            shareability="private",
            sortBy="updatedAt",
            sortDirection="desc",
            visibleFields=["title", "status", "priority", "sourceType", "businessCategory", "eventLine", "evidenceCount"],
            filterSet=filter_set,
            builtIn=False,
            createdAt=now_iso(),
            updatedAt=now_iso(),
        )

    def _tasks_for_task_view(
        view: TaskViewDefinitionRecord,
        *,
        extra_filters: dict[str, object] | None = None,
    ) -> list[TaskRecord]:
        records = _task_records_for_views()
        filtered = [task for task in records if _matches_task_view(task, view, extra_filters)]
        return _sort_tasks_for_view(filtered, view)

    def _support_requests_for_tasks(tasks: list[TaskRecord]) -> list[SupportRequestRecord]:
        if not get_cloud_token():
            return []
        request_ids: set[str] = set()
        records: list[SupportRequestRecord] = []
        for task in tasks:
            if not task.id:
                continue
            response = cloud_request("GET", f"/api/v1/support-requests?taskId={quote(task.id)}")
            if not isinstance(response, list):
                continue
            for item in response:
                if not isinstance(item, dict):
                    continue
                record = SupportRequestRecord(**item)
                if record.id in request_ids:
                    continue
                request_ids.add(record.id)
                records.append(record)
        records.sort(key=lambda item: item.createdAt, reverse=True)
        return records

    def _drill_target_response_for_event_line(
        target: ReviewDashboardCardTargetRecord,
    ) -> ReviewDashboardDrillTargetResponse:
        event_line_id = _normalize_event_line_reference(target.targetId)
        local_row = state.db.fetchone("SELECT * FROM event_lines WHERE id = ?", (event_line_id,))
        detail: EventLineDetailRecord | None = None
        if get_cloud_token():
            response = cloud_request("GET", f"/api/v1/event-lines/{event_line_id}")
            if isinstance(response, dict):
                detail = build_cloud_event_line_detail(response)
        if detail is None:
            if not local_row:
                raise HTTPException(status_code=404, detail="Event line not found")
            detail = build_event_line_detail(local_row)
        memory_response = get_event_line_memory_response(state.db, event_line_id) if local_row else None
        attachments = _attachments_for_tasks(detail.tasks, cloud=bool(get_cloud_token()))
        meetings = _meeting_summaries_for_tasks(detail.tasks)
        support_requests = _support_requests_for_tasks(detail.tasks)
        return ReviewDashboardDrillTargetResponse(
            target=target,
            eventLineDetail=detail,
            eventLineMemory=memory_response.eventLineMemorySnapshot if memory_response else None,
            tasks=detail.tasks,
            meetings=meetings,
            supportRequests=support_requests,
            attachments=attachments,
        )

    def _drill_target_response_for_task_view(
        target: ReviewDashboardCardTargetRecord,
    ) -> ReviewDashboardDrillTargetResponse:
        view = _load_task_view_definition(target.targetId)
        if view is None:
            view = _build_ad_hoc_task_view(
                target_id=target.targetId,
                target_label=target.targetLabel,
                target_filters=target.targetFilters,
            )
        tasks = _tasks_for_task_view(view, extra_filters=target.targetFilters)
        attachments = _attachments_for_tasks(tasks, cloud=bool(get_cloud_token()))
        meetings = _meeting_summaries_for_tasks(tasks)
        support_requests = _support_requests_for_tasks(tasks)
        return ReviewDashboardDrillTargetResponse(
            target=target,
            tasks=tasks,
            meetings=meetings,
            supportRequests=support_requests,
            attachments=attachments,
        )

    def _drill_target_response_for_meeting(
        target: ReviewDashboardCardTargetRecord,
    ) -> ReviewDashboardDrillTargetResponse:
        meeting = _meeting_summary_for_id(target.targetId)
        tasks = [
            task
            for task in _task_records_for_views()
            if task.sourceId == meeting.id and task.sourceType in {"meeting", "meeting_publish", "meeting_action_item_publish"}
        ]
        attachments = _attachments_for_tasks(tasks, cloud=bool(get_cloud_token()))
        support_requests = _support_requests_for_tasks(tasks)
        return ReviewDashboardDrillTargetResponse(
            target=target,
            tasks=tasks,
            meetings=[meeting],
            supportRequests=support_requests,
            attachments=attachments,
        )

    def _drill_target_response_for_support_request(
        target: ReviewDashboardCardTargetRecord,
    ) -> ReviewDashboardDrillTargetResponse:
        request_record = _support_request_by_id(target.targetId)
        tasks = _task_records_for_ids([request_record.taskId] if request_record.taskId else [])
        attachments = _attachments_for_tasks(tasks, cloud=bool(get_cloud_token()))
        meetings = _meeting_summaries_for_tasks(tasks)
        return ReviewDashboardDrillTargetResponse(
            target=target,
            tasks=tasks,
            meetings=meetings,
            supportRequests=[request_record],
            attachments=attachments,
        )

    def _drill_target_response_for_attachment_group(
        target: ReviewDashboardCardTargetRecord,
    ) -> ReviewDashboardDrillTargetResponse:
        attachment_ids = [
            str(item)
            for item in target.targetFilters.get("attachmentIds", [])
            if isinstance(item, str) and item.strip()
        ]
        task_ids = [
            str(item)
            for item in target.targetFilters.get("taskIds", [])
            if isinstance(item, str) and item.strip()
        ]
        event_line_id = str(target.targetFilters.get("eventLineId") or "").strip()

        attachments = _attachments_for_ids(attachment_ids)
        if not attachments and task_ids:
            tasks = _task_records_for_ids(task_ids)
            attachments = _attachments_for_tasks(tasks, cloud=bool(get_cloud_token()))
        elif not attachments and event_line_id:
            attachments = _attachments_for_event_line(event_line_id)

        related_task_ids = task_ids or [attachment.taskId for attachment in attachments if attachment.taskId]
        tasks = _task_records_for_ids(related_task_ids)
        meetings = _meeting_summaries_for_tasks(tasks)
        support_requests = _support_requests_for_tasks(tasks)
        return ReviewDashboardDrillTargetResponse(
            target=target,
            tasks=tasks,
            meetings=meetings,
            supportRequests=support_requests,
            attachments=attachments,
        )

    def normalize_local_task_tags(tag_ids: list[str] | None, legacy_names: list[str] | None) -> list[TaskTagRecord]:
        _ = tag_ids, legacy_names
        return []

    def sync_local_tasks_for_tag_change(tag_id: str) -> None:
        tag_row = state.db.fetchone("SELECT * FROM task_tags WHERE id = ?", (tag_id,))
        task_rows = state.db.fetchall("SELECT id, tag_ids_json FROM tasks")
        for row in task_rows:
            tag_ids = _parse_json_list(row["tag_ids_json"])
            if tag_id not in tag_ids:
                continue
            if tag_row:
                resolved_tags = [
                    _local_task_tag_record(item)
                    for item in _local_tag_rows_by_ids(state.db, tag_ids)
                ]
            else:
                next_tag_ids = [item for item in tag_ids if item != tag_id]
                resolved_tags = [
                    _local_task_tag_record(item)
                    for item in _local_tag_rows_by_ids(state.db, next_tag_ids)
                ]
            state.db.execute(
                "UPDATE tasks SET tag_ids_json = ?, tags_json = ?, updated_at = ? WHERE id = ?",
                (to_json([item.id for item in resolved_tags]), to_json([item.name for item in resolved_tags]), now_iso(), str(row["id"])),
            )

    def build_client_summary(client_id: str):
        row = state.db.fetchone("SELECT * FROM clients WHERE id = ?", (client_id,))
        if not row:
            raise HTTPException(status_code=404, detail="Client not found")
        return ClientSummary(
            id=str(row["id"]),
            name=str(row["name"]),
            alias=str(row["alias"]),
            domain=str(row["domain"]),
            type=str(row["type"]),
            intro=str(row["intro"]),
            stage=str(row["stage"]),
            folderCount=len(HUMAN_VISIBLE_CATEGORIES),
            documentCount=state.db.scalar("SELECT COUNT(1) AS count FROM documents WHERE client_id = ?", (client_id,)),
            taskCount=state.db.scalar(
                """
                SELECT COUNT(1) AS count
                FROM tasks
                WHERE source_id = ?
                   OR source_id IN (SELECT id FROM meetings WHERE client_id = ?)
                """,
                (client_id, client_id),
            ),
            lastActivityAt=str(row["updated_at"]),
        )

    def build_task_project_context(
        client_id: str | None,
        source_type: str | None = None,
        source_id: str | None = None,
        task_title: str = "",
        task_desc: str = "",
        project_module_id: str | None = None,
        project_flow_id: str | None = None,
    ) -> TaskProjectContextRecord | None:
        if not client_id:
            return None
        client_row = state.db.fetchone("SELECT * FROM clients WHERE id = ?", (client_id,))
        if not client_row:
            return None
        client_name = str(client_row["name"])
        stage = str(client_row["stage"]) if client_row["stage"] else None
        intro = str(client_row["intro"] or "").strip()
        dna_modules = {module.moduleKey: module for module in list_client_dna_modules(client_id)}
        business_module = dna_modules.get("business_intro")
        organization_module = dna_modules.get("organization_intro")
        team_module = dna_modules.get("team_intro")
        market_module = dna_modules.get("market_intro")
        project_module, project_flow = resolve_project_structure_refs(client_id, project_module_id, project_flow_id, strict=False)
        goal_rows = state.db.fetchall(
            "SELECT title, quarter, progress FROM goal_records WHERE client_id = ? ORDER BY updated_at DESC LIMIT 3",
            (client_id,),
        )
        goals = [str(row["title"]).strip() for row in goal_rows if str(row["title"]).strip()]
        meeting_rows = state.db.fetchall(
            "SELECT title FROM meetings WHERE client_id = ? ORDER BY updated_at DESC LIMIT 2",
            (client_id,),
        )
        meetings = [str(row["title"]).strip() for row in meeting_rows if str(row["title"]).strip()]
        excerpt_row = state.db.fetchone(
            "SELECT excerpt FROM documents WHERE client_id = ? AND excerpt IS NOT NULL AND TRIM(excerpt) != '' ORDER BY created_at DESC LIMIT 1",
            (client_id,),
        )
        document_excerpt = str(excerpt_row["excerpt"]).strip() if excerpt_row and excerpt_row["excerpt"] else ""
        open_related_task_count = int(state.db.scalar(
            """
            SELECT COUNT(1) AS count
            FROM tasks
            WHERE status != 'done'
              AND (source_id = ? OR source_id IN (SELECT id FROM meetings WHERE client_id = ?))
            """,
            (client_id, client_id),
        ) or 0)

        dna_ready_count = sum(
            1
            for module in [organization_module, business_module, team_module, market_module]
            if module and module.hasDocument and module.normalizedText.strip()
        )
        completeness_points = dna_ready_count
        if goals:
            completeness_points += 1
        if project_module or project_flow:
            completeness_points += 1
        info_completeness = 'high' if completeness_points >= 5 else 'medium' if completeness_points >= 3 else 'low'

        normalized_task_desc = re.sub(r"\s+", " ", task_desc or "").strip()
        task_clauses = [
            item.strip()
            for item in re.split(r"\n|[。；;|｜]", task_desc or "")
            if item and item.strip()
        ]

        def pick_task_clause(patterns: list[str], fallback_to_first: bool = False) -> str:
            for clause in task_clauses:
                if any(pattern in clause for pattern in patterns):
                    return clause.strip()[:120]
            if fallback_to_first and task_clauses:
                return task_clauses[0].strip()[:120]
            return ""

        desc_background = pick_task_clause(["背景", "对象", "关系", "合作", "机构", "联系人"], fallback_to_first=True)
        desc_focus = pick_task_clause(["推进", "目标", "输出", "形成", "确认", "介绍", "合作"], fallback_to_first=True)
        desc_risk = pick_task_clause(["阻塞", "卡", "风险", "待", "未", "缺", "补", "没有"])
        desc_next = pick_task_clause(["下一步", "继续", "先", "补", "推进", "确认"])

        def summarize_module_text(module: ClientDnaModuleRecord | None, fallback_length: int = 220) -> str:
            if not module:
                return ""
            summary_text = (module.summary or "").strip()
            if summary_text:
                return summary_text
            normalized = re.sub(r"\s+", " ", module.normalizedText or "").strip()
            return normalized[:fallback_length]

        background_parts = [
            desc_background,
            summarize_module_text(business_module, 240),
            summarize_module_text(organization_module, 180),
            summarize_module_text(team_module, 160),
            summarize_module_text(market_module, 160),
        ]
        background_summary = "；".join(part for part in background_parts if part)[:160]
        if not background_summary:
            background_summary = intro or document_excerpt or f"{client_name}当前处于{stage or '推进中'}，建议继续补充项目背景和关键资料。"

        module_summary = None
        if project_module:
            module_summary = "；".join(part for part in [project_module.goal.strip(), project_module.description.strip()] if part)[:140] or None
        flow_summary = None
        if project_flow:
            flow_summary = "；".join(
                part for part in [
                    project_flow.scenario.strip(),
                    " / ".join(project_flow.steps[:3]).strip(),
                    " / ".join(project_flow.riskPoints[:2]).strip(),
                ] if part
            )[:140] or None

        if project_module and project_module.goal.strip():
            goal_summary = project_module.goal.strip()[:120]
        elif goals:
            goal_summary = '；'.join(goals[:2])[:120]
        elif desc_focus:
            goal_summary = desc_focus[:120]
        elif business_module and summarize_module_text(business_module, 180):
            goal_summary = summarize_module_text(business_module, 180)[:120]
        else:
            goal_summary = '当前还没有写入明确的项目目标。'

        if project_flow and project_flow.riskPoints:
            risk_summary = f"当前流程风险：{'；'.join(project_flow.riskPoints[:2])}"[:120]
        elif desc_risk:
            risk_summary = desc_risk[:120]
        elif market_module and summarize_module_text(market_module, 180):
            risk_summary = f"外部环境提示：{summarize_module_text(market_module, 180)}"[:120]
        elif info_completeness == 'low':
            risk_summary = '当前项目背景资料仍偏薄，建议补齐组织、项目、团队和市场四张资料卡后再做更深判断。'
        elif open_related_task_count > 0:
            risk_summary = f'当前仍有{open_related_task_count}条关联任务待推进，需持续跟进项目节奏。'
        elif meetings:
            risk_summary = f"最近会议聚焦于：{' / '.join(meetings[:2])}。"[:120]
        else:
            risk_summary = '当前暂无明显风险信号，可围绕既定目标继续推进。'

        if project_module and project_module.name.strip():
            current_focus = (
                f"{project_module.name.strip()}：{(project_module.goal or '').strip()}"
                if (project_module.goal or "").strip()
                else project_module.name.strip()
            )[:120]
        elif desc_focus:
            current_focus = desc_focus[:120]
        elif goals:
            current_focus = f"当前主要在推进：{goals[0]}"[:120]
        elif meetings:
            current_focus = f"当前讨论集中在：{meetings[0]}"[:120]
        else:
            current_focus = f"{client_name}当前重点仍待补充，建议先明确这一阶段的核心事项。"[:120]

        if project_flow and project_flow.riskPoints:
            current_blocker = f"当前阻塞：{'；'.join(project_flow.riskPoints[:2])}"[:120]
        elif desc_risk:
            current_blocker = desc_risk[:120]
        elif info_completeness == 'low':
            current_blocker = '当前阻塞更像资料不足，项目背景、目标和流程线索都还不完整。'
        elif open_related_task_count > 0:
            current_blocker = f"当前仍有 {open_related_task_count} 条关联任务未收束，推进节奏容易被拉长。"[:120]
        else:
            current_blocker = '当前没有特别突出的阻塞，但仍需盯住推进收束。'

        if project_flow and project_flow.steps:
            next_action = f"下一步动作：{project_flow.steps[0]}"[:120]
        elif desc_next:
            next_action = desc_next[:120]
        elif project_module and project_module.description.strip():
            next_action = f"下一步动作：围绕{project_module.name.strip()}继续细化并推进落地。"[:120]
        elif goals:
            next_action = f"下一步动作：继续围绕“{goals[0]}”推进具体动作。"[:120]
        elif meetings:
            next_action = f"下一步动作：根据最近会议“{meetings[0]}”形成明确后续安排。"[:120]
        else:
            next_action = '下一步动作：先补齐项目背景，再明确这一阶段最核心的推进事项。'

        if meetings:
            recent_progress = f"最近进展：{' / '.join(meetings[:2])}"[:120]
        elif normalized_task_desc:
            recent_progress = f"任务补充：{desc_focus or desc_background or normalized_task_desc}"[:120]
        elif open_related_task_count > 0:
            recent_progress = f"最近进展：已有 {open_related_task_count} 条关联任务在持续推进。"[:120]
        elif document_excerpt:
            recent_progress = f"最近线索：{document_excerpt}"[:120]
        else:
            recent_progress = f"最近进展仍待补充，建议尽快沉淀会议或推进记录。"[:120]

        source_evidence: list[str] = []
        if source_type == 'meeting':
            source_evidence.append('会议来源')
        elif source_type == 'goal':
            source_evidence.append('目标来源')
        elif source_id == client_id:
            source_evidence.append('客户工作台来源')
        else:
            source_evidence.append('任务关联客户')
        if normalized_task_desc:
            source_evidence.append('任务描述补充')
        for module in [organization_module, business_module, team_module, market_module]:
            if module and module.hasDocument and module.normalizedText.strip():
                source_evidence.append(module.title)
        if goals:
            source_evidence.append('项目目标')
        if project_module:
            source_evidence.append(f"任务模块：{project_module.name}")
        if project_flow:
            source_evidence.append(f"流程：{project_flow.name}")
        if intro or document_excerpt:
            source_evidence.append('历史项目摘要')

        return TaskProjectContextRecord(
            clientId=client_id,
            clientName=client_name,
            stage=stage,
            projectModuleId=project_module.id if project_module else None,
            projectModuleName=project_module.name if project_module else None,
            projectModuleSummary=module_summary,
            projectFlowId=project_flow.id if project_flow else None,
            projectFlowName=project_flow.name if project_flow else None,
            projectFlowSummary=flow_summary,
            backgroundSummary=background_summary[:160],
            goalSummary=goal_summary[:120],
            riskSummary=risk_summary[:120],
            currentFocus=current_focus[:120],
            currentBlocker=current_blocker[:120],
            nextAction=next_action[:120],
            recentProgress=recent_progress[:120],
            infoCompleteness=info_completeness,
            sourceEvidence=source_evidence,
        )

    def ensure_standard_client_folders(client_id: str) -> None:
        folders = ensure_client_workspace(state.data_dir, client_id)
        timestamp = now_iso()
        for label, path in folders.items():
            existing = state.db.fetchone(
                "SELECT id FROM client_folders WHERE client_id = ? AND label = ?",
                (client_id, label),
            )
            file_count = int(
                state.db.scalar(
                    """
                    SELECT COUNT(1) AS count
                    FROM knowledge_documents
                    WHERE client_id = ? AND human_folder_category = ?
                    """,
                    (client_id, label),
                )
            )
            if file_count > 0:
                unhide_client_folder_label(client_id, label)
            if existing:
                state.db.execute(
                    "UPDATE client_folders SET path = ?, file_count = ?, last_scanned_at = ? WHERE id = ?",
                    (str(path), file_count, timestamp, str(existing["id"])),
                )
            else:
                state.db.execute(
                    """
                    INSERT INTO client_folders(id, client_id, label, path, file_count, last_scanned_at, created_at)
                    VALUES(?, ?, ?, ?, ?, ?, ?)
                    """,
                    (new_id("fld"), client_id, label, str(path), file_count, timestamp, timestamp),
                    )

    def build_document_card_record(payload: dict[str, object]) -> DocumentCardRecord:
        return DocumentCardRecord(**payload)

    def build_knowledge_status_record(client_id: str) -> KnowledgeStatusRecord:
        return KnowledgeStatusRecord(**compute_knowledge_status(state.db, client_id, state.data_dir))

    def build_retrieval_bundle(client_id: str, prompt: str):
        hint_terms = build_client_dna_retrieval_hint(client_id, prompt)
        retrieval_prompt = prompt if not hint_terms else f"{prompt}\n背景关注：{' '.join(hint_terms)}"
        bundle = retrieve_knowledge_bundle(state.db, state.data_dir, client_id, retrieval_prompt)
        retrieval_summary = bundle.retrieval_summary if isinstance(bundle.retrieval_summary, dict) else {}
        bundle.retrieval_summary = {
            **retrieval_summary,
            "sourcePrompt": prompt,
            "clientDnaHintTerms": hint_terms,
        }
        return bundle

    def persist_retrieval_bundle(client_id: str, prompt: str, thread_id: str | None, bundle, retrieval_elapsed_ms: float) -> str:
        search_id = new_id("ks")
        timestamp = now_iso()
        payload = serialize_retrieval_bundle(bundle)
        payload["timing"] = {"retrievalMs": retrieval_elapsed_ms}
        normalized_thread_id: str | None = None
        if thread_id:
            existing_thread = state.db.fetchone(
                "SELECT id FROM chat_threads WHERE id = ? AND client_id = ?",
                (thread_id, client_id),
            )
            if existing_thread:
                normalized_thread_id = str(existing_thread["id"])
        state.db.execute(
            """
            INSERT INTO knowledge_search_runs(id, client_id, thread_id, prompt, status, retrieval_json, created_at, updated_at)
            VALUES(?, ?, ?, ?, 'ready', ?, ?, ?)
            """,
            (search_id, client_id, normalized_thread_id, prompt, to_json(payload), timestamp, timestamp),
        )
        return search_id

    def load_cached_retrieval_bundle(client_id: str, search_id: str, prompt: str):
        row = state.db.fetchone(
            """
            SELECT *
            FROM knowledge_search_runs
            WHERE id = ? AND client_id = ? AND status = 'ready'
            """,
            (search_id, client_id),
        )
        if not row:
            return None, 0.0
        if str(row["prompt"]).strip() != prompt.strip():
            return None, 0.0
        payload = from_json(str(row["retrieval_json"]), {})
        if not isinstance(payload, dict):
            return None, 0.0
        timing = payload.get("timing", {})
        retrieval_elapsed_ms = float(timing.get("retrievalMs", 0.0) or 0.0) if isinstance(timing, dict) else 0.0
        return deserialize_retrieval_bundle(payload), retrieval_elapsed_ms

    def fetch_chat_message_for_client(client_id: str, message_id: str) -> ChatMessageRecord:
        row = state.db.fetchone(
            """
            SELECT m.*
            FROM chat_messages m
            JOIN chat_threads t ON t.id = m.thread_id
            WHERE m.id = ? AND t.client_id = ?
            """,
            (message_id, client_id),
        )
        if not row:
            raise HTTPException(status_code=404, detail="Chat message not found")
        return build_chat_message(row)

    def build_analysis_evidence_summary(client_id: str, prompt: str, retrieval_bundle) -> ClientAnalysisEvidenceSummaryRecord:
        retrieval_meta = retrieval_bundle.retrieval_summary if isinstance(retrieval_bundle.retrieval_summary, dict) else {}
        evidence = [
            EvidenceItem(
                id=new_id("ev"),
                title=item.title,
                excerpt=item.excerpt,
                sourceType="knowledge_chunk" if item.chunk_id else "knowledge_document",
                documentId=item.knowledge_document_id,
                path=item.path,
                score=item.score,
                coverage=item.coverage,
                sectionLabel=item.section_label,
                retrievalStage={"master_index": "master_index", "surrogate": "surrogate"}.get(item.source_stage, "raw_chunk"),
                isFallback=item.source_stage == "master_index",
                matchedTerms=item.matched_terms,
            )
            for item in retrieval_bundle.citations
            if item.source_stage == "raw_chunk"
        ]
        preferred_categories = [
            str(item)
            for item in retrieval_meta.get("preferredCategories", [])
            if str(item).strip()
        ] if isinstance(retrieval_meta.get("preferredCategories"), list) else []
        curated_hits = select_high_signal_evidence(
            evidence,
            limit=8,
            prompt=prompt,
            preferred_categories=preferred_categories,
        )
        hits = [
            KnowledgeSearchHitRecord(
                title=item.title,
                excerpt=item.excerpt,
                score=item.score,
                stage=item.retrievalStage,  # type: ignore[arg-type]
                path=item.path,
                sectionLabel=item.sectionLabel,
                matchedTerms=item.matchedTerms,
            )
            for item in curated_hits
        ]
        preview_summary = str(retrieval_meta.get("previewSummary") or "").strip()
        if not preview_summary:
            preview_summary = build_retrieval_preview_summary(client_id, prompt, evidence, retrieval_bundle)
        covered_categories = [
            str(item)
            for item in retrieval_meta.get("categoryCoverage", [])
            if str(item).strip()
        ] if isinstance(retrieval_meta.get("categoryCoverage"), list) else []
        missing_categories = [item for item in preferred_categories if item not in covered_categories]
        return ClientAnalysisEvidenceSummaryRecord(
            summaryText=preview_summary,
            masterHitCount=int(retrieval_meta.get("masterHitCount", 0) or 0),
            surrogateHitCount=int(retrieval_meta.get("surrogateHitCount", 0) or 0),
            rawChunkHitCount=int(retrieval_meta.get("rawChunkHitCount", 0) or 0),
            drillthroughUsed=bool(retrieval_meta.get("drillthroughUsed", False)),
            coveredCategories=covered_categories,
            missingCategories=missing_categories,
            evidenceList=hits,
        )

    def create_client_analysis_run(
        client_id: str,
        thread_id: str,
        user_message_id: str,
        assistant_message_id: str,
        question: str,
        created_at: str,
    ) -> ClientAnalysisRunRecord:
        run_id = new_id("analysis")
        state.db.execute(
            """
            INSERT INTO client_analysis_runs(
                id, client_id, thread_id, user_message_id, assistant_message_id, question,
                status, phase, progress, progress_floor, progress_ceiling, stage_label, elapsed_ms,
                evidence_summary_json, long_answer, structured_summary_json, long_answer_status, summary_status,
                answer_mode, llm_invoked, provider_used, failure_reason, timing_json, created_at, updated_at
            )
            VALUES(?, ?, ?, ?, ?, ?, 'queued', 'queued', 0, 0, 30, '等待开始整理背景材料', 0, '{}', NULL, '{}', 'pending', 'pending', NULL, 0, NULL, NULL, '{}', ?, ?)
            """,
            (run_id, client_id, thread_id, user_message_id, assistant_message_id, question, created_at, created_at),
        )
        row = state.db.fetchone("SELECT * FROM client_analysis_runs WHERE id = ?", (run_id,))
        assert row is not None
        return build_client_analysis_run(row)

    _ANALYSIS_FIELD_UNSET = object()

    def update_client_analysis_run(
        run_id: str,
        *,
        status: str | None = None,
        phase: str | None = None,
        progress: float | None = None,
        progress_floor: float | None = None,
        progress_ceiling: float | None = None,
        stage_label: str | None = None,
        elapsed_ms: float | None = None,
        evidence_summary: ClientAnalysisEvidenceSummaryRecord | dict[str, object] | object = _ANALYSIS_FIELD_UNSET,
        long_answer: str | None | object = _ANALYSIS_FIELD_UNSET,
        structured_summary: AiStructuredResponse | dict[str, object] | None | object = _ANALYSIS_FIELD_UNSET,
        long_answer_status: str | None = None,
        summary_status: str | None = None,
        answer_mode: str | None = None,
        llm_invoked: bool | None = None,
        provider_used: str | None = None,
        failure_reason: str | None = None,
        timing: dict[str, float] | None = None,
    ) -> None:
        row = state.db.fetchone("SELECT * FROM client_analysis_runs WHERE id = ?", (run_id,))
        if not row:
            return
        next_status = status or str(row["status"])
        next_phase = phase or str(row["phase"])
        next_progress = float(progress if progress is not None else row["progress"])
        next_floor = float(progress_floor if progress_floor is not None else row["progress_floor"])
        next_ceiling = float(progress_ceiling if progress_ceiling is not None else row["progress_ceiling"])
        next_stage_label = stage_label if stage_label is not None else (str(row["stage_label"]) if row["stage_label"] else None)
        next_elapsed = float(elapsed_ms if elapsed_ms is not None else row["elapsed_ms"])
        next_evidence_summary = (
            evidence_summary.model_dump()
            if isinstance(evidence_summary, ClientAnalysisEvidenceSummaryRecord)
            else evidence_summary
        )
        if next_evidence_summary is _ANALYSIS_FIELD_UNSET:
            stored_summary = from_json(str(row["evidence_summary_json"] or "{}"), {})
            next_evidence_summary = stored_summary if isinstance(stored_summary, dict) else {}
        next_structured_summary = (
            structured_summary.model_dump()
            if isinstance(structured_summary, AiStructuredResponse)
            else structured_summary
        )
        if next_structured_summary is _ANALYSIS_FIELD_UNSET:
            stored_structured = from_json(str(row["structured_summary_json"] or "{}"), {})
            next_structured_summary = stored_structured if isinstance(stored_structured, dict) else {}
        elif next_structured_summary is None:
            next_structured_summary = {}
        next_timing = timing
        if next_timing is None:
            stored_timing = from_json(str(row["timing_json"] or "{}"), {})
            next_timing = stored_timing if isinstance(stored_timing, dict) else {}
        next_long_answer = str(row["long_answer"]) if row["long_answer"] is not None else None
        if long_answer is not _ANALYSIS_FIELD_UNSET:
            next_long_answer = str(long_answer) if isinstance(long_answer, str) else None
        state.db.execute(
            """
            UPDATE client_analysis_runs
            SET status = ?, phase = ?, progress = ?, progress_floor = ?, progress_ceiling = ?, stage_label = ?,
                elapsed_ms = ?, evidence_summary_json = ?, long_answer = ?,
                structured_summary_json = ?, long_answer_status = ?, summary_status = ?, answer_mode = ?,
                llm_invoked = ?, provider_used = ?, failure_reason = ?, timing_json = ?, updated_at = ?
            WHERE id = ?
            """,
            (
                next_status,
                next_phase,
                next_progress,
                next_floor,
                next_ceiling,
                next_stage_label,
                next_elapsed,
                to_json(next_evidence_summary),
                next_long_answer,
                to_json(next_structured_summary),
                long_answer_status or str(row["long_answer_status"]),
                summary_status or str(row["summary_status"]),
                answer_mode if answer_mode is not None else (str(row["answer_mode"]) if row["answer_mode"] else None),
                1 if (llm_invoked if llm_invoked is not None else bool(row["llm_invoked"])) else 0,
                provider_used if provider_used is not None else (str(row["provider_used"]) if row["provider_used"] else None),
                failure_reason if failure_reason is not None else (str(row["failure_reason"]) if row["failure_reason"] else None),
                to_json(next_timing),
                now_iso(),
                run_id,
            ),
        )

    def build_client_analysis_run(row) -> ClientAnalysisRunRecord:
        evidence_summary_data = from_json(row["evidence_summary_json"], {})
        structured_summary_data = from_json(row["structured_summary_json"], {})
        timing_data = from_json(row["timing_json"], {})
        assistant_message = None
        if row["assistant_message_id"]:
            assistant_row = state.db.fetchone("SELECT * FROM chat_messages WHERE id = ?", (str(row["assistant_message_id"]),))
            if assistant_row:
                assistant_message = build_chat_message(assistant_row)
        return ClientAnalysisRunRecord(
            id=str(row["id"]),
            clientId=str(row["client_id"]),
            threadId=str(row["thread_id"]),
            userMessageId=str(row["user_message_id"]),
            assistantMessageId=str(row["assistant_message_id"]),
            question=str(row["question"]),
            status=str(row["status"]),  # type: ignore[arg-type]
            phase=str(row["phase"]),  # type: ignore[arg-type]
            progress=float(row["progress"] or 0.0),
            progressFloor=float(row["progress_floor"] or 0.0),
            progressCeiling=float(row["progress_ceiling"] or 0.0),
            stageLabel=str(row["stage_label"]) if row["stage_label"] else None,
            elapsedMs=float(row["elapsed_ms"] or 0.0),
            evidenceSummary=ClientAnalysisEvidenceSummaryRecord(**(evidence_summary_data if isinstance(evidence_summary_data, dict) else {})),
            longAnswerStatus=str(row["long_answer_status"]),  # type: ignore[arg-type]
            summaryStatus=str(row["summary_status"]),  # type: ignore[arg-type]
            longAnswer=str(row["long_answer"]) if row["long_answer"] else None,
            structuredSummary=AiStructuredResponse(**structured_summary_data) if isinstance(structured_summary_data, dict) and structured_summary_data else None,
            answerMode=str(row["answer_mode"]) if row["answer_mode"] else None,  # type: ignore[arg-type]
            llmInvoked=bool(row["llm_invoked"]),
            providerUsed=str(row["provider_used"]) if row["provider_used"] else None,
            failureReason=str(row["failure_reason"]) if row["failure_reason"] else None,
            timing=timing_data if isinstance(timing_data, dict) else {},
            assistantMessage=assistant_message,
            createdAt=str(row["created_at"]),
            updatedAt=str(row["updated_at"]),
        )

    def fetch_analysis_run_for_client(client_id: str, run_id: str) -> ClientAnalysisRunRecord:
        row = state.db.fetchone("SELECT * FROM client_analysis_runs WHERE id = ? AND client_id = ?", (run_id, client_id))
        if not row:
            raise HTTPException(status_code=404, detail="Analysis run not found")
        return build_client_analysis_run(row)

    def is_client_analysis_run_canceled(run_id: str | None) -> bool:
        if not run_id:
            return False
        row = state.db.fetchone("SELECT status, failure_reason FROM client_analysis_runs WHERE id = ?", (run_id,))
        if not row:
            return False
        return str(row["status"] or "") == "canceled" or str(row["failure_reason"] or "") == "user_canceled"

    def cancel_analysis_run_for_client(client_id: str, run_id: str) -> ClientAnalysisRunRecord:
        run = fetch_analysis_run_for_client(client_id, run_id)
        if run.status in {"completed", "failed", "canceled"}:
            return run
        timestamp = now_iso()
        structured = AiStructuredResponse(
            content="本次回答已停止。",
            judgment="你手动停止了这次生成，因此没有继续产出正式回答。",
            analysis="当前计算已从前台中断，系统不会再继续展示这次回答结果。",
            actions="如果需要，可以直接重新提问，或换一个更明确的问题再次生成。",
            timeline=f"停止时间：{timestamp}",
        )
        state.db.execute(
            """
            UPDATE chat_messages
            SET content = ?, structured_data_json = ?, model_route = ?, llm_invoked = 0, provider_used = NULL,
                answer_mode = NULL, evidence_status = 'none', failure_reason = 'user_canceled',
                timing_json = COALESCE(timing_json, '{}'), retrieval_summary_json = COALESCE(retrieval_summary_json, '{}'),
                evidence_json = '[]', status = 'success', created_at = ?
            WHERE id = ?
            """,
            (
                structured.content,
                to_json(structured.model_dump()),
                "已手动停止",
                timestamp,
                run.assistantMessageId,
            ),
        )
        update_client_analysis_run(
            run_id,
            status="canceled",
            phase="canceled",
            progress=100.0,
            progress_floor=100.0,
            progress_ceiling=100.0,
            stage_label="已停止当前回答",
            long_answer=None,
            structured_summary=structured,
            long_answer_status="failed",
            summary_status="failed",
            answer_mode=None,
            provider_used=None,
            failure_reason="user_canceled",
            timing=run.timing,
        )
        state.db.execute("UPDATE chat_threads SET updated_at = ? WHERE id = ?", (timestamp, run.threadId))
        log_activity("chat.cancel", "chat_thread", run.threadId, {"clientId": client_id, "runId": run_id, "prompt": run.question})
        return fetch_analysis_run_for_client(client_id, run_id)

    def build_answer_memory_markdown(client_id: str, message: ChatMessageRecord) -> str:
        client_name = build_client_summary(client_id).name
        lines = [
            f"# {client_name} · 战略陪伴记忆",
            "",
            f"- 客户：{client_name}",
            f"- 生成时间：{now_iso()}",
            "",
            "## 回答摘要",
            message.content.strip() or "（无内容）",
        ]
        if message.structuredData:
            lines.extend(
                [
                    "",
                    "## 核心判断",
                    (message.structuredData.judgment or "").strip() or "【待补充】",
                    "",
                    "## 结构化分析",
                    (message.structuredData.analysis or "").strip() or "【待补充】",
                    "",
                    "## 建议动作",
                    (message.structuredData.actions or "").strip() or "【待补充】",
                    "",
                    "## 关键时间线",
                    (message.structuredData.timeline or "").strip() or "【待补充】",
                ]
            )
        if message.evidence:
            lines.extend(["", "## 证据来源"])
            for item in message.evidence:
                lines.extend(
                    [
                        f"- {item.title} | {item.sectionLabel or item.sourceType or '证据'}",
                        f"  - 摘录：{(item.excerpt or '').strip()}",
                    ]
                )
        return "\n".join(lines).strip() + "\n"

    def register_generated_workspace_document(
        client_id: str,
        *,
        target_path: Path,
        title: str,
        kind: str,
        source: str,
        excerpt: str,
        folder_label: str = "战略陪伴",
    ) -> ClientTextDocumentResponse:
        ensure_standard_client_folders(client_id)
        folder_row = state.db.fetchone(
            "SELECT id FROM client_folders WHERE client_id = ? AND label = ?",
            (client_id, folder_label),
        )
        document_id = new_id("doc")
        timestamp_iso = now_iso()
        normalized_excerpt = (excerpt or "").strip()[:140] or f"{title} 已加入当前项目文档库。"
        state.db.execute(
            """
            INSERT INTO documents(id, client_id, folder_id, title, path, original_source_path, kind, source, excerpt, tags_json, created_at)
            VALUES(?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                document_id,
                client_id,
                str(folder_row["id"]) if folder_row else None,
                target_path.name,
                str(target_path),
                str(target_path),
                kind,
                source,
                normalized_excerpt,
                to_json([source, kind, folder_label]),
                timestamp_iso,
            ),
        )
        ingest_document_knowledge(
            state.db,
            data_dir=state.data_dir,
            client_id=client_id,
            import_id=None,
            document_id=document_id,
            source_path=target_path,
            original_source_path=target_path,
            title=target_path.name,
            kind=kind,
            source=source,
            fallback_excerpt=normalized_excerpt,
            created_at=timestamp_iso,
            ai_service=None,
        )
        document_row = state.db.fetchone("SELECT path FROM documents WHERE id = ?", (document_id,))
        resolved_path = str(document_row["path"]) if document_row and document_row["path"] else str(target_path)
        return ClientTextDocumentResponse(
            clientId=client_id,
            documentId=document_id,
            title=title,
            fileName=Path(resolved_path).name,
            path=resolved_path,
        )

    def create_answer_memory_markdown_document(client_id: str, message: ChatMessageRecord) -> ClientTextDocumentResponse:
        folders = ensure_client_workspace(state.data_dir, client_id)
        target_dir = folders["战略陪伴"]
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        target_path = target_dir / f"{timestamp}_{safe_filename(message.content[:18] or '战略陪伴记忆')}.md"
        target_path.write_text(build_answer_memory_markdown(client_id, message), encoding="utf-8")
        return register_generated_workspace_document(
            client_id,
            target_path=target_path,
            title=f"{build_client_summary(client_id).name} · 战略陪伴记忆",
            kind="md",
            source="answer_memory_doc",
            excerpt=message.content,
        )

    def export_answer_to_docx(client_id: str, message: ChatMessageRecord) -> Path:
        folders = ensure_client_workspace(state.data_dir, client_id)
        target_dir = folders["战略陪伴"]
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        target_path = target_dir / f"{timestamp}_{safe_filename(message.content[:18] or '战略陪伴沉淀')}.docx"
        document = WordDocument()
        document.add_heading("战略陪伴沉淀", level=1)
        document.add_paragraph(f"客户：{build_client_summary(client_id).name}")
        document.add_paragraph(f"生成时间：{now_iso()}")
        document.add_paragraph("")
        document.add_heading("回答摘要", level=2)
        document.add_paragraph(message.content)
        if message.structuredData:
            document.add_heading("核心判断", level=2)
            document.add_paragraph(message.structuredData.judgment)
            document.add_heading("结构化分析", level=2)
            document.add_paragraph(message.structuredData.analysis)
            document.add_heading("建议动作", level=2)
            document.add_paragraph(message.structuredData.actions)
            document.add_heading("关键时间线", level=2)
            document.add_paragraph(message.structuredData.timeline)
        if message.evidence:
            document.add_heading("证据来源", level=2)
            for item in message.evidence:
                paragraph = document.add_paragraph(style="List Bullet")
                paragraph.add_run(item.title).bold = True
                paragraph.add_run(f" | {item.sectionLabel or item.sourceType or '证据'}")
                document.add_paragraph(item.excerpt)
        document.save(target_path)
        return target_path

    def create_answer_export_document(client_id: str, message: ChatMessageRecord) -> ClientTextDocumentResponse:
        target_path = export_answer_to_docx(client_id, message)
        return register_generated_workspace_document(
            client_id,
            target_path=target_path,
            title="战略陪伴沉淀",
            kind="docx",
            source="answer_export_doc",
            excerpt=message.content,
        )

    def build_consultation_knowledge_title(request: ConsultationKnowledgeRequestRecord) -> str:
        for candidate in (request.question.strip(), request.answer.strip()):
            if not candidate:
                continue
            normalized = re.sub(r"\s+", " ", candidate)
            return normalized[:28]
        return "咨询沉淀"

    def build_consultation_context_lines(request: ConsultationKnowledgeRequestRecord) -> list[str]:
        lines: list[str] = []
        if request.clientName or request.clientId:
            lines.append(f"- 客户：{request.clientName or request.clientId}")
        if request.taskId:
            lines.append(f"- 关联任务：{request.taskId}")
        if request.eventLineId:
            lines.append(f"- 关联事件线：{request.eventLineId}")
        if request.requestedByName or request.requestedByUserId:
            lines.append(f"- 发起人：{request.requestedByName or request.requestedByUserId}")
        lines.append("- 来源：手机咨询助手")
        return lines

    def build_consultation_memory_markdown(request: ConsultationKnowledgeRequestRecord) -> str:
        lines = [
            "# 手机咨询沉淀记忆",
            f"生成时间：{now_iso()}",
            "",
            "## 关联上下文",
            *build_consultation_context_lines(request),
            "",
        ]
        if request.question.strip():
            lines.extend(["## 原始问题", request.question.strip(), ""])
        lines.extend(["## 答案内容", request.answer.strip(), ""])
        return "\n".join(lines).strip() + "\n"

    def build_consultation_archive_content(request: ConsultationKnowledgeRequestRecord) -> str:
        lines = ["## 关联上下文", *build_consultation_context_lines(request), ""]
        if request.question.strip():
            lines.extend(["## 原始问题", request.question.strip(), ""])
        lines.extend(["## 答案内容", request.answer.strip()])
        return "\n".join(lines).strip()

    def create_consultation_memory_document(
        client_id: str,
        request: ConsultationKnowledgeRequestRecord,
    ) -> ClientTextDocumentResponse:
        client = build_client_summary(client_id)
        folders = ensure_client_workspace(state.data_dir, client_id)
        target_dir = folders["战略陪伴"]
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_stem = safe_filename(build_consultation_knowledge_title(request)[:18] or "咨询沉淀记忆")
        target_path = target_dir / f"{timestamp}_{safe_stem}.md"
        if target_path.exists():
            target_path = target_dir / f"{timestamp}_{safe_stem}_{uuid4().hex[:6]}.md"
        target_path.write_text(build_consultation_memory_markdown(request), encoding="utf-8")
        return register_generated_workspace_document(
            client_id,
            target_path=target_path,
            title=f"{client.name} · 咨询沉淀记忆",
            kind="md",
            source="consultation_knowledge_memory",
            excerpt=request.answer,
        )

    def sink_consultation_knowledge_request(
        request: ConsultationKnowledgeRequestRecord,
    ) -> ClientTextDocumentResponse:
        client_id = (request.clientId or "").strip()
        if not client_id:
            raise HTTPException(status_code=400, detail="沉淀请求缺少 clientId，无法写入本地综合库")

        client = build_client_summary(client_id)
        title = build_consultation_knowledge_title(request)
        if request.target == "vector_memory":
            source_links = []
            if request.taskId:
                source_links.append(
                    {
                        "title": f"关联任务 {request.taskId}",
                        "documentId": request.taskId,
                        "path": None,
                        "sectionLabel": "任务上下文",
                    }
                )
            if request.eventLineId:
                source_links.append(
                    {
                        "title": f"关联事件线 {request.eventLineId}",
                        "documentId": request.eventLineId,
                        "path": None,
                        "sectionLabel": "事件线上下文",
                    }
                )
            create_memory_surrogate_from_answer(
                state.db,
                data_dir=state.data_dir,
                client_id=client_id,
                title=f"{client.name} · {title}",
                content=request.answer,
                actions="",
                analysis=request.question.strip(),
                source_links=source_links,
                created_at=now_iso(),
                ai_service=state.ai,
            )
            generated = create_consultation_memory_document(client_id, request)
            log_activity(
                "consultation.knowledge.vector_memory",
                "knowledge_memory",
                generated.documentId,
                {
                    "requestId": request.id,
                    "clientId": client_id,
                    "taskId": request.taskId or "",
                    "eventLineId": request.eventLineId or "",
                    "path": generated.path,
                },
            )
            return generated

        generated = create_client_text_document(
            client_id,
            ClientTextDocumentPayload(
                title=f"{client.name} · {title}",
                content=build_consultation_archive_content(request),
            ),
        )
        log_activity(
            "consultation.knowledge.document_archive",
            "document",
            generated.documentId,
            {
                "requestId": request.id,
                "clientId": client_id,
                "taskId": request.taskId or "",
                "eventLineId": request.eventLineId or "",
                "path": generated.path,
            },
        )
        return generated

    def infer_text_document_title(content: str) -> str:
        normalized = re.sub(r"\s+", " ", content).strip()
        if not normalized:
            return "未命名新增文档"
        for raw_line in content.splitlines():
            line = re.sub(r"^[#>*\-\d\.\)\s]+", "", raw_line).strip()
            if len(line) < 4:
                continue
            candidate = re.split(r"[。！？!?]", line, maxsplit=1)[0].strip() or line
            return candidate[:28]
        return normalized[:28]

    def create_client_text_document(client_id: str, payload: ClientTextDocumentPayload) -> ClientTextDocumentResponse:
        client = build_client_summary(client_id)
        ensure_standard_client_folders(client_id)
        folders = ensure_client_workspace(state.data_dir, client_id)
        target_dir = (folders.get("项目与业务") or next(iter(folders.values()))) / "手动新增文档"
        target_dir.mkdir(parents=True, exist_ok=True)

        normalized_content = str(payload.content or "").replace("\r\n", "\n").strip()
        if not normalized_content:
            raise HTTPException(status_code=400, detail="请先粘贴文档内容。")

        resolved_title = str(payload.title or "").strip() or infer_text_document_title(normalized_content)
        safe_stem = safe_filename(resolved_title or "新增文档")
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        target_path = target_dir / f"{timestamp}_{safe_stem}.docx"
        if target_path.exists():
            target_path = target_dir / f"{timestamp}_{safe_stem}_{uuid4().hex[:6]}.docx"

        document = WordDocument()
        document.add_heading(resolved_title, level=1)
        for block in re.split(r"\n\s*\n+", normalized_content):
            text = block.strip()
            if not text:
                continue
            heading_match = re.match(r"^(#{1,4})\s+(.+)$", text)
            if heading_match:
                level = min(len(heading_match.group(1)) + 1, 4)
                document.add_heading(heading_match.group(2).strip(), level=level)
                continue
            document.add_paragraph(text)
        document.save(target_path)

        timestamp_iso = now_iso()
        folder_row = state.db.fetchone(
            "SELECT id FROM client_folders WHERE client_id = ? AND label = ?",
            (client_id, "项目与业务"),
        )
        document_id = new_id("doc")
        excerpt = normalized_content[:140] or f"{resolved_title} 已进入当前项目文档库。"
        state.db.execute(
            """
            INSERT INTO documents(id, client_id, folder_id, title, path, original_source_path, kind, source, excerpt, tags_json, created_at)
            VALUES(?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                document_id,
                client_id,
                str(folder_row["id"]) if folder_row else None,
                target_path.name,
                str(target_path),
                str(target_path),
                "docx",
                "manual_text_doc",
                excerpt,
                to_json(["manual_text_doc", "docx"]),
                timestamp_iso,
            ),
        )
        ingest_document_knowledge(
            state.db,
            data_dir=state.data_dir,
            client_id=client_id,
            import_id=None,
            document_id=document_id,
            source_path=target_path,
            original_source_path=target_path,
            title=target_path.name,
            kind="docx",
            source="manual_text_doc",
            fallback_excerpt=excerpt,
            created_at=timestamp_iso,
            ai_service=None,
        )
        log_activity(
            "client.document.create_from_text",
            "document",
            document_id,
            {
                "clientId": client_id,
                "clientName": client.name,
                "title": resolved_title,
                "path": str(target_path),
            },
        )
        document_row = state.db.fetchone("SELECT path FROM documents WHERE id = ?", (document_id,))
        resolved_path = str(document_row["path"]) if document_row and document_row["path"] else str(target_path)
        return ClientTextDocumentResponse(
            clientId=client_id,
            documentId=document_id,
            title=resolved_title,
            fileName=Path(resolved_path).name,
            path=resolved_path,
        )

    def build_template_fill_context(
        client_id: str,
        template_name: str,
        field_label: str,
        *,
        field_type: str = "general",
        evidence_limit: int = 12,
        excerpt_limit: int = 2200,
        evidence_char_budget: int = 18000,
        dna_max_chars: int = 2200,
    ) -> tuple[str, list[EvidenceItem], list[TemplateWebSource]]:
        def collect_template_fill_public_hints(max_rows: int = 60) -> tuple[list[str], list[str]]:
            rows = state.db.fetchall(
                """
                SELECT file_name, preview_text
                FROM v2_documents
                WHERE client_id = ?
                  AND COALESCE(parse_status, 'ready') = 'ready'
                ORDER BY
                  CASE
                    WHEN preview_text LIKE '%http%' OR preview_text LIKE '%.org%' OR preview_text LIKE '%.cn%' THEN 0
                    WHEN file_name LIKE '%中国%' OR preview_text LIKE '%中国%' THEN 1
                    WHEN file_name LIKE '%基金会论坛%' OR preview_text LIKE '%基金会论坛%' THEN 1
                    ELSE 2
                  END,
                  updated_at DESC,
                  id DESC
                LIMIT ?
                """,
                (client_id, max_rows),
            )
            titles: list[str] = []
            snippets: list[str] = []
            for row in rows or []:
                try:
                    title = str(row["file_name"] or "").strip()
                except Exception:
                    title = str((row[0] if len(row) > 0 else "") or "").strip()
                try:
                    snippet = str(row["preview_text"] or "").strip()
                except Exception:
                    snippet = str((row[1] if len(row) > 1 else "") or "").strip()
                if title and title not in titles:
                    titles.append(title)
                if snippet and snippet not in snippets:
                    snippets.append(snippet)
            return titles, snippets

        client_summary = build_client_summary(client_id)
        retrieval_bundle = build_retrieval_bundle(
            client_id,
            build_template_fill_retrieval_query(
                client_name=client_summary.name,
                template_name=template_name,
                field_label=field_label,
                field_type=field_type,
            ),
        )
        evidence = [
            EvidenceItem(
                id=new_id("ev"),
                title=item.title,
                excerpt=item.excerpt,
                sourceType="knowledge_chunk" if item.chunk_id else "knowledge_document",
                documentId=item.knowledge_document_id,
                path=item.path,
                score=item.score,
                coverage=item.coverage,
                sectionLabel=item.section_label,
                retrievalStage={"master_index": "master_index", "surrogate": "surrogate"}.get(item.source_stage, "raw_chunk"),
                isFallback=item.source_stage == "master_index",
                matchedTerms=item.matched_terms,
            )
            for item in retrieval_bundle.citations
        ]
        retrieval_meta = retrieval_bundle.retrieval_summary if isinstance(retrieval_bundle.retrieval_summary, dict) else {}
        preferred_categories = [
            str(item)
            for item in retrieval_meta.get("preferredCategories", [])
            if str(item).strip()
        ] if isinstance(retrieval_meta.get("preferredCategories"), list) else []
        curated_evidence = select_high_signal_evidence(
            evidence,
            limit=evidence_limit,
            prompt=field_label,
            preferred_categories=preferred_categories,
        )
        client_name = client_summary.name
        lines = [
            f"当前客户：{client_name}",
            f"模板：{template_name}",
            f"待填写字段：{field_label}",
        ]
        client_dna_context = build_client_dna_context(client_id, field_label, max_chars=dna_max_chars)
        if client_dna_context:
            lines.append(client_dna_context)
        if curated_evidence:
            blocks: list[str] = []
            used_chars = 0
            for index, item in enumerate(curated_evidence, start=1):
                label = item.title
                if item.sectionLabel:
                    label = f"{label} / {item.sectionLabel}"
                compact_excerpt = re.sub(r"\s+", " ", item.excerpt or "").strip()[:excerpt_limit]
                block = (
                    f"[参考证据 {index}]\n"
                    f"标题：{label}\n"
                    f"片段：{compact_excerpt}"
                )
                if blocks and used_chars + len(block) > evidence_char_budget:
                    break
                blocks.append(block)
                used_chars += len(block)
            if blocks:
                lines.append("可参考的客户资料：\n" + "\n\n".join(blocks))
        web_sources: list[TemplateWebSource] = []
        if should_enable_template_fill_web_supplement(field_type, len(curated_evidence), field_label=field_label):
            public_hint_titles, public_hint_snippets = collect_template_fill_public_hints()
            combined_titles = list(
                dict.fromkeys([item.title for item in curated_evidence] + public_hint_titles)
            )
            combined_snippets = list(
                dict.fromkeys([item.excerpt for item in curated_evidence if item.excerpt] + public_hint_snippets)
            )
            web_sources = fetch_template_fill_web_sources(
                client_name=client_name,
                field_label=field_label,
                template_name=template_name,
                client_domain=client_summary.domain,
                evidence_titles=combined_titles,
                evidence_snippets=combined_snippets,
                max_items=2,
                field_type=field_type,
            )
            if web_sources:
                web_blocks = []
                for index, item in enumerate(web_sources, start=1):
                    web_blocks.append(
                        f"[网页补充 {index}]\n"
                        f"标题：{item.title}\n"
                        f"链接：{item.url}\n"
                        f"摘要：{item.snippet}"
                    )
                lines.append("联网补充（公开网页，仅作弱证据）：\n" + "\n\n".join(web_blocks))
        return "\n\n".join(lines).strip(), curated_evidence, web_sources

    def summarize_template_field_basis(value: str, evidence_titles: list[str], web_titles: list[str] | None = None) -> str:
        web_titles = [str(item).strip() for item in (web_titles or []) if str(item).strip()]
        if str(value or "").startswith("【待确认】"):
            if evidence_titles:
                return f"当前已检索到 {len(evidence_titles)} 份相关资料，但不足以直接确认该字段。"
            if web_titles:
                return f"当前仅补到 {len(web_titles)} 条公开网页线索，仍不足以正式确认该字段。"
            return "当前未检索到可直接支撑该字段的客户资料。"
        if evidence_titles:
            return "主要参考：" + "；".join(evidence_titles[:2])
        if web_titles:
            return "网页补充：" + "；".join(web_titles[:2])
        return "本字段由模板链路自动生成，但未记录到明确证据标题。"

    def estimate_template_field_confidence(
        *,
        field_type: str,
        value_kind: str,
        evidence_count: int,
        review_required: bool,
    ) -> float:
        if value_kind == "missing":
            return 0.0
        base = {
            "precise_fact": 0.9,
            "quantitative_result": 0.82,
            "governance_mechanism": 0.62,
            "structural_summary": 0.72,
            "attachment_material": 0.5,
            "general": 0.6,
        }.get(field_type, 0.6)
        if evidence_count <= 0:
            base -= 0.22
        elif evidence_count == 1:
            base -= 0.08
        if review_required:
            base -= 0.18
        if value_kind == "inference":
            base -= 0.15
        return round(max(0.0, min(0.98, base)), 2)

    def create_client_template_fill_run(client_id: str, template_path_raw: str) -> ClientTemplateFillRunRecord:
        template_path = Path(template_path_raw).expanduser()
        timestamp = now_iso()
        run_id = new_id("tmplfill")
        state.db.execute(
            """
            INSERT INTO client_template_fill_runs(
                id, client_id, template_name, template_path, status, phase, progress, stage_label, elapsed_ms,
                field_count, processed_count, filled_count, missing_count, current_field_label, evidence_titles_json, fields_json, output_path, error_message,
                created_at, updated_at
            )
            VALUES(?, ?, ?, ?, 'queued', 'queued', 0, '等待开始识别模板字段', 0, 0, 0, 0, 0, NULL, '[]', '[]', NULL, NULL, ?, ?)
            """,
            (run_id, client_id, template_path.name, str(template_path), timestamp, timestamp),
        )
        row = state.db.fetchone("SELECT * FROM client_template_fill_runs WHERE id = ?", (run_id,))
        assert row is not None
        return build_client_template_fill_run(row)

    def _normalize_template_fill_path(template_path_raw: str) -> str:
        return str(Path(template_path_raw).expanduser().resolve(strict=False))

    def fetch_active_client_template_fill_run(
        client_id: str,
        *,
        template_path_raw: str | None = None,
    ) -> ClientTemplateFillRunRecord | None:
        active_rows = state.db.fetchall(
            """
            SELECT *
            FROM client_template_fill_runs
            WHERE client_id = ? AND status IN ('queued', 'running')
            ORDER BY created_at DESC
            """,
            (client_id,),
        )
        if not active_rows:
            return None
        if template_path_raw is None:
            return build_client_template_fill_run(active_rows[0])
        normalized_target = _normalize_template_fill_path(template_path_raw)
        for row in active_rows:
            existing_path = str(row["template_path"] or "")
            if _normalize_template_fill_path(existing_path) == normalized_target:
                return build_client_template_fill_run(row)
        return None

    def update_client_template_fill_run(
        run_id: str,
        *,
        status: str | None = None,
        phase: str | None = None,
        progress: float | None = None,
        stage_label: str | None = None,
        elapsed_ms: float | None = None,
        field_count: int | None = None,
        processed_count: int | None = None,
        filled_count: int | None = None,
        missing_count: int | None = None,
        current_field_label: str | None = None,
        clear_current_field_label: bool = False,
        evidence_titles: list[str] | None = None,
        fields: list[ClientTemplateFillFieldRecord] | None = None,
        output_path: str | None = None,
        error_message: str | None = None,
    ) -> None:
        row = state.db.fetchone("SELECT * FROM client_template_fill_runs WHERE id = ?", (run_id,))
        if not row:
            return
        state.db.execute(
            """
            UPDATE client_template_fill_runs
            SET status = ?, phase = ?, progress = ?, stage_label = ?, elapsed_ms = ?,
                field_count = ?, processed_count = ?, filled_count = ?, missing_count = ?, current_field_label = ?, evidence_titles_json = ?, fields_json = ?,
                output_path = ?, error_message = ?, updated_at = ?
            WHERE id = ?
            """,
            (
                status or str(row["status"]),
                phase or str(row["phase"]),
                float(progress if progress is not None else row["progress"]),
                stage_label if stage_label is not None else (str(row["stage_label"]) if row["stage_label"] else None),
                float(elapsed_ms if elapsed_ms is not None else row["elapsed_ms"]),
                int(field_count if field_count is not None else row["field_count"]),
                int(processed_count if processed_count is not None else row["processed_count"]),
                int(filled_count if filled_count is not None else row["filled_count"]),
                int(missing_count if missing_count is not None else row["missing_count"]),
                None
                if clear_current_field_label
                else (
                    current_field_label
                    if current_field_label is not None
                    else (str(row["current_field_label"]) if row["current_field_label"] else None)
                ),
                to_json(evidence_titles if evidence_titles is not None else from_json(str(row["evidence_titles_json"] or "[]"), [])),
                to_json([field.model_dump() for field in fields] if fields is not None else from_json(str(row["fields_json"] or "[]"), [])),
                output_path if output_path is not None else (str(row["output_path"]) if row["output_path"] else None),
                error_message if error_message is not None else (str(row["error_message"]) if row["error_message"] else None),
                now_iso(),
                run_id,
            ),
        )

    def build_client_template_fill_run(row) -> ClientTemplateFillRunRecord:
        evidence_titles = from_json(str(row["evidence_titles_json"] or "[]"), [])
        fields_data = from_json(str(row["fields_json"] or "[]"), [])
        template_path = Path(str(row["template_path"] or "")).expanduser()
        fields = [
            ClientTemplateFillFieldRecord(**item)
            for item in fields_data
            if isinstance(item, dict)
        ]
        review_field_count = sum(1 for item in fields if item.reviewRequired or item.status == "missing")
        attachment_checklist = extract_docx_attachment_checklist(template_path) if template_path.exists() and template_path.suffix.lower() == ".docx" else []
        return ClientTemplateFillRunRecord(
            id=str(row["id"]),
            clientId=str(row["client_id"]),
            templateName=str(row["template_name"]),
            templatePath=str(row["template_path"]),
            status=str(row["status"]),  # type: ignore[arg-type]
            phase=str(row["phase"]),  # type: ignore[arg-type]
            progress=float(row["progress"] or 0.0),
            stageLabel=str(row["stage_label"]) if row["stage_label"] else None,
            elapsedMs=float(row["elapsed_ms"] or 0.0),
            fieldCount=int(row["field_count"] or 0),
            processedCount=int(row["processed_count"] or 0),
            filledCount=int(row["filled_count"] or 0),
            missingCount=int(row["missing_count"] or 0),
            reviewFieldCount=review_field_count,
            currentFieldLabel=str(row["current_field_label"]) if row["current_field_label"] else None,
            evidenceTitles=[str(item) for item in evidence_titles] if isinstance(evidence_titles, list) else [],
            attachmentChecklist=attachment_checklist,
            fields=fields,
            outputPath=str(row["output_path"]) if row["output_path"] else None,
            errorMessage=str(row["error_message"]) if row["error_message"] else None,
            createdAt=str(row["created_at"]),
            updatedAt=str(row["updated_at"]),
        )

    def fetch_client_template_fill_run(client_id: str, run_id: str) -> ClientTemplateFillRunRecord:
        row = state.db.fetchone(
            "SELECT * FROM client_template_fill_runs WHERE id = ? AND client_id = ?",
            (run_id, client_id),
        )
        if not row:
            raise HTTPException(status_code=404, detail="Template fill run not found")
        return build_client_template_fill_run(row)

    def _fill_client_template_docx_impl(
        client_id: str,
        template_path_raw: str,
        progress_callback: Callable[[str, float, str, int, int, int, int, str | None, list[str], list[ClientTemplateFillFieldRecord]], None] | None = None,
    ) -> ClientTemplateFillResponse:
        client = build_client_summary(client_id)
        template_path = Path(template_path_raw).expanduser()
        if not template_path.exists() or not template_path.is_file():
            raise HTTPException(status_code=400, detail="模板文件不存在。")
        if template_path.suffix.lower() != ".docx":
            raise HTTPException(status_code=400, detail="当前模板自动填写 MVP 只支持 .docx。")
        fields = extract_docx_template_fields(template_path)
        if not fields:
            raise HTTPException(status_code=400, detail="没有识别到可自动填写的字段。请在 docx 中使用 {{字段名}} 占位符，或使用“标题列 + 空白答案列”的两列表格。")

        ordered_labels: list[str] = []
        seen_labels: set[str] = set()
        for field in fields:
            if field.label in seen_labels:
                continue
            seen_labels.add(field.label)
            ordered_labels.append(field.label)

        values: dict[str, str] = {}
        field_records: list[ClientTemplateFillFieldRecord] = []
        total_fields = len(ordered_labels)
        field_types = {label: infer_template_field_type(label) for label in ordered_labels}
        attachment_checklist = extract_docx_attachment_checklist(template_path)
        batch_size = 4
        if progress_callback:
            progress_callback("parsing", 12.0, f"已识别 {total_fields} 个待填写字段", total_fields, 0, 0, 0, ordered_labels[0] if ordered_labels else None, [], [])
        for batch_start in range(0, total_fields, batch_size):
            batch_labels = ordered_labels[batch_start : batch_start + batch_size]
            batch_contexts: list[tuple[str, str]] = []
            batch_evidence: dict[str, list[EvidenceItem]] = {}
            batch_web_sources: dict[str, list[TemplateWebSource]] = {}
            batch_label = "、".join(batch_labels[:2]) + (" 等" if len(batch_labels) > 2 else "")
            for label in batch_labels:
                context_summary, evidence, web_sources = build_template_fill_context(
                    client_id,
                    template_path.name,
                    label,
                    field_type=field_types.get(label, "general"),
                    evidence_limit=4,
                    excerpt_limit=900,
                    evidence_char_budget=4200,
                    dna_max_chars=900,
                )
                batch_contexts.append((label, context_summary))
                batch_evidence[label] = evidence
                batch_web_sources[label] = web_sources
            if progress_callback:
                current_titles = list(
                    dict.fromkeys(
                        [
                            *[title for record in field_records for title in record.evidenceTitles],
                            *[title for record in field_records for title in record.webSourceTitles],
                            *[item.title for label in batch_labels for item in batch_evidence.get(label, [])[:2]],
                            *[item.title for label in batch_labels for item in batch_web_sources.get(label, [])[:1]],
                        ]
                    )
                )[:8]
                processed = len(field_records)
                progress = min(78.0, 18.0 + (processed / max(total_fields, 1)) * 52.0)
                progress_callback(
                    "retrieving",
                    progress,
                    f"正在检索第 {batch_start + 1}-{batch_start + len(batch_labels)} 个字段所需资料",
                    total_fields,
                    batch_start,
                    sum(1 for item in field_records if item.status == "filled"),
                    sum(1 for item in field_records if item.status == "missing"),
                    batch_label,
                    current_titles,
                    field_records,
                )
            try:
                batch_values = state.ai.generate_template_field_values_batch(
                    template_name=template_path.name,
                    client_name=client.name,
                    field_contexts=batch_contexts,
                    field_types=field_types,
                )
                for label in batch_labels:
                    value = str(batch_values.get(label) or "【待确认】当前缺少可直接填写该字段的资料。").strip()
                    field_type = field_types.get(label, "general")
                    evidence_titles = list(dict.fromkeys(item.title for item in batch_evidence.get(label, [])[:3]))
                    web_titles = list(dict.fromkeys(item.title for item in batch_web_sources.get(label, [])[:2]))
                    value_kind = infer_template_value_kind(value, field_type)
                    review_required = value_kind in {"missing", "inference"} or value.startswith("【待确认】")
                    values[label] = value
                    field_records.append(
                        ClientTemplateFillFieldRecord(
                            label=label,
                            value=value,
                            status="missing" if value.startswith("【待确认】") else "filled",
                            evidenceTitles=evidence_titles,
                            webSourceTitles=web_titles,
                            fieldType=field_type,
                            valueKind=value_kind,
                            confidence=estimate_template_field_confidence(
                                field_type=field_type,
                                value_kind=value_kind,
                                evidence_count=len(evidence_titles),
                                review_required=review_required,
                            ),
                            basisSummary=summarize_template_field_basis(value, evidence_titles, web_titles),
                            followUpQuestion=build_template_follow_up_question(field_type, label) if review_required else None,
                            suggestedSources=build_template_suggested_sources(field_type, label),
                            reviewRequired=review_required,
                        )
                    )
                    if progress_callback:
                        filled_count = sum(1 for item in field_records if item.status == "filled")
                        missing_count = len(field_records) - filled_count
                        progress = min(92.0, 42.0 + (len(field_records) / max(total_fields, 1)) * 48.0)
                        progress_callback(
                            "writing",
                            progress,
                            f"正在写入字段：{label}",
                            total_fields,
                            len(field_records),
                            filled_count,
                            missing_count,
                            label,
                            list(
                                dict.fromkeys(
                                    [title for record in field_records for title in record.evidenceTitles]
                                    + [title for record in field_records for title in record.webSourceTitles]
                                )
                            )[:8],
                            field_records,
                        )
            except AiInvocationError:
                for label in batch_labels:
                    try:
                        context_summary = next(summary for current_label, summary in batch_contexts if current_label == label)
                        evidence = batch_evidence.get(label, [])
                        web_titles = list(dict.fromkeys(item.title for item in batch_web_sources.get(label, [])[:2]))
                        value = state.ai.generate_template_field_value(
                            field_label=label,
                            template_name=template_path.name,
                            client_name=client.name,
                            context_summary=context_summary,
                            field_type=field_types.get(label, "general"),
                        )
                    except AiInvocationError as error:
                        raise HTTPException(
                            status_code=504,
                            detail=f"字段“{label}”填写超时或模型未返回结果（{batch_start + batch_labels.index(label) + 1}/{total_fields}）。{error.detail}",
                        ) from error
                    except HTTPException:
                        raise
                    except Exception as error:
                        raise HTTPException(
                            status_code=500,
                            detail=f"字段“{label}”填写失败（{batch_start + batch_labels.index(label) + 1}/{total_fields}）。{error}",
                        ) from error
                    field_type = field_types.get(label, "general")
                    evidence_titles = list(dict.fromkeys(item.title for item in evidence[:3]))
                    value_kind = infer_template_value_kind(value, field_type)
                    review_required = value_kind in {"missing", "inference"} or value.startswith("【待确认】")
                    values[label] = value
                    field_records.append(
                        ClientTemplateFillFieldRecord(
                            label=label,
                            value=value,
                            status="missing" if value.startswith("【待确认】") else "filled",
                            evidenceTitles=evidence_titles,
                            webSourceTitles=web_titles,
                            fieldType=field_type,
                            valueKind=value_kind,
                            confidence=estimate_template_field_confidence(
                                field_type=field_type,
                                value_kind=value_kind,
                                evidence_count=len(evidence_titles),
                                review_required=review_required,
                            ),
                            basisSummary=summarize_template_field_basis(value, evidence_titles, web_titles),
                            followUpQuestion=build_template_follow_up_question(field_type, label) if review_required else None,
                            suggestedSources=build_template_suggested_sources(field_type, label),
                            reviewRequired=review_required,
                        )
                    )
                    if progress_callback:
                        filled_count = sum(1 for item in field_records if item.status == "filled")
                        missing_count = len(field_records) - filled_count
                        progress = min(92.0, 42.0 + (len(field_records) / max(total_fields, 1)) * 48.0)
                        progress_callback(
                            "writing",
                            progress,
                            f"正在写入字段：{label}",
                            total_fields,
                            len(field_records),
                            filled_count,
                            missing_count,
                            label,
                            list(
                                dict.fromkeys(
                                    [title for record in field_records for title in record.evidenceTitles]
                                    + [title for record in field_records for title in record.webSourceTitles]
                                )
                            )[:8],
                            field_records,
                        )
            except HTTPException:
                raise
            except Exception as error:
                batch_end = min(total_fields, batch_start + len(batch_labels))
                raise HTTPException(
                    status_code=500,
                    detail=f"字段批次填写失败（{batch_start + 1}-{batch_end}/{total_fields}）。{error}",
                ) from error
        folders = ensure_client_workspace(state.data_dir, client_id)
        target_dir = folders["战略陪伴"] / "自动填写文档"
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        target_path = target_dir / f"{safe_filename(template_path.stem)}_已填写_{timestamp}.docx"
        apply_docx_template_values(template_path, target_path, values)
        filled_count = sum(1 for item in field_records if item.status == "filled")
        missing_count = len(field_records) - filled_count
        if progress_callback:
            progress_callback(
                "completed",
                100.0,
                f"已生成结果文档，共 {filled_count} 项自动填写，{missing_count} 项待确认",
                len(field_records),
                len(field_records),
                filled_count,
                missing_count,
                None,
                list(
                    dict.fromkeys(
                        [title for record in field_records for title in record.evidenceTitles]
                        + [title for record in field_records for title in record.webSourceTitles]
                    )
                )[:8],
                field_records,
            )
        log_activity(
            "document.template_fill",
            "document",
            str(target_path),
            {"clientId": client_id, "templatePath": str(template_path), "fieldCount": len(field_records), "filledCount": filled_count},
        )
        return ClientTemplateFillResponse(
            path=str(target_path),
            fileName=target_path.name,
            fieldCount=len(field_records),
            filledCount=filled_count,
            missingCount=missing_count,
            reviewFieldCount=sum(1 for item in field_records if item.reviewRequired or item.status == "missing"),
            attachmentChecklist=attachment_checklist,
            fields=field_records,
        )

    def fill_client_template_docx(client_id: str, template_path_raw: str) -> ClientTemplateFillResponse:
        return _fill_client_template_docx_impl(client_id, template_path_raw)

    def run_client_template_fill(client_id: str, run_id: str, template_path_raw: str) -> None:
        started_at = perf_counter()
        update_client_template_fill_run(
            run_id,
            status="running",
            phase="parsing",
            progress=4.0,
            stage_label="正在识别模板字段",
            elapsed_ms=0.0,
            processed_count=0,
            current_field_label=None,
        )

        def _progress(
            phase: str,
            progress: float,
            stage_label: str,
            field_count: int,
            processed_count: int,
            filled_count: int,
            missing_count: int,
            current_field_label: str | None,
            evidence_titles: list[str],
            fields: list[ClientTemplateFillFieldRecord],
        ) -> None:
            update_client_template_fill_run(
                run_id,
                status="running" if phase not in {"completed", "failed"} else ("completed" if phase == "completed" else "failed"),
                phase=phase,
                progress=progress,
                stage_label=stage_label,
                elapsed_ms=(perf_counter() - started_at) * 1000,
                field_count=field_count,
                processed_count=processed_count,
                filled_count=filled_count,
                missing_count=missing_count,
                current_field_label=current_field_label,
                evidence_titles=evidence_titles,
                fields=fields,
            )

        try:
            result = _fill_client_template_docx_impl(client_id, template_path_raw, _progress)
            update_client_template_fill_run(
                run_id,
                status="completed",
                phase="completed",
                progress=100.0,
                stage_label="填写完成",
                elapsed_ms=(perf_counter() - started_at) * 1000,
                field_count=result.fieldCount,
                processed_count=result.fieldCount,
                filled_count=result.filledCount,
                missing_count=result.missingCount,
                clear_current_field_label=True,
                evidence_titles=list(dict.fromkeys(title for field in result.fields for title in field.evidenceTitles))[:8],
                fields=result.fields,
                output_path=result.path,
                error_message=None,
            )
        except HTTPException as error:
            last_phase = str(state.db.scalar("SELECT phase FROM client_template_fill_runs WHERE id = ?", (run_id,)) or "parsing")
            update_client_template_fill_run(
                run_id,
                status="failed",
                phase=last_phase,
                progress=max(8.0, min(float(state.db.scalar("SELECT progress FROM client_template_fill_runs WHERE id = ?", (run_id,)) or 0.0), 96.0)),
                stage_label="模板填写失败",
                elapsed_ms=(perf_counter() - started_at) * 1000,
                processed_count=int(state.db.scalar("SELECT processed_count FROM client_template_fill_runs WHERE id = ?", (run_id,)) or 0),
                clear_current_field_label=True,
                error_message=str(error.detail),
            )
        except Exception as error:
            last_phase = str(state.db.scalar("SELECT phase FROM client_template_fill_runs WHERE id = ?", (run_id,)) or "parsing")
            update_client_template_fill_run(
                run_id,
                status="failed",
                phase=last_phase,
                progress=max(8.0, min(float(state.db.scalar("SELECT progress FROM client_template_fill_runs WHERE id = ?", (run_id,)) or 0.0), 96.0)),
                stage_label="模板填写失败",
                elapsed_ms=(perf_counter() - started_at) * 1000,
                processed_count=int(state.db.scalar("SELECT processed_count FROM client_template_fill_runs WHERE id = ?", (run_id,)) or 0),
                clear_current_field_label=True,
                error_message=str(error),
            )

    def workspace_for_client(client_id: str) -> ClientWorkspaceResponse:
        client = build_client_summary(client_id)
        ensure_standard_client_folders(client_id)
        hidden_labels = get_hidden_client_folders(client_id)
        document_limit = 200
        folder_rows = {
            str(row["label"]): row
            for row in state.db.fetchall(
                "SELECT * FROM client_folders WHERE client_id = ?",
                (client_id,),
            )
            if str(row["label"]) in HUMAN_VISIBLE_CATEGORIES and str(row["label"]) not in hidden_labels
        }
        folders = [
            ClientFolder(
                id=str(folder_rows[label]["id"]),
                clientId=str(folder_rows[label]["client_id"]),
                label=str(folder_rows[label]["label"]),
                path=str(folder_rows[label]["path"]),
                fileCount=int(folder_rows[label]["file_count"]),
                lastScannedAt=str(folder_rows[label]["last_scanned_at"]) if folder_rows[label]["last_scanned_at"] else None,
            )
            for label in HUMAN_VISIBLE_CATEGORIES
            if label not in hidden_labels
            if label in folder_rows
        ]
        documents = [
            DocumentRecord(
                id=str(row["id"]),
                clientId=str(row["client_id"]),
                folderId=str(row["folder_id"]) if row["folder_id"] else None,
                title=str(row["title"]),
                path=str(row["path"]),
                kind=str(row["kind"]),
                source=str(row["source"]),
                excerpt=str(row["excerpt"]),
                tags=_parse_json_list(row["tags_json"]),
                importedAt=str(row["created_at"]),
            )
            for row in state.db.fetchall("SELECT * FROM documents WHERE client_id = ? ORDER BY created_at DESC LIMIT ?", (client_id, document_limit))
        ]
        imports = [
            ImportRecord(
                id=str(row["id"]),
                clientId=str(row["client_id"]),
                sourcePath=str(row["source_path"]),
                mode=str(row["mode"]),  # type: ignore[arg-type]
                status=str(row["status"]),  # type: ignore[arg-type]
                importedCount=int(row["imported_count"]),
                skippedCount=int(row["skipped_count"]),
                createdAt=str(row["created_at"]),
            )
            for row in state.db.fetchall("SELECT * FROM imports WHERE client_id = ? ORDER BY created_at DESC LIMIT 10", (client_id,))
        ]
        threads = [
            ChatThread(
                id=str(row["id"]),
                clientId=str(row["client_id"]),
                title=str(row["title"]),
                createdAt=str(row["created_at"]),
                updatedAt=str(row["updated_at"]),
            )
            for row in state.db.fetchall("SELECT * FROM chat_threads WHERE client_id = ? ORDER BY updated_at DESC", (client_id,))
        ]
        messages = [
            build_chat_message(row)
            for row in state.db.fetchall(
                """
                SELECT recent.*
                FROM (
                    SELECT m.*
                    FROM chat_messages m
                    JOIN chat_threads t ON t.id = m.thread_id
                    WHERE t.client_id = ?
                    ORDER BY m.created_at DESC
                    LIMIT 50
                ) recent
                ORDER BY recent.created_at ASC
                """,
                (client_id,),
            )
        ]
        analysis_runs = [
            build_client_analysis_run(row)
            for row in state.db.fetchall(
                "SELECT * FROM client_analysis_runs WHERE client_id = ? ORDER BY updated_at DESC LIMIT 12",
                (client_id,),
            )
        ]
        meetings = [
            build_meeting_summary(row)
            for row in state.db.fetchall("SELECT * FROM meetings WHERE client_id = ? ORDER BY updated_at DESC", (client_id,))
        ]
        goals = [
            GoalRecord(
                id=str(row["id"]),
                clientId=str(row["client_id"]),
                title=str(row["title"]),
                quarter=str(row["quarter"]),
                progress=int(row["progress"]),
                ownerName=str(row["owner_name"]),
            )
            for row in state.db.fetchall("SELECT * FROM goal_records WHERE client_id = ? ORDER BY updated_at DESC", (client_id,))
        ]
        dna_modules = list_client_dna_modules(client_id)
        project_modules = list_project_modules(client_id)
        project_flows = list_project_flows(client_id)
        dna_terms = [
            DnaTerm(
                id=str(row["id"]),
                clientId=str(row["client_id"]),
                category=str(row["category"]),
                canonicalName=str(row["canonical_name"]),
                aliases=_parse_json_list(row["aliases_json"]),
                description=str(row["description"]),
                sourceLevel="client",
            )
            for row in state.db.fetchall("SELECT * FROM dna_terms WHERE client_id = ? ORDER BY updated_at DESC", (client_id,))
        ]
        document_cards = [
            build_document_card_record(item)
            for item in fetch_document_cards(state.db, client_id, data_dir=state.data_dir, limit=document_limit)
        ]
        knowledge_jobs = [KnowledgeJobRecord(**item) for item in fetch_recent_knowledge_jobs(state.db, client_id, limit=8)]
        recent_reclass_events = [FileReclassEventRecord(**item) for item in fetch_recent_reclass_events(state.db, client_id, limit=8)]
        knowledge_status = build_knowledge_status_record(client_id)
        notebook_summary = get_client_notebook_response(state.db, client_id).organizationNotebookSnapshot
        memory_status = get_client_memory_status(state.db, client_id)
        related_tasks = fetch_tasks(
            "t.source_id = ? OR t.source_id IN (SELECT id FROM meetings WHERE client_id = ?)",
            (client_id, client_id),
        )
        return ClientWorkspaceResponse(
            client=client,
            folders=folders,
            documents=documents,
            documentCards=document_cards,
            imports=imports,
            knowledgeStatus=knowledge_status,
            knowledgeJobs=knowledge_jobs,
            recentReclassEvents=recent_reclass_events,
            surrogateCount=knowledge_status.surrogateCount,
            memoryDocCount=knowledge_status.memoryDocCount,
            threads=threads,
            recentMessages=messages,
            analysisRuns=analysis_runs,
            meetings=meetings,
            goals=goals,
            dnaModules=dna_modules,
            projectModules=project_modules,
            projectFlows=project_flows,
            dnaTerms=dna_terms,
            relatedTasks=related_tasks,
            notebookSummary=notebook_summary,
            memoryStatus=memory_status,
        )

    def _strategic_unique_non_empty(values: list[str | None]) -> list[str]:
        seen: set[str] = set()
        ordered: list[str] = []
        for raw in values:
            text = (raw or "").strip()
            if not text or text in seen:
                continue
            seen.add(text)
            ordered.append(text)
        return ordered

    def _strategic_first_non_empty(values: list[str | None], fallback: str) -> str:
        ordered = _strategic_unique_non_empty(values)
        return ordered[0] if ordered else fallback

    def _strategic_truncate(value: str | None, limit: int = 88) -> str:
        text = (value or "").strip()
        if not text:
            return ""
        return text if len(text) <= limit else f"{text[:limit - 1]}…"

    def _strategic_format_date_label(value: str | None) -> str:
        if not value:
            return "待补"
        try:
            return datetime.fromisoformat(value).strftime("%Y-%m-%d")
        except ValueError:
            return value[:10]

    def _strategic_is_placeholder_context_text(value: str | None) -> bool:
        text = (value or "").strip()
        if not text:
            return True
        return any(pattern in text for pattern in STRATEGIC_PLACEHOLDER_CONTEXT_PATTERNS)

    def _strategic_is_relationship_task(task: TaskRecord) -> bool:
        task_text = " ".join([task.title, task.desc or ""])
        return any(keyword in task_text for keyword in STRATEGIC_RELATIONSHIP_TASK_KEYWORDS)

    def _strategic_has_contextual_description(task: TaskRecord) -> bool:
        description = (task.desc or "").strip()
        if len(description) < 20 or _strategic_is_placeholder_context_text(description):
            return False
        keyword_hits = sum(1 for keyword in STRATEGIC_CONTEXTUAL_DESCRIPTION_KEYWORDS if keyword in description)
        return keyword_hits >= 2 or len(description) >= 42

    def _strategic_background_signal_score(task: TaskRecord) -> float:
        score = 0.0
        readiness = task.backgroundReadiness
        if readiness:
            score = max(score, float(readiness.score or 0.0))
            if readiness.backgroundSources:
                score += 0.08
            if readiness.level == "high":
                score += 0.08
            elif readiness.level == "medium":
                score += 0.04
        if task.eventLineId or task.eventLineName:
            score += 0.22
        if task.projectModuleId or task.projectModuleName or task.projectFlowId or task.projectFlowName:
            score += 0.16
        if task.memoryHints:
            score += 0.08
        if task.linkedFactsPreview:
            score += 0.12
        if task.attachments:
            score += 0.08
        if _strategic_has_contextual_description(task):
            score += 0.18
        return round(min(score, 1.0), 2)

    def _strategic_has_meaningful_task_background(task: TaskRecord) -> bool:
        project_context = task.projectContext
        if task.eventLineId or task.eventLineName or task.projectModuleId or task.projectModuleName or task.projectFlowId or task.projectFlowName:
            return True
        background_score = _strategic_background_signal_score(task)
        if _strategic_is_relationship_task(task):
            readiness_sources = set(task.backgroundReadiness.backgroundSources) if task.backgroundReadiness else set()
            has_non_self_memory = any(
                fact.sourceType != "task" or fact.sourceId != task.id
                for fact in task.linkedFactsPreview
            )
            if has_non_self_memory and (
                {"event_line_memory", "event_line_facts", "client_facts"} & readiness_sources
            ) and background_score >= 0.55:
                return True
            if _strategic_has_contextual_description(task):
                return True
            return False
        if background_score >= 0.5:
            return True
        if _strategic_has_contextual_description(task):
            return True
        if (task.tags or task.attachments) and background_score >= 0.28:
            return True
        if not project_context:
            return False
        return any(
            not _strategic_is_placeholder_context_text(item)
            for item in [
                project_context.backgroundSummary,
                project_context.currentFocus,
                project_context.currentBlocker,
                project_context.nextAction,
                project_context.recentProgress,
            ]
        )

    def _load_org_model_profile_safe() -> OrgModelProfileRecord | None:
        if not get_cloud_token():
            return None
        try:
            payload = cloud_request("GET", "/api/v1/settings/org-model/profile")
        except HTTPException:
            return None
        if not isinstance(payload, dict):
            return None
        try:
            return OrgModelProfileRecord(**payload)
        except Exception:
            return None

    def _load_strategic_snapshot_row(client_id: str):
        return state.db.fetchone("SELECT * FROM strategic_cockpit_snapshots WHERE client_id = ?", (client_id,))

    def _strategic_meeting_event_line_ids(client_id: str, meeting_title: str, *, meeting_id: str | None = None) -> list[str]:
        line_ids = _strategic_unique_non_empty(
            [
                *(
                    [
                        str(row["event_line_id"]).strip()
                        for row in state.db.fetchall(
                            """
                            SELECT DISTINCT event_line_id
                            FROM tasks
                            WHERE source_type = 'meeting' AND source_id = ? AND event_line_id IS NOT NULL AND TRIM(event_line_id) <> ''
                            """,
                            (meeting_id,),
                        )
                    ]
                    if meeting_id
                    else []
                ),
                *(
                    [item.id for item in list_linked_event_lines(state.db, client_id)]
                    if any(keyword in meeting_title for keyword in STRATEGIC_WEEKLY_MEETING_KEYWORDS)
                    else []
                ),
            ]
        )
        return line_ids[:12]

    def _strategic_task_pool(client_id: str, workspace: ClientWorkspaceResponse) -> list[TaskRecord]:
        task_map: dict[str, TaskRecord] = {task.id: task for task in workspace.relatedTasks}
        try:
            task_candidates = cloud_task_board().tasks if get_cloud_token() else fetch_tasks("t.source_type != ?", (AGENT_AUTO_SOURCE_TYPE,))
        except HTTPException:
            task_candidates = fetch_tasks("t.source_type != ?", (AGENT_AUTO_SOURCE_TYPE,))
        for task in task_candidates:
            if task.clientId == client_id or (task.projectContext and task.projectContext.clientId == client_id):
                task_map[task.id] = task
        return sorted(task_map.values(), key=lambda item: item.updatedAt, reverse=True)

    def _strategic_priority(probability: str | None) -> Literal["high", "medium", "low"]:
        if probability == "high":
            return "high"
        if probability == "medium":
            return "medium"
        return "low"

    def _strategic_health_status(score: int, *, calibrated: bool) -> Literal["healthy", "watch", "risk", "uncalibrated"]:
        if not calibrated:
            return "uncalibrated"
        if score >= 75:
            return "healthy"
        if score >= 45:
            return "watch"
        return "risk"

    def build_strategic_cockpit_snapshot(client_id: str) -> StrategicCockpitSnapshotRecord:
        workspace = workspace_for_client(client_id)
        client = workspace.client
        notebook_response = get_client_notebook_response(state.db, client_id)
        notebook_summary = notebook_response.organizationNotebookSnapshot
        memory_status = get_client_memory_status(state.db, client_id)
        review_dashboard = local_review_dashboard()
        review_analysis = review_dashboard.workAnalysis
        client_tasks = _strategic_task_pool(client_id, workspace)
        analyzable_tasks = sorted(
            [task for task in client_tasks if _strategic_has_meaningful_task_background(task)],
            key=lambda item: (_strategic_background_signal_score(item), item.updatedAt),
            reverse=True,
        )
        background_thin_tasks = sorted(
            [task for task in client_tasks if not _strategic_has_meaningful_task_background(task)],
            key=lambda item: item.updatedAt,
            reverse=True,
        )
        linked_event_lines = list_linked_event_lines(state.db, client_id)
        event_line_ids = {task.eventLineId for task in analyzable_tasks if task.eventLineId}
        event_line_ids.update({item.id for item in linked_event_lines})
        module_names = {item.name for item in workspace.projectModules}
        event_line_names = {task.eventLineName for task in analyzable_tasks if task.eventLineName}

        review_event_line_summaries = [
            item
            for item in (review_analysis.eventLineSummaries if review_analysis else [])
            if item.eventLineId in event_line_ids
            or item.projectName == client.name
            or (item.moduleName and item.moduleName in module_names)
            or item.title in event_line_names
        ]
        review_completeness = [
            item for item in (review_analysis.eventLineCompleteness if review_analysis else []) if item.eventLineId in event_line_ids
        ]
        review_risk_cards = [
            item for item in (review_analysis.riskCards if review_analysis else []) if item.eventLineId in event_line_ids
        ]
        review_opportunity_cards = [
            item for item in (review_analysis.opportunityCards if review_analysis else []) if item.eventLineId in event_line_ids
        ]

        notebook_stage = notebook_summary.currentStage if notebook_summary else ""
        notebook_intro = notebook_summary.organizationIntro if notebook_summary else ""
        notebook_recent_facts = list(notebook_summary.recentFacts) if notebook_summary else []
        notebook_gaps = list(notebook_summary.informationGaps) if notebook_summary else []
        snapshot_row = _load_strategic_snapshot_row(client_id)
        session_user = get_cached_session_user()
        org_model_profile = _load_org_model_profile_safe()
        leader_user_id = org_model_profile.organization.leaderUserId if org_model_profile else None
        is_ceo = bool(session_user and leader_user_id and session_user.id == leader_user_id)
        permission = StrategicPermissionRecord(
            canEdit=is_ceo,
            isCeo=is_ceo,
            leaderUserId=leader_user_id,
            notice="请先在组织设置中确认 CEO 账号" if not leader_user_id else (None if is_ceo else "当前页面仅 CEO 可确认和改写经营判断"),
        )

        recent_meeting = workspace.meetings[0] if workspace.meetings else None
        recent_analysis = workspace.analysisRuns[0] if workspace.analysisRuns else None
        active_task = next((item for item in analyzable_tasks if item.status == "doing"), analyzable_tasks[0] if analyzable_tasks else None)
        completed_tasks = [item for item in analyzable_tasks if item.status == "done"]
        blocked_tasks = [item for item in analyzable_tasks if item.status in {"todo", "rejected"}]

        readiness_checks = [
            any(module.hasDocument for module in workspace.dnaModules),
            bool(workspace.analysisRuns),
            bool(workspace.meetings),
            bool(workspace.goals),
            bool(workspace.projectModules),
            bool(workspace.projectFlows),
            bool(review_event_line_summaries or review_risk_cards or review_opportunity_cards),
            bool(analyzable_tasks or linked_event_lines),
        ]
        readiness_count = len([item for item in readiness_checks if item])
        readiness_status: Literal["ready", "insufficient"] = "ready" if readiness_count >= 4 else "insufficient"
        readiness_gaps = _strategic_unique_non_empty([
            None if workspace.meetings else "当前还没有会议沉淀，缺少经营讨论后的正式信号。",
            None if workspace.goals else "当前还没有业务目标锚点，页面无法判断什么算真正推进。",
            None if workspace.projectModules else "当前还没有业务模块定义，任务和资料难以挂到稳定承接位。",
            None if workspace.projectFlows else "当前还没有关键流程，阻塞和下一步只能停在泛描述。",
            None if (review_event_line_summaries or review_risk_cards or review_opportunity_cards) else "周复盘里还没有事件线摘要、风险卡或机会卡，预测基础不足。",
            f"当前有 {len(background_thin_tasks)} 条任务只写了动作名，没有补对象背景、合作关系和推进目的，暂不纳入洞察与预测。"
            if background_thin_tasks
            else None,
            *notebook_gaps,
        ])
        readiness = StrategicReadinessRecord(
            status=readiness_status,
            score=round(readiness_count / len(readiness_checks) * 100),
            summary=(
                "当前已具备基本判断条件，可以把任务、资料、会议和分析压成经营判断。"
                if readiness_status == "ready"
                else f"当前判断准备度不足。优先补齐：{('；'.join(readiness_gaps[:3])) if readiness_gaps else '会议、目标、模块和流程信号。'}"
            ),
            gaps=readiness_gaps,
        )

        dossier_summary = _strategic_first_non_empty(
            [
                notebook_intro,
                client.intro,
                *[item.summary for item in workspace.dnaModules],
                *[item.summary for item in workspace.documentCards],
            ],
            "当前还没有足够厚的业务底稿。建议优先复用客户 DNA、战略分析、业务分析和战略规划资料，形成对这条业务的稳定描述。",
        )
        dossier_cards = [
            StrategicEvidenceCardRecord(label="机构阶段", value=(notebook_stage or client.stage or "待判断")),
            StrategicEvidenceCardRecord(label="DNA 模块", value=f"{len([item for item in workspace.dnaModules if item.hasDocument])} 个已接入"),
            StrategicEvidenceCardRecord(label="项目模块", value=f"{len(workspace.projectModules)} 个模块"),
            StrategicEvidenceCardRecord(label="项目流程", value=f"{len(workspace.projectFlows)} 条流程"),
            StrategicEvidenceCardRecord(label="资料卡", value=f"{len(workspace.documentCards)} 份"),
            StrategicEvidenceCardRecord(label="分析运行", value=f"{len(workspace.analysisRuns)} 次"),
        ]
        dossier_boundaries = _strategic_unique_non_empty([
            *[risk for flow in workspace.projectFlows for risk in flow.riskPoints],
            *[f"DNA 缺口：{missing}" for module in workspace.dnaModules for missing in module.missingInfo],
            *notebook_gaps,
            f"当前仍有 {workspace.knowledgeStatus.reviewPendingDocuments} 份资料待复核，关键判断不能说太满。"
            if workspace.knowledgeStatus and workspace.knowledgeStatus.reviewPendingDocuments
            else None,
        ])[:4]

        key_facts = _strategic_unique_non_empty([
            f"最近会议“{recent_meeting.title}”更新于 {_strategic_format_date_label(recent_meeting.updatedAt)}。" if recent_meeting else None,
            f"最近分析围绕“{_strategic_truncate(recent_analysis.question, 32)}”，当前状态 {recent_analysis.status}。" if recent_analysis else None,
            f"当前最明确的业务锚点是“{workspace.goals[0].title}”。" if workspace.goals else None,
            f"最近完成的关键动作包括“{completed_tasks[0].title}”。" if completed_tasks else None,
            f"最近一次资料导入来自 {Path(workspace.imports[0].sourcePath).name}。" if workspace.imports else None,
            f"当前最稳定的事件线是“{linked_event_lines[0].name}”。" if linked_event_lines else None,
            *notebook_recent_facts[:2],
        ])[:5]
        key_warnings = _strategic_unique_non_empty([
            readiness.summary if readiness_status == "insufficient" else None,
            review_risk_cards[0].statement if review_risk_cards else None,
            review_risk_cards[1].statement if len(review_risk_cards) > 1 else None,
            "客户 DNA 还不完整，宏观判断容易受短期热度影响。" if not any(item.hasDocument for item in workspace.dnaModules) else None,
            "当前还没有稳定项目模块，经营讨论缺少承载位。" if not workspace.projectModules else None,
            "当前还没有关键流程结构，阻塞容易反复出现。" if not workspace.projectFlows else None,
            f"有 {len(background_thin_tasks)} 条任务缺少背景描述，当前只保留在事实层，不进入洞察与预测。" if background_thin_tasks else None,
        ])[:5]

        contradiction_auto = "当前业务已经有推进痕迹，但还需要把过程信号稳定压缩成经营判断。"
        contradiction_sources = _strategic_unique_non_empty([
            review_risk_cards[0].title if review_risk_cards else None,
            workspace.goals[0].title if workspace.goals else None,
            active_task.title if active_task else None,
        ])
        if readiness_status == "insufficient":
            contradiction_auto = f"当前先不强行判断，主要因为{readiness_gaps[0] if readiness_gaps else '业务结构和周判断信号都还偏薄'}。"
        elif not any(item.hasDocument for item in workspace.dnaModules):
            contradiction_auto = "运行信号在增加，但业务底稿仍然偏薄，容易把短期现象误当成长期方向。"
        elif not workspace.projectModules or not workspace.projectFlows:
            contradiction_auto = "业务目标已经出现，但承接结构还不够清楚，目标难以稳定落到模块与流程。"
        elif len(workspace.documentCards) >= 6 and not workspace.analysisRuns:
            contradiction_auto = "资料已经在变厚，但还没有被收敛成正式分析与管理判断。"
        elif review_risk_cards:
            contradiction_auto = review_risk_cards[0].statement
        elif len(blocked_tasks) > max(2, len(completed_tasks)):
            contradiction_auto = "当前推进动作不少，但卡点没有被持续拆解，管理注意力容易停留在事务层。"

        core_breakthrough_auto = "先把最重要的一条业务线说清楚，再围绕它组织下一次周会与行动推进。"
        core_breakthrough_sources = _strategic_unique_non_empty([
            contradiction_auto,
            recent_analysis.question if recent_analysis else None,
            workspace.goals[0].title if workspace.goals else None,
        ])
        if readiness_status == "insufficient":
            core_breakthrough_auto = "先补“业务目标 + 模块/流程 + 周会沉淀 + 任务背景说明”这四件基础设施，再谈洞察和预测。"
        elif not any(item.hasDocument for item in workspace.dnaModules):
            core_breakthrough_auto = "先补齐机构说明、项目说明和团队说明，让后续判断不被短期波动带偏。"
        elif not workspace.projectModules:
            core_breakthrough_auto = "把当前业务拆成稳定的项目模块，让经营判断有明确承载位。"
        elif not workspace.projectFlows:
            core_breakthrough_auto = "优先补齐关键流程，让阻塞不再只停留在口头描述。"
        elif recent_analysis:
            core_breakthrough_auto = f"围绕“{_strategic_truncate(recent_analysis.question, 30)}”收敛出 1 个本周期主问题，并让任务链围绕它推进。"
        elif workspace.goals:
            core_breakthrough_auto = f"围绕“{workspace.goals[0].title}”组织下一轮动作与证据，不要同时铺开太多陪伴线。"

        focus_items_auto = _strategic_unique_non_empty([
            f"先把“{workspace.goals[0].title}”落成真正可跟踪的业务目标。" if readiness_status == "insufficient" and workspace.goals else None,
            "先明确这条业务本周期只抓的一个核心目标，不再让任务各走各的。" if readiness_status == "insufficient" and not workspace.goals else None,
            "先把当前业务拆成 2 到 4 个稳定模块，别再直接拿零散任务做判断。" if readiness_status == "insufficient" and not workspace.projectModules else None,
            "先补一条关键推进流程，让“当前阻塞 / 下一步”有稳定承接位。" if readiness_status == "insufficient" and not workspace.projectFlows else None,
            "关系推进类任务必须在描述里写清：对象是谁、当前关系、这次动作想推动什么。" if readiness_status == "insufficient" and background_thin_tasks else None,
            "本周至少形成一次业务盘点会或推进会，让战略页开始有正式会议信号。" if readiness_status == "insufficient" and not recent_meeting else None,
            f"围绕“{workspace.goals[0].title}”校准当前业务的真实优先级。" if readiness_status == "ready" and workspace.goals else None,
            f"把推进动作“{active_task.title}”和本周主问题挂钩。" if readiness_status == "ready" and active_task else None,
            f"下一次会谈先对齐最近会议“{recent_meeting.title}”里仍未闭环的问题。" if readiness_status == "ready" and recent_meeting else None,
            f"把最近分析“{_strategic_truncate(recent_analysis.question, 32)}”从洞察变成动作。" if readiness_status == "ready" and recent_analysis else None,
        ])[:3]

        snapshot_focus_items = _parse_json_list(snapshot_row["focus_items_json"]) if snapshot_row and snapshot_row["focus_items_json"] else []
        week_summary_confirmed = str(snapshot_row["week_summary"]).strip() if snapshot_row and snapshot_row["week_summary"] else ""
        main_contradiction_confirmed = str(snapshot_row["main_contradiction"]).strip() if snapshot_row and snapshot_row["main_contradiction"] else ""
        core_breakthrough_confirmed = str(snapshot_row["core_breakthrough"]).strip() if snapshot_row and snapshot_row["core_breakthrough"] else ""

        week_summary_auto = (
            "当前先把业务底稿补厚，暂不把零散动作抬成经营判断。"
            if readiness_status == "insufficient"
            else _strategic_first_non_empty([
                review_dashboard.orgReport.headline if review_dashboard.orgReport else None,
                f"这条业务当前更需要把“{_strategic_truncate(workspace.goals[0].title if workspace.goals else recent_analysis.question if recent_analysis else review_event_line_summaries[0].title if review_event_line_summaries else '主问题', 28)}”收敛成稳定经营判断。",
            ], "当前业务还缺经营摘要，需要先补齐底稿和周判断。")
        )
        freshness = " / ".join(_strategic_unique_non_empty([
            f"会议 {_strategic_format_date_label(recent_meeting.updatedAt)}" if recent_meeting else None,
            f"分析 {_strategic_format_date_label(recent_analysis.updatedAt)}" if recent_analysis else None,
            f"导入 {_strategic_format_date_label(workspace.imports[0].createdAt)}" if workspace.imports else None,
            f"任务 {_strategic_format_date_label(client_tasks[0].updatedAt)}" if client_tasks else None,
            f"复盘 {_strategic_format_date_label(review_dashboard.currentReview.updatedAt)}" if review_dashboard.currentReview else None,
        ])[:4]) or "当前还没有足够新的更新信号"

        headline = StrategicHeadlineRecord(
            weekSummary=StrategicJudgmentRecord(
                value=week_summary_confirmed or week_summary_auto,
                status="confirmed" if week_summary_confirmed else "system_draft",
                sources=["CEO 已确认"] if week_summary_confirmed else _strategic_unique_non_empty([
                    recent_meeting.title if recent_meeting else None,
                    recent_analysis.question if recent_analysis else None,
                    workspace.goals[0].title if workspace.goals else None,
                ])[:3],
            ),
            mainContradiction=StrategicJudgmentRecord(
                value=main_contradiction_confirmed or contradiction_auto,
                status="confirmed" if main_contradiction_confirmed else ("waiting" if readiness_status == "insufficient" else "system_draft"),
                sources=["CEO 已确认"] if main_contradiction_confirmed else contradiction_sources[:3],
            ),
            coreBreakthrough=StrategicJudgmentRecord(
                value=core_breakthrough_confirmed or core_breakthrough_auto,
                status="confirmed" if core_breakthrough_confirmed else ("waiting" if readiness_status == "insufficient" else "system_draft"),
                sources=["CEO 已确认"] if core_breakthrough_confirmed else core_breakthrough_sources[:3],
            ),
            focusItems=snapshot_focus_items[:3] if snapshot_focus_items else focus_items_auto,
            focusStatus="confirmed" if snapshot_focus_items else ("waiting" if readiness_status == "insufficient" else "system_draft"),
            freshness=freshness,
        )

        event_line_memory_rows = {
            str(row["event_line_id"]): row
            for row in state.db.fetchall(
                "SELECT * FROM event_line_memory_snapshots WHERE event_line_id IN ({})".format(",".join("?" for _ in event_line_ids)) if event_line_ids else "SELECT * FROM event_line_memory_snapshots WHERE 0",
                tuple(event_line_ids),
            )
        }

        def _strategic_line_id(title: str, module_name: str | None = None, flow_name: str | None = None) -> str:
            payload = "::".join(
                [
                    client_id,
                    (module_name or "").strip(),
                    (flow_name or "").strip(),
                    title.strip(),
                ]
            )
            return f"sl_{hashlib.sha1(payload.encode('utf-8')).hexdigest()[:12]}"

        strategic_lines: list[StrategicLineRecord] = []
        if workspace.projectModules:
            for module in workspace.projectModules:
                module_tasks = [item for item in client_tasks if item.projectModuleId == module.id or item.projectModuleName == module.name]
                module_summary = next((item for item in review_event_line_summaries if item.moduleName == module.name), None)
                flow = next((item for item in workspace.projectFlows if item.moduleId == module.id), None)
                module_event_line_id = next((item.eventLineId for item in module_tasks if item.eventLineId), None)
                memory_row = event_line_memory_rows.get(module_event_line_id or "") if module_event_line_id else None
                momentum: Literal["加码", "稳住", "收口", "暂停"] = (
                    "暂停"
                    if module_summary and module_summary.status == "blocked"
                    else "收口"
                    if module_tasks and all(item.status == "done" for item in module_tasks)
                    else "加码"
                    if module_summary and module_summary.predictionReadiness == "strong_forecast" and review_opportunity_cards
                    else "稳住"
                )
                strategic_lines.append(
                    StrategicLineRecord(
                        id=_strategic_line_id(module.name, module.name, flow.name if flow else None),
                        title=module.name,
                        summary=_strategic_first_non_empty([
                            str(memory_row["current_work"]).strip() if memory_row and memory_row["current_work"] else None,
                            module.goal,
                            module.description,
                            module_summary.whatThisLineIs if module_summary else None,
                        ], "这条业务线需要进一步明确它到底在推进什么。"),
                        module=module.name,
                        flow=flow.name if flow else None,
                        stage=str(memory_row["current_stage"]).strip() if memory_row and memory_row["current_stage"] else (module_tasks[0].projectContext.stage if module_tasks and module_tasks[0].projectContext else None),
                        blocker=_strategic_first_non_empty([
                            str(memory_row["current_blocker"]).strip() if memory_row and memory_row["current_blocker"] else None,
                            module_summary.mainBlocker if module_summary else None,
                            *[item.projectContext.currentBlocker for item in module_tasks if item.projectContext],
                        ], "当前还没有稳定识别到这条线最主要的阻塞。"),
                        decision=_strategic_first_non_empty([
                            str(memory_row["recent_decision"]).strip() if memory_row and memory_row["recent_decision"] else None,
                            module_summary.currentState if module_summary else None,
                            *[item.eventLineName for item in module_tasks],
                        ], "最近关键决策仍待补充。"),
                        nextStep=_strategic_first_non_empty([
                            str(memory_row["next_step"]).strip() if memory_row and memory_row["next_step"] else None,
                            module_summary.nextCriticalMove if module_summary else None,
                            *[item.projectContext.nextAction for item in module_tasks if item.projectContext],
                        ], "先把下一步动作拆清楚，再进入会谈推进。"),
                        momentum=momentum,
                        evidence=_strategic_unique_non_empty([
                            module_summary.whatHappenedThisWeek if module_summary else None,
                            flow.description if flow else None,
                            *(_parse_json_list(memory_row["evidence_refs_json"])[:2] if memory_row and memory_row["evidence_refs_json"] else []),
                            module_tasks[0].title if module_tasks else None,
                            module_tasks[1].title if len(module_tasks) > 1 else None,
                        ])[:3],
                        memoryConfidence=float(memory_row["confidence"] or 0.0) if memory_row else None,
                        predictionReadiness=float(memory_row["prediction_readiness"] or 0.0) if memory_row else None,
                        clarificationNeeds=_parse_json_list(memory_row["clarification_needs_json"]) if memory_row and memory_row["clarification_needs_json"] else [],
                    )
                )
        else:
            for item in review_event_line_summaries[:6]:
                memory_row = event_line_memory_rows.get(item.eventLineId)
                strategic_lines.append(
                    StrategicLineRecord(
                        id=_strategic_line_id(item.title, item.moduleName, item.flowName),
                        title=item.title,
                        summary=_strategic_first_non_empty([
                            str(memory_row["current_work"]).strip() if memory_row and memory_row["current_work"] else None,
                            item.whatThisLineIs,
                        ], "当前还缺稳定战略线。"),
                        module=item.moduleName,
                        flow=item.flowName,
                        stage=str(memory_row["current_stage"]).strip() if memory_row and memory_row["current_stage"] else None,
                        blocker=_strategic_first_non_empty([
                            str(memory_row["current_blocker"]).strip() if memory_row and memory_row["current_blocker"] else None,
                            item.mainBlocker,
                        ], "当前阻塞仍待澄清。"),
                        decision=_strategic_first_non_empty([
                            str(memory_row["recent_decision"]).strip() if memory_row and memory_row["recent_decision"] else None,
                            item.currentState,
                        ], "当前状态仍待澄清。"),
                        nextStep=_strategic_first_non_empty([
                            str(memory_row["next_step"]).strip() if memory_row and memory_row["next_step"] else None,
                            item.nextCriticalMove,
                        ], "先补下一步动作。"),
                        momentum="暂停" if item.status == "blocked" else "收口" if item.status == "done" else "加码" if item.predictionReadiness == "strong_forecast" else "稳住",
                        evidence=_strategic_unique_non_empty([
                            item.whatHappenedThisWeek,
                            *item.evidencePreview[:2],
                            *(_parse_json_list(memory_row["evidence_refs_json"])[:2] if memory_row and memory_row["evidence_refs_json"] else []),
                        ])[:3],
                        memoryConfidence=float(memory_row["confidence"] or 0.0) if memory_row else None,
                        predictionReadiness=float(memory_row["prediction_readiness"] or 0.0) if memory_row else None,
                        clarificationNeeds=_parse_json_list(memory_row["clarification_needs_json"]) if memory_row and memory_row["clarification_needs_json"] else [],
                    )
                )
        if not strategic_lines:
            for event_line in linked_event_lines[:6]:
                memory_row = event_line_memory_rows.get(event_line.id)
                momentum: Literal["加码", "稳住", "收口", "暂停"] = (
                    "暂停"
                    if event_line.status == "blocked"
                    else "收口"
                    if event_line.status == "done"
                    else "加码"
                    if memory_row and float(memory_row["prediction_readiness"] or 0.0) >= 0.72
                    else "稳住"
                )
                strategic_lines.append(
                    StrategicLineRecord(
                        id=_strategic_line_id(event_line.name, event_line.primaryClientName, None),
                        title=event_line.name,
                        summary=_strategic_first_non_empty(
                            [
                                str(memory_row["current_work"]).strip() if memory_row and memory_row["current_work"] else None,
                                event_line.summary,
                                event_line.intent,
                            ],
                            "当前还缺这条业务线的稳定摘要。",
                        ),
                        module=event_line.primaryClientName,
                        flow=None,
                        stage=str(memory_row["current_stage"]).strip() if memory_row and memory_row["current_stage"] else event_line.stage,
                        blocker=_strategic_first_non_empty(
                            [
                                str(memory_row["current_blocker"]).strip() if memory_row and memory_row["current_blocker"] else None,
                                event_line.currentBlocker,
                            ],
                            "当前阻塞仍待澄清。",
                        ),
                        decision=_strategic_first_non_empty(
                            [
                                str(memory_row["recent_decision"]).strip() if memory_row and memory_row["recent_decision"] else None,
                                event_line.recentDecision,
                            ],
                            "最近关键决策仍待补充。",
                        ),
                        nextStep=_strategic_first_non_empty(
                            [
                                str(memory_row["next_step"]).strip() if memory_row and memory_row["next_step"] else None,
                                event_line.nextStep,
                            ],
                            "先补下一步动作。",
                        ),
                        momentum=momentum,
                        evidence=_strategic_unique_non_empty(
                            [
                                *(_parse_json_list(memory_row["evidence_refs_json"])[:3] if memory_row and memory_row["evidence_refs_json"] else []),
                                event_line.intent,
                            ]
                        )[:3],
                        memoryConfidence=float(memory_row["confidence"] or 0.0) if memory_row else None,
                        predictionReadiness=float(memory_row["prediction_readiness"] or 0.0) if memory_row else None,
                        clarificationNeeds=_parse_json_list(memory_row["clarification_needs_json"]) if memory_row and memory_row["clarification_needs_json"] else [],
                    )
                )

        linked_event_line_memories = [
            response.eventLineMemorySnapshot
            for line in linked_event_lines[:8]
            for response in [get_event_line_memory_response(state.db, line.id)]
            if response.eventLineMemorySnapshot is not None
        ]

        direction_calibrated = bool(workspace.goals or strategic_lines or (review_analysis and review_analysis.nextWeekFocus))
        direction_score = (30 if workspace.goals else 0) + (25 if strategic_lines else 0) + (25 if any(task.projectContext and task.projectContext.currentFocus for task in client_tasks) else 0) + (20 if review_analysis and review_analysis.nextWeekFocus else 0)
        carrying_calibrated = bool(workspace.projectModules or workspace.projectFlows or analyzable_tasks)
        carrying_score = (25 if workspace.projectModules else 0) + (25 if workspace.projectFlows else 0) + (20 if any(task.ownerName.strip() for task in client_tasks) else 0) + (30 if len(blocked_tasks) <= max(1, (len(client_tasks) + 1) // 2) else 10)
        collaboration_calibrated = bool(workspace.meetings or any(task.collaborators for task in client_tasks) or review_risk_cards)
        collaboration_score = (25 if workspace.meetings else 0) + (25 if any(task.collaborators for task in client_tasks) else 0) + (10 if any(item.riskType == "collaboration_friction" for item in review_risk_cards) else 25) + (25 if review_event_line_summaries else 10)
        decision_calibrated = bool(workspace.meetings or review_risk_cards or review_completeness or workspace.analysisRuns)
        decision_score = (25 if recent_meeting else 0) + (10 if any(item.riskType == "decision_lag" for item in review_risk_cards) else 30) + (25 if any(item.status in {"forecast_ready", "high_confidence"} for item in review_completeness) else 10) + (20 if recent_analysis else 10)
        deposition_calibrated = bool(workspace.dnaModules or workspace.documentCards or workspace.analysisRuns or client_tasks)
        deposition_score = (25 if any(item.hasDocument for item in workspace.dnaModules) else 0) + (25 if len(workspace.documentCards) >= 3 else 15 if workspace.documentCards else 0) + (20 if workspace.analysisRuns else 0) + (30 if workspace.knowledgeStatus and workspace.knowledgeStatus.reviewPendingDocuments == 0 else 15)

        health = [
            StrategicHealthLineRecord(
                key="direction",
                title="方向健康",
                status=_strategic_health_status(direction_score, calibrated=direction_calibrated),
                trend="正在收敛" if workspace.goals and strategic_lines else "待校准",
                summary="目标、任务与当前业务主线已经开始对齐。" if workspace.goals else "当前业务仍缺少足够清晰的阶段锚点。",
                evidence=_strategic_unique_non_empty([workspace.goals[0].title if workspace.goals else None, strategic_lines[0].title if strategic_lines else None, review_analysis.nextWeekFocus[0] if review_analysis and review_analysis.nextWeekFocus else None])[:3],
            ),
            StrategicHealthLineRecord(
                key="carrying",
                title="承接健康",
                status=_strategic_health_status(carrying_score, calibrated=carrying_calibrated),
                trend="结构在变清楚" if workspace.projectModules and workspace.projectFlows else "承接位不足",
                summary="业务已经开始落到模块与流程。" if workspace.projectModules else "还缺稳定模块，经营动作容易只停留在会议和任务表层。",
                evidence=_strategic_unique_non_empty([workspace.projectModules[0].name if workspace.projectModules else None, workspace.projectFlows[0].name if workspace.projectFlows else None, active_task.title if active_task else None])[:3],
            ),
            StrategicHealthLineRecord(
                key="collaboration",
                title="协同健康",
                status=_strategic_health_status(collaboration_score, calibrated=collaboration_calibrated),
                trend="有会谈节律" if workspace.meetings else "协同节律偏弱",
                summary="当前已有跨人协作推进。" if any(task.collaborators for task in client_tasks) else "关键事项仍更多依赖单点推进。",
                evidence=_strategic_unique_non_empty([recent_meeting.title if recent_meeting else None, next((task.title for task in client_tasks if task.collaborators), None), next((item.statement for item in review_risk_cards if item.riskType == "collaboration_friction"), None)])[:3],
            ),
            StrategicHealthLineRecord(
                key="decision",
                title="决策健康",
                status=_strategic_health_status(decision_score, calibrated=decision_calibrated),
                trend="拍板偏慢" if any(item.riskType == "decision_lag" for item in review_risk_cards) else "决策节奏可用",
                summary="存在关键问题未被及时拍板。" if any(item.riskType == "decision_lag" for item in review_risk_cards) else "当前尚未观察到显著决策拖延。",
                evidence=_strategic_unique_non_empty([next((item.statement for item in review_risk_cards if item.riskType == "decision_lag"), None), recent_meeting.title if recent_meeting else None, recent_analysis.question if recent_analysis else None])[:3],
            ),
            StrategicHealthLineRecord(
                key="deposition",
                title="沉淀健康",
                status=_strategic_health_status(deposition_score, calibrated=deposition_calibrated),
                trend="证据在变厚" if workspace.documentCards else "底稿偏薄",
                summary="资料、分析与 DNA 已开始形成可复用底座。" if workspace.documentCards else "还没有形成足够稳定的资料沉淀与分析层。",
                evidence=_strategic_unique_non_empty([next((item.title for item in workspace.dnaModules if item.hasDocument), None), workspace.documentCards[0].title if workspace.documentCards else None, recent_analysis.question if recent_analysis else None])[:3],
            ),
        ]

        ambiguity_rows = state.db.fetchall(
            """
            SELECT a.raw_text, m.title AS meeting_title
            FROM ambiguities a
            JOIN meetings m ON m.id = a.meeting_id
            WHERE m.client_id = ? AND a.status = 'pending'
            ORDER BY m.updated_at DESC
            LIMIT 4
            """,
            (client_id,),
        )
        clarify_items_records = [
            *[
                StrategicChecklistItemRecord(
                    title=f"{item.title}：{slot.label}待澄清",
                    detail=slot.summary,
                    source="周复盘 / 事件线完整度",
                    priority="high" if slot.key in {"blocker", "next_action"} else "medium",
                )
                for item in review_completeness
                for slot in item.slots
                if slot.recommendedFix == "clarify_now" and slot.coverage != "full"
            ],
            *[
                StrategicChecklistItemRecord(
                    title=f"{module.title}：补充关键背景",
                    detail=missing,
                    source="客户 DNA 缺口",
                    priority="medium",
                )
                for module in workspace.dnaModules
                for missing in module.missingInfo
            ],
        ]
        if not workspace.projectModules:
            clarify_items_records.append(StrategicChecklistItemRecord(title="当前业务线还缺稳定模块定义", detail="建议在周会上先确认：这条业务到底按哪几个模块来看，而不是继续按零散事项推进。", source="项目结构缺口", priority="high"))
        if background_thin_tasks:
            clarify_items_records.append(StrategicChecklistItemRecord(title="关系推进任务需要补背景说明", detail=f"当前有 {len(background_thin_tasks)} 条任务只写了动作名。请在任务描述里补对象是谁、当前关系、这次动作想推动什么。", source="任务描述缺口", priority="high"))
        clarify_items_records = list({f"{item.title}::{item.detail}": item for item in clarify_items_records}.values())[:8]

        decision_items_records = [
            *([StrategicChecklistItemRecord(title="拍板：本周期这条业务到底只抓什么", detail="先定一个主问题，再决定这条业务按哪些模块和流程推进；否则战略页只能继续堆事实，不能形成判断。", source="经营结构缺口", priority="high")] if readiness_status == "insufficient" else []),
            *[
                StrategicChecklistItemRecord(
                    title=f"拍板：{item.title}",
                    detail=item.statement,
                    source="周复盘风险卡",
                    priority=_strategic_priority(item.probability),
                )
                for item in review_risk_cards
                if item.riskType in {"decision_lag", "goal_drift"}
            ],
            *[
                StrategicChecklistItemRecord(
                    title=f"拍板：{_strategic_truncate(str(row['meeting_title']), 20)}",
                    detail=str(row["raw_text"]),
                    source="会议待澄清项",
                    priority="high",
                )
                for row in ambiguity_rows
            ],
            *[
                StrategicChecklistItemRecord(
                    title=f"拍板：{task.title}",
                    detail=task.projectContext.currentBlocker if task.projectContext and task.projectContext.currentBlocker else (task.desc or "这条推进链还缺关键拍板。"),
                    source="任务推进阻塞",
                    priority="medium",
                )
                for task in blocked_tasks[:2]
            ],
        ]
        decision_items_records = list({f"{item.title}::{item.detail}": item for item in decision_items_records}.values())[:5]

        material_items_records = [
            *[
                StrategicChecklistItemRecord(
                    title=f"{item.title}：补充资料",
                    detail=slot.summary,
                    source="周复盘资料缺口",
                    priority="medium",
                )
                for item in review_completeness
                for slot in item.slots
                if slot.recommendedFix == "upload_docs" and slot.coverage != "full"
            ],
            *([StrategicChecklistItemRecord(title="复核待确认资料", detail=f"当前仍有 {workspace.knowledgeStatus.reviewPendingDocuments} 份资料待复核，会影响经营判断置信度。", source="知识状态", priority="medium")] if workspace.knowledgeStatus and workspace.knowledgeStatus.reviewPendingDocuments else []),
            *([StrategicChecklistItemRecord(title="为关系推进任务补背景", detail="以后凡是“吃饭 / 见面 / 介绍 / 合作推进”类任务，都要在描述里补对象背景、合作关系和预期结果。", source="任务描述规范", priority="high")] if background_thin_tasks else []),
        ]
        material_items_records = list({f"{item.title}::{item.detail}": item for item in material_items_records}.values())[:6]

        asset_candidates = [
            *[
                StrategicAssetCandidateRecord(
                    title=item.title,
                    source="资料卡",
                    summary=_strategic_truncate(item.shortSummary or item.retrievalSummary or item.summary, 80),
                    nextAction=item.coreQuestions[0] if item.coreQuestions else item.queryHints[0] if item.queryHints else "继续抽出一版稳定摘要或模板。",
                )
                for item in workspace.documentCards[:3]
            ],
            *[
                StrategicAssetCandidateRecord(
                    title=_strategic_truncate(item.question, 42),
                    source="分析运行",
                    summary=f"当前状态 {item.status}，已有 {len(item.evidenceSummary.evidenceList)} 条证据命中。",
                    nextAction="从这次分析里抽一版可复用的方法、框架或顾问判断模板。",
                )
                for item in workspace.analysisRuns[:2]
            ],
            *([StrategicAssetCandidateRecord(title=recent_meeting.title, source="会议", summary=f"最近会谈更新于 {_strategic_format_date_label(recent_meeting.updatedAt)}。", nextAction="把会谈里有效的问题结构沉淀成下次可复用的共创议程。")] if recent_meeting else []),
        ]
        asset_candidates = list({item.title: item for item in asset_candidates}.values())[:6]

        facts_group_items = [
            *([StrategicChecklistItemRecord(title="同步最近会议变化", detail=f"最近会议是“{recent_meeting.title}”，更新时间 {_strategic_format_date_label(recent_meeting.updatedAt)}。", source="会议", priority="low")] if recent_meeting else []),
            *([StrategicChecklistItemRecord(title="同步最近分析主题", detail=f"最近分析围绕“{_strategic_truncate(recent_analysis.question, 36)}”，状态 {recent_analysis.status}。", source="分析运行", priority="low")] if recent_analysis else []),
            *([StrategicChecklistItemRecord(title="同步当前业务锚点", detail=f"当前最明确的目标是“{workspace.goals[0].title}”。", source="目标", priority="low")] if workspace.goals else []),
            *([StrategicChecklistItemRecord(title="同步当前推进动作", detail=f"当前最活跃的动作是“{active_task.title}”，状态 {active_task.status}。", source="任务", priority="low")] if active_task else []),
        ][:5]
        observe_group_items = [
            *[
                StrategicChecklistItemRecord(
                    title=f"观察：{item.title}",
                    detail=item.ifIgnored,
                    source="风险卡",
                    priority=_strategic_priority(item.probability),
                )
                for item in review_risk_cards[:2]
            ],
            *[
                StrategicChecklistItemRecord(
                    title=f"观察：{item.title}",
                    detail=item.recommendedAmplifier,
                    source="机会卡",
                    priority="medium" if item.confidence == "high" else "low",
                )
                for item in review_opportunity_cards[:2]
            ],
        ][:5]
        meeting_pack_groups = [
            StrategicChecklistGroupRecord(key="facts", title="先同步的事实", description="先把真正改变业务状态的事实说清楚，再进入判断。", items=facts_group_items),
            StrategicChecklistGroupRecord(key="clarify", title="必须澄清的问题", description="这些未知项不问清楚，会直接拖低判断质量。", items=clarify_items_records),
            StrategicChecklistGroupRecord(key="decision", title="必须拍板的事项", description="这些地方不拍板，下周推进大概率会继续卡住。", items=decision_items_records),
            StrategicChecklistGroupRecord(key="material", title="必须补的资料", description="这些材料不是锦上添花，而是判断证据缺口。", items=material_items_records),
            StrategicChecklistGroupRecord(key="asset", title="必须沉淀的资产", description="这条业务除了要往前推，也要给益语留下可复用能力。", items=[StrategicChecklistItemRecord(title=item.title, detail=item.nextAction, source=item.source, priority="medium" if item.source != "分析运行" else "high") for item in asset_candidates[:3]]),
            StrategicChecklistGroupRecord(key="observe", title="下周观察点", description="这些信号会决定局面是在转好，还是在继续失衡。", items=observe_group_items),
        ]
        meeting_pack_agenda = _strategic_unique_non_empty([
            f"先对齐主矛盾：{headline.mainContradiction.value}",
            f"再确认核心突破：{headline.coreBreakthrough.value}",
            f"重点澄清：{clarify_items_records[0].title}" if clarify_items_records else None,
            f"重点拍板：{decision_items_records[0].title}" if decision_items_records else None,
            "补关系推进任务背景：对象是谁、当前关系、推进目标分别是什么" if background_thin_tasks else None,
        ])[:4]

        two_week_changes: list[StrategicChangePointRecord] = []
        if readiness_status == "insufficient":
            two_week_changes.append(StrategicChangePointRecord(title="当前不做经营预测", summary="当前先不输出经营预测。现在最有价值的不是猜下周会怎样，而是把能支撑判断的结构化信号补出来。", confidence="等待更多信号", signals=readiness_gaps[:3]))
        else:
            if review_risk_cards:
                risk = review_risk_cards[0]
                two_week_changes.append(StrategicChangePointRecord(title=f"如果 {risk.title} 不补", summary=f"如果当前阻塞不拆，2 周内最可能 {risk.ifIgnored}", confidence="中等置信" if risk.probability == "high" else "保守判断", signals=_strategic_unique_non_empty([risk.whyNow, *risk.triggerSignals[:2]])[:3]))
            if review_opportunity_cards:
                opportunity = review_opportunity_cards[0]
                two_week_changes.append(StrategicChangePointRecord(title=f"如果 {opportunity.title} 被放大", summary=f"如果相关动作被拍板并继续推进，这条线最可能转向 {opportunity.upside}", confidence="中等置信" if opportunity.confidence == "high" else "观察中", signals=_strategic_unique_non_empty([*opportunity.supportingSignals[:2], opportunity.recommendedAmplifier])[:3]))
            if workspace.knowledgeStatus and workspace.knowledgeStatus.reviewPendingDocuments:
                two_week_changes.append(StrategicChangePointRecord(title="如果关键资料继续不补", summary=f"如果这 {workspace.knowledgeStatus.reviewPendingDocuments} 份待复核资料继续悬空，接下来两周很多判断都会停在保守层。", confidence="高依赖资料", signals=["待复核资料数量偏多", "证据链稳定性不足"]))
        two_week_changes = two_week_changes[:3]

        return StrategicCockpitSnapshotRecord(
            clientId=client_id,
            clientName=client.name,
            clientTagline=" · ".join(_strategic_unique_non_empty([client.type, client.domain])) or "业务发展驾驶台",
            stageLabel=notebook_stage or client.stage or "待判断",
            permission=permission,
            readiness=readiness,
            headline=headline,
            health=health,
            strategicLines=strategic_lines[:6],
            twoWeekChanges=two_week_changes,
            pendingDecisions=decision_items_records[:3],
            pendingMaterials=material_items_records[:3],
            meetingPackDraft=StrategicMeetingPackDraftRecord(title=f"{client.name} 周盘点会", agenda=meeting_pack_agenda, groups=meeting_pack_groups),
            evidencePreview=StrategicEvidencePreviewRecord(summary=dossier_summary, cards=dossier_cards, boundaries=dossier_boundaries, keyFacts=key_facts, keyWarnings=key_warnings),
            assetCandidates=asset_candidates,
            notebookSummary=notebook_summary,
            memoryStatus=memory_status,
            linkedEventLineMemories=linked_event_line_memories,
        )

    def save_strategic_cockpit_snapshot(client_id: str, payload: StrategicCockpitConfirmPayload, session_user: SessionUserRecord) -> None:
        existing = _load_strategic_snapshot_row(client_id)
        focus_items = _strategic_unique_non_empty([item for item in payload.focusItems])[:3]
        week_summary = (payload.weekSummary or "").strip()
        main_contradiction = (payload.mainContradiction or "").strip()
        core_breakthrough = (payload.coreBreakthrough or "").strip()
        if not week_summary and not main_contradiction and not core_breakthrough and not focus_items:
            state.db.execute("DELETE FROM strategic_cockpit_snapshots WHERE client_id = ?", (client_id,))
            return
        created_at = str(existing["created_at"]) if existing and existing["created_at"] else now_iso()
        state.db.execute(
            """
            INSERT INTO strategic_cockpit_snapshots(
                client_id, week_summary, main_contradiction, core_breakthrough, focus_items_json,
                confirmed_by_user_id, confirmed_by_user_name, created_at, updated_at
            ) VALUES(?, ?, ?, ?, ?, ?, ?, ?, ?)
            ON CONFLICT(client_id) DO UPDATE SET
                week_summary = excluded.week_summary,
                main_contradiction = excluded.main_contradiction,
                core_breakthrough = excluded.core_breakthrough,
                focus_items_json = excluded.focus_items_json,
                confirmed_by_user_id = excluded.confirmed_by_user_id,
                confirmed_by_user_name = excluded.confirmed_by_user_name,
                updated_at = excluded.updated_at
            """,
            (
                client_id,
                week_summary,
                main_contradiction,
                core_breakthrough,
                to_json(focus_items),
                session_user.id,
                session_user.fullName,
                created_at,
                now_iso(),
            ),
        )

    def _render_strategic_meeting_pack_text(snapshot: StrategicCockpitSnapshotRecord) -> str:
        lines: list[str] = [snapshot.meetingPackDraft.title, ""]
        if snapshot.meetingPackDraft.agenda:
            lines.append("建议议程")
            for item in snapshot.meetingPackDraft.agenda:
                lines.append(f"- {item}")
            lines.append("")
        for group in snapshot.meetingPackDraft.groups:
            lines.append(group.title)
            lines.append(group.description)
            for item in group.items:
                lines.append(f"- {item.title}")
                lines.append(f"  说明：{item.detail}")
                lines.append(f"  来源：{item.source}")
            lines.append("")
        return "\n".join(lines).strip()

    def _require_strategic_ceo() -> SessionUserRecord:
        session_user = get_cached_session_user()
        if session_user is None:
            session_user = require_session_user()
        org_model_profile = _load_org_model_profile_safe()
        leader_user_id = org_model_profile.organization.leaderUserId if org_model_profile else None
        if not leader_user_id:
            raise HTTPException(status_code=403, detail="请先在组织设置中确认 CEO 账号")
        if session_user.id != leader_user_id:
            raise HTTPException(status_code=403, detail="当前页面只有 CEO 可以确认经营判断")
        return session_user

    def _create_strategic_meeting_pack(client_id: str) -> MeetingDetail:
        snapshot = build_strategic_cockpit_snapshot(client_id)
        meeting_id = new_id("meeting")
        timestamp = now_iso()
        notes_text = _render_strategic_meeting_pack_text(snapshot)
        state.db.execute(
            """
            INSERT INTO meetings(id, client_id, title, stage, scheduled_at, transcript_text, notes, created_at, updated_at)
            VALUES(?, ?, ?, 'prepared', NULL, '', ?, ?, ?)
            """,
            (meeting_id, client_id, snapshot.meetingPackDraft.title, notes_text, timestamp, timestamp),
        )
        for index, agenda_item in enumerate(snapshot.meetingPackDraft.agenda[:8]):
            state.db.execute(
                "INSERT INTO agenda_items(id, meeting_id, title, description, sort_order) VALUES(?, ?, ?, ?, ?)",
                (new_id("agenda"), meeting_id, agenda_item[:80], "战略陪伴周会清单议程", index),
            )
        state.db.execute(
            "INSERT INTO meeting_sources(id, meeting_id, title, content_text, created_at) VALUES(?, ?, ?, ?, ?)",
            (new_id("ms"), meeting_id, "战略陪伴周会清单草案", notes_text, timestamp),
        )
        log_activity("meeting.prepare_from_strategic_cockpit", "meeting", meeting_id, {"clientId": client_id, "agendaItems": len(snapshot.meetingPackDraft.agenda)})
        return build_meeting_detail(meeting_id)

    def _build_strategic_payload_from_meeting(client_id: str, meeting: MeetingDetail) -> StrategicCockpitConfirmPayload:
        current_snapshot = build_strategic_cockpit_snapshot(client_id)
        note_lines = [
            line.strip(" -\t")
            for line in (meeting.notes or "").splitlines()
            if line.strip() and "建议议程" not in line and "来源：" not in line and "说明：" not in line and line.strip() not in {
                meeting.title,
                "先同步的事实",
                "必须澄清的问题",
                "必须拍板的事项",
                "必须补的资料",
                "必须沉淀的资产",
                "下周观察点",
            }
        ]
        week_summary = _strategic_first_non_empty(
            [
                meeting.decisions[0].summary if meeting.decisions else None,
                note_lines[0] if note_lines else None,
                current_snapshot.headline.weekSummary.value,
            ],
            current_snapshot.headline.weekSummary.value,
        )
        main_contradiction = _strategic_first_non_empty(
            [
                meeting.ambiguities[0].rawText if meeting.ambiguities else None,
                meeting.risks[0].summary if meeting.risks else None,
                current_snapshot.headline.mainContradiction.value,
            ],
            current_snapshot.headline.mainContradiction.value,
        )
        core_breakthrough = _strategic_first_non_empty(
            [
                meeting.actionItems[0].title if meeting.actionItems else None,
                meeting.decisions[0].summary if meeting.decisions else None,
                current_snapshot.headline.coreBreakthrough.value,
            ],
            current_snapshot.headline.coreBreakthrough.value,
        )
        focus_items = _strategic_unique_non_empty(
            [
                *[item.title for item in meeting.actionItems[:3]],
                *[item.title for item in meeting.agendaItems[:2]],
                *current_snapshot.headline.focusItems,
            ]
        )[:3]
        return StrategicCockpitConfirmPayload(
            weekSummary=week_summary,
            mainContradiction=main_contradiction,
            coreBreakthrough=core_breakthrough,
            focusItems=focus_items,
        )

    def fetch_chat_thread_for_client(client_id: str, thread_id: str) -> ChatThread:
        row = state.db.fetchone(
            "SELECT * FROM chat_threads WHERE id = ? AND client_id = ?",
            (thread_id, client_id),
        )
        if not row:
            raise HTTPException(status_code=404, detail="Chat thread not found")
        return ChatThread(
            id=str(row["id"]),
            clientId=str(row["client_id"]),
            title=str(row["title"]),
            createdAt=str(row["created_at"]),
            updatedAt=str(row["updated_at"]),
        )

    def list_chat_messages_for_thread(client_id: str, thread_id: str) -> list[ChatMessageRecord]:
        fetch_chat_thread_for_client(client_id, thread_id)
        return [
            build_chat_message(row)
            for row in state.db.fetchall(
                """
                SELECT m.*
                FROM chat_messages m
                JOIN chat_threads t ON t.id = m.thread_id
                WHERE t.client_id = ? AND m.thread_id = ?
                ORDER BY m.created_at ASC
                """,
                (client_id, thread_id),
            )
        ]

    def build_chat_message(row) -> ChatMessageRecord:
        structured = from_json(row["structured_data_json"], None)
        data = AiStructuredResponse(**structured) if structured else None
        evidence_data = from_json(row["evidence_json"], [])
        evidence = [EvidenceItem(**item) for item in evidence_data] if isinstance(evidence_data, list) else []
        retrieval_summary = from_json(row["retrieval_summary_json"], {})
        timing = from_json(row["timing_json"], {})
        return ChatMessageRecord(
            id=str(row["id"]),
            threadId=str(row["thread_id"]),
            role=str(row["role"]),  # type: ignore[arg-type]
            content=str(row["content"]),
            createdAt=str(row["created_at"]),
            status=str(row["status"]),  # type: ignore[arg-type]
            modelRoute=str(row["model_route"]) if row["model_route"] else None,
            llmInvoked=bool(row["llm_invoked"]),
            providerUsed=str(row["provider_used"]) if row["provider_used"] else None,
            answerMode=str(row["answer_mode"]) if row["answer_mode"] else None,  # type: ignore[arg-type]
            evidenceStatus=str(row["evidence_status"]) if row["evidence_status"] else None,  # type: ignore[arg-type]
            failureReason=str(row["failure_reason"]) if row["failure_reason"] else None,
            timing=timing if isinstance(timing, dict) else {},
            retrievalSummary=retrieval_summary if isinstance(retrieval_summary, dict) else {},
            structuredData=data,
            evidence=evidence,
        )

    def build_meeting_summary(row) -> MeetingSummary:
        return MeetingSummary(
            id=str(row["id"]),
            clientId=str(row["client_id"]),
            title=str(row["title"]),
            stage=str(row["stage"]),  # type: ignore[arg-type]
            scheduledAt=str(row["scheduled_at"]) if row["scheduled_at"] else None,
            updatedAt=str(row["updated_at"]),
        )

    def build_meeting_detail(meeting_id: str) -> MeetingDetail:
        row = state.db.fetchone("SELECT * FROM meetings WHERE id = ?", (meeting_id,))
        if not row:
            raise HTTPException(status_code=404, detail="Meeting not found")
        agenda = [
            AgendaItem(id=str(item["id"]), title=str(item["title"]), description=str(item["description"]))
            for item in state.db.fetchall("SELECT * FROM agenda_items WHERE meeting_id = ? ORDER BY sort_order", (meeting_id,))
        ]
        decisions = [
            DecisionItem(id=str(item["id"]), summary=str(item["summary"]))
            for item in state.db.fetchall("SELECT * FROM decisions WHERE meeting_id = ? ORDER BY created_at", (meeting_id,))
        ]
        action_items = []
        for item in state.db.fetchall("SELECT * FROM action_items WHERE meeting_id = ? ORDER BY created_at", (meeting_id,)):
            action_items.append(
                TaskRecord(
                    id=str(item["id"]),
                    title=str(item["title"]),
                    desc="来自会议抽取的行动项",
                    status="todo" if item["publish_status"] == "published" else "inbox",
                    priority="normal",
                    listId="list-0",
                    listName="会议草稿",
                    listColor="#888681",
                    ddl=str(item["due_date"]),
                    ownerName=str(item["owner_name"]),
                    sourceType="meeting",
                    sourceId=str(meeting_id),
                    tags=[
                        TaskTagRecord(
                            id="tag_meeting_builtin",
                            name="会议",
                            color="#F59E0B",
                            scope="org",
                            ownerUserId=None,
                            createdBy="system",
                            updatedAt=str(item["created_at"]),
                            archivedAt=None,
                        )
                    ],
                    note=None,
                    createdAt=str(item["created_at"]),
                    updatedAt=str(item["created_at"]),
                )
            )
        risks = [
            RiskItem(id=str(item["id"]), summary=str(item["summary"]), severity=str(item["severity"]))  # type: ignore[arg-type]
            for item in state.db.fetchall("SELECT * FROM risks WHERE meeting_id = ? ORDER BY created_at", (meeting_id,))
        ]
        ambiguities = [
            AmbiguityItem(
                id=str(item["id"]),
                rawText=str(item["raw_text"]),
                candidates=_parse_json_list(item["candidates_json"]),
                status=str(item["status"]),  # type: ignore[arg-type]
            )
            for item in state.db.fetchall("SELECT * FROM ambiguities WHERE meeting_id = ? ORDER BY created_at", (meeting_id,))
        ]
        return MeetingDetail(
            **build_meeting_summary(row).model_dump(),
            transcriptText=str(row["transcript_text"]),
            notes=str(row["notes"]),
            agendaItems=agenda,
            decisions=decisions,
            actionItems=action_items,
            risks=risks,
            ambiguities=ambiguities,
        )

    def select_evidence(client_id: str, prompt: str) -> list[EvidenceItem]:
        bundle = build_retrieval_bundle(client_id, prompt)
        evidence = [
            EvidenceItem(
                id=new_id("ev"),
                title=item.title,
                excerpt=item.excerpt,
                sourceType="knowledge_chunk" if item.chunk_id else "knowledge_document",
                documentId=item.knowledge_document_id,
                path=item.path,
                score=item.score,
                coverage=item.coverage,
                sectionLabel=item.section_label,
                retrievalStage={"master_index": "doc_index", "surrogate": "section_index"}.get(item.source_stage, item.source_stage),
                isFallback=item.source_stage in {"master_index", "doc_index"},
                matchedTerms=item.matched_terms,
            )
            for item in bundle.citations
        ]
        for goal in state.db.fetchall("SELECT * FROM goal_records WHERE client_id = ? ORDER BY updated_at DESC LIMIT 1", (client_id,)):
            evidence.append(
                EvidenceItem(
                    id=new_id("ev"),
                    title=f"目标：{goal['title']}",
                    excerpt=f"季度 {goal['quarter']}，进度 {goal['progress']}%，负责人 {goal['owner_name']}",
                    sourceType="goal",
                    coverage=bundle.coverage,
                    matchedTerms=[],
                )
            )
        return evidence[:4]

    def context_summary(client_id: str, prompt: str, evidence: list[EvidenceItem]) -> str:
        client = build_client_summary(client_id)
        dna_count = state.db.scalar("SELECT COUNT(1) AS count FROM dna_terms WHERE client_id = ?", (client_id,))
        dna_doc_count = state.db.scalar("SELECT COUNT(1) AS count FROM client_dna_documents WHERE client_id = ?", (client_id,))
        return (
            f"客户={client.name}/{client.domain}；当前阶段={client.stage}；"
            f"客户DNA文档={dna_doc_count}；组织补充词条={dna_count}"
        )

    def build_chat_answer_context(
        client_id: str,
        prompt: str,
        evidence: list[EvidenceItem],
        retrieval_bundle,
        *,
        memory_background_context: str = "",
    ) -> str:
        retrieval_meta = retrieval_bundle.retrieval_summary if isinstance(retrieval_bundle.retrieval_summary, dict) else {}
        client_dna_context = build_client_dna_context(client_id, prompt)
        preferred_categories = [
            str(item)
            for item in retrieval_meta.get("preferredCategories", [])
            if str(item).strip()
        ] if isinstance(retrieval_meta.get("preferredCategories"), list) else []
        curated_evidence = select_high_signal_evidence(
            evidence,
            limit=72,
            prompt=prompt,
            preferred_categories=preferred_categories,
        )
        evidence_blocks: list[str] = []
        evidence_char_budget = 220000
        evidence_char_used = 0
        for index, item in enumerate(curated_evidence, start=1):
            label = item.title
            if item.sectionLabel:
                label = f"{label} / {item.sectionLabel}"
            excerpt = re.sub(r"\s+", " ", item.excerpt or "").strip()
            block = (
                f"[原始证据 {index}]\n"
                f"标题：{label}\n"
                f"片段：{excerpt[:4800]}"
            ).strip()
            if evidence_blocks and evidence_char_used + len(block) > evidence_char_budget:
                break
            evidence_blocks.append(block)
            evidence_char_used += len(block)
        client_name = build_client_summary(client_id).name
        summary_lines = [
            f"用户问题：{prompt}",
            f"当前对象：{client_name}",
            (
                "请直接阅读下面的原始材料并回答问题。"
                "不要把答案写成资料摘要、证据罗列或系统说明。"
                "允许你基于多条材料共同指向的信号做高强度归纳、抽象和深层判断。"
                "只有材料里不存在的具体事实、数字、人名、时间和身份，不要写成已被证实。"
                "除非用户明确要求简短，否则请尽量讲透。"
            ),
        ]
        if memory_background_context:
            summary_lines.append(memory_background_context)
        if client_dna_context:
            summary_lines.append(client_dna_context)
        if evidence_blocks:
            summary_lines.append(
                "原始证据包（可用于正式判断）：\n"
                + "\n\n".join(evidence_blocks)
            )
        return "\n\n".join(summary_lines).strip()

    def build_retrieval_preview_summary(client_id: str, prompt: str, evidence: list[EvidenceItem], retrieval_bundle) -> str:
        client = build_client_summary(client_id)
        retrieval_meta = retrieval_bundle.retrieval_summary if isinstance(retrieval_bundle.retrieval_summary, dict) else {}
        category_coverage = [str(item) for item in retrieval_meta.get("categoryCoverage", []) if str(item).strip()] if isinstance(retrieval_meta.get("categoryCoverage"), list) else []
        preferred_categories = [
            str(item)
            for item in retrieval_meta.get("preferredCategories", [])
            if str(item).strip()
        ] if isinstance(retrieval_meta.get("preferredCategories"), list) else []
        top_evidence = select_high_signal_evidence(
            evidence,
            limit=3,
            prompt=prompt,
            preferred_categories=preferred_categories or category_coverage,
        )
        if not top_evidence:
            return (
                f"围绕“{prompt}”，当前还没有命中足够可支撑正式判断的原始材料。"
                f"如果继续回答，只能作为基于通用背景的初步判断，不属于基于 {client.name} 原始资料的正式分析。"
            )
        titles = "、".join(dict.fromkeys(item.title for item in top_evidence))
        if is_strategy_analysis_query(prompt):
            dimension_text = "、".join(category_coverage) if category_coverage else "项目与业务、品牌与传播、财务与筹款、组织与战略"
            return (
                f"围绕“{prompt}”，我已经先从 {client.name} 的原始材料里命中了几组最相关的证据，包括：{titles}。"
                f"这些证据大致覆盖 {dimension_text} 等维度。下面会在这些原始证据之上继续组织更完整的战略分析。"
            )
        return (
            f"围绕“{prompt}”，我已经先从 {client.name} 的原始材料里命中了几份最相关的证据，包括：{titles}。"
            "下面会基于这些原始证据继续生成更完整的分析回答。"
        )

    def build_answer_work_trace(prompt: str, evidence: list[EvidenceItem], retrieval_bundle) -> dict[str, object]:
        retrieval_meta = retrieval_bundle.retrieval_summary if isinstance(retrieval_bundle.retrieval_summary, dict) else {}
        preferred_categories = [
            str(item)
            for item in retrieval_meta.get("preferredCategories", [])
            if str(item).strip()
        ] if isinstance(retrieval_meta.get("preferredCategories"), list) else []
        covered_categories = [
            str(item)
            for item in retrieval_meta.get("categoryCoverage", [])
            if str(item).strip()
        ] if isinstance(retrieval_meta.get("categoryCoverage"), list) else []
        strategic_mode = bool(retrieval_meta.get("strategicMode", False))

        def stage_label(stage: str | None) -> str:
            return {
                "doc_index": "文档索引",
                "master_index": "文档索引",
                "section_index": "章节定位",
                "surrogate": "章节定位",
                "background": "背景材料",
                "raw_chunk": "原文片段",
            }.get(stage or "", stage or "资料")

        focus = preferred_categories or covered_categories
        if not focus:
            focus = ["机构定位", "核心业务", "战略张力", "价值落地"]
        web_trail = [
            item
            for item in retrieval_meta.get("webTrail", [])
            if isinstance(item, dict)
        ][:6] if isinstance(retrieval_meta.get("webTrail"), list) else []
        return {
            "note": "这里展示的是本次回答如何利用背景底稿、联网补充和原始证据，不展示模型原始思维全文。",
            "problemFrame": (
                f"围绕“{prompt}”，先用背景底稿和联网补充建立客户语境，再确认原始证据能支撑哪些事实，最后在此基础上形成顾问式判断。"
            ),
            "analysisPlan": (
                f"优先围绕 {'、'.join(focus)} 组织分析，把背景理解、联网补充和原始证据整合成一版顾问式回答。"
                if strategic_mode
                else "优先整理机构定位、核心业务、推进线索和最值得展开的判断，并用原始证据校准判断。"
            ),
            "analysisFocus": focus[:6],
            "backgroundTrail": [
                item
                for item in retrieval_meta.get("backgroundTrail", [])
                if isinstance(item, dict)
            ][:8] if isinstance(retrieval_meta.get("backgroundTrail"), list) else [],
            "materialTrail": [
                {
                    "title": item.title,
                    "stage": stage_label(item.retrievalStage),
                    "sectionLabel": item.sectionLabel,
                    "path": item.path,
                    "excerpt": item.excerpt,
                }
                for item in select_high_signal_evidence(
                    evidence,
                    limit=6,
                    prompt=prompt,
                    preferred_categories=preferred_categories or covered_categories,
                )
            ],
            "clientDnaTrail": [str(item) for item in retrieval_meta.get("clientDnaTrail", []) if str(item).strip()] if isinstance(retrieval_meta.get("clientDnaTrail"), list) else [],
            "webTrail": web_trail,
        }

    def prompt_targets_org_content(prompt: str) -> bool:
        return any(token in prompt for token in ("益语", "你们", "顾问方法", "工作方法", "服务方式"))

    INTRO_QUERY_HINTS = ("介绍", "简介", "概况", "概览", "背景", "定位", "做什么", "业务", "团队", "历史")
    INTRO_PRIORITY_HINTS = (
        "介绍",
        "简介",
        "概览",
        "定位",
        "核心业务",
        "团队",
        "访谈",
        "纪要",
        "理事会",
        "工作坊",
        "战略框架",
        "业务介绍",
        "组织介绍",
    )
    INTRO_NOISE_HINTS = (
        "文件导入",
        "完整解决方案",
        "上传说明",
        "目录重分类",
        "重建知识索引",
        "导入飞书",
        "缓冲池",
        "精简版",
        "完整版",
        "第8稿",
        "第7稿",
        "click to edit master",
        "master title style",
        "工作台",
    )

    def is_intro_profile_query(prompt: str) -> bool:
        normalized = re.sub(r"\s+", "", (prompt or "").lower())
        return any(token in normalized for token in INTRO_QUERY_HINTS)

    IDENTITY_ROLE_TERMS = (
        "创始人",
        "联合创始人",
        "创办人",
        "发起人",
        "负责人",
        "理事长",
        "董事长",
        "秘书长",
        "CEO",
        "主席",
    )

    def is_identity_role_query(prompt: str) -> bool:
        return any(token in prompt for token in IDENTITY_ROLE_TERMS)

    def organization_identity_names(max_items: int = 12) -> list[str]:
        modules = [module for module in list_organization_dna_modules() if module.hasDocument and module.normalizedText.strip()]
        if not modules:
            return []
        text = "\n".join(module.normalizedText[:3200] for module in modules)
        names: list[str] = []
        seen: set[str] = set()
        stopwords = {
            "益语智库",
            "我们",
            "客户",
            "团队",
            "公司",
            "业务",
            "战略",
            "组织",
            "方向",
            "判断",
            "品牌",
            "市场",
            "核心",
            "现阶段",
            "人类团队",
            "人类同事",
            "角色",
            "支点",
            "三位",
        }

        def append_candidate(value: str) -> None:
            candidate = value.strip("：:，,。；;、 ")
            if not re.fullmatch(r"[\u4e00-\u9fff]{2,4}", candidate):
                return
            if candidate in stopwords or candidate in seen:
                return
            seen.add(candidate)
            names.append(candidate)

        for pattern in (
            r"(?:同事有|成员有|核心人类同事有|现阶段.*?有三位)[:：]?\s*([^\n。]{2,40})",
            r"([\u4e00-\u9fff]{2,4})：",
            r"([\u4e00-\u9fff]{2,4})是益语智库",
        ):
            for match in re.finditer(pattern, text):
                if match.lastindex and match.lastindex >= 1:
                    value = match.group(1)
                    if "、" in value or "，" in value or "," in value:
                        for part in re.split(r"[、，,和及\s]+", value):
                            append_candidate(part)
                    else:
                        append_candidate(value)
                if len(names) >= max_items:
                    return names[:max_items]
        return names[:max_items]

    def prompt_identity_role_terms(prompt: str) -> list[str]:
        matched = [token for token in IDENTITY_ROLE_TERMS if token in prompt]
        return matched or ["创始人", "负责人"]

    def evidence_text(item: EvidenceItem) -> str:
        return re.sub(r"\s+", " ", f"{item.title} {item.sectionLabel or ''} {item.excerpt or ''}").strip()

    def evidence_has_explicit_role_binding(item: EvidenceItem, *, prompt: str, names: list[str] | None = None) -> bool:
        text = evidence_text(item)
        role_terms = prompt_identity_role_terms(prompt)
        role_pattern = "|".join(re.escape(token) for token in role_terms)
        if not re.search(role_pattern, text):
            return False
        candidate_names = names or organization_identity_names()
        person_pattern = "|".join(re.escape(name) for name in candidate_names) if candidate_names else r"[\u4e00-\u9fff]{2,4}"
        binding_patterns = (
            rf"(?:{person_pattern}).{{0,10}}(?:是|为|担任|作为)?(?:[^。；，\n]{{0,8}})?(?:{role_pattern})",
            rf"(?:{role_pattern})[:： ]?(?:[^。；，\n]{{0,8}})?(?:{person_pattern})",
        )
        return any(re.search(pattern, text) for pattern in binding_patterns)

    def evidence_mentions_org_identity_name(item: EvidenceItem, names: list[str] | None = None) -> bool:
        org_names = names or organization_identity_names()
        if not org_names:
            return False
        text = evidence_text(item)
        return any(name in text for name in org_names)

    def build_identity_guard_response(client_id: str, prompt: str, evidence: list[EvidenceItem], retrieval_bundle) -> AiStructuredResponse:
        client = build_client_summary(client_id)
        retrieval_meta = retrieval_bundle.retrieval_summary if isinstance(retrieval_bundle.retrieval_summary, dict) else {}
        preferred_categories = [
            str(item)
            for item in retrieval_meta.get("preferredCategories", [])
            if str(item).strip()
        ] if isinstance(retrieval_meta.get("preferredCategories"), list) else []
        top_evidence = select_high_signal_evidence(
            evidence,
            limit=4,
            prompt=prompt,
            preferred_categories=preferred_categories,
        )
        org_names = organization_identity_names()
        mentioned_org_people = [
            name
            for name in org_names
            if any(name in evidence_text(item) for item in top_evidence)
        ]
        title = f"{client.name} 人物角色仍待确认"
        content_lines = [
            f"{title}",
            "",
            f"当前资料不足以直接确认 {client.name} 的{prompt_identity_role_terms(prompt)[0]}是谁。现有命中材料更多是在讨论机构战略、工作坊判断、项目设计或外部顾问参与，并没有出现一条可以直接把具体人名与该角色绑定起来的原文证据。",
        ]
        if mentioned_org_people:
            people_text = "、".join(mentioned_org_people[:3])
            content_lines.append(
                f"值得特别排除的是，当前材料里确实出现了 {people_text} 等益语侧人物，但这些内容体现的是外部顾问、访谈参与者或发言人角色，不能据此推断其就是 {client.name} 的{prompt_identity_role_terms(prompt)[0]}。"
            )
        if top_evidence:
            content_lines.append("")
            content_lines.append("当前更像背景材料而非角色证据的命中文档包括：")
            for index, item in enumerate(top_evidence[:4], start=1):
                label = item.title
                if item.sectionLabel:
                    label = f"{label} / {item.sectionLabel}"
                content_lines.append(f"{index}. {label}")
        judgment = f"当前证据不足以确认 {client.name} 的{prompt_identity_role_terms(prompt)[0]}身份，任何具体人名结论都不可靠。"
        analysis = (
            "这类问题必须依赖明确的人名-角色绑定证据，比如机构介绍、注册资料、正式署名访谈、年报或直接写出“某人是创始人/负责人”的原文。"
            "当前命中的高频发言人与外部顾问材料，不足以支撑这种身份判断。"
        )
        actions = "建议优先补充机构介绍、明确署名访谈、注册/年报资料，或直接检索包含“创始人/负责人/秘书长”等角色词的原文。"
        timeline = "补入显式角色证据后，可立即重新生成更可靠的人物分析。"
        return AiStructuredResponse(
            content="\n".join(content_lines).strip(),
            judgment=judgment,
            analysis=analysis,
            actions=actions,
            timeline=timeline,
        )

    def infer_evidence_category(item: EvidenceItem) -> str | None:
        haystack = " ".join(
            part for part in (
                item.path or "",
                item.title or "",
                item.sectionLabel or "",
                item.excerpt[:180] if item.excerpt else "",
            )
            if part
        )
        for category in ("组织与战略", "项目与业务", "品牌与传播", "财务与筹款", "其他资料"):
            if category in haystack:
                return category
        haystack_lower = haystack.lower()
        keyword_map = {
            "组织与战略": ("战略", "组织", "治理", "团队", "人力", "负责人", "路线图"),
            "项目与业务": ("业务", "项目", "交付", "会员", "产品", "运营", "执行"),
            "品牌与传播": ("品牌", "传播", "媒体", "内容", "活动", "公关"),
            "财务与筹款": ("财务", "筹款", "预算", "募资", "捐赠", "现金流"),
        }
        for category, keywords in keyword_map.items():
            if any(keyword in haystack_lower for keyword in keywords):
                return category
        return None

    def evidence_mentions_service_provider(item: EvidenceItem) -> bool:
        haystack = f"{item.title} {item.excerpt or ''}".lower()
        return any(
            marker in haystack
            for marker in (
                "益语智库",
                "我们不卖",
                "标准答案",
                "长期陪伴",
                "增长式咨询",
                "战略陪伴",
                "顾问方法",
                "导师介绍",
                "我们的业务",
                "我们的服务",
            )
        )

    def evidence_is_intro_noise(item: EvidenceItem) -> bool:
        haystack = " ".join(
            part.lower()
            for part in (
                item.title or "",
                item.path or "",
                item.sectionLabel or "",
                item.excerpt[:220] if item.excerpt else "",
            )
            if part
        )
        return any(marker in haystack for marker in INTRO_NOISE_HINTS)

    def evidence_is_intro_priority(item: EvidenceItem) -> bool:
        haystack = " ".join(
            part.lower()
            for part in (
                item.title or "",
                item.path or "",
                item.sectionLabel or "",
                item.excerpt[:220] if item.excerpt else "",
            )
            if part
        )
        if "核心业务介绍" in haystack or "业务介绍" in haystack or "组织介绍" in haystack:
            return True
        if any(marker in haystack for marker in INTRO_PRIORITY_HINTS):
            return True
        inferred_category = infer_evidence_category(item)
        return inferred_category in {"组织与战略", "项目与业务"}

    def evidence_is_noisy_for_fallback(item: EvidenceItem, *, prompt: str = "") -> bool:
        haystack = f"{item.title} {item.excerpt or ''}".lower()
        if any(
            marker in haystack
            for marker in (
                "click to edit master",
                "master title style",
                "second level third level fourth level",
                "pdf 文档：原文件",
                "总页数:",
                "代理文档",
            )
        ):
            return True
        excerpt = re.sub(r"\s+", " ", item.excerpt or "").strip()
        if re.search(r"[!@#$%^&*()_+=<>]{4,}", excerpt):
            return True
        chinese_chars = len(re.findall(r"[\u4e00-\u9fff]", excerpt))
        latin_tokens = len(re.findall(r"[A-Za-z]{4,}", excerpt))
        if chinese_chars < 18 and latin_tokens >= 8:
            return True
        if not prompt_targets_org_content(prompt) and evidence_mentions_service_provider(item):
            return True
        if is_identity_role_query(prompt):
            org_names = organization_identity_names()
            if evidence_mentions_org_identity_name(item, org_names) and not evidence_has_explicit_role_binding(item, prompt=prompt, names=org_names):
                return True
        return False

    def select_high_signal_evidence(
        evidence: list[EvidenceItem],
        limit: int = 8,
        *,
        prompt: str = "",
        preferred_categories: list[str] | None = None,
    ) -> list[EvidenceItem]:
        intro_mode = is_intro_profile_query(prompt)
        finance_mode = is_finance_query(prompt)
        finance_statement_mode = is_finance_statement_query(prompt)

        def document_key(item: EvidenceItem) -> str:
            base = (item.documentId or item.path or item.title or "").strip().lower()
            return re.sub(r"\s+", " ", base)

        def section_key(item: EvidenceItem) -> str:
            label = (item.sectionLabel or "").strip().lower()
            return re.sub(r"\s+", " ", label)

        def score(item: EvidenceItem) -> float:
            ranking = float(item.score or 0.0)
            if item.retrievalStage == "raw_chunk":
                ranking += 1.2
            if item.retrievalStage == "surrogate":
                ranking -= 0.6
            if item.isFallback:
                ranking -= 0.4
            if item.sectionLabel and item.sectionLabel not in {"概览", "目录索引", "代理文档"}:
                ranking += 0.15
            inferred_category = infer_evidence_category(item)
            if preferred_categories and inferred_category in preferred_categories:
                ranking += 0.22
            if finance_mode:
                if inferred_category == "财务与筹款":
                    ranking += 0.95
                if is_finance_priority_text(item.title, item.sectionLabel, item.excerpt[:320] if item.excerpt else "", item.path):
                    ranking += 0.68
                if finance_statement_mode:
                    if is_finance_statement_priority_text(item.title, item.sectionLabel, item.excerpt[:320] if item.excerpt else "", item.path):
                        ranking += 0.9
                    elif inferred_category != "财务与筹款":
                        ranking -= 0.28
                elif inferred_category != "财务与筹款":
                    ranking -= 0.22
            if is_identity_role_query(prompt):
                org_names = organization_identity_names()
                if evidence_has_explicit_role_binding(item, prompt=prompt, names=org_names):
                    ranking += 2.0
                else:
                    ranking -= 0.8
                if evidence_mentions_org_identity_name(item, org_names) and not evidence_has_explicit_role_binding(item, prompt=prompt, names=org_names):
                    ranking -= 2.4
            if intro_mode:
                if evidence_is_intro_priority(item):
                    ranking += 0.55
                if evidence_is_intro_noise(item):
                    ranking -= 3.2
            if evidence_is_noisy_for_fallback(item, prompt=prompt):
                ranking -= 3.0
            return ranking

        ranked = sorted(evidence, key=score, reverse=True)
        selected = [
            item
            for item in ranked
            if not evidence_is_noisy_for_fallback(item, prompt=prompt)
            and not (intro_mode and evidence_is_intro_noise(item))
        ]
        if intro_mode and not selected:
            selected = [
                item
                for item in ranked
                if evidence_is_intro_priority(item) and not evidence_is_intro_noise(item)
            ]
        if not selected:
            selected = [
                item
                for item in ranked
                if item.retrievalStage == "raw_chunk"
                and not (intro_mode and evidence_is_intro_noise(item))
            ] or [
                item
                for item in ranked
                if not (intro_mode and evidence_is_intro_noise(item))
            ] or ranked
        concentrated: list[EvidenceItem] = []
        seen_units: set[str] = set()

        for item in selected:
            key = document_key(item)
            section_value = section_key(item)
            excerpt_key = re.sub(r"\s+", " ", (item.excerpt or "")[:320].strip().lower())
            unit_key = f"{key}::{section_value}::{excerpt_key}"
            if unit_key in seen_units:
                continue
            concentrated.append(item)
            seen_units.add(unit_key)
            if len(concentrated) >= limit:
                break

        return concentrated[:limit] if concentrated else selected[:limit]

    def build_local_retrieval_fallback(client_id: str, prompt: str, evidence: list[EvidenceItem], retrieval_bundle, failure_detail: str) -> AiStructuredResponse:
        retrieval_meta = retrieval_bundle.retrieval_summary if isinstance(retrieval_bundle.retrieval_summary, dict) else {}
        preferred_categories = [
            str(item)
            for item in retrieval_meta.get("preferredCategories", [])
            if str(item).strip()
        ] if isinstance(retrieval_meta.get("preferredCategories"), list) else []
        top_evidence = select_high_signal_evidence(
            evidence,
            limit=6,
            prompt=prompt,
            preferred_categories=preferred_categories,
        )
        if not top_evidence:
            return AiStructuredResponse(
                content="正式长回答阶段没有成功完成，当前也缺少可用于兜底的原始证据。",
                judgment="这次回答没有拿到足够的原始证据来组织一版可靠兜底稿。",
                analysis=f"错误信息：{failure_detail}",
                actions="建议直接重试；如果反复失败，请检查本地 AI 配置与当前证据链路。",
                timeline="恢复后可立即重新生成。",
            )
        evidence_lines = []
        for index, item in enumerate(top_evidence, start=1):
            label = item.title
            if item.sectionLabel:
                label = f"{label} / {item.sectionLabel}"
            evidence_lines.append(f"{index}. {label}：{item.excerpt[:180]}")
        content = (
            "先基于当前已命中的高信号原始证据，给出一版可继续推进讨论的判断稿。\n\n"
            "当前最值得抓住的原始观察包括：\n"
            + "\n".join(evidence_lines)
        )
        judgment = "当前已经命中了一批高信号原始证据，可以先形成一版判断；正式长文稿暂未完成。"
        analysis = (
            "这版内容优先提取最有信息量的原始观察，避免把低质量包装页、模板页或噪音片段混进来。"
            "它的作用是先把关键判断抓出来，便于继续追问、扩写或要求重生成。"
            f"\n\n失败详情：{failure_detail}"
        )
        actions = "建议继续在同一问题上重试正式生成，或围绕这里列出的原始观察逐条追问。"
        timeline = "当前处于证据已就位、正式成文未完成的状态。"
        return AiStructuredResponse(
            content=content,
            judgment=judgment,
            analysis=analysis,
            actions=actions,
            timeline=timeline,
        )

    def build_partial_generation_fallback(
        prompt: str,
        partial_content: str,
        failure_detail: str,
        *,
        partial_structured: dict[str, object] | None = None,
    ) -> AiStructuredResponse:
        cleaned_content = partial_content.strip()
        if not cleaned_content:
            return AiStructuredResponse(
                content="正式成文阶段没有完整完成，当前也没有保留到足够可读的已生成正文。",
                judgment="这次回答进入了成文阶段，但没有留下可直接交付的部分正文。",
                analysis=f"错误信息：{failure_detail}",
                actions="建议稍后重试，或把问题拆小后继续追问。",
                timeline="恢复后可立即重新扩写。",
            )
        note = "注：后续扩写阶段未完整完成，当前先保留已生成的核心正文。"
        if note not in cleaned_content:
            cleaned_content = f"{cleaned_content}\n\n{note}"
        payload = partial_structured if isinstance(partial_structured, dict) else {}
        judgment = str(payload.get("judgment") or "").strip() or "当前已经生成一版可读的核心判断，但后续扩写阶段没有完整完成。"
        analysis = str(payload.get("analysis") or "").strip()
        actions = str(payload.get("actions") or "").strip() or "建议围绕当前已生成正文继续追问，或稍后重试扩写。"
        if analysis:
            analysis = f"{analysis}\n\n后续扩写失败详情：{failure_detail}"
        else:
            analysis = (
                "这次回答已经保留了前面成功生成的正文，失败发生在后续主体扩写或建议动作阶段。"
                "为避免旧兜底稿覆盖当前判断，系统当前会优先保留这版正文。"
                f"\n\n失败详情：{failure_detail}"
            )
        return AiStructuredResponse(
            content=cleaned_content,
            judgment=judgment,
            analysis=analysis,
            actions=actions,
            timeline="当前处于部分成文已完成、后续扩写未完成的状态。",
        )

    def build_compact_grounded_note(client_id: str, prompt: str, evidence: list[EvidenceItem], retrieval_bundle) -> str:
        client = build_client_summary(client_id)
        retrieval_meta = retrieval_bundle.retrieval_summary if isinstance(retrieval_bundle.retrieval_summary, dict) else {}
        preferred_categories = [
            str(item)
            for item in retrieval_meta.get("preferredCategories", [])
            if str(item).strip()
        ] if isinstance(retrieval_meta.get("preferredCategories"), list) else []
        curated_evidence = select_high_signal_evidence(
            evidence,
            limit=8,
            prompt=prompt,
            preferred_categories=preferred_categories,
        )
        summary_lines = [
            f"客户：{client.name}",
            f"问题：{prompt}",
            f"当前已命中证据数：{len(curated_evidence)}",
            "以下是已经命中的原始证据，请基于这些证据快速组织一版完整回答：",
        ]
        for index, item in enumerate(curated_evidence, start=1):
            label = item.title
            if item.sectionLabel:
                label = f"{label} / {item.sectionLabel}"
            compact_excerpt = re.sub(r"\s+", " ", item.excerpt or "").strip()[:220]
            summary_lines.append(f"[证据{index}] {label}\n{compact_excerpt}")
        return "\n\n".join(summary_lines)[:2600]

    def upsert_task_note(task_id: str, note: str) -> None:
        table_name = "task_notes" if state.db.fetchone("SELECT 1 FROM tasks WHERE id = ?", (task_id,)) else "task_notes_cloud"
        existing = state.db.fetchone(f"SELECT id FROM {table_name} WHERE task_id = ?", (task_id,))
        timestamp = now_iso()
        if existing:
            state.db.execute(
                f"UPDATE {table_name} SET note = ?, updated_at = ? WHERE task_id = ?",
                (note, timestamp, task_id),
            )
        else:
            state.db.execute(
                f"INSERT INTO {table_name}(id, task_id, note, created_at, updated_at) VALUES(?, ?, ?, ?, ?)",
                (new_id("tn"), task_id, note, timestamp, timestamp),
            )

    def create_task(payload: TaskPayload, status: str = "todo") -> TaskRecord:
        scope_mode = payload.scopeMode or "COLLAB_SHARED"
        requested_client_id = None if scope_mode == "PERSONAL_ONLY" else payload.clientId
        requested_event_line_id = None if scope_mode == "PERSONAL_ONLY" else payload.eventLineId
        normalized_client_id, normalized_event_line_id = _normalize_task_client_and_event_line_refs(
            requested_client_id,
            requested_event_line_id,
        )
        requested_project_module_id = None if scope_mode == "PERSONAL_ONLY" else payload.projectModuleId
        requested_project_flow_id = None if scope_mode == "PERSONAL_ONLY" else payload.projectFlowId
        project_module, project_flow = resolve_project_structure_refs(normalized_client_id, requested_project_module_id, requested_project_flow_id)
        project_context = build_task_project_context(
            normalized_client_id,
            payload.sourceType,
            payload.sourceId,
            task_title=payload.title,
            task_desc=payload.desc,
            project_module_id=project_module.id if project_module else None,
            project_flow_id=project_flow.id if project_flow else None,
        )
        event_line_context = _event_line_snapshot_context(state.db, normalized_event_line_id, None)
        (
            business_category,
            current_blocker,
            next_action,
            recent_decision,
            evidence_count,
        ) = _resolve_task_action_os_fields(
            title=payload.title,
            desc=payload.desc,
            source_type=payload.sourceType,
            business_category=payload.businessCategory,
            current_blocker=payload.currentBlocker,
            next_action=payload.nextAction,
            recent_decision=payload.recentDecision,
            evidence_count=payload.evidenceCount,
            project_context=project_context,
            event_line_context=event_line_context,
            attachment_count=0,
        )
        if get_cloud_token():
            session_user = get_cached_session_user()
            collaborator_ids = payload.collaboratorIds or ([session_user.id] if session_user else [])
            owner_id = payload.ownerId or (collaborator_ids[0] if collaborator_ids else None)
            response = cloud_request(
                "POST",
                "/api/v1/tasks",
                json_body={
                    "title": payload.title,
                    "description": payload.desc,
                    "priority": payload.priority,
                    "listId": payload.listId,
                    "dueDate": payload.dueDate or normalize_due_date_input(payload.ddl),
                    "durationMinutes": payload.durationMinutes,
                    "scopeMode": scope_mode,
                    "clientId": normalized_client_id,
                    "eventLineId": normalized_event_line_id,
                    "projectModuleId": project_module.id if project_module else None,
                    "projectFlowId": project_flow.id if project_flow else None,
                    "collaboratorIds": collaborator_ids,
                    "ownerId": owner_id,
                    "sourceType": payload.sourceType,
                    "sourceId": payload.sourceId,
                    "businessCategory": business_category,
                    "currentBlocker": current_blocker,
                    "nextAction": next_action,
                    "recentDecision": recent_decision,
                    "evidenceCount": evidence_count,
                },
            )
            if not isinstance(response, dict):
                raise HTTPException(status_code=502, detail="Invalid cloud task payload")
            log_activity("task.create", "task", str(response.get("id", "unknown")), payload.model_dump())
            created_task = build_cloud_task(response, {})
            growth_user_id, growth_user_name = resolve_growth_actor()
            ingest_task_growth_candidate(
                state.db,
                user_id=growth_user_id,
                user_name=growth_user_name,
                task=created_task,
                source_type="task_context_candidate",
                created_at=now_iso(),
            )
            return created_task
        timestamp = now_iso()
        task_id = new_id("task")
        list_id = payload.listId or (_get_local_task_settings().defaultListId or "list-0")
        list_row = state.db.fetchone("SELECT * FROM task_lists WHERE id = ?", (list_id,))
        if not list_row or list_row["archived_at"]:
            raise HTTPException(status_code=400, detail="任务清单无效")
        resolved_tags = normalize_local_task_tags(payload.tagIds, payload.tags)
        state.db.execute(
            """
            INSERT INTO tasks(
                id, title, description, status, priority, list_id, owner_name, ddl, due_date, duration_minutes, event_line_id, source_type, source_id,
                client_id, project_module_id, project_flow_id, scope_mode, business_category, current_blocker, next_action, recent_decision, evidence_count,
                tags_json, tag_ids_json, created_at, updated_at
            )
            VALUES(?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                task_id,
                payload.title,
                payload.desc,
                status,
                payload.priority,
                list_id,
                payload.ownerName,
                payload.ddl,
                payload.dueDate or normalize_due_date_input(payload.ddl),
                payload.durationMinutes,
                normalized_event_line_id,
                payload.sourceType,
                payload.sourceId,
                normalized_client_id,
                project_module.id if project_module else None,
                project_flow.id if project_flow else None,
                scope_mode,
                business_category,
                current_blocker,
                next_action,
                recent_decision,
                evidence_count,
                to_json([tag.name for tag in resolved_tags]),
                to_json([tag.id for tag in resolved_tags]),
                timestamp,
                timestamp,
            ),
        )
        if normalized_event_line_id:
            state.db.execute(
                """
                INSERT INTO event_line_activities(
                    id, event_line_id, source_type, source_id, happened_at, actor_id, actor_name, title, summary, metadata_json
                ) VALUES(?, ?, 'task_activity', ?, ?, NULL, ?, ?, ?, ?)
                """,
                    (
                        new_id("ela"),
                    normalized_event_line_id,
                    task_id,
                    timestamp,
                    payload.ownerName or "",
                    "新增任务",
                    f"创建任务：{payload.title}",
                    to_json({"eventType": "created"}),
                ),
            )
        log_activity("task.create", "task", task_id, payload.model_dump())
        created_task = fetch_tasks("t.id = ?", (task_id,))[0]
        record_task_writeback(
            state.db,
            task_id=created_task.id,
            title=created_task.title,
            description=created_task.desc,
            status=created_task.status,
            due_date=created_task.dueDate,
            client_id=created_task.clientId,
            event_line_id=created_task.eventLineId,
        )
        growth_user_id, growth_user_name = resolve_growth_actor()
        ingest_task_growth_candidate(
            state.db,
            user_id=growth_user_id,
            user_name=growth_user_name,
            task=created_task,
            source_type="task_context_candidate",
            created_at=timestamp,
        )
        return created_task

    def extract_meeting_content(text: str) -> tuple[list[str], list[str], list[tuple[str, str, float]], list[tuple[str, str]], list[tuple[str, list[str]]]]:
        sentences = [segment.strip(" -") for segment in re.split(r"[\n。！？!?]+", text) if segment.strip()]
        agenda = sentences[:2] or ["会前目标梳理", "关键问题确认"]
        decisions = [item for item in sentences if "决定" in item or "确定" in item][:3]
        if not decisions and sentences:
            decisions = [f"优先推进：{sentences[0][:24]}"]
        actions = [item for item in sentences if any(keyword in item for keyword in ["负责", "跟进", "行动", "任务", "推进"])][:3]
        if not actions:
            actions = [f"跟进 {item[:20]}" for item in sentences[:2]] or ["补齐会议待办"]
        parsed_actions = [(item, current_operator_row()["name"], 0.84) for item in actions]
        risks = [(item, "high" if any(word in item for word in ["风险", "阻力", "卡点"]) else "normal") for item in sentences if any(word in item for word in ["风险", "阻力", "卡点", "资源", "预算"])]
        ambiguities = [(item, ["待确认责任人", "待确认时间点"]) for item in sentences if any(word in item for word in ["待确认", "待补", "需要再看", "?"])]
        return agenda, decisions, parsed_actions, risks[:3], ambiguities[:3]

    def is_private_task(task: TaskRecord) -> bool:
        return task.scopeMode == "PERSONAL_ONLY" or any(tag.scope == "self" for tag in task.tags)

    def local_review_row_for_week(week_label: str):
        operator_id = str(current_operator_row()["id"])
        return state.db.fetchone(
            "SELECT * FROM weekly_reviews WHERE week_label = ? AND operator_id = ? ORDER BY created_at DESC LIMIT 1",
            (week_label, operator_id),
        )

    def local_review_history() -> ReviewHistoryResponse:
        operator_id = str(current_operator_row()["id"])
        rows = state.db.fetchall(
            """
            SELECT
                r.week_label,
                COALESCE(r.updated_at, r.created_at) AS submitted_at,
                (
                    SELECT COUNT(*)
                    FROM weekly_review_task_entries e
                    WHERE e.review_id = r.id AND e.content_domain = 'work'
                ) AS work_item_count,
                (
                    SELECT COUNT(*)
                    FROM weekly_review_task_entries e
                    WHERE e.review_id = r.id AND e.content_domain = 'personal'
                ) AS personal_item_count
            FROM weekly_reviews r
            WHERE r.operator_id = ?
            ORDER BY COALESCE(r.updated_at, r.created_at) DESC, r.week_label DESC
            """,
            (operator_id,),
        )
        return ReviewHistoryResponse(
            items=[
                ReviewHistoryEntryRecord(
                    weekLabel=str(row["week_label"] or ""),
                    submittedAt=str(row["submitted_at"] or ""),
                    workItemCount=int(row["work_item_count"] or 0),
                    personalItemCount=int(row["personal_item_count"] or 0),
                )
                for row in rows
                if str(row["week_label"] or "").strip()
            ]
        )

    def build_local_review_record(row) -> WeeklyReviewRecord:
        operator = current_operator_row()
        return WeeklyReviewRecord(
            id=str(row["id"]),
            userId=str(operator["id"]),
            userName=str(operator["name"]),
            weekLabel=str(row["week_label"]),
            workFreeNote=str(row["work_free_note"] or row["summary"] or ""),
            personalGrowthNote=str(row["personal_growth_note"] or ""),
            personalPrivateNote=str(row["personal_private_note"] or ""),
            submittedAt=str(row["updated_at"] or row["created_at"]),
            createdAt=str(row["created_at"]),
            updatedAt=str(row["updated_at"] or row["created_at"]),
        )

    def build_preview_review_record(week_label: str) -> WeeklyReviewRecord:
        operator = current_operator_row()
        timestamp = now_iso()
        return WeeklyReviewRecord(
            id=f"review_preview::{week_label}",
            userId=str(operator["id"]),
            userName=str(operator["name"]),
            weekLabel=week_label,
            workFreeNote="",
            personalGrowthNote="",
            personalPrivateNote="",
            submittedAt=timestamp,
            createdAt=timestamp,
            updatedAt=timestamp,
        )

    def local_review_entries_by_task(review_id: str) -> dict[str, dict[str, object]]:
        rows = state.db.fetchall(
            "SELECT * FROM weekly_review_task_entries WHERE review_id = ? ORDER BY reviewed_at DESC",
            (review_id,),
        )
        return {str(row["task_id"]): dict(row) for row in rows}

    def summarize_local_review_notes(items: list[WeeklyReviewTaskEntryRecord]) -> str:
        if not items:
            return "本周还没有填写任务复盘。"
        issue_keywords = ("卡住", "阻力", "困难", "问题", "风险", "不足")
        harvest_keywords = ("收获", "学到", "发现", "有效", "清楚")
        support_keywords = ("需要支持", "需要帮助", "资源", "协同")
        overload_count = sum(1 for item in items if item.structuredNote.lightweightTag == "工作过度饱和")
        issue_count = sum(
            1
            for item in items
            if item.structuredNote.lightweightTag
            or item.structuredNote.blockerReason
            or any(keyword in item.note for keyword in issue_keywords)
        )
        harvest_count = sum(
            1
            for item in items
            if item.structuredNote.reflection
            or item.structuredNote.successReason
            or item.structuredNote.successExperience
            or any(keyword in item.note for keyword in harvest_keywords)
        )
        support_count = sum(
            1
            for item in items
            if item.structuredNote.lightweightTag
            or item.structuredNote.supportNeeded
            or any(keyword in item.note for keyword in support_keywords)
        )
        insight_count = sum(
            1
            for item in items
            if item.structuredNote.reflection or item.structuredNote.failureInsight or item.structuredNote.blockerReason
        )
        overload_text = f"，其中明确提到工作过度饱和 {overload_count} 次" if overload_count else ""
        return f"共记录 {len(items)} 条任务复盘，其中提到问题/阻力 {issue_count} 次，收获/有效经验 {harvest_count} 次，失败心得 {insight_count} 次，需要支持 {support_count} 次{overload_text}。"

    def current_review_week_label() -> str:
        iso = datetime.now().isocalendar()
        return f"{iso.year}-W{iso.week:02d}"

    def build_client_business_context_modules(
        work_items: list[WeeklyReviewTaskEntryRecord],
    ) -> list[OrganizationDnaModuleRecord]:
        discovered_client_ids: list[str] = []
        seen_client_ids: set[str] = set()
        for item in work_items:
            task_rows = fetch_tasks("t.id = ?", (item.taskId,))
            if not task_rows:
                continue
            task = task_rows[0]
            source_id = (task.sourceId or "").strip()
            source_type = (task.sourceType or "").strip()
            client_id: str | None = (task.clientId or "").strip() or None
            if source_type == "meeting" and source_id:
                row = state.db.fetchone("SELECT client_id FROM meetings WHERE id = ?", (source_id,))
                client_id = str(row["client_id"]) if row else None
            elif source_id and state.db.fetchone("SELECT 1 FROM clients WHERE id = ?", (source_id,)):
                client_id = source_id
            elif source_id:
                row = state.db.fetchone("SELECT client_id FROM goal_records WHERE id = ?", (source_id,))
                client_id = str(row["client_id"]) if row else None
            elif task.clientName:
                row = state.db.fetchone("SELECT id FROM clients WHERE name = ?", (task.clientName,))
                client_id = str(row["id"]) if row and row["id"] else None
            if client_id and client_id not in seen_client_ids:
                seen_client_ids.add(client_id)
                discovered_client_ids.append(client_id)

        modules: list[OrganizationDnaModuleRecord] = []
        for client_id in discovered_client_ids[:3]:
            try:
                client = build_client_summary(client_id)
            except HTTPException:
                continue
            goals = list_client_goals(client_id)[:2]
            recent_meeting_rows = state.db.fetchall(
                "SELECT id FROM meetings WHERE client_id = ? ORDER BY updated_at DESC LIMIT 2",
                (client_id,),
            )
            decision_summaries: list[str] = []
            risk_summaries: list[str] = []
            for meeting_row in recent_meeting_rows:
                detail = build_meeting_detail(str(meeting_row["id"]))
                for decision in detail.decisions:
                    summary = decision.summary.strip()
                    if summary and summary not in decision_summaries:
                        decision_summaries.append(summary)
                    if len(decision_summaries) >= 2:
                        break
                for risk in detail.risks:
                    summary = risk.summary.strip()
                    if summary and summary not in risk_summaries:
                        risk_summaries.append(summary)
                    if len(risk_summaries) >= 2:
                        break
                if len(decision_summaries) >= 2 and len(risk_summaries) >= 2:
                    break
            document_cards = fetch_document_cards(state.db, client_id, data_dir=state.data_dir, limit=2)
            card_titles = [str(item.get("title") or "").strip() for item in document_cards if str(item.get("title") or "").strip()]
            client_dna_summary = next(
                (
                    module.summary.strip()
                    for module in list_client_dna_modules(client_id)
                    if module.summary.strip()
                ),
                "",
            )
            summary_parts = [f"客户阶段：{client.stage}"]
            if goals:
                summary_parts.append(f"关键目标：{'；'.join(goal.title for goal in goals)}")
            if decision_summaries:
                summary_parts.append(f"近期会议决策：{'；'.join(decision_summaries[:2])}")
            if risk_summaries:
                summary_parts.append(f"当前阻力：{'；'.join(risk_summaries[:2])}")
            if card_titles:
                summary_parts.append(f"资料线索：{'、'.join(card_titles[:2])}")
            if client_dna_summary:
                summary_parts.append(f"客户背景：{client_dna_summary}")
            normalized = " ".join(part for part in summary_parts if part).strip()
            if not normalized:
                continue
            modules.append(
                OrganizationDnaModuleRecord(
                    moduleKey="business_intro",
                    title=f"{client.name} 业务背景",
                    markdownContent=normalized,
                    normalizedText=normalized,
                    summary=normalized,
                    fileName=None,
                    contentHash=None,
                    updatedAt=None,
                    updatedBy="client_workspace_sync",
                    hasDocument=True,
                )
            )
        modules.sort(key=lambda item: ((item.title or "").strip(), (item.summary or "").strip()))
        return modules

    def build_review_context_modules(
        work_items: list[WeeklyReviewTaskEntryRecord],
        organization_modules: list[OrganizationDnaModuleRecord] | None = None,
    ) -> list[OrganizationDnaModuleRecord]:
        base_modules = list(organization_modules) if organization_modules is not None else list_organization_dna_modules()
        client_modules = build_client_business_context_modules(work_items)
        return [*client_modules, *base_modules]

    def build_review_analyses(
        week_label: str,
        work_items: list[WeeklyReviewTaskEntryRecord],
        personal_items: list[WeeklyReviewTaskEntryRecord],
        organization_modules: list[OrganizationDnaModuleRecord] | None = None,
        org_model_profile: OrgModelProfileRecord | None = None,
        viewer_role: Literal["employee", "department_lead", "admin"] = "employee",
    ) -> tuple[WeeklyReviewAnalysisRecord, WeeklyReviewAnalysisRecord]:
        work_modules = build_review_context_modules(work_items, organization_modules)
        personal_modules = list(organization_modules) if organization_modules is not None else list_organization_dna_modules()
        work_analysis = build_weekly_review_analysis(
            "work",
            week_label,
            work_items,
            work_modules,
            org_model_profile=org_model_profile,
            viewer_role=viewer_role,
        )
        personal_analysis = build_weekly_review_analysis("personal", week_label, personal_items, personal_modules, org_model_profile=org_model_profile)
        return work_analysis, personal_analysis

    def _memory_slot_label(slot_key: str) -> str:
        mapping = {
            "current_stage": "当前阶段",
            "current_work": "当前事项",
            "current_blocker": "当前阻塞",
            "recent_decision": "最近关键决策",
            "next_step": "下一步",
        }
        return mapping.get(slot_key.strip(), slot_key.strip())

    def _normalize_event_line_reference(value: str | None) -> str:
        normalized = str(value or "").strip()
        if normalized.startswith("event_line::"):
            return normalized.split("::", 1)[1].strip()
        return normalized

    def _memory_background_source_labels(
        *,
        has_notebook: bool,
        has_snapshot: bool,
        has_evidence: bool,
        has_review_signals: bool,
        has_linked_facts: bool,
        has_pending_clarification: bool,
    ) -> list[str]:
        labels: list[str] = []
        if has_notebook:
            labels.append("组织笔记")
        if has_snapshot:
            labels.append("事件线记忆")
        if has_review_signals:
            labels.append("周复盘信号")
        if has_linked_facts:
            labels.append("统一事实池")
        if has_evidence:
            labels.append("任务/附件证据")
        if has_pending_clarification:
            labels.append("待澄清槽位")
        return labels

    def _compact_business_line(value: str | None, *, max_length: int = 120) -> str:
        return sanitize_memory_background_text(value, reject_generic=True, max_length=max_length)

    def _task_clause_candidates(value: str | None) -> list[str]:
        return [
            clause
            for clause in (
                _compact_business_line(item, max_length=140)
                for item in re.split(r"[\n。；;]", str(value or ""))
            )
            if clause
        ]

    def _dedupe_lines(values: list[str], *, limit: int = 4) -> list[str]:
        result: list[str] = []
        seen: set[str] = set()
        for value in values:
            normalized = _compact_business_line(value, max_length=140)
            if not normalized or normalized in seen:
                continue
            seen.add(normalized)
            result.append(normalized)
            if len(result) >= limit:
                break
        return result

    def _first_meaningful_line(*values: str | None) -> str:
        return next((item for item in (_compact_business_line(value, max_length=140) for value in values) if item), "")

    def _infer_bundle_operating_mode(bundle: EventLineContextBundleRecord) -> Literal["meeting", "delivery", "materials", "analysis", "coordination", "general"]:
        text = " ".join(
            [
                bundle.lineName,
                bundle.summary,
                bundle.intent,
                bundle.currentWork,
                bundle.currentBlocker,
                bundle.nextStep,
                *[fact.title for fact in bundle.taskFacts[:3]],
                *[fact.summary for fact in bundle.taskFacts[:3]],
                *[fact.title for fact in bundle.attachmentFacts[:3]],
                *[fact.summary for fact in bundle.attachmentFacts[:2]],
            ]
        )
        if _contains_any_keyword(text, ("见面", "会谈", "拜访", "演示", "电话会", "讨论", "交流", "看系统", "对接")):
            return "meeting"
        if _contains_any_keyword(text, ("方案", "报告", "提纲", "清单", "输出", "成稿", "交付", "文稿", "版本")):
            return "delivery"
        if _contains_any_keyword(text, ("资料", "材料", "台账", "导入", "归档", "工具包", "附件", "底稿")):
            return "materials"
        if _contains_any_keyword(text, ("诊断", "分析", "研判", "判断", "策略", "洞察")):
            return "analysis"
        if _contains_any_keyword(text, ("确认", "审批", "复核", "拍板", "协同", "流程")):
            return "coordination"
        return "general"

    def _bundle_key_evidence_lines(bundle: EventLineContextBundleRecord, *, limit: int = 3) -> list[str]:
        candidates = [
            *[fact.summary for fact in bundle.taskFacts[:2]],
            *[fact.summary for fact in bundle.meetingFacts[:2]],
            *[fact.title for fact in bundle.attachmentFacts[:2]],
            *bundle.recentFacts[:2],
            *[fact.summary for fact in bundle.clarificationFacts[:2]],
        ]
        return _dedupe_lines(candidates, limit=limit)

    def _bundle_people_or_products_hint(bundle: EventLineContextBundleRecord) -> str:
        people = _dedupe_lines(bundle.keyPeople, limit=2)
        products = _dedupe_lines(bundle.keyProducts, limit=2)
        if people and products:
            return f"关键对象是 {'、'.join(people)}，当前讨论对象涉及 {'、'.join(products)}。"
        if people:
            return f"关键对象是 {'、'.join(people)}。"
        if products:
            return f"当前讨论对象涉及 {'、'.join(products)}。"
        return ""

    def _bundle_missing_context_lines(bundle: EventLineContextBundleRecord) -> list[str]:
        missing: list[str] = []
        if not bundle.keyPeople:
            missing.append("关键人角色")
        if not bundle.collaborationRelationship:
            missing.append("合作关系")
        if not bundle.currentWork:
            missing.append("当前事项")
        if not bundle.nextStep:
            missing.append("下周动作")
        return missing[:3]

    def _bundle_coverage_score(bundle: EventLineContextBundleRecord) -> int:
        checks = [
            bool(bundle.summary or bundle.intent),
            bool(bundle.currentWork),
            bool(bundle.currentBlocker),
            bool(bundle.nextStep or bundle.recentDecision),
            bool(bundle.projectName or bundle.collaborationRelationship or bundle.organizationIntro),
            bool(bundle.currentChallenges or bundle.collaborationGoals),
            bool(bundle.keyPeople or bundle.keyProducts),
            bool(bundle.taskFacts),
            bool(bundle.meetingFacts or bundle.attachmentFacts or bundle.clarificationFacts),
            bool(bundle.evidenceRefs),
        ]
        return int(round((sum(1 for item in checks if item) / len(checks)) * 100))

    def _bundle_confidence_score(
        bundle: EventLineContextBundleRecord,
        *,
        coverage_score: int,
    ) -> int:
        readiness_weight = {"low": 0.35, "medium": 0.6, "high": 0.82}.get(bundle.readiness, 0.35)
        evidence_density = min(
            1.0,
            (
                min(bundle.taskCount, 3) * 0.16
                + min(bundle.meetingCount, 2) * 0.16
                + min(bundle.attachmentCount, 3) * 0.12
                + min(bundle.supportRequestCount, 2) * 0.08
                + min(len(bundle.clarificationFacts), 2) * 0.1
                + min(len(bundle.evidenceRefs), 4) * 0.06
            ),
        )
        confidence = min(
            1.0,
            (coverage_score / 100.0) * 0.52 + readiness_weight * 0.28 + evidence_density * 0.2,
        )
        return int(round(confidence * 100))

    def _bundle_safe_output_mode(
        bundle: EventLineContextBundleRecord,
        *,
        coverage_score: int,
        confidence_score: int,
    ) -> Literal["needs_input", "summary_only", "full_judgment"]:
        if coverage_score >= 70 and confidence_score >= 65:
            return "full_judgment"
        if coverage_score >= 40 and confidence_score >= 35:
            return "summary_only"
        return "needs_input"

    def _bundle_publish_state(
        *,
        viewer_role: ReviewViewerRole,
        safe_output_mode: Literal["needs_input", "summary_only", "full_judgment"],
    ) -> Literal["local_preview", "publish_ready", "published_by_human", "published_by_robot", "stale"]:
        if viewer_role in {"admin", "department_lead"} and safe_output_mode == "full_judgment":
            return "publish_ready"
        return "local_preview"

    def _bundle_fingerprint(bundle: EventLineContextBundleRecord) -> str:
        payload = {
            "eventLineId": bundle.eventLineId,
            "lineName": bundle.lineName,
            "businessCategory": bundle.businessCategory,
            "stage": bundle.stage,
            "summary": bundle.summary,
            "intent": bundle.intent,
            "currentWork": bundle.currentWork,
            "currentBlocker": bundle.currentBlocker,
            "recentDecision": bundle.recentDecision,
            "nextStep": bundle.nextStep,
            "recentProgress": bundle.recentProgress,
            "projectName": bundle.projectName,
            "collaborationRelationship": bundle.collaborationRelationship,
            "organizationIntro": bundle.organizationIntro,
            "currentChallenges": bundle.currentChallenges,
            "collaborationGoals": bundle.collaborationGoals,
            "keyPeople": bundle.keyPeople,
            "keyProducts": bundle.keyProducts,
            "recentFacts": bundle.recentFacts,
            "taskFacts": [fact.model_dump(mode="json") for fact in bundle.taskFacts[:6]],
            "meetingFacts": [fact.model_dump(mode="json") for fact in bundle.meetingFacts[:4]],
            "attachmentFacts": [fact.model_dump(mode="json") for fact in bundle.attachmentFacts[:6]],
            "clarificationFacts": [fact.model_dump(mode="json") for fact in bundle.clarificationFacts[:4]],
            "evidenceRefs": [ref.model_dump(mode="json") for ref in bundle.evidenceRefs[:8]],
            "trendSignals": [
                {
                    "key": signal.key,
                    "signalType": signal.signalType,
                    "severity": signal.severity,
                    "statement": signal.statement,
                }
                for signal in bundle.trendSignals[:6]
            ],
            "taskCount": bundle.taskCount,
            "meetingCount": bundle.meetingCount,
            "attachmentCount": bundle.attachmentCount,
            "supportRequestCount": bundle.supportRequestCount,
            "readiness": bundle.readiness,
        }
        serialized = json.dumps(payload, ensure_ascii=False, sort_keys=True, separators=(",", ":"))
        return hashlib.sha1(serialized.encode("utf-8")).hexdigest()[:16]

    def _event_line_bundle_readiness(
        *,
        notebook_confidence: float,
        snapshot_confidence: float,
        task_count: int,
        attachment_count: int,
        support_request_count: int,
        clarification_count: int,
        current_work: str,
        current_blocker: str,
        next_step: str,
    ) -> Literal["low", "medium", "high"]:
        score = 0
        if notebook_confidence >= 0.5:
            score += 1
        if snapshot_confidence >= 0.5:
            score += 1
        if task_count >= 2:
            score += 1
        if attachment_count >= 1:
            score += 1
        if support_request_count >= 1:
            score += 1
        if clarification_count >= 1:
            score += 1
        if current_work:
            score += 1
        if current_blocker:
            score += 1
        if next_step:
            score += 1
        if score >= 6:
            return "high"
        if score >= 3:
            return "medium"
        return "low"

    def _event_line_context_fact(
        *,
        source_type: Literal["task", "meeting", "attachment", "support_request", "clarification", "notebook", "event_line_memory"],
        source_id: str,
        title: str,
        summary: str,
        happened_at: str | None = None,
    ) -> EventLineContextFactRecord | None:
        normalized_title = str(title or "").strip()
        normalized_summary = _compact_business_line(summary, max_length=160)
        if not normalized_title and not normalized_summary:
            return None
        return EventLineContextFactRecord(
            sourceType=source_type,
            sourceId=source_id,
            title=normalized_title or normalized_summary,
            summary=normalized_summary or normalized_title,
            happenedAt=happened_at,
        )

    def _context_fact_to_evidence_ref(fact: EventLineContextFactRecord) -> ReviewDashboardEvidenceRefRecord:
        return ReviewDashboardEvidenceRefRecord(
            sourceType=fact.sourceType,
            sourceId=fact.sourceId,
            title=fact.title,
            summary=fact.summary,
        )

    def _dedupe_dashboard_evidence_refs(
        refs: list[ReviewDashboardEvidenceRefRecord],
        *,
        limit: int | None = None,
    ) -> list[ReviewDashboardEvidenceRefRecord]:
        deduped: list[ReviewDashboardEvidenceRefRecord] = []
        seen: set[tuple[str, str, str, str]] = set()
        for ref in refs:
            source_type = str(ref.sourceType or "").strip()
            source_id = str(ref.sourceId or "").strip()
            title = str(ref.title or "").strip()
            summary = str(ref.summary or "").strip()
            if not source_type or not source_id or not title:
                continue
            key = (source_type, source_id, title, summary)
            if key in seen:
                continue
            seen.add(key)
            deduped.append(
                ReviewDashboardEvidenceRefRecord(
                    sourceType=ref.sourceType,
                    sourceId=source_id,
                    title=title,
                    summary=summary or None,
                )
            )
            if limit is not None and len(deduped) >= limit:
                break
        return deduped

    def _event_line_context_bundle(
        event_line_id: str,
        *,
        analysis: WeeklyReviewAnalysisRecord | None = None,
    ) -> EventLineContextBundleRecord | None:
        normalized_event_line_id = _normalize_event_line_reference(event_line_id)
        if not normalized_event_line_id:
            return None
        row = state.db.fetchone("SELECT * FROM event_lines WHERE id = ?", (normalized_event_line_id,))
        if not row:
            return None
        detail = build_event_line_detail(row)
        snapshot = detail.memorySnapshot
        notebook = (
            get_client_notebook_response(state.db, detail.eventLine.primaryClientId).organizationNotebookSnapshot
            if detail.eventLine.primaryClientId
            else None
        )
        clarification_rows = state.db.fetchall(
            """
            SELECT *
            FROM clarification_records
            WHERE (scope_type = 'event_line' AND scope_id = ?)
               OR (scope_type = 'client' AND scope_id = ?)
            ORDER BY updated_at DESC, created_at DESC
            LIMIT 8
            """,
            (
                normalized_event_line_id,
                detail.eventLine.primaryClientId or "",
            ),
        )
        clarification_facts = [
            fact
            for fact in (
                _event_line_context_fact(
                    source_type="clarification",
                    source_id=str(item["id"]),
                    title=_first_nonempty_text(item["question"], item["slot_key"], "澄清记录") or "澄清记录",
                    summary=_first_nonempty_text(item["answer_text"], item["question"], item["slot_key"]) or "",
                    happened_at=str(item["updated_at"] or item["created_at"] or ""),
                )
                for item in clarification_rows
            )
            if fact is not None
        ]
        task_facts = [
            fact
            for fact in (
                _event_line_context_fact(
                    source_type="task",
                    source_id=task.id,
                    title=task.title,
                    summary=_first_nonempty_text(
                        _task_clause_candidates(task.desc)[0] if _task_clause_candidates(task.desc) else None,
                        task.recentDecision,
                        task.nextAction,
                        task.currentBlocker,
                    ) or task.title,
                    happened_at=task.updatedAt,
                )
                for task in detail.tasks[:8]
            )
            if fact is not None
        ]
        meeting_facts = [
            fact
            for fact in (
                _event_line_context_fact(
                    source_type="meeting",
                    source_id=activity.sourceId,
                    title=activity.title,
                    summary=_first_nonempty_text(
                        activity.summary,
                        activity.detail,
                        activity.title,
                    ) or activity.title,
                    happened_at=activity.happenedAt,
                )
                for activity in detail.activities
                if activity.sourceType == "meeting"
            )
            if fact is not None
        ][:6]
        attachment_facts = [
            fact
            for fact in (
                _event_line_context_fact(
                    source_type="attachment",
                    source_id=attachment.id,
                    title=attachment.title,
                    summary=_first_nonempty_text(
                        attachment.summary if hasattr(attachment, "summary") else None,
                        attachment.title,
                        attachment.path,
                    ) or attachment.title,
                    happened_at=attachment.createdAt,
                )
                for task in detail.tasks
                for attachment in task.attachments
            )
            if fact is not None
        ][:8]
        support_request_records = _support_requests_for_tasks(detail.tasks)
        support_request_count = len(support_request_records) or len(
            [activity for activity in detail.activities if activity.sourceType == "support_request"]
        )
        trend_signals = [
            signal
            for signal in (analysis.trendSignals if analysis else [])
            if (
                _normalize_event_line_reference(signal.relatedEventLineId or "") == normalized_event_line_id
                or bool(set(signal.relatedTaskIds).intersection({task.id for task in detail.tasks}))
            )
        ]
        current_work = _first_nonempty_text(
            snapshot.currentWork if snapshot else None,
            detail.eventLine.intent,
            detail.eventLine.summary,
            task_facts[0].summary if task_facts else None,
            meeting_facts[0].summary if meeting_facts else None,
        ) or ""
        current_blocker = _first_nonempty_text(
            snapshot.currentBlocker if snapshot else None,
            detail.eventLine.currentBlocker,
            notebook.currentChallenges[0] if notebook and notebook.currentChallenges else None,
            clarification_facts[0].summary if clarification_facts else None,
        ) or ""
        recent_decision = _first_nonempty_text(
            snapshot.recentDecision if snapshot else None,
            detail.eventLine.recentDecision,
            meeting_facts[0].summary if meeting_facts else None,
            clarification_facts[0].summary if clarification_facts else None,
        ) or ""
        next_step = _first_nonempty_text(
            snapshot.nextStep if snapshot else None,
            detail.eventLine.nextStep,
            detail.tasks[0].nextAction if detail.tasks and detail.tasks[0].nextAction else None,
            notebook.collaborationGoals[0] if notebook and notebook.collaborationGoals else None,
        ) or ""
        recent_progress = _first_nonempty_text(
            detail.tasks[0].recentDecision if detail.tasks and detail.tasks[0].recentDecision else None,
            meeting_facts[0].summary if meeting_facts else None,
            task_facts[0].summary if task_facts else None,
            attachment_facts[0].title if attachment_facts else None,
        ) or ""
        notebook_facts = [
            fact
            for fact in (
                _event_line_context_fact(
                    source_type="notebook",
                    source_id=f"{detail.eventLine.primaryClientId or normalized_event_line_id}-fact-{index}",
                    title="组织业务笔记",
                    summary=fact_text,
                )
                for index, fact_text in enumerate(
                    _dedupe_lines(
                        [
                            notebook.collaborationRelationship if notebook else "",
                            notebook.organizationIntro if notebook else "",
                            *(notebook.recentFacts[:2] if notebook else []),
                            *([f"关键人：{'、'.join(notebook.keyPeople[:2])}"] if notebook and notebook.keyPeople else []),
                            *([f"关键产品：{'、'.join(notebook.keyProducts[:2])}"] if notebook and notebook.keyProducts else []),
                        ],
                        limit=4,
                    )
                )
            )
            if fact is not None
        ]
        evidence_refs = _dedupe_dashboard_evidence_refs(
            [
                *_story_evidence_refs(
                    {
                        "id": normalized_event_line_id,
                        "eventLineId": normalized_event_line_id,
                        "taskTitles": [task.title for task in detail.tasks],
                        "taskIds": [task.id for task in detail.tasks],
                    }
                ),
                *[_context_fact_to_evidence_ref(fact) for fact in task_facts[:3]],
                *[_context_fact_to_evidence_ref(fact) for fact in meeting_facts[:2]],
                *[_context_fact_to_evidence_ref(fact) for fact in attachment_facts[:3]],
                *[_context_fact_to_evidence_ref(fact) for fact in clarification_facts[:2]],
                *[_context_fact_to_evidence_ref(fact) for fact in notebook_facts[:2]],
            ]
        )[:8]
        readiness = _event_line_bundle_readiness(
            notebook_confidence=notebook.confidence if notebook else 0.0,
            snapshot_confidence=snapshot.confidence if snapshot else 0.0,
            task_count=len(detail.tasks),
            attachment_count=len(attachment_facts),
            support_request_count=support_request_count,
            clarification_count=len(clarification_facts),
            current_work=current_work,
            current_blocker=current_blocker,
            next_step=next_step,
        )
        return EventLineContextBundleRecord(
            eventLineId=normalized_event_line_id,
            lineName=detail.eventLine.name,
            businessCategory=detail.eventLine.businessCategory or "",
            stage=_first_nonempty_text(detail.eventLine.stage, snapshot.currentStage if snapshot else None) or "",
            summary=detail.eventLine.summary or "",
            intent=detail.eventLine.intent or "",
            currentWork=current_work,
            currentBlocker=current_blocker,
            recentDecision=recent_decision,
            nextStep=next_step,
            recentProgress=recent_progress,
            projectName=detail.eventLine.primaryClientName or "",
            collaborationRelationship=notebook.collaborationRelationship if notebook else "",
            organizationIntro=notebook.organizationIntro if notebook else "",
            currentChallenges=list(notebook.currentChallenges) if notebook else [],
            collaborationGoals=list(notebook.collaborationGoals) if notebook else [],
            keyPeople=list(notebook.keyPeople) if notebook else [],
            keyProducts=list(notebook.keyProducts) if notebook else [],
            recentFacts=list(notebook.recentFacts)[:5] if notebook else [],
            taskFacts=task_facts,
            meetingFacts=meeting_facts,
            attachmentFacts=attachment_facts,
            clarificationFacts=clarification_facts,
            evidenceRefs=evidence_refs,
            trendSignals=trend_signals,
            taskCount=len(detail.tasks),
            meetingCount=len(meeting_facts),
            attachmentCount=len(attachment_facts),
            supportRequestCount=support_request_count,
            readiness=readiness,
        )

    def _infer_bundle_blocker_type(bundle: EventLineContextBundleRecord) -> Literal["business", "collaboration", "decision", "structure", "capacity", "evidence"]:
        text = " ".join(
            [
                bundle.currentBlocker,
                bundle.recentDecision,
                bundle.nextStep,
                bundle.currentWork,
                *bundle.currentChallenges[:2],
                *[signal.signalType for signal in bundle.trendSignals],
            ]
        )
        rich_evidence = bundle.attachmentCount >= 2 or len(bundle.meetingFacts) >= 1 or len(bundle.recentFacts) >= 2
        if _contains_any_keyword(text, ("目标", "价值", "会谈", "演示", "方案", "定位", "合作", "客户", "场景", "判断", "收束", "落点")):
            return "business"
        if _contains_any_keyword(text, ("复核", "审批", "确认", "拍板", "决策")):
            return "decision"
        if _contains_any_keyword(text, ("跨部门", "协同", "对接", "等待他人", "支持")):
            return "collaboration"
        if _contains_any_keyword(text, ("带宽", "排期", "容量", "过载", "改期")):
            return "capacity"
        if _contains_any_keyword(text, ("模块", "流程", "结构", "归属", "口径")) and not rich_evidence:
            return "structure"
        if _contains_any_keyword(text, ("资料", "证据", "附件", "摘要", "背景包", "信息不足")) and not rich_evidence:
            return "evidence"
        return "business"

    def _build_event_line_judgment(
        bundle: EventLineContextBundleRecord,
        *,
        viewer_role: ReviewViewerRole = "employee",
    ) -> EventLineJudgmentRecord:
        judgment_version = "event_line_judgment_v1"
        bundle_fingerprint = _bundle_fingerprint(bundle)
        coverage_score = _bundle_coverage_score(bundle)
        confidence_score = _bundle_confidence_score(bundle, coverage_score=coverage_score)
        safe_output_mode = _bundle_safe_output_mode(
            bundle,
            coverage_score=coverage_score,
            confidence_score=confidence_score,
        )
        publish_state = _bundle_publish_state(
            viewer_role=viewer_role,
            safe_output_mode=safe_output_mode,
        )
        operating_mode = _infer_bundle_operating_mode(bundle)
        happened_basis = _first_meaningful_line(
            bundle.currentWork,
            bundle.recentDecision,
            bundle.recentProgress,
            bundle.taskFacts[0].summary if bundle.taskFacts else None,
            bundle.summary,
            bundle.intent,
        ) or f"{bundle.lineName} 本周仍在推进。"
        what_happened = happened_basis
        if viewer_role == "admin":
            why_it_matters = _first_meaningful_line(
                bundle.collaborationGoals[0] if bundle.collaborationGoals else None,
                bundle.collaborationRelationship and f"这条线直接影响 {bundle.projectName or '当前项目'} 的合作判断是否能继续往前推。",
                f"这条线决定 {bundle.projectName or '当前项目'} 这一阶段能否形成值得继续投入的明确结论。",
            )
            manager_implication = (
                "管理层现在要判断的不是资料量，而是这条线能否收成明确结论、责任边界和下轮动作。"
            )
        elif viewer_role == "department_lead":
            why_it_matters = _first_meaningful_line(
                bundle.collaborationGoals[0] if bundle.collaborationGoals else None,
                "这条线会直接占用部门带宽；如果不及时收束，部门会继续消耗在来回确认和补背景上。",
            )
            manager_implication = "部门负责人要盯的是接口、带宽和收束动作，不是继续加任务。"
        else:
            why_it_matters = _first_meaningful_line(
                bundle.collaborationGoals[0] if bundle.collaborationGoals else None,
                "这条线如果不尽快收成结论，你本周的推进会停在交流和准备层，而不是形成结果。",
            )
            manager_implication = "现在最关键的不是再补一轮泛背景，而是把这次推进后的判断和下一步钉住。"
        core_blocker = _first_meaningful_line(
            bundle.currentBlocker,
            bundle.currentChallenges[0] if bundle.currentChallenges else None,
        )
        blocker_type = _infer_bundle_blocker_type(bundle)
        missing_context = _bundle_missing_context_lines(bundle)
        if not core_blocker:
            if missing_context:
                core_blocker = f"现在还缺 {'、'.join(missing_context)}，所以这条线难以被准确判断和推进。"
                blocker_type = "evidence"
            elif operating_mode == "meeting":
                core_blocker = "真正的阻碍不是见面本身，而是会后要形成什么结论、谁接动作还不够明确。"
            elif operating_mode == "delivery":
                core_blocker = "真正的阻碍不是继续补话术，而是交付边界、收口标准和责任人还不够明确。"
            elif operating_mode == "materials":
                core_blocker = "真正的阻碍不是资料数量，而是哪些资料足以支撑判断、哪些还只是堆放并不清楚。"
            else:
                core_blocker = "当前最关键的阻碍还没有被明确写成一句业务判断。"
        key_evidence = _bundle_key_evidence_lines(bundle, limit=3)
        evidence_summary = "；".join(key_evidence) if key_evidence else (
            f"已关联 {bundle.taskCount} 条任务、{bundle.meetingCount} 次会议、{bundle.attachmentCount} 份附件"
            + (f"、{bundle.supportRequestCount} 条支持请求" if bundle.supportRequestCount else "")
        )
        next_week_focus = _first_meaningful_line(
            bundle.nextStep,
            bundle.collaborationGoals[0] if bundle.collaborationGoals else None,
            bundle.intent,
        ) or (
            "把这条线压成一句明确结论和一个最小动作。"
            if operating_mode != "meeting"
            else "把会谈要确认的结论、关键人反馈和会后动作一次性钉住。"
        )
        minimum_action = _first_meaningful_line(
            bundle.nextStep,
            bundle.recentDecision and f"把“{bundle.recentDecision}”转成明确跟进行动",
        ) or (
            "明确这次推进后的责任人、时间点和会后动作。"
            if operating_mode == "meeting"
            else "把这条线收成可执行结论、责任人和最小后续动作。"
        )
        risk_if_ignored = {
            "business": "如果继续不收束，这条线会继续停在泛交流或泛合作表达层，管理层看不清是否值得继续加码。",
            "collaboration": "如果继续不处理协同接口，这条线会继续卡在等待与来回确认上，并很快变成跨周风险。",
            "decision": "如果继续不拍板，这条线会把推进变成反复确认，后续动作很难真正启动。",
            "structure": "如果继续不把归属和产出链挂清楚，资料会越来越多，但判断仍然会继续发虚。",
            "capacity": "如果继续不做取舍，这条线会继续和别的事项争抢带宽，导致推进质量下滑。",
            "evidence": "如果继续不补关键证据，后续所有判断都会停在大概正确但无法指导动作的层面。",
        }.get(blocker_type, "如果继续放着不管，这条线会继续停在推进表面，难以形成真正结果。")
        opportunity_if_amplified = (
            "如果现在就把这条线沉成结论、证据和后续动作，它能从一次推进变成可复制的合作样板。"
            if operating_mode == "meeting"
            else "如果现在就把这条线沉成结论、证据和后续动作，它就能从一次推进变成可复制的正向样板。"
        )
        people_or_products_hint = _bundle_people_or_products_hint(bundle)
        if people_or_products_hint:
            why_it_matters = f"{why_it_matters} {people_or_products_hint}".strip()
        return EventLineJudgmentRecord(
            eventLineId=bundle.eventLineId,
            title=bundle.lineName,
            viewerRole=viewer_role if viewer_role in {"employee", "department_lead", "admin"} else "employee",
            judgmentVersion=judgment_version,
            bundleFingerprint=bundle_fingerprint,
            coverageScore=coverage_score,
            confidenceScore=confidence_score,
            safeOutputMode=safe_output_mode,
            publishState=publish_state,
            whatHappened=what_happened,
            whyItMatters=why_it_matters,
            coreBlocker=core_blocker,
            blockerType=blocker_type,
            evidenceSummary=evidence_summary,
            managerImplication=manager_implication,
            nextWeekFocus=next_week_focus,
            minimumAction=minimum_action,
            riskIfIgnored=risk_if_ignored,
            opportunityIfAmplified=opportunity_if_amplified,
            evidenceRefs=bundle.evidenceRefs[:6],
            target=ReviewDashboardCardTargetRecord(
                targetType="event_line",
                targetId=bundle.eventLineId,
                targetLabel=bundle.lineName,
                evidenceRefs=bundle.evidenceRefs[:4],
            ),
        )

    def _build_ad_hoc_task_context_bundle(task: TaskRecord) -> EventLineContextBundleRecord:
        notebook = (
            get_client_notebook_response(state.db, task.clientId).organizationNotebookSnapshot
            if task.clientId
            else None
        )
        task_clauses = _task_clause_candidates(task.desc)
        task_fact = _event_line_context_fact(
            source_type="task",
            source_id=task.id,
            title=task.title,
            summary=_first_nonempty_text(
                task_clauses[0] if task_clauses else None,
                task.recentDecision,
                task.nextAction,
                task.currentBlocker,
            ) or task.title,
            happened_at=task.updatedAt,
        )
        attachment_facts = [
            fact
            for fact in (
                _event_line_context_fact(
                    source_type="attachment",
                    source_id=item.id,
                    title=item.title,
                    summary=_first_nonempty_text(
                        item.summary if hasattr(item, "summary") else None,
                        item.title,
                        item.path,
                    ),
                    happened_at=item.createdAt,
                )
                for item in task.attachments
            )
            if fact is not None
        ]
        clarification_facts = [
            _event_line_context_fact(
                source_type="clarification",
                source_id=f"{task.id}-memory-hint-{index}",
                title="已关联背景",
                summary=hint,
            )
            for index, hint in enumerate((task.memoryHints or [])[:3])
        ]
        clarification_facts = [fact for fact in clarification_facts if fact is not None]
        evidence_refs = _dedupe_dashboard_evidence_refs(
            [
                ReviewDashboardEvidenceRefRecord(
                    sourceType="task",
                    sourceId=task.id,
                    title=task.title,
                    summary=_first_nonempty_text(
                        task_clauses[0] if task_clauses else None,
                        task.recentDecision,
                        task.nextAction,
                        task.currentBlocker,
                    ),
                ),
                *[_context_fact_to_evidence_ref(fact) for fact in attachment_facts[:4]],
                *[_context_fact_to_evidence_ref(fact) for fact in clarification_facts[:2]],
            ]
        )
        return EventLineContextBundleRecord(
            eventLineId=task.eventLineId or f"task::{task.id}",
            lineName=task.eventLineName or task.title,
            businessCategory=task.businessCategory or "",
            stage=_first_nonempty_text(task.projectContext.stage if task.projectContext else None, task.orgContext.organizationFocusKey if task.orgContext else None) or "",
            summary=_first_nonempty_text(task.desc, task.projectContext.backgroundSummary if task.projectContext else None) or task.title,
            intent=_first_nonempty_text(task.projectContext.currentFocus if task.projectContext else None, task.nextAction, task_clauses[0] if task_clauses else None, task.title) or task.title,
            currentWork=_first_nonempty_text(task.projectContext.currentFocus if task.projectContext else None, task_clauses[0] if task_clauses else None, task.title) or task.title,
            currentBlocker=_first_nonempty_text(task.currentBlocker, task.projectContext.currentBlocker if task.projectContext else None) or "",
            recentDecision=_first_nonempty_text(task.recentDecision, task.projectContext.recentProgress if task.projectContext else None) or "",
            nextStep=_first_nonempty_text(task.nextAction, task.projectContext.nextAction if task.projectContext else None) or "",
            recentProgress=_first_nonempty_text(task.recentDecision, task.projectContext.recentProgress if task.projectContext else None) or "",
            projectName=task.clientName or "",
            collaborationRelationship=notebook.collaborationRelationship if notebook else "",
            organizationIntro=notebook.organizationIntro if notebook else "",
            currentChallenges=list(notebook.currentChallenges) if notebook else [],
            collaborationGoals=list(notebook.collaborationGoals) if notebook else [],
            keyPeople=list(notebook.keyPeople) if notebook else [],
            keyProducts=list(notebook.keyProducts) if notebook else [],
            recentFacts=list(notebook.recentFacts)[:5] if notebook else [],
            taskFacts=[task_fact] if task_fact else [],
            meetingFacts=[],
            attachmentFacts=attachment_facts[:6],
            clarificationFacts=clarification_facts[:4],
            evidenceRefs=evidence_refs[:8],
            trendSignals=[],
            taskCount=1,
            meetingCount=0,
            attachmentCount=len(attachment_facts),
            supportRequestCount=0,
            readiness=_event_line_bundle_readiness(
                notebook_confidence=notebook.confidence if notebook else 0.0,
                snapshot_confidence=0.0,
                task_count=1,
                attachment_count=len(attachment_facts),
                support_request_count=0,
                clarification_count=len(clarification_facts),
                current_work=_first_nonempty_text(task.projectContext.currentFocus if task.projectContext else None, task.desc, task.title) or "",
                current_blocker=_first_nonempty_text(task.currentBlocker, task.projectContext.currentBlocker if task.projectContext else None) or "",
                next_step=_first_nonempty_text(task.nextAction, task.projectContext.nextAction if task.projectContext else None) or "",
            ),
        )

    def _build_task_context_preview(task: TaskRecord) -> TaskContextPreviewRecord:
        session_user = get_cached_session_user()
        governance = _review_governance_with_members() if session_user else None
        viewer_role = _resolve_review_viewer_role(session_user, governance)
        bundle = _event_line_context_bundle(task.eventLineId) if task.eventLineId else None
        if bundle is None:
            bundle = _build_ad_hoc_task_context_bundle(task)
        judgment = _build_event_line_judgment(bundle, viewer_role=viewer_role)
        summary_chips = _dedupe_texts(
            [
                bundle.businessCategory or task.businessCategory or "",
                f"项目 · {bundle.projectName}" if bundle.projectName else "",
                f"事件线 · {bundle.lineName}" if bundle.lineName and task.eventLineId else "",
                f"阶段 · {bundle.stage}" if bundle.stage else "",
            ],
            limit=4,
        )
        return TaskContextPreviewRecord(
            taskId=task.id,
            clientId=task.clientId,
            clientName=task.clientName,
            contextBundle=bundle,
            judgment=judgment,
            judgmentVersion=judgment.judgmentVersion,
            bundleFingerprint=judgment.bundleFingerprint,
            coverageScore=judgment.coverageScore,
            confidenceScore=judgment.confidenceScore,
            safeOutputMode=judgment.safeOutputMode,
            publishState=judgment.publishState,
            summaryChips=summary_chips,
            readiness=bundle.readiness,
        )

    def _enrich_weekly_review_analysis_with_memory(
        analysis: WeeklyReviewAnalysisRecord | None,
        *,
        viewer_role: ReviewViewerRole = "employee",
    ) -> WeeklyReviewAnalysisRecord | None:
        if analysis is None or not analysis.eventLineSummaries:
            return analysis

        summary_updates: dict[str, dict[str, object]] = {}
        completeness_updates: dict[str, dict[str, object]] = {}
        bundles: list[EventLineContextBundleRecord] = []
        judgments: list[EventLineJudgmentRecord] = []

        event_line_ids = list(
            {
                _normalize_event_line_reference(item.eventLineId)
                for item in analysis.eventLineSummaries
                if _normalize_event_line_reference(item.eventLineId)
            }
        )
        for event_line_id in event_line_ids:
            memory_response = get_event_line_memory_response(state.db, event_line_id)
            snapshot = memory_response.eventLineMemorySnapshot
            if snapshot is None:
                continue
            clean_current_work = sanitize_memory_background_text(snapshot.currentWork, reject_generic=True, max_length=140)
            clean_blocker = sanitize_memory_background_text(snapshot.currentBlocker, reject_generic=True, max_length=140)
            clean_recent_decision = sanitize_memory_background_text(snapshot.recentDecision, reject_generic=True, max_length=140)
            clean_next_step = sanitize_memory_background_text(snapshot.nextStep, reject_generic=True, max_length=140)
            event_line_row = state.db.fetchone(
                "SELECT primary_client_id FROM event_lines WHERE id = ?",
                (event_line_id,),
            )
            client_id = str(event_line_row["primary_client_id"]) if event_line_row and event_line_row["primary_client_id"] else None
            notebook = get_client_notebook_response(state.db, client_id).organizationNotebookSnapshot if client_id else None
            linked_facts_preview = get_task_memory_enrichment(
                state.db,
                task_id=f"event-line:{event_line_id}",
                client_id=client_id,
                event_line_id=event_line_id,
            )[2]
            merged_missing_slots = list(
                dict.fromkeys(
                    [
                        *[
                            slot
                            for item in analysis.eventLineSummaries
                            if _normalize_event_line_reference(item.eventLineId) == event_line_id
                            for slot in item.missingSlots
                        ],
                        *[_memory_slot_label(slot) for slot in snapshot.clarificationNeeds],
                    ]
                )
            )[:5]
            background_sources = _memory_background_source_labels(
                has_notebook=bool(notebook and notebook.confidence > 0),
                has_snapshot=True,
                has_evidence=bool(snapshot.evidenceRefs),
                has_review_signals=bool(snapshot.analysisSignals),
                has_linked_facts=bool(linked_facts_preview),
                has_pending_clarification=bool(snapshot.clarificationNeeds),
            )
            summary_update: dict[str, object] = {
                "eventLineId": event_line_id,
                "memoryConfidence": round(snapshot.confidence, 2),
                "backgroundSources": background_sources,
                "missingSlots": merged_missing_slots,
            }
            if clean_current_work:
                summary_update["whatThisLineIs"] = f"这条线当前主要在推进：{clean_current_work}"
                summary_update["whatHappenedThisWeek"] = f"本周主要在推进：{clean_current_work}。"
            elif clean_recent_decision:
                summary_update["whatHappenedThisWeek"] = f"本周形成的关键决策是：{clean_recent_decision}。"
            if clean_blocker:
                summary_update["mainBlocker"] = clean_blocker
            if clean_next_step:
                summary_update["nextCriticalMove"] = clean_next_step
            summary_updates[event_line_id] = summary_update
            completeness_updates[event_line_id] = {
                "eventLineId": event_line_id,
                "memoryConfidence": round(snapshot.confidence, 2),
                "backgroundSources": background_sources,
                "missingSlots": merged_missing_slots,
            }
            bundle = _event_line_context_bundle(event_line_id, analysis=analysis)
            if bundle is not None:
                bundles.append(bundle)
                judgment = _build_event_line_judgment(bundle, viewer_role=viewer_role)
                judgments.append(judgment)
                clean_line_identity = sanitize_memory_background_text(
                    bundle.summary or bundle.intent,
                    reject_generic=True,
                    max_length=140,
                )
                summary_updates[event_line_id].update(
                    {
                        "whatThisLineIs": clean_line_identity or summary_updates[event_line_id].get("whatThisLineIs"),
                        "whatHappenedThisWeek": judgment.whatHappened,
                        "currentState": judgment.whyItMatters,
                        "mainBlocker": judgment.coreBlocker,
                        "nextCriticalMove": judgment.minimumAction,
                        "evidencePreview": _dedupe_texts(
                            [
                                judgment.evidenceSummary,
                                *bundle.recentFacts[:2],
                                *(fact.summary for fact in bundle.taskFacts[:2]),
                            ],
                            limit=4,
                        ),
                        "target": judgment.target,
                        "evidenceRefs": judgment.evidenceRefs,
                        "publishState": judgment.publishState,
                        "publishedAt": judgment.publishedAt,
                        "publishedBy": judgment.publishedBy,
                        "invalidatedAt": judgment.invalidatedAt,
                        "projectName": bundle.projectName or summary_update.get("projectName"),
                    }
                )
                completeness_updates[event_line_id].update(
                    {
                        "strongestSlots": _dedupe_texts(
                            [
                                "当前推进",
                                "当前阻碍" if bundle.currentBlocker else "",
                                "下一步" if bundle.nextStep else "",
                                "会议证据" if bundle.meetingCount else "",
                                "附件证据" if bundle.attachmentCount else "",
                            ],
                            limit=4,
                        )
                    }
                )

        if not summary_updates and not completeness_updates and not bundles and not judgments:
            return analysis

        judgment_by_id = {item.eventLineId: item for item in judgments}
        risk_cards = list(analysis.riskCards or [])
        risk_by_id = {item.eventLineId: item for item in risk_cards}
        for judgment in judgments:
            if judgment.eventLineId in risk_by_id:
                item = risk_by_id[judgment.eventLineId]
                risk_by_id[judgment.eventLineId] = item.model_copy(
                    update={
                        "statement": judgment.coreBlocker,
                        "whyNow": judgment.whyItMatters,
                        "ifIgnored": judgment.riskIfIgnored,
                        "suggestedAction": judgment.minimumAction,
                        "target": judgment.target,
                        "evidenceRefs": judgment.evidenceRefs,
                        "publishState": judgment.publishState,
                        "publishedAt": judgment.publishedAt,
                        "publishedBy": judgment.publishedBy,
                        "invalidatedAt": judgment.invalidatedAt,
                    }
                )
            else:
                risk_by_id[judgment.eventLineId] = EventLineRiskCardRecord(
                    eventLineId=judgment.eventLineId,
                    title=judgment.title,
                    riskType=(
                        "collaboration_friction"
                        if judgment.blockerType == "collaboration"
                        else "decision_lag"
                        if judgment.blockerType == "decision"
                        else "workflow_breakdown"
                        if judgment.blockerType in {"structure", "evidence"}
                        else "overload"
                        if judgment.blockerType == "capacity"
                        else "goal_drift"
                    ),
                    statement=judgment.coreBlocker,
                    forecastWindow="1w",
                    probability="high" if judgment.blockerType in {"decision", "capacity", "structure"} else "medium",
                    impactScope="org" if viewer_role == "admin" else "team" if viewer_role == "department_lead" else "project",
                    triggerSignals=_dedupe_texts([judgment.coreBlocker, judgment.evidenceSummary], limit=3),
                    whyNow=judgment.whyItMatters,
                    ifIgnored=judgment.riskIfIgnored,
                    suggestedAction=judgment.minimumAction,
                    ownerRole="该线负责人",
                    publishState=judgment.publishState,
                    publishedAt=judgment.publishedAt,
                    publishedBy=judgment.publishedBy,
                    invalidatedAt=judgment.invalidatedAt,
                    target=judgment.target,
                    evidenceRefs=judgment.evidenceRefs,
                )

        opportunity_cards = list(analysis.opportunityCards or [])
        opportunity_by_id = {item.eventLineId: item for item in opportunity_cards}
        for judgment in judgments:
            if judgment.eventLineId in opportunity_by_id:
                item = opportunity_by_id[judgment.eventLineId]
                opportunity_by_id[judgment.eventLineId] = item.model_copy(
                    update={
                        "statement": judgment.opportunityIfAmplified,
                        "upside": judgment.managerImplication,
                        "recommendedAmplifier": judgment.minimumAction,
                        "target": judgment.target,
                        "evidenceRefs": judgment.evidenceRefs,
                        "publishState": judgment.publishState,
                        "publishedAt": judgment.publishedAt,
                        "publishedBy": judgment.publishedBy,
                        "invalidatedAt": judgment.invalidatedAt,
                    }
                )
            else:
                opportunity_by_id[judgment.eventLineId] = EventLineOpportunityCardRecord(
                    eventLineId=judgment.eventLineId,
                    title=judgment.title,
                    opportunityType="momentum_building",
                    statement=judgment.opportunityIfAmplified,
                    forecastWindow="1w",
                    confidence="medium" if judgments else "low",
                    upside=judgment.managerImplication,
                    supportingSignals=_dedupe_texts([judgment.evidenceSummary, judgment.whatHappened], limit=3),
                    recommendedAmplifier=judgment.minimumAction,
                    ownerRole="该线负责人",
                    publishState=judgment.publishState,
                    publishedAt=judgment.publishedAt,
                    publishedBy=judgment.publishedBy,
                    invalidatedAt=judgment.invalidatedAt,
                    target=judgment.target,
                    evidenceRefs=judgment.evidenceRefs,
                )

        return analysis.model_copy(
            update={
                "eventLineSummaries": [
                    item.model_copy(
                        update=summary_updates.get(
                            _normalize_event_line_reference(item.eventLineId),
                            {"eventLineId": _normalize_event_line_reference(item.eventLineId)},
                        )
                    )
                    for item in analysis.eventLineSummaries
                ],
                "eventLineCompleteness": [
                    item.model_copy(
                        update=completeness_updates.get(
                            _normalize_event_line_reference(item.eventLineId),
                            {"eventLineId": _normalize_event_line_reference(item.eventLineId)},
                        )
                    )
                    for item in analysis.eventLineCompleteness
                ],
                "eventLineContextBundles": bundles,
                "eventLineJudgments": judgments,
                "riskCards": list(risk_by_id.values())[:6],
                "opportunityCards": list(opportunity_by_id.values())[:6],
                "nextWeekFocus": _dedupe_texts(
                    [
                        *[f"{item.title}｜{item.nextWeekFocus}" for item in judgments],
                        *(analysis.nextWeekFocus or []),
                    ],
                    limit=6,
                ),
            }
        )

    def _format_review_week_label(value: datetime.date) -> str:
        iso = value.isocalendar()
        return f"{iso.year}-W{iso.week:02d}"

    def _previous_review_week_labels(week_label: str, *, count: int = 2) -> list[str]:
        bounds = _week_bounds(week_label)
        if bounds is None:
            return []
        start, _ = bounds
        return [
            _format_review_week_label(start - timedelta(days=7 * offset))
            for offset in range(1, count + 1)
        ]

    def _parse_iso_datetime(value: str | None) -> datetime | None:
        trimmed = str(value or "").strip()
        if not trimmed:
            return None
        normalized = trimmed[:-1] + "+00:00" if trimmed.endswith("Z") else trimmed
        try:
            return datetime.fromisoformat(normalized)
        except ValueError:
            return None

    def _historical_work_review_rows(week_labels: list[str]) -> list[dict[str, object]]:
        if not week_labels:
            return []
        operator_id = str(current_operator_row()["id"])
        placeholders = _sql_placeholders(tuple(week_labels))
        rows = state.db.fetchall(
            f"""
            SELECT e.*
            FROM weekly_review_task_entries e
            INNER JOIN weekly_reviews r ON r.id = e.review_id
            WHERE r.operator_id = ?
              AND e.content_domain = 'work'
              AND e.week_label IN ({placeholders})
            ORDER BY e.reviewed_at DESC, e.created_at DESC
            """,
            (operator_id, *week_labels),
        )
        return [dict(row) for row in rows]

    def _review_snapshot_dict(row: dict[str, object]) -> dict[str, object]:
        payload = from_json(str(row.get("task_snapshot_json") or "{}"), {})
        return payload if isinstance(payload, dict) else {}

    def _review_snapshot_org_context(snapshot: dict[str, object]) -> dict[str, object]:
        payload = snapshot.get("orgContext")
        return payload if isinstance(payload, dict) else {}

    def _review_snapshot_project_context(snapshot: dict[str, object]) -> dict[str, object]:
        payload = snapshot.get("projectContext")
        return payload if isinstance(payload, dict) else {}

    def _review_snapshot_event_line_context(snapshot: dict[str, object]) -> dict[str, object]:
        payload = snapshot.get("eventLineContext")
        return payload if isinstance(payload, dict) else {}

    def _review_row_event_line_id(row: dict[str, object]) -> str:
        snapshot = _review_snapshot_dict(row)
        event_line_context = _review_snapshot_event_line_context(snapshot)
        return _normalize_event_line_reference(
            _first_nonempty_text(
                snapshot.get("eventLineId"),
                event_line_context.get("id"),
            )
        )

    def _review_row_support_requested(row: dict[str, object]) -> bool:
        structured = coerce_review_structured_note(row.get("structured_note_json"))
        return bool(
            structured.supportNeeded.strip()
            or structured.lightweightTag.strip() in {"需要支持", "资源不够", "等待他人"}
        )

    def _review_row_blocker_text(row: dict[str, object]) -> str:
        structured = coerce_review_structured_note(row.get("structured_note_json"))
        snapshot = _review_snapshot_dict(row)
        event_line_context = _review_snapshot_event_line_context(snapshot)
        project_context = _review_snapshot_project_context(snapshot)
        return _first_nonempty_text(
            structured.blockerReason,
            event_line_context.get("currentBlocker"),
            project_context.get("currentBlocker"),
            row.get("note"),
        )

    def _review_row_needs_review(row: dict[str, object]) -> bool:
        snapshot = _review_snapshot_dict(row)
        org_context = _review_snapshot_org_context(snapshot)
        return bool(org_context.get("needsReview"))

    def _recent_due_date_change_stats(task_ids: list[str], *, lookback_days: int = 21) -> dict[str, dict[str, object]]:
        wanted = [task_id for task_id in task_ids if task_id]
        if not wanted:
            return {}
        since = (datetime.now() - timedelta(days=lookback_days)).isoformat(timespec="seconds")
        placeholders = _sql_placeholders(tuple(wanted))
        rows = state.db.fetchall(
            f"""
            SELECT entity_id, detail_json, created_at
            FROM activity_logs
            WHERE entity_type = 'task'
              AND action = 'task.update'
              AND entity_id IN ({placeholders})
              AND created_at >= ?
            ORDER BY created_at DESC
            """,
            (*wanted, since),
        )
        stats: dict[str, dict[str, object]] = {}
        for row in rows:
            detail = from_json(row["detail_json"], {}) if row["detail_json"] else {}
            if not isinstance(detail, dict):
                continue
            if "dueDate" not in detail and "ddl" not in detail:
                continue
            task_id = str(row["entity_id"])
            current = stats.setdefault(task_id, {"changeCount": 0, "weeks": set(), "lastChangedAt": None})
            current["changeCount"] = int(current["changeCount"]) + 1
            changed_at = _parse_iso_datetime(str(row["created_at"]))
            if changed_at is not None:
                weeks = current["weeks"]
                if isinstance(weeks, set):
                    weeks.add(_format_review_week_label(changed_at.date()))
                current["lastChangedAt"] = str(row["created_at"])
        return stats

    def _merge_trend_signals(
        existing: list[TrendSignalRecord],
        additions: list[TrendSignalRecord],
    ) -> list[TrendSignalRecord]:
        merged: dict[str, TrendSignalRecord] = {signal.key: signal for signal in existing}
        for signal in additions:
            merged[signal.key] = signal
        severity_rank = {"high": 0, "medium": 1, "low": 2}
        return sorted(
            merged.values(),
            key=lambda item: (severity_rank.get(item.severity, 3), item.title),
        )[:8]

    def _enrich_weekly_review_analysis_with_operational_trends(
        analysis: WeeklyReviewAnalysisRecord | None,
        items: list[WeeklyReviewTaskEntryRecord],
        week_label: str,
    ) -> WeeklyReviewAnalysisRecord | None:
        if analysis is None or not items:
            return analysis

        previous_weeks = _previous_review_week_labels(week_label, count=2)
        if not previous_weeks:
            return analysis

        current_task_ids = [item.taskId for item in items if item.taskId]
        current_titles = {item.taskId: item.taskSnapshot.title for item in items if item.taskId}
        historical_rows = _historical_work_review_rows(previous_weeks)
        if not historical_rows and not current_task_ids:
            return analysis

        historical_review_tasks: set[str] = set()
        historical_review_event_lines: set[str] = set()
        historical_support_tasks: set[str] = set()
        historical_support_event_lines: set[str] = set()
        historical_blockers: dict[str, str] = {}

        for row in historical_rows:
            task_id = str(row.get("task_id") or "").strip()
            event_line_id = _review_row_event_line_id(row)
            if _review_row_needs_review(row):
                if task_id:
                    historical_review_tasks.add(task_id)
                if event_line_id:
                    historical_review_event_lines.add(event_line_id)
            if _review_row_support_requested(row):
                if task_id:
                    historical_support_tasks.add(task_id)
                if event_line_id:
                    historical_support_event_lines.add(event_line_id)
            blocker_text = _review_row_blocker_text(row)
            if blocker_text and event_line_id:
                historical_blockers[event_line_id] = blocker_text

        new_signals: list[TrendSignalRecord] = []

        recent_due_date_change_stats = _recent_due_date_change_stats(current_task_ids)
        repeated_reschedules: list[tuple[WeeklyReviewTaskEntryRecord, dict[str, object]]] = []
        for item in items:
            stats = recent_due_date_change_stats.get(item.taskId)
            if not stats:
                continue
            change_count = int(stats.get("changeCount") or 0)
            week_count = len(stats.get("weeks") or set())
            if change_count >= 2 and week_count >= 2:
                repeated_reschedules.append((item, stats))
        if repeated_reschedules:
            related_task_ids = [item.taskId for item, _ in repeated_reschedules[:5]]
            sample_titles = "、".join(current_titles.get(task_id, task_id) for task_id in related_task_ids[:2])
            new_signals.append(
                TrendSignalRecord(
                    key="repeat_reschedule::cross_week",
                    title="改期开始连续化",
                    statement=f"过去 2-3 周内至少 {len(repeated_reschedules)} 条任务反复改期，其中 {sample_titles or '相关任务'} 已跨周调整，说明排期不只是本周波动，而是推进路径尚未稳定。",
                    signalType="repeat_reschedule",
                    severity="high" if len(repeated_reschedules) >= 2 else "medium",
                    windowLabel="连续 2-3 周",
                    relatedEventLineId=None,
                    relatedTaskIds=related_task_ids,
                    evidenceRefs=[
                        ReviewDashboardEvidenceRefRecord(
                            sourceType="task",
                            sourceId=item.taskId,
                            title=item.taskSnapshot.title,
                            summary=f"过去 2-3 周改期 {int(stats.get('changeCount') or 0)} 次",
                        )
                        for item, stats in repeated_reschedules[:4]
                    ],
                    target=ReviewDashboardCardTargetRecord(
                        targetType="task_view",
                        targetId="builtin:risk",
                        targetLabel="风险视图",
                        targetFilters={"onlyRisky": True, "relatedTaskIds": related_task_ids},
                    ),
                )
            )

        repeated_review_items = [
            item
            for item in items
            if bool(item.taskSnapshot.orgContext and item.taskSnapshot.orgContext.needsReview)
            and (
                item.taskId in historical_review_tasks
                or _normalize_event_line_reference(item.taskSnapshot.eventLineId or (item.taskSnapshot.eventLineContext.id if item.taskSnapshot.eventLineContext else None)) in historical_review_event_lines
            )
        ]
        if repeated_review_items:
            related_task_ids = [item.taskId for item in repeated_review_items[:5]]
            new_signals.append(
                TrendSignalRecord(
                    key="repeat_review_pending::cross_week",
                    title="待复核事项连续两周未收束",
                    statement=f"当前仍有 {len(repeated_review_items)} 条任务连续两周停在复核/确认链上，问题已经从一次性审批延迟，变成持续拖慢推进的结构性回收链。",
                    signalType="repeat_review_pending",
                    severity="high" if len(repeated_review_items) >= 2 else "medium",
                    windowLabel="连续 2 周",
                    relatedEventLineId=None,
                    relatedTaskIds=related_task_ids,
                    evidenceRefs=[
                        ReviewDashboardEvidenceRefRecord(
                            sourceType="task",
                            sourceId=item.taskId,
                            title=item.taskSnapshot.title,
                            summary="连续两周待复核或待确认",
                        )
                        for item in repeated_review_items[:4]
                    ],
                    target=ReviewDashboardCardTargetRecord(
                        targetType="task_view",
                        targetId="builtin:risk",
                        targetLabel="风险视图",
                        targetFilters={"onlyRisky": True, "needsReview": True, "relatedTaskIds": related_task_ids},
                    ),
                )
            )

        repeated_support_items = [
            item
            for item in items
            if (
                item.structuredNote.supportNeeded.strip()
                or item.structuredNote.lightweightTag.strip() in {"需要支持", "资源不够", "等待他人"}
            )
            and (
                item.taskId in historical_support_tasks
                or _normalize_event_line_reference(item.taskSnapshot.eventLineId or (item.taskSnapshot.eventLineContext.id if item.taskSnapshot.eventLineContext else None)) in historical_support_event_lines
            )
        ]
        if repeated_support_items:
            related_task_ids = [item.taskId for item in repeated_support_items[:5]]
            new_signals.append(
                TrendSignalRecord(
                    key="repeat_support_request::cross_week",
                    title="支持依赖连续两周未化解",
                    statement=f"当前至少 {len(repeated_support_items)} 条任务连续两周提到支持或外部依赖，这不再是偶发协作缺口，而是需要管理层介入收束的协作链问题。",
                    signalType="repeat_support_request",
                    severity="medium" if len(repeated_support_items) == 1 else "high",
                    windowLabel="连续 2 周",
                    relatedEventLineId=None,
                    relatedTaskIds=related_task_ids,
                    evidenceRefs=[
                        ReviewDashboardEvidenceRefRecord(
                            sourceType="task",
                            sourceId=item.taskId,
                            title=item.taskSnapshot.title,
                            summary=_first_nonempty_text(item.structuredNote.supportNeeded, item.structuredNote.lightweightTag, item.note),
                        )
                        for item in repeated_support_items[:4]
                    ],
                    target=ReviewDashboardCardTargetRecord(
                        targetType="task_view",
                        targetId="builtin:risk",
                        targetLabel="风险视图",
                        targetFilters={"onlyRisky": True, "relatedTaskIds": related_task_ids},
                    ),
                )
            )

        escalating_blocker_items: list[tuple[WeeklyReviewTaskEntryRecord, str]] = []
        for item in items:
            event_line_id = _normalize_event_line_reference(
                item.taskSnapshot.eventLineId or (item.taskSnapshot.eventLineContext.id if item.taskSnapshot.eventLineContext else None)
            )
            if not event_line_id:
                continue
            current_blocker = _first_nonempty_text(
                item.structuredNote.blockerReason,
                item.taskSnapshot.eventLineContext.currentBlocker if item.taskSnapshot.eventLineContext else None,
                item.taskSnapshot.projectContext.currentBlocker if item.taskSnapshot.projectContext else None,
            )
            if not current_blocker or event_line_id not in historical_blockers:
                continue
            escalating_blocker_items.append((item, current_blocker))
        if escalating_blocker_items:
            related_task_ids = [item.taskId for item, _ in escalating_blocker_items[:5]]
            primary_item, primary_blocker = escalating_blocker_items[0]
            primary_event_line_id = _normalize_event_line_reference(
                primary_item.taskSnapshot.eventLineId or (primary_item.taskSnapshot.eventLineContext.id if primary_item.taskSnapshot.eventLineContext else None)
            )
            new_signals.append(
                TrendSignalRecord(
                    key=f"escalating_blocker::{primary_event_line_id or primary_item.taskId}",
                    title=f"{primary_item.taskSnapshot.eventLineName or primary_item.taskSnapshot.title} 阻塞升级",
                    statement=f"{primary_item.taskSnapshot.eventLineName or primary_item.taskSnapshot.title} 连续两周都卡在“{primary_blocker}”，如果下一周还不收束，这条线会继续从局部问题升级成趋势性风险。",
                    signalType="escalating_blocker",
                    severity="high",
                    windowLabel="连续 2 周",
                    relatedEventLineId=primary_event_line_id or None,
                    relatedTaskIds=related_task_ids,
                    evidenceRefs=[
                        ReviewDashboardEvidenceRefRecord(
                            sourceType="task",
                            sourceId=item.taskId,
                            title=item.taskSnapshot.title,
                            summary=blocker_text,
                        )
                        for item, blocker_text in escalating_blocker_items[:4]
                    ],
                    target=ReviewDashboardCardTargetRecord(
                        targetType="event_line" if primary_event_line_id else "task_view",
                        targetId=primary_event_line_id or "builtin:risk",
                        targetLabel=primary_item.taskSnapshot.eventLineName or primary_item.taskSnapshot.title,
                        targetFilters={} if primary_event_line_id else {"onlyRisky": True, "relatedTaskIds": related_task_ids},
                    ),
                )
            )

        if not new_signals:
            return analysis

        return analysis.model_copy(
            update={
                "trendSignals": _merge_trend_signals(list(analysis.trendSignals or []), new_signals),
            }
        )

    def _narratives_from_event_line_judgments(
        analysis: WeeklyReviewAnalysisRecord | None,
    ) -> list[NarrativeAnalysisRecord]:
        if analysis is None or not analysis.eventLineJudgments:
            return []
        bundle_by_id = {
            bundle.eventLineId: bundle
            for bundle in (analysis.eventLineContextBundles or [])
            if bundle.eventLineId
        }
        narratives: list[NarrativeAnalysisRecord] = []
        for judgment in analysis.eventLineJudgments[:4]:
            bundle = bundle_by_id.get(judgment.eventLineId)
            missing_lines: list[str] = []
            if bundle:
                if not bundle.keyPeople:
                    missing_lines.append("关键对象和角色")
                if not bundle.collaborationRelationship:
                    missing_lines.append("合作关系")
                if not bundle.recentFacts and not bundle.meetingFacts and not bundle.attachmentFacts:
                    missing_lines.append("近期证据")
            current_progress = judgment.nextWeekFocus
            if bundle and bundle.recentProgress:
                current_progress = bundle.recentProgress
            narratives.append(
                NarrativeAnalysisRecord(
                    eventLineId=judgment.eventLineId,
                    eventLineName=judgment.title,
                    clientName=bundle.projectName if bundle and bundle.projectName else None,
                    whatThisIs=judgment.whatHappened,
                    whyImportant=judgment.whyItMatters,
                    currentProgress=current_progress,
                    missingUnderstanding="、".join(missing_lines) if missing_lines else judgment.evidenceSummary,
                    riskNote=judgment.riskIfIgnored or None,
                    minimumAction=judgment.minimumAction or None,
                    managementAdvice=judgment.managerImplication or None,
                    contextLayersUsed=[
                        label
                        for label, available in [
                            ("organization_dna", True),
                            ("client_profile", bool(bundle and bundle.projectName)),
                            ("cooperation_relationship", bool(bundle and bundle.collaborationRelationship)),
                            ("event_line_history", bool(bundle and bundle.recentFacts)),
                            ("current_tasks", True),
                        ]
                        if available
                    ],
                    confidenceLevel="high"
                    if judgment.safeOutputMode == "full_judgment"
                    else "medium"
                    if judgment.safeOutputMode == "summary_only"
                    else "low",
                )
            )
        narratives.sort(key=lambda item: ((item.eventLineName or "").strip(), (item.eventLineId or "").strip()))
        return narratives

    def _weekly_overview_cache_payload(
        *,
        week_label: str,
        items: list[WeeklyReviewTaskEntryRecord],
        narratives: list[NarrativeAnalysisRecord],
        org_modules: list[OrganizationDnaModuleRecord],
    ) -> dict[str, object]:
        cache_version = "v2-line-cards"
        task_signatures = sorted(
            [
            {
                "taskId": item.taskId,
                "title": item.taskSnapshot.title,
                "note": item.note,
                "client": item.taskSnapshot.clientName,
                "eventLine": item.taskSnapshot.eventLineName,
            }
            for item in items
            ],
            key=lambda item: (
                str(item.get("eventLine") or ""),
                str(item.get("client") or ""),
                str(item.get("title") or ""),
                str(item.get("taskId") or ""),
            ),
        )
        narrative_signatures = sorted(
            [
            {
                "eventLineId": item.eventLineId,
                "eventLineName": item.eventLineName,
                "whatThisIs": item.whatThisIs,
                "whyImportant": item.whyImportant,
                "currentProgress": item.currentProgress,
                "missingUnderstanding": item.missingUnderstanding,
            }
            for item in narratives
            ],
            key=lambda item: (
                str(item.get("eventLineName") or ""),
                str(item.get("eventLineId") or ""),
            ),
        )
        module_signatures = sorted(
            [
            {
                "moduleKey": item.moduleKey,
                "title": item.title,
                "updatedAt": item.updatedAt,
                "summary": item.summary[:160],
            }
            for item in org_modules[:6]
            ],
            key=lambda item: (
                str(item.get("title") or ""),
                str(item.get("moduleKey") or ""),
            ),
        )
        fingerprint_source = json.dumps(
            {
                "cacheVersion": cache_version,
                "weekLabel": week_label,
                "tasks": task_signatures,
                "narratives": narrative_signatures,
                "modules": module_signatures,
            },
            ensure_ascii=False,
            sort_keys=True,
        )
        return {
            "cacheVersion": cache_version,
            "fingerprint": hashlib.sha1(fingerprint_source.encode("utf-8")).hexdigest(),
        }

    def _load_cached_weekly_overview(
        *,
        week_label: str,
        fingerprint: str,
    ) -> tuple[str, list[str], list[str]] | None:
        raw = state.db.get_setting(f"weekly_overview_cache::{week_label}", "")
        if not raw:
            return None
        try:
            payload = json.loads(raw)
        except Exception:
            return None
        if not isinstance(payload, dict) or str(payload.get("fingerprint") or "") != fingerprint:
            return None
        return (
            str(payload.get("overview") or ""),
            [str(item) for item in payload.get("focusLines") or [] if str(item).strip()],
            [str(item) for item in payload.get("nextFocus") or [] if str(item).strip()],
        )

    def _save_cached_weekly_overview(
        *,
        week_label: str,
        fingerprint: str,
        overview: str,
        focus_lines: list[str],
        next_focus: list[str],
    ) -> None:
        state.db.set_setting(
            f"weekly_overview_cache::{week_label}",
            json.dumps(
                {
                    "cacheVersion": "v2-line-cards",
                    "fingerprint": fingerprint,
                    "overview": overview,
                    "focusLines": focus_lines,
                    "nextFocus": next_focus,
                    "updatedAt": now_iso(),
                },
                ensure_ascii=False,
            ),
        )

    def local_rollup_work_items(week_label: str) -> list[WeeklyReviewTaskEntryRecord]:
        rows = state.db.fetchall(
            """
            SELECT e.*
            FROM weekly_review_task_entries e
            INNER JOIN weekly_reviews r ON r.id = e.review_id
            WHERE e.week_label = ? AND e.content_domain = 'work'
            ORDER BY e.reviewed_at DESC, e.created_at DESC
            """,
            (week_label,),
        )
        items: list[WeeklyReviewTaskEntryRecord] = []
        for row in rows:
            task_id = str(row["task_id"])
            task_rows = fetch_tasks("t.id = ?", (task_id,))
            note = str(row["note"] or "")
            structured_note = coerce_review_structured_note(row["structured_note_json"])
            snapshot = from_json(str(row["task_snapshot_json"] or "{}"), {})
            if task_rows:
                task = task_rows[0]
                merged_snapshot = {
                    **_task_snapshot_from_task(task, state.db),
                    **(snapshot if isinstance(snapshot, dict) else {}),
                }
                items.append(
                    _review_entry_from_task(
                        task=task,
                        week_label=week_label,
                        content_domain="work",
                        review_id=str(row["review_id"]),
                        note=note,
                        structured_note=structured_note,
                        reviewed_at=str(row["reviewed_at"]) if row["reviewed_at"] else None,
                        snapshot=merged_snapshot,
                        db=state.db,
                    )
                )
                continue
            if isinstance(snapshot, dict) and snapshot:
                items.append(
                    WeeklyReviewTaskEntryRecord(
                        id=str(row["id"]),
                        reviewId=str(row["review_id"]),
                        taskId=task_id,
                        weekLabel=week_label,
                        contentDomain="work",
                        note=note,
                        structuredNote=structured_note,
                        reviewedAt=str(row["reviewed_at"]) if row["reviewed_at"] else None,
                        taskSnapshot=snapshot,  # type: ignore[arg-type]
                    )
                )
        return items

    def build_executive_review_overlay(
        week_label: str,
    ) -> tuple[HierarchyReportRecord | None, list[HierarchyReportRecord], ReviewSimulationBundleRecord | None]:
        session_user = get_cached_session_user()
        if not session_user:
            return None, [], None
        governance = _review_governance_with_members()
        viewer_role = _resolve_review_viewer_role(session_user, governance)
        base_org_modules = list_organization_dna_modules()
        org_model_profile: OrgModelProfileRecord | None = None
        try:
            org_model_payload = cloud_request("GET", "/api/v1/settings/org-model/profile")
            if isinstance(org_model_payload, dict):
                org_model_profile = OrgModelProfileRecord(**org_model_payload)
        except HTTPException:
            org_model_profile = None
        work_items = local_rollup_work_items(week_label)
        agent_work_items = build_agent_weekly_review_items(
            db=state.db,
            week_label=week_label,
            thread_sync_path=THREAD_SYNC_DOC_PATH,
        )
        org_modules = build_review_context_modules([*work_items, *agent_work_items], base_org_modules)
        executive_org_report, department_reports = build_executive_review_rollup(
            week_label=week_label,
            work_items=[*work_items, *agent_work_items],
            governance=governance,
            organization_dna_modules=org_modules,
            org_model_profile=org_model_profile,
        )
        if viewer_role != "admin":
            lead_department = _review_department_for_session_user(session_user, governance)
            if not lead_department:
                return None, [], None
            return None, [report for report in department_reports if report.scopeRefId == lead_department.name], None
        simulation_bundle = build_review_simulation_bundle(
            week_label=week_label,
            organization_dna_modules=base_org_modules,
        )
        return executive_org_report, department_reports, simulation_bundle

    def augment_review_response(response: ReviewResponse, week_label: str | None = None) -> ReviewResponse:
        target_week = week_label or (response.currentReview.weekLabel if response.currentReview else current_review_week_label())
        work_analysis = response.workAnalysis
        personal_analysis = response.personalAnalysis
        session_user = get_cached_session_user()
        governance = _review_governance_with_members() if session_user else None
        viewer_role = _resolve_review_viewer_role(session_user, governance)
        org_model_profile: OrgModelProfileRecord | None = None
        try:
            org_model_payload = cloud_request("GET", "/api/v1/settings/org-model/profile")
            if isinstance(org_model_payload, dict):
                org_model_profile = OrgModelProfileRecord(**org_model_payload)
        except HTTPException:
            org_model_profile = None
        if work_analysis is None or personal_analysis is None:
            computed_work_analysis, computed_personal_analysis = build_review_analyses(
                target_week,
                response.workItems,
                response.personalItems,
                org_model_profile=org_model_profile,
                viewer_role=viewer_role,
            )
            work_analysis = work_analysis or computed_work_analysis
            personal_analysis = personal_analysis or computed_personal_analysis
        work_analysis = _enrich_weekly_review_analysis_with_memory(work_analysis, viewer_role=viewer_role)
        work_analysis = _enrich_weekly_review_analysis_with_operational_trends(
            work_analysis,
            response.workItems,
            target_week,
        )
        if work_analysis is not None and response.workItems:
            narrative_modules = build_review_context_modules(response.workItems, list_organization_dna_modules())
            narrative_analyses = _narratives_from_event_line_judgments(work_analysis)
            cache_payload = _weekly_overview_cache_payload(
                week_label=target_week,
                items=response.workItems,
                narratives=narrative_analyses,
                org_modules=narrative_modules,
            )
            fingerprint = str(cache_payload["fingerprint"])
            cached_overview = _load_cached_weekly_overview(
                week_label=target_week,
                fingerprint=fingerprint,
            )
            if cached_overview is not None:
                weekly_overview, weekly_focus_lines, weekly_next_focus = cached_overview
            else:
                weekly_overview, weekly_focus_lines, weekly_next_focus = build_weekly_overview_draft(
                    ai=state.ai,
                    week_label=target_week,
                    items=response.workItems,
                    org_dna_modules=narrative_modules,
                    narratives=narrative_analyses,
                    fallback_overview=work_analysis.weeklyOverview,
                    fallback_focus_lines=work_analysis.weeklyFocusLines,
                    fallback_next_focus=work_analysis.weeklyNextFocus,
                )
                _save_cached_weekly_overview(
                    week_label=target_week,
                    fingerprint=fingerprint,
                    overview=weekly_overview,
                    focus_lines=weekly_focus_lines,
                    next_focus=weekly_next_focus,
                )
            work_analysis = work_analysis.model_copy(
                update={
                    "narrativeAnalyses": narrative_analyses,
                    "weeklyOverview": weekly_overview,
                    "weeklyFocusLines": weekly_focus_lines,
                    "weeklyNextFocus": weekly_next_focus,
                }
            )
        self_report = response.selfReport
        if self_report is None and work_analysis is not None and response.workItems:
            self_report = build_employee_review_report(
                week_label=target_week,
                scope_ref_id=(response.currentReview.userId if response.currentReview else "self"),
                items=response.workItems,
                analysis=work_analysis,
                org_model_profile=org_model_profile,
                viewer_role=viewer_role,
            )
        executive_org_report, department_reports, simulation_bundle = build_executive_review_overlay(target_week)
        if session_user is None:
            return response.model_copy(
                update={
                    "workAnalysis": work_analysis,
                    "personalAnalysis": personal_analysis,
                    "selfReport": self_report,
                }
            )
        if session_user.primaryRole != "admin":
            return response.model_copy(
                update={
                    "workAnalysis": work_analysis,
                    "personalAnalysis": personal_analysis,
                    "selfReport": self_report,
                    "executiveOrgReport": None,
                    "departmentReports": department_reports,
                    "agentDepartmentDigests": [],
                    "agentDepartmentPlans": [],
                    "simulationBundle": None,
                }
            )
        agent_department_digests = build_agent_weekly_digests(
            db=state.db,
            week_label=target_week,
            thread_sync_path=THREAD_SYNC_DOC_PATH,
        )
        agent_department_plans = build_agent_weekly_plans(
            db=state.db,
            week_label=target_week,
            thread_sync_path=THREAD_SYNC_DOC_PATH,
        )
        return response.model_copy(
            update={
                "workAnalysis": work_analysis,
                "personalAnalysis": personal_analysis,
                "selfReport": self_report,
                "executiveOrgReport": executive_org_report,
                "departmentReports": department_reports,
                "agentDepartmentDigests": agent_department_digests,
                "agentDepartmentPlans": agent_department_plans,
                "simulationBundle": simulation_bundle,
            }
        )

    def local_review_dashboard(week_label: str | None = None) -> ReviewResponse:
        target_week = week_label or current_review_week_label()
        review_row = local_review_row_for_week(target_week)
        review_entries = local_review_entries_by_task(str(review_row["id"])) if review_row else {}
        tasks_in_week = [task for task in fetch_tasks() if _task_in_week(task, target_week)]
        tasks_by_id = {task.id: task for task in tasks_in_week}
        for task_id in review_entries:
            if task_id in tasks_by_id:
                continue
            task_rows = fetch_tasks("t.id = ?", (task_id,))
            if task_rows:
                tasks_by_id[task_id] = task_rows[0]
        org_modules = list_organization_dna_modules()
        session_user = get_cached_session_user()
        governance = _review_governance_with_members() if session_user else None
        viewer_role = _resolve_review_viewer_role(session_user, governance)
        work_items: list[WeeklyReviewTaskEntryRecord] = []
        personal_items: list[WeeklyReviewTaskEntryRecord] = []
        for task in tasks_by_id.values():
            stored = review_entries.get(task.id)
            note = str(stored["note"]) if stored else ""
            structured_note = coerce_review_structured_note(stored.get("structured_note_json")) if stored else empty_review_structured_note()
            if not note and stored:
                note = compose_review_note(structured_note, "")
            reviewed_at = str(stored["reviewed_at"]) if stored else None
            snapshot = from_json(str(stored["task_snapshot_json"]), {}) if stored else None
            merged_snapshot = {
                **_task_snapshot_from_task(task, state.db),
                **(snapshot if isinstance(snapshot, dict) else {}),
            } if stored else None
            content_domain = "personal" if is_private_task(task) else "work"
            item = _review_entry_from_task(
                task=task,
                week_label=target_week,
                content_domain=content_domain,
                review_id=str(review_row["id"]) if review_row else None,
                note=note,
                structured_note=structured_note,
                reviewed_at=reviewed_at,
                snapshot=merged_snapshot,
                db=state.db,
            )
            if content_domain == "personal":
                personal_items.append(item)
            else:
                work_items.append(item)
        for task_id, stored in review_entries.items():
            if task_id in tasks_by_id:
                continue
            note = str(stored["note"] or "")
            structured_note = coerce_review_structured_note(stored.get("structured_note_json"))
            if not note:
                note = compose_review_note(structured_note, "")
            snapshot = from_json(str(stored["task_snapshot_json"] or "{}"), {})
            if not isinstance(snapshot, dict) or not snapshot:
                continue
            content_domain = str(stored.get("content_domain") or "work")
            item = WeeklyReviewTaskEntryRecord(
                id=str(stored["id"]),
                reviewId=str(stored["review_id"]),
                taskId=task_id,
                weekLabel=target_week,
                contentDomain=content_domain,  # type: ignore[arg-type]
                note=note,
                structuredNote=structured_note,
                reviewedAt=str(stored["reviewed_at"]) if stored["reviewed_at"] else None,
                taskSnapshot=snapshot,  # type: ignore[arg-type]
            )
            if content_domain == "personal":
                personal_items.append(item)
            else:
                work_items.append(item)
        current_review = build_local_review_record(review_row) if review_row else build_preview_review_record(target_week)
        work_analysis, personal_analysis = build_review_analyses(
            target_week,
            work_items,
            personal_items,
            org_modules,
            viewer_role=viewer_role,
        )
        self_report = None
        base_response = ReviewResponse(
            currentReview=current_review,
            workItems=work_items,
            personalItems=personal_items,
            workAnalysis=work_analysis,
            personalAnalysis=personal_analysis,
            selfReport=self_report,
            plans=[],
        )
        return augment_review_response(base_response, target_week)

    @app.get("/api/public/task-attachments/{attachment_id}")
    def proxy_cloud_task_attachment(attachment_id: str) -> Response:
        if not get_cloud_token():
            raise HTTPException(status_code=404, detail="Attachment not found")
        try:
            token = get_cloud_token()
            headers = {"Authorization": f"Bearer {token}"} if token else {}
            resp = httpx.get(
                f"{state.cloud_api_url}/api/public/task-attachments/{attachment_id}",
                headers=headers,
                timeout=30.0,
                follow_redirects=True,
            )
            if resp.status_code >= 400:
                raise HTTPException(status_code=resp.status_code, detail="Attachment not found")
            content_type = resp.headers.get("content-type", "application/octet-stream")
            content_disposition = resp.headers.get("content-disposition", "")
            response_headers = {}
            if content_disposition:
                response_headers["Content-Disposition"] = content_disposition
            return Response(content=resp.content, media_type=content_type, headers=response_headers)
        except httpx.HTTPError as exc:
            raise HTTPException(status_code=502, detail=f"Cloud attachment unavailable: {exc}") from exc

    @app.get("/api/v1/system/health", response_model=HealthResponse)
    def get_health() -> HealthResponse:
        return build_health()

    def _sync_org_ai_config_from_cloud() -> None:
        """Pull org-level AI config from cloud and apply locally (background, non-blocking)."""
        try:
            secret_payload = cloud_request("GET", "/api/v1/settings/org-ai-config/secret")
            if not isinstance(secret_payload, dict):
                return
            provider = str(secret_payload.get("aiProvider", "")).strip()
            model = str(secret_payload.get("aiModel", "")).strip()
            api_key = str(secret_payload.get("apiKey", "")).strip()
            if not provider or provider == "mock":
                return
            state.ai.configure(provider, model, api_key, False)
        except Exception:
            pass  # 云端不可用时保留本地配置

    @app.get("/api/v1/auth/me", response_model=AuthStateResponse)
    def auth_me() -> AuthStateResponse:
        token = get_cloud_token()
        refresh_token = get_cloud_refresh_token()
        if not token and not refresh_token:
            return AuthStateResponse(authenticated=False)
        cached_user = get_cached_session_user()
        try:
            user = require_session_user()
        except HTTPException as exc:
            if exc.status_code in {401, 403}:
                clear_cloud_session()
                return AuthStateResponse(authenticated=False, message=str(exc.detail))
            if exc.status_code in {502, 503, 504} and cached_user is not None:
                return AuthStateResponse(
                    authenticated=True,
                    user=cached_user,
                    message="云端暂时不可用，已保留当前设备上的登录状态。",
                )
            raise
        import threading
        threading.Thread(target=_sync_org_ai_config_from_cloud, daemon=True).start()
        return AuthStateResponse(authenticated=True, user=user)

    @app.get("/api/v1/auth/department-options", response_model=list[DepartmentOptionRecord])
    def auth_department_options() -> list[DepartmentOptionRecord]:
        return [
            DepartmentOptionRecord(id=item.id, name=item.name, color=item.color)
            for item in list_department_catalog()
        ]

    @app.post("/api/v1/auth/register", response_model=AuthStateResponse)
    def auth_register(payload: AuthRegisterPayload) -> AuthStateResponse:
        response = cloud_request(
            "POST",
            "/api/v1/auth/register",
            json_body=payload.model_dump(),
            allow_unauthenticated=True,
        )
        message = response.get("message") if isinstance(response, dict) else "你的账号已提交，正在等待管理员审核。"
        return AuthStateResponse(authenticated=False, message=str(message))

    @app.post("/api/v1/auth/login", response_model=AuthStateResponse)
    def auth_login(payload: AuthLoginPayload) -> AuthStateResponse:
        response = cloud_request(
            "POST",
            "/api/v1/auth/login",
            json_body=payload.model_dump(),
            allow_unauthenticated=True,
        )
        if not isinstance(response, dict):
            raise HTTPException(status_code=502, detail="Invalid auth payload")
        token = str(response.get("accessToken", ""))
        refresh_token = str(response.get("refreshToken", ""))
        user_payload = response.get("user")
        if not token or not refresh_token or not isinstance(user_payload, dict):
            raise HTTPException(status_code=502, detail="Cloud auth payload missing session data")
        user = SessionUserRecord(**user_payload)
        set_cloud_session(token, user, persist=payload.rememberMe)
        set_cloud_refresh_token(refresh_token, persist=payload.rememberMe)
        log_activity("auth.login", "session", user.id, {"email": user.email})
        return AuthStateResponse(authenticated=True, user=user)

    @app.post("/api/v1/auth/logout", response_model=AuthStateResponse)
    def auth_logout() -> AuthStateResponse:
        if get_cloud_token():
            try:
                cloud_request("POST", "/api/v1/auth/logout")
            except HTTPException:
                pass
        clear_cloud_session()
        log_activity("auth.logout", "session", "current", {})
        return AuthStateResponse(authenticated=False)

    def process_pending_consultation_knowledge_requests_impl() -> ConsultationKnowledgeProcessSummaryResponse:
        all_requests = list_cloud_consultation_knowledge_requests()
        now_ts = time.time()
        eligible = [
            item
            for item in all_requests
            if _should_retry_consultation_knowledge_request(item, now_ts=now_ts)
        ]
        eligible.sort(
            key=lambda item: (
                _parse_iso_moment(item.createdAt).timestamp() if _parse_iso_moment(item.createdAt) else now_ts,
                item.id,
            )
        )
        summary = ConsultationKnowledgeProcessSummaryResponse(
            totalPending=len(eligible),
            updatedAt=now_iso(),
        )
        if not eligible:
            return summary

        processed_items: list[ConsultationKnowledgeRequestRecord] = []
        for item in eligible:
            summary.processedCount += 1
            try:
                update_cloud_consultation_knowledge_request_status(item.id, status="processing")
                generated = sink_consultation_knowledge_request(item)
                completed = update_cloud_consultation_knowledge_request_status(
                    item.id,
                    status="completed",
                    local_document_id=generated.documentId,
                    local_document_path=generated.path,
                )
                summary.completedCount += 1
                processed_items.append(completed)
            except Exception as exc:
                detail = exc.detail if isinstance(exc, HTTPException) else exc
                error_message = str(detail or "本地综合库写入失败")
                try:
                    failed = update_cloud_consultation_knowledge_request_status(
                        item.id,
                        status="failed",
                        error_message=error_message,
                    )
                except Exception:
                    failed = ConsultationKnowledgeRequestRecord(
                        **{
                            **item.model_dump(),
                            "status": "failed",
                            "errorMessage": error_message,
                            "completedAt": None,
                            "updatedAt": now_iso(),
                        }
                    )
                summary.failedCount += 1
                processed_items.append(failed)
        summary.updatedAt = now_iso()
        summary.items = processed_items
        return summary

    @app.get("/api/v1/consultation/knowledge-requests", response_model=list[ConsultationKnowledgeRequestRecord])
    def list_consultation_knowledge_requests(
        status_filter: Literal["pending", "processing", "completed", "failed"] | None = Query(default=None, alias="status"),
    ) -> list[ConsultationKnowledgeRequestRecord]:
        require_session_user()
        return list_cloud_consultation_knowledge_requests(status_filter)

    @app.post(
        "/api/v1/consultation/knowledge-requests/process-pending",
        response_model=ConsultationKnowledgeProcessSummaryResponse,
    )
    def process_pending_consultation_knowledge_requests() -> ConsultationKnowledgeProcessSummaryResponse:
        require_session_user()
        if state.consultation_knowledge_sync_running:
            return ConsultationKnowledgeProcessSummaryResponse(updatedAt=now_iso())
        state.consultation_knowledge_sync_running = True
        try:
            return process_pending_consultation_knowledge_requests_impl()
        finally:
            state.consultation_knowledge_sync_running = False

    @app.get("/api/v1/admin/employees", response_model=list[EmployeeRecord])
    def list_employee_reviews() -> list[EmployeeRecord]:
        payload = cloud_request("GET", "/api/v1/admin/employees")
        if not isinstance(payload, list):
            raise HTTPException(status_code=502, detail="Invalid employee payload")
        return [EmployeeRecord(**item) for item in payload if isinstance(item, dict)]

    @app.get("/api/v1/settings/org-model/profile", response_model=OrgModelProfileRecord)
    def read_org_model_profile() -> OrgModelProfileRecord:
        payload = cloud_request("GET", "/api/v1/settings/org-model/profile")
        if not isinstance(payload, dict):
            raise HTTPException(status_code=502, detail="Invalid org model payload")
        return OrgModelProfileRecord(**payload)

    @app.post("/api/v1/settings/org-model/profile", response_model=OrgModelProfileRecord)
    def update_org_model_profile(payload: OrgModelProfileRecord) -> OrgModelProfileRecord:
        response = cloud_request("POST", "/api/v1/settings/org-model/profile", json_body=payload.model_dump())
        if not isinstance(response, dict):
            raise HTTPException(status_code=502, detail="Invalid org model payload")
        return OrgModelProfileRecord(**response)

    @app.post("/api/v1/settings/org-model/backfill-task-links", response_model=TaskOrgBackfillResultRecord)
    def backfill_org_task_links() -> TaskOrgBackfillResultRecord:
        response = cloud_request("POST", "/api/v1/settings/org-model/backfill-task-links")
        if not isinstance(response, dict):
            raise HTTPException(status_code=502, detail="Invalid task org backfill payload")
        return TaskOrgBackfillResultRecord(**response)

    @app.get("/api/v1/event-lines", response_model=list[EventLineRecord])
    def list_event_lines() -> list[EventLineRecord]:
        if get_cloud_token():
            response = cloud_request("GET", "/api/v1/event-lines")
            if not isinstance(response, list):
                raise HTTPException(status_code=502, detail="Invalid event line payload")
            return [build_cloud_event_line(item) for item in response if isinstance(item, dict)]
        rows = state.db.fetchall(
            """
            SELECT *
            FROM event_lines
            ORDER BY CASE status WHEN 'active' THEN 0 WHEN 'blocked' THEN 1 WHEN 'paused' THEN 2 WHEN 'done' THEN 3 ELSE 4 END,
                     updated_at DESC
            """,
        )
        return [build_event_line(row) for row in rows]

    @app.post("/api/v1/event-lines", response_model=EventLineRecord)
    def create_event_line(payload: EventLineCreatePayload) -> EventLineRecord:
        if get_cloud_token():
            response = cloud_request("POST", "/api/v1/event-lines", json_body=payload.model_dump())
            if not isinstance(response, dict):
                raise HTTPException(status_code=502, detail="Invalid event line payload")
            return build_cloud_event_line(response)
        owner_row = state.db.fetchone(
            "SELECT full_name, email FROM employee_accounts WHERE id = ?",
            (payload.ownerId,),
        ) if payload.ownerId else None
        client_row = state.db.fetchone("SELECT name FROM clients WHERE id = ?", (payload.primaryClientId,)) if payload.primaryClientId else None
        department_row = state.db.fetchone("SELECT name FROM org_departments WHERE id = ?", (payload.primaryDepartmentId,)) if payload.primaryDepartmentId else None
        owner_name = (
            str(owner_row["full_name"])
            if owner_row and owner_row["full_name"]
            else str(owner_row["email"])
            if owner_row and owner_row["email"]
            else current_operator_name()
        )
        timestamp = now_iso()
        event_line_id = new_id("eline")
        participant_ids = list(dict.fromkeys([item for item in (payload.participantIds or []) if item]))
        state.db.execute(
            """
            INSERT INTO event_lines(
                id, name, kind, status, business_category, stage, summary, intent, current_blocker, recent_decision, next_step, evidence_count, owner_id, owner_name,
                primary_client_id, primary_client_name, primary_department_id, primary_department_name,
                participant_ids_json, created_at, updated_at
            ) VALUES(?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                event_line_id,
                payload.name.strip(),
                payload.kind,
                payload.status,
                payload.businessCategory,
                payload.stage,
                payload.summary,
                payload.intent,
                payload.currentBlocker,
                payload.recentDecision,
                payload.nextStep,
                int(payload.evidenceCount or 0),
                payload.ownerId,
                owner_name,
                payload.primaryClientId,
                str(client_row["name"]) if client_row else None,
                payload.primaryDepartmentId,
                str(department_row["name"]) if department_row else None,
                to_json(participant_ids),
                timestamp,
                timestamp,
            ),
        )
        state.db.execute(
            """
            INSERT INTO event_line_activities(
                id, event_line_id, source_type, source_id, happened_at, actor_id, actor_name, title, summary, metadata_json, created_at
            ) VALUES(?, ?, 'manual_note', ?, ?, NULL, ?, ?, ?, ?, ?)
            """,
            (
                new_id("ela"),
                event_line_id,
                event_line_id,
                timestamp,
                current_operator_name(),
                "创建事件线",
                f"创建事件线：{payload.name.strip()}",
                to_json({"kind": payload.kind, "status": payload.status}),
                timestamp,
            ),
        )
        row = state.db.fetchone("SELECT * FROM event_lines WHERE id = ?", (event_line_id,))
        if not row:
            raise HTTPException(status_code=500, detail="Event line creation failed")
        refresh_event_line_memory_snapshot(state.db, event_line_id)
        if payload.primaryClientId:
            get_client_memory_status(state.db, payload.primaryClientId)
        return build_event_line(row)

    @app.get("/api/v1/event-lines/{event_line_id}", response_model=EventLineDetailRecord)
    def get_event_line(event_line_id: str) -> EventLineDetailRecord:
        if get_cloud_token():
            response = cloud_request("GET", f"/api/v1/event-lines/{event_line_id}")
            if not isinstance(response, dict):
                raise HTTPException(status_code=502, detail="Invalid event line detail payload")
            return build_cloud_event_line_detail(response)
        row = state.db.fetchone("SELECT * FROM event_lines WHERE id = ?", (event_line_id,))
        if not row:
            raise HTTPException(status_code=404, detail="Event line not found")
        return build_event_line_detail(row)

    @app.post("/api/v1/event-lines/{event_line_id}/clarification-draft", response_model=EventLineClarificationDraftRecord)
    def generate_event_line_clarification_draft(
        event_line_id: str,
        payload: EventLineClarificationDraftPayload,
    ) -> EventLineClarificationDraftRecord:
        detail = get_event_line(event_line_id)
        conversation_text = payload.conversationText.strip()
        if len(conversation_text) < 8:
            raise HTTPException(status_code=400, detail="请先粘贴至少一小段聊天记录，再让 AI 整理。")
        activity_lines = [
            f"{item.happenedAt[:10]} {item.title}：{item.summary}"
            for item in detail.activities[:6]
            if item.title or item.summary
        ]
        task_lines = [
            f"任务《{task.title}》：{task.desc or task.ddl or task.status}"
            for task in detail.tasks[:5]
            if task.title
        ]
        draft = state.ai.generate_event_line_clarification_draft(
            event_line_name=detail.eventLine.name,
            conversation_text=conversation_text,
            current_summary=detail.eventLine.summary or "",
            current_stage=detail.eventLine.stage or "",
            current_intent=detail.eventLine.intent or "",
            current_blocker=detail.eventLine.currentBlocker or "",
            current_next_step=detail.eventLine.nextStep or "",
            current_recent_decision=detail.eventLine.recentDecision or "",
            recent_activity_lines=[*activity_lines, *task_lines],
        )
        return EventLineClarificationDraftRecord(**draft)

    @app.get("/api/v1/event-lines/{event_line_id}/memory", response_model=EventLineMemoryResponse)
    def get_event_line_memory(event_line_id: str) -> EventLineMemoryResponse:
        row = state.db.fetchone("SELECT id FROM event_lines WHERE id = ?", (event_line_id,))
        if not row:
            raise HTTPException(status_code=404, detail="Event line not found")
        return get_event_line_memory_response(state.db, event_line_id)

    @app.get("/api/v1/event-lines/{event_line_id}/context-bundle", response_model=EventLineContextBundleRecord)
    def get_event_line_context_bundle(event_line_id: str) -> EventLineContextBundleRecord:
        bundle = _event_line_context_bundle(event_line_id)
        if bundle is None:
            raise HTTPException(status_code=404, detail="Event line context bundle not found")
        return bundle

    @app.get("/api/v1/tasks/{task_id}/context-preview", response_model=TaskContextPreviewRecord)
    def get_task_context_preview(task_id: str) -> TaskContextPreviewRecord:
        if get_cloud_token():
            task = fetch_cloud_task_by_id(task_id)
        else:
            task = next(iter(fetch_tasks("t.id = ?", (task_id,))), None)
            if task is None:
                raise HTTPException(status_code=404, detail="Task not found")
        return _build_task_context_preview(task)

    @app.get("/api/v1/event-lines/{event_line_id}/report-snapshot")
    def get_event_line_report_snapshot(event_line_id: str) -> dict:
        if not get_cloud_token():
            raise HTTPException(status_code=400, detail="需要登录云端才能获取事件线汇报快照")
        payload = cloud_request("GET", f"/api/v1/event-lines/{event_line_id}/report-snapshot")
        if not isinstance(payload, dict):
            raise HTTPException(status_code=502, detail="Invalid report snapshot payload")
        return payload

    @app.post("/api/v1/event-lines/{event_line_id}/export-word")
    def export_event_line_word(event_line_id: str, draft: dict) -> Response:
        doc = WordDocument()

        event_line_name = str(draft.get("eventLineName", "事件线汇报"))
        summary = str(draft.get("summary", ""))
        participants = draft.get("participantNames", [])
        snapshot_at = str(draft.get("snapshotAt", now_iso()))

        doc.add_heading(event_line_name, level=1)
        if summary:
            doc.add_paragraph(summary)
        meta_parts = [f"导出时间：{snapshot_at[:16].replace('T', ' ')}"]
        if participants:
            meta_parts.append(f"参与者：{', '.join(str(name) for name in participants)}")
        doc.add_paragraph(" | ".join(meta_parts)).style = doc.styles["Subtitle"]

        doc.add_heading("事件时间线", level=2)
        activities = draft.get("activities", [])
        visible_activities = [a for a in activities if not a.get("hidden")]
        if not visible_activities:
            doc.add_paragraph("（无活动记录）")
        for activity in visible_activities:
            title = str(activity.get("editedTitle") or activity.get("title", ""))
            summary_text = str(activity.get("editedSummary") or activity.get("summary", ""))
            happened_at = str(activity.get("happenedAt", ""))[:16].replace("T", " ")
            actor = str(activity.get("actorName", ""))
            source_type = str(activity.get("sourceType", ""))
            source_labels = {
                "task_activity": "任务",
                "meeting": "会议",
                "support_request": "支持请求",
                "review": "复核",
                "attachment": "附件",
                "manual_note": "备注",
            }
            label = source_labels.get(source_type, source_type)
            heading = f"[{happened_at}] [{label}] {title}"
            if actor:
                heading += f"（{actor}）"
            p = doc.add_paragraph()
            run = p.add_run(heading)
            run.bold = True
            if summary_text:
                doc.add_paragraph(summary_text)

        attachments = draft.get("attachments", [])
        if attachments:
            doc.add_heading("附件清单", level=2)
            table = doc.add_table(rows=1, cols=4)
            table.style = "Table Grid"
            header_cells = table.rows[0].cells
            header_cells[0].text = "文件名"
            header_cells[1].text = "类型"
            header_cells[2].text = "大小"
            header_cells[3].text = "上传时间"
            for att in attachments:
                row = table.add_row()
                row.cells[0].text = str(att.get("title", ""))
                row.cells[1].text = str(att.get("kind", ""))
                size_bytes = int(att.get("sizeBytes", 0))
                if size_bytes < 1024:
                    size_label = f"{size_bytes} B"
                elif size_bytes < 1024 * 1024:
                    size_label = f"{size_bytes // 1024} KB"
                else:
                    size_label = f"{size_bytes / (1024 * 1024):.1f} MB"
                row.cells[2].text = size_label
                row.cells[3].text = str(att.get("createdAt", ""))[:16].replace("T", " ")

        from io import BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        safe_name = safe_filename(f"{event_line_name[:30]}_汇报.docx")
        return Response(
            content=buffer.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{safe_name}"'},
        )

    @app.patch("/api/v1/event-lines/{event_line_id}", response_model=EventLineRecord)
    def update_event_line(event_line_id: str, payload: EventLineUpdatePayload) -> EventLineRecord:
        if get_cloud_token():
            response = cloud_request("PATCH", f"/api/v1/event-lines/{event_line_id}", json_body=payload.model_dump(exclude_unset=True))
            if not isinstance(response, dict):
                raise HTTPException(status_code=502, detail="Invalid event line payload")
            return build_cloud_event_line(response)
        row = state.db.fetchone("SELECT * FROM event_lines WHERE id = ?", (event_line_id,))
        if not row:
            raise HTTPException(status_code=404, detail="Event line not found")
        next_owner_id = payload.ownerId if "ownerId" in payload.model_fields_set else (str(row["owner_id"]) if row["owner_id"] else None)
        owner_row = state.db.fetchone(
            "SELECT full_name, email FROM employee_accounts WHERE id = ?",
            (next_owner_id,),
        ) if next_owner_id else None
        next_client_id = payload.primaryClientId if "primaryClientId" in payload.model_fields_set else (str(row["primary_client_id"]) if row["primary_client_id"] else None)
        client_row = state.db.fetchone("SELECT name FROM clients WHERE id = ?", (next_client_id,)) if next_client_id else None
        next_department_id = payload.primaryDepartmentId if "primaryDepartmentId" in payload.model_fields_set else (str(row["primary_department_id"]) if row["primary_department_id"] else None)
        department_row = state.db.fetchone("SELECT name FROM org_departments WHERE id = ?", (next_department_id,)) if next_department_id else None
        merged = {
            "name": payload.name.strip() if payload.name is not None else str(row["name"]),
            "kind": payload.kind or str(row["kind"]),
            "status": payload.status or str(row["status"]),
            "business_category": payload.businessCategory if "businessCategory" in payload.model_fields_set else row["business_category"],
            "stage": payload.stage if "stage" in payload.model_fields_set else row["stage"],
            "summary": payload.summary if "summary" in payload.model_fields_set else row["summary"],
            "intent": payload.intent if "intent" in payload.model_fields_set else row["intent"],
            "current_blocker": payload.currentBlocker if "currentBlocker" in payload.model_fields_set else row["current_blocker"],
            "recent_decision": payload.recentDecision if "recentDecision" in payload.model_fields_set else row["recent_decision"],
            "next_step": payload.nextStep if "nextStep" in payload.model_fields_set else row["next_step"],
            "evidence_count": payload.evidenceCount if "evidenceCount" in payload.model_fields_set and payload.evidenceCount is not None else int(row["evidence_count"] or 0),
            "owner_id": next_owner_id,
            "owner_name": (
                str(owner_row["full_name"])
                if owner_row and owner_row["full_name"]
                else str(owner_row["email"])
                if owner_row and owner_row["email"]
                else (str(row["owner_name"]) if row["owner_name"] else current_operator_name())
            ),
            "primary_client_id": next_client_id,
            "primary_client_name": str(client_row["name"]) if client_row else (str(row["primary_client_name"]) if row["primary_client_name"] else None),
            "primary_department_id": next_department_id,
            "primary_department_name": str(department_row["name"]) if department_row else (str(row["primary_department_name"]) if row["primary_department_name"] else None),
            "participant_ids_json": to_json(payload.participantIds if payload.participantIds is not None else from_json(row["participant_ids_json"], [])),
            "updated_at": now_iso(),
        }
        state.db.execute(
            """
            UPDATE event_lines
            SET name = ?, kind = ?, status = ?, business_category = ?, stage = ?, summary = ?, intent = ?, current_blocker = ?, recent_decision = ?, next_step = ?, evidence_count = ?, owner_id = ?, owner_name = ?,
                primary_client_id = ?, primary_client_name = ?, primary_department_id = ?, primary_department_name = ?, participant_ids_json = ?, updated_at = ?
            WHERE id = ?
            """,
            (
                merged["name"],
                merged["kind"],
                merged["status"],
                merged["business_category"],
                merged["stage"],
                merged["summary"],
                merged["intent"],
                merged["current_blocker"],
                merged["recent_decision"],
                merged["next_step"],
                merged["evidence_count"],
                merged["owner_id"],
                merged["owner_name"],
                merged["primary_client_id"],
                merged["primary_client_name"],
                merged["primary_department_id"],
                merged["primary_department_name"],
                merged["participant_ids_json"],
                merged["updated_at"],
                event_line_id,
            ),
        )
        state.db.execute(
            """
            INSERT INTO event_line_activities(
                id, event_line_id, source_type, source_id, happened_at, actor_id, actor_name, title, summary, metadata_json, created_at
            ) VALUES(?, ?, 'manual_note', ?, ?, NULL, ?, ?, ?, ?, ?)
            """,
            (
                new_id("ela"),
                event_line_id,
                event_line_id,
                merged["updated_at"],
                current_operator_name(),
                "更新事件线",
                f"更新事件线：{merged['name']}",
                to_json(
                    {
                        "status": merged["status"],
                        "stage": merged["stage"],
                        "currentBlocker": merged["current_blocker"],
                        "recentDecision": merged["recent_decision"],
                    }
                ),
                merged["updated_at"],
            ),
        )
        updated_row = state.db.fetchone("SELECT * FROM event_lines WHERE id = ?", (event_line_id,))
        if not updated_row:
            raise HTTPException(status_code=500, detail="Event line update failed")
        refresh_event_line_memory_snapshot(state.db, event_line_id)
        if merged["primary_client_id"]:
            get_client_memory_status(state.db, str(merged["primary_client_id"]))
        return build_event_line(updated_row)

    @app.post("/api/v1/event-lines/{event_line_id}/notes")
    def add_event_line_note(event_line_id: str, payload: dict = Body(...)) -> dict:
        """Add a manual observation/note to an event line."""
        note_text = str(payload.get("text", "")).strip()
        if not note_text:
            raise HTTPException(status_code=400, detail="Note text is required")
        el_row = state.db.fetchone("SELECT * FROM event_lines WHERE id = ?", (event_line_id,))
        if not el_row:
            raise HTTPException(status_code=404, detail="Event line not found")
        note_ts = now_iso()
        note_id = new_id("ela")
        state.db.execute(
            """
            INSERT INTO event_line_activities(
                id, event_line_id, source_type, source_id, happened_at, actor_id, actor_name, title, summary, metadata_json, created_at
            ) VALUES(?, ?, 'manual_note', ?, ?, NULL, ?, ?, ?, ?, ?)
            """,
            (
                note_id,
                event_line_id,
                note_id,
                note_ts,
                current_operator_name(),
                "手动备注",
                note_text[:500],
                to_json({"kind": "user_note"}),
                note_ts,
            ),
        )
        log_activity("event_line.note", "event_line", event_line_id, {"noteLength": len(note_text)})
        return {"id": note_id, "eventLineId": event_line_id, "text": note_text[:500], "createdAt": note_ts}

    @app.get("/api/v1/task-views", response_model=TaskViewsResponse)
    def list_task_views() -> TaskViewsResponse:
        views = _ensure_builtin_task_views()
        rows = state.db.fetchall(
            """
            SELECT *
            FROM task_views
            ORDER BY built_in DESC,
                     CASE kind WHEN 'custom' THEN 1 ELSE 0 END,
                     updated_at DESC,
                     name ASC
            """
        )
        views = [_task_view_record_from_row(row) for row in rows]
        return TaskViewsResponse(
            views=views,
            presets=_task_view_presets(views),
        )

    @app.post("/api/v1/task-views", response_model=TaskViewDefinitionRecord)
    def create_task_view(payload: TaskViewMutationPayload) -> TaskViewDefinitionRecord:
        _ensure_builtin_task_views()
        if payload.kind != "custom":
            raise HTTPException(status_code=400, detail="Only custom task views can be created")
        name = payload.name.strip()
        if not name:
            raise HTTPException(status_code=400, detail="Task view name is required")
        timestamp = now_iso()
        view_id = new_id("tview")
        state.db.execute(
            """
            INSERT INTO task_views(
                id, name, kind, description, calendar_scope, shareability, sort_by, sort_direction,
                visible_fields_json, filter_set_json, built_in, created_at, updated_at
            ) VALUES(?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 0, ?, ?)
            """,
            (
                view_id,
                name,
                payload.kind,
                payload.description,
                payload.calendarScope,
                payload.shareability,
                payload.sortBy,
                payload.sortDirection,
                to_json(payload.visibleFields),
                to_json(payload.filterSet.model_dump()),
                timestamp,
                timestamp,
            ),
        )
        row = state.db.fetchone("SELECT * FROM task_views WHERE id = ?", (view_id,))
        if not row:
            raise HTTPException(status_code=500, detail="Task view creation failed")
        return _task_view_record_from_row(row)

    @app.patch("/api/v1/task-views/{view_id}", response_model=TaskViewDefinitionRecord)
    def update_task_view(view_id: str, payload: TaskViewMutationPayload) -> TaskViewDefinitionRecord:
        _ensure_builtin_task_views()
        row = state.db.fetchone("SELECT * FROM task_views WHERE id = ?", (view_id,))
        if not row:
            raise HTTPException(status_code=404, detail="Task view not found")
        if str(row["kind"]) != "custom":
            raise HTTPException(status_code=400, detail="Built-in task views cannot be edited")
        timestamp = now_iso()
        next_name = payload.name.strip() if payload.name is not None else str(row["name"])
        if not next_name:
            raise HTTPException(status_code=400, detail="Task view name is required")
        next_filter_set = payload.filterSet.model_dump() if payload.filterSet is not None else from_json(row["filter_set_json"], {})
        next_visible_fields = payload.visibleFields if payload.visibleFields is not None else from_json(row["visible_fields_json"], [])
        state.db.execute(
            """
            UPDATE task_views
            SET name = ?, description = ?, filter_set_json = ?, sort_by = ?, sort_direction = ?, visible_fields_json = ?,
                calendar_scope = ?, shareability = ?, updated_at = ?
            WHERE id = ?
            """,
            (
                next_name,
                payload.description if payload.description is not None else str(row["description"] or ""),
                to_json(next_filter_set),
                payload.sortBy or str(row["sort_by"]),
                payload.sortDirection or str(row["sort_direction"]),
                to_json(next_visible_fields),
                payload.calendarScope or str(row["calendar_scope"]),
                payload.shareability or str(row["shareability"]),
                timestamp,
                view_id,
            ),
        )
        updated_row = state.db.fetchone("SELECT * FROM task_views WHERE id = ?", (view_id,))
        if not updated_row:
            raise HTTPException(status_code=500, detail="Task view update failed")
        return _task_view_record_from_row(updated_row)

    @app.get("/api/v1/reviews/dashboard/drill-target", response_model=ReviewDashboardDrillTargetResponse)
    def get_review_dashboard_drill_target(
        targetType: str,
        targetId: str,
        targetLabel: str | None = None,
        targetFilters: str | None = None,
    ) -> ReviewDashboardDrillTargetResponse:
        parsed_filters: dict[str, object] = {}
        if targetFilters:
            try:
                candidate = json.loads(targetFilters)
            except json.JSONDecodeError as exc:
                raise HTTPException(status_code=400, detail="Invalid targetFilters payload") from exc
            if not isinstance(candidate, dict):
                raise HTTPException(status_code=400, detail="targetFilters must be an object")
            parsed_filters = candidate
        target = ReviewDashboardCardTargetRecord(
            targetType=targetType,
            targetId=targetId,
            targetLabel=targetLabel,
            targetFilters=parsed_filters,
        )
        if target.targetType == "event_line":
            return _drill_target_response_for_event_line(target)
        if target.targetType == "task_view":
            return _drill_target_response_for_task_view(target)
        if target.targetType == "meeting":
            return _drill_target_response_for_meeting(target)
        if target.targetType == "support_request":
            return _drill_target_response_for_support_request(target)
        if target.targetType == "attachment_group":
            return _drill_target_response_for_attachment_group(target)
        raise HTTPException(status_code=400, detail=f"Unsupported dashboard drill target: {target.targetType}")

    @app.get("/api/v1/tasks/{task_id}/plan-link", response_model=TaskPlanLinkRecord | None)
    def read_task_plan_link(task_id: str) -> TaskPlanLinkRecord | None:
        response = cloud_request("GET", f"/api/v1/tasks/{task_id}/plan-link")
        if response is None:
            return None
        if not isinstance(response, dict):
            raise HTTPException(status_code=502, detail="Invalid task plan link payload")
        return TaskPlanLinkRecord(**response)

    @app.post("/api/v1/tasks/{task_id}/plan-link/recompute", response_model=TaskPlanLinkRecord | None)
    def recompute_task_plan_link(task_id: str) -> TaskPlanLinkRecord | None:
        response = cloud_request("POST", f"/api/v1/tasks/{task_id}/plan-link/recompute")
        if response is None:
            return None
        if not isinstance(response, dict):
            raise HTTPException(status_code=502, detail="Invalid task plan link payload")
        return TaskPlanLinkRecord(**response)

    @app.patch("/api/v1/tasks/{task_id}/plan-link", response_model=TaskPlanLinkRecord | None)
    def patch_task_plan_link(task_id: str, payload: TaskPlanLinkUpsertPayload) -> TaskPlanLinkRecord | None:
        response = cloud_request("PATCH", f"/api/v1/tasks/{task_id}/plan-link", json_body=payload.model_dump())
        if response is None:
            return None
        if not isinstance(response, dict):
            raise HTTPException(status_code=502, detail="Invalid task plan link payload")
        return TaskPlanLinkRecord(**response)

    @app.get("/api/v1/support-requests", response_model=list[SupportRequestRecord])
    def list_support_requests(status: str | None = Query(default=None), taskId: str | None = Query(default=None)) -> list[SupportRequestRecord]:
        query = []
        if status:
            query.append(f"status={quote(status)}")
        if taskId:
            query.append(f"taskId={quote(taskId)}")
        suffix = f"?{'&'.join(query)}" if query else ""
        response = cloud_request("GET", f"/api/v1/support-requests{suffix}")
        if not isinstance(response, list):
            raise HTTPException(status_code=502, detail="Invalid support request payload")
        return [SupportRequestRecord(**item) for item in response if isinstance(item, dict)]

    @app.post("/api/v1/support-requests", response_model=SupportRequestRecord)
    def create_support_request(payload: SupportRequestCreatePayload) -> SupportRequestRecord:
        response = cloud_request("POST", "/api/v1/support-requests", json_body=payload.model_dump())
        if not isinstance(response, dict):
            raise HTTPException(status_code=502, detail="Invalid support request payload")
        return SupportRequestRecord(**response)

    @app.post("/api/v1/support-requests/{request_id}/resolve", response_model=SupportRequestRecord)
    def resolve_support_request(request_id: str, payload: SupportRequestResolvePayload) -> SupportRequestRecord:
        response = cloud_request("POST", f"/api/v1/support-requests/{request_id}/resolve", json_body=payload.model_dump())
        if not isinstance(response, dict):
            raise HTTPException(status_code=502, detail="Invalid support request payload")
        record = SupportRequestRecord(**response)
        # Write support_request activity to the related event line (if task has one)
        if record.taskId:
            task_row = state.db.fetchone("SELECT event_line_id, title FROM tasks WHERE id = ?", (record.taskId,))
            if task_row and task_row["event_line_id"]:
                sr_ts = now_iso()
                state.db.execute(
                    """
                    INSERT INTO event_line_activities(
                        id, event_line_id, source_type, source_id, happened_at, actor_id, actor_name, title, summary, metadata_json, created_at
                    ) VALUES(?, ?, 'support_request', ?, ?, NULL, ?, ?, ?, ?, ?)
                    """,
                    (
                        new_id("ela"),
                        str(task_row["event_line_id"]),
                        request_id,
                        sr_ts,
                        current_operator_name(),
                        f"支持请求已处理：{record.requestType}",
                        f"针对任务「{task_row['title']}」的{record.requestType}请求已{record.status}。" + (f" 处理说明：{record.resolutionNote[:80]}" if record.resolutionNote else ""),
                        to_json({"taskId": record.taskId, "requestType": record.requestType, "status": record.status}),
                        sr_ts,
                    ),
                )
        return record

    @app.post("/api/v1/admin/employees/{employee_id}/approve", response_model=EmployeeRecord)
    def approve_employee(employee_id: str, payload: EmployeeRolePayload) -> EmployeeRecord:
        response = cloud_request("POST", f"/api/v1/admin/employees/{employee_id}/approve", json_body=payload.model_dump())
        if not isinstance(response, dict):
            raise HTTPException(status_code=502, detail="Invalid employee payload")
        return EmployeeRecord(**response)

    @app.post("/api/v1/admin/employees/{employee_id}/reject", response_model=EmployeeRecord)
    def reject_employee(employee_id: str, payload: EmployeeRejectPayload) -> EmployeeRecord:
        response = cloud_request("POST", f"/api/v1/admin/employees/{employee_id}/reject", json_body=payload.model_dump())
        if not isinstance(response, dict):
            raise HTTPException(status_code=502, detail="Invalid employee payload")
        return EmployeeRecord(**response)

    @app.post("/api/v1/admin/employees/{employee_id}/disable", response_model=EmployeeRecord)
    def disable_employee(employee_id: str) -> EmployeeRecord:
        response = cloud_request("POST", f"/api/v1/admin/employees/{employee_id}/disable")
        if not isinstance(response, dict):
            raise HTTPException(status_code=502, detail="Invalid employee payload")
        return EmployeeRecord(**response)

    @app.patch("/api/v1/admin/employees/{employee_id}/role", response_model=EmployeeRecord)
    def patch_employee_role(employee_id: str, payload: EmployeeRolePayload) -> EmployeeRecord:
        response = cloud_request("PATCH", f"/api/v1/admin/employees/{employee_id}/role", json_body=payload.model_dump())
        if not isinstance(response, dict):
            raise HTTPException(status_code=502, detail="Invalid employee payload")
        return EmployeeRecord(**response)

    @app.patch("/api/v1/admin/employees/{employee_id}/department", response_model=EmployeeRecord)
    def patch_employee_department(employee_id: str, payload: EmployeeDepartmentPayload) -> EmployeeRecord:
        response = cloud_request("PATCH", f"/api/v1/admin/employees/{employee_id}/department", json_body=payload.model_dump())
        if not isinstance(response, dict):
            raise HTTPException(status_code=502, detail="Invalid employee payload")
        return EmployeeRecord(**response)

    @app.get("/api/v1/employees/mention-candidates", response_model=list[MentionCandidateRecord])
    def get_mention_candidates(q: str = Query(default="")) -> list[MentionCandidateRecord]:
        payload = cloud_request("GET", f"/api/v1/employees/mention-candidates?q={q}")
        if not isinstance(payload, list):
            raise HTTPException(status_code=502, detail="Invalid mention payload")
        return [MentionCandidateRecord(**item) for item in payload if isinstance(item, dict)]

    @app.get("/api/v1/settings", response_model=SettingsResponse)
    def get_settings() -> SettingsResponse:
        return build_settings_response()

    @app.get("/api/v1/settings/logs", response_model=list[ActivityLogRecord])
    def get_activity_logs() -> list[ActivityLogRecord]:
        return [
            ActivityLogRecord(
                id=str(row["id"]),
                actorName=str(row["actor_name"]),
                action=str(row["action"]),
                entityType=str(row["entity_type"]),
                entityId=str(row["entity_id"]),
                detail=from_json(row["detail_json"], {}),
                createdAt=str(row["created_at"]),
            )
            for row in state.db.fetchall("SELECT * FROM activity_logs ORDER BY created_at DESC LIMIT 30")
        ]

    @app.get("/api/v1/tasks/agent-worklogs", response_model=AgentWorklogResponse)
    def get_agent_worklogs(month: str | None = Query(default=None)) -> AgentWorklogResponse:
        target_month = month or datetime.now().strftime("%Y-%m")
        if not current_session_is_admin():
            return AgentWorklogResponse(month=target_month)
        try:
            return build_agent_worklog_response(
                db=state.db,
                month_label=target_month,
                thread_sync_path=THREAD_SYNC_DOC_PATH,
            )
        except ValueError as exc:
            raise HTTPException(status_code=400, detail=str(exc)) from exc

    @app.put("/api/v1/tasks/agent-weekly-plans/{week_label}/{agent_key}", response_model=AgentWeeklyPlanRecord)
    def update_agent_weekly_plan(
        week_label: str,
        agent_key: str,
        payload: AgentWeeklyPlanPayload,
    ) -> AgentWeeklyPlanRecord:
        if not current_session_is_admin():
            raise HTTPException(status_code=403, detail="只有机构负责人可以调整机器人部门周计划。")
        if payload.weekLabel != week_label or payload.agentKey != agent_key:
            raise HTTPException(status_code=400, detail="路径参数和计划内容不一致。")
        user = require_session_user()
        upsert_agent_weekly_plan_override(
            db=state.db,
            payload=payload,
            updated_by=user.fullName,
        )
        log_activity(
            "agent.plan.update",
            "agent_weekly_plan",
            f"{week_label}:{agent_key}",
            {
                "weekLabel": week_label,
                "agentKey": agent_key,
                "planItemCount": len(payload.planItems),
                "updatedBy": user.fullName,
            },
        )
        plans = build_agent_weekly_plans(
            db=state.db,
            week_label=week_label,
            thread_sync_path=THREAD_SYNC_DOC_PATH,
        )
        for plan in plans:
            if plan.agentKey == agent_key:
                return plan
        raise HTTPException(status_code=404, detail="未找到对应的机器人部门计划。")

    @app.post("/api/v1/settings", response_model=SettingsResponse)
    def update_settings(payload: AppSettingsPayload) -> SettingsResponse:
        wants_sensitive_update = any(
            value is not None and value != ""
            for value in [payload.aiProvider, payload.aiModel, payload.apiKey]
        ) or payload.clearApiKey
        if wants_sensitive_update and get_system_admin_settings().protectAiAndCloud:
            ensure_admin_for_sensitive_settings()
        elif payload.currentOperatorId:
            ensure_business_settings_editable()
        if payload.currentOperatorId:
            operator = state.db.fetchone("SELECT * FROM operators WHERE id = ?", (payload.currentOperatorId,))
            if not operator:
                raise HTTPException(status_code=404, detail="Operator not found")
            state.db.set_setting("current_operator_id", payload.currentOperatorId)
            state.db.execute("UPDATE operators SET is_current = CASE WHEN id = ? THEN 1 ELSE 0 END", (payload.currentOperatorId,))
        state.ai.configure(payload.aiProvider, payload.aiModel, payload.apiKey, payload.clearApiKey)
        log_activity("settings.update", "settings", "app", payload.model_dump(exclude_none=True))
        return build_settings_response()

    @app.get("/api/v1/settings/tasks", response_model=TaskSettingsRecord)
    def get_task_settings() -> TaskSettingsRecord:
        if get_cloud_token():
            payload = cloud_request("GET", "/api/v1/settings/tasks")
            if not isinstance(payload, dict):
                raise HTTPException(status_code=502, detail="Invalid task settings payload")
            return TaskSettingsRecord(**payload)
        return _get_local_task_settings()

    @app.post("/api/v1/settings/tasks", response_model=TaskSettingsRecord)
    def update_task_settings(payload: TaskSettingsPayload) -> TaskSettingsRecord:
        ensure_business_settings_editable()
        if get_cloud_token():
            response = cloud_request("POST", "/api/v1/settings/tasks", json_body=payload.model_dump(exclude_none=True))
            if not isinstance(response, dict):
                raise HTTPException(status_code=502, detail="Invalid task settings payload")
            return TaskSettingsRecord(**response)
        operator_id = str(current_operator_row()["id"])
        current = _get_local_task_settings(operator_id)
        next_default_list_id = payload.defaultListId if payload.defaultListId is not None else current.defaultListId
        if next_default_list_id:
            list_row = state.db.fetchone("SELECT * FROM task_lists WHERE id = ?", (next_default_list_id,))
            if not list_row or list_row["archived_at"]:
                raise HTTPException(status_code=400, detail="默认清单无效")
        timestamp = now_iso()
        next_record = TaskSettingsRecord(
            defaultListId=next_default_list_id,
            defaultPriority=payload.defaultPriority or current.defaultPriority,
            defaultDueDatePreset=payload.defaultDueDatePreset or current.defaultDueDatePreset,
            defaultViewMode=payload.defaultViewMode or current.defaultViewMode,
            listSortMode=payload.listSortMode or current.listSortMode,
            showCompletedTasks=payload.showCompletedTasks if payload.showCompletedTasks is not None else current.showCompletedTasks,
            defaultReviewScope=payload.defaultReviewScope or current.defaultReviewScope,
            autoAssignSelf=payload.autoAssignSelf if payload.autoAssignSelf is not None else current.autoAssignSelf,
            updatedAt=timestamp,
        )
        state.db.execute(
            """
            INSERT OR REPLACE INTO task_settings(
                operator_id, default_list_id, default_priority, default_due_date_preset,
                default_view_mode, list_sort_mode, show_completed_tasks, default_review_scope,
                auto_assign_self, updated_at
            ) VALUES(?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                operator_id,
                next_record.defaultListId,
                next_record.defaultPriority,
                next_record.defaultDueDatePreset,
                next_record.defaultViewMode,
                next_record.listSortMode,
                1 if next_record.showCompletedTasks else 0,
                next_record.defaultReviewScope,
                1 if next_record.autoAssignSelf else 0,
                next_record.updatedAt,
            ),
        )
        if next_record.defaultListId:
            state.db.execute(
                "UPDATE task_lists SET is_default = CASE WHEN id = ? THEN 1 ELSE 0 END",
                (next_record.defaultListId,),
        )
        log_activity("settings.tasks.update", "settings", operator_id, payload.model_dump(exclude_none=True))
        return _get_local_task_settings(operator_id)

    @app.get("/api/v1/settings/review-governance", response_model=ReviewGovernanceSettingsRecord)
    def read_review_governance_settings() -> ReviewGovernanceSettingsRecord:
        ensure_admin_for_sensitive_settings()
        return _review_governance_with_members()

    @app.post("/api/v1/settings/review-governance", response_model=ReviewGovernanceSettingsRecord)
    def update_review_governance_settings(payload: ReviewGovernanceSettingsPayload) -> ReviewGovernanceSettingsRecord:
        ensure_admin_for_sensitive_settings()
        record = _sanitize_review_governance_settings(payload.departments)
        state.db.set_setting("settings.review_governance", to_json(record.model_dump()))
        log_activity(
            "settings.review_governance.update",
            "settings",
            "review_governance",
            {"departmentCount": len(record.departments)},
        )
        return record

    @app.get("/api/v1/settings/org-dna", response_model=OrganizationDnaResponse)
    def get_organization_dna() -> OrganizationDnaResponse:
        return OrganizationDnaResponse(modules=list_organization_dna_modules())

    @app.get("/api/v1/settings/org-dna/{module_key}", response_model=OrganizationDnaModuleRecord)
    def get_organization_dna_module(module_key: str) -> OrganizationDnaModuleRecord:
        modules = {module.moduleKey: module for module in list_organization_dna_modules()}
        module = modules.get(module_key)
        if not module:
            raise HTTPException(status_code=404, detail="未知的组织 DNA 模块")
        return module

    @app.post("/api/v1/settings/org-dna/{module_key}", response_model=OrganizationDnaModuleRecord)
    def update_organization_dna_module(module_key: str, payload: OrganizationDnaUploadPayload) -> OrganizationDnaModuleRecord:
        ensure_org_dna_editable()
        record = upsert_organization_dna_module(module_key, payload)
        log_activity("settings.org_dna.update", "settings", module_key, {"moduleKey": module_key, "fileName": record.fileName, "contentHash": record.contentHash})
        return record

    @app.get("/api/v1/settings/client-workspace", response_model=ClientWorkspaceSettingsRecord)
    def read_client_workspace_settings() -> ClientWorkspaceSettingsRecord:
        return get_client_workspace_settings()

    @app.post("/api/v1/settings/client-workspace", response_model=ClientWorkspaceSettingsRecord)
    def update_client_workspace_settings(payload: ClientWorkspaceSettingsPayload) -> ClientWorkspaceSettingsRecord:
        ensure_business_settings_editable()
        current = get_client_workspace_settings()
        next_payload = current.model_dump()
        next_payload.update(payload.model_dump(exclude_none=True))
        next_payload["updatedAt"] = now_iso()
        next_record = ClientWorkspaceSettingsRecord(**next_payload)
        if next_record.meetingPublishDefaultListId:
            list_row = state.db.fetchone("SELECT id, archived_at FROM task_lists WHERE id = ?", (next_record.meetingPublishDefaultListId,))
            if not list_row or list_row["archived_at"]:
                raise HTTPException(status_code=400, detail="默认会议任务清单无效")
        _save_json_settings_record("settings.client_workspace", next_record)
        log_activity("settings.client_workspace.update", "settings", "client_workspace", payload.model_dump(exclude_none=True))
        return next_record

    @app.get("/api/v1/settings/topics", response_model=TopicsSettingsRecord)
    def read_topics_settings() -> TopicsSettingsRecord:
        return get_topics_settings()

    @app.post("/api/v1/settings/topics", response_model=TopicsSettingsRecord)
    def update_topics_settings(payload: TopicsSettingsPayload) -> TopicsSettingsRecord:
        ensure_business_settings_editable()
        current = get_topics_settings()
        next_payload = current.model_dump()
        next_payload.update(payload.model_dump(exclude_none=True))
        next_payload["updatedAt"] = now_iso()
        next_record = TopicsSettingsRecord(**next_payload)
        _save_json_settings_record("settings.topics", next_record)
        log_activity("settings.topics.update", "settings", "topics", payload.model_dump(exclude_none=True))
        return next_record

    @app.get("/api/v1/settings/analysis-workbench", response_model=AnalysisWorkbenchSettingsRecord)
    def read_analysis_workbench_settings() -> AnalysisWorkbenchSettingsRecord:
        return get_analysis_workbench_settings()

    @app.post("/api/v1/settings/analysis-workbench", response_model=AnalysisWorkbenchSettingsRecord)
    def update_analysis_workbench_settings(payload: AnalysisWorkbenchSettingsPayload) -> AnalysisWorkbenchSettingsRecord:
        ensure_business_settings_editable()
        current = get_analysis_workbench_settings()
        next_payload = current.model_dump()
        next_payload.update(payload.model_dump(exclude_none=True))
        next_payload["updatedAt"] = now_iso()
        enabled_ids = [str(item) for item in next_payload.get("enabledTemplateIds", [])]
        known_ids = {str(row["id"]) for row in state.db.fetchall("SELECT id FROM analysis_templates")}
        if enabled_ids:
            unknown_ids = [item for item in enabled_ids if item not in known_ids]
            if unknown_ids:
                raise HTTPException(status_code=400, detail=f"未知的分析模板：{'、'.join(unknown_ids)}")
        default_template_id = next_payload.get("defaultTemplateId")
        if default_template_id and default_template_id not in known_ids:
            raise HTTPException(status_code=400, detail="默认分析模板无效")
        next_record = AnalysisWorkbenchSettingsRecord(**next_payload)
        _save_json_settings_record("settings.analysis_workbench", next_record)
        log_activity("settings.analysis_workbench.update", "settings", "analysis_workbench", payload.model_dump(exclude_none=True))
        return next_record

    @app.get("/api/v1/settings/handbook", response_model=HandbookSettingsRecord)
    def read_handbook_settings() -> HandbookSettingsRecord:
        return get_handbook_settings()

    @app.post("/api/v1/settings/handbook", response_model=HandbookSettingsRecord)
    def update_handbook_settings(payload: HandbookSettingsPayload) -> HandbookSettingsRecord:
        ensure_business_settings_editable()
        current = get_handbook_settings()
        next_payload = current.model_dump()
        next_payload.update(payload.model_dump(exclude_none=True))
        next_payload["updatedAt"] = now_iso()
        next_record = HandbookSettingsRecord(**next_payload)
        _save_json_settings_record("settings.handbook", next_record)
        log_activity("settings.handbook.update", "settings", "handbook", payload.model_dump(exclude_none=True))
        return next_record

    @app.get("/api/v1/settings/system-admin", response_model=SystemAdminSettingsRecord)
    def read_system_admin_settings() -> SystemAdminSettingsRecord:
        return get_system_admin_settings()

    @app.post("/api/v1/settings/system-admin", response_model=SystemAdminSettingsRecord)
    def update_system_admin_settings(payload: SystemAdminSettingsPayload) -> SystemAdminSettingsRecord:
        ensure_admin_for_sensitive_settings()
        current = get_system_admin_settings()
        next_payload = current.model_dump()
        updates = payload.model_dump(exclude_none=True)
        if "brandLogoDataUrl" in updates:
            updates["brandLogoDataUrl"] = _normalize_brand_logo_data_url(updates.get("brandLogoDataUrl"))
        next_payload.update(updates)
        next_payload["updatedAt"] = now_iso()
        next_record = SystemAdminSettingsRecord(**next_payload)
        _save_json_settings_record("settings.system_admin", next_record)
        logged_updates = dict(updates)
        if "brandLogoDataUrl" in logged_updates:
            logged_updates["brandLogoDataUrl"] = "[PNG data URL omitted]" if logged_updates["brandLogoDataUrl"] else None
        log_activity("settings.system_admin.update", "settings", "system_admin", logged_updates)
        return next_record

    @app.get("/api/v1/settings/feishu-bot", response_model=FeishuBotSettingsRecord)
    def read_feishu_bot_settings() -> FeishuBotSettingsRecord:
        return get_feishu_bot_settings()

    @app.post("/api/v1/settings/feishu-bot", response_model=FeishuBotSettingsRecord)
    def save_feishu_bot_settings(payload: FeishuBotSettingsPayload) -> FeishuBotSettingsRecord:
        ensure_admin_for_sensitive_settings()
        return update_feishu_bot_settings(payload)

    @app.get("/api/v1/settings/feishu-user-binding", response_model=FeishuUserBindingRecord)
    def read_feishu_user_binding() -> FeishuUserBindingRecord:
        user = require_session_user()
        sync_feishu_user_binding_from_cloud_relay(user.id)
        return get_feishu_user_binding(user.id)

    @app.post("/api/v1/settings/feishu-user-binding/start", response_model=FeishuUserBindingStartResponse)
    def start_feishu_user_binding(request: Request) -> FeishuUserBindingStartResponse:
        user = require_session_user()
        settings = get_feishu_bot_settings()
        if not settings.appId.strip():
            raise HTTPException(status_code=400, detail="请先在系统设置里配置飞书 App ID。")
        try:
            app_secret = state.feishu_secret_store.get_api_key().strip()
        except Exception:
            app_secret = ""
        if not app_secret:
            raise HTTPException(status_code=400, detail="请先在系统设置里保存飞书 App Secret。")
        _clear_feishu_cloud_relay_session(user.id)
        expires_at = (datetime.now() + timedelta(minutes=10)).replace(microsecond=0).isoformat()
        state_token = new_id("fs_state")
        save_feishu_oauth_state(state_token, user.id, expires_at)
        configured_callback_url = settings.userBindingCallbackUrl.strip()
        local_callback_url = f"{str(request.base_url).rstrip('/')}/api/v1/auth/feishu/callback"
        cloud_callback_url = _feishu_cloud_relay_callback_url()
        callback_url = configured_callback_url or local_callback_url
        pending_mode = "local"
        if _is_public_feishu_callback_url(configured_callback_url):
            callback_url = configured_callback_url
        elif _is_public_feishu_callback_url(cloud_callback_url):
            cloud_request(
                "POST",
                "/api/v1/integrations/feishu/user-binding/sessions",
                json_body={"state": state_token, "expiresAt": expires_at},
            )
            callback_url = cloud_callback_url
            pending_mode = "cloud_relay"
        save_feishu_user_binding_pending(
            user.id,
            state_token=state_token,
            expires_at=expires_at,
            callback_url=callback_url,
            mode=pending_mode,
        )
        authorize_url = build_user_authorize_url(app_id=settings.appId.strip(), redirect_uri=callback_url, state=state_token)
        qr_ready = _is_public_feishu_callback_url(callback_url)
        qr_blocked_reason = None if qr_ready else "当前回调地址仍是本机地址或非 HTTPS 地址，手机扫码后无法把授权结果回传到这台工作台。请先配置可公网访问的 HTTPS 回调地址，或直接在当前电脑浏览器完成授权。"
        log_activity("feishu.user_binding.start", "settings", user.id, {"callbackUrl": callback_url})
        return FeishuUserBindingStartResponse(
            authorizeUrl=authorize_url,
            state=state_token,
            expiresAt=expires_at,
            callbackUrl=callback_url,
            qrReady=qr_ready,
            qrBlockedReason=qr_blocked_reason,
        )

    @app.delete("/api/v1/settings/feishu-user-binding", response_model=FeishuUserBindingRecord)
    def delete_feishu_user_binding() -> FeishuUserBindingRecord:
        user = require_session_user()
        _clear_feishu_cloud_relay_session(user.id)
        cleared = clear_feishu_user_binding(user.id)
        log_activity("feishu.user_binding.clear", "settings", user.id, {})
        return cleared

    @app.get("/api/v1/auth/feishu/callback", response_class=HTMLResponse)
    def receive_feishu_auth_callback(
        code: str | None = Query(default=None),
        state_token: str | None = Query(default=None, alias="state"),
    ) -> HTMLResponse:
        if not state_token:
            return _render_feishu_binding_callback_page("飞书绑定失败", "缺少 state，无法确认这次授权属于哪个工作台会话。", success=False)
        oauth_state = pop_feishu_oauth_state(state_token)
        if not oauth_state:
            return _render_feishu_binding_callback_page("飞书绑定失败", "这次授权状态已失效，请回到工作台重新发起绑定。", success=False)
        expires_at = oauth_state.get("expiresAt", "")
        if expires_at:
            try:
                if datetime.fromisoformat(expires_at) <= datetime.now():
                    return _render_feishu_binding_callback_page("飞书绑定失败", "这次授权请求已经过期，请回到工作台重新发起绑定。", success=False)
            except ValueError:
                return _render_feishu_binding_callback_page("飞书绑定失败", "授权状态已损坏，请回到工作台重新发起绑定。", success=False)
        user_id = oauth_state.get("userId", "").strip()
        if not user_id:
            return _render_feishu_binding_callback_page("飞书绑定失败", "这次授权缺少用户信息，请重新发起绑定。", success=False)
        if not code or not code.strip():
            return _render_feishu_binding_callback_page("飞书绑定失败", "飞书没有返回有效授权码，请重新发起绑定。", success=False)

        try:
            binding = _finalize_feishu_user_binding(user_id, code.strip())
            clear_feishu_user_binding_pending(user_id)
            return _render_feishu_binding_callback_page("飞书账号绑定成功", f"已绑定 {binding.name or binding.email or binding.openId}。后续在任务与日历里发起飞书会议时，会优先按当前登录员工的绑定身份发送。", success=True)
        except FeishuApiError as exc:
            _save_feishu_user_binding_error(user_id, str(exc))
            clear_feishu_user_binding_pending(user_id)
            return _render_feishu_binding_callback_page("飞书绑定失败", str(exc), success=False)
        except HTTPException as exc:
            _save_feishu_user_binding_error(user_id, str(exc.detail))
            clear_feishu_user_binding_pending(user_id)
            return _render_feishu_binding_callback_page("飞书绑定失败", str(exc.detail), success=False)

    @app.post("/api/v1/channels/feishu/events")
    async def receive_feishu_events(request: Request) -> dict[str, object]:
        payload = await request.json()
        if not isinstance(payload, dict):
            raise HTTPException(status_code=400, detail="Invalid Feishu event payload")
        return handle_feishu_event(payload)

    @app.post("/api/v1/settings/backup", response_model=BackupResponse)
    def create_backup() -> BackupResponse:
        timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
        backup_path = state.backup_dir / f"yiyu-workbench-{timestamp}.db"
        state.db.backup_to(backup_path)
        state.db.set_setting("last_backup_at", now_iso())
        log_activity("settings.backup", "backup", timestamp, {"path": str(backup_path)})
        return BackupResponse(backupPath=str(backup_path), createdAt=now_iso())

    @app.post("/api/v1/settings/demo-data/load", response_model=DemoDataResponse)
    def load_demo_dataset_endpoint() -> DemoDataResponse:
        response = load_demo_dataset(state)
        log_activity("settings.demo.load", "settings", "demo_data", response.model_dump())
        return response

    @app.post("/api/v1/settings/demo-data/clear", response_model=DemoDataResponse)
    def clear_demo_dataset_endpoint() -> DemoDataResponse:
        response = clear_demo_dataset(state)
        log_activity("settings.demo.clear", "settings", "demo_data", response.model_dump())
        return response

    @app.post("/api/v1/settings/legacy-scan", response_model=LegacyScanResponse)
    def legacy_scan(payload: LegacyScanRequest) -> LegacyScanResponse:
        target = Path(payload.path).expanduser()
        if not target.exists():
            raise HTTPException(status_code=404, detail="Path not found")
        entries: list[LegacyScanEntry] = []
        for path in target.rglob("*"):
            if path.is_file() and path.suffix.lower() in {".db", ".sqlite", ".json", ".csv"}:
                suffix = path.suffix.lower()
                entries.append(
                    LegacyScanEntry(
                        path=str(path),
                        kind=suffix.lstrip("."),
                        importable=suffix in LEGACY_IMPORT_EXTENSIONS,
                    )
                )
        return LegacyScanResponse(
            path=str(target),
            found=[item.path for item in entries[:30]],
            entries=entries[:30],
            message="已完成旧数据候选扫描。JSON/CSV 可导入到资料缓冲池，DB/SQLite 仅保留扫描结果供后续适配。",
        )

    @app.get("/api/v1/clients", response_model=list[ClientSummary])
    def list_clients() -> list[ClientSummary]:
        return [build_client_summary(str(row["id"])) for row in state.db.fetchall("SELECT id FROM clients ORDER BY updated_at DESC")]

    @app.post("/api/v1/clients", response_model=ClientSummary)
    def create_client(payload: ClientMutationPayload) -> ClientSummary:
        client_id = new_id("client")
        timestamp = now_iso()
        state.db.execute(
            """
            INSERT INTO clients(id, name, alias, domain, type, intro, stage, created_at, updated_at)
            VALUES(?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (client_id, payload.name, payload.alias, payload.domain, payload.type, payload.intro, payload.stage, timestamp, timestamp),
        )
        thread_id = new_id("thread")
        state.db.execute(
            "INSERT INTO chat_threads(id, client_id, title, created_at, updated_at) VALUES(?, ?, ?, ?, ?)",
            (thread_id, client_id, "默认研判线程", timestamp, timestamp),
        )
        ensure_standard_client_folders(client_id)
        log_activity("client.create", "client", client_id, payload.model_dump())
        return build_client_summary(client_id)

    @app.put("/api/v1/clients/{client_id}", response_model=ClientSummary)
    def update_client(client_id: str, payload: ClientMutationPayload) -> ClientSummary:
        if not state.db.fetchone("SELECT 1 FROM clients WHERE id = ?", (client_id,)):
            raise HTTPException(status_code=404, detail="Client not found")
        state.db.execute(
            """
            UPDATE clients
            SET name = ?, alias = ?, domain = ?, type = ?, intro = ?, stage = ?, updated_at = ?
            WHERE id = ?
            """,
            (payload.name, payload.alias, payload.domain, payload.type, payload.intro, payload.stage, now_iso(), client_id),
        )
        ensure_standard_client_folders(client_id)
        log_activity("client.update", "client", client_id, payload.model_dump())
        return build_client_summary(client_id)

    @app.delete("/api/v1/clients/{client_id}/folders/{folder_id}")
    def delete_client_folder(client_id: str, folder_id: str) -> dict[str, bool]:
        build_client_summary(client_id)
        ensure_standard_client_folders(client_id)
        row = state.db.fetchone("SELECT * FROM client_folders WHERE id = ? AND client_id = ?", (folder_id, client_id))
        if not row:
            raise HTTPException(status_code=404, detail="Folder not found")
        label = str(row["label"])
        file_count = int(row["file_count"] or 0)
        if file_count > 0:
            raise HTTPException(status_code=400, detail="该文件夹下还有文件，暂时不能移除")
        hide_client_folder_label(client_id, label)
        log_activity("client.folder.hide", "client_folder", folder_id, {"clientId": client_id, "label": label})
        return {"deleted": True}

    @app.delete("/api/v1/clients/{client_id}")
    def delete_client(client_id: str) -> dict[str, bool]:
        row = state.db.fetchone("SELECT name FROM clients WHERE id = ?", (client_id,))
        if not row:
            raise HTTPException(status_code=404, detail="Client not found")
        client_name = str(row["name"] or client_id)
        workspace_root = state.data_dir / "client_workspace" / client_id
        vector_root = state.data_dir / "vector_store" / client_id
        state.db.execute("DELETE FROM clients WHERE id = ?", (client_id,))
        for target in (workspace_root, vector_root):
            if target.exists():
                shutil.rmtree(target, ignore_errors=True)
        log_activity("client.delete", "client", client_id, {"name": client_name})
        return {"deleted": True}

    @app.get("/api/v1/clients/{client_id}/workspace", response_model=ClientWorkspaceResponse)
    def get_client_workspace(client_id: str) -> ClientWorkspaceResponse:
        return workspace_for_client(client_id)

    @app.get("/api/v1/clients/{client_id}/strategic-cockpit", response_model=StrategicCockpitSnapshotRecord)
    def get_strategic_cockpit(client_id: str) -> StrategicCockpitSnapshotRecord:
        build_client_summary(client_id)
        return build_strategic_cockpit_snapshot(client_id)

    @app.get("/api/v1/clients/{client_id}/strategic-cockpit/lines/{line_id}", response_model=StrategicLineDetailRecord)
    def get_strategic_line_detail(client_id: str, line_id: str) -> StrategicLineDetailRecord:
        build_client_summary(client_id)
        snapshot = build_strategic_cockpit_snapshot(client_id)
        line = next((item for item in snapshot.strategicLines if item.id == line_id), None)
        if line is None:
            raise HTTPException(status_code=404, detail="战略线不存在")
        task_rows = state.db.fetchall(
            """
            SELECT id, title
            FROM tasks
            WHERE client_id = ?
              AND (
                LOWER(COALESCE(title, '')) LIKE '%' || LOWER(?) || '%'
                OR LOWER(COALESCE(description, '')) LIKE '%' || LOWER(?) || '%'
                OR LOWER(COALESCE(recent_decision, '')) LIKE '%' || LOWER(?) || '%'
                OR LOWER(COALESCE(current_blocker, '')) LIKE '%' || LOWER(?) || '%'
              )
            ORDER BY updated_at DESC, created_at DESC
            LIMIT 8
            """,
            (client_id, line.title, line.title, line.title, line.title),
        )
        return StrategicLineDetailRecord(
            **line.model_dump(),
            clientId=snapshot.clientId,
            clientName=snapshot.clientName,
            stageLabel=snapshot.stageLabel,
            relatedTaskIds=[str(row["id"]) for row in task_rows],
            relatedTaskTitles=[str(row["title"]) for row in task_rows],
            contextSummary=f"战略线「{line.title}」当前处在「{line.stage or snapshot.stageLabel}」，下一步是「{line.nextStep}」。",
        )

    @app.post("/api/v1/clients/{client_id}/strategic-cockpit/confirm", response_model=StrategicCockpitSnapshotRecord)
    def confirm_strategic_cockpit(client_id: str, payload: StrategicCockpitConfirmPayload) -> StrategicCockpitSnapshotRecord:
        build_client_summary(client_id)
        session_user = _require_strategic_ceo()
        save_strategic_cockpit_snapshot(client_id, payload, session_user)
        snapshot = build_strategic_cockpit_snapshot(client_id)
        growth_user_id, growth_user_name = resolve_growth_actor()
        ingest_strategic_growth_candidate(
            state.db,
            user_id=growth_user_id,
            user_name=growth_user_name,
            snapshot=snapshot,
            source_type="strategic_confirm",
            source_id=client_id,
            created_at=now_iso(),
        )
        return snapshot

    @app.post("/api/v1/clients/{client_id}/strategic-cockpit/meeting-pack", response_model=MeetingPipelineResponse)
    def create_strategic_meeting_pack(client_id: str) -> MeetingPipelineResponse:
        build_client_summary(client_id)
        _require_strategic_ceo()
        meeting = _create_strategic_meeting_pack(client_id)
        return MeetingPipelineResponse(meeting=meeting, message="战略陪伴周会草稿已创建，并已把周会清单正式写入会议对象。")

    @app.post("/api/v1/clients/{client_id}/strategic-cockpit/meeting-pack/{meeting_id}/apply", response_model=StrategicCockpitSnapshotRecord)
    def apply_strategic_meeting_pack(client_id: str, meeting_id: str) -> StrategicCockpitSnapshotRecord:
        build_client_summary(client_id)
        session_user = _require_strategic_ceo()
        meeting = build_meeting_detail(meeting_id)
        if meeting.clientId != client_id:
            raise HTTPException(status_code=404, detail="Meeting not found")
        has_post_meeting_content = bool(
            meeting.transcriptText.strip()
            or meeting.decisions
            or meeting.actionItems
            or meeting.risks
            or meeting.ambiguities
        )
        if not has_post_meeting_content:
            raise HTTPException(status_code=400, detail="当前会议还没有会后内容，无法回填战略判断。")
        payload = _build_strategic_payload_from_meeting(client_id, meeting)
        save_strategic_cockpit_snapshot(client_id, payload, session_user)
        record_meeting_publish_writeback(
            state.db,
            client_id=client_id,
            meeting_id=meeting_id,
            meeting_title=meeting.title,
            event_line_ids=_strategic_meeting_event_line_ids(client_id, meeting.title, meeting_id=meeting_id),
        )
        log_activity("strategic_cockpit.apply_meeting_pack", "meeting", meeting_id, {"clientId": client_id})
        snapshot = build_strategic_cockpit_snapshot(client_id)
        growth_user_id, growth_user_name = resolve_growth_actor()
        ingest_strategic_growth_candidate(
            state.db,
            user_id=growth_user_id,
            user_name=growth_user_name,
            snapshot=snapshot,
            source_type="strategic_meeting_apply",
            source_id=meeting_id,
            meeting_id=meeting_id,
            created_at=now_iso(),
        )
        return snapshot

    @app.get("/api/v1/clients/{client_id}/notebook", response_model=ClientNotebookResponse)
    def get_client_notebook(client_id: str) -> ClientNotebookResponse:
        build_client_summary(client_id)
        return get_client_notebook_response(state.db, client_id)

    @app.get("/api/v1/clients/{client_id}/memory-status", response_model=MemoryStatus)
    def get_client_memory_status_route(client_id: str) -> MemoryStatus:
        build_client_summary(client_id)
        return get_client_memory_status(state.db, client_id)

    @app.post("/api/v1/memory/backfill", response_model=MemoryBackfillResultRecord)
    def backfill_memory_foundation_route() -> MemoryBackfillResultRecord:
        return backfill_memory_foundation(state.db)

    @app.post("/api/v1/clarifications", response_model=ClarificationRecord)
    def create_clarification(payload: ClarificationCreatePayload) -> ClarificationRecord:
        if payload.scopeType == "client":
            build_client_summary(payload.scopeId)
        elif payload.scopeType == "event_line" and not state.db.fetchone("SELECT id FROM event_lines WHERE id = ?", (payload.scopeId,)):
            raise HTTPException(status_code=404, detail="Event line not found")
        return create_clarification_record(state.db, payload)

    @app.post("/api/v1/clarifications/{clarification_id}/answer", response_model=ClarificationRecord)
    def answer_clarification(clarification_id: str, payload: ClarificationAnswerPayload) -> ClarificationRecord:
        try:
            return answer_clarification_record(state.db, clarification_id, payload)
        except KeyError as exc:
            raise HTTPException(status_code=404, detail="Clarification not found") from exc
        except ValueError as exc:
            raise HTTPException(status_code=400, detail=str(exc)) from exc

    @app.get("/api/v1/clients/{client_id}/knowledge/reclass-events", response_model=list[FileReclassEventRecord])
    def list_reclass_events(client_id: str) -> list[FileReclassEventRecord]:
        build_client_summary(client_id)
        return [FileReclassEventRecord(**item) for item in fetch_recent_reclass_events(state.db, client_id, limit=50)]

    @app.get("/api/v1/clients/{client_id}/knowledge/status", response_model=KnowledgeStatusRecord)
    def get_client_knowledge_status(client_id: str) -> KnowledgeStatusRecord:
        build_client_summary(client_id)
        ensure_standard_client_folders(client_id)
        return build_knowledge_status_record(client_id)

    @app.post("/api/v1/clients/{client_id}/knowledge/rebuild", response_model=KnowledgeJobRecord)
    def rebuild_client_knowledge(client_id: str) -> KnowledgeJobRecord:
        build_client_summary(client_id)
        pending = state.db.fetchone(
            "SELECT * FROM knowledge_jobs WHERE client_id = ? AND status IN ('queued', 'running') ORDER BY created_at DESC LIMIT 1",
            (client_id,),
        )
        if pending:
            return KnowledgeJobRecord(
                id=str(pending["id"]),
                clientId=str(pending["client_id"]),
                jobType=str(pending["job_type"]),
                status=str(pending["status"]),  # type: ignore[arg-type]
                totalItems=int(pending["total_items"]),
                processedItems=int(pending["processed_items"]),
                lastError=str(pending["last_error"]) if pending["last_error"] else None,
                createdAt=str(pending["created_at"]),
                startedAt=str(pending["started_at"]) if pending["started_at"] else None,
                finishedAt=str(pending["finished_at"]) if pending["finished_at"] else None,
                updatedAt=str(pending["updated_at"]),
            )
        total_items = int(state.db.scalar("SELECT COUNT(1) AS count FROM documents WHERE client_id = ?", (client_id,)))
        job = enqueue_knowledge_job(
            client_id,
            "rebuild_client_knowledge",
            {"clientId": client_id},
            total_items=total_items,
        )
        log_activity("knowledge.rebuild", "client", client_id, {"jobId": job.id})
        return job

    @app.post("/api/v1/clients/{client_id}/knowledge/search", response_model=KnowledgeSearchResponse)
    def search_client_knowledge(client_id: str, payload: ChatRequest) -> KnowledgeSearchResponse:
        build_client_summary(client_id)
        query = payload.prompt.strip()
        retrieval_started = perf_counter()
        bundle = build_retrieval_bundle(client_id, query)
        retrieval_elapsed_ms = round((perf_counter() - retrieval_started) * 1000, 2)
        hits = [
            KnowledgeSearchHitRecord(
                title=item.title,
                excerpt=item.excerpt,
                score=item.score,
                stage={"master_index": "master_index", "surrogate": "surrogate"}.get(item.source_stage, "raw_chunk"),  # type: ignore[arg-type]
                path=item.path,
                sectionLabel=item.section_label,
                matchedTerms=item.matched_terms,
            )
            for item in bundle.citations
        ]
        preview_evidence = [
            EvidenceItem(
                id=new_id("ev"),
                title=item.title,
                excerpt=item.excerpt,
                sourceType="knowledge_chunk" if item.chunk_id else "knowledge_document",
                documentId=item.knowledge_document_id,
                path=item.path,
                score=item.score,
                coverage=item.coverage,
                sectionLabel=item.section_label,
                retrievalStage={"master_index": "master_index", "surrogate": "surrogate"}.get(item.source_stage, "raw_chunk"),
                isFallback=item.source_stage == "master_index",
                matchedTerms=item.matched_terms,
            )
            for item in bundle.citations
        ]
        retrieval_meta = bundle.retrieval_summary if isinstance(bundle.retrieval_summary, dict) else {}
        preview_summary = build_retrieval_preview_summary(client_id, query, preview_evidence, bundle)
        work_trace = build_answer_work_trace(query, preview_evidence, bundle)
        bundle.retrieval_summary = {
            **retrieval_meta,
            "previewSummary": preview_summary,
            "workTrace": work_trace,
        }
        retrieval_meta = bundle.retrieval_summary
        search_id = persist_retrieval_bundle(client_id, query, payload.threadId, bundle, retrieval_elapsed_ms)
        return KnowledgeSearchResponse(
            searchId=search_id,
            clientId=client_id,
            query=query,
            coverage=bundle.coverage,
            matchedTerms=bundle.matched_terms,
            masterHitCount=int(retrieval_meta.get("masterHitCount", 0) or 0),
            surrogateHitCount=int(retrieval_meta.get("surrogateHitCount", 0) or 0),
            rawChunkHitCount=int(retrieval_meta.get("rawChunkHitCount", 0) or 0),
            drillthroughUsed=bool(retrieval_meta.get("drillthroughUsed", False)),
            strategicMode=bool(retrieval_meta.get("strategicMode", False)),
            categoryCoverage=[str(item) for item in retrieval_meta.get("categoryCoverage", []) if str(item).strip()] if isinstance(retrieval_meta.get("categoryCoverage"), list) else [],
            preferredCategories=[str(item) for item in retrieval_meta.get("preferredCategories", []) if str(item).strip()] if isinstance(retrieval_meta.get("preferredCategories"), list) else [],
            phase="grounding",
            progress=38.0,
            progressFloor=25.0,
            progressCeiling=55.0,
            stageLabel="庆华已经整理好当前问题所需的背景材料，准备调用千问组织答案",
            lastUpdatedAt=now_iso(),
            failureReason=bundle.failure_reason,
            hits=hits,
            previewSummary=preview_summary,
        )

    @app.get("/api/v1/clients/{client_id}/goals", response_model=list[GoalRecord])
    def list_client_goals(client_id: str) -> list[GoalRecord]:
        return workspace_for_client(client_id).goals

    @app.post("/api/v1/clients/{client_id}/goals", response_model=GoalRecord)
    def create_goal(client_id: str, payload: GoalPayload) -> GoalRecord:
        build_client_summary(client_id)
        goal_id = new_id("goal")
        timestamp = now_iso()
        state.db.execute(
            """
            INSERT INTO goal_records(id, client_id, title, quarter, progress, owner_name, created_at, updated_at)
            VALUES(?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (goal_id, client_id, payload.title, payload.quarter, payload.progress, payload.ownerName, timestamp, timestamp),
        )
        log_activity("goal.create", "goal", goal_id, payload.model_dump())
        return GoalRecord(id=goal_id, clientId=client_id, title=payload.title, quarter=payload.quarter, progress=payload.progress, ownerName=payload.ownerName)

    @app.get("/api/v1/clients/{client_id}/dna-documents", response_model=ClientDnaModulesResponse)
    def list_client_dna_documents(client_id: str) -> ClientDnaModulesResponse:
        build_client_summary(client_id)
        return ClientDnaModulesResponse(modules=list_client_dna_modules(client_id))

    @app.post("/api/v1/clients/{client_id}/dna-documents/generate", response_model=KnowledgeJobRecord)
    def generate_client_dna_documents(client_id: str, payload: ClientDnaGeneratePayload) -> KnowledgeJobRecord:
        build_client_summary(client_id)
        ensure_standard_client_folders(client_id)
        job = maybe_enqueue_client_dna_generation_job(client_id, refresh_generated=payload.refreshGenerated)
        if job is not None:
            log_activity(
                "client.dna_document.generate",
                "client",
                client_id,
                {"refreshGenerated": payload.refreshGenerated, "jobId": job.id},
            )
            return job
        return KnowledgeJobRecord(
            id=new_id("kjob"),
            clientId=client_id,
            jobType="generate_client_dna_candidates",
            status="completed",
            totalItems=0,
            processedItems=0,
            createdAt=now_iso(),
            updatedAt=now_iso(),
        )

    @app.get("/api/v1/clients/{client_id}/dna-documents/{module_key}", response_model=ClientDnaModuleRecord)
    def get_client_dna_document(client_id: str, module_key: str) -> ClientDnaModuleRecord:
        build_client_summary(client_id)
        modules = {module.moduleKey: module for module in list_client_dna_modules(client_id)}
        module = modules.get(module_key)
        if not module:
            raise HTTPException(status_code=404, detail="未知的客户 DNA 模块")
        return module

    @app.post("/api/v1/clients/{client_id}/dna-documents/{module_key}", response_model=ClientDnaModuleRecord)
    def update_client_dna_document(client_id: str, module_key: str, payload: OrganizationDnaUploadPayload) -> ClientDnaModuleRecord:
        build_client_summary(client_id)
        record = upsert_client_dna_module(client_id, module_key, payload)
        log_activity(
            "client.dna_document.update",
            "dna_document",
            f"{client_id}:{module_key}",
            {"clientId": client_id, "moduleKey": module_key, "fileName": record.fileName, "contentHash": record.contentHash},
        )
        return record

    @app.get("/api/v1/clients/{client_id}/project-structure", response_model=ProjectStructureResponse)
    def get_client_project_structure(client_id: str) -> ProjectStructureResponse:
        return build_project_structure(client_id)

    @app.get("/api/v1/clients/{client_id}/project-modules/{module_id}", response_model=ProjectModuleDetailRecord)
    def get_client_project_module_detail(client_id: str, module_id: str) -> ProjectModuleDetailRecord:
        build_client_summary(client_id)
        return get_project_module_detail(client_id, module_id)

    @app.post("/api/v1/clients/{client_id}/project-modules", response_model=ProjectModuleRecord)
    def create_client_project_module(client_id: str, payload: ProjectModulePayload) -> ProjectModuleRecord:
        build_client_summary(client_id)
        timestamp = now_iso()
        module_id = new_id("pmodule")
        alias = payload.alias.strip() if payload.alias else None
        owner_name = payload.ownerName.strip() if payload.ownerName else None
        deliverables = _sanitize_text_list(payload.deliverables)
        keywords = _sanitize_text_list(payload.keywords)
        state.db.execute(
            """
            INSERT INTO project_modules(
                id, client_id, name, alias, goal, description, owner_name, deliverables_json, keywords_json, created_at, updated_at
            ) VALUES(?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                module_id,
                client_id,
                payload.name.strip(),
                alias,
                (payload.goal or "").strip(),
                (payload.description or "").strip(),
                owner_name,
                to_json(deliverables),
                to_json(keywords),
                timestamp,
                timestamp,
            ),
        )
        row = state.db.fetchone("SELECT * FROM project_modules WHERE id = ?", (module_id,))
        if not row:
            raise HTTPException(status_code=500, detail="任务模块创建失败")
        log_activity("project.module.create", "project_module", module_id, {"clientId": client_id, "name": payload.name.strip()})
        return _project_module_record(row)

    @app.patch("/api/v1/clients/{client_id}/project-modules/{module_id}", response_model=ProjectModuleRecord)
    def update_client_project_module(client_id: str, module_id: str, payload: ProjectModulePayload) -> ProjectModuleRecord:
        build_client_summary(client_id)
        row = state.db.fetchone("SELECT * FROM project_modules WHERE id = ? AND client_id = ?", (module_id, client_id))
        if not row:
            raise HTTPException(status_code=404, detail="任务模块不存在")
        merged = {
            "name": payload.name.strip(),
            "alias": payload.alias.strip() if payload.alias else None,
            "goal": (payload.goal or "").strip(),
            "description": (payload.description or "").strip(),
            "owner_name": payload.ownerName.strip() if payload.ownerName else None,
            "deliverables_json": to_json(_sanitize_text_list(payload.deliverables)),
            "keywords_json": to_json(_sanitize_text_list(payload.keywords)),
            "updated_at": now_iso(),
        }
        state.db.execute(
            """
            UPDATE project_modules
            SET name = ?, alias = ?, goal = ?, description = ?, owner_name = ?, deliverables_json = ?, keywords_json = ?, updated_at = ?
            WHERE id = ?
            """,
            (
                merged["name"],
                merged["alias"],
                merged["goal"],
                merged["description"],
                merged["owner_name"],
                merged["deliverables_json"],
                merged["keywords_json"],
                merged["updated_at"],
                module_id,
            ),
        )
        updated = state.db.fetchone("SELECT * FROM project_modules WHERE id = ?", (module_id,))
        if not updated:
            raise HTTPException(status_code=500, detail="任务模块更新失败")
        log_activity("project.module.update", "project_module", module_id, {"clientId": client_id, "name": merged["name"]})
        return _project_module_record(updated)

    @app.get("/api/v1/clients/{client_id}/project-flows/{flow_id}", response_model=ProjectFlowDetailRecord)
    def get_client_project_flow_detail(client_id: str, flow_id: str) -> ProjectFlowDetailRecord:
        build_client_summary(client_id)
        return get_project_flow_detail(client_id, flow_id)

    @app.post("/api/v1/clients/{client_id}/project-flows", response_model=ProjectFlowRecord)
    def create_client_project_flow(client_id: str, payload: ProjectFlowPayload) -> ProjectFlowRecord:
        build_client_summary(client_id)
        module_row = state.db.fetchone("SELECT * FROM project_modules WHERE id = ? AND client_id = ?", (payload.moduleId, client_id))
        if not module_row:
            raise HTTPException(status_code=400, detail="请先选择当前项目下的任务模块")
        timestamp = now_iso()
        flow_id = new_id("pflow")
        state.db.execute(
            """
            INSERT INTO project_flows(
                id, client_id, module_id, name, description, scenario, trigger_condition, steps_json, inputs_json, outputs_json, collaborators_json, risk_points_json, created_at, updated_at
            ) VALUES(?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                flow_id,
                client_id,
                payload.moduleId,
                payload.name.strip(),
                (payload.description or "").strip(),
                (payload.scenario or "").strip(),
                (payload.triggerCondition or "").strip(),
                to_json(_sanitize_text_list(payload.steps)),
                to_json(_sanitize_text_list(payload.inputs)),
                to_json(_sanitize_text_list(payload.outputs)),
                to_json(_sanitize_text_list(payload.collaborators)),
                to_json(_sanitize_text_list(payload.riskPoints)),
                timestamp,
                timestamp,
            ),
        )
        row = state.db.fetchone(
            """
            SELECT f.*, m.name AS module_name
            FROM project_flows f
            LEFT JOIN project_modules m ON m.id = f.module_id
            WHERE f.id = ?
            """,
            (flow_id,),
        )
        if not row:
            raise HTTPException(status_code=500, detail="流程创建失败")
        log_activity("project.flow.create", "project_flow", flow_id, {"clientId": client_id, "moduleId": payload.moduleId, "name": payload.name.strip()})
        return _project_flow_record(row, str(row["module_name"]) if row["module_name"] else None)

    @app.patch("/api/v1/clients/{client_id}/project-flows/{flow_id}", response_model=ProjectFlowRecord)
    def update_client_project_flow(client_id: str, flow_id: str, payload: ProjectFlowPayload) -> ProjectFlowRecord:
        build_client_summary(client_id)
        module_row = state.db.fetchone("SELECT * FROM project_modules WHERE id = ? AND client_id = ?", (payload.moduleId, client_id))
        if not module_row:
            raise HTTPException(status_code=400, detail="请先选择当前项目下的任务模块")
        row = state.db.fetchone("SELECT * FROM project_flows WHERE id = ? AND client_id = ?", (flow_id, client_id))
        if not row:
            raise HTTPException(status_code=404, detail="流程不存在")
        state.db.execute(
            """
            UPDATE project_flows
            SET module_id = ?, name = ?, description = ?, scenario = ?, trigger_condition = ?, steps_json = ?, inputs_json = ?, outputs_json = ?, collaborators_json = ?, risk_points_json = ?, updated_at = ?
            WHERE id = ?
            """,
            (
                payload.moduleId,
                payload.name.strip(),
                (payload.description or "").strip(),
                (payload.scenario or "").strip(),
                (payload.triggerCondition or "").strip(),
                to_json(_sanitize_text_list(payload.steps)),
                to_json(_sanitize_text_list(payload.inputs)),
                to_json(_sanitize_text_list(payload.outputs)),
                to_json(_sanitize_text_list(payload.collaborators)),
                to_json(_sanitize_text_list(payload.riskPoints)),
                now_iso(),
                flow_id,
            ),
        )
        updated = state.db.fetchone(
            """
            SELECT f.*, m.name AS module_name
            FROM project_flows f
            LEFT JOIN project_modules m ON m.id = f.module_id
            WHERE f.id = ?
            """,
            (flow_id,),
        )
        if not updated:
            raise HTTPException(status_code=500, detail="流程更新失败")
        log_activity("project.flow.update", "project_flow", flow_id, {"clientId": client_id, "moduleId": payload.moduleId, "name": payload.name.strip()})
        return _project_flow_record(updated, str(updated["module_name"]) if updated["module_name"] else None)

    @app.get("/api/v1/clients/{client_id}/dna", response_model=list[DnaTerm])
    def list_client_dna(client_id: str) -> list[DnaTerm]:
        return workspace_for_client(client_id).dnaTerms

    @app.post("/api/v1/clients/{client_id}/dna", response_model=DnaTerm)
    def upsert_client_dna(client_id: str, payload: DnaTermPayload) -> DnaTerm:
        build_client_summary(client_id)
        existing = state.db.fetchone(
            "SELECT * FROM dna_terms WHERE client_id = ? AND canonical_name = ?",
            (client_id, payload.canonicalName),
        )
        timestamp = now_iso()
        if existing:
            state.db.execute(
                """
                UPDATE dna_terms
                SET category = ?, aliases_json = ?, description = ?, updated_at = ?
                WHERE id = ?
                """,
                (payload.category, to_json(payload.aliases), payload.description, timestamp, existing["id"]),
            )
            term_id = str(existing["id"])
        else:
            term_id = new_id("dna")
            state.db.execute(
                """
                INSERT INTO dna_terms(id, client_id, category, canonical_name, aliases_json, description, created_at, updated_at)
                VALUES(?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (term_id, client_id, payload.category, payload.canonicalName, to_json(payload.aliases), payload.description, timestamp, timestamp),
            )
        log_activity("dna.upsert", "dna", term_id, payload.model_dump())
        return DnaTerm(
            id=term_id,
            clientId=client_id,
            category=payload.category,
            canonicalName=payload.canonicalName,
            aliases=payload.aliases,
            description=payload.description,
            sourceLevel="client",
        )

    @app.post("/api/v1/imports", response_model=list[ImportRecord])
    def import_documents(payload: ImportPayload) -> list[ImportRecord]:
        build_client_summary(payload.clientId)
        ensure_standard_client_folders(payload.clientId)
        results: list[ImportRecord] = []
        allowed_extensions = set(SUPPORTED_IMPORT_EXTENSIONS)
        if payload.allowLegacy:
            allowed_extensions.update(LEGACY_IMPORT_EXTENSIONS)
        for raw_path in payload.paths:
            source_path = Path(raw_path).expanduser()
            if not source_path.exists():
                continue
            import_id = new_id("imp")
            timestamp = now_iso()
            queued = 0
            skipped = 0
            state.db.execute(
                """
                INSERT INTO imports(id, client_id, source_path, mode, status, imported_count, skipped_count, created_at)
                VALUES(?, ?, ?, ?, 'queued', 0, 0, ?)
                """,
                (import_id, payload.clientId, str(source_path), payload.mode, timestamp),
            )
            if payload.mode == "folder":
                candidates = [path for path in source_path.rglob("*") if path.is_file()]
            else:
                candidates = [source_path]
            ensure_source_tree_snapshot(
                state.db,
                import_id=import_id,
                client_id=payload.clientId,
                source_path=source_path,
                mode=payload.mode,
                created_at=timestamp,
            )
            queued_documents: list[dict[str, object]] = []
            for path in candidates:
                if path.suffix.lower() not in allowed_extensions:
                    skipped += 1
                    continue
                if state.db.fetchone(
                    """
                    SELECT 1
                    FROM knowledge_documents
                    WHERE client_id = ?
                      AND (import_source_path = ? OR current_human_path = ? OR original_path = ?)
                    """,
                    (payload.clientId, str(path), str(path), str(path)),
                ):
                    skipped += 1
                    continue
                managed_import_path = stage_import_copy(state.data_dir, payload.clientId, import_id, path)
                excerpt = build_excerpt(path)
                document_id = new_id("doc")
                state.db.execute(
                    """
                    INSERT INTO documents(id, client_id, folder_id, title, path, original_source_path, kind, source, excerpt, tags_json, created_at)
                    VALUES(?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """,
                    (
                        document_id,
                        payload.clientId,
                        None,
                        path.name,
                        str(managed_import_path),
                        str(path),
                        path.suffix.lower().lstrip("."),
                        payload.mode,
                        excerpt,
                        to_json([path.suffix.lower().lstrip(".")]),
                        timestamp,
                    ),
                )
                queued_documents.append(
                    {
                        "documentId": document_id,
                        "sourcePath": str(managed_import_path),
                        "originalSourcePath": str(path),
                        "title": path.name,
                        "kind": path.suffix.lower().lstrip("."),
                        "source": payload.mode,
                        "createdAt": timestamp,
                    }
                )
                queued += 1
            state.db.execute(
                "UPDATE imports SET skipped_count = ? WHERE id = ?",
                (skipped, import_id),
            )
            if queued == 0:
                state.db.execute(
                    "UPDATE imports SET status = 'completed', imported_count = 0, skipped_count = ? WHERE id = ?",
                    (skipped, import_id),
                )
                log_activity("import.create", "import", import_id, {"clientId": payload.clientId, "sourcePath": str(source_path), "queued": queued, "skipped": skipped, "jobId": None})
                results.append(
                    ImportRecord(
                        id=import_id,
                        clientId=payload.clientId,
                        sourcePath=str(source_path),
                        mode=payload.mode,
                        status="completed",
                        importedCount=0,
                        skippedCount=skipped,
                        createdAt=timestamp,
                    )
                )
                continue
            job = enqueue_knowledge_job(
                payload.clientId,
                "ingest_import",
                {
                    "clientId": payload.clientId,
                    "importId": import_id,
                    "mode": payload.mode,
                    "documents": queued_documents,
                },
                total_items=queued,
            )
            log_activity("import.create", "import", import_id, {"clientId": payload.clientId, "sourcePath": str(source_path), "queued": queued, "skipped": skipped, "jobId": job.id})
            results.append(
                ImportRecord(
                    id=import_id,
                    clientId=payload.clientId,
                    sourcePath=str(source_path),
                    mode=payload.mode,
                    status="queued",
                    importedCount=queued,
                    skippedCount=skipped,
                    createdAt=timestamp,
                )
            )
        return results

    def ensure_chat_thread(client_id: str, thread_id: str | None, prompt: str, timestamp: str) -> str:
        existing_thread_id = thread_id
        if existing_thread_id and state.db.fetchone("SELECT 1 FROM chat_threads WHERE id = ?", (existing_thread_id,)):
            return existing_thread_id
        next_thread_id = new_id("thread")
        state.db.execute(
            "INSERT INTO chat_threads(id, client_id, title, created_at, updated_at) VALUES(?, ?, ?, ?, ?)",
            (next_thread_id, client_id, prompt[:16], timestamp, timestamp),
        )
        return next_thread_id

    def insert_user_chat_message(thread_id: str, prompt: str, timestamp: str) -> str:
        message_id = new_id("msg")
        state.db.execute(
            """
            INSERT INTO chat_messages(
                id, thread_id, role, content, structured_data_json, model_route, llm_invoked, provider_used,
                answer_mode, evidence_status, failure_reason, timing_json, retrieval_summary_json, evidence_json, status, created_at
            )
            VALUES(?, ?, 'user', ?, NULL, NULL, 0, NULL, NULL, NULL, NULL, '{}', '{}', '[]', 'success', ?)
            """,
            (message_id, thread_id, prompt, timestamp),
        )
        return message_id

    def phase_progress_window(phase: str) -> tuple[float, float]:
        if phase == "retrieving":
            return 0.0, 25.0
        if phase == "grounding":
            return 25.0, 55.0
        if phase == "generating":
            return 55.0, 92.0
        return 100.0, 100.0

    def insert_loading_assistant_message(thread_id: str, retrieval_summary: dict[str, object], timestamp: str) -> str:
        assistant_id = new_id("msg")
        provider_used = state.ai.current_provider()
        merged_summary = dict(retrieval_summary)
        merged_summary.setdefault("startedAt", timestamp)
        phase = str(merged_summary.get("phase") or "").strip()
        if not phase:
            has_hits = any(int(merged_summary.get(key, 0) or 0) > 0 for key in ("masterHitCount", "surrogateHitCount", "rawChunkHitCount"))
            phase = "grounding" if has_hits else "retrieving"
        merged_summary["phase"] = phase
        merged_summary.setdefault("progress", 36.0 if phase == "grounding" else 6.0)
        floor, ceiling = phase_progress_window(phase)
        merged_summary["progressFloor"] = float(merged_summary.get("progressFloor", floor) or floor)
        merged_summary["progressCeiling"] = float(merged_summary.get("progressCeiling", ceiling) or ceiling)
        merged_summary.setdefault(
            "stageLabel",
            "背景材料已整理完成，正在准备调用千问组织答案" if phase == "grounding" else "正在整理客户背景材料",
        )
        merged_summary["lastUpdatedAt"] = timestamp
        loading_content = str(merged_summary.get("stageLabel") or "庆华正在整理背景材料，并组织分析答案……")
        state.db.execute(
            """
            INSERT INTO chat_messages(
                id, thread_id, role, content, structured_data_json, model_route, llm_invoked, provider_used,
                answer_mode, evidence_status, failure_reason, timing_json, retrieval_summary_json, evidence_json, status, created_at
            )
            VALUES(?, ?, 'assistant', ?, NULL, ?, 0, ?, NULL, NULL, NULL, '{}', ?, '[]', 'loading', ?)
            """,
            (
                assistant_id,
                thread_id,
                loading_content,
                f"AI · {provider_used}",
                provider_used,
                to_json(merged_summary),
                timestamp,
            ),
        )
        return assistant_id

    def update_loading_assistant_message(
        assistant_id: str,
        *,
        retrieval_summary: dict[str, object] | None = None,
        timing: dict[str, float] | None = None,
        content: str | None = None,
    ) -> None:
        current_row = state.db.fetchone("SELECT retrieval_summary_json, timing_json, content FROM chat_messages WHERE id = ?", (assistant_id,))
        if not current_row:
            return
        merged_summary = from_json(str(current_row["retrieval_summary_json"] or "{}"), {})
        if not isinstance(merged_summary, dict):
            merged_summary = {}
        if retrieval_summary:
            merged_summary.update(retrieval_summary)
        merged_timing = from_json(str(current_row["timing_json"] or "{}"), {})
        if not isinstance(merged_timing, dict):
            merged_timing = {}
        if timing:
            merged_timing.update(timing)
        phase = str(merged_summary.get("phase") or "retrieving").strip() or "retrieving"
        floor, ceiling = phase_progress_window(phase)
        merged_summary["progressFloor"] = float(merged_summary.get("progressFloor", floor) or floor)
        merged_summary["progressCeiling"] = float(merged_summary.get("progressCeiling", ceiling) or ceiling)
        merged_summary["lastUpdatedAt"] = now_iso()
        next_content = content if content is not None else str(current_row["content"] or "")
        state.db.execute(
            """
            UPDATE chat_messages
            SET content = ?, retrieval_summary_json = ?, timing_json = ?
            WHERE id = ? AND status = 'loading'
            """,
            (next_content, to_json(merged_summary), to_json(merged_timing), assistant_id),
        )

    def resolve_chat_answer(
        client_id: str,
        thread_id: str,
        prompt: str,
        assistant_id: str,
        search_id: str | None,
        request_started: float,
        run_id: str | None = None,
    ) -> ChatMessageRecord:
        if is_client_analysis_run_canceled(run_id):
            return fetch_chat_message_for_client(client_id, assistant_id)
        retrieval_bundle = None
        retrieval_elapsed_ms = 0.0
        cached_retrieval = False
        if search_id:
            retrieval_bundle, retrieval_elapsed_ms = load_cached_retrieval_bundle(client_id, search_id, prompt)
            cached_retrieval = retrieval_bundle is not None
        if retrieval_bundle is None:
            retrieval_started = perf_counter()
            retrieval_bundle = build_retrieval_bundle(client_id, prompt)
            retrieval_elapsed_ms = round((perf_counter() - retrieval_started) * 1000, 2)
        if search_id:
            retrieval_bundle.retrieval_summary = {
                **(retrieval_bundle.retrieval_summary if isinstance(retrieval_bundle.retrieval_summary, dict) else {}),
                "searchId": search_id,
                "cacheHit": cached_retrieval,
            }
        evidence = [
            EvidenceItem(
                id=new_id("ev"),
                title=item.title,
                excerpt=item.excerpt,
                sourceType="knowledge_chunk" if item.chunk_id else "knowledge_document",
                documentId=item.knowledge_document_id,
                path=item.path,
                score=item.score,
                coverage=item.coverage,
                sectionLabel=item.section_label,
                retrievalStage={"master_index": "master_index", "surrogate": "surrogate"}.get(item.source_stage, "raw_chunk"),
                isFallback=item.source_stage == "master_index",
                matchedTerms=item.matched_terms,
            )
            for item in retrieval_bundle.citations
        ]
        retrieval_meta = retrieval_bundle.retrieval_summary if isinstance(retrieval_bundle.retrieval_summary, dict) else {}
        client_dna_priority_note = build_client_dna_priority_note(client_id, prompt)
        if client_dna_priority_note:
            retrieval_bundle.retrieval_summary = {
                **retrieval_meta,
                "clientDnaPriority": client_dna_priority_note,
                "clientDnaTrail": client_dna_priority_note.replace("本次已优先参考客户 DNA 背景：", "").split("、"),
            }
            retrieval_meta = retrieval_bundle.retrieval_summary
        surrogate_hit_count = int(retrieval_meta.get("surrogateHitCount", 0) or 0)
        raw_chunk_hit_count = int(retrieval_meta.get("rawChunkHitCount", 0) or 0)
        has_grounded_evidence = bool(evidence)

        if has_grounded_evidence:
            answer_mode = "grounded_answer"
            evidence_status = "sufficient"
            retrieval_stage = "raw_chunk"
        else:
            answer_mode = "general_answer"
            evidence_status = "none"
            retrieval_stage = "background_only"
            evidence = []
        evidence_summary = build_analysis_evidence_summary(client_id, prompt, retrieval_bundle)
        work_trace = build_answer_work_trace(prompt, evidence, retrieval_bundle)
        if is_client_analysis_run_canceled(run_id):
            return fetch_chat_message_for_client(client_id, assistant_id)
        if run_id:
            update_client_analysis_run(
                run_id,
                status="running",
                phase="evidence_ready",
                progress=45.0,
                progress_floor=30.0,
                progress_ceiling=45.0,
                stage_label="背景材料已整理，正在组织长回答",
                elapsed_ms=retrieval_elapsed_ms,
                evidence_summary=evidence_summary,
                long_answer_status="pending",
                summary_status="pending",
                timing={"retrievalMs": retrieval_elapsed_ms},
            )
        pre_llm_meta = {
            **retrieval_meta,
            "retrievalStage": retrieval_stage,
            "answerMode": answer_mode,
            "evidenceStatus": evidence_status,
            "llmInvoked": False,
            "previewSummary": str(retrieval_meta.get("previewSummary") or build_retrieval_preview_summary(client_id, prompt, evidence, retrieval_bundle)),
            "workTrace": work_trace,
            "phase": "generating",
            "progress": 58.0 if answer_mode != "general_answer" else 46.0,
            "progressFloor": 55.0 if answer_mode != "general_answer" else 25.0,
            "progressCeiling": 92.0,
            "stageLabel": (
                "庆华已经整理好当前问题所需的背景材料，正在调用千问组织完整分析"
                if answer_mode != "general_answer"
                else "当前没有命中足够的原始材料，庆华正在生成通用背景判断"
            ),
        }
        update_loading_assistant_message(
            assistant_id,
            retrieval_summary=pre_llm_meta,
            timing={"retrievalMs": retrieval_elapsed_ms},
            content="庆华正在整理背景材料，并组织分析答案……",
        )
        if is_client_analysis_run_canceled(run_id):
            return fetch_chat_message_for_client(client_id, assistant_id)
        if run_id:
            update_client_analysis_run(
                run_id,
                status="running",
                phase="generating_long_answer",
                progress=float(pre_llm_meta.get("progress", 58.0) or 58.0),
                progress_floor=45.0,
                progress_ceiling=85.0,
                stage_label=str(pre_llm_meta.get("stageLabel") or "正在生成长回答"),
                elapsed_ms=retrieval_elapsed_ms,
                evidence_summary=evidence_summary,
                timing={"retrievalMs": retrieval_elapsed_ms},
            )

        llm_started = perf_counter()
        provider_used = state.ai.current_provider()
        llm_invoked = True
        model_route = f"AI · {provider_used}"
        local_fallback_used = False
        compact_model_fallback_used = False
        partial_generation_preserved = False
        memory_background_context, memory_background_meta = build_client_memory_background_context(client_id, prompt)
        if memory_background_meta.get("memoryBackgroundUsed"):
            retrieval_bundle.retrieval_summary = {
                **retrieval_meta,
                **memory_background_meta,
            }
            retrieval_meta = retrieval_bundle.retrieval_summary
        answer_context = build_chat_answer_context(
            client_id,
            prompt,
            evidence,
            retrieval_bundle,
            memory_background_context=memory_background_context,
        )
        identity_role_insufficient = False
        latest_partial_content = ""
        latest_partial_structured: dict[str, object] | None = None
        grounded_system_instruction = (
            "请直接基于给定的客户原始材料回答问题。"
            "不要把答案写成证据罗列、材料摘要或系统说明。"
            "不要预设固定格式、固定结构、固定段数或固定栏目。"
            "由模型自己决定如何组织回答。"
            "允许基于多条材料共同指向的信号做更高层的综合判断。"
            "只有原始材料里没有出现过的具体事实、数字、人名、时间和身份，不要写成已被证实。"
        )
        identity_sensitive_instruction = (
            grounded_system_instruction
            + "如果问题涉及创始人、负责人、理事长、秘书长等具体人物身份，只有在原始证据明确把人名与角色绑定时，才能下结论；否则只能说明当前证据不足以确认。"
        )
        grounded_fallback_instruction = (
            "请只基于当前已经整理出的原始证据继续回答。"
            "不要编造原始证据里没有的确定性事实。"
            "除此之外，不要替自己预设固定格式、固定结构、固定段数或固定栏目。"
            "不要停留在表层摘录；只要证据允许，就尽量把深层判断讲透。"
        )

        def load_preserved_partial() -> tuple[str, dict[str, object] | None]:
            preserved_content = latest_partial_content
            preserved_structured = latest_partial_structured
            if run_id and not preserved_content:
                row = state.db.fetchone(
                    "SELECT long_answer, structured_summary_json FROM client_analysis_runs WHERE id = ?",
                    (run_id,),
                )
                if row:
                    if row["long_answer"]:
                        preserved_content = str(row["long_answer"]).strip()
                    structured_payload = from_json(str(row["structured_summary_json"] or "{}"), {})
                    if isinstance(structured_payload, dict) and structured_payload:
                        preserved_structured = structured_payload
            return preserved_content, preserved_structured

        def push_partial_analysis(partial: dict[str, object]) -> None:
            nonlocal latest_partial_content, latest_partial_structured
            if is_client_analysis_run_canceled(run_id):
                return
            partial_content = str(partial.get("content") or "").strip()
            if not partial_content:
                return
            latest_partial_content = partial_content
            latest_partial_structured = partial.get("structured") if isinstance(partial.get("structured"), dict) else None
            partial_stage = str(partial.get("stageLabel") or "正在生成长回答")
            partial_progress = float(partial.get("progress") or 62.0)
            partial_structured = latest_partial_structured
            elapsed_now_ms = round((perf_counter() - request_started) * 1000, 2)
            update_client_analysis_run(
                run_id,
                status="running",
                phase="generating_long_answer",
                progress=partial_progress,
                progress_floor=58.0,
                progress_ceiling=95.0,
                stage_label=partial_stage,
                elapsed_ms=elapsed_now_ms,
                evidence_summary=evidence_summary,
                long_answer=partial_content,
                structured_summary=partial_structured if isinstance(partial_structured, dict) else None,
                long_answer_status="pending",
                summary_status="pending",
                answer_mode=answer_mode,
                llm_invoked=True,
                provider_used=provider_used,
                timing={
                    "retrievalMs": retrieval_elapsed_ms,
                    "llmMs": round((perf_counter() - llm_started) * 1000, 2),
                    "totalMs": elapsed_now_ms,
                },
            )

        def build_compact_model_fallback(failure_detail: str) -> AiStructuredResponse:
            note = build_compact_grounded_note(client_id, prompt, evidence, retrieval_bundle)
            return state.ai.generate_compact_grounded_fallback(
                prompt,
                f"{note}\n\n本轮正式成文未完整完成。失败详情：{failure_detail}",
            )

        if answer_mode == "grounded_answer" and is_identity_role_query(prompt):
            org_names = organization_identity_names()
            explicit_role_support = [
                item
                for item in evidence
                if evidence_has_explicit_role_binding(item, prompt=prompt, names=org_names)
            ]
            if not explicit_role_support:
                structured = build_identity_guard_response(client_id, prompt, evidence, retrieval_bundle)
                answer_mode = "grounded_fallback"
                evidence_status = "partial"
                llm_invoked = False
                provider_used = None
                model_route = "证据校验"
                identity_role_insufficient = True
            else:
                try:
                    structured = state.ai.generate_chat_response(
                        prompt,
                        identity_sensitive_instruction,
                        answer_context,
                        on_partial=push_partial_analysis,
                    )
                except AiInvocationError as error:
                    answer_mode = "system_failure"
                    model_route = f"AI · {error.provider}"
                    structured = AiStructuredResponse(
                        content="庆华暂时没能完成这次回答。",
                        judgment="模型调用失败，本次回答未成功生成。",
                        analysis=f"错误信息：{error.detail}",
                        actions="建议稍后重试；如果持续失败，请检查本地后端与 AI 配置。",
                        timeline="恢复后可立即重新生成。",
                    )
                    if evidence:
                        preserved_content, preserved_structured = load_preserved_partial()
                        if preserved_content:
                            structured = build_partial_generation_fallback(
                                prompt,
                                preserved_content,
                                error.detail,
                                partial_structured=preserved_structured,
                            )
                            partial_generation_preserved = True
                        else:
                            try:
                                structured = build_compact_model_fallback(error.detail)
                                compact_model_fallback_used = True
                            except AiInvocationError:
                                structured = build_local_retrieval_fallback(
                                    client_id,
                                    prompt,
                                    evidence,
                                    retrieval_bundle,
                                    error.detail,
                                )
                                local_fallback_used = True
                        answer_mode = "grounded_fallback"
                        evidence_status = "partial"
                    else:
                        structured = AiStructuredResponse(
                            content="当前没有命中足够的原始材料，且本次通用回答阶段也未成功完成。",
                            judgment="这次请求没有整理出足够直接的原始证据，同时大模型通用回答阶段超时。",
                            analysis=f"错误信息：{error.detail}",
                            actions="建议先换一个更明确的问题重试；如果持续失败，请检查本地 AI 配置与网络状态。",
                            timeline="恢复后可立即重新生成。",
                        )
                except Exception as error:
                    llm_invoked = False
                    provider_used = None
                    answer_mode = "system_failure"
                    model_route = "AI 调用失败"
                    structured = AiStructuredResponse(
                        content="庆华暂时没能完成这次回答。",
                        judgment="模型调用失败，本次回答未成功生成。",
                        analysis=f"错误信息：{str(error)}",
                        actions="建议稍后重试；如果持续失败，请检查本地后端与 AI 配置。",
                        timeline="恢复后可立即重新生成。",
                    )
        else:
            try:
                if answer_mode == "general_answer":
                    structured = state.ai.generate_general_fallback(
                        prompt,
                        (
                            "当前没有命中足够可支撑正式分析的原始材料。"
                            "请明确把这次回答写成基于通用背景的初步判断，而不是客户资料结论；不要伪造本地背景里不存在的事实、数据、会议结论或项目状态。"
                        ),
                        subject_name=build_client_summary(client_id).name,
                    )
                elif answer_mode == "grounded_fallback":
                    structured = state.ai.generate_chat_response(
                        prompt,
                        grounded_fallback_instruction,
                        answer_context,
                        on_partial=push_partial_analysis,
                    )
                else:
                    structured = state.ai.generate_chat_response(
                        prompt,
                        grounded_system_instruction,
                        answer_context,
                        on_partial=push_partial_analysis,
                    )
            except AiInvocationError as error:
                answer_mode = "system_failure"
                model_route = f"AI · {error.provider}"
                structured = AiStructuredResponse(
                    content="庆华暂时没能完成这次回答。",
                    judgment="模型调用失败，本次回答未成功生成。",
                    analysis=f"错误信息：{error.detail}",
                    actions="建议稍后重试；如果持续失败，请检查本地后端与 AI 配置。",
                    timeline="恢复后可立即重新生成。",
                )
                if evidence:
                    preserved_content, preserved_structured = load_preserved_partial()
                    if preserved_content:
                        structured = build_partial_generation_fallback(
                            prompt,
                            preserved_content,
                            error.detail,
                            partial_structured=preserved_structured,
                        )
                        partial_generation_preserved = True
                    else:
                        try:
                            structured = build_compact_model_fallback(error.detail)
                            compact_model_fallback_used = True
                        except AiInvocationError:
                            structured = build_local_retrieval_fallback(
                                client_id,
                                prompt,
                                evidence,
                                retrieval_bundle,
                                error.detail,
                            )
                            local_fallback_used = True
                    answer_mode = "grounded_fallback"
                    evidence_status = "partial"
                else:
                    structured = AiStructuredResponse(
                        content="当前没有命中足够的原始材料，且本次通用回答阶段也未成功完成。",
                        judgment="这次请求没有整理出足够直接的原始证据，同时大模型通用回答阶段超时。",
                        analysis=f"错误信息：{error.detail}",
                        actions="建议先换一个更明确的问题重试；如果持续失败，请检查本地 AI 配置与网络状态。",
                        timeline="恢复后可立即重新生成。",
                    )
            except Exception as error:
                llm_invoked = False
                provider_used = None
                answer_mode = "system_failure"
                model_route = "AI 调用失败"
                structured = AiStructuredResponse(
                    content="庆华暂时没能完成这次回答。",
                    judgment="模型调用失败，本次回答未成功生成。",
                    analysis=f"错误信息：{str(error)}",
                    actions="建议稍后重试；如果持续失败，请检查本地后端与 AI 配置。",
                    timeline="恢复后可立即重新生成。",
                )
        if is_client_analysis_run_canceled(run_id):
            return fetch_chat_message_for_client(client_id, assistant_id)
        if answer_mode == "general_answer":
            disclaimer = "以下内容不是基于当前客户原始资料的正式分析，而是基于通用背景的初步判断。"
            if not structured.content.startswith(disclaimer):
                structured.content = f"{disclaimer}\n\n{structured.content.strip()}"
            if disclaimer not in structured.judgment:
                structured.judgment = f"{disclaimer}{structured.judgment}"
        run_long_answer = structured.content if answer_mode != "system_failure" else None
        run_structured_summary = structured if answer_mode != "system_failure" else None
        if answer_mode == "system_failure":
            run_long_answer_status = "failed"
            run_summary_status = "failed"
        elif answer_mode == "grounded_fallback":
            run_long_answer_status = "fallback"
            run_summary_status = "fallback"
        else:
            run_long_answer_status = "ready"
            run_summary_status = "ready"
        llm_elapsed_ms = round((perf_counter() - llm_started) * 1000, 2) if llm_invoked or answer_mode == "system_failure" else 0.0
        total_elapsed_ms = round((perf_counter() - request_started) * 1000, 2)
        failure_reason = None
        if answer_mode == "grounded_fallback":
            if identity_role_insufficient:
                failure_reason = "identity_role_evidence_insufficient"
            elif partial_generation_preserved:
                failure_reason = "llm_partial_preserved_after_retry"
            elif compact_model_fallback_used:
                failure_reason = "llm_compact_fallback_after_retry"
            elif local_fallback_used:
                failure_reason = "llm_local_fallback_after_retry"
            else:
                failure_reason = "partial_materials"
        elif answer_mode == "general_answer":
            failure_reason = "no_relevant_materials"
        elif answer_mode == "system_failure":
            failure_reason = "llm_failure"
        response_meta = {
            **retrieval_meta,
            "retrievalStage": retrieval_stage,
            "answerMode": answer_mode,
            "evidenceStatus": evidence_status,
            "failureReason": failure_reason,
            "workTrace": work_trace,
            "phase": "completed" if answer_mode != "system_failure" else "failed",
            "progress": 100.0,
            "progressFloor": 100.0,
            "progressCeiling": 100.0,
            "stageLabel": "回答已生成" if answer_mode != "system_failure" else "回答生成失败",
            "timing": {
                "totalMs": total_elapsed_ms,
                "retrievalMs": retrieval_elapsed_ms,
                "llmMs": llm_elapsed_ms,
            },
        }
        timestamp = now_iso()
        if run_id:
            update_client_analysis_run(
                run_id,
                status="running" if answer_mode != "system_failure" else "failed",
                phase="generating_summary" if answer_mode != "system_failure" else "failed",
                progress=90.0 if answer_mode != "system_failure" else 100.0,
                progress_floor=85.0 if answer_mode != "system_failure" else 100.0,
                progress_ceiling=100.0,
                stage_label="正在整理最终答案" if answer_mode != "system_failure" else "回答生成失败",
                elapsed_ms=total_elapsed_ms,
                evidence_summary=evidence_summary,
                long_answer=run_long_answer,
                long_answer_status=run_long_answer_status,
                summary_status="pending" if answer_mode != "system_failure" else "failed",
                answer_mode=answer_mode,
                llm_invoked=llm_invoked,
                provider_used=provider_used,
                failure_reason=failure_reason,
                timing={"totalMs": total_elapsed_ms, "retrievalMs": retrieval_elapsed_ms, "llmMs": llm_elapsed_ms},
            )
        state.db.execute(
            """
            UPDATE chat_messages
            SET content = ?, structured_data_json = ?, model_route = ?, llm_invoked = ?, provider_used = ?,
                answer_mode = ?, evidence_status = ?, failure_reason = ?, timing_json = ?, retrieval_summary_json = ?,
                evidence_json = ?, status = 'success', created_at = ?
            WHERE id = ?
            """,
            (
                structured.content,
                to_json(structured.model_dump()),
                model_route,
                1 if llm_invoked else 0,
                provider_used,
                answer_mode,
                evidence_status,
                failure_reason,
                to_json({"totalMs": total_elapsed_ms, "retrievalMs": retrieval_elapsed_ms, "llmMs": llm_elapsed_ms}),
                to_json(response_meta),
                to_json([item.model_dump() for item in evidence]),
                timestamp,
                assistant_id,
            ),
        )
        answer_run_id = new_id("ans")
        state.db.execute(
            """
            INSERT INTO answer_runs(
                id, client_id, thread_id, prompt, status, coverage_score, retrieval_mode, llm_invoked,
                provider_used, failure_reason, retrieval_json, created_at
            )
            VALUES(?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                answer_run_id,
                client_id,
                thread_id,
                prompt,
                "completed" if answer_mode in {"grounded_answer", "grounded_fallback", "general_answer"} else "failed",
                retrieval_bundle.coverage,
                "master_chunk",
                1 if llm_invoked else 0,
                provider_used,
                failure_reason,
                to_json(response_meta),
                timestamp,
            ),
        )
        for item in retrieval_bundle.citations:
            citation_document_id: str | None = item.knowledge_document_id
            if citation_document_id.startswith("v2doc_"):
                mapped_document = state.db.fetchone(
                    "SELECT document_id FROM v2_documents WHERE id = ?",
                    (citation_document_id,),
                )
                if mapped_document and mapped_document["document_id"]:
                    legacy_document = state.db.fetchone(
                        "SELECT id FROM knowledge_documents WHERE document_id = ?",
                        (str(mapped_document["document_id"]),),
                    )
                    citation_document_id = str(legacy_document["id"]) if legacy_document and legacy_document["id"] else None
                else:
                    citation_document_id = None
            if not citation_document_id:
                continue
            state.db.execute(
                """
                INSERT INTO answer_citations(
                    id, answer_run_id, knowledge_document_id, chunk_id, source_stage, drillthrough_used, title, excerpt, score, coverage_contribution, section_label, matched_terms_json, path, created_at
                )
                VALUES(?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    new_id("cit"),
                    answer_run_id,
                    citation_document_id,
                    item.chunk_id,
                    item.source_stage,
                    1 if item.drillthrough_used else 0,
                    item.title,
                    item.excerpt,
                    item.score,
                    retrieval_bundle.coverage,
                    item.section_label,
                    to_json(item.matched_terms),
                    item.path,
                    timestamp,
                ),
            )
            state.db.execute(
                "UPDATE knowledge_documents SET deep_read = 1, last_hit_question = ?, updated_at = ? WHERE id = ?",
                (prompt[:120], timestamp, citation_document_id),
        )
        state.db.execute("UPDATE chat_threads SET updated_at = ? WHERE id = ?", (timestamp, thread_id))
        if run_id:
            update_client_analysis_run(
                run_id,
                status="completed" if answer_mode != "system_failure" else "failed",
                phase="completed" if answer_mode != "system_failure" else "failed",
                progress=100.0,
                progress_floor=100.0,
                progress_ceiling=100.0,
                stage_label="分析已完成" if answer_mode != "system_failure" else "回答生成失败",
                elapsed_ms=total_elapsed_ms,
                evidence_summary=evidence_summary,
                long_answer=run_long_answer,
                structured_summary=run_structured_summary,
                long_answer_status=run_long_answer_status,
                summary_status=run_summary_status,
                answer_mode=answer_mode,
                llm_invoked=llm_invoked,
                provider_used=provider_used,
                failure_reason=failure_reason,
                timing={"totalMs": total_elapsed_ms, "retrievalMs": retrieval_elapsed_ms, "llmMs": llm_elapsed_ms},
            )
        log_activity(
            "chat.reply",
            "chat_thread",
            thread_id,
            {
                "clientId": client_id,
                "prompt": prompt,
                "coverage": retrieval_bundle.coverage,
                "citationCount": len(retrieval_bundle.citations),
                "answerMode": answer_mode,
            },
        )
        row = state.db.fetchone("SELECT * FROM chat_messages WHERE id = ?", (assistant_id,))
        assert row is not None
        return build_chat_message(row)

    def background_resolve_chat_answer(client_id: str, thread_id: str, prompt: str, assistant_id: str, search_id: str | None, request_started: float, run_id: str | None = None) -> None:
        try:
            resolve_chat_answer(client_id, thread_id, prompt, assistant_id, search_id, request_started, run_id)
        except Exception as error:
            if is_client_analysis_run_canceled(run_id):
                return
            timestamp = now_iso()
            current_row = state.db.fetchone("SELECT retrieval_summary_json FROM chat_messages WHERE id = ?", (assistant_id,))
            existing_summary = from_json(str(current_row["retrieval_summary_json"] or "{}"), {}) if current_row else {}
            if not isinstance(existing_summary, dict):
                existing_summary = {}
            existing_summary.update(
                {
                    "phase": "failed",
                    "progress": 100.0,
                    "progressFloor": 100.0,
                    "progressCeiling": 100.0,
                    "stageLabel": "回答生成失败",
                    "failureReason": str(error),
                    "lastUpdatedAt": timestamp,
                }
            )
            state.db.execute(
                """
                UPDATE chat_messages
                SET content = ?, structured_data_json = ?, model_route = ?, llm_invoked = 0, provider_used = NULL,
                    answer_mode = 'system_failure', evidence_status = 'none', failure_reason = ?, timing_json = ?,
                    retrieval_summary_json = ?, evidence_json = '[]', status = 'success', created_at = ?
                WHERE id = ?
                """,
                (
                    "庆华暂时没能完成这次回答。",
                    to_json(
                        AiStructuredResponse(
                            content="庆华暂时没能完成这次回答。",
                            judgment="模型调用失败，本次回答未成功生成。",
                            analysis=f"错误信息：{str(error)}",
                            actions="建议稍后重试；如果持续失败，请检查本地后端与 AI 配置。",
                            timeline="恢复后可立即重新生成。",
                        ).model_dump()
                    ),
                    "AI 调用失败",
                    str(error),
                    to_json({"totalMs": round((perf_counter() - request_started) * 1000, 2), "retrievalMs": 0.0, "llmMs": 0.0}),
                    to_json(existing_summary),
                    timestamp,
                    assistant_id,
                ),
            )
            if run_id:
                update_client_analysis_run(
                    run_id,
                    status="failed",
                    phase="failed",
                    progress=100.0,
                    progress_floor=100.0,
                    progress_ceiling=100.0,
                    stage_label="回答生成失败",
                    failure_reason=str(error),
                    long_answer_status="failed",
                    summary_status="failed",
                    elapsed_ms=round((perf_counter() - request_started) * 1000, 2),
                    timing={"totalMs": round((perf_counter() - request_started) * 1000, 2), "retrievalMs": 0.0, "llmMs": 0.0},
                )

    @app.post("/api/v1/clients/{client_id}/workspace/chat/start", response_model=ChatStartResponse)
    def start_chat_message(client_id: str, payload: ChatRequest) -> ChatStartResponse:
        build_client_summary(client_id)
        timestamp = now_iso()
        thread_id = ensure_chat_thread(client_id, payload.threadId, payload.prompt, timestamp)
        user_message_id = insert_user_chat_message(thread_id, payload.prompt, timestamp)
        retrieval_summary: dict[str, object] = {}
        if payload.searchId:
            bundle, _ = load_cached_retrieval_bundle(client_id, payload.searchId, payload.prompt)
            if bundle and isinstance(bundle.retrieval_summary, dict):
                retrieval_summary = {
                    **bundle.retrieval_summary,
                    "searchId": payload.searchId,
                }
        assistant_id = insert_loading_assistant_message(thread_id, retrieval_summary, timestamp)
        analysis_run = create_client_analysis_run(client_id, thread_id, user_message_id, assistant_id, payload.prompt, timestamp)
        state.db.execute("UPDATE chat_threads SET updated_at = ? WHERE id = ?", (timestamp, thread_id))
        user_row = state.db.fetchone("SELECT * FROM chat_messages WHERE id = ?", (user_message_id,))
        assistant_row = state.db.fetchone("SELECT * FROM chat_messages WHERE id = ?", (assistant_id,))
        assert user_row is not None and assistant_row is not None
        if state.chat_answer_executor is None:
            raise HTTPException(status_code=500, detail="聊天执行器不可用")
        state.chat_answer_executor.submit(
            background_resolve_chat_answer,
            client_id,
            thread_id,
            payload.prompt,
            assistant_id,
            payload.searchId,
            perf_counter(),
            analysis_run.id,
        )
        return ChatStartResponse(
            threadId=thread_id,
            userMessage=build_chat_message(user_row),
            assistantMessage=build_chat_message(assistant_row),
            analysisRun=analysis_run,
        )

    @app.get("/api/v1/clients/{client_id}/analysis-runs/{run_id}", response_model=ClientAnalysisRunRecord)
    def get_client_analysis_run(client_id: str, run_id: str) -> ClientAnalysisRunRecord:
        build_client_summary(client_id)
        return fetch_analysis_run_for_client(client_id, run_id)

    @app.post("/api/v1/clients/{client_id}/analysis-runs/{run_id}/cancel", response_model=ClientAnalysisRunRecord)
    def cancel_client_analysis_run(client_id: str, run_id: str) -> ClientAnalysisRunRecord:
        build_client_summary(client_id)
        return cancel_analysis_run_for_client(client_id, run_id)

    @app.get("/api/v1/clients/{client_id}/workspace/chat/messages/{message_id}", response_model=ChatMessageRecord)
    def get_chat_message(client_id: str, message_id: str) -> ChatMessageRecord:
        return fetch_chat_message_for_client(client_id, message_id)

    @app.get("/api/v1/clients/{client_id}/workspace/chat/threads/{thread_id}", response_model=ChatThreadDetailResponse)
    def get_chat_thread_detail(client_id: str, thread_id: str) -> ChatThreadDetailResponse:
        build_client_summary(client_id)
        return ChatThreadDetailResponse(
            thread=fetch_chat_thread_for_client(client_id, thread_id),
            messages=list_chat_messages_for_thread(client_id, thread_id),
        )

    @app.post("/api/v1/clients/{client_id}/workspace/chat", response_model=ChatMessageRecord)
    def send_chat_message(client_id: str, payload: ChatRequest) -> ChatMessageRecord:
        build_client_summary(client_id)
        timestamp = now_iso()
        thread_id = ensure_chat_thread(client_id, payload.threadId, payload.prompt, timestamp)
        insert_user_chat_message(thread_id, payload.prompt, timestamp)
        assistant_id = insert_loading_assistant_message(thread_id, {}, timestamp)
        return resolve_chat_answer(client_id, thread_id, payload.prompt, assistant_id, payload.searchId, perf_counter())

    @app.post("/api/v1/clients/{client_id}/knowledge/vectorize-answer", response_model=ClientTextDocumentResponse)
    def vectorize_answer(client_id: str, payload: VectorizeAnswerPayload) -> ClientTextDocumentResponse:
        build_client_summary(client_id)
        message = fetch_chat_message_for_client(client_id, payload.messageId)
        timestamp = now_iso()
        memory = create_memory_surrogate_from_answer(
            state.db,
            data_dir=state.data_dir,
            client_id=client_id,
            title=f"{build_client_summary(client_id).name} · 战略陪伴记忆",
            content=message.content,
            actions=message.structuredData.actions if message.structuredData else "",
            analysis=message.structuredData.analysis if message.structuredData else "",
            source_links=[
                {
                    "title": item.title,
                    "documentId": item.documentId,
                    "path": item.path,
                    "sectionLabel": item.sectionLabel,
                }
                for item in message.evidence
            ],
            created_at=timestamp,
            ai_service=state.ai,
        )
        generated = create_answer_memory_markdown_document(client_id, message)
        log_activity(
            "knowledge.vectorize_answer",
            "knowledge_memory",
            str(memory["id"]),
            {
                "clientId": client_id,
                "messageId": payload.messageId,
                "documentId": generated.documentId,
                "path": generated.path,
            },
        )
        return generated

    @app.post("/api/v1/clients/{client_id}/knowledge/export-answer", response_model=ClientTextDocumentResponse)
    def export_answer(client_id: str, payload: ExportAnswerPayload) -> ClientTextDocumentResponse:
        build_client_summary(client_id)
        message = fetch_chat_message_for_client(client_id, payload.messageId)
        exported = create_answer_export_document(client_id, message)
        log_activity(
            "knowledge.export_answer",
            "document",
            exported.documentId,
            {"clientId": client_id, "messageId": payload.messageId, "path": exported.path},
        )
        return exported

    @app.post("/api/v1/clients/{client_id}/documents/from-text", response_model=ClientTextDocumentResponse)
    def create_client_document_from_text(client_id: str, payload: ClientTextDocumentPayload) -> ClientTextDocumentResponse:
        return create_client_text_document(client_id, payload)

    @app.post("/api/v1/clients/{client_id}/documents/fill-template", response_model=ClientTemplateFillResponse)
    def fill_client_template(client_id: str, payload: ClientTemplateFillPayload) -> ClientTemplateFillResponse:
        build_client_summary(client_id)
        return fill_client_template_docx(client_id, payload.templatePath)

    @app.post("/api/v1/clients/{client_id}/documents/fill-template/start", response_model=ClientTemplateFillRunRecord)
    def start_client_template_fill(client_id: str, payload: ClientTemplateFillPayload) -> ClientTemplateFillRunRecord:
        build_client_summary(client_id)
        active_run = fetch_active_client_template_fill_run(client_id, template_path_raw=payload.templatePath)
        if active_run:
            return active_run
        other_active_run = fetch_active_client_template_fill_run(client_id)
        if other_active_run:
            raise HTTPException(
                status_code=409,
                detail=f"已有模板填写任务正在运行：{other_active_run.templateName}。请等待完成后再发起新的模板填写。",
            )
        run = create_client_template_fill_run(client_id, payload.templatePath)
        if state.template_fill_executor is None:
            raise HTTPException(status_code=503, detail="模板填写执行器不可用。")
        state.template_fill_executor.submit(run_client_template_fill, client_id, run.id, payload.templatePath)
        return fetch_client_template_fill_run(client_id, run.id)

    @app.get("/api/v1/clients/{client_id}/template-fill-runs/{run_id}", response_model=ClientTemplateFillRunRecord)
    def get_client_template_fill_run(client_id: str, run_id: str) -> ClientTemplateFillRunRecord:
        build_client_summary(client_id)
        return fetch_client_template_fill_run(client_id, run_id)

    @app.post("/api/v1/clients/{client_id}/workspace/backfill-imports", response_model=WorkspaceImportBackfillResponse)
    def backfill_client_workspace_imports(client_id: str) -> WorkspaceImportBackfillResponse:
        build_client_summary(client_id)
        summary = backfill_workspace_import(
            state.db,
            data_dir=state.data_dir,
            client_id=client_id,
        )
        log_activity("knowledge.backfill_workspace_import", "client", client_id, summary)
        return WorkspaceImportBackfillResponse(**summary)

    @app.get("/api/v1/clients/{client_id}/meetings", response_model=list[MeetingSummary])
    def list_meetings(client_id: str) -> list[MeetingSummary]:
        build_client_summary(client_id)
        return [
            build_meeting_summary(row)
            for row in state.db.fetchall("SELECT * FROM meetings WHERE client_id = ? ORDER BY updated_at DESC", (client_id,))
        ]

    @app.post("/api/v1/clients/{client_id}/meetings", response_model=MeetingPipelineResponse)
    def prepare_meeting(client_id: str, payload: MeetingCreatePayload) -> MeetingPipelineResponse:
        build_client_summary(client_id)
        meeting_id = new_id("meeting")
        timestamp = now_iso()
        state.db.execute(
            """
            INSERT INTO meetings(id, client_id, title, stage, scheduled_at, transcript_text, notes, created_at, updated_at)
            VALUES(?, ?, ?, 'prepared', ?, '', '', ?, ?)
            """,
            (meeting_id, client_id, payload.title, payload.scheduledAt, timestamp, timestamp),
        )
        agenda_source = workspace_for_client(client_id).goals[:2] or [GoalRecord(id="seed", clientId=client_id, title="明确本周推进重点", quarter="本季度", progress=0, ownerName=current_operator_row()["name"])]
        for index, goal in enumerate(agenda_source):
            state.db.execute(
                "INSERT INTO agenda_items(id, meeting_id, title, description, sort_order) VALUES(?, ?, ?, ?, ?)",
                (new_id("agenda"), meeting_id, goal.title, "会前准备议题", index),
            )
        log_activity("meeting.prepare", "meeting", meeting_id, {"title": payload.title})
        return MeetingPipelineResponse(meeting=build_meeting_detail(meeting_id), message="会议已准备，可继续入库会议原文。")

    @app.post("/api/v1/clients/{client_id}/meetings/launch-feishu", response_model=FeishuMeetingLaunchResponse)
    def launch_feishu_meeting(client_id: str, payload: FeishuMeetingLaunchPayload) -> FeishuMeetingLaunchResponse:
        prepared = prepare_meeting(client_id, MeetingCreatePayload(title=payload.title, scheduledAt=payload.scheduledAt))
        client_summary = build_client_summary(client_id)
        scheduled_label = payload.scheduledAt or "待补充"
        command_hint = f"纪要回写 {prepared.meeting.id}\\n请把会议纪要正文粘贴在第二行开始。"
        notice_text = (
            f"【会议草稿】{prepared.meeting.title}\n"
            f"客户/项目：{client_summary.name}\n"
            f"计划时间：{scheduled_label}\n"
            f"会议编号：{prepared.meeting.id}\n\n"
            f"纪要回写格式：\n{command_hint}"
        )
        delivery_status: Literal["sent", "skipped", "failed"] = "skipped"
        delivery_message = "已创建会议草稿，但尚未发送到飞书。"
        delivery_mode: Literal["bound_user", "configured_receiver", "none"] = "none"
        delivery_target: str | None = None
        receive_id_type, receive_id = None, None
        delivery_mode, receive_id_type, receive_id, delivery_target = _resolve_feishu_meeting_delivery()
        if receive_id_type and receive_id:
            try:
                _send_feishu_text_message(receive_id_type, receive_id, notice_text)
                delivery_status = "sent"
                if delivery_mode == "bound_user":
                    delivery_message = "已创建会议草稿，并按当前登录员工绑定的飞书账号发送会议通知。"
                else:
                    delivery_message = "已创建会议草稿，并把纪要回写指令发送到飞书。"
                log_activity("feishu.meeting.launch", "meeting", prepared.meeting.id, {"clientId": client_id, "sourceTaskId": payload.sourceTaskId, "deliveryMode": delivery_mode, "deliveryTarget": delivery_target})
            except FeishuApiError as exc:
                delivery_status = "failed"
                delivery_message = f"会议草稿已创建，但飞书发送失败：{exc}"
                log_activity("feishu.meeting.launch_failed", "meeting", prepared.meeting.id, {"clientId": client_id, "sourceTaskId": payload.sourceTaskId, "deliveryMode": delivery_mode, "deliveryTarget": delivery_target, "error": str(exc)})
        else:
            delivery_message = "会议草稿已创建，但当前登录员工还没有绑定飞书账号，且全局飞书接收方也未配置完整；请先完成飞书绑定或补齐 App ID、Secret 和接收方。"
        return FeishuMeetingLaunchResponse(
            meeting=prepared.meeting,
            deliveryStatus=delivery_status,
            deliveryMessage=delivery_message,
            commandHint=command_hint,
            noticeText=notice_text,
            deliveryMode=delivery_mode,
            deliveryTarget=delivery_target,
        )

    @app.get("/api/v1/clients/{client_id}/meetings/{meeting_id}", response_model=MeetingDetail)
    def get_meeting_detail(client_id: str, meeting_id: str) -> MeetingDetail:
        meeting = build_meeting_detail(meeting_id)
        if meeting.clientId != client_id:
            raise HTTPException(status_code=404, detail="Meeting not found")
        return meeting

    @app.post("/api/v1/clients/{client_id}/meetings/{meeting_id}/ingest", response_model=MeetingPipelineResponse)
    def ingest_meeting(client_id: str, meeting_id: str, payload: MeetingIngestPayload) -> MeetingPipelineResponse:
        meeting = build_meeting_detail(meeting_id)
        if meeting.clientId != client_id:
            raise HTTPException(status_code=404, detail="Meeting not found")
        state.db.execute(
            "UPDATE meetings SET transcript_text = ?, notes = ?, stage = 'ingested', updated_at = ? WHERE id = ?",
            (payload.transcriptText, payload.notes, now_iso(), meeting_id),
        )
        state.db.execute("DELETE FROM meeting_sources WHERE meeting_id = ?", (meeting_id,))
        source_text = payload.transcriptText or payload.notes
        if source_text:
            state.db.execute(
                "INSERT INTO meeting_sources(id, meeting_id, title, content_text, created_at) VALUES(?, ?, ?, ?, ?)",
                (new_id("ms"), meeting_id, "会议原文", source_text, now_iso()),
            )
        log_activity("meeting.ingest", "meeting", meeting_id, {"transcriptLength": len(payload.transcriptText), "notesLength": len(payload.notes)})
        return MeetingPipelineResponse(meeting=build_meeting_detail(meeting_id), message="会议原文已入库，可继续抽取。")

    @app.post("/api/v1/clients/{client_id}/meetings/{meeting_id}/extract", response_model=MeetingPipelineResponse)
    def extract_meeting(client_id: str, meeting_id: str) -> MeetingPipelineResponse:
        meeting = build_meeting_detail(meeting_id)
        if meeting.clientId != client_id:
            raise HTTPException(status_code=404, detail="Meeting not found")
        text = f"{meeting.transcriptText}\n{meeting.notes}".strip()
        agenda, decisions, actions, risks, ambiguities = extract_meeting_content(text)
        state.db.execute("DELETE FROM agenda_items WHERE meeting_id = ?", (meeting_id,))
        state.db.execute("DELETE FROM decisions WHERE meeting_id = ?", (meeting_id,))
        state.db.execute("DELETE FROM action_items WHERE meeting_id = ?", (meeting_id,))
        state.db.execute("DELETE FROM risks WHERE meeting_id = ?", (meeting_id,))
        state.db.execute("DELETE FROM ambiguities WHERE meeting_id = ?", (meeting_id,))
        for index, item in enumerate(agenda):
            state.db.execute("INSERT INTO agenda_items(id, meeting_id, title, description, sort_order) VALUES(?, ?, ?, ?, ?)", (new_id("agenda"), meeting_id, item[:28], "抽取后的议程点", index))
        for item in decisions:
            state.db.execute("INSERT INTO decisions(id, meeting_id, summary, created_at) VALUES(?, ?, ?, ?)", (new_id("dec"), meeting_id, item[:120], now_iso()))
        for item, owner, confidence in actions:
            state.db.execute(
                "INSERT INTO action_items(id, meeting_id, title, owner_name, due_date, confidence, publish_status, created_at) VALUES(?, ?, ?, ?, ?, ?, 'draft', ?)",
                (new_id("act"), meeting_id, item[:120], owner, "本周", confidence, now_iso()),
            )
        for item, severity in risks:
            state.db.execute("INSERT INTO risks(id, meeting_id, summary, severity, created_at) VALUES(?, ?, ?, ?, ?)", (new_id("risk"), meeting_id, item[:120], severity, now_iso()))
        for item, candidates in ambiguities:
            state.db.execute(
                "INSERT INTO ambiguities(id, meeting_id, raw_text, candidates_json, status, created_at) VALUES(?, ?, ?, ?, 'pending', ?)",
                (new_id("amb"), meeting_id, item[:120], to_json(candidates), now_iso()),
            )
        state.db.execute("UPDATE meetings SET stage = 'extracted', updated_at = ? WHERE id = ?", (now_iso(), meeting_id))
        log_activity("meeting.extract", "meeting", meeting_id, {"decisions": len(decisions), "actions": len(actions), "risks": len(risks)})
        return MeetingPipelineResponse(meeting=build_meeting_detail(meeting_id), message="结构化抽取完成，下一步可消歧。")

    @app.post("/api/v1/clients/{client_id}/meetings/{meeting_id}/resolve", response_model=MeetingPipelineResponse)
    def resolve_meeting(client_id: str, meeting_id: str) -> MeetingPipelineResponse:
        meeting = build_meeting_detail(meeting_id)
        if meeting.clientId != client_id:
            raise HTTPException(status_code=404, detail="Meeting not found")
        state.db.execute("UPDATE ambiguities SET status = 'resolved' WHERE meeting_id = ?", (meeting_id,))
        state.db.execute("UPDATE meetings SET stage = 'resolved', updated_at = ? WHERE id = ?", (now_iso(), meeting_id))
        log_activity("meeting.resolve", "meeting", meeting_id, {})
        return MeetingPipelineResponse(meeting=build_meeting_detail(meeting_id), message="低置信点已标记处理，可正式发布行动项。")

    @app.post("/api/v1/clients/{client_id}/meetings/{meeting_id}/publish", response_model=MeetingPipelineResponse)
    def publish_meeting(client_id: str, meeting_id: str) -> MeetingPipelineResponse:
        meeting = build_meeting_detail(meeting_id)
        if meeting.clientId != client_id:
            raise HTTPException(status_code=404, detail="Meeting not found")
        workspace_settings = get_client_workspace_settings()
        default_list_id = workspace_settings.meetingPublishDefaultListId or _get_local_task_settings().defaultListId or "list-0"
        for item in state.db.fetchall("SELECT * FROM action_items WHERE meeting_id = ? AND publish_status != 'published'", (meeting_id,)):
            payload = TaskPayload(
                title=str(item["title"]),
                desc="来自会议发布的行动项",
                priority=workspace_settings.meetingPublishDefaultPriority,
                listId=default_list_id,
                ddl=str(item["due_date"]),
                ownerName=str(item["owner_name"]),
                tags=["会议"],
                sourceType="meeting",
                sourceId=meeting_id,
            )
            create_task(payload, status="inbox")
            state.db.execute("UPDATE action_items SET publish_status = 'published' WHERE id = ?", (str(item["id"]),))
        for decision in state.db.fetchall("SELECT * FROM decisions WHERE meeting_id = ? ORDER BY created_at LIMIT 2", (meeting_id,)):
            state.db.execute(
                "INSERT INTO evidence_refs(id, client_id, meeting_id, document_id, title, excerpt, source_type, path, created_at) VALUES(?, ?, ?, NULL, ?, ?, 'meeting', NULL, ?)",
                (new_id("evr"), client_id, meeting_id, f"会议结论 · {meeting.title}", str(decision["summary"]), now_iso()),
            )
        state.db.execute("UPDATE meetings SET stage = 'published', updated_at = ? WHERE id = ?", (now_iso(), meeting_id))
        strategic_event_line_ids = _strategic_meeting_event_line_ids(client_id, meeting.title, meeting_id=meeting_id)
        record_meeting_publish_writeback(
            state.db,
            client_id=client_id,
            meeting_id=meeting_id,
            meeting_title=meeting.title,
            event_line_ids=strategic_event_line_ids,
        )
        growth_user_id, growth_user_name = resolve_growth_actor()
        ingest_meeting_growth_candidate(
            state.db,
            user_id=growth_user_id,
            user_name=growth_user_name,
            client_id=client_id,
            meeting=meeting,
            event_line_ids=strategic_event_line_ids,
            created_at=now_iso(),
        )
        # Write meeting activity to each related event line
        meeting_ts = now_iso()
        action_count = len([i for i in state.db.fetchall("SELECT id FROM action_items WHERE meeting_id = ? AND publish_status = 'published'", (meeting_id,))])
        decision_rows = state.db.fetchall("SELECT summary FROM decisions WHERE meeting_id = ? ORDER BY created_at LIMIT 3", (meeting_id,))
        decision_summary = "; ".join(str(d["summary"])[:60] for d in decision_rows) if decision_rows else ""
        for el_id in strategic_event_line_ids:
            state.db.execute(
                """
                INSERT INTO event_line_activities(
                    id, event_line_id, source_type, source_id, happened_at, actor_id, actor_name, title, summary, metadata_json, created_at
                ) VALUES(?, ?, 'meeting', ?, ?, NULL, ?, ?, ?, ?, ?)
                """,
                (
                    new_id("ela"),
                    el_id,
                    meeting_id,
                    meeting_ts,
                    current_operator_name(),
                    f"会议发布：{meeting.title}",
                    f"会议已发布，产生 {action_count} 条行动项。" + (f" 关键决策：{decision_summary}" if decision_summary else ""),
                    to_json({"clientId": client_id, "actionCount": action_count, "meetingTitle": meeting.title}),
                    meeting_ts,
                ),
            )
        log_activity("meeting.publish", "meeting", meeting_id, {"tasksWritten": len(meeting.actionItems)})
        return MeetingPipelineResponse(meeting=build_meeting_detail(meeting_id), message="会议已发布，行动项已写入任务收件箱。")

    @app.get("/api/v1/tasks", response_model=TaskBoardResponse)
    def list_tasks() -> TaskBoardResponse:
        if not get_cloud_token():
            return TaskBoardResponse(tasks=fetch_tasks("t.source_type != ?", (AGENT_AUTO_SOURCE_TYPE,)), lists=task_lists(), tags=task_tags())
        return cloud_task_board()

    @app.get("/api/v1/task-lists", response_model=TaskListLibraryResponse)
    def list_task_lists() -> TaskListLibraryResponse:
        if get_cloud_token():
            payload = cloud_request("GET", "/api/v1/task-lists")
            if not isinstance(payload, dict):
                raise HTTPException(status_code=502, detail="Invalid task list payload")
            return TaskListLibraryResponse(lists=[TaskListRecord(**item) for item in payload.get("lists", []) if isinstance(item, dict)])
        return TaskListLibraryResponse(lists=task_lists())

    @app.post("/api/v1/task-lists", response_model=TaskListRecord)
    def create_task_list(payload: TaskListMutationPayload) -> TaskListRecord:
        if get_cloud_token():
            response = cloud_request("POST", "/api/v1/task-lists", json_body=payload.model_dump(exclude_none=True))
            if not isinstance(response, dict):
                raise HTTPException(status_code=502, detail="Invalid task list payload")
            return TaskListRecord(**response)
        session_user = get_cached_session_user()
        if session_user and session_user.primaryRole != "admin" and (payload.scope or "org") != "personal":
            raise HTTPException(status_code=403, detail="Only admin can create public task lists")
        trimmed_name = payload.name.strip()
        if not trimmed_name:
            raise HTTPException(status_code=400, detail="清单名称不能为空")
        timestamp = now_iso()
        list_id = new_id("list")
        next_scope = payload.scope or "org"
        is_default = bool(payload.isDefault) or state.db.scalar("SELECT COUNT(1) AS count FROM task_lists WHERE scope = ?", (next_scope,)) == 0
        sort_order = payload.sortOrder if payload.sortOrder is not None else state.db.scalar("SELECT COALESCE(MAX(sort_order), -1) + 1 AS count FROM task_lists")
        if is_default:
            state.db.execute("UPDATE task_lists SET is_default = 0 WHERE scope = ?", (next_scope,))
        state.db.execute(
            """
            INSERT INTO task_lists(id, name, color, sort_order, is_default, scope, archived_at)
            VALUES(?, ?, ?, ?, ?, ?, NULL)
            """,
            (list_id, trimmed_name, payload.color.strip(), sort_order, 1 if is_default else 0, next_scope),
        )
        row = state.db.fetchone("SELECT * FROM task_lists WHERE id = ?", (list_id,))
        assert row is not None
        log_activity("task-list.create", "task_list", list_id, payload.model_dump(exclude_none=True))
        return _local_task_list_record(row)

    @app.patch("/api/v1/task-lists/{list_id}", response_model=TaskListRecord)
    def update_task_list(list_id: str, payload: TaskListMutationPayload) -> TaskListRecord:
        if get_cloud_token():
            response = cloud_request("PATCH", f"/api/v1/task-lists/{list_id}", json_body=payload.model_dump(exclude_none=True))
            if not isinstance(response, dict):
                raise HTTPException(status_code=502, detail="Invalid task list payload")
            return TaskListRecord(**response)
        session_user = get_cached_session_user()
        if session_user and session_user.primaryRole != "admin":
            row_scope = None
            row = state.db.fetchone("SELECT scope FROM task_lists WHERE id = ?", (list_id,))
            if row:
                row_scope = str(row["scope"] or "org")
            if (payload.scope or row_scope or "org") != "personal":
                raise HTTPException(status_code=403, detail="Only admin can update public task lists")
        row = state.db.fetchone("SELECT * FROM task_lists WHERE id = ?", (list_id,))
        if not row:
            raise HTTPException(status_code=404, detail="Task list not found")
        trimmed_name = payload.name.strip()
        timestamp = now_iso()
        next_archived_at = str(row["archived_at"]) if row["archived_at"] else None
        next_scope = payload.scope or str(row["scope"] or "org")
        if payload.archived is True:
            active_list_count = state.db.scalar(
                "SELECT COUNT(1) AS count FROM task_lists WHERE scope = ? AND (archived_at IS NULL OR archived_at = '')",
                (next_scope,),
            )
            if active_list_count <= 1 and not row["archived_at"]:
                raise HTTPException(status_code=400, detail="至少保留一个可用清单")
            next_archived_at = timestamp
        elif payload.archived is False:
            next_archived_at = None
        next_is_default = bool(payload.isDefault) if payload.isDefault is not None else bool(int(row["is_default"] or 0))
        if next_archived_at:
            next_is_default = False
        if next_is_default:
            state.db.execute("UPDATE task_lists SET is_default = 0 WHERE scope = ?", (next_scope,))
        state.db.execute(
            """
            UPDATE task_lists
            SET name = ?, color = ?, sort_order = ?, is_default = ?, scope = ?, archived_at = ?
            WHERE id = ?
            """,
            (
                trimmed_name,
                payload.color.strip(),
                payload.sortOrder if payload.sortOrder is not None else int(row["sort_order"] or 0),
                1 if next_is_default else 0,
                next_scope,
                next_archived_at,
                list_id,
            ),
        )
        if not next_is_default and not next_archived_at and state.db.scalar(
            "SELECT COUNT(1) AS count FROM task_lists WHERE scope = ? AND is_default = 1",
            (next_scope,),
        ) == 0:
            state.db.execute(
                "UPDATE task_lists SET is_default = 1 WHERE id = ?",
                (list_id,),
            )
        if next_archived_at and bool(int(row["is_default"] or 0)):
            fallback_row = state.db.fetchone(
                "SELECT id FROM task_lists WHERE scope = ? AND id != ? AND (archived_at IS NULL OR archived_at = '') ORDER BY sort_order ASC, name COLLATE NOCASE ASC LIMIT 1",
                (next_scope, list_id),
            )
            if fallback_row:
                state.db.execute(
                    "UPDATE task_lists SET is_default = CASE WHEN id = ? THEN 1 ELSE 0 END WHERE scope = ?",
                    (str(fallback_row["id"]), next_scope),
                )
        updated = state.db.fetchone("SELECT * FROM task_lists WHERE id = ?", (list_id,))
        assert updated is not None
        log_activity("task-list.update", "task_list", list_id, payload.model_dump(exclude_none=True))
        return _local_task_list_record(updated)

    @app.delete("/api/v1/task-lists/{list_id}")
    def delete_task_list(list_id: str) -> dict[str, bool]:
        if get_cloud_token():
            cloud_request("DELETE", f"/api/v1/task-lists/{list_id}")
            return {"deleted": True}
        session_user = get_cached_session_user()
        if session_user and session_user.primaryRole != "admin":
            row = state.db.fetchone("SELECT scope FROM task_lists WHERE id = ?", (list_id,))
            if not row or str(row["scope"] or "org") != "personal":
                raise HTTPException(status_code=403, detail="Only admin can delete public task lists")
        row = state.db.fetchone("SELECT * FROM task_lists WHERE id = ?", (list_id,))
        if not row:
            raise HTTPException(status_code=404, detail="Task list not found")
        task_count = state.db.scalar("SELECT COUNT(1) AS count FROM tasks WHERE list_id = ?", (list_id,))
        if task_count > 0:
            raise HTTPException(status_code=400, detail="该清单已有任务，请先归档，不支持直接删除")
        if state.db.scalar("SELECT COUNT(1) AS count FROM task_lists WHERE scope = ?", (str(row["scope"] or "org"),)) <= 1:
            raise HTTPException(status_code=400, detail="至少保留一个清单")
        if bool(int(row["is_default"] or 0)):
            fallback_row = state.db.fetchone(
                "SELECT id FROM task_lists WHERE scope = ? AND id != ? ORDER BY sort_order ASC, name COLLATE NOCASE ASC LIMIT 1",
                (str(row["scope"] or "org"), list_id),
            )
            if fallback_row:
                state.db.execute(
                    "UPDATE task_lists SET is_default = CASE WHEN id = ? THEN 1 ELSE 0 END WHERE scope = ?",
                    (str(fallback_row["id"]), str(row["scope"] or "org")),
                )
        state.db.execute("DELETE FROM task_lists WHERE id = ?", (list_id,))
        log_activity("task-list.delete", "task_list", list_id, {})
        return {"deleted": True}

    @app.get("/api/v1/task-tags", response_model=TaskTagLibraryResponse)
    def list_task_tags() -> TaskTagLibraryResponse:
        if get_cloud_token():
            response = cloud_request("GET", "/api/v1/task-tags")
            if isinstance(response, dict):
                return TaskTagLibraryResponse(
                    tags=[build_cloud_task_tag(item) for item in response.get("tags", []) if isinstance(item, dict)]
                )
        operator_row = current_operator_row()
        return TaskTagLibraryResponse(tags=_visible_local_task_tags(state.db, str(operator_row["id"])))

    @app.post("/api/v1/task-tags", response_model=TaskTagRecord)
    def create_task_tag(payload: TaskTagMutationPayload) -> TaskTagRecord:
        if get_cloud_token():
            response = cloud_request("POST", "/api/v1/task-tags", payload.model_dump())
            if isinstance(response, dict):
                return build_cloud_task_tag(response)
        operator_row = current_operator_row()
        return _ensure_local_tag(state.db, str(operator_row["id"]), payload.name, payload.scope, payload.color)

    @app.patch("/api/v1/task-tags/{tag_id}", response_model=TaskTagRecord)
    def update_task_tag(tag_id: str, payload: TaskTagMutationPayload) -> TaskTagRecord:
        if get_cloud_token():
            response = cloud_request("PATCH", f"/api/v1/task-tags/{tag_id}", payload.model_dump())
            if isinstance(response, dict):
                return build_cloud_task_tag(response)
        row = state.db.fetchone("SELECT * FROM task_tags WHERE id = ?", (tag_id,))
        if not row:
            raise HTTPException(status_code=404, detail="标签不存在")
        archived_at = now_iso() if payload.archived else None
        state.db.execute(
            "UPDATE task_tags SET name = ?, color = ?, scope = ?, archived_at = ?, updated_at = ? WHERE id = ?",
            (payload.name, payload.color or str(row["color"]), payload.scope, archived_at, now_iso(), tag_id),
        )
        updated = state.db.fetchone("SELECT * FROM task_tags WHERE id = ?", (tag_id,))
        assert updated is not None
        return _local_task_tag_record(updated)

    @app.delete("/api/v1/task-tags/{tag_id}")
    def delete_task_tag(tag_id: str) -> dict[str, bool]:
        if get_cloud_token():
            cloud_request("DELETE", f"/api/v1/task-tags/{tag_id}")
            return {"deleted": True}
        row = state.db.fetchone("SELECT * FROM task_tags WHERE id = ?", (tag_id,))
        if not row:
            raise HTTPException(status_code=404, detail="标签不存在")
        state.db.execute("DELETE FROM task_tags WHERE id = ?", (tag_id,))
        return {"deleted": True}

    @app.post("/api/v1/tasks/refresh-contexts", response_model=TaskContextRefreshResultRecord)
    def refresh_task_contexts() -> TaskContextRefreshResultRecord:
        task_records = cloud_task_board().tasks if get_cloud_token() else fetch_tasks()
        clients = [build_client_summary(str(row["id"])) for row in state.db.fetchall("SELECT id FROM clients ORDER BY updated_at DESC")]
        event_lines = list_event_lines()
        project_structures: dict[str, ProjectStructureResponse] = {}
        updated_tasks = 0
        unchanged_tasks = 0
        failed_tasks = 0
        client_updated_tasks = 0
        event_line_updated_tasks = 0
        module_updated_tasks = 0
        flow_updated_tasks = 0
        for task in task_records:
            payload = _build_task_scope_refresh_payload(task, clients, event_lines, project_structures)
            if not payload:
                unchanged_tasks += 1
                continue
            try:
                update_task(task.id, TaskUpdatePayload(**payload))
                updated_tasks += 1
                if "clientId" in payload:
                    client_updated_tasks += 1
                if "eventLineId" in payload:
                    event_line_updated_tasks += 1
                if "projectModuleId" in payload:
                    module_updated_tasks += 1
                if "projectFlowId" in payload:
                    flow_updated_tasks += 1
            except Exception:
                failed_tasks += 1
        unchanged_tasks = max(0, len(task_records) - updated_tasks - failed_tasks)
        backfill_memory_foundation(state.db)
        return TaskContextRefreshResultRecord(
            totalTasks=len(task_records),
            updatedTasks=updated_tasks,
            unchangedTasks=unchanged_tasks,
            failedTasks=failed_tasks,
            clientUpdatedTasks=client_updated_tasks,
            eventLineUpdatedTasks=event_line_updated_tasks,
            moduleUpdatedTasks=module_updated_tasks,
            flowUpdatedTasks=flow_updated_tasks,
            updatedAt=now_iso(),
        )

    @app.post("/api/v1/tasks/bootstrap-event-lines", response_model=TaskEventLineBootstrapResultRecord)
    def bootstrap_task_event_lines() -> TaskEventLineBootstrapResultRecord:
        task_records = cloud_task_board().tasks if get_cloud_token() else fetch_tasks()
        existing_event_lines = list_event_lines()
        event_line_by_signature = {
            f"{(item.primaryClientId or '').strip()}::{item.name.strip()}": item
            for item in existing_event_lines
            if item.name.strip()
        }
        created_event_lines = 0
        linked_tasks = 0
        skipped_tasks = 0
        failed_tasks = 0

        for task in task_records:
            if not _task_eligible_for_event_line_bootstrap(task):
                skipped_tasks += 1
                continue
            payload = _build_bootstrap_event_line_payload(task)
            signature = f"{(payload.primaryClientId or '').strip()}::{payload.name.strip()}"
            event_line = event_line_by_signature.get(signature)
            if event_line is None:
                try:
                    event_line = create_event_line(payload)
                    event_line_by_signature[signature] = event_line
                    created_event_lines += 1
                except Exception:
                    failed_tasks += 1
                    continue
            try:
                update_task(task.id, TaskUpdatePayload(eventLineId=event_line.id))
                linked_tasks += 1
            except Exception:
                failed_tasks += 1

        backfill_memory_foundation(state.db)
        return TaskEventLineBootstrapResultRecord(
            totalTasks=len(task_records),
            createdEventLines=created_event_lines,
            linkedTasks=linked_tasks,
            skippedTasks=skipped_tasks,
            failedTasks=failed_tasks,
            updatedAt=now_iso(),
        )

    @app.post("/api/v1/tasks", response_model=TaskRecord)
    def create_manual_task(payload: TaskPayload) -> TaskRecord:
        return create_task(payload)

    @app.patch("/api/v1/tasks/{task_id}", response_model=TaskRecord)
    def update_task(task_id: str, payload: TaskUpdatePayload) -> TaskRecord:
        if not get_cloud_token():
            row = state.db.fetchone("SELECT * FROM tasks WHERE id = ?", (task_id,))
            if not row:
                raise HTTPException(status_code=404, detail="Task not found")
            if payload.listId:
                list_row = state.db.fetchone("SELECT * FROM task_lists WHERE id = ?", (payload.listId,))
                if not list_row or list_row["archived_at"]:
                    raise HTTPException(status_code=400, detail="任务清单无效")
            resolved_tags = normalize_local_task_tags(
                payload.tagIds if payload.tagIds is not None else _parse_json_list(row["tag_ids_json"]),
                payload.tags if payload.tags is not None else None,
            )
            next_client_id = payload.clientId if "clientId" in payload.model_fields_set else (str(row["client_id"]) if row["client_id"] else None)
            next_event_line_id = payload.eventLineId if "eventLineId" in payload.model_fields_set else (str(row["event_line_id"]) if row["event_line_id"] else None)
            next_scope_mode = payload.scopeMode if payload.scopeMode is not None else str(row["scope_mode"] or "COLLAB_SHARED")
            if next_scope_mode == "PERSONAL_ONLY":
                next_client_id = None
                next_event_line_id = None
            client_id, event_line_id = _normalize_task_client_and_event_line_refs(
                next_client_id,
                next_event_line_id,
            )
            project_module_id = None if next_scope_mode == "PERSONAL_ONLY" else (payload.projectModuleId if "projectModuleId" in payload.model_fields_set else (str(row["project_module_id"]) if row["project_module_id"] else None))
            project_flow_id = None if next_scope_mode == "PERSONAL_ONLY" else (payload.projectFlowId if "projectFlowId" in payload.model_fields_set else (str(row["project_flow_id"]) if row["project_flow_id"] else None))
            project_module, project_flow = resolve_project_structure_refs(client_id, project_module_id, project_flow_id)
            project_context = build_task_project_context(
                client_id,
                str(row["source_type"]),
                str(row["source_id"]) if row["source_id"] else None,
                task_title=payload.title or str(row["title"]),
                task_desc=payload.desc if payload.desc is not None else str(row["description"]),
                project_module_id=project_module.id if project_module else None,
                project_flow_id=project_flow.id if project_flow else None,
            )
            event_line_context = _event_line_snapshot_context(state.db, event_line_id, None)
            attachment_count = int(state.db.scalar("SELECT COUNT(1) FROM task_attachments WHERE task_id = ?", (task_id,)) or 0)
            (
                business_category,
                current_blocker,
                next_action,
                recent_decision,
                evidence_count,
            ) = _resolve_task_action_os_fields(
                title=payload.title or str(row["title"]),
                desc=payload.desc if payload.desc is not None else str(row["description"]),
                source_type=str(row["source_type"]),
                business_category=payload.businessCategory if "businessCategory" in payload.model_fields_set else (str(row["business_category"]) if row["business_category"] else None),
                current_blocker=payload.currentBlocker if "currentBlocker" in payload.model_fields_set else (str(row["current_blocker"]) if row["current_blocker"] else None),
                next_action=payload.nextAction if "nextAction" in payload.model_fields_set else (str(row["next_action"]) if row["next_action"] else None),
                recent_decision=payload.recentDecision if "recentDecision" in payload.model_fields_set else (str(row["recent_decision"]) if row["recent_decision"] else None),
                evidence_count=payload.evidenceCount if "evidenceCount" in payload.model_fields_set else int(row["evidence_count"] or 0),
                project_context=project_context,
                event_line_context=event_line_context,
                attachment_count=attachment_count,
            )
            merged = {
                "title": payload.title or row["title"],
                "description": payload.desc if payload.desc is not None else row["description"],
                "status": payload.status or row["status"],
                "priority": payload.priority or row["priority"],
                "list_id": payload.listId or row["list_id"],
                "scope_mode": next_scope_mode,
                "client_id": client_id,
                "event_line_id": event_line_id,
                "project_module_id": project_module.id if project_module else None,
                "project_flow_id": project_flow.id if project_flow else None,
                "ddl": payload.ddl or row["ddl"],
                "due_date": payload.dueDate if payload.dueDate is not None else row["due_date"],
                "duration_minutes": payload.durationMinutes if payload.durationMinutes is not None else int(row["duration_minutes"] or 60),
                "owner_name": payload.ownerName or row["owner_name"],
                "business_category": business_category,
                "current_blocker": current_blocker,
                "next_action": next_action,
                "recent_decision": recent_decision,
                "evidence_count": evidence_count,
                "tags_json": to_json([tag.name for tag in resolved_tags]),
                "tag_ids_json": to_json([tag.id for tag in resolved_tags]),
                "updated_at": now_iso(),
            }
            state.db.execute(
                """
                UPDATE tasks
                SET title = ?, description = ?, status = ?, priority = ?, list_id = ?, scope_mode = ?, client_id = ?, event_line_id = ?, project_module_id = ?, project_flow_id = ?, ddl = ?, due_date = ?, duration_minutes = ?, owner_name = ?, business_category = ?, current_blocker = ?, next_action = ?, recent_decision = ?, evidence_count = ?, tags_json = ?, tag_ids_json = ?, updated_at = ?
                WHERE id = ?
                """,
                (
                    merged["title"],
                    merged["description"],
                    merged["status"],
                    merged["priority"],
                    merged["list_id"],
                    merged["scope_mode"],
                    merged["client_id"],
                    merged["event_line_id"],
                    merged["project_module_id"],
                    merged["project_flow_id"],
                    merged["ddl"],
                    merged["due_date"],
                    merged["duration_minutes"],
                    merged["owner_name"],
                    merged["business_category"],
                    merged["current_blocker"],
                    merged["next_action"],
                    merged["recent_decision"],
                    merged["evidence_count"],
                    merged["tags_json"],
                    merged["tag_ids_json"],
                    merged["updated_at"],
                    task_id,
                ),
            )
            _sync_task_attachment_scope(
                state.db,
                state.data_dir,
                build_task_attachment,
                build_attachment_event_line_activity,
                ensure_standard_client_folders,
                task_id,
                client_id,
                event_line_id,
                cloud=False,
            )
            if merged["event_line_id"]:
                old_status = str(row["status"])
                new_status = merged["status"]
                status_changed = old_status != new_status
                if status_changed:
                    status_labels = {
                        "done": ("任务完成", f"任务已完成：{merged['title']}"),
                        "doing": ("任务开始执行", f"任务进入执行中：{merged['title']}"),
                        "todo": ("任务已排入计划", f"任务进入待办：{merged['title']}"),
                        "rejected": ("任务已退回", f"任务被退回：{merged['title']}"),
                        "inbox": ("任务退回收件箱", f"任务退回收件箱：{merged['title']}"),
                    }
                    ela_title, ela_summary = status_labels.get(
                        new_status,
                        ("任务状态变更", f"任务状态变更为 {new_status}：{merged['title']}"),
                    )
                    state.db.execute(
                        """
                        INSERT INTO event_line_activities(
                            id, event_line_id, source_type, source_id, happened_at, actor_id, actor_name, title, summary, metadata_json
                        ) VALUES(?, ?, 'task_activity', ?, ?, NULL, ?, ?, ?, ?)
                        """,
                        (
                            new_id("ela"),
                            merged["event_line_id"],
                            task_id,
                            merged["updated_at"],
                            merged["owner_name"],
                            ela_title,
                            ela_summary,
                            to_json({"eventType": "status_change", "from": old_status, "to": new_status}),
                        ),
                    )
                else:
                    state.db.execute(
                        """
                        INSERT INTO event_line_activities(
                            id, event_line_id, source_type, source_id, happened_at, actor_id, actor_name, title, summary, metadata_json
                        ) VALUES(?, ?, 'task_activity', ?, ?, NULL, ?, ?, ?, ?)
                        """,
                        (
                            new_id("ela"),
                            merged["event_line_id"],
                            task_id,
                            merged["updated_at"],
                            merged["owner_name"],
                            "任务更新",
                            f"更新任务：{merged['title']}",
                            to_json({"eventType": "updated"}),
                        ),
                    )
            log_activity("task.update", "task", task_id, payload.model_dump(exclude_none=True))
            updated_task = fetch_tasks("t.id = ?", (task_id,))[0]
            record_task_writeback(
                state.db,
                task_id=updated_task.id,
                title=updated_task.title,
                description=updated_task.desc,
                status=updated_task.status,
                due_date=updated_task.dueDate,
                client_id=updated_task.clientId,
                event_line_id=updated_task.eventLineId,
            )
            growth_user_id, growth_user_name = resolve_growth_actor()
            ingest_task_growth_candidate(
                state.db,
                user_id=growth_user_id,
                user_name=growth_user_name,
                task=updated_task,
                source_type="task_context_candidate",
                created_at=str(merged["updated_at"]),
            )
            return updated_task
        cloud_status_map = {"todo": "todo", "doing": "doing", "done": "done", "inbox": "inbox", "rejected": "rejected"}
        response = cloud_request(
            "PATCH",
            f"/api/v1/tasks/{task_id}",
            json_body={
                "title": payload.title,
                "description": payload.desc,
                "priority": payload.priority,
                "listId": payload.listId,
                "dueDate": payload.dueDate if payload.dueDate is not None else normalize_due_date_input(payload.ddl),
                "durationMinutes": payload.durationMinutes,
                "scopeMode": payload.scopeMode,
                "clientId": payload.clientId,
                "eventLineId": payload.eventLineId,
                "projectModuleId": payload.projectModuleId,
                "projectFlowId": payload.projectFlowId,
                "progressStatus": cloud_status_map.get(payload.status) if payload.status else None,
                "collaboratorIds": payload.collaboratorIds,
                "ownerId": payload.ownerId,
                "businessCategory": payload.businessCategory,
                "currentBlocker": payload.currentBlocker,
                "nextAction": payload.nextAction,
                "recentDecision": payload.recentDecision,
                "evidenceCount": payload.evidenceCount,
            },
        )
        if not isinstance(response, dict):
            raise HTTPException(status_code=502, detail="Invalid cloud task payload")
        log_activity("task.update", "task", task_id, payload.model_dump(exclude_none=True))
        updated_task = build_cloud_task(response, {})
        growth_user_id, growth_user_name = resolve_growth_actor()
        ingest_task_growth_candidate(
            state.db,
            user_id=growth_user_id,
            user_name=growth_user_name,
            task=updated_task,
            source_type="task_context_candidate",
            created_at=now_iso(),
        )
        return updated_task

    @app.delete("/api/v1/tasks/{task_id}")
    def delete_task(task_id: str) -> dict[str, bool]:
        if not get_cloud_token():
            row = state.db.fetchone("SELECT * FROM tasks WHERE id = ?", (task_id,))
            if not row:
                raise HTTPException(status_code=404, detail="Task not found")
            task_title = str(row["title"] or "任务")
            event_line_id = str(row["event_line_id"]) if row["event_line_id"] else None
            client_id = str(row["client_id"]) if row["client_id"] else None
            state.db.execute("DELETE FROM activity_logs WHERE entity_type = 'task' AND entity_id = ?", (task_id,))
            state.db.execute("DELETE FROM event_line_activities WHERE source_type = 'task_activity' AND source_id = ?", (task_id,))
            state.db.execute("DELETE FROM memory_facts WHERE scope_type = 'task' AND scope_id = ?", (task_id,))
            state.db.execute(
                """
                DELETE FROM memory_facts
                WHERE source_type = 'task'
                  AND source_id = ?
                """,
                (task_id,),
            )
            state.db.execute("DELETE FROM growth_signal_events WHERE task_id = ?", (task_id,))
            state.db.execute("DELETE FROM growth_evidence_records WHERE task_id = ?", (task_id,))
            state.db.execute("DELETE FROM tasks WHERE id = ?", (task_id,))
            log_activity("task.delete", "task", task_id, {"title": task_title, "eventLineId": event_line_id, "clientId": client_id})
            if event_line_id:
                refresh_event_line_memory_snapshot(state.db, event_line_id)
            if client_id:
                refresh_organization_notebook_snapshot(state.db, client_id)
            return {"deleted": True}
        cloud_request("DELETE", f"/api/v1/tasks/{task_id}")
        log_activity("task.delete", "task", task_id, {})
        return {"deleted": True}

    @app.post("/api/v1/tasks/{task_id}/attachments", response_model=TaskRecord)
    def upload_task_attachment(
        task_id: str,
        file: UploadFile = File(...),
        clientId: str | None = Form(default=None),
        eventLineId: str | None = Form(default=None),
        taskTitle: str | None = Form(default=None),
    ) -> TaskRecord:
        local_task_row = state.db.fetchone("SELECT * FROM tasks WHERE id = ?", (task_id,))
        is_cloud_task = local_task_row is None
        if is_cloud_task and not get_cloud_token():
            raise HTTPException(status_code=404, detail="Task not found")

        resolved_client_id = (
            str(local_task_row["client_id"]) if local_task_row and local_task_row["client_id"] else clientId
        )
        if not resolved_client_id:
            raise HTTPException(status_code=400, detail="请先关联客户/项目后再上传附件。")
        build_client_summary(resolved_client_id)
        ensure_standard_client_folders(resolved_client_id)

        resolved_event_line_id = (
            str(local_task_row["event_line_id"]) if local_task_row and local_task_row["event_line_id"] else eventLineId
        )
        resolved_task_title = (
            str(local_task_row["title"]) if local_task_row and local_task_row["title"] else (taskTitle or file.filename or "任务附件")
        )
        upload_name = safe_filename(file.filename or f"{resolved_task_title}-附件")
        content = file.file.read()
        if not content:
            raise HTTPException(status_code=400, detail="上传失败：附件内容为空。")
        staged_path = stage_task_attachment_upload(resolved_client_id, upload_name, content)
        timestamp = now_iso()
        folder_row = state.db.fetchone(
            "SELECT id FROM client_folders WHERE client_id = ? AND label = ?",
            (resolved_client_id, "项目与业务"),
        )
        document_id = new_id("doc")
        excerpt_seed = ""
        if staged_path.suffix.lower() in {".md", ".txt", ".csv", ".json", ".yaml", ".yml", ".py", ".ts", ".tsx", ".js", ".jsx"}:
            excerpt_seed = content.decode("utf-8", errors="ignore")
        fallback_excerpt = (
            excerpt_seed.strip()[:140] or f"{upload_name} 已作为任务附件进入项目资料库，可用于后续检索、问答与事件线证据引用。"
        ) if excerpt_seed else f"{upload_name} 已作为任务附件进入项目资料库，可用于后续检索、问答与事件线证据引用。"
        state.db.execute(
            """
            INSERT INTO documents(id, client_id, folder_id, title, path, original_source_path, kind, source, excerpt, tags_json, created_at)
            VALUES(?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                document_id,
                resolved_client_id,
                str(folder_row["id"]) if folder_row else None,
                upload_name,
                str(staged_path),
                str(staged_path),
                staged_path.suffix.lower().lstrip(".") or "bin",
                "task_attachment",
                fallback_excerpt,
                to_json(["task_attachment", staged_path.suffix.lower().lstrip(".") or "bin"]),
                timestamp,
            ),
        )
        ingest_document_knowledge(
            state.db,
            data_dir=state.data_dir,
            client_id=resolved_client_id,
            import_id=None,
            document_id=document_id,
            source_path=staged_path,
            original_source_path=staged_path,
            title=upload_name,
            kind=staged_path.suffix.lower().lstrip(".") or "bin",
            source="task_attachment",
            fallback_excerpt=fallback_excerpt,
            created_at=timestamp,
            ai_service=None,
        )
        document_row = state.db.fetchone("SELECT * FROM documents WHERE id = ?", (document_id,))
        if not document_row:
            raise HTTPException(status_code=500, detail="附件归档失败。")
        attachment_id = new_id("tatt")
        table_name = "task_attachments_cloud" if is_cloud_task else "task_attachments"
        state.db.execute(
            f"""
            INSERT INTO {table_name}(id, task_id, client_id, event_line_id, document_id, title, path, kind, source, size_bytes, created_at)
            VALUES(?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                attachment_id,
                task_id,
                resolved_client_id,
                resolved_event_line_id,
                document_id,
                str(document_row["title"]),
                str(document_row["path"]),
                str(document_row["kind"]),
                str(document_row["source"]),
                len(content),
                timestamp,
            ),
        )
        state.db.execute(
            """
            INSERT INTO evidence_refs(id, client_id, meeting_id, document_id, title, excerpt, source_type, path, created_at)
            VALUES(?, ?, NULL, ?, ?, ?, ?, ?, ?)
            """,
            (
                new_id("ev"),
                resolved_client_id,
                document_id,
                str(document_row["title"]),
                str(document_row["excerpt"]),
                "task_attachment",
                str(document_row["path"]),
                timestamp,
            ),
        )
        if not is_cloud_task and resolved_event_line_id:
            state.db.execute(
                """
                INSERT INTO event_line_activities(
                    id, event_line_id, source_type, source_id, happened_at, actor_id, actor_name, title, summary, metadata_json, created_at
                ) VALUES(?, ?, 'attachment', ?, ?, NULL, ?, ?, ?, ?, ?)
                """,
                (
                    new_id("ela"),
                    resolved_event_line_id,
                    attachment_id,
                    timestamp,
                    current_operator_name(),
                    f"上传附件：{document_row['title']}",
                    f"任务附件已进入项目资料库：{document_row['title']}",
                    to_json(
                        {
                            "taskId": task_id,
                            "documentId": document_id,
                            "clientId": resolved_client_id,
                            "path": str(document_row["path"]),
                        }
                    ),
                    timestamp,
                ),
            )
        log_activity(
            "task.attachment.upload",
            "task",
            task_id,
            {
                "attachmentId": attachment_id,
                "documentId": document_id,
                "clientId": resolved_client_id,
                "eventLineId": resolved_event_line_id,
                "title": str(document_row["title"]),
            },
        )
        record_task_attachment_writeback(
            state.db,
            task_id=task_id,
            client_id=resolved_client_id,
            event_line_id=resolved_event_line_id,
            attachment_title=str(document_row["title"]),
            attachment_path=str(document_row["path"]),
        )
        if resolved_event_line_id and get_cloud_token():
            import threading

            def _bg_upload():
                try:
                    cloud_upload_file(
                        f"/api/v1/tasks/{task_id}/attachments",
                        file_name=upload_name,
                        file_content=content,
                        content_type=file.content_type or "application/octet-stream",
                        form_fields={
                            "clientId": resolved_client_id or "",
                            "eventLineId": resolved_event_line_id,
                            "title": str(document_row["title"]),
                            "taskTitle": resolved_task_title,
                        },
                    )
                except Exception:
                    pass  # 云端上传失败不阻断本地流程

            threading.Thread(target=_bg_upload, daemon=True).start()
        task_after_upload = fetch_tasks("t.id = ?", (task_id,))[0] if not is_cloud_task else fetch_cloud_task_by_id(task_id)
        growth_user_id, growth_user_name = resolve_growth_actor()
        ingest_task_growth_candidate(
            state.db,
            user_id=growth_user_id,
            user_name=growth_user_name,
            task=task_after_upload,
            source_type="task_attachment_candidate",
            created_at=timestamp,
        )
        return task_after_upload

    @app.post("/api/v1/tasks/{task_id}/confirm", response_model=TaskRecord)
    def confirm_task(task_id: str) -> TaskRecord:
        if not get_cloud_token():
            state.db.execute("UPDATE tasks SET status = 'doing', updated_at = ? WHERE id = ?", (now_iso(), task_id))
            log_activity("task.confirm", "task", task_id, {})
            return fetch_tasks("t.id = ?", (task_id,))[0]
        user = require_session_user()
        response = cloud_request("POST", f"/api/v1/tasks/{task_id}/collaborators/{user.id}/accept")
        if not isinstance(response, dict):
            raise HTTPException(status_code=502, detail="Invalid cloud task payload")
        log_activity("task.confirm", "task", task_id, {"userId": user.id})
        return build_cloud_task(response, {})

    @app.post("/api/v1/tasks/{task_id}/reject", response_model=TaskRecord)
    def reject_task(task_id: str, payload: TaskRejectPayload) -> TaskRecord:
        if not get_cloud_token():
            state.db.execute("UPDATE tasks SET status = 'rejected', updated_at = ? WHERE id = ?", (now_iso(), task_id))
            upsert_task_note(task_id, payload.reason)
            log_activity("task.reject", "task", task_id, {"reason": payload.reason})
            return fetch_tasks("t.id = ?", (task_id,))[0]
        user = require_session_user()
        response = cloud_request(
            "POST",
            f"/api/v1/tasks/{task_id}/collaborators/{user.id}/return",
            json_body={"reason": payload.reason},
        )
        upsert_task_note(task_id, payload.reason)
        log_activity("task.reject", "task", task_id, {"reason": payload.reason, "userId": user.id})
        if not isinstance(response, dict):
            raise HTTPException(status_code=502, detail="Invalid cloud task payload")
        return build_cloud_task(response, {})

    @app.post("/api/v1/tasks/{task_id}/complete-with-review", response_model=TaskRecord)
    def complete_task_with_review(task_id: str, payload: TaskCompletionReviewPayload) -> TaskRecord:
        if not get_cloud_token():
            raise HTTPException(status_code=400, detail="当前环境未启用组织复核链")
        response = cloud_request("POST", f"/api/v1/tasks/{task_id}/complete-with-review", payload.model_dump())
        if not isinstance(response, dict):
            raise HTTPException(status_code=502, detail="Invalid cloud task payload")
        log_activity("task.complete-with-review", "task", task_id, {"reviewNote": payload.reviewNote[:60]})
        return build_cloud_task(response, {})

    @app.post("/api/v1/tasks/{task_id}/review/approve", response_model=TaskRecord)
    def approve_task_review(task_id: str) -> TaskRecord:
        if not get_cloud_token():
            raise HTTPException(status_code=400, detail="当前环境未启用组织复核链")
        response = cloud_request("POST", f"/api/v1/tasks/{task_id}/review/approve")
        if not isinstance(response, dict):
            raise HTTPException(status_code=502, detail="Invalid cloud task payload")
        log_activity("task.review.approve", "task", task_id, {})
        return build_cloud_task(response, {})

    @app.post("/api/v1/tasks/{task_id}/review/return", response_model=TaskRecord)
    def return_task_review(task_id: str, payload: TaskRejectPayload) -> TaskRecord:
        if not get_cloud_token():
            raise HTTPException(status_code=400, detail="当前环境未启用组织复核链")
        response = cloud_request(
            "POST",
            f"/api/v1/tasks/{task_id}/review/return",
            json_body={"reason": payload.reason},
        )
        if not isinstance(response, dict):
            raise HTTPException(status_code=502, detail="Invalid cloud task payload")
        log_activity("task.review.return", "task", task_id, {"reason": payload.reason})
        return build_cloud_task(response, {})

    @app.post("/api/v1/tasks/{task_id}/note", response_model=TaskRecord)
    def save_task_note(task_id: str, payload: TaskNotePayload) -> TaskRecord:
        upsert_task_note(task_id, payload.note)
        log_activity("task.note", "task", task_id, {"noteLength": len(payload.note)})
        if not get_cloud_token():
            state.db.execute("UPDATE tasks SET updated_at = ? WHERE id = ?", (now_iso(), task_id))
            return fetch_tasks("t.id = ?", (task_id,))[0]
        try:
            cloud_request("POST", f"/api/v1/tasks/{task_id}/note", {"note": payload.note})
        except HTTPException:
            pass  # 云端保存失败时保留本地备注，不阻断用户操作
        board = cloud_task_board()
        task = next((item for item in board.tasks if item.id == task_id), None)
        if not task:
            raise HTTPException(status_code=404, detail="Task not found")
        return task

    @app.get("/api/v1/tasks/{task_id}/activity", response_model=list[TaskActivityRecord])
    def get_task_activity(task_id: str) -> list[TaskActivityRecord]:
        if not get_cloud_token():
            if task_id.startswith("agent_task_"):
                ensure_admin_for_sensitive_settings()
                return build_agent_execution_task_activity(
                    db=state.db,
                    task_id=task_id,
                    thread_sync_path=THREAD_SYNC_DOC_PATH,
                )
            return [
                TaskActivityRecord(
                    id=str(row["id"]),
                    taskId=task_id,
                    actorId=str(row["actor_name"] or "local"),
                    actorName=str(row["actor_name"] or "本地用户"),
                    eventType=str(row["action"]),
                    payload=from_json(row["detail_json"], {}) if row["detail_json"] else {},
                    createdAt=str(row["created_at"]),
                )
                for row in state.db.fetchall(
                    """
                    SELECT *
                    FROM activity_logs
                    WHERE entity_type = 'task' AND entity_id = ?
                    ORDER BY created_at DESC
                    LIMIT 50
                    """,
                    (task_id,),
                )
            ]
        payload = cloud_request("GET", f"/api/v1/tasks/{task_id}/activity")
        if not isinstance(payload, list):
            raise HTTPException(status_code=502, detail="Invalid task activity payload")
        return [TaskActivityRecord(**item) for item in payload if isinstance(item, dict)]

    @app.get("/api/v1/tasks/agent-execution", response_model=list[TaskRecord])
    def list_agent_execution_tasks(week: str | None = None, department: str | None = None) -> list[TaskRecord]:
        session_user = get_cached_session_user()
        if get_cloud_token() and session_user is None:
            session_user = require_session_user()
        governance = _review_governance_with_members()
        allowed_department = _resolve_agent_execution_department_scope(
            session_user=session_user,
            governance=governance,
            requested_department=department,
        )
        target_week = week or current_review_week_label()
        sync_agent_execution_tasks(
            db=state.db,
            week_label=target_week,
            thread_sync_path=THREAD_SYNC_DOC_PATH,
        )
        tasks = build_agent_execution_tasks(
            db=state.db,
            week_label=target_week,
            thread_sync_path=THREAD_SYNC_DOC_PATH,
        )
        if not allowed_department:
            return tasks
        normalized_allowed_department = _normalize_department_name(allowed_department)
        filtered: list[TaskRecord] = []
        for task in tasks:
            if any(_normalize_department_name(tag.name) == normalized_allowed_department for tag in task.tags):
                filtered.append(task)
        return filtered

    @app.post("/api/v1/local/tasks/tag-suggestions", response_model=AiTagSuggestionResponse)
    def suggest_task_tags(payload: AiTagSuggestionPayload) -> AiTagSuggestionResponse:
        _ = payload
        return AiTagSuggestionResponse(suggestedTags=[])

    @app.get("/api/v1/reviews", response_model=ReviewResponse)
    def list_reviews(weekLabel: str | None = Query(default=None)) -> ReviewResponse:
        if get_cloud_token():
            suffix = f"?weekLabel={quote(weekLabel)}" if weekLabel else ""
            payload = cloud_request("GET", f"/api/v1/reviews/dashboard{suffix}")
            if not isinstance(payload, dict):
                raise HTTPException(status_code=502, detail="Invalid review payload")
            return augment_review_response(ReviewResponse(**payload), weekLabel)
        return local_review_dashboard(weekLabel)

    @app.get("/api/v1/reviews/history", response_model=ReviewHistoryResponse)
    def list_review_history() -> ReviewHistoryResponse:
        if get_cloud_token():
            payload = cloud_request("GET", "/api/v1/reviews/history")
            if not isinstance(payload, dict):
                raise HTTPException(status_code=502, detail="Invalid review history payload")
            return ReviewHistoryResponse(**payload)
        return local_review_history()

    @app.post("/api/v1/reviews/weekly", response_model=ReviewResponse)
    def create_weekly_review(payload: WeeklyReviewPayload) -> ReviewResponse:
        if get_cloud_token():
            response_payload = cloud_request("POST", "/api/v1/reviews/weekly", json_body=payload.model_dump())
            if not isinstance(response_payload, dict):
                raise HTTPException(status_code=502, detail="Invalid review payload")
            log_activity(
                "review.create",
                "weekly_review",
                str(response_payload.get("currentReview", {}).get("id", "review")),
                {"weekLabel": payload.weekLabel, "personalExcludedFromAggregation": True},
            )
            response = augment_review_response(ReviewResponse(**response_payload), payload.weekLabel)
            if response.currentReview:
                user_id, user_name = resolve_growth_actor()
                ingest_review_growth(
                    state.db,
                    user_id=user_id,
                    user_name=user_name,
                    review=response.currentReview,
                    task_entries=[*response.workItems, *response.personalItems],
                )
            return response
        created_at = now_iso()
        operator_id = str(current_operator_row()["id"])
        existing = local_review_row_for_week(payload.weekLabel)
        if existing:
            review_id = str(existing["id"])
            state.db.execute(
                """
                UPDATE weekly_reviews
                SET work_free_note = ?, personal_growth_note = ?, personal_private_note = ?, updated_at = ?
                WHERE id = ?
                """,
                (
                    payload.workFreeNote.strip(),
                    payload.personalGrowthNote.strip(),
                    payload.personalPrivateNote.strip(),
                    created_at,
                    review_id,
                ),
            )
        else:
            review_id = new_id("review")
            state.db.execute(
                """
                INSERT INTO weekly_reviews(
                    id, week_label, operator_id, summary, work_free_note, personal_growth_note, personal_private_note, created_at, updated_at
                ) VALUES(?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    review_id,
                    payload.weekLabel,
                    operator_id,
                    "",
                    payload.workFreeNote.strip(),
                    payload.personalGrowthNote.strip(),
                    payload.personalPrivateNote.strip(),
                    created_at,
                    created_at,
                ),
            )
        reviewed_work_items: list[WeeklyReviewTaskEntryRecord] = []
        for entry in payload.taskEntries:
            task_id = str(entry.get("taskId", "")).strip()
            if not task_id:
                continue
            task_items = fetch_tasks("t.id = ?", (task_id,))
            if not task_items:
                continue
            task_row = task_items[0]
            structured_note = coerce_review_structured_note(entry.get("structuredNote"))
            note = compose_review_note(structured_note, str(entry.get("note", "")).strip())
            existing_entry = state.db.fetchone(
                "SELECT id FROM weekly_review_task_entries WHERE review_id = ? AND task_id = ?",
                (review_id, task_id),
            )
            if not note:
                if existing_entry:
                    state.db.execute("DELETE FROM weekly_review_task_entries WHERE id = ?", (str(existing_entry["id"]),))
                continue
            content_domain = str(entry.get("contentDomain") or ("personal" if is_private_task(task_row) else "work"))
            snapshot = _task_snapshot_from_task(task_row, state.db)
            if existing_entry:
                state.db.execute(
                    """
                    UPDATE weekly_review_task_entries
                    SET content_domain = ?, note = ?, structured_note_json = ?, reviewed_at = ?, task_snapshot_json = ?, updated_at = ?
                    WHERE id = ?
                    """,
                    (content_domain, note, to_json(structured_note.model_dump()), created_at, to_json(snapshot), created_at, str(existing_entry["id"])),
                )
            else:
                state.db.execute(
                    """
                    INSERT INTO weekly_review_task_entries(
                        id, review_id, task_id, week_label, content_domain, note, structured_note_json, reviewed_at, task_snapshot_json, created_at, updated_at
                    ) VALUES(?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """,
                    (
                        new_id("review_item"),
                        review_id,
                        task_id,
                        payload.weekLabel,
                        content_domain,
                        note,
                        to_json(structured_note.model_dump()),
                        created_at,
                        to_json(snapshot),
                        created_at,
                        created_at,
                    ),
                )
            if content_domain == "work":
                reviewed_work_items.append(
                    _review_entry_from_task(
                        task=task_row,
                        week_label=payload.weekLabel,
                        content_domain=content_domain,
                        review_id=review_id,
                        note=note,
                        structured_note=structured_note,
                        reviewed_at=created_at,
                        snapshot=snapshot,
                        db=state.db,
                    )
                )
        state.db.execute(
            """
            UPDATE weekly_reviews
            SET summary = ?, work_free_note = ?, personal_growth_note = ?, personal_private_note = ?, updated_at = ?
            WHERE id = ?
            """,
            (
                summarize_local_review_notes(reviewed_work_items),
                payload.workFreeNote.strip(),
                payload.personalGrowthNote.strip(),
                payload.personalPrivateNote.strip(),
                created_at,
                review_id,
            ),
        )
        log_activity("review.create", "weekly_review", review_id, {"weekLabel": payload.weekLabel})
        record_weekly_review_writeback(state.db, review_id=review_id)
        response = local_review_dashboard(payload.weekLabel)
        if response.currentReview:
            _, user_name = resolve_growth_actor()
            ingest_review_growth(
                state.db,
                user_id=operator_id,
                user_name=user_name,
                review=response.currentReview,
                task_entries=[*response.workItems, *response.personalItems],
                created_at=created_at,
            )
            response = local_review_dashboard(payload.weekLabel)
        return response

    def build_topic_candidate(row) -> TopicCandidateRecord:
        return TopicCandidateRecord(
            id=str(row["id"]),
            radarId=str(row["radar_id"]),
            title=str(row["title"]),
            summary=str(row["summary"]),
            source=str(row["source"]),
            sourceUrl=str(row["source_url"]) if row["source_url"] else None,
            publishedAt=str(row["published_at"]) if row["published_at"] else None,
            captureMethod=str(row["capture_method"]) if row["capture_method"] else "manual",
            capturedBy=str(row["captured_by"]) if row["captured_by"] else None,
            status=str(row["status"]),  # type: ignore[arg-type]
            insightStatus=str(row["insight_status"] or "pending"),  # type: ignore[arg-type]
            insightUpdatedAt=str(row["insight_updated_at"]) if row["insight_updated_at"] else None,
            createdAt=str(row["created_at"]),
        )

    def build_topic_candidate_insight(row) -> TopicCandidateInsightRecord:
        return TopicCandidateInsightRecord(
            candidateId=str(row["candidate_id"]),
            overview=str(row["overview"]),
            keyPoints=_parse_json_list(row["key_points_json"]),
            recommendationReasons=_parse_json_list(row["recommendation_reasons_json"]),
            practicalUses=_parse_json_list(row["practical_uses_json"]),
            editorialNote=str(row["editorial_note"] or ""),
            discussionPrompts=_parse_json_list(row["discussion_prompts_json"]),
            createdAt=str(row["created_at"]),
            updatedAt=str(row["updated_at"]),
        )

    def _fallback_topic_source_label(url: str) -> str:
        parsed = urlparse(url)
        host = parsed.netloc.lower().replace("www.", "")
        primary = host.split(".")[0] if host else ""
        primary = re.sub(r"[^a-zA-Z0-9\u4e00-\u9fff]+", "", primary)
        if 2 <= len(primary) <= 6:
            return primary
        if primary:
            return primary[:6]
        return "优先站点"

    def normalize_topic_radar_source_url(value: str) -> str:
        raw = value.strip()
        if not raw:
            raise HTTPException(status_code=400, detail="网址不能为空")
        candidate = raw if "://" in raw else f"https://{raw}"
        parsed = urlparse(candidate)
        if not parsed.netloc and parsed.path and "." in parsed.path:
            candidate = f"https://{parsed.path}"
            parsed = urlparse(candidate)
        if parsed.scheme not in {"http", "https"} or not parsed.netloc:
            raise HTTPException(status_code=400, detail="网址格式无效")
        normalized_path = parsed.path.rstrip("/")
        return urlunparse((parsed.scheme, parsed.netloc.lower(), normalized_path, "", "", ""))

    def suggest_topic_radar_source_label(url: str) -> str:
        parsed = urlparse(url)
        host = parsed.netloc.lower().replace("www.", "")
        path = parsed.path.strip("/")
        descriptor = f"资讯来源网址：{host}{('/' + path) if path else ''}，请提炼成 2 到 6 个字的中文来源标签。"
        label = state.ai.suggest_short_title(descriptor).strip()
        if label:
            return label[:10]
        return _fallback_topic_source_label(url)

    def normalize_topic_radar_preferred_sources(items: list[TopicRadarPreferredSourceRecord] | None) -> list[TopicRadarPreferredSourceRecord]:
        normalized: list[TopicRadarPreferredSourceRecord] = []
        seen_urls: set[str] = set()
        for item in items or []:
            url = normalize_topic_radar_source_url(item.url)
            if url in seen_urls:
                continue
            seen_urls.add(url)
            label = item.label.strip() if item.label.strip() else suggest_topic_radar_source_label(url)
            normalized.append(TopicRadarPreferredSourceRecord(url=url, label=label[:10] or _fallback_topic_source_label(url)))
        return normalized

    def parse_topic_radar_preferred_sources(raw_value: object) -> list[TopicRadarPreferredSourceRecord]:
        parsed = from_json(raw_value, [])
        if not isinstance(parsed, list):
            return []
        normalized: list[TopicRadarPreferredSourceRecord] = []
        for item in parsed:
            if not isinstance(item, dict):
                continue
            url = str(item.get("url") or "").strip()
            label = str(item.get("label") or "").strip()
            if not url:
                continue
            normalized.append(TopicRadarPreferredSourceRecord(url=url, label=label or _fallback_topic_source_label(url)))
        return normalized

    def build_topic_radar(row) -> TopicRadarRecord:
        return TopicRadarRecord(
            id=str(row["id"]),
            title=str(row["title"]),
            prompt=str(row["prompt"]),
            timeRange=str(row["time_range"]),
            preferredSources=parse_topic_radar_preferred_sources(row["preferred_sources_json"]),
            createdAt=str(row["created_at"]),
        )

    def fetch_topic_candidates() -> list[TopicCandidateRecord]:
        return [
            build_topic_candidate(row)
            for row in state.db.fetchall("SELECT * FROM topic_candidates ORDER BY created_at DESC")
        ]

    def save_topic_candidate_insight(
        *,
        candidate_id: str,
        overview: str,
        key_points: list[str],
        recommendation_reasons: list[str],
        practical_uses: list[str],
        editorial_note: str,
        discussion_prompts: list[str],
        source_excerpt: str,
    ) -> TopicCandidateInsightRecord:
        existing = state.db.fetchone("SELECT * FROM topic_candidate_insights WHERE candidate_id = ?", (candidate_id,))
        timestamp = now_iso()
        if existing:
            state.db.execute(
                """
                UPDATE topic_candidate_insights
                SET overview = ?, key_points_json = ?, recommendation_reasons_json = ?, practical_uses_json = ?, editorial_note = ?, discussion_prompts_json = ?, source_excerpt = ?, updated_at = ?
                WHERE candidate_id = ?
                """,
                (
                    overview,
                    to_json(key_points),
                    to_json(recommendation_reasons),
                    to_json(practical_uses),
                    editorial_note,
                    to_json(discussion_prompts),
                    source_excerpt,
                    timestamp,
                    candidate_id,
                ),
            )
        else:
            state.db.execute(
                """
                INSERT INTO topic_candidate_insights(
                    id, candidate_id, overview, key_points_json, recommendation_reasons_json, practical_uses_json, editorial_note, discussion_prompts_json, source_excerpt, created_at, updated_at
                ) VALUES(?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    new_id("insight"),
                    candidate_id,
                    overview,
                    to_json(key_points),
                    to_json(recommendation_reasons),
                    to_json(practical_uses),
                    editorial_note,
                    to_json(discussion_prompts),
                    source_excerpt,
                    timestamp,
                    timestamp,
                ),
            )
        row = state.db.fetchone("SELECT * FROM topic_candidate_insights WHERE candidate_id = ?", (candidate_id,))
        assert row is not None
        return build_topic_candidate_insight(row)

    def update_topic_candidate_insight_state(candidate_id: str, status: str, *, error: str | None = None) -> None:
        state.db.execute(
            """
            UPDATE topic_candidates
            SET insight_status = ?, insight_updated_at = ?, insight_error = ?
            WHERE id = ?
            """,
            (status, now_iso(), error or None, candidate_id),
        )

    def topic_candidate_insight_needs_refresh(candidate_row, cached_row) -> bool:
        overview = str(cached_row["overview"] or "").strip()
        key_points = _parse_json_list(cached_row["key_points_json"])
        reasons = _parse_json_list(cached_row["recommendation_reasons_json"])
        uses = _parse_json_list(cached_row["practical_uses_json"])
        editorial_note = str(cached_row["editorial_note"] or "").strip()
        discussion_prompts = _parse_json_list(cached_row["discussion_prompts_json"])
        source_excerpt = str(cached_row["source_excerpt"] or "").strip()

        if not overview or not key_points or not reasons or not uses or not editorial_note or not discussion_prompts:
            return True
        if candidate_row["source_url"] and not source_excerpt:
            return True
        if not state.ai._has_sufficient_cjk(overview):
            return True
        if not state.ai._has_sufficient_cjk(editorial_note):
            return True
        if len(overview) < 120:
            return True
        if len(editorial_note) < 120:
            return True
        if state.ai._looks_generic_topic_overview(overview):
            return True
        if state.ai._looks_stale_topic_editorial_note(editorial_note):
            return True
        if any(not state.ai._has_sufficient_cjk(item) for item in key_points):
            return True
        if any(not state.ai._has_sufficient_cjk(item) for item in discussion_prompts):
            return True
        return False

    def ensure_topic_candidate_insight(candidate_row) -> tuple[TopicCandidateInsightRecord, str]:
        topics_settings = get_topics_settings()
        cached = state.db.fetchone("SELECT * FROM topic_candidate_insights WHERE candidate_id = ?", (str(candidate_row["id"]),))
        if cached and not topic_candidate_insight_needs_refresh(candidate_row, cached):
            update_topic_candidate_insight_state(str(candidate_row["id"]), "ready")
            return build_topic_candidate_insight(cached), str(cached["source_excerpt"] or "")
        source_content = fetch_topic_source_excerpt(str(candidate_row["source_url"])) if candidate_row["source_url"] else ""
        insight_payload = state.ai.build_topic_candidate_insight(
            candidate_title=str(candidate_row["title"]),
            candidate_summary=str(candidate_row["summary"]),
            source=str(candidate_row["source"]),
            published_at=str(candidate_row["published_at"]) if candidate_row["published_at"] else None,
            source_url=str(candidate_row["source_url"]) if candidate_row["source_url"] else None,
            source_content=source_content,
            organization_context=build_organization_dna_context() if topics_settings.useOrgDnaForInsight else "",
        )
        insight = save_topic_candidate_insight(
            candidate_id=str(candidate_row["id"]),
            overview=str(insight_payload.get("overview") or "").strip(),
            key_points=[str(item) for item in insight_payload.get("keyPoints", []) if str(item).strip()] if isinstance(insight_payload.get("keyPoints"), list) else [],
            recommendation_reasons=[str(item) for item in insight_payload.get("recommendationReasons", []) if str(item).strip()] if isinstance(insight_payload.get("recommendationReasons"), list) else [],
            practical_uses=[str(item) for item in insight_payload.get("practicalUses", []) if str(item).strip()] if isinstance(insight_payload.get("practicalUses"), list) else [],
            editorial_note=str(insight_payload.get("editorialNote") or "").strip(),
            discussion_prompts=[str(item) for item in insight_payload.get("discussionPrompts", []) if str(item).strip()] if isinstance(insight_payload.get("discussionPrompts"), list) else [],
            source_excerpt=source_content,
        )
        update_topic_candidate_insight_state(str(candidate_row["id"]), "ready")
        log_activity("topic.candidate.insight", "topic_candidate", str(candidate_row["id"]), {"keyPoints": len(insight.keyPoints)})
        return insight, source_content

    def build_topic_candidate_chat_context(
        *,
        candidate_row,
        radar_title: str,
        insight: TopicCandidateInsightRecord | None,
        source_excerpt: str,
        history: list[TopicCandidateChatMessageRecord],
    ) -> str:
        key_points = insight.keyPoints if insight else []
        writing_angles = insight.practicalUses if insight else []
        discussion_prompts = insight.discussionPrompts if insight else []
        recommendation_reasons = insight.recommendationReasons if insight else []
        editorial_note = insight.editorialNote.strip() if insight else ""
        overview = insight.overview.strip() if insight else ""

        blocks = [
            f"当前情报标题：{str(candidate_row['title'])}",
            f"关联雷达：{radar_title}",
            f"来源：{str(candidate_row['source'])}",
        ]
        if candidate_row["published_at"]:
            blocks.append(f"发布时间：{str(candidate_row['published_at'])}")
        if candidate_row["source_url"]:
            blocks.append(f"原文链接：{str(candidate_row['source_url'])}")

        blocks.extend(
            [
                "",
                "候选摘要：",
                str(candidate_row["summary"] or "").strip() or "暂无摘要。",
            ]
        )
        if overview:
            blocks.extend(["", "大周对文章本身的概述：", overview])
        if recommendation_reasons:
            blocks.extend(["", "为什么这篇内容值得看："])
            blocks.extend(f"{index + 1}. {item}" for index, item in enumerate(recommendation_reasons[:4]))
        if key_points:
            blocks.extend(["", "核心观点："])
            blocks.extend(f"{index + 1}. {item}" for index, item in enumerate(key_points[:6]))
        if editorial_note:
            blocks.extend(["", "大周前哨判断：", editorial_note])
        if writing_angles:
            blocks.extend(["", "可继续展开成文的角度："])
            blocks.extend(f"{index + 1}. {item}" for index, item in enumerate(writing_angles[:4]))
        if discussion_prompts:
            blocks.extend(["", "原本建议继续追问的问题："])
            blocks.extend(f"{index + 1}. {item}" for index, item in enumerate(discussion_prompts[:4]))
        if source_excerpt.strip():
            blocks.extend(["", "原文关键摘录：", source_excerpt.strip()[:2800]])
        if history:
            blocks.extend(["", "已发生的对话："])
            for item in history[-8:]:
                speaker = "用户" if item.role == "user" else "大周"
                content = re.sub(r"\s+", " ", item.content.strip())[:500]
                if content:
                    blocks.append(f"{speaker}：{content}")
        return "\n".join(part for part in blocks if part is not None).strip()

    def prefetch_topic_candidate_insight(candidate_id: str) -> None:
        row = state.db.fetchone("SELECT * FROM topic_candidates WHERE id = ?", (candidate_id,))
        if not row:
            return
        update_topic_candidate_insight_state(candidate_id, "pending")
        try:
            ensure_topic_candidate_insight(row)
        except Exception as error:
            update_topic_candidate_insight_state(candidate_id, "failed", error=str(error)[:240])
            log_activity(
                "topic.candidate.insight.prefetch_failed",
                "topic_candidate",
                candidate_id,
                {"error": str(error)[:240]},
            )

    def schedule_topic_candidate_insight(candidate_id: str) -> None:
        if state.topic_insight_executor is None:
            prefetch_topic_candidate_insight(candidate_id)
            return
        state.topic_insight_executor.submit(prefetch_topic_candidate_insight, candidate_id)

    def create_topic_candidate_record(
        *,
        radar_id: str,
        title: str,
        summary: str,
        source: str,
        status: str,
        source_url: str | None = None,
        published_at: str | None = None,
        capture_method: str = "manual",
        captured_by: str | None = None,
    ) -> TopicCandidateRecord:
        candidate_id = new_id("cand")
        created_at = now_iso()
        state.db.execute(
            """
            INSERT INTO topic_candidates(
                id, radar_id, title, summary, source, source_url, published_at, capture_method, captured_by, status, created_at, updated_at
            )
            VALUES(?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                candidate_id,
                radar_id,
                title,
                summary,
                source,
                source_url,
                published_at,
                capture_method,
                captured_by,
                status,
                created_at,
                created_at,
            ),
        )
        row = state.db.fetchone("SELECT * FROM topic_candidates WHERE id = ?", (candidate_id,))
        assert row is not None
        remember_topic_candidate_seen(
            radar_id=radar_id,
            source_url=source_url,
            title=title,
            source=source,
        )
        schedule_topic_candidate_insight(candidate_id)
        return build_topic_candidate(row)

    def normalize_topic_candidate_source_url(value: str | None) -> str:
        raw = str(value or "").strip()
        if not raw:
            return ""
        try:
            parsed = urlparse(raw)
            path = (parsed.path or "").rstrip("/")
            normalized = urlunparse(
                (
                    (parsed.scheme or "https").lower(),
                    parsed.netloc.lower(),
                    path,
                    "",
                    "",
                    "",
                )
            ).strip()
            return normalized or raw.lower()
        except Exception:
            return raw.lower()

    def build_topic_candidate_match_keys(*, source_url: str | None, title: str, source: str) -> tuple[str, str]:
        source_url_key = normalize_topic_candidate_source_url(source_url)
        normalized_title = re.sub(r"\s+", " ", str(title or "").strip()).lower()
        normalized_source = re.sub(r"\s+", " ", str(source or "").strip()).lower()
        title_source_key = f"{normalized_title}||{normalized_source}".strip("|")
        return source_url_key, title_source_key

    def remember_topic_candidate_seen(
        *,
        radar_id: str,
        source_url: str | None,
        title: str,
        source: str,
        deleted_at: str | None = None,
    ) -> None:
        source_url_key, title_source_key = build_topic_candidate_match_keys(
            source_url=source_url,
            title=title,
            source=source,
        )
        if not source_url_key and not title_source_key:
            return
        existing = state.db.fetchone(
            """
            SELECT id
            FROM topic_candidate_seen
            WHERE radar_id = ?
              AND (
                (? <> '' AND source_url_key = ?)
                OR (? <> '' AND title_source_key = ?)
              )
            LIMIT 1
            """,
            (radar_id, source_url_key, source_url_key, title_source_key, title_source_key),
        )
        if existing:
            state.db.execute(
                """
                UPDATE topic_candidate_seen
                SET source_url_key = ?, title_source_key = ?, source_url = ?, title = ?, source = ?, deleted_at = ?
                WHERE id = ?
                """,
                (source_url_key, title_source_key, source_url, title, source, deleted_at, str(existing["id"])),
            )
            return
        state.db.execute(
            """
            INSERT INTO topic_candidate_seen(
                id, radar_id, source_url_key, title_source_key, source_url, title, source, created_at, deleted_at
            ) VALUES(?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (new_id("seen"), radar_id, source_url_key, title_source_key, source_url, title, source, now_iso(), deleted_at),
        )

    def topic_candidate_already_seen(*, radar_id: str, source_url: str | None, title: str, source: str) -> bool:
        source_url_key, title_source_key = build_topic_candidate_match_keys(
            source_url=source_url,
            title=title,
            source=source,
        )
        existing_candidate = state.db.fetchone(
            """
            SELECT id
            FROM topic_candidates
            WHERE radar_id = ?
              AND (
                (source_url IS NOT NULL AND source_url = ?)
                OR (LOWER(title) = LOWER(?) AND LOWER(source) = LOWER(?))
              )
            LIMIT 1
            """,
            (radar_id, source_url, title, source),
        )
        if existing_candidate is not None:
            return True
        if not source_url_key and not title_source_key:
            return False
        existing = state.db.fetchone(
            """
            SELECT id
            FROM topic_candidate_seen
            WHERE radar_id = ?
              AND (
                (? <> '' AND source_url_key = ?)
                OR (? <> '' AND title_source_key = ?)
              )
            LIMIT 1
            """,
            (radar_id, source_url_key, source_url_key, title_source_key, title_source_key),
        )
        return existing is not None

    def capture_topic_radar_internal(radar_row) -> TopicCaptureRunRecord:
        capture_limit = 5
        preferred_sources = parse_topic_radar_preferred_sources(radar_row["preferred_sources_json"])
        hits = fetch_topic_candidates_from_web(
            state.ai,
            radar_title=str(radar_row["title"]),
            radar_prompt=str(radar_row["prompt"]),
            time_range=str(radar_row["time_range"]),
            preferred_source_urls=[item.url for item in preferred_sources],
            max_items=capture_limit,
        )

        created_candidates: list[TopicCandidateRecord] = []
        skipped_count = 0
        query = hits[0].query if hits else ""

        for hit in hits:
            if len(created_candidates) >= capture_limit:
                break
            localized = state.ai.localize_topic_hit(
                title=hit.title,
                summary=hit.summary,
                radar_title=str(radar_row["title"]),
                radar_prompt=str(radar_row["prompt"]),
            )
            normalized_title = str(localized.get("title") or hit.title).strip()
            normalized_summary = str(localized.get("summary") or hit.summary).strip()
            if topic_candidate_already_seen(
                radar_id=str(radar_row["id"]),
                source_url=hit.source_url,
                title=normalized_title,
                source=hit.source,
            ):
                skipped_count += 1
                continue
            created_candidates.append(
                create_topic_candidate_record(
                    radar_id=str(radar_row["id"]),
                    title=normalized_title,
                    summary=normalized_summary,
                    source=hit.source,
                    status="tracking",
                    source_url=hit.source_url,
                    published_at=hit.published_at,
                    capture_method="web_search",
                    captured_by="大周",
                )
            )

        log_activity(
            "topic.radar.capture",
            "topic_radar",
            str(radar_row["id"]),
            {
                "radarTitle": str(radar_row["title"]),
                "query": query,
                "fetchedCount": len(hits),
                "createdCount": len(created_candidates),
                "skippedCount": skipped_count,
            },
        )
        return TopicCaptureRunRecord(
            radarId=str(radar_row["id"]),
            radarTitle=str(radar_row["title"]),
            query=query,
            fetchedCount=len(hits),
            createdCount=len(created_candidates),
            skippedCount=skipped_count,
            candidates=created_candidates,
        )

    @app.get("/api/v1/topics", response_model=TopicsResponse)
    def list_topics() -> TopicsResponse:
        radars = [build_topic_radar(row) for row in state.db.fetchall("SELECT * FROM topic_radars ORDER BY created_at ASC")]
        return TopicsResponse(radars=radars, candidates=fetch_topic_candidates())

    @app.post("/api/v1/topics/radars", response_model=TopicRadarRecord)
    def create_radar(payload: TopicRadarPayload) -> TopicRadarRecord:
        radar_id = new_id("radar")
        created_at = now_iso()
        preferred_sources = normalize_topic_radar_preferred_sources(payload.preferredSources)
        state.db.execute(
            "INSERT INTO topic_radars(id, title, prompt, time_range, preferred_sources_json, created_at) VALUES(?, ?, ?, ?, ?, ?)",
            (radar_id, payload.title, payload.prompt, payload.timeRange, to_json([item.model_dump() for item in preferred_sources]), created_at),
        )
        log_activity("topic.radar.create", "topic_radar", radar_id, payload.model_dump())
        return TopicRadarRecord(
            id=radar_id,
            title=payload.title,
            prompt=payload.prompt,
            timeRange=payload.timeRange,
            preferredSources=preferred_sources,
            createdAt=created_at,
        )

    @app.put("/api/v1/topics/radars/{radar_id}", response_model=TopicRadarRecord)
    def update_radar(radar_id: str, payload: TopicRadarPayload) -> TopicRadarRecord:
        row = state.db.fetchone("SELECT * FROM topic_radars WHERE id = ?", (radar_id,))
        if not row:
            raise HTTPException(status_code=404, detail="Radar not found")
        preferred_sources = normalize_topic_radar_preferred_sources(payload.preferredSources)
        state.db.execute(
            "UPDATE topic_radars SET title = ?, prompt = ?, time_range = ?, preferred_sources_json = ? WHERE id = ?",
            (payload.title, payload.prompt, payload.timeRange, to_json([item.model_dump() for item in preferred_sources]), radar_id),
        )
        log_activity("topic.radar.update", "topic_radar", radar_id, payload.model_dump())
        return TopicRadarRecord(
            id=radar_id,
            title=payload.title,
            prompt=payload.prompt,
            timeRange=payload.timeRange,
            preferredSources=preferred_sources,
            createdAt=str(row["created_at"]),
        )

    @app.post("/api/v1/topics/radars/generate-title", response_model=TitleSuggestionResponse)
    def generate_radar_title(payload: TopicTitlePayload) -> TitleSuggestionResponse:
        return TitleSuggestionResponse(title=state.ai.suggest_short_title(payload.prompt))

    @app.post("/api/v1/topics/radars/source-label", response_model=TopicRadarSourceLabelResponse)
    def suggest_radar_source_label(payload: TopicRadarSourceLabelPayload) -> TopicRadarSourceLabelResponse:
        normalized_url = normalize_topic_radar_source_url(payload.url)
        return TopicRadarSourceLabelResponse(url=normalized_url, label=suggest_topic_radar_source_label(normalized_url))

    def _build_assisted_radar_prompt(*, prompt: str, title: str, queries: list[str], time_range: str) -> str:
        cleaned_prompt = prompt.strip()
        window_label = {
            "1_day": "近 1 天",
            "3_days": "近 3 天",
            "7_days": "近 7 天",
            "30_days": "近 30 天",
        }.get(time_range, "最近一周")
        query_text = "、".join(f"“{item}”" for item in queries[:3] if item.strip())
        focus_title = title.strip() or "这个主题"
        guidance = (
            f"重点追踪 {focus_title} 在 {window_label} 内的最新动态，优先留意政策发布、项目案例、方法总结、争议讨论与行业信号，"
            "并在整理时明确发布时间、适用场景、关键数据、执行门槛、涉及机构与可复用做法。"
        )
        if query_text:
            guidance = f"{guidance} 可优先使用 {query_text} 这些搜索表达。"
        return f"{cleaned_prompt}\n\n{guidance}".strip()

    @app.post("/api/v1/topics/radars/assist", response_model=TopicRadarAssistResponse)
    def assist_radar_draft(payload: TopicRadarAssistPayload) -> TopicRadarAssistResponse:
        prompt = payload.prompt.strip()
        if not prompt:
            raise HTTPException(status_code=400, detail="Prompt is required")
        title = state.ai.suggest_short_title(prompt)
        queries = state.ai.suggest_topic_search_queries(title=title, prompt=prompt, time_range=payload.timeRange)
        assisted_prompt = _build_assisted_radar_prompt(prompt=prompt, title=title, queries=queries, time_range=payload.timeRange)
        return TopicRadarAssistResponse(title=title, prompt=assisted_prompt, queries=queries)

    @app.post("/api/v1/topics/candidates", response_model=TopicCandidateRecord)
    def create_candidate(payload: TopicCandidatePayload) -> TopicCandidateRecord:
        if not state.db.fetchone("SELECT 1 FROM topic_radars WHERE id = ?", (payload.radarId,)):
            raise HTTPException(status_code=404, detail="Radar not found")
        created = create_topic_candidate_record(
            radar_id=payload.radarId,
            title=payload.title,
            summary=payload.summary,
            source=payload.source,
            status="candidate",
        )
        log_activity("topic.candidate.create", "topic_candidate", created.id, payload.model_dump())
        return created

    @app.post("/api/v1/topics/candidates/{candidate_id}/insights", response_model=TopicCandidateInsightRecord)
    def get_candidate_insights(candidate_id: str) -> TopicCandidateInsightRecord:
        row = state.db.fetchone("SELECT * FROM topic_candidates WHERE id = ?", (candidate_id,))
        if not row:
            raise HTTPException(status_code=404, detail="Candidate not found")
        if str(row["insight_status"] or "pending") != "ready":
            raise HTTPException(status_code=409, detail="候选解析尚未完成")
        insight, _ = ensure_topic_candidate_insight(row)
        return insight

    @app.post("/api/v1/topics/candidates/{candidate_id}/chat", response_model=TopicCandidateChatResponse)
    def chat_with_topic_candidate(candidate_id: str, payload: TopicCandidateChatPayload) -> TopicCandidateChatResponse:
        row = state.db.fetchone("SELECT * FROM topic_candidates WHERE id = ?", (candidate_id,))
        if not row:
            raise HTTPException(status_code=404, detail="Candidate not found")
        question = payload.question.strip()
        if not question:
            raise HTTPException(status_code=400, detail="问题不能为空")

        radar_row = state.db.fetchone("SELECT * FROM topic_radars WHERE id = ?", (str(row["radar_id"]),))
        radar_title = str(radar_row["title"]) if radar_row else "未命名雷达"

        insight: TopicCandidateInsightRecord | None = None
        source_excerpt = ""
        try:
            insight, source_excerpt = ensure_topic_candidate_insight(row)
        except Exception:
            cached = state.db.fetchone("SELECT * FROM topic_candidate_insights WHERE candidate_id = ?", (candidate_id,))
            if cached:
                insight = build_topic_candidate_insight(cached)
                source_excerpt = str(cached["source_excerpt"] or "")

        context_summary = build_topic_candidate_chat_context(
            candidate_row=row,
            radar_title=radar_title,
            insight=insight,
            source_excerpt=source_excerpt,
            history=payload.history,
        )
        system_instruction = (
            "你是资讯情报站里的大周。"
            "你现在只围绕当前这篇情报继续回答问题。"
            "回答时优先基于这篇新闻本身、大周已有的解析和原文摘录，不要脱离当前材料泛泛而谈。"
            "你可以做更高层的判断，但不要编造当前材料里没有出现过的具体事实、数据、人名或时间。"
            "如果用户的问题已经超出当前材料，就直接说明材料还不够，并指出下一步该补什么信息。"
            "请用中文直接回答，不要解释系统过程，不要输出 JSON。"
        )
        try:
            structured = state.ai.generate_topic_candidate_chat_response(question, system_instruction, context_summary)
        except AiInvocationError as error:
            raise HTTPException(status_code=502, detail=str(error)) from error

        answer = structured.content.strip() or structured.analysis.strip() or structured.judgment.strip()
        if not answer:
            answer = "我暂时还没法基于这篇情报给出稳定回答，建议先点开原文再继续追问。"
        generated_at = now_iso()
        message = TopicCandidateChatMessageRecord(role="assistant", content=answer, createdAt=generated_at)
        log_activity(
            "topic.candidate.chat",
            "topic_candidate",
            candidate_id,
            {
                "questionLength": len(question),
                "answerLength": len(answer),
            },
        )
        return TopicCandidateChatResponse(
            candidateId=candidate_id,
            question=question,
            answer=answer,
            generatedAt=generated_at,
            message=message,
        )

    @app.post("/api/v1/topics/candidates/{candidate_id}/task-plan", response_model=TopicTaskPlanResponse)
    def build_candidate_task_plan(candidate_id: str) -> TopicTaskPlanResponse:
        row = state.db.fetchone("SELECT * FROM topic_candidates WHERE id = ?", (candidate_id,))
        if not row:
            raise HTTPException(status_code=404, detail="Candidate not found")
        topics_settings = get_topics_settings()
        if topics_settings.requireInsightBeforeActions and str(row["insight_status"] or "pending") != "ready":
            raise HTTPException(status_code=409, detail="候选解析尚未完成")
        insight, source_content = ensure_topic_candidate_insight(row)
        plan = state.ai.build_topic_task_plan(
            candidate_title=str(row["title"]),
            candidate_summary=str(row["summary"]),
            source=str(row["source"]),
            published_at=str(row["published_at"]) if row["published_at"] else None,
            source_url=str(row["source_url"]) if row["source_url"] else None,
            source_content=source_content,
            candidate_insight=insight.model_dump(),
            organization_context=build_organization_dna_context() if topics_settings.useOrgDnaForTaskPlan else "",
        )
        tasks = [
            TopicTaskSuggestionRecord(**task)
            for task in plan.get("tasks", [])
            if isinstance(task, dict)
        ]
        if not tasks:
            fallback_title = str(row["title"])[:60]
            tasks = [
                TopicTaskSuggestionRecord(
                    title=fallback_title,
                    desc=str(row["summary"])[:180],
                    ddl="待确认",
                    note=f"来源：{row['source']}",
                    priority="normal",
                    tags=["资讯机会"],
                )
            ]
        return TopicTaskPlanResponse(
            candidateId=str(row["id"]),
            candidateTitle=str(row["title"]),
            candidateSummary=str(row["summary"]),
            candidateSource=str(row["source"]),
            candidateSourceUrl=str(row["source_url"]) if row["source_url"] else None,
            overview=str(plan.get("overview") or "").strip(),
            tasks=tasks,
        )

    @app.post("/api/v1/topics/candidates/{candidate_id}/promote-tasks", response_model=TopicTaskPromotionResponse)
    def promote_candidate_to_tasks(candidate_id: str, payload: TopicTaskPromotionPayload) -> TopicTaskPromotionResponse:
        row = state.db.fetchone("SELECT * FROM topic_candidates WHERE id = ?", (candidate_id,))
        if not row:
            raise HTTPException(status_code=404, detail="Candidate not found")
        if not payload.tasks:
            raise HTTPException(status_code=400, detail="至少要选择一条任务")
        session_user = get_cached_session_user() if get_cloud_token() else None
        default_owner_name = session_user.fullName if session_user else current_operator_row()["name"]
        created_tasks: list[TaskRecord] = []
        for item in payload.tasks:
            if not item.title.strip():
                continue
            task = create_task(
                TaskPayload(
                    title=item.title.strip(),
                    desc=item.desc.strip(),
                    priority=item.priority,
                    listId=item.listId,
                    dueDate=item.dueDate,
                    ddl=item.ddl.strip() or item.dueDate or "待确认",
                    ownerId=item.ownerId,
                    ownerName=item.ownerName.strip() or default_owner_name,
                    collaboratorIds=item.collaboratorIds,
                    tagIds=item.tagIds,
                    tags=item.tags,
                    sourceType="topic_candidate",
                    sourceId=candidate_id,
                ),
                status="inbox",
            )
            if item.note.strip():
                upsert_task_note(task.id, item.note.strip())
            created_tasks.append(task)
        if not created_tasks:
            raise HTTPException(status_code=400, detail="没有可创建的任务")
        state.db.execute("UPDATE topic_candidates SET status = 'promoted', updated_at = ? WHERE id = ?", (now_iso(), candidate_id))
        log_activity(
            "topic.promote.tasks",
            "topic_candidate",
            candidate_id,
            {
                "taskIds": [task.id for task in created_tasks],
                "count": len(created_tasks),
            },
        )
        return TopicTaskPromotionResponse(tasks=created_tasks, createdCount=len(created_tasks))

    @app.post("/api/v1/topics/capture", response_model=TopicCaptureBatchResponse)
    def capture_all_topic_radars() -> TopicCaptureBatchResponse:
        radar_rows = state.db.fetchall("SELECT * FROM topic_radars ORDER BY created_at ASC")
        with ThreadPoolExecutor(max_workers=min(4, max(1, len(radar_rows)))) as executor:
            runs = list(executor.map(capture_topic_radar_internal, radar_rows))
        return TopicCaptureBatchResponse(
            runs=runs,
            totalCreated=sum(item.createdCount for item in runs),
            totalSkipped=sum(item.skippedCount for item in runs),
        )

    @app.post("/api/v1/topics/candidates/{candidate_id}/promote-task", response_model=TaskRecord)
    def promote_candidate_to_task(candidate_id: str, payload: dict = Body(default_factory=dict)) -> TaskRecord:
        row = state.db.fetchone("SELECT * FROM topic_candidates WHERE id = ?", (candidate_id,))
        if not row:
            raise HTTPException(status_code=404, detail="Candidate not found")
        event_line_id = str(payload.get("eventLineId", "") or row.get("event_line_id", "") or "").strip() or None
        state.db.execute("UPDATE topic_candidates SET status = 'promoted', updated_at = ? WHERE id = ?", (now_iso(), candidate_id))
        task = create_task(
            TaskPayload(
                title=str(row["title"]),
                desc=str(row["summary"]),
                priority="normal",
                listId="list-0",
                ddl="本周",
                ownerName=current_operator_row()["name"],
                tags=["选题"],
                sourceType="topic_candidate",
                sourceId=candidate_id,
                eventLineId=event_line_id,
            ),
            status="inbox",
        )
        log_activity("topic.promote.task", "topic_candidate", candidate_id, {"taskId": task.id, "eventLineId": event_line_id})
        return task

    @app.delete("/api/v1/topics/candidates/{candidate_id}")
    def delete_topic_candidate(candidate_id: str) -> dict[str, bool]:
        row = state.db.fetchone("SELECT * FROM topic_candidates WHERE id = ?", (candidate_id,))
        if not row:
            raise HTTPException(status_code=404, detail="Candidate not found")
        remember_topic_candidate_seen(
            radar_id=str(row["radar_id"]),
            source_url=str(row["source_url"]) if row["source_url"] else None,
            title=str(row["title"]),
            source=str(row["source"]),
            deleted_at=now_iso(),
        )
        state.db.execute("DELETE FROM topic_candidates WHERE id = ?", (candidate_id,))
        log_activity("topic.candidate.delete", "topic_candidate", candidate_id, {"title": str(row["title"])})
        return {"deleted": True}

    @app.get("/api/v1/analysis-tools", response_model=AnalysisToolsResponse)
    def list_analysis_tools() -> AnalysisToolsResponse:
        templates = [
            AnalysisTemplateRecord(
                id=str(row["id"]),
                title=str(row["title"]),
                description=str(row["description"]),
                templateKey=str(row["template_key"]),
            )
            for row in state.db.fetchall("SELECT * FROM analysis_templates ORDER BY created_at ASC")
        ]
        runs = [
            AnalysisRunRecord(
                id=str(row["id"]),
                templateId=str(row["template_id"]),
                title=str(row["title"]),
                inputText=str(row["input_text"]),
                output=AiStructuredResponse(**from_json(row["output_json"], {})),
                createdAt=str(row["created_at"]),
                status=str(row["status"]),  # type: ignore[arg-type]
            )
            for row in state.db.fetchall("SELECT * FROM analysis_runs ORDER BY created_at DESC")
        ]
        return AnalysisToolsResponse(templates=templates, runs=runs)

    @app.post("/api/v1/analysis-tools/runs", response_model=AnalysisRunRecord)
    def run_analysis(payload: AnalysisRunPayload) -> AnalysisRunRecord:
        template = state.db.fetchone("SELECT * FROM analysis_templates WHERE id = ?", (payload.templateId,))
        if not template:
            raise HTTPException(status_code=404, detail="Analysis template not found")
        analysis_settings = get_analysis_workbench_settings()
        if analysis_settings.enabledTemplateIds and payload.templateId not in analysis_settings.enabledTemplateIds:
            raise HTTPException(status_code=403, detail="当前模板已在系统设置中停用")
        org_context = build_organization_dna_context() if analysis_settings.useOrgDna else ""
        output = state.ai.generate_structured(
            payload.inputText,
            f"你是咨询分析助手，请根据模板 {template['title']} 输出结构化结果。",
            "\n\n".join(
                item
                for item in [
                    f"模板说明：{template['description']}",
                    org_context,
                ]
                if item
            ),
        )
        run_id = new_id("run")
        created_at = now_iso()
        state.db.execute(
            "INSERT INTO analysis_runs(id, template_id, title, input_text, output_json, status, created_at) VALUES(?, ?, ?, ?, ?, 'success', ?)",
            (run_id, payload.templateId, payload.title, payload.inputText, to_json(output.model_dump()), created_at),
        )
        log_activity("analysis.run", "analysis_run", run_id, {"templateId": payload.templateId, "title": payload.title})
        return AnalysisRunRecord(id=run_id, templateId=payload.templateId, title=payload.title, inputText=payload.inputText, output=output, createdAt=created_at, status="success")

    def build_handbook_entry_record(row) -> HandbookEntryRecord:
        client_id = str(row["client_id"]) if row["client_id"] else None
        client_name = None
        if client_id:
            client_row = state.db.fetchone("SELECT name FROM clients WHERE id = ?", (client_id,))
            client_name = str(client_row["name"]) if client_row and client_row["name"] else None
        linked_contexts: list[dict[str, object]] = []
        linked_contexts.append({"objectType": "handbook", "objectId": str(row["id"]), "label": str(row["title"]), "subtitle": str(row["source_type"] or ""), "tab": "growth", "statusLabel": ""})
        if row["source_object_id"] and row["source_title"]:
            linked_contexts.append(
                {
                    "objectType": str(row["source_object_type"] or "unknown"),
                    "objectId": str(row["source_object_id"]),
                    "label": str(row["source_title"]),
                    "subtitle": "",
                    "tab": "growth",
                    "statusLabel": "",
                }
            )
        if client_id and client_name:
            linked_contexts.append({"objectType": "client", "objectId": client_id, "label": client_name, "subtitle": "", "tab": "client_workspace", "statusLabel": ""})
        if row["event_line_id"] and row["event_line_name"]:
            linked_contexts.append(
                {
                    "objectType": "event_line",
                    "objectId": str(row["event_line_id"]),
                    "label": str(row["event_line_name"]),
                    "subtitle": str(row["project_stage"] or ""),
                    "tab": "tasks",
                    "statusLabel": "",
                }
            )
        return HandbookEntryRecord(
            id=str(row["id"]),
            title=str(row["title"]),
            summary=str(row["summary"]),
            tags=_parse_json_list(row["tags_json"]),
            sourceType=str(row["source_type"]),
            clientId=client_id,
            clientName=client_name,
            sourceObjectType=str(row["source_object_type"]) if row["source_object_type"] else None,
            sourceObjectId=str(row["source_object_id"]) if row["source_object_id"] else None,
            sourceTitle=str(row["source_title"]) if row["source_title"] else None,
            eventLineId=str(row["event_line_id"]) if row["event_line_id"] else None,
            eventLineName=str(row["event_line_name"]) if row["event_line_name"] else None,
            projectModuleId=str(row["project_module_id"]) if row["project_module_id"] else None,
            projectModuleName=str(row["project_module_name"]) if row["project_module_name"] else None,
            projectFlowId=str(row["project_flow_id"]) if row["project_flow_id"] else None,
            projectFlowName=str(row["project_flow_name"]) if row["project_flow_name"] else None,
            projectStage=str(row["project_stage"]) if row["project_stage"] else None,
            businessCategory=str(row["business_category"]) if row["business_category"] else None,
            abilityKeys=from_json(row["ability_keys_json"], []),
            evidenceRefs=from_json(row["evidence_refs_json"], []),
            contextSummary=str(row["context_summary"] or ""),
            reuseCount=int(row["reuse_count"] or 0),
            lastReusedAt=str(row["last_reused_at"]) if row["last_reused_at"] else None,
            linkedContexts=linked_contexts,
            createdAt=str(row["created_at"]),
        )

    def build_handbook_detail(entry_id: str, user_id: str) -> HandbookEntryDetailRecord:
        row = state.db.fetchone("SELECT * FROM handbook_entries WHERE id = ?", (entry_id,))
        if not row:
            raise HTTPException(status_code=404, detail="成长手册条目不存在")
        entry = build_handbook_entry_record(row)
        related_rows = state.db.fetchall(
            """
            SELECT
                l.*,
                e.reason,
                e.evidence_type,
                e.metadata_json,
                e.contribution_tags_json,
                e.validation_state,
                e.org_contribution_score,
                e.review_id,
                e.task_id,
                e.handbook_entry_id,
                s.source_type,
                s.source_id,
                s.context_json
            FROM xp_ledger l
            INNER JOIN growth_evidence_records e ON e.id = l.evidence_id
            INNER JOIN growth_signal_events s ON s.id = e.signal_id
            WHERE l.user_id = ? AND l.reversed_at IS NULL AND e.handbook_entry_id = ?
            ORDER BY l.created_at DESC
            LIMIT 20
            """,
            (user_id, entry_id),
        )
        from app.services.growth_engine import _build_ledger_entry, _fetch_profile_map

        profile_map = _fetch_profile_map(state.db)
        related_entries = [_build_ledger_entry(profile_map, ledger_row) for ledger_row in related_rows]

        def merge_context_dicts(*items: object) -> dict[str, object]:
            merged: dict[str, object] = {}
            for item in items:
                if not isinstance(item, dict):
                    continue
                for key, value in item.items():
                    if key not in merged or merged[key] in (None, "", [], {}):
                        merged[key] = value
            return merged

        def build_context_links(raw_links: object) -> list[GrowthContextLinkRecord]:
            if not isinstance(raw_links, list):
                return []
            links: list[GrowthContextLinkRecord] = []
            for raw in raw_links:
                if not isinstance(raw, dict):
                    continue
                object_type = str(raw.get("objectType") or "").strip()
                object_id = str(raw.get("objectId") or "").strip()
                label = str(raw.get("label") or "").strip()
                if not object_type or not object_id or not label:
                    continue
                links.append(
                    GrowthContextLinkRecord(
                        objectType=object_type,
                        objectId=object_id,
                        label=label,
                        subtitle=str(raw.get("subtitle") or ""),
                        tab=str(raw.get("tab") or ""),
                        statusLabel=str(raw.get("statusLabel") or ""),
                    )
                )
            return links

        def dedupe_context_links(*groups: list[GrowthContextLinkRecord]) -> list[GrowthContextLinkRecord]:
            deduped: list[GrowthContextLinkRecord] = []
            seen: set[tuple[str, str]] = set()
            for group in groups:
                for link in group:
                    key = (link.objectType, link.objectId)
                    if key in seen:
                        continue
                    seen.add(key)
                    deduped.append(link)
            return deduped

        origin_contexts = dedupe_context_links(
            entry.linkedContexts,
            *[ledger_entry.linkedContexts for ledger_entry in related_entries],
        )

        reuse_rows = state.db.fetchall(
            """
            SELECT
                v.*,
                l.delta,
                l.total_xp,
                e.metadata_json,
                s.context_json
            FROM growth_validation_events v
            INNER JOIN growth_evidence_records e ON e.id = v.evidence_id
            INNER JOIN growth_signal_events s ON s.id = e.signal_id
            LEFT JOIN xp_ledger l ON l.evidence_id = e.id AND l.reversed_at IS NULL
            WHERE v.user_id = ? AND e.handbook_entry_id = ? AND v.event_type = 'handbook_reused'
            ORDER BY v.created_at DESC
            """,
            (user_id, entry_id),
        )
        reuse_buckets: dict[str, dict[str, object]] = {}
        for reuse_row in reuse_rows:
            detail = from_json(reuse_row["detail_json"], {})
            metadata = from_json(reuse_row["metadata_json"], {})
            signal_context = from_json(reuse_row["context_json"], {})
            merged_context = merge_context_dicts(signal_context, metadata, detail)
            linked_contexts = build_context_links(detail.get("linkedContexts") if isinstance(detail, dict) else None)
            if not linked_contexts:
                linked_contexts = build_context_links(merged_context.get("linkedContexts"))
            bucket_id = f"{reuse_row['source_type']}::{reuse_row['source_id']}::{reuse_row['created_at']}"
            bucket = reuse_buckets.get(bucket_id)
            if bucket is None:
                source_label = (
                    str(detail.get("sourceLabel") or "").strip()
                    if isinstance(detail, dict)
                    else ""
                )
                if not source_label:
                    source_label = next((item.label for item in linked_contexts if item.label), "").strip()
                if not source_label:
                    source_label = str(reuse_row["source_id"] or "已记录复用").strip()
                bucket = {
                    "id": bucket_id,
                    "sourceType": str(reuse_row["source_type"] or ""),
                    "sourceId": str(reuse_row["source_id"] or ""),
                    "sourceLabel": source_label,
                    "note": str(detail.get("note") or "").strip() if isinstance(detail, dict) else "",
                    "contextSummary": str((detail.get("contextSummary") if isinstance(detail, dict) else "") or merged_context.get("contextSummary") or "").strip(),
                    "gainedXp": 0,
                    "createdAt": str(reuse_row["created_at"]),
                    "linkedContexts": linked_contexts,
                }
                reuse_buckets[bucket_id] = bucket
            bucket["gainedXp"] = int(bucket["gainedXp"]) + int(reuse_row["total_xp"] or reuse_row["delta"] or 0)
            bucket["linkedContexts"] = dedupe_context_links(bucket["linkedContexts"], linked_contexts)  # type: ignore[arg-type]

        reuse_history = [
            HandbookReuseRecord(
                id=str(item["id"]),
                sourceType=str(item["sourceType"]),
                sourceId=str(item["sourceId"]),
                sourceLabel=str(item["sourceLabel"]),
                note=str(item["note"]),
                contextSummary=str(item["contextSummary"]),
                gainedXp=int(item["gainedXp"]),
                createdAt=str(item["createdAt"]),
                linkedContexts=list(item["linkedContexts"]),  # type: ignore[arg-type]
            )
            for item in sorted(
                reuse_buckets.values(),
                key=lambda current: str(current["createdAt"]),
                reverse=True,
            )
        ]

        return HandbookEntryDetailRecord(
            **entry.model_dump(),
            relatedLedgerEntries=related_entries,
            originContexts=origin_contexts,
            reuseHistory=reuse_history,
        )

    def create_handbook_entry(payload: HandbookPayload) -> HandbookEntryRecord:
        handbook_settings = get_handbook_settings()
        if payload.sourceType == "task" and not handbook_settings.allowTaskSource:
            raise HTTPException(status_code=403, detail="当前系统设置禁止从任务沉淀进入成长手册")
        if payload.sourceType == "analysis" and not handbook_settings.allowAnalysisSource:
            raise HTTPException(status_code=403, detail="当前系统设置禁止从分析结论沉淀进入成长手册")
        resolved_tags = payload.tags or handbook_settings.defaultTags
        entry_id = new_id("handbook")
        created_at = now_iso()
        state.db.execute(
            """
            INSERT INTO handbook_entries(
                id, title, summary, tags_json, source_type, client_id, source_object_type, source_object_id, source_title,
                event_line_id, event_line_name, project_module_id, project_module_name, project_flow_id, project_flow_name,
                project_stage, business_category, ability_keys_json, evidence_refs_json, context_summary, reuse_count, last_reused_at,
                author_user_id, author_user_name, created_at
            ) VALUES(?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 0, NULL, ?, ?, ?)
            """,
            (
                entry_id,
                payload.title,
                payload.summary,
                to_json(resolved_tags),
                payload.sourceType,
                payload.clientId,
                payload.sourceObjectType,
                payload.sourceObjectId,
                payload.sourceTitle,
                payload.eventLineId,
                payload.eventLineName,
                payload.projectModuleId,
                payload.projectModuleName,
                payload.projectFlowId,
                payload.projectFlowName,
                payload.projectStage,
                payload.businessCategory,
                to_json(payload.abilityKeys),
                to_json(payload.evidenceRefs),
                payload.contextSummary.strip(),
                resolve_growth_actor()[0],
                resolve_growth_actor()[1],
                created_at,
            ),
        )
        row = state.db.fetchone("SELECT * FROM handbook_entries WHERE id = ?", (entry_id,))
        assert row is not None
        return build_handbook_entry_record(row)

    @app.get("/api/v1/handbook", response_model=HandbookResponse)
    def list_handbook() -> HandbookResponse:
        entries = [build_handbook_entry_record(row) for row in state.db.fetchall("SELECT * FROM handbook_entries ORDER BY created_at DESC")]
        return HandbookResponse(entries=entries)

    @app.get("/api/v1/handbook/{entry_id}", response_model=HandbookEntryDetailRecord)
    def get_handbook_entry_detail(entry_id: str) -> HandbookEntryDetailRecord:
        user_id, _user_name = resolve_growth_actor()
        return build_handbook_detail(entry_id, user_id)

    @app.post("/api/v1/handbook", response_model=HandbookEntryRecord)
    def create_handbook(payload: HandbookPayload) -> HandbookEntryRecord:
        entry = create_handbook_entry(payload)
        user_id, user_name = resolve_growth_actor()
        ingest_handbook_codification(state.db, user_id=user_id, user_name=user_name, entry=entry, created_at=entry.createdAt)
        log_activity("handbook.create", "handbook_entry", entry.id, payload.model_dump())
        return entry

    @app.post("/api/v1/growth/handbook/{entry_id}/mark-reused", response_model=GrowthValidationActionResponse)
    def mark_growth_handbook_reused(entry_id: str, payload: GrowthValidationPayload) -> GrowthValidationActionResponse:
        user_id, user_name = resolve_growth_actor()
        handbook_entries = list_handbook().entries
        entry = next((item for item in handbook_entries if item.id == entry_id), None)
        if entry is None:
            raise HTTPException(status_code=404, detail="成长手册条目不存在")
        return mark_handbook_entry_reused(
            state.db,
            user_id=user_id,
            user_name=user_name,
            entry=entry,
            week_label=resolve_growth_week_label(user_id, None),
            source_type=(payload.sourceType or "handbook_manual_reuse").strip() or "handbook_manual_reuse",
            source_id=(payload.sourceId or "").strip() or resolve_growth_week_label(user_id, None),
            source_label=payload.sourceLabel or "",
            context_summary=payload.contextSummary or "",
            linked_contexts=[item.model_dump() for item in payload.linkedContexts] if payload.linkedContexts else None,
            note=payload.note,
            created_at=now_iso(),
        )

    @app.get("/api/v1/growth/overview", response_model=GrowthOverviewRecord)
    def get_growth_overview(weekLabel: str | None = Query(default=None)) -> GrowthOverviewRecord:
        user_id, user_name = resolve_growth_actor()
        handbook_entries = list_handbook().entries
        if handbook_entries:
            backfill_handbook_entries(state.db, user_id=user_id, user_name=user_name, entries=handbook_entries)
        build_badge_board(state.db, user_id=user_id, user_name=user_name, auto_sync=True)
        return build_growth_overview(
            state.db,
            user_id=user_id,
            user_name=user_name,
            week_label=resolve_growth_week_label(user_id, weekLabel),
        )

    @app.get("/api/v1/growth/workbench", response_model=GrowthWorkbenchSnapshotRecord)
    def get_growth_workbench(weekLabel: str | None = Query(default=None)) -> GrowthWorkbenchSnapshotRecord:
        user_id, user_name = resolve_growth_actor()
        handbook_entries = list_handbook().entries
        if handbook_entries:
            backfill_handbook_entries(state.db, user_id=user_id, user_name=user_name, entries=handbook_entries)
        build_badge_board(state.db, user_id=user_id, user_name=user_name, auto_sync=True)
        return build_growth_workbench_snapshot(week_label=resolve_growth_week_label(user_id, weekLabel))

    @app.get("/api/v1/growth/ledger", response_model=GrowthLedgerResponse)
    def get_growth_ledger(
        abilityKey: str | None = Query(default=None),
        weekLabel: str | None = Query(default=None),
    ) -> GrowthLedgerResponse:
        user_id, user_name = resolve_growth_actor()
        handbook_entries = list_handbook().entries
        if handbook_entries:
            backfill_handbook_entries(state.db, user_id=user_id, user_name=user_name, entries=handbook_entries)
        build_badge_board(state.db, user_id=user_id, user_name=user_name, auto_sync=True)
        return build_growth_ledger(state.db, user_id=user_id, ability_key=abilityKey, week_label=weekLabel)

    @app.get("/api/v1/growth/badges", response_model=BadgeBoardResponse)
    def get_growth_badges() -> BadgeBoardResponse:
        user_id, user_name = resolve_growth_actor()
        handbook_entries = list_handbook().entries
        if handbook_entries:
            backfill_handbook_entries(state.db, user_id=user_id, user_name=user_name, entries=handbook_entries)
        return build_badge_board(state.db, user_id=user_id, user_name=user_name, auto_sync=True)

    @app.get("/api/v1/growth/recommendations", response_model=list[LearningRecommendationRecord])
    def get_growth_recommendations() -> list[LearningRecommendationRecord]:
        user_id, user_name = resolve_growth_actor()
        handbook_entries = list_handbook().entries
        if handbook_entries:
            backfill_handbook_entries(state.db, user_id=user_id, user_name=user_name, entries=handbook_entries)
        return list_learning_recommendations(state.db, user_id)

    @app.post("/api/v1/growth/pending-captures/{capture_id}/state", response_model=GrowthPendingCaptureActionResponse)
    def update_growth_pending_capture(capture_id: str, payload: GrowthPendingCaptureActionPayload) -> GrowthPendingCaptureActionResponse:
        user_id, _user_name = resolve_growth_actor()
        updated = update_pending_capture_state(
            state.db,
            user_id=user_id,
            capture_id=capture_id,
            status=payload.status,
            reason=payload.reason,
            handbook_entry_id=payload.handbookEntryId,
            created_at=now_iso(),
        )
        if updated is None:
            raise HTTPException(status_code=404, detail="待放大的成长信号不存在或已失效")
        return GrowthPendingCaptureActionResponse(capture=updated)

    @app.post("/api/v1/growth/recommendations/{recommendation_id}/accept", response_model=GrowthRecommendationActionResponse)
    def accept_growth_recommendation(recommendation_id: str) -> GrowthRecommendationActionResponse:
        user_id, user_name = resolve_growth_actor()
        recommendation = next((item for item in list_learning_recommendations(state.db, user_id) if item.id == recommendation_id), None)
        if recommendation is None:
            raise HTTPException(status_code=404, detail="成长练习推荐不存在或已失效")
        task_settings = _get_local_task_settings()
        task = create_task(
            TaskPayload(
                title=f"成长练习：{recommendation.title}",
                desc="\n".join(
                    line
                    for line in (
                        recommendation.summary.strip(),
                        f"推荐原因：{recommendation.reason.strip()}",
                        f"行动目标：{recommendation.practiceTask.strip()}",
                    )
                    if line
                ),
                priority="normal",
                listId=task_settings.defaultListId or "list-0",
                ownerName=user_name,
                tags=["成长练习", recommendation.abilityLabel],
                sourceType="growth_recommendation",
                sourceId=recommendation.id,
            )
        )
        updated = mark_recommendation_accepted(state.db, recommendation_id, task.id)
        if updated is None:
            raise HTTPException(status_code=404, detail="成长练习推荐不存在或已失效")
        return GrowthRecommendationActionResponse(recommendation=updated, task=task)

    @app.post("/api/v1/growth/recommendations/{recommendation_id}/dismiss", response_model=GrowthRecommendationActionResponse)
    def dismiss_growth_recommendation(
        recommendation_id: str,
        payload: GrowthRecommendationDismissPayload,
    ) -> GrowthRecommendationActionResponse:
        user_id, _ = resolve_growth_actor()
        recommendation = next((item for item in list_learning_recommendations(state.db, user_id) if item.id == recommendation_id), None)
        if recommendation is None:
            raise HTTPException(status_code=404, detail="成长练习推荐不存在或已失效")
        updated = mark_recommendation_dismissed(state.db, recommendation_id, payload.reason)
        if updated is None:
            raise HTTPException(status_code=404, detail="成长练习推荐不存在或已失效")
        return GrowthRecommendationActionResponse(recommendation=updated, task=None)

    return app


def build_excerpt(path: Path) -> str:
    if path.suffix.lower() in {".md", ".txt", ".json", ".csv"}:
        try:
            return path.read_text(encoding="utf-8", errors="ignore")[:140] or f"{path.name} 已导入。"
        except Exception:
            pass
    return f"{path.name} 已进入资料缓冲池，可作为后续问答与证据引用来源。"


def backfill_local_task_tag_ids(state: AppState) -> None:
    timestamp = now_iso()
    operator_id = state.db.get_setting("current_operator_id", "") or "op_qh"
    state.db.execute(
        """
        UPDATE task_tags
        SET scope = COALESCE(NULLIF(scope, ''), 'org'),
            color = COALESCE(NULLIF(color, ''), CASE WHEN scope = 'self' THEN '#9CA3AF' ELSE '#5B7BFE' END),
            owner_operator_id = COALESCE(owner_operator_id, ''),
            created_by = COALESCE(NULLIF(created_by, ''), '系统'),
            created_at = COALESCE(NULLIF(created_at, ''), ?),
            updated_at = COALESCE(NULLIF(updated_at, ''), ?)
        """,
        (timestamp, timestamp),
    )
    state.db.execute(
        "UPDATE task_lists SET is_default = CASE WHEN id = 'list-0' THEN 1 ELSE COALESCE(is_default, 0) END WHERE is_default IS NULL OR is_default = ''"
    )
    for row in state.db.fetchall("SELECT id, tags_json, tag_ids_json FROM tasks"):
        tag_ids = _parse_json_list(row["tag_ids_json"])
        if tag_ids:
            continue
        tag_names = _parse_json_list(row["tags_json"])
        if not tag_names:
            continue
        resolved = [_ensure_local_tag(state.db, operator_id, name, "org") for name in tag_names if name.strip()]
        state.db.execute(
            "UPDATE tasks SET tag_ids_json = ?, tags_json = ?, updated_at = ? WHERE id = ?",
            (to_json([item.id for item in resolved]), to_json([item.name for item in resolved]), timestamp, str(row["id"])),
        )
    state.db.execute(
        "UPDATE weekly_reviews SET operator_id = COALESCE(NULLIF(operator_id, ''), ?), updated_at = COALESCE(NULLIF(updated_at, ''), created_at) WHERE operator_id = '' OR updated_at = ''",
        (operator_id,),
    )


def seed_defaults(state: AppState) -> None:
    timestamp = now_iso()
    state.db.set_setting("folders_root_label", state.db.get_setting("folders_root_label", "桌面客户资料"))
    state.db.set_setting("ai_provider", state.db.get_setting("ai_provider", DEFAULT_PROVIDER))
    state.db.set_setting("ai_model", state.db.get_setting("ai_model", DEFAULT_MODEL))
    state.db.set_setting("demo_data_loaded", state.db.get_setting("demo_data_loaded", "0"))
    if state.db.scalar("SELECT COUNT(1) AS count FROM operators") == 0:
        operators = [
            ("op_qh", "庆华", "首席咨询助理", "咨询策略", "#5B7BFE", 1),
            ("op_ys", "一朔", "研究分析师", "洞察研究", "#10B981", 0),
            ("op_jn", "嘉宁", "项目推进", "交付协同", "#F59E0B", 0),
        ]
        state.db.executemany(
            "INSERT INTO operators(id, name, role, team, color, is_current, created_at, updated_at) VALUES(?, ?, ?, ?, ?, ?, ?, ?)",
            [(item[0], item[1], item[2], item[3], item[4], item[5], timestamp, timestamp) for item in operators],
        )
        state.db.set_setting("current_operator_id", "op_qh")
    if state.db.scalar("SELECT COUNT(1) AS count FROM task_lists") == 0:
        state.db.executemany(
            "INSERT INTO task_lists(id, name, color, sort_order, is_default, scope, archived_at) VALUES(?, ?, ?, ?, ?, ?, NULL)",
            [
                ("list-0", "收集箱", "#888681", 0, 1, "org"),
                ("list-1", "客户推进", "#5B7BFE", 1, 0, "org"),
                ("list-2", "研究洞察", "#F59E0B", 2, 0, "org"),
                ("list-3", "交付沉淀", "#10B981", 3, 0, "org"),
                ("plist-1", "健身", "#5B7BFE", 10, 1, "personal"),
                ("plist-2", "约会", "#EC4899", 11, 0, "personal"),
                ("plist-3", "吃饭", "#F59E0B", 12, 0, "personal"),
                ("plist-4", "学习", "#10B981", 13, 0, "personal"),
            ],
        )
    state.db.execute("UPDATE task_lists SET scope = 'org' WHERE scope IS NULL OR scope = ''")
    if state.db.scalar("SELECT COUNT(1) AS count FROM task_lists WHERE scope = 'personal'") == 0:
        state.db.executemany(
            "INSERT INTO task_lists(id, name, color, sort_order, is_default, scope, archived_at) VALUES(?, ?, ?, ?, ?, ?, NULL)",
            [
                ("plist-1", "健身", "#5B7BFE", 10, 1, "personal"),
                ("plist-2", "约会", "#EC4899", 11, 0, "personal"),
                ("plist-3", "吃饭", "#F59E0B", 12, 0, "personal"),
                ("plist-4", "学习", "#10B981", 13, 0, "personal"),
            ],
        )
    if state.db.scalar("SELECT COUNT(1) AS count FROM task_tags") == 0:
        state.db.executemany(
            "INSERT INTO task_tags(id, name, scope, color, owner_operator_id, created_by, created_at, updated_at, archived_at) VALUES(?, ?, 'org', ?, '', '系统', ?, ?, NULL)",
            [
                (new_id("tag"), "高优", "#EF4444", timestamp, timestamp),
                (new_id("tag"), "会议", "#5B7BFE", timestamp, timestamp),
                (new_id("tag"), "待跟进", "#F59E0B", timestamp, timestamp),
                (new_id("tag"), "跨部门", "#10B981", timestamp, timestamp),
                (new_id("tag"), "选题", "#8B5CF6", timestamp, timestamp),
            ],
        )
    if state.db.scalar("SELECT COUNT(1) AS count FROM task_settings") == 0:
        state.db.executemany(
            """
            INSERT INTO task_settings(
                operator_id, default_list_id, default_priority, default_due_date_preset,
                default_view_mode, list_sort_mode, show_completed_tasks, default_review_scope,
                auto_assign_self, updated_at
            ) VALUES(?, 'list-0', 'normal', 'today', 'calendar', 'manual', 0, 'work', 1, ?)
            """,
            [(str(row["id"]), timestamp) for row in state.db.fetchall("SELECT id FROM operators")],
        )
    if state.db.scalar("SELECT COUNT(1) AS count FROM analysis_templates") == 0:
        state.db.executemany(
            "INSERT INTO analysis_templates(id, title, description, template_key, created_at) VALUES(?, ?, ?, ?, ?)",
            [
                ("tpl_fundraising", "筹款分析", "聚焦筹资路径、渠道效率、节奏与风险。", "fundraising", timestamp),
                ("tpl_systemic", "系统分析", "聚焦组织问题、依赖关系、根因与推进建议。", "systemic", timestamp),
            ],
        )
    backfill_local_task_tag_ids(state)


app = create_app()
