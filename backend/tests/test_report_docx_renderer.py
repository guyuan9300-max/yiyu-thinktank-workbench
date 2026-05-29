"""R3 · docx_renderer 单测。

只 stub 掉 LibreOffice（外部进程依赖），其它都用真实 python-docx 渲染。
"""

from __future__ import annotations

import base64
from pathlib import Path
from unittest.mock import patch

import pytest
from docx import Document

from app.models import (
    ChartHint,
    CitationRef,
    GeneratedChart,
    ReportArtifact,
    ReportBlueprint,
    SectionContent,
    SectionPlan,
)
from app.services.report_docx_renderer import (
    DocxRenderError,
    _is_separator_row,
    _parse_table_row,
    _replace_chart_placeholders_with_data_uri,
    convert_docx_to_pdf_via_libreoffice,
    render_report_artifact_to_docx,
    render_report_artifact_to_markdown,
)


_PNG_1x1 = (
    b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00"
    b"\x01\x08\x06\x00\x00\x00\x1f\x15\xc4\x89\x00\x00\x00\rIDATx\x9cc"
    b"\xfc\xff\xff?\x03\x00\x05\xfe\x02\xfe\xa3<\x9ej\x00\x00\x00\x00"
    b"IEND\xaeB`\x82"
)


def _png_base64() -> str:
    return base64.b64encode(_PNG_1x1).decode("ascii")


def _make_blueprint() -> ReportBlueprint:
    return ReportBlueprint(
        title="测试报告",
        subtitle="副标题",
        report_kind="测试",
        audience="决策层",
        tone="客观",
        period_start="2026-01-01",
        period_end="2026-03-31",
        sections=[],
        inferred_theme="测试主题",
        confidence=0.8,
        open_questions_for_human=["问题 A？", "问题 B？"],
        event_line_id="el-1",
        client_id="c-1",
        generated_at="2026-05-12T10:00:00+00:00",
    )


def _make_plan(title: str = "总览", level: int = 1) -> SectionPlan:
    return SectionPlan(
        level=level,
        title=title,
        goal="x",
        data_sources=["events"],
        chart_hints=[
            ChartHint(
                kind="timeline",
                title="时间线",
                caption="节奏",
                data_source_hint="x",
            )
        ],
        citation_budget=3,
        estimated_words=200,
    )


def _make_section_with_chart() -> SectionContent:
    return SectionContent(
        plan=_make_plan(),
        markdown=(
            "本季度推进顺利。\n\n"
            "[CHART:0]\n\n"
            "下一步聚焦优化。"
        ),
        citations=[
            CitationRef(
                type="event", id="a1", label="启动会", excerpt="确定主题"
            )
        ],
        charts=[
            GeneratedChart(
                hint=ChartHint(
                    kind="timeline",
                    title="时间线",
                    caption="节奏",
                    data_source_hint="x",
                ),
                png_bytes_base64=_png_base64(),
                width_cm=14.5,
            )
        ],
        data_source_annotation="事件线 a1..a5",
        confidence=0.85,
        warnings=[],
    )


@pytest.mark.unit
def test_render_docx_writes_file(tmp_path: Path) -> None:
    artifact = ReportArtifact(
        blueprint=_make_blueprint(),
        sections=[_make_section_with_chart()],
        output_files={},
        generated_at="2026-05-12T10:00:00Z",
    )
    out = tmp_path / "out.docx"
    result_path = render_report_artifact_to_docx(
        artifact, out, client_name="A组织"
    )
    assert result_path == out.resolve()
    assert out.exists()
    assert out.stat().st_size > 5000  # 含图的 docx 通常 > 5KB

    doc = Document(str(out))
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "测试报告" in all_text
    assert "A组织" in all_text
    assert "总览" in all_text
    assert "本季度推进顺利" in all_text
    assert "事件线 a1..a5" in all_text


@pytest.mark.unit
def test_render_docx_no_sections_raises(tmp_path: Path) -> None:
    artifact = ReportArtifact(
        blueprint=_make_blueprint(),
        sections=[],
        output_files={},
        generated_at="x",
    )
    with pytest.raises(DocxRenderError, match="没有任何章节"):
        render_report_artifact_to_docx(artifact, tmp_path / "x.docx")


@pytest.mark.unit
def test_render_markdown_outputs_self_contained(tmp_path: Path) -> None:
    artifact = ReportArtifact(
        blueprint=_make_blueprint(),
        sections=[_make_section_with_chart()],
        output_files={},
        generated_at="x",
    )
    out = tmp_path / "out.md"
    result_path = render_report_artifact_to_markdown(
        artifact, out, client_name="客户 A"
    )
    assert result_path == out.resolve()
    content = out.read_text(encoding="utf-8")
    assert "# 测试报告" in content
    assert "客户 A" in content
    assert "![节奏](data:image/png;base64," in content
    assert "[CHART:0]" not in content


@pytest.mark.unit
def test_render_handles_empty_chart_placeholder(tmp_path: Path) -> None:
    """callout_only / table_only 的 chart png_bytes_base64='' 时不该崩。"""
    plan = _make_plan()
    plan = plan.model_copy(
        update={
            "chart_hints": [
                ChartHint(
                    kind="callout_only",
                    title="提醒",
                    caption=None,
                    data_source_hint="",
                )
            ]
        }
    )
    sc = SectionContent(
        plan=plan,
        markdown="本节是纯文字。\n\n[CHART:0]\n\n继续正文。",
        citations=[],
        charts=[
            GeneratedChart(
                hint=plan.chart_hints[0],
                png_bytes_base64="",
                width_cm=14.5,
            )
        ],
        data_source_annotation="",
        confidence=0.8,
        warnings=[],
    )
    artifact = ReportArtifact(
        blueprint=_make_blueprint(),
        sections=[sc],
        output_files={},
        generated_at="x",
    )
    out = tmp_path / "x.docx"
    render_report_artifact_to_docx(artifact, out)
    assert out.exists()


@pytest.mark.unit
def test_render_chart_idx_out_of_range_records_placeholder(
    tmp_path: Path,
) -> None:
    sc = SectionContent(
        plan=_make_plan(),
        markdown="正文 [CHART:99] 末尾",
        citations=[],
        charts=[],
        data_source_annotation="",
        confidence=0.8,
        warnings=[],
    )
    artifact = ReportArtifact(
        blueprint=_make_blueprint(),
        sections=[sc],
        output_files={},
        generated_at="x",
    )
    out = tmp_path / "x.docx"
    render_report_artifact_to_docx(artifact, out)
    doc = Document(str(out))
    all_text = "\n".join(p.text for p in doc.paragraphs)
    # 整段（不是 [CHART:99] 单独成行）当作普通段落处理
    assert "[CHART:99]" in all_text


@pytest.mark.unit
def test_render_markdown_table(tmp_path: Path) -> None:
    md = (
        "前言\n"
        "\n"
        "| 模块 | 完成度 | 负责人 |\n"
        "| --- | --- | --- |\n"
        "| 心盛 | 80% | 张三 |\n"
        "| 心晴 | 60% | 李四 |\n"
        "\n"
        "尾声"
    )
    sc = SectionContent(
        plan=_make_plan().model_copy(update={"chart_hints": []}),
        markdown=md,
        citations=[],
        charts=[],
        data_source_annotation="",
        confidence=0.8,
        warnings=[],
    )
    artifact = ReportArtifact(
        blueprint=_make_blueprint(),
        sections=[sc],
        output_files={},
        generated_at="x",
    )
    out = tmp_path / "t.docx"
    render_report_artifact_to_docx(artifact, out)
    doc = Document(str(out))
    assert len(doc.tables) >= 1
    table = doc.tables[0]
    headers = [c.text.strip() for c in table.rows[0].cells]
    assert headers == ["模块", "完成度", "负责人"]
    assert table.rows[1].cells[0].text.strip() == "心盛"
    assert table.rows[2].cells[2].text.strip() == "李四"


@pytest.mark.unit
def test_render_callout_block(tmp_path: Path) -> None:
    md = (
        "正文 1\n"
        "\n"
        "> 这是 callout 的第一行\n"
        "> 第二行延续\n"
        "\n"
        "正文 2"
    )
    sc = SectionContent(
        plan=_make_plan().model_copy(update={"chart_hints": []}),
        markdown=md,
        citations=[],
        charts=[],
        data_source_annotation="",
        confidence=0.8,
        warnings=[],
    )
    artifact = ReportArtifact(
        blueprint=_make_blueprint(),
        sections=[sc],
        output_files={},
        generated_at="x",
    )
    out = tmp_path / "c.docx"
    render_report_artifact_to_docx(artifact, out)
    doc = Document(str(out))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "这是 callout 的第一行" in text
    assert "第二行延续" in text


@pytest.mark.unit
def test_render_bullets_and_numbered_lists(tmp_path: Path) -> None:
    md = (
        "无序列表：\n"
        "- 第一项\n"
        "- 第二项\n"
        "  - 子项\n"
        "\n"
        "编号列表：\n"
        "1. 步骤一\n"
        "2. 步骤二"
    )
    sc = SectionContent(
        plan=_make_plan().model_copy(update={"chart_hints": []}),
        markdown=md,
        citations=[],
        charts=[],
        data_source_annotation="",
        confidence=0.8,
        warnings=[],
    )
    artifact = ReportArtifact(
        blueprint=_make_blueprint(),
        sections=[sc],
        output_files={},
        generated_at="x",
    )
    out = tmp_path / "b.docx"
    render_report_artifact_to_docx(artifact, out)
    doc = Document(str(out))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "第一项" in text
    assert "子项" in text
    assert "步骤一" in text


@pytest.mark.unit
def test_render_includes_warnings_appendix(tmp_path: Path) -> None:
    sc = SectionContent(
        plan=_make_plan().model_copy(update={"chart_hints": []}),
        markdown="正文",
        citations=[],
        charts=[],
        data_source_annotation="",
        confidence=0.8,
        warnings=["事实基础不足，建议主理人复核"],
    )
    artifact = ReportArtifact(
        blueprint=_make_blueprint(),
        sections=[sc],
        output_files={},
        generated_at="x",
    )
    out = tmp_path / "w.docx"
    render_report_artifact_to_docx(artifact, out)
    doc = Document(str(out))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "起草警告" in text
    assert "事实基础不足" in text


@pytest.mark.unit
def test_no_warnings_no_appendix(tmp_path: Path) -> None:
    sc = SectionContent(
        plan=_make_plan().model_copy(update={"chart_hints": []}),
        markdown="正文",
        citations=[],
        charts=[],
        data_source_annotation="",
        confidence=0.8,
        warnings=[],
    )
    artifact = ReportArtifact(
        blueprint=_make_blueprint(),
        sections=[sc],
        output_files={},
        generated_at="x",
    )
    out = tmp_path / "n.docx"
    render_report_artifact_to_docx(artifact, out)
    doc = Document(str(out))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "起草警告" not in text


@pytest.mark.unit
def test_parse_table_row() -> None:
    assert _parse_table_row("| a | b | c |") == ["a", "b", "c"]
    assert _parse_table_row("|a|b|") == ["a", "b"]


@pytest.mark.unit
def test_is_separator_row() -> None:
    assert _is_separator_row(["---", "---"]) is True
    assert _is_separator_row([":---", "---:", ":---:"]) is True
    assert _is_separator_row(["a", "b"]) is False
    assert _is_separator_row([]) is False


@pytest.mark.unit
def test_replace_chart_placeholders_with_data_uri() -> None:
    charts = [
        GeneratedChart(
            hint=ChartHint(
                kind="timeline",
                title="t",
                caption="说明",
                data_source_hint="x",
            ),
            png_bytes_base64=_png_base64(),
            width_cm=14.5,
        )
    ]
    md = "前 [CHART:0] 后"
    result = _replace_chart_placeholders_with_data_uri(md, charts)
    assert "data:image/png;base64," in result
    assert "[CHART:0]" not in result


@pytest.mark.unit
def test_replace_chart_placeholders_skips_empty_base64() -> None:
    charts = [
        GeneratedChart(
            hint=ChartHint(
                kind="callout_only",
                title="t",
                caption=None,
                data_source_hint="",
            ),
            png_bytes_base64="",
            width_cm=14.5,
        )
    ]
    md = "前 [CHART:0] 后"
    result = _replace_chart_placeholders_with_data_uri(md, charts)
    assert "[CHART:0]" not in result
    assert "前  后" in result  # 占位被替换为空


@pytest.mark.unit
def test_inline_chart_placeholder_renders_image(tmp_path: Path) -> None:
    """LLM B 经常把 [CHART:0] 嵌在句子里（如"参考时间线[CHART:0]。"），
    渲染器应：emit 段落（去占位符）+ 段落后追加图，inline_shapes 计数 ≥ 1。"""
    plan = _make_plan()
    sc = SectionContent(
        plan=plan,
        markdown="本季度推进顺利，全周期核心节点参考关键里程碑时间线[CHART:0]。",
        citations=[],
        charts=[
            GeneratedChart(
                hint=plan.chart_hints[0],
                png_bytes_base64=_png_base64(),
                width_cm=14.5,
            )
        ],
        data_source_annotation="x",
        confidence=0.8,
        warnings=[],
    )
    artifact = ReportArtifact(
        blueprint=_make_blueprint(),
        sections=[sc],
        output_files={},
        generated_at="x",
    )
    out = tmp_path / "i.docx"
    render_report_artifact_to_docx(artifact, out)
    doc = Document(str(out))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "[CHART:0]" not in text  # 占位符已删
    assert "全周期核心节点" in text
    assert len(doc.inline_shapes) >= 1  # 图已嵌入


@pytest.mark.unit
def test_inline_multiple_charts_in_one_paragraph(tmp_path: Path) -> None:
    plan = _make_plan().model_copy(
        update={
            "chart_hints": [
                ChartHint(kind="timeline", title="a", caption=None, data_source_hint=""),
                ChartHint(kind="timeline", title="b", caption=None, data_source_hint=""),
            ]
        }
    )
    sc = SectionContent(
        plan=plan,
        markdown="进展见时间线[CHART:0]，完成度见对比[CHART:1]。",
        citations=[],
        charts=[
            GeneratedChart(
                hint=plan.chart_hints[0],
                png_bytes_base64=_png_base64(),
                width_cm=14.5,
            ),
            GeneratedChart(
                hint=plan.chart_hints[1],
                png_bytes_base64=_png_base64(),
                width_cm=14.5,
            ),
        ],
        data_source_annotation="x",
        confidence=0.8,
        warnings=[],
    )
    artifact = ReportArtifact(
        blueprint=_make_blueprint(),
        sections=[sc],
        output_files={},
        generated_at="x",
    )
    out = tmp_path / "m.docx"
    render_report_artifact_to_docx(artifact, out)
    doc = Document(str(out))
    assert len(doc.inline_shapes) == 2


@pytest.mark.unit
def test_libreoffice_not_installed_returns_none(tmp_path: Path) -> None:
    """没有 soffice 时不应 raise，应返回 None。"""
    docx = tmp_path / "fake.docx"
    docx.write_bytes(b"PK\x03\x04")  # 一个 zip header 占位，让 exists() 过
    with patch(
        "app.services.report_docx_renderer.shutil.which", return_value=None
    ):
        result = convert_docx_to_pdf_via_libreoffice(docx, tmp_path)
    assert result is None


@pytest.mark.unit
def test_libreoffice_missing_source_returns_none(tmp_path: Path) -> None:
    """源 docx 不存在 → None。"""
    result = convert_docx_to_pdf_via_libreoffice(
        tmp_path / "nope.docx", tmp_path
    )
    assert result is None
