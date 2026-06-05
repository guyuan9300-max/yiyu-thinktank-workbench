"""P0a(拆字段名/提示)+ P0b(字典直填兜底)的针对性测试。

重点:
  1. split_field_label 拆分正确 + 单行向后兼容(==normalize_template_label)。
  2. 拆后字段类型不再被说明里的"登记证书"污染成 attachment_material。
  3. ★关键约束★ 拆字段名后,extract 的 label 与 apply 的写回匹配键仍一致 →
     值能正确写回单元格(用真模板跑端到端)。
  4. match_field_to_verified_glossary:schema 别名归一 → 命中已核验属性。
"""
from pathlib import Path

import pytest
from docx import Document

from app.services.template_fill import (
    apply_docx_template_values,
    extract_docx_template_fields,
    infer_template_field_type,
    match_field_to_verified_glossary,
    normalize_template_label,
    split_field_label,
)


def test_split_field_label_basic():
    name, hint = split_field_label("1. 机构名称\n填写登记证书上的全称。")
    assert name == "1. 机构名称"
    assert hint == "填写登记证书上的全称。"


def test_split_field_label_single_line_backward_compatible():
    # 无换行的格:拆出来的字段名必须等于旧 normalize_template_label(写回兼容的根本)
    for s in ["机构名称", "2021 年", "7. 项目名称", "组织类型及登记信息"]:
        assert split_field_label(s)[0] == normalize_template_label(s)
        assert split_field_label(s)[1] is None


def test_split_field_label_empty_first_line_fallback():
    name, hint = split_field_label("\n填写说明在这")
    assert name  # 不为空(回退用整体)
    assert hint is None


def test_classification_not_polluted_by_hint():
    full = "1. 机构名称\n填写登记证书上的全称。"
    name, _ = split_field_label(full)
    # 旧:整坨 → attachment_material(被"登记证书"污染);新:拆后字段名 → 不是 attachment
    assert infer_template_field_type(normalize_template_label(full)) == "attachment_material"
    assert infer_template_field_type(name) != "attachment_material"


def _make_namehint_docx(path: Path) -> None:
    """造一个"字段名\n填写说明 | 空答案列"的两列表,模拟民政模板结构。"""
    doc = Document()
    table = doc.add_table(rows=3, cols=2)
    table.rows[0].cells[0].text = "1. 机构名称\n填写登记证书上的全称。"
    table.rows[1].cells[0].text = "3. 联系人及联系方式\n填写姓名、职务、手机号、邮箱。"
    table.rows[2].cells[0].text = "12. 项目目标\n填写年度目标、服务目标。"
    # 答案列留空
    doc.save(str(path))


def test_extract_sets_clean_label_and_hint(tmp_path):
    tpl = tmp_path / "nh.docx"
    _make_namehint_docx(tpl)
    fields = extract_docx_template_fields(tpl)
    by_label = {f.label: f for f in fields}
    # label 应是干净字段名(不含说明)
    assert "1. 机构名称" in by_label
    assert "3. 联系人及联系方式" in by_label
    # hint 被拆出来挂到 field.hint 上
    assert by_label["1. 机构名称"].hint == "填写登记证书上的全称。"
    assert "姓名" in (by_label["3. 联系人及联系方式"].hint or "")


def test_apply_writeback_consistent_after_split(tmp_path):
    """★核心约束★:extract 拆了 label 后,apply 仍能用 split 的首行键找回单元格并写值。"""
    tpl = tmp_path / "nh.docx"
    out = tmp_path / "nh_filled.docx"
    _make_namehint_docx(tpl)
    fields = extract_docx_template_fields(tpl)
    values = {f.label: f"答案-{f.label}" for f in fields}
    applied, _missing = apply_docx_template_values(tpl, out, values)
    assert applied >= 3
    doc = Document(str(out))
    t = doc.tables[0]
    # 三行答案列都应被写入(证明 extract label 与 apply 匹配键一致)
    assert t.rows[0].cells[1].text.strip().startswith("答案-1. 机构名称")
    assert t.rows[1].cells[1].text.strip().startswith("答案-3. 联系人")
    assert t.rows[2].cells[1].text.strip().startswith("答案-12. 项目目标")


def test_apply_on_real_template_writeback(tmp_path):
    """真模板:确认拆字段名后 apply 不破(每个字段给值,断言被写回)。"""
    real = Path("/Users/guyuanyuan/Downloads/广州市民政局公益创投项目征集填报表_可填格版.docx")
    if not real.exists():
        pytest.skip("真模板不在")
    out = tmp_path / "real_filled.docx"
    fields = extract_docx_template_fields(real)
    values = {f.label: f"TESTVAL[{f.label[:10]}]" for f in fields}
    applied, _missing = apply_docx_template_values(real, out, values)
    # 应写回相当数量(机构名称/联系人/项目… 这些两列字段都该被填上)
    assert applied >= 15, f"写回数量异常: {applied}"
    # 机构名称那格确实被填(拆字段名没破写回)
    doc = Document(str(out))
    org_cell = doc.tables[1].rows[0].cells[1].text.strip()
    assert org_cell.startswith("TESTVAL"), f"机构名称未写回: {org_cell!r}"


class _Attr:
    def __init__(self, term, attribute_name, value_text, scope="", as_of_date=None):
        self.term = term
        self.attribute_name = attribute_name
        self.value_text = value_text
        self.scope = scope
        self.as_of_date = as_of_date


def test_glossary_match_via_schema_alias():
    attrs = [
        _Attr("为爱黔行公益", "基金会全称", "贵州省为爱黔行公益服务中心"),
        _Attr("为爱黔行公益", "成立时间", "2019年12月24日"),
        _Attr("为爱黔行公益", "法定代表人", "吴建林"),
    ]
    # 机构名称 → 基金会全称(schema 别名)
    assert match_field_to_verified_glossary(attrs, "1. 机构名称").attribute_name == "基金会全称"
    # 成立时间 → 精确
    assert match_field_to_verified_glossary(attrs, "成立时间").attribute_name == "成立时间"
    # 理事长 → 法定代表人(schema 别名)
    assert match_field_to_verified_glossary(attrs, "理事长").attribute_name == "法定代表人"
    # 不相关字段不应误命中
    assert match_field_to_verified_glossary(attrs, "预算概算和资金使用说明") is None


def test_glossary_match_empty_inputs():
    assert match_field_to_verified_glossary([], "机构名称") is None
    assert match_field_to_verified_glossary([_Attr("t", "a", "")], "机构名称") is None  # 空值不命中


def test_map_field_to_narrative_segment():
    from app.services.template_fill import map_field_to_narrative_segment as m
    # 复合字段 → 对应叙事段
    assert m("13. 主要服务内容/活动安排") == "business_intro"
    assert m("14. 项目特点与创新点") == "business_intro"
    assert m("15. 实施计划与关键节点") == "timeline"
    assert m("16. 合作资源与协同机制") == "cooperation"
    assert m("9. 项目负责人及团队") == "people"
    assert m("4. 机构简介") == "essence"
    assert m("12. 项目目标") == "next_steps"
    # 精确字段 / 年度矩阵 / 盖章 / 预算 → 不注入叙事(None)
    assert m("1. 机构名称") is None
    assert m("2021 年") is None
    assert m("39. 填报单位盖章") is None
    assert m("17. 预算概算和资金使用说明") is None
