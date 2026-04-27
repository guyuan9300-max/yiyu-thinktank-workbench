from __future__ import annotations

import sys
from pathlib import Path

from docx import Document as WordDocument

sys.path.append(str(Path(__file__).resolve().parents[1]))

from app.services.template_fill import (
    apply_docx_template_values,
    build_template_follow_up_question,
    build_template_fill_retrieval_query,
    build_template_fill_web_queries,
    build_template_suggested_sources,
    derive_template_fill_public_names,
    derive_template_fill_public_domain,
    extract_template_milestone_year,
    extract_docx_attachment_checklist,
    extract_docx_template_fields,
    fetch_template_fill_web_sources,
    infer_template_field_type,
    infer_template_value_kind,
    normalize_template_public_domain,
    should_enable_template_fill_web_supplement,
)


def test_extract_docx_template_fields_detects_placeholders_and_blank_table_cells(tmp_path: Path):
    target = tmp_path / "template.docx"
    document = WordDocument()
    document.add_paragraph("机构名称：{{机构名称}}")
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "机构简介"
    table.rows[0].cells[1].text = ""
    table.rows[1].cells[0].text = "服务区域"
    table.rows[1].cells[1].text = "{{服务区域}}"
    document.save(target)

    fields = extract_docx_template_fields(target)
    labels = [item.label for item in fields]

    assert "机构名称" in labels
    assert "机构简介" in labels
    assert "服务区域" in labels


def test_apply_docx_template_values_fills_placeholders_and_table_cells(tmp_path: Path):
    source = tmp_path / "template.docx"
    target = tmp_path / "filled.docx"
    document = WordDocument()
    document.add_paragraph("机构名称：{{机构名称}}")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "机构简介"
    table.rows[0].cells[1].text = ""
    document.save(source)

    applied, missing = apply_docx_template_values(
        source,
        target,
        {
            "机构名称": "日慈基金会",
            "机构简介": "专注于青少年心理健康与社会情感学习。",
        },
    )

    assert applied >= 2
    assert missing == 0

    filled = WordDocument(target)
    assert "日慈基金会" in filled.paragraphs[0].text
    assert filled.tables[0].rows[0].cells[1].text == "专注于青少年心理健康与社会情感学习。"


def test_extract_docx_template_fields_detects_right_side_field_fill_pairs(tmp_path: Path):
    target = tmp_path / "paired-template.docx"
    document = WordDocument()
    table = document.add_table(rows=2, cols=4)
    table.rows[0].cells[0].text = "字段"
    table.rows[0].cells[1].text = "填写内容"
    table.rows[0].cells[2].text = "字段"
    table.rows[0].cells[3].text = "填写内容"
    table.rows[1].cells[0].text = "组织全称"
    table.rows[1].cells[1].text = ""
    table.rows[1].cells[2].text = "法定代表人"
    table.rows[1].cells[3].text = ""
    document.save(target)

    fields = extract_docx_template_fields(target)
    labels = [item.label for item in fields]

    assert "组织全称" in labels
    assert "法定代表人" in labels


def test_apply_docx_template_values_fills_right_side_field_fill_pairs(tmp_path: Path):
    source = tmp_path / "paired-template.docx"
    target = tmp_path / "paired-filled.docx"
    document = WordDocument()
    table = document.add_table(rows=2, cols=4)
    table.rows[0].cells[0].text = "字段"
    table.rows[0].cells[1].text = "填写内容"
    table.rows[0].cells[2].text = "字段"
    table.rows[0].cells[3].text = "填写内容"
    table.rows[1].cells[0].text = "组织全称"
    table.rows[1].cells[1].text = ""
    table.rows[1].cells[2].text = "法定代表人"
    table.rows[1].cells[3].text = ""
    document.save(source)

    applied, missing = apply_docx_template_values(
        source,
        target,
        {
            "组织全称": "北京基业长青社会组织服务中心",
            "法定代表人": "顾源源",
        },
    )

    assert applied == 2
    assert missing == 0

    filled = WordDocument(target)
    assert filled.tables[0].rows[1].cells[1].text == "北京基业长青社会组织服务中心"
    assert filled.tables[0].rows[1].cells[3].text == "顾源源"


def test_extract_docx_template_fields_detects_service_object_column(tmp_path: Path):
    target = tmp_path / "module-template.docx"
    document = WordDocument()
    table = document.add_table(rows=2, cols=4)
    table.rows[0].cells[0].text = "业务模块"
    table.rows[0].cells[1].text = "主要内容"
    table.rows[0].cells[2].text = "服务对象/覆盖对象"
    table.rows[0].cells[3].text = "可提取资料来源"
    table.rows[1].cells[0].text = "政策倡导"
    table.rows[1].cells[1].text = ""
    table.rows[1].cells[2].text = ""
    table.rows[1].cells[3].text = "新闻稿"
    document.save(target)

    labels = [item.label for item in extract_docx_template_fields(target)]

    assert "政策倡导" in labels
    assert "政策倡导（服务对象/覆盖对象）" in labels


def test_apply_docx_template_values_fills_service_object_column(tmp_path: Path):
    source = tmp_path / "module-template.docx"
    target = tmp_path / "module-filled.docx"
    document = WordDocument()
    table = document.add_table(rows=2, cols=4)
    table.rows[0].cells[0].text = "业务模块"
    table.rows[0].cells[1].text = "主要内容"
    table.rows[0].cells[2].text = "服务对象/覆盖对象"
    table.rows[0].cells[3].text = "可提取资料来源"
    table.rows[1].cells[0].text = "政策倡导"
    table.rows[1].cells[1].text = ""
    table.rows[1].cells[2].text = ""
    table.rows[1].cells[3].text = "新闻稿"
    document.save(source)

    applied, missing = apply_docx_template_values(
        source,
        target,
        {
            "政策倡导": "面向公益行业的政策倡导与法规建言。",
            "政策倡导（服务对象/覆盖对象）": "基金会、行业平台与政策参与方。",
        },
    )

    assert applied == 2
    assert missing == 0

    filled = WordDocument(target)
    assert filled.tables[0].rows[1].cells[1].text == "面向公益行业的政策倡导与法规建言。"
    assert filled.tables[0].rows[1].cells[2].text == "基金会、行业平台与政策参与方。"


def test_infer_template_field_type_covers_core_field_classes():
    assert infer_template_field_type("统一社会信用代码") == "precise_fact"
    assert infer_template_field_type("机构定位") == "structural_summary"
    assert infer_template_field_type("党建与业务工作的结合方式") == "governance_mechanism"
    assert infer_template_field_type("近三年代表性活动/会议数量") == "quantitative_result"
    assert infer_template_field_type("章程（含党建条款页）") == "attachment_material"


def test_infer_template_value_kind_marks_missing_values_conservatively():
    assert infer_template_value_kind("【待确认】当前缺少可直接填写该字段的资料。", "precise_fact") == "missing"
    assert infer_template_value_kind("2017年6月登记注册", "precise_fact") == "fact"
    assert infer_template_value_kind("围绕行业透明、公信力、治理能力建设开展工作。", "governance_mechanism") == "summary"


def test_extract_docx_attachment_checklist_reads_attachment_table(tmp_path: Path):
    target = tmp_path / "attachment-template.docx"
    document = WordDocument()
    table = document.add_table(rows=3, cols=4)
    table.rows[0].cells[0].text = "序号"
    table.rows[0].cells[1].text = "附件名称"
    table.rows[0].cells[2].text = "是否已备"
    table.rows[1].cells[0].text = "1"
    table.rows[1].cells[1].text = "机构登记证书复印件"
    table.rows[2].cells[0].text = "2"
    table.rows[2].cells[1].text = "章程（含党建条款页）"
    document.save(target)

    attachments = extract_docx_attachment_checklist(target)

    assert attachments == ["机构登记证书复印件", "章程（含党建条款页）"]


def test_missing_field_helpers_produce_follow_up_and_source_hints():
    follow_up = build_template_follow_up_question("governance_mechanism", "党组织在重大事项决策中的作用描述")
    sources = build_template_suggested_sources("governance_mechanism", "党组织在重大事项决策中的作用描述")

    assert "章程" in follow_up
    assert "会议纪要" in follow_up
    assert "章程" in sources
    assert "制度文件" in sources


def test_template_fill_web_supplement_only_enables_for_sparse_supported_fields():
    assert should_enable_template_fill_web_supplement("precise_fact", 0) is True
    assert should_enable_template_fill_web_supplement("structural_summary", 2) is True
    assert should_enable_template_fill_web_supplement("governance_mechanism", 0) is False
    assert should_enable_template_fill_web_supplement("structural_summary", 3) is False
    assert should_enable_template_fill_web_supplement("structural_summary", 3, field_label="2008年重大事件/里程碑") is True
    assert should_enable_template_fill_web_supplement("structural_summary", 4, field_label="2008年重大事件/里程碑") is False


def test_normalize_template_public_domain_extracts_host_safely():
    assert normalize_template_public_domain("https://www.cff.org.cn/about") == "cff.org.cn"
    assert normalize_template_public_domain("WWW.EXAMPLE.ORG") == "example.org"
    assert normalize_template_public_domain("invalid-host") is None


def test_extract_template_milestone_year_reads_year_from_label():
    assert extract_template_milestone_year("2008年重大事件/里程碑") == "2008"
    assert extract_template_milestone_year("2017重大事件/里程碑") == "2017"
    assert extract_template_milestone_year("机构定位") is None


def test_build_template_fill_retrieval_query_enriches_milestone_fields():
    query = build_template_fill_retrieval_query(
        client_name="CFFC",
        template_name="模板.docx",
        field_label="2008年重大事件/里程碑",
        field_type="structural_summary",
    )

    assert "2008年" in query
    assert "大事记" in query
    assert "发展历程" in query
    assert "模板.docx" not in query


def test_build_template_fill_web_queries_prioritizes_history_terms_for_milestones():
    queries = build_template_fill_web_queries(
        client_name="中国基金会发展论坛",
        field_label="2008年重大事件/里程碑",
        template_name="模板.docx",
        client_domain="cff.org.cn",
    )

    assert queries[0].startswith("中国基金会发展论坛 2008")
    assert any("大事记" in item for item in queries)


def test_derive_template_fill_public_names_prefers_chinese_org_name_from_local_titles():
    names = derive_template_fill_public_names(
        "CFFC",
        [
            "2016-2020年中国基金会发展论坛年会评估报告_CFFC_20260211.docx",
            "北京基业长青社会组织服务中心品牌使用指南（2025年过渡期版）_CFFC_20260211.docx",
        ],
    )

    assert names[0] == "CFFC"
    assert "中国基金会发展论坛" in names


def test_derive_template_fill_public_names_strips_rule_like_suffixes_from_titles():
    names = derive_template_fill_public_names(
        "CFFC",
        [
            "附件二：中国基金会发展论坛组委会运行规则-2023年度组委会第三次会议决议版_CFFC_20260211.pdf",
        ],
    )

    assert "中国基金会发展论坛" in names


def test_derive_template_fill_public_names_can_also_use_local_snippets():
    names = derive_template_fill_public_names(
        "CFFC",
        [],
        [
            "中国基金会发展论坛（英文名称 China Foundation Forum，中文简称基金会论坛，英文简称 CFF）是由多家机构共同发起的行业平台。",
        ],
    )

    assert "中国基金会发展论坛" in names
    assert "基金会论坛" in names


def test_derive_template_fill_public_names_ignores_generic_title_candidates():
    names = derive_template_fill_public_names(
        "CFFC",
        [
            "基金会实务工具包_CFFC_20260211.docx",
            "组织架构图_CFFC_20260211.docx",
            "中国基金会发展论坛2023年会项目总结报告_CFFC_20260211.docx",
        ],
    )

    assert "基金会实务工具包" not in names
    assert "组织架构图" not in names
    assert "中国基金会发展论坛" in names


def test_derive_template_fill_public_names_prefers_snippet_org_name_before_generic_titles():
    names = derive_template_fill_public_names(
        "CFFC",
        [
            "基金会实务工具包_CFFC_20260211.docx",
            "组织架构图_CFFC_20260211.docx",
        ],
        [
            "中国基金会发展论坛（英文名称 China Foundation Forum，中文简称基金会论坛，英文简称 CFF）是有志于追求机构卓越、行业发展的社会组织自愿发起的行业平台。",
        ],
    )

    assert names[:3] == ["CFFC", "中国基金会发展论坛", "基金会论坛"]


def test_derive_template_fill_public_domain_falls_back_to_domains_in_local_snippets():
    domain = derive_template_fill_public_domain(
        None,
        [
            "更多信息见 https://www.cfforum.org.cn/about",
            "备用网址 cfforum.org",
        ],
    )

    assert domain == "cfforum.org.cn"


def test_derive_template_fill_public_domain_prefers_domain_in_same_snippet_as_public_name():
    domain = derive_template_fill_public_domain(
        None,
        [
            "南都观察站更多信息见 https://nandu.org.cn/about",
            "中国基金会发展论坛（英文名称 China Foundation Forum，简称基金会论坛，CFF）秘书处邮箱：mishuchu@cfforum.org.cn",
        ],
        public_names=["中国基金会发展论坛", "基金会论坛"],
        client_name="CFFC",
    )

    assert domain == "cfforum.org.cn"


def test_derive_template_fill_public_domain_prefers_org_cn_over_email_only_org_variant():
    domain = derive_template_fill_public_domain(
        None,
        [
            "中国基金会发展论坛秘书处邮箱：mishuchu@cfforum.org",
            "中国基金会发展论坛官网：cfforum.org.cn",
        ],
        public_names=["中国基金会发展论坛", "基金会论坛"],
        client_name="CFFC",
    )

    assert domain == "cfforum.org.cn"


def test_derive_template_fill_public_domain_prefers_org_cn_when_only_org_is_email_sibling():
    domain = derive_template_fill_public_domain(
        None,
        [
            "中国基金会发展论坛秘书处邮箱：mishuchu@cfforum.org",
            "北京基业长青社会组织服务中心邮箱：mishuchu@cfforum.org.cn",
        ],
        public_names=["中国基金会发展论坛", "基金会论坛"],
        client_name="CFFC",
    )

    assert domain == "cfforum.org.cn"


def test_derive_template_fill_public_domain_ignores_generic_meeting_hosts():
    domain = derive_template_fill_public_domain(
        None,
        [
            "中国基金会发展论坛2024年度会议线上地址：https://meeting.dingtalk.com/abc",
            "中国基金会发展论坛官网：cfforum.org.cn",
        ],
        public_names=["中国基金会发展论坛", "基金会论坛"],
        client_name="CFFC",
    )

    assert domain == "cfforum.org.cn"


def test_fetch_template_fill_web_sources_prefers_official_internal_history_pages_for_milestones(monkeypatch):
    homepage_url = "https://cfforum.org.cn"
    history_url = "https://cfforum.org.cn/category/21"
    about_url = "https://cfforum.org.cn/category/20"

    html_map = {
        homepage_url: """
        <html><body>
          <a href="/category/20">关于我们</a>
          <a href="/category/21">大事记</a>
          <a href="/category/26">年度盛会</a>
        </body></html>
        """,
        history_url: "<html><body>2008年，在当时的民政部民间组织管理局指导下发起中国非公募基金会发展论坛。</body></html>",
        about_url: "<html><body>2016年转型为中国基金会发展论坛，2017年秘书处完成注册。</body></html>",
    }

    from app.services import template_fill as template_fill_module

    template_fill_module._fetch_url_html.cache_clear()
    template_fill_module._fetch_url_snippet.cache_clear()
    template_fill_module._search_duckduckgo_html.cache_clear()

    def fake_fetch_url_html(url: str) -> str:
        return html_map.get(url, "")

    monkeypatch.setattr(template_fill_module, "_fetch_url_html", fake_fetch_url_html)
    monkeypatch.setattr(template_fill_module, "_fetch_url_snippet", lambda url: template_fill_module._strip_web_html(fake_fetch_url_html(url))[:900])
    monkeypatch.setattr(template_fill_module, "_search_duckduckgo_html", lambda query: ())

    sources = fetch_template_fill_web_sources(
        client_name="中国基金会发展论坛",
        field_label="2008年重大事件/里程碑",
        template_name="模板.docx",
        client_domain="cfforum.org.cn",
        evidence_titles=["中国基金会发展论坛2023年会项目总结报告.docx"],
        evidence_snippets=["中国基金会发展论坛（简称基金会论坛）是行业平台。"],
        max_items=3,
        field_type="structural_summary",
    )

    urls = [item.url for item in sources]
    assert homepage_url in urls
    assert history_url in urls or about_url in urls
