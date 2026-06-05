"""本次改动的针对性测试:
  1. strip_emoji —— 剔除 emoji,保留中文 / ①②③ / 【】/ 全角标点。
  2. 宽表(矩阵)多列识别 —— "第一列锚点 + 多个维度列"的每个数据格都成为可填目标,
     而不是只认第一列(根治"多列只识别第一列空格")。
  3. 端到端 apply —— 矩阵表每列都被写入,且写入值已去 emoji。
  4. 普通两列表 / 单行表不被矩阵逻辑误伤。
"""
from pathlib import Path

from docx import Document

from app.services.template_fill import (
    apply_docx_template_values,
    extract_docx_template_fields,
    strip_emoji,
)


def test_strip_emoji_removes_emoji_keeps_chinese():
    assert strip_emoji("项目目标🎯：提升能力✅") == "项目目标：提升能力"
    assert strip_emoji("①政府背书 ②资金到位💰") == "①政府背书 ②资金到位"
    assert strip_emoji("【待确认】当前缺少资料。") == "【待确认】当前缺少资料。"
    assert strip_emoji("里程碑➡️第一阶段⭐") == "里程碑第一阶段"
    assert strip_emoji("") == ""
    # 无 emoji 文本保持不变(不误伤)
    assert strip_emoji("贵州省为爱黔行公益服务中心（慈善组织）") == "贵州省为爱黔行公益服务中心（慈善组织）"


def _make_matrix_docx(path: Path) -> None:
    doc = Document()
    table = doc.add_table(rows=4, cols=5)
    headers = ["年度", "服务规模", "项目数量", "资金数据", "代表性成果"]
    for c, h in enumerate(headers):
        table.rows[0].cells[c].text = h
    for r, year in enumerate(["2021 年", "2022 年", "2023 年"], start=1):
        table.rows[r].cells[0].text = year
        # 维度列留空(待填)
    doc.save(str(path))


def test_matrix_wide_table_recognizes_all_answer_columns(tmp_path):
    tpl = tmp_path / "matrix.docx"
    _make_matrix_docx(tpl)
    fields = extract_docx_template_fields(tpl)
    matrix_fields = [f for f in fields if "·" in f.label]
    # 3 行数据 × 4 个维度列 = 12 个可填格(旧逻辑只会得到 3 个=每行第一维度列)
    assert len(matrix_fields) == 12, f"应识别 12 个矩阵格,实际 {len(matrix_fields)}"
    # 标签形如 "2021 年·服务规模"
    assert any(f.label.startswith("2021 年·服务规模") for f in matrix_fields)
    assert any(f.label.startswith("2023 年·代表性成果") for f in matrix_fields)
    # 覆盖到 col2/3/4(旧逻辑漏掉的)
    cols = {f.cell_index for f in matrix_fields}
    assert cols == {1, 2, 3, 4}, f"应覆盖 1-4 列,实际 {cols}"


def test_matrix_apply_fills_all_columns_and_strips_emoji(tmp_path):
    tpl = tmp_path / "matrix.docx"
    out = tmp_path / "matrix_filled.docx"
    _make_matrix_docx(tpl)
    fields = extract_docx_template_fields(tpl)
    values = {f.label: f"答案🎯[{f.cell_index}]" for f in fields if "·" in f.label}
    applied, _missing = apply_docx_template_values(tpl, out, values)
    assert applied >= 12
    doc = Document(str(out))
    t = doc.tables[0]
    for r in range(1, 4):
        for c in range(1, 5):
            cell_text = (t.rows[r].cells[c].text or "").strip()
            assert cell_text.startswith("答案["), f"行{r}列{c} 未填或未去emoji: {cell_text!r}"
            assert "🎯" not in cell_text


def test_plain_two_col_table_not_treated_as_matrix(tmp_path):
    """普通两列表(标签|答案)不应触发矩阵逻辑,标签里不应出现 '·' 拼接。"""
    tpl = tmp_path / "two_col.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "机构名称"
    table.rows[0].cells[1].text = ""
    table.rows[1].cells[0].text = "联系人"
    table.rows[1].cells[1].text = ""
    doc.save(str(tpl))
    fields = extract_docx_template_fields(tpl)
    assert all("·" not in f.label for f in fields), "两列表不应产生矩阵式 '锚点·维度' 标签"


def test_single_row_wide_table_not_matrix(tmp_path):
    """单行宽表(无数据行)不应触发矩阵逻辑。"""
    tpl = tmp_path / "single_row.docx"
    doc = Document()
    table = doc.add_table(rows=1, cols=4)
    for c, h in enumerate(["盖章", "", "日期", ""]):
        table.rows[0].cells[c].text = h
    doc.save(str(tpl))
    fields = extract_docx_template_fields(tpl)
    assert all("·" not in f.label for f in fields)
