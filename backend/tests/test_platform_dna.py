from __future__ import annotations

import sys
from pathlib import Path

from docx import Document as WordDocument

sys.path.append(str(Path(__file__).resolve().parents[1]))

from app.services.platform_dna import extract_platform_dna_text


def _write_simple_pdf(path: Path, text: str) -> None:
    objects: list[bytes] = []

    def add_object(payload: bytes) -> int:
        objects.append(payload)
        return len(objects)

    add_object(b"<< /Type /Catalog /Pages 2 0 R >>")
    add_object(b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>")
    add_object(b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 300 144] /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>")
    stream = f"BT\n/F1 18 Tf\n36 96 Td\n({text}) Tj\nET".encode("latin-1")
    add_object(f"<< /Length {len(stream)} >>\nstream\n".encode("latin-1") + stream + b"\nendstream")
    add_object(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")

    chunks = [b"%PDF-1.4\n"]
    offsets = [0]
    for index, payload in enumerate(objects, start=1):
        offsets.append(sum(len(chunk) for chunk in chunks))
        chunks.append(f"{index} 0 obj\n".encode("latin-1"))
        chunks.append(payload)
        chunks.append(b"\nendobj\n")
    xref_offset = sum(len(chunk) for chunk in chunks)
    chunks.append(f"xref\n0 {len(objects) + 1}\n".encode("latin-1"))
    chunks.append(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        chunks.append(f"{offset:010d} 00000 n \n".encode("latin-1"))
    chunks.append(
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF\n".encode("latin-1")
    )
    path.write_bytes(b"".join(chunks))


def test_extract_platform_dna_text_reads_markdown(tmp_path: Path) -> None:
    target = tmp_path / "platform.md"
    target.write_text("# 腾讯公益\n\n核心偏好\n- 真实\n- 预算清楚\n", encoding="utf-8")

    result = extract_platform_dna_text(target)

    assert "腾讯公益" in result
    assert "预算清楚" in result


def test_extract_platform_dna_text_reads_docx(tmp_path: Path) -> None:
    target = tmp_path / "platform.docx"
    document = WordDocument()
    document.add_heading("抖音公益", level=1)
    document.add_paragraph("核心偏好")
    document.add_paragraph("更看重真实故事与具体行动。")
    document.save(target)

    result = extract_platform_dna_text(target)

    assert "抖音公益" in result
    assert "真实故事与具体行动" in result


def test_extract_platform_dna_text_reads_pdf(tmp_path: Path) -> None:
    target = tmp_path / "platform.pdf"
    _write_simple_pdf(target, "Hello PDF DNA")

    result = extract_platform_dna_text(target)

    assert "Hello PDF DNA" in result
